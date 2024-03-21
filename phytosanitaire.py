import arabic_reshaper
import bidi.algorithm
import locale
import os
import tempfile
import docx
import sqlite3
from docx.enum.section import WD_ORIENTATION
from docx.shared import Inches, Pt, RGBColor
from PyQt6 import QtCore, QtGui, QtWidgets
from PyQt6.QtCore import QDate, Qt, QLocale, QTimer
from PyQt6.QtPrintSupport import QPrintDialog
from PyQt6.QtWidgets import QMessageBox
from PyQt6.QtWidgets import *
from docx.enum.section import WD_ORIENTATION
from docx.shared import Inches, Pt, RGBColor
from docxtpl import DocxTemplate
import mysql.connector

userslist = ["nadir", "qualite"]
passwordlist = ["Nadir206@", "qualite48"]

for user ,password in zip(userslist,passwordlist):
    try:
        dattabase = mysql.connector.connect(host="localhost", user=user, password=password)
        curs = dattabase.cursor()
        curs.execute("CREATE DATABASE IF NOT EXISTS datta")
        curs.execute("USE datta")
        curs.execute("""create table if not exists phytotable
                        (
                          
                          LIEU_DE_TRAITEMENT text,
                          NATURE_ET_QUANTITE_DE_PRODUIT_TRAITE text,
                          DATE_DE_PEREMPTION text,
                          STOCK_FINAL INTEGER,
                          QUANTITE_UTILISEE INTEGER,
                          QUANTITE_ACHETE INTEGER,
                          STOCK_INITIAL INTEGER,
                          PRODUITS_PHYTOSANITAIRES TEXT ,
                          LIEU_DE_STOCKAGE text,
                          le text,
                          id INT AUTO_INCREMENT PRIMARY KEY ,
                          STOCK_INITIAL_ACTELLIC integer, 
                          QUANTITE_UTILISEE_ACTELLIC integer,
                          QUANTITE_ACHETE_ACTELLIC integer,
                          STOCK_INITIAL_CIRATHRINE integer, 
                          QUANTITE_UTILISEE_CIRATHRINE integer,
                          QUANTITE_ACHETE_CIRATHRINE integer,
                          STOCK_INITIAL_DEKATRINE integer, 
                          QUANTITE_UTILISEE_DEKATRINE integer,
                          QUANTITE_ACHETE_DEKATRINE integer,
                          STOCK_INITIAL_PHOSTOXIN_PH3 integer, 
                          QUANTITE_UTILISEE_PHOSTOXIN_PH3 integer,
                          QUANTITE_ACHETE_PHOSTOXIN_PH3 integer,
                          STOCK_INITIAL_RATICIDE integer, 
                          QUANTITE_UTILISEE_RATICIDE integer,
                          QUANTITE_ACHETE_RATICIDE integer,
                          STOCK_INITIAL_TEXTO integer,
                          QUANTITE_UTILISEE_TEXTO integer,
                          QUANTITE_ACHETE_TEXTO integer,
                          STOCK_FINAL_ACTELLIC REAL,
                          STOCK_FINAL_CIRATHRINE REAL,
                          STOCK_FINAL_PHOSTOXIN_PH3 REAL,
                          STOCK_FINAL_RATICIDE REAL,
                          STOCK_FINAL_TEXTO REAL
                        )
                        """)
        curs.execute("""create table if not exists phytotablebidi
                        (
                          
                          LIEU_DE_TRAITEMENT text,
                          NATURE_ET_QUANTITE_DE_PRODUIT_TRAITE text,
                          DATE_DE_PEREMPTION text,
                          STOCK_FINAL INTEGER,
                          QUANTITE_UTILISEE INTEGER,
                          QUANTITE_ACHETE INTEGER,
                          STOCK_INITIAL INTEGER,
                          PRODUITS_PHYTOSANITAIRES TEXT ,
                          LIEU_DE_STOCKAGE text,
                          le text,
                          idbidi INT AUTO_INCREMENT PRIMARY KEY ,
                          STOCK_INITIAL_ACTELLIC integer, 
                          QUANTITE_UTILISEE_ACTELLIC integer,
                          QUANTITE_ACHETE_ACTELLIC integer,
                          STOCK_INITIAL_CIRATHRINE integer, 
                          QUANTITE_UTILISEE_CIRATHRINE integer,
                          QUANTITE_ACHETE_CIRATHRINE integer,
                          STOCK_INITIAL_DEKATRINE integer, 
                          QUANTITE_UTILISEE_DEKATRINE integer,
                          QUANTITE_ACHETE_DEKATRINE integer,
                          STOCK_INITIAL_PHOSTOXIN_PH3 integer, 
                          QUANTITE_UTILISEE_PHOSTOXIN_PH3 integer,
                          QUANTITE_ACHETE_PHOSTOXIN_PH3 integer,
                          STOCK_INITIAL_RATICIDE integer, 
                          QUANTITE_UTILISEE_RATICIDE integer,
                          QUANTITE_ACHETE_RATICIDE integer,
                          STOCK_INITIAL_TEXTO integer,
                          QUANTITE_UTILISEE_TEXTO integer,
                          QUANTITE_ACHETE_TEXTO integer,
                          STOCK_FINAL_ACTELLIC REAL,
                          STOCK_FINAL_CIRATHRINE REAL,
                          STOCK_FINAL_PHOSTOXIN_PH3 REAL,
                          STOCK_FINAL_RATICIDE REAL,
                          STOCK_FINAL_TEXTO REAL
                        )
                        """)
        dattabase.commit()
        dattabase.close()

    except mysql.connector.Error as e:
        print("mysql error:" + e)


try:
    class Phyto_Window():
        def phyoto_produit(self, MainWindow):
            MainWindow.setObjectName("Window")
            MainWindow.resize(1350, 700)

            MainWindow.setStyleSheet("""QToolTip
            {
                border: 1px solid #76797C;
                background-color: rgb(90, 102, 117);;
                color: white;
                padding: 5px;
                opacity: 200;
            }

            QWidget
            {
                color: #000000;
                background-color: #ffffff;
                selection-background-color:#3daee9;
                selection-color: #3daee9;
                background-clip: border;
                border-image: none;
                border: 0px transparent black;
                outline: 0;
            }

            QWidget:item:hover
            {
                background-color: #3daee9;
                color: #eff0f1;
            }

            QWidget:item:selected
            {
                background-color: #3daee9;
            }



            QWidget:disabled
            {
                color: #454545;
                background-color: #31363b;
            }

            QAbstractItemView
            {
                alternate-background-color: #31363b;
                color: #eff0f1;
                border: 1px solid 3A3939;
                border-radius: 2px;
            }

            QWidget:focus, QMenuBar:focus
            {
                border: 1px solid #3daee9;
            }

            QTabWidget:focus, QCheckBox:focus, QRadioButton:focus, QSlider:focus
            {
                border: none;
            }

            QLineEdit
            {
                background-color: #FDFEFE;
                padding: 1px;
                border-style: solid;
                border: 1px solid #76797C;
                border-radius: 0px;
                color: #000000;
            }
            QDoubleSpinBox
            {
                background-color: #FDFEFE;
                padding: 1px;
                border-style: solid;
                border: 1px solid #76797C;
                border-radius: 0px;
                color:#000000;
                font-size: 11px;
                font-weight: bold;

            }
            QDoubleSpinBox:focus{
                background-color: #FDFEFE;
                border-style: solid;
                border: 2px solid #76797C;
                border-radius: 4px;
                border-color: #ff8c00;
            }
            QDoubleSpinBox::drop-down
            {
                subcontrol-origin: padding;
                subcontrol-position: top right;
                width: 1px;

                border-left-width: 0px;
                border-left-color: #232629;
                border-left-style: solid;
                border-top-right-radius: 1px;
                border-bottom-right-radius: 1px;
            }



            QGroupBox {
                border:1px solid #76797C;
                border-radius: 2px;
                margin-top: 5px;
            }

            QGroupBox::title {
                subcontrol-origin: margin;
                subcontrol-position: top center;
                padding-left: 4px;
                padding-right: 4px;
                padding-top: 4px;
            }

            QAbstractScrollArea
            {
                border-radius: 2px;
                border: 1px solid #76797C;
                background-color: transparent;
            }

            QScrollBar:horizontal
            {
                height: 15px;
                margin: 3px 15px 3px 15px;
                border: 1px transparent #2A2929;
                border-radius: 4px;
                background-color: #2A2929;
            }

            QScrollBar::handle:horizontal
            {
                background-color: #605F5F;
                min-width: 5px;
                border-radius: 4px;
            }

            QScrollBar::add-line:horizontal
            {
                margin: 0px 3px 0px 3px;
                border-image: url(:/qss_icons/Dark_rc/right_arrow_disabled.png);
                width: 10px;
                height: 10px;
                subcontrol-position: right;
                subcontrol-origin: margin;
            }

            QScrollBar::sub-line:horizontal
            {
                margin: 0px 3px 0px 3px;
                border-image: url(:/qss_icons/Dark_rc/left_arrow_disabled.png);
                height: 10px;
                width: 10px;
                subcontrol-position: left;
                subcontrol-origin: margin;
            }

            QScrollBar::add-line:horizontal:hover,QScrollBar::add-line:horizontal:on
            {
                border-image: url(:/qss_icons/Dark_rc/right_arrow.png);
                height: 10px;
                width: 10px;
                subcontrol-position: right;
                subcontrol-origin: margin;
            }


            QScrollBar::sub-line:horizontal:hover, QScrollBar::sub-line:horizontal:on
            {
                border-image: url(:/qss_icons/Dark_rc/left_arrow.png);
                height: 10px;
                width: 10px;
                subcontrol-position: left;
                subcontrol-origin: margin;
            }

            QScrollBar::up-arrow:horizontal, QScrollBar::down-arrow:horizontal
            {
                background: none;
            }


            QScrollBar::add-page:horizontal, QScrollBar::sub-page:horizontal
            {
                background: none;
            }

            QScrollBar:vertical
            {
                background-color: #2A2929;
                width: 15px;
                margin: 15px 3px 15px 3px;
                border: 1px transparent #2A2929;
                border-radius: 4px;
            }

            QScrollBar::handle:vertical
            {
                background-color: #605F5F;
                min-height: 5px;
                border-radius: 4px;
            }

            QScrollBar::sub-line:vertical
            {
                margin: 3px 0px 3px 0px;
                border-image: url(:/qss_icons/Dark_rc/up_arrow_disabled.png);
                height: 10px;
                width: 10px;
                subcontrol-position: top;
                subcontrol-origin: margin;
            }

            QScrollBar::add-line:vertical
            {
                margin: 3px 0px 3px 0px;
                border-image: url(:/qss_icons/Dark_rc/down_arrow_disabled.png);
                height: 10px;
                width: 10px;
                subcontrol-position: bottom;
                subcontrol-origin: margin;
            }

            QScrollBar::sub-line:vertical:hover,QScrollBar::sub-line:vertical:on
            {

                border-image: url(:/qss_icons/Dark_rc/up_arrow.png);
                height: 10px;
                width: 10px;
                subcontrol-position: top;
                subcontrol-origin: margin;
            }


            QScrollBar::add-line:vertical:hover, QScrollBar::add-line:vertical:on
            {
                border-image: url(:/qss_icons/Dark_rc/down_arrow.png);
                height: 10px;
                width: 10px;
                subcontrol-position: bottom;
                subcontrol-origin: margin;
            }

            QScrollBar::up-arrow:vertical, QScrollBar::down-arrow:vertical
            {
                background: none;
            }


            QScrollBar::add-page:vertical, QScrollBar::sub-page:vertical
            {
                background: none;
            }

            QTextEdit
            {
                background-color: #FDFEFE;
                color: #000000;
                border: 1px solid #76797C;
            }

            QPlainTextEdit
            {
                background-color: #232629;;
                color: #eff0f1;
                border-radius: 2px;
                border: 1px solid #76797C;
            }

            QHeaderView::section
            {
                background-color: #76797C;
                color: #eff0f1;
                padding: 1px;
                border: 1px solid #76797C;
            }

            QSizeGrip {
                width: 12px;
                height: 12px;
            }


            QMainWindow::separator
            {
                background-color: #31363b;
                color: white;
                padding-left: 4px;
                spacing: 2px;
                border: 1px dashed #76797C;
            }

            QMainWindow::separator:hover
            {

                background-color: #787876;
                color: white;
                padding-left: 4px;
                border: 1px solid #76797C;
                spacing: 2px;
            }


            QMenu::separator
            {
                height: 1px;
                background-color: #76797C;
                color: white;
                padding-left: 4px;
                margin-left: 10px;
                margin-right: 5px;
            }


            QFrame
            {
                border-radius: 2px;
                border: 1px solid #76797C;
            }

            QFrame[frameShape="0"]
            {
                border-radius: 2px;
                border: 1px transparent #76797C;
            }

            QStackedWidget
            {
                border: 1px transparent black;
            }


            QPushButton
            {
                color: #00000;
                background-color:#ade3e7;
                border-width: 1px;
                border-color: #1e1e1e;
                border-style: solid;
                border-radius: 6;
                padding: 3px;
                font-size: 12px;
                padding-left: 5px;
                padding-right: 5px;
                min-width: 40px;

            }

            QPushButton:disabled
            {
                background-color:#03ecff;
                border-width: 1px;
                border-color: #454545;
                border-style: solid;
                padding-top: 5px;
                padding-bottom: 5px;
                padding-left: 10px;
                padding-right: 10px;
                border-radius: 2px;
                color: #454545;
            }
            QPushButton:pressed
            {
                background-color: #3daee9;
                padding-top: -15px;
                padding-bottom: -17px;
            }

            QComboBox {
        background-color: #FDFEFE;
        border: 1px solid #76797C;
        color:#000000;
        border-radius: 0.25em;
        padding: 0.0em 0.0em;
        font-size: 1.25rem;
        cursor: pointer;
    }

    QComboBox::drop-down {
        subcontrol-origin: padding;
        subcontrol-position: top right;
        width: 1.3em;
        border-left: 0px solid #777;
        border-radius: 0.25em;
    }

    QComboBox::drop-down::icon {
        image: url('E:/pythonProject_moullin-application.3.5/images/down-arroww.png');
    }

            QPushButton:checked{
                background-color: #76797C;
                border-color: #6A6969;
            }
            QComboBox:on
            {
                padding-top: 1px;
                padding-left: 1px;
                selection-background-color: #FDFEFE;
            }
            QComboBox QAbstractItemView
            {
                background-color: #FDFEFE;
                border-radius: 2px;
                border: 1px solid #76797C;
                color:#000000;
                selection-background-color: #000000;
            }
            QComboBox:hover,QDoubleSpinBox:Hover,QPushButton:hover,QAbstractSpinBox:hover,QLineEdit:hover,QTextEdit:hover,QPlainTextEdit:hover,QAbstractView:hover,QTreeView:hover
            {
                border: 1px solid #ff8c00;
                color: #000000;
            }




            QLabel
            {
                border: 2px solid black;
            }

            QTabWidget{
                border: 0px transparent black;
            }

            QTabWidget::pane {
                border: 1px solid #76797C;
                padding: 5px;
                margin: 0px;
            }

            QTabBar
            {
                qproperty-drawBase: 0;
                left: 5px; /* move to the right by 5px */
                border-radius: 3px;
            }

            QTabBar:focus
            {
                border: 0px transparent black;
            }

            QTabBar::close-button  {
                image: url(:/qss_icons/Dark_rc/close.png);
                background: transparent;
            }

            QTabBar::close-button:hover
            {
                image: url(:/qss_icons/Dark_rc/close-hover.png);
                background: transparent;
            }

            QTabBar::close-button:pressed {
                image: url(:/qss_icons/Dark_rc/close-pressed.png);
                background: transparent;
            }

            /* TOP TABS */
            QTabBar::tab:top {
                color: #eff0f1;
                border: 1px solid #76797C;
                border-bottom: 1px transparent black;
                background-color: #31363b;
                padding: 5px;
                min-width: 10px;
                border-top-left-radius: 2px;
                border-top-right-radius: 2px;
            }

            QTabBar::tab:top:!selected
            {
                color: #eff0f1;
                background-color: #54575B;
                border: 1px solid #76797C;
                border-bottom: 1px transparent black;
                border-top-left-radius: 2px;
                border-top-right-radius: 2px;    
            }

            QTabBar::tab:top:!selected:hover {
                background-color: #3daee9;
            }

            /* BOTTOM TABS */
            QTabBar::tab:bottom {
                color: #eff0f1;
                border: 1px solid #76797C;
                border-top: 1px transparent black;
                background-color: #31363b;
                padding: 5px;
                border-bottom-left-radius: 2px;
                border-bottom-right-radius: 2px;
                min-width: 10px;
            }

            QTabBar::tab:bottom:!selected
            {
                color: #eff0f1;
                background-color: #54575B;
                border: 1px solid #76797C;
                border-top: 1px transparent black;
                border-bottom-left-radius: 2px;
                border-bottom-right-radius: 2px;
            }

            QTabBar::tab:bottom:!selected:hover {
                background-color: #3daee9;
            }

            /* LEFT TABS */
            QTabBar::tab:left {
                color: #eff0f1;
                border: 1px solid #76797C;
                border-left: 1px transparent black;
                background-color: #31363b;
                padding: 5px;
                border-top-right-radius: 2px;
                border-bottom-right-radius: 2px;
                min-height: 50px;
            }

            QTabBar::tab:left:!selected
            {
                color: #eff0f1;
                background-color: #54575B;
                border: 1px solid #76797C;
                border-left: 1px transparent black;
                border-top-right-radius: 2px;
                border-bottom-right-radius: 2px;
            }

            QTabBar::tab:left:!selected:hover {
                background-color: #3daee9;
            }


            /* RIGHT TABS */
            QTabBar::tab:right {
                color: #eff0f1;
                border: 1px solid #76797C;
                border-right: 1px transparent black;
                background-color: #31363b;
                padding: 5px;
                border-top-left-radius: 2px;
                border-bottom-left-radius: 2px;
                min-height: 50px;
            }

            QTabBar::tab:right:!selected
            {
                color: #eff0f1;
                background-color: #54575B;
                border: 1px solid #76797C;
                border-right: 1px transparent black;
                border-top-left-radius: 2px;
                border-bottom-left-radius: 2px;
            }





            QSlider::groove:horizontal {
                border: 1px solid #565a5e;
                height: 4px;
                background: #565a5e;
                margin: 0px;
                border-radius: 2px;
            }

            QSlider::handle:horizontal {
                background: #232629;
                border: 1px solid #565a5e;
                width: 16px;
                height: 16px;
                margin: -8px 0;
                border-radius: 9px;
            }

            QSlider::groove:vertical {
                border: 1px solid #565a5e;
                width: 4px;
                background: #565a5e;
                margin: 0px;
                border-radius: 3px;
            }

            QSlider::handle:vertical {
                background: #232629;
                border: 1px solid #565a5e;
                width: 16px;
                height: 16px;
                margin: 0 -8px;
                border-radius: 9px;
            }

            QToolButton {
                background-color: transparent;
                border: 1px transparent #76797C;
                border-radius: 2px;
                margin: 3px;
                padding: 5px;
            }

            QToolButton[popupMode="1"] { /* only for MenuButtonPopup */
             padding-right: 20px; /* make way for the popup button */
             border: 1px #76797C;
             border-radius: 5px;
            }

            QToolButton[popupMode="2"] { /* only for InstantPopup */
             padding-right: 10px; /* make way for the popup button */
             border: 1px #76797C;
            }


            QToolButton:hover, QToolButton::menu-button:hover {
                background-color: transparent;
                border: 1px solid #3daee9;
                padding: 5px;
            }

            QToolButton:checked, QToolButton:pressed,
                    QToolButton::menu-button:pressed {
                background-color: #3daee9;
                border: 1px solid #3daee9;
                padding: 5px;
            }

            /* the subcontrol below is used only in the InstantPopup or DelayedPopup mode */
            QToolButton::menu-indicator {
                background-color:ff8c00;
                top: -7px; left: -2px; /* shift it a bit */
            }

            /* the subcontrols below are used only in the MenuButtonPopup mode */
            QToolButton::menu-button {
                border: 1px transparent #76797C;
                border-top-right-radius: 6px;
                border-bottom-right-radius: 6px;
                /* 16px width + 4px for border = 20px allocated above */
                width: 16px;
                outline: none;
            }

            QToolButton::menu-arrow {
               background-color:ff8c00;
            }

            QToolButton::menu-arrow:open {
                border: 1px solid #76797C;
            }

            QPushButton::menu-indicator  {
                subcontrol-origin: padding;
                subcontrol-position: bottom right;
                left: 8px;
            }

            QTableView
            {
                border: 1px solid #76797C;
                gridline-color: #31363b;
                background-color: #FDFEFE;
                color:#000000;
            }


            QTableView, QHeaderView
            {
                background-color: #FDFEFE;
                color:#000000;
                border-radius: 0px;
            }

            QTableView::item:pressed, QListView::item:pressed, QTreeView::item:pressed  {
                background: #FDFEFE;
                color: #000000;
            }

            QTableView::item:selected:active, QTreeView::item:selected:active, QListView::item:selected:active  {
                background: #3daee9;
                color: #000000;
            }


            QHeaderView
            {
                background-color: #FDFEFE;
                border: 1px transparent;
                border-radius: 0px;
                margin: 0px;
                padding: 0px;

            }

            QHeaderView::section  {
                background-color:#80f1f9;
                color: #000000;
                padding: 5px;
                border: 1px solid #76797C;
                border-radius: 0px;
                text-align: center;
            }

            QHeaderView::section::vertical::first, QHeaderView::section::vertical::only-one
            {
                border-top: 1px solid #76797C;
            }

            QHeaderView::section::vertical
            {
                border-top: transparent;
            }

            QHeaderView::section::horizontal::first, QHeaderView::section::horizontal::only-one
            {
                border-left: 1px solid #76797C;
            }

            QHeaderView::section::horizontal
            {
                border-left: transparent;
            }


            QHeaderView::section:checked
             {
                color: #000000;
                background-color: #3daee9;
             }

             /* style the sort indicator */
            QHeaderView::down-arrow {

            }

            QHeaderView::up-arrow {

            }


            QTableCornerButton::section {
                background-color: #31363b;
                border: 1px transparent #76797C;
                border-radius: 0px;
            }

            QToolBox  {
                padding: 5px;
                border: 1px transparent black;
            }

            QToolBox::tab {
                color: #eff0f1;
                background-color: #31363b;
                border: 1px solid #76797C;
                border-bottom: 1px transparent #31363b;
                border-top-left-radius: 5px;
                border-top-right-radius: 5px;
            }

            QToolBox::tab:selected { /* italicize selected tabs */
                font: italic;
                background-color: #31363b;
                border-color: #3daee9;
             }

            QStatusBar::item {
                border: 0px transparent dark;
             }


            QFrame[height="3"], QFrame[width="3"] {
                background-color: #76797C;
            }




            QDateEdit
            {
                background-color: #232629;;
                border-style: solid;
                border: 1px solid #76797C;
                border-radius: 2px;
                padding: 1px;
                min-width: 75px;
            }

            QDateEdit:on
            {
                padding-top: 2px;
                padding-left: 2px;
                selection-background-color: #4a4a4a;
            }

            QDateEdit QAbstractItemView
            {
                background-color: #ff8c00;
                border-radius: 2px;
                border: 1px solid #3375A3;
                selection-background-color:ff8c00;
            }

            QDateEdit::drop-down
            {
                subcontrol-origin: padding;
                subcontrol-position: top right;
                width: 15px;
                border-left-width: 0px;
                border-left-color: darkgray;
                border-left-style: solid;
                border-top-right-radius: 3px;
                border-bottom-right-radius: 3px;
            }   
            QDateTimeEdit
            {
                background-color: #232629;;
                border-style: solid;
                border: 1px solid #76797C;
                border-radius: 2px;
                padding: 1px;
                min-width: 75px;

            }    
            """)

            self.addbtn = QtWidgets.QPushButton("اضف", MainWindow, clicked=lambda: self.add_datta_save())
            self.addbtn.setGeometry(QtCore.QRect(620, 275, 100, 40))

            ##################################################sortie###############
            # self.deletebtn = QtWidgets.QPushButton("احذف", MainWindow, clicked=lambda: self.delete_item())
            # self.deletebtn.setGeometry(QtCore.QRect(260, 275, 100, 40))

            self.printbtn = QtWidgets.QPushButton("السجل", MainWindow, clicked=lambda: self.print_docx())
            self.printbtn.setGeometry(QtCore.QRect(380, 275, 100, 40))

            self.printbtnfich = QtWidgets.QPushButton(" بطاقة\nالمعالجة", MainWindow,
                                                      clicked=lambda: self.fiche_traitement())
            self.printbtnfich.setGeometry(QtCore.QRect(500, 275, 100, 40))

            self.filtertxt = QtWidgets.QLabel(" اختر التاريخ ", MainWindow)
            self.filtertxt.setFrameShape(QtWidgets.QFrame.Shape.WinPanel.VLine)
            self.filtertxt.setGeometry(QtCore.QRect(740, 268, 590, 51))

            self.datefilter = QtWidgets.QDateEdit(MainWindow)
            self.datefilter.setGeometry(QtCore.QRect(1130, 275, 130, 40))
            self.datefilter.setDisplayFormat("yyyy-MM-dd")
            self.dateday = QDate.currentDate()
            self.datefilter.setDate(self.dateday)
            self.datefilter.setStyleSheet(
                " background-color: #FDFEFE;padding: 1px;border-style: solid;border: 1px solid #76797C;border-radius: 0px;color: #000000;")

            self.produitphytofl = QtWidgets.QComboBox(MainWindow)
            self.produitphytofl.setGeometry(QtCore.QRect(900, 275, 130, 40))
            self.produitphytofl.addItem("--------------")
            self.produitphytofl.addItem("PHOSTOXIN(PH3)")
            self.produitphytofl.addItem("ACTELLIC")
            self.produitphytofl.addItem("CIRATHRINE")
            self.produitphytofl.addItem("DEKATRINE")
            self.produitphytofl.addItem("RATICIDE")
            self.produitphytofl.addItem("TEXTO")

            self.btnfilter = QtWidgets.QPushButton("ابدء", MainWindow, clicked=lambda: self.impot_filter())
            self.btnfilter.setGeometry((QtCore.QRect(830, 275, 40, 40)))

            self.btnfcncl = QtWidgets.QPushButton("خروج", MainWindow, clicked=lambda: self.impot_all())
            self.btnfcncl.setGeometry((QtCore.QRect(760, 275, 40, 40)))

            self.datafiltertxt = QtWidgets.QLabel("اختر المادة:", MainWindow)
            self.datafiltertxt.setGeometry(QtCore.QRect(1040, 285, 60, 20))

            self.filtertxt = QtWidgets.QLabel("تصفية البيانات:", MainWindow)
            self.filtertxt.setGeometry(QtCore.QRect(1200, 254, 100, 20))

            self.cclstxt = QtWidgets.QLabel("<h2>تعاونية الحبوب والبقول الجافة بغليزان<h2/>", MainWindow)
            self.cclstxt.setGeometry(QtCore.QRect(430, 0, 500, 41))
            font = QtGui.QFont()
            font.setPointSize(14)
            font.setBold(True)
            font.setWeight(75)
            self.cclstxt.setFont(font)
            self.cclstxt.setMouseTracking(False)
            self.cclstxt.setLayoutDirection(QtCore.Qt.LayoutDirection.LeftToRight)
            self.cclstxt.setAutoFillBackground(False)
            self.cclstxt.setFrameShape(QtWidgets.QFrame.Shape.WinPanel)
            self.cclstxt.setLineWidth(0)
            self.cclstxt.setMidLineWidth(0)
            self.cclstxt.setTextFormat(QtCore.Qt.TextFormat.AutoText)
            self.cclstxt.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
            self.entrielabel = QtWidgets.QLabel("<h3>حالة مخزون مواد الصحة التباتية <h3/>", MainWindow)
            self.entrielabel.setGeometry(170, 75, 650, 40)
            self.entrielabel.setFont(font)
            self.titletxt = QtWidgets.QLabel(MainWindow)
            self.titletxt.setGeometry(QtCore.QRect(540, 50, 350, 30))
            font = QtGui.QFont()
            font.setPointSize(14)
            font.setBold(True)
            font.setUnderline(True)
            font.setWeight(75)
            self.titletxt.setFont(font)
            self.titletxt.setAlignment(
                QtCore.Qt.AlignmentFlag.AlignLeading | QtCore.Qt.AlignmentFlag.AlignLeft | QtCore.Qt.AlignmentFlag.AlignVCenter)
            self.titletxt.setObjectName("titletxt")
            #############################################
            self.txtdate = QtWidgets.QLabel("التاريخ", MainWindow)
            self.txtdate.setGeometry(QtCore.QRect(1215, 135, 110, 16))

            ####################################################
            self.dateedite = QtWidgets.QDateEdit(MainWindow)
            self.dateedite.setDisplayFormat("yyyy-MM-dd")

            self.dateedite.setStyleSheet(
                " background-color: #FDFEFE;padding: 1px;border-style: solid;border: 1px solid #76797C;border-radius: 0px;color: #000000;")
            self.dateedite.setGeometry(QtCore.QRect(1090, 130, 150, 35))
            self.dateedite.setDate(self.dateday)
            self.dateedite.setAlignment(QtCore.Qt.AlignmentFlag.AlignRight)
            ########################################################
            self.lieutxt = QtWidgets.QLabel("المخزن", MainWindow)
            self.lieutxt.setGeometry(QtCore.QRect(1225, 195, 100, 20))
            self.lieucombo = QtWidgets.QComboBox(MainWindow)
            self.lieucombo.setGeometry(QtCore.QRect(1090, 190, 150, 35))

            self.lieucombo.addItem(" ")
            self.lieucombo.addItem("المخزن الرئيسي غليزان")
            #####################################################################################
            self.produitphytotxt = QtWidgets.QLabel("مواد الحماية النباتية", MainWindow)
            self.produitphytotxt.setGeometry(QtCore.QRect(900, 125, 116, 40))
            ###################################################################
            self.produitphyto = QtWidgets.QComboBox(MainWindow, editable=True)
            self.produitphyto.setGeometry(QtCore.QRect(760, 130, 150, 35))
            self.produitphyto.addItem("")
            self.produitphyto.addItem("PHOSTOXIN(PH3)")
            self.produitphyto.addItem("ACTELLIC")
            self.produitphyto.addItem("CIRATHRINE")
            self.produitphyto.addItem("DEKATRINE")
            self.produitphyto.addItem("RATICIDE")
            self.produitphyto.addItem("TEXTO")

            ###############################################################################
            self.txtSTOCKINITIAL = QtWidgets.QLabel("الكمية المشترات", MainWindow)
            self.txtSTOCKINITIAL.setGeometry(QtCore.QRect(550, 138, 120, 16))

            self.STOCKINITIAL = QtWidgets.QDoubleSpinBox(MainWindow)
            self.STOCKINITIAL.setSpecialValueText(' ')
            self.STOCKINITIAL.setDecimals(3)
            self.STOCKINITIAL.setRange(0, 100000)
            self.STOCKINITIAL.setGeometry(QtCore.QRect(430, 133, 150, 35))
            ############################################################################
            self.quantitéutilsetxt = QtWidgets.QLabel("الكمية المستخدمة", MainWindow)
            self.quantitéutilsetxt.setGeometry(QtCore.QRect(235, 138, 100, 22))

            self.quantitéutilsé = QtWidgets.QDoubleSpinBox(MainWindow)
            self.quantitéutilsé.setGeometry(QtCore.QRect(80, 133, 150, 35))
            self.quantitéutilsé.setSpecialValueText(" ")
            self.quantitéutilsé.setDecimals(3)

            ###########################################################################
            self.dateprempotxt = QtWidgets.QLabel("تاريخ نهاية الصلاحية", MainWindow)
            self.dateprempotxt.setGeometry(QtCore.QRect(865, 195, 150, 22))

            self.dateprempo = QtWidgets.QDateEdit(MainWindow)
            self.dateprempo.setDisplayFormat("yyyy-MM-dd")
            self.dateprempo.setDate(self.dateday)
            self.dateprempo.setStyleSheet(
                " background-color: #FDFEFE;padding: 1px;border-style: solid;border: 1px solid #76797C;border-radius: 0px;color: #000000;")
            self.dateprempo.setGeometry(QtCore.QRect(760, 190, 150, 35))

            self.naturelieutxt = QtWidgets.QLabel("نوع وكمية المادة المعالجة", MainWindow)
            self.naturelieutxt.setGeometry(QtCore.QRect(217, 182, 155, 45))
            self.naturelieu = QtWidgets.QTextEdit(MainWindow)
            self.naturelieu.setGeometry(QtCore.QRect(80, 190, 150, 35))
            self.naturelieu.setAlignment(Qt.AlignmentFlag.AlignRight)

            #########################################################################
            ########################################################################
            self.lieudetraitementtxt = QtWidgets.QLabel("مكان المعالجة", MainWindow)
            self.lieudetraitementtxt.setGeometry(QtCore.QRect(507, 195, 160, 22))

            self.lieudetraitementcombo = QtWidgets.QComboBox(MainWindow)
            self.lieudetraitementcombo.setGeometry(QtCore.QRect(430, 190, 150, 35))

            self.lieudetraitementcombo.addItem("")
            self.lieudetraitementcombo.addItem("المخزن الرئيسي غليزان")
            self.lieudetraitementcombo.addItem("مخزن ماسرة")
            self.lieudetraitementcombo.addItem("مخزن الكهف الازرق")
            self.lieudetraitementcombo.addItem("مخزن زمورة")
            self.lieudetraitementcombo.addItem("المحطة الجديدة منداس")
            self.lieudetraitementcombo.addItem("محطة منداس")
            self.lieudetraitementcombo.addItem("مخزن اوفلا ")
            self.lieudetraitementcombo.addItem("مخزن بلعسل")

            ########################################################################
            ####################################################################

            self.textEdit = QtWidgets.QTableWidget(MainWindow)
            self.textEdit.setRowCount(0)
            self.textEdit.setColumnCount(11)
            self.textEdit.setColumnWidth(0, 100)
            self.textEdit.setColumnWidth(1, 160)
            self.textEdit.setColumnWidth(2, 120)
            self.textEdit.setColumnWidth(3, 120)
            self.textEdit.setColumnWidth(4, 120)
            self.textEdit.setColumnWidth(5, 100)
            self.textEdit.setColumnWidth(6, 100)
            self.textEdit.setColumnWidth(7, 150)
            self.textEdit.setColumnWidth(8, 160)
            self.textEdit.setColumnWidth(9, 100)
            self.textEdit.setColumnWidth(10, 60)

            self.textEdit.setHorizontalHeaderLabels(("مكان المعالجة", "نوع وكمية\n المادة المعالجة",
                                                     "تاريخ نهاية الصلاحية", "المخزون النهائي", "الكمية المستخدمة",
                                                     "الكمية المشترات ", "المخزون الاولي", "منتجات\nالحماية النباتية",
                                                     "المخزن", "تاريخ المعاجة","الرقم°",))
            self.textEdit.verticalHeader().setVisible(False)
            self.textEdit.setGeometry(QtCore.QRect(20, 320, 1323, 310))
            self.textEdit.setObjectName("textEdit")

            self.totalfont = QtGui.QFont()
            self.totalfont.setPointSize(8)
            self.totalfont.setBold(True)
            self.totalfont.bold()

            self.phtxt = QtWidgets.QLabel("PHOSTOXIN(PH3)", MainWindow)
            # self.phtxt.setStyleSheet("background-color:QLinearGradient( x1: 0, y1: 0, x2: 0, y2: 1, stop: 0 #565656, stop: 0.1 #525252, stop: 0.5 #4e4e4e, stop: 0.9 #4a4a4a, stop: 1 #464646);")
            self.phtxt.setFrameShape(QtWidgets.QFrame.Shape.WinPanel.VLine)
            self.phtxt.setGeometry(20, 640, 245, 50)
            self.phtxt.setFont(self.totalfont)

            self.phinitial = QtWidgets.QDoubleSpinBox(readOnly=True)
            self.phinitial.setPrefix("المخزون الاولي: ")
            self.phinitial.setGeometry(160, 645, 140, 20)
            self.phinitial.setRange(0, 900000)
            self.phinitial.setDecimals(3)

            self.phfinal = QtWidgets.QDoubleSpinBox(MainWindow, readOnly=True)
            self.phfinal.setPrefix("المخزون النهائي: ")
            self.phfinal.setDecimals(3)

            self.phfinal.setGeometry(118, 656, 140, 20)
            self.phfinal.setRange(0, 900000)

            self.ACTELLICtxt = QtWidgets.QLabel("ACTELLIC", MainWindow)
            # self.ACTELLICtxt.setStyleSheet("background-color:QLinearGradient( x1: 0, y1: 0, x2: 0, y2: 1, stop: 0 #565656, stop: 0.1 #525252, stop: 0.5 #4e4e4e, stop: 0.9 #4a4a4a, stop: 1 #464646);")
            self.ACTELLICtxt.setFrameShape(QtWidgets.QFrame.Shape.WinPanel.VLine)
            self.ACTELLICtxt.setGeometry(267, 640, 205, 50)
            self.ACTELLICtxt.setFont(self.totalfont)

            self.ACTELLICinitial = QtWidgets.QDoubleSpinBox(readOnly=True)
            self.ACTELLICinitial.setPrefix("المخزون الاولي: ")

            self.ACTELLICinitial.setGeometry(420, 645, 140, 20)
            self.ACTELLICinitial.setRange(0, 90000.00)

            self.ACTELLICfinal = QtWidgets.QDoubleSpinBox(MainWindow, readOnly=True)
            self.ACTELLICfinal.setPrefix("المخزون النهائي: ")

            self.ACTELLICfinal.setGeometry(323, 656, 140, 20)
            self.ACTELLICfinal.setRange(0, 90000.00)

            self.CIRATHRINEtxt = QtWidgets.QLabel("CIRATHRINE", MainWindow)
            # self.CIRATHRINEtxt.setStyleSheet("background-color:QLinearGradient( x1: 0, y1: 0, x2: 0, y2: 1, stop: 0 #565656, stop: 0.1 #525252, stop: 0.5 #4e4e4e, stop: 0.9 #4a4a4a, stop: 1 #464646);")
            self.CIRATHRINEtxt.setFrameShape(QtWidgets.QFrame.Shape.WinPanel.VLine)
            self.CIRATHRINEtxt.setGeometry(693, 640, 220, 50)
            self.CIRATHRINEtxt.setFont(self.totalfont)

            self.CIRATHRINEinitial = QtWidgets.QDoubleSpinBox(readOnly=True)
            self.CIRATHRINEinitial.setPrefix("المخزون الاولي: ")

            self.CIRATHRINEinitial.setGeometry(690, 645, 140, 20)
            self.CIRATHRINEinitial.setRange(0, 90000.00)

            self.CIRATHRINEfinal = QtWidgets.QDoubleSpinBox(MainWindow, readOnly=True)
            self.CIRATHRINEfinal.setPrefix("المخزون النهائي: ")

            self.CIRATHRINEfinal.setGeometry(763, 656, 140, 20)
            self.CIRATHRINEfinal.setRange(0, 90000.00)

            self.DEKATRINEtxt = QtWidgets.QLabel("DEKATRINE", MainWindow)
            # self.DEKATRINEtxt.setStyleSheet("background-color:QLinearGradient( x1: 0, y1: 0, x2: 0, y2: 1, stop: 0 #565656, stop: 0.1 #525252, stop: 0.5 #4e4e4e, stop: 0.9 #4a4a4a, stop: 1 #464646);")
            self.DEKATRINEtxt.setFrameShape(QtWidgets.QFrame.Shape.WinPanel.VLine)
            self.DEKATRINEtxt.setGeometry(475, 640, 215, 50)
            self.DEKATRINEtxt.setFont(self.totalfont)

            self.DEKATRINEinitial = QtWidgets.QDoubleSpinBox(readOnly=True)
            self.DEKATRINEinitial.setPrefix("المخزون الاولي: ")

            self.DEKATRINEinitial.setGeometry(690, 645, 140, 20)
            self.DEKATRINEinitial.setRange(0, 90000.00)

            self.DEKATRINEfinal = QtWidgets.QDoubleSpinBox(MainWindow, readOnly=True)
            self.DEKATRINEfinal.setPrefix("المخزون النهائي: ")

            self.DEKATRINEfinal.setGeometry(540, 656, 140, 20)
            self.DEKATRINEfinal.setRange(0, 90000.00)

            self.RATICIDEtxt = QtWidgets.QLabel("RATICIDE", MainWindow)
            # self.RATICIDEtxt.setStyleSheet("background-color:QLinearGradient( x1: 0, y1: 0, x2: 0, y2: 1, stop: 0 #565656, stop: 0.1 #525252, stop: 0.5 #4e4e4e, stop: 0.9 #4a4a4a, stop: 1 #464646);")
            self.RATICIDEtxt.setFrameShape(QtWidgets.QFrame.Shape.WinPanel.VLine)
            self.RATICIDEtxt.setGeometry(916, 640, 203, 50)
            self.RATICIDEtxt.setFont(self.totalfont)

            self.RATICIDEinitial = QtWidgets.QDoubleSpinBox(readOnly=True)
            self.RATICIDEinitial.setPrefix("المخزون الاولي: ")

            self.RATICIDEinitial.setGeometry(940, 645, 140, 20)
            self.RATICIDEinitial.setRange(0, 90000.00)

            self.RATICIDEfinal = QtWidgets.QDoubleSpinBox(MainWindow, readOnly=True)
            self.RATICIDEfinal.setPrefix("المخزون النهائي: ")

            self.RATICIDEfinal.setGeometry(970, 656, 140, 20)
            self.RATICIDEfinal.setRange(0, 90000.00)

            self.TEXTOtxt = QtWidgets.QLabel("TEXTO", MainWindow)
            self.TEXTOtxt.setFrameShape(QtWidgets.QFrame.Shape.WinPanel.VLine)
            # self.TEXTOtxt.setStyleSheet("background-color:QLinearGradient( x1: 0, y1: 0, x2: 0, y2: 1, stop: 0 #565656, stop: 0.1 #525252, stop: 0.5 #4e4e4e, stop: 0.9 #4a4a4a, stop: 1 #464646);")
            self.TEXTOtxt.setGeometry(1123, 640, 220, 50)
            self.TEXTOtxt.setFont(self.totalfont)

            self.TEXTOinitial = QtWidgets.QDoubleSpinBox(readOnly=True)
            self.TEXTOinitial.setPrefix("المخزون الاولي: ")

            self.TEXTOinitial.setGeometry(1185, 645, 140, 20)
            self.TEXTOinitial.setRange(0, 90000.00)

            self.TEXTOfinal = QtWidgets.QDoubleSpinBox(MainWindow, readOnly=True)
            self.TEXTOfinal.setPrefix("المخزون النهائي:")

            # self.TEXTOfinal.setStyleSheet("background-color: #88ffaa;color:#000000; border: 2px solid bleu ;border-radius: 4px;padding: 0px;")
            self.TEXTOfinal.setGeometry(1165, 656, 140, 20)
            self.TEXTOfinal.setRange(0, 90000.00)

            self.impot_all()
            self.importph()
            self.importactilic()
            self.importraticide()
            self.importtexto()
            self.importcirathrine()
            self.importdekatrine()


        ####################################################################################################################################
        def SITUATION_phyto(self, columns=None):
            try:
                selected_date = self.date_edit.date().toString("yyyy-MM")
                if self.produitphytofiltre.currentText() not in ['PHOSTOXIN(PH3)', 'ACTELLIC', 'CIRATHRINE', 'RATICIDE',
                                                                 'TEXTO'] and selected_date == "2000-01":
                    conn = mysql.connector.connect(host="localhost", user=user, password=password,database="datta")
                    cursor = conn.cursor()
                    cursor.execute(
                        "SELECT   LIEU_DE_TRAITEMENT, NATURE_ET_QUANTITE_DE_PRODUIT_TRAITE, DATE_DE_PEREMPTION, STOCK_FINAL, QUANTITE_UTILISEE, QUANTITE_ACHETE, STOCK_INITIAL, PRODUITS_PHYTOSANITAIRES, LIEU_DE_STOCKAGE,le FROM phytotablebidi")
                    rows = cursor.fetchall()
                    ########################################################################
                    # Create a new document and add a table
                    self.doc = docx.Document()
                    section = self.doc.sections[0]
                    section.orientation = WD_ORIENTATION.LANDSCAPE
                    section.page_width = docx.shared.Cm(29.7)
                    section.page_height = docx.shared.Cm(21.0)
                    section.top_margin = docx.shared.Cm(1.5)
                    section.bottom_margin = docx.shared.Cm(1.5)
                    section.left_margin = docx.shared.Cm(1)
                    section.right_margin = docx.shared.Cm(1)
                    # add second heading
                    paragraph = self.doc.add_paragraph()
                    paragraph.style.font.name = 'Times New Roman'
                    paragraph.style.font.size = Pt(20)
                    paragraph.style.font.bold = True
                    paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                    reshap1 = arabic_reshaper.reshape('حالة مخزون مواد الصحة النباتية')
                    bidi2 = bidi.algorithm.get_display(reshap1)
                    left_run1 = paragraph.add_run(bidi2)
                    reshap = arabic_reshaper.reshape('تعاونية الحبوب والبقول الجافة غليزان\n \n')
                    bidi1 = bidi.algorithm.get_display(reshap)
                    heading1 = self.doc.add_heading(bidi1, level=1)
                    heading1.style.font.name = 'Times New Roman'
                    heading1.style.font.size = Pt(14)
                    heading1.style.font.bold = True
                    heading1.style.font.color.rgb = RGBColor(0, 0, 0)
                    heading1.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT

                    table = self.doc.add_table(rows=1, cols=10)
                    table.style = "Table Grid"  # set the table style
                    hdr_cells = table.rows[0].cells
                    t1 = arabic_reshaper.reshape("مكان المعالجة")
                    bidi0 = bidi.algorithm.get_display(t1)
                    t2 = arabic_reshaper.reshape("نوع وكمية المادة المعالجة")
                    bidi2 = bidi.algorithm.get_display(t2)
                    t3 = arabic_reshaper.reshape("تاريخ نهاية الصلاحية")
                    bidi3 = bidi.algorithm.get_display(t3)
                    t4 = arabic_reshaper.reshape("المخزون النهائي")
                    bidi4 = bidi.algorithm.get_display(t4)
                    t5 = arabic_reshaper.reshape("الكمية المستعملة")
                    bidi5 = bidi.algorithm.get_display(t5)
                    t6 = arabic_reshaper.reshape("الكمية المشترات")
                    bidi6 = bidi.algorithm.get_display(t6)
                    t7 = arabic_reshaper.reshape("المخزون الابتدائي")
                    bidi7 = bidi.algorithm.get_display(t7)
                    t8 = arabic_reshaper.reshape(" المواد الكيماوية")
                    bidi8 = bidi.algorithm.get_display(t8)
                    t9 = arabic_reshaper.reshape("مكان التخزين")
                    bidi9 = bidi.algorithm.get_display(t9)
                    t10 = arabic_reshaper.reshape("تاريخ المعالجة")
                    bidi10 = bidi.algorithm.get_display(t10)
                    # Add text to the header cells
                    hdr_cells[0].text = bidi0
                    hdr_cells[1].text = bidi2
                    hdr_cells[2].text = bidi3
                    hdr_cells[3].text = bidi4
                    hdr_cells[4].text = bidi5
                    hdr_cells[5].text = bidi6
                    hdr_cells[6].text = bidi7
                    hdr_cells[7].text = bidi8
                    hdr_cells[8].text = bidi9
                    hdr_cells[9].text = bidi10
                    # Set the width of the header cells
                    hdr_cells[0].width = Inches(2)  # set the width of the first column to 1.5 inches
                    hdr_cells[1].width = Inches(2.4)
                    hdr_cells[2].width = Inches(1.8)
                    hdr_cells[3].width = Inches(1.4)
                    hdr_cells[4].width = Inches(1.4)
                    hdr_cells[5].width = Inches(1.4)
                    hdr_cells[6].width = Inches(1.5)
                    hdr_cells[7].width = Inches(1.2)
                    hdr_cells[8].width = Inches(2)
                    hdr_cells[9].width = Inches(1.5)

                    for cell in table.rows[0].cells:
                        cell.paragraphs[0].runs[0].font.name = 'Arabic Typesetting'
                        cell.paragraphs[0].runs[0].font.size = Pt(9)
                        cell.paragraphs[0].runs[0].font.bold = True

                    # Add the data to the table
                    if rows:  # Check if the 'rows' list is not empty
                        table_cols = len(rows[0])
                        for row in range(len(rows)):
                            table.add_row()
                            for col in range(table_cols):
                                cell = table.cell(row + 1, col)
                                cell_text = str(rows[row][col]).replace(" ", "\n")
                                cell.text = str(rows[row][col])

                                # Set font properties
                                cell.paragraphs[0].runs[0].font.name = 'Times New Roman (Titres CS)'
                                cell.paragraphs[0].runs[0].font.size = Pt(11)
                                cell.paragraphs[0].runs[0].font.bold = True
                                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                            for row in table.rows:
                                for cell in row.cells:
                                    for paragraph in cell.paragraphs:
                                        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT


                    else:
                        # Handle the case when there are no rows in the 'rows' list
                        table.add_row()
                        cell = table.cell(1, 0)
                        celltext = "لا توجد بيانات"
                        err = arabic_reshaper.reshape(celltext)
                        errbidi = bidi.algorithm.get_display(err)
                        cell.text = errbidi

                    for row in table.rows:
                        row.height = Inches(0.6)
                    # Save and open the document for printing
                    path, _ = QFileDialog.getSaveFileName(None, "Enregistrer la fiche", "", "Fichiers DOCX (*.docx)")
                    if path:
                        self.doc.save(path)
                    conn.close()
                    msgbox = QtWidgets.QMessageBox()
                    msgbox.setWindowTitle('تاكيد العملية')
                    msgbox.setText('تم حفظ الملف بنجاح')
                    msgbox.exec()
                elif selected_date and self.produitphytofiltre.currentText() not in ['PHOSTOXIN(PH3)', 'ACTELLIC',
                                                                                     'CIRATHRINE', 'RATICIDE', 'TEXTO']:

                    selected_month = self.date_edit.date().month()
                    # Create a list of Arabic month names
                    arabic_month_names = ["جانفي", "فيفري", "مارس", "أفريل", "ماي", "جوان", "جويلية", "أوت", "سبتمبر",
                                          "أكتوبر", "نوفمبر", "ديسمبر"]
                    # Create a custom formatted Arabic month string
                    formatted_month = arabic_month_names[selected_month - 1]
                    mois = formatted_month
                    conn = mysql.connector.connect(host="localhost", user=user, password=password,database="datta")
                    cursor = conn.cursor()
                    cursor.execute(
                        """SELECT
                                      REPLACE(group_concat(LIEU_DE_TRAITEMENT), ',', '                                     ') AS LIEU_DE_TRAITEMENT,
                                      REPLACE(group_concat(NATURE_ET_QUANTITE_DE_PRODUIT_TRAITE), ',', '                ') AS NATURE_ET_QUANTITE_DE_PRODUIT_TRAITE,
                                      DATE_DE_PEREMPTION,
                                      STOCK_FINAL,
                                      SUM(QUANTITE_UTILISEE) AS QUANTITE_UTILISEE,
                                      QUANTITE_ACHETE,
                                      STOCK_INITIAL,
                                      PRODUITS_PHYTOSANITAIRES,   
                                      LIEU_DE_STOCKAGE
                                      FROM phytotablebidi WHERE DATE_FORMAT('%Y-%m', le) = %s GROUP BY PRODUITS_PHYTOSANITAIRES;""",
                        (selected_date,))
                    rows = cursor.fetchall()
                    ########################################################################
                    # Create a new document and add a table
                    self.doc = docx.Document()
                    section = self.doc.sections[0]
                    section.orientation = WD_ORIENTATION.LANDSCAPE
                    section.page_width = docx.shared.Cm(29.7)
                    section.page_height = docx.shared.Cm(21.0)
                    section.top_margin = docx.shared.Cm(0.5)
                    section.bottom_margin = docx.shared.Cm(1.5)
                    section.left_margin = docx.shared.Cm(1)
                    section.right_margin = docx.shared.Cm(1)
                    # add second heading

                    reshap = arabic_reshaper.reshape('حالة مخزون مواد الصحة النباتية  \n ')
                    bidi1 = bidi.algorithm.get_display(reshap)
                    heading1 = self.doc.add_heading(bidi1, level=1)
                    heading1.style.font.name = 'Times New Roman'
                    heading1.style.font.size = Pt(20)
                    heading1.style.font.bold = True
                    heading1.style.font.color.rgb = RGBColor(0, 0, 0)
                    heading1.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

                    paragraph = self.doc.add_paragraph()
                    paragraph.style.font.name = 'Times New Roman'
                    paragraph.style.font.size = Pt(14)
                    paragraph.style.font.bold = True
                    paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                    reshapl = arabic_reshaper.reshape('شهر:' + mois)
                    bidil = bidi.algorithm.get_display(reshapl)
                    left_run = paragraph.add_run(bidil)

                    reshap1 = arabic_reshaper.reshape(
                        'تعاونية الحبوب والبقول الجافة غليزان\t\t\t\t\t\t\t\t\t\t\t\t\t\n')
                    bidi2 = bidi.algorithm.get_display(reshap1)
                    left_run1 = paragraph.add_run(bidi2)
                    table = self.doc.add_table(rows=1, cols=9)
                    table.style = "Table Grid"  # set the table style
                    hdr_cells = table.rows[0].cells
                    t1 = arabic_reshaper.reshape("مكان المعالجة")
                    bidi0 = bidi.algorithm.get_display(t1)
                    t2 = arabic_reshaper.reshape("نوع وكمية المادة المعالجة")
                    bidi2 = bidi.algorithm.get_display(t2)
                    t3 = arabic_reshaper.reshape(" تاريخ نهاية الصلاحية")
                    bidi3 = bidi.algorithm.get_display(t3)
                    t4 = arabic_reshaper.reshape(" المخزون النهائي")
                    bidi4 = bidi.algorithm.get_display(t4)
                    t5 = arabic_reshaper.reshape("الكمية المستعملة")
                    bidi5 = bidi.algorithm.get_display(t5)
                    t6 = arabic_reshaper.reshape("الكمية المشترات")
                    bidi6 = bidi.algorithm.get_display(t6)
                    t7 = arabic_reshaper.reshape("المخزون الابتدائي")
                    bidi7 = bidi.algorithm.get_display(t7)
                    t8 = arabic_reshaper.reshape("المواد الكيماوية")
                    bidi8 = bidi.algorithm.get_display(t8)
                    t9 = arabic_reshaper.reshape("مكان التخزين")
                    bidi9 = bidi.algorithm.get_display(t9)
                    # Add text to the header cells
                    hdr_cells[0].text = bidi0
                    hdr_cells[1].text = bidi2
                    hdr_cells[2].text = bidi3
                    hdr_cells[3].text = bidi4
                    hdr_cells[4].text = bidi5
                    hdr_cells[5].text = bidi6
                    hdr_cells[6].text = bidi7
                    hdr_cells[7].text = bidi8
                    hdr_cells[8].text = bidi9
                    # Set the width of the header cells
                    hdr_cells[0].width = Inches(2)  # set the width of the first column to 1.5 inches
                    hdr_cells[1].width = Inches(2.4)
                    hdr_cells[2].width = Inches(1.8)
                    hdr_cells[3].width = Inches(1.4)
                    hdr_cells[4].width = Inches(1.4)
                    hdr_cells[5].width = Inches(1.4)
                    hdr_cells[6].width = Inches(1.5)
                    hdr_cells[7].width = Inches(1.2)
                    hdr_cells[8].width = Inches(2)

                    for cell in table.rows[0].cells:
                        cell.paragraphs[0].runs[0].font.name = 'Calibri'
                        cell.paragraphs[0].runs[0].font.size = Pt(11)
                        cell.paragraphs[0].runs[0].font.bold = True
                    # Add the data to the table
                    if rows:  # Check if the 'rows' list is not empty
                        table_cols = len(rows[0])
                        for row in range(len(rows)):
                            table.add_row()
                            for col in range(table_cols):
                                cell = table.cell(row + 1, col)
                                cell.text = str(rows[row][col])
                                # Set font properties
                                cell.paragraphs[0].runs[0].font.name = 'Calibri'
                                cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                                cell.paragraphs[0].runs[0].font.bold = False
                                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                            for row in table.rows:
                                for cell in row.cells:
                                    for paragraph in cell.paragraphs:
                                        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        # Handle the case when there are no rows in the 'rows' list
                        table.add_row()
                        cell = table.cell(1, 0)
                        cell.text = "No data available."
                    for row in table.rows:
                        row.height = Inches(0.6)
                    # Save and open the document for printing
                    path, _ = QFileDialog.getSaveFileName(None, "Enregistrer la fiche", "", "Fichiers DOCX (*.docx)")
                    if path:
                        self.doc.save(path)
                    conn.close()
                    msgbox = QtWidgets.QMessageBox()
                    msgbox.setWindowTitle('تاكيد العملية')
                    msgbox.setText('تم حفظ الملف بنجاح')
                    msgbox.exec()


                elif self.produitphytofiltre.currentText() and selected_date == "01-01-2000":
                    produitphyto = self.produitphytofiltre.currentText()
                    selected_date = self.date_edit.date().toString("dd-MM-yyyy")
                    locale.setlocale(locale.LC_ALL, 'fr_FR.utf8')
                    QtCore.QLocale.setDefault(QtCore.QLocale(QtCore.QLocale.Language.French))
                    selected_date1 = self.date_edit.date().toPyDate()
                    mois = selected_date1.DATE_FORMAT("%B").lower()

                    conn = mysql.connector.connect(host="localhost", user=user, password=password,database="datta")
                    cursor = conn.cursor()
                    cursor.execute(
                        """"SELECT
                                      REPLACE(group_concat(LIEU_DE_TRAITEMENT), ',', '                                     ') AS LIEU_DE_TRAITEMENT,
                                      REPLACE(group_concat(NATURE_ET_QUANTITE_DE_PRODUIT_TRAITE), ',', '  ') AS NATURE_ET_QUANTITE_DE_PRODUIT_TRAITE,
                                      REPLACE(group_concat(DATE_DE_PEREMPTION), ',', ',') AS DATE_DE_PEREMPTION,
                                      STOCK_FINAL,
                                      SUM(QUANTITE_UTILISEE) AS QUANTITE_UTILISEE,
                                      QUANTITE_ACHETE,
                                      STOCK_INITIAL,
                                      PRODUITS_PHYTOSANITAIRES,
                                      LIEU_DE_STOCKAGE
                                      FROM phytotablebidi  WHERE   PRODUITS_PHYTOSANITAIRES= %s GROUP BY PRODUITS_PHYTOSANITAIRES;""",
                        (produitphyto,))
                    rows = cursor.fetchall()

                    ########################################################################
                    # Create a new document and add a table
                    self.doc = docx.Document()
                    section = self.doc.sections[0]
                    section.orientation = WD_ORIENTATION.LANDSCAPE
                    section.page_width = docx.shared.Cm(29.7)
                    section.page_height = docx.shared.Cm(21.0)
                    section.top_margin = docx.shared.Cm(0.5)
                    section.bottom_margin = docx.shared.Cm(1.5)
                    section.left_margin = docx.shared.Cm(1)
                    section.right_margin = docx.shared.Cm(1)
                    # add second heading

                    reshap = arabic_reshaper.reshape('حالة مخزون مواد الصحة النباتية  \n ')
                    bidi1 = bidi.algorithm.get_display(reshap)
                    heading1 = self.doc.add_heading(bidi1, level=1)
                    heading1.style.font.name = 'Times New Roman'
                    heading1.style.font.size = Pt(20)
                    heading1.style.font.bold = True
                    heading1.style.font.color.rgb = RGBColor(0, 0, 0)
                    heading1.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

                    paragraph = self.doc.add_paragraph()
                    paragraph.style.font.name = 'Times New Roman'
                    paragraph.style.font.size = Pt(14)
                    paragraph.style.font.bold = True
                    paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                    reshapl = arabic_reshaper.reshape('منتج:' + produitphyto)
                    bidil = bidi.algorithm.get_display(reshapl)
                    left_run = paragraph.add_run(bidil)

                    reshap1 = arabic_reshaper.reshape('تعاونية الحبوب والبقول الجافة غليزان\t\t\t\t\t\t\t\t\t\t\t\t\n')
                    bidi2 = bidi.algorithm.get_display(reshap1)
                    left_run1 = paragraph.add_run(bidi2)

                    table = self.doc.add_table(rows=1, cols=9)
                    table.style = "Table Grid"  # set the table style
                    hdr_cells = table.rows[0].cells
                    t1 = arabic_reshaper.reshape("مكان المعالجة")
                    bidi0 = bidi.algorithm.get_display(t1)
                    t2 = arabic_reshaper.reshape("نوع وكمية المادة المعالجة")
                    bidi2 = bidi.algorithm.get_display(t2)
                    t3 = arabic_reshaper.reshape("تاريخ نهاية الصلاحية")
                    bidi3 = bidi.algorithm.get_display(t3)
                    t4 = arabic_reshaper.reshape("المخزون النهائي")
                    bidi4 = bidi.algorithm.get_display(t4)
                    t5 = arabic_reshaper.reshape("الكمية المستعملة")
                    bidi5 = bidi.algorithm.get_display(t5)
                    t6 = arabic_reshaper.reshape("الكمية المشترات")
                    bidi6 = bidi.algorithm.get_display(t6)
                    t7 = arabic_reshaper.reshape("المخزون الابتدائي")
                    bidi7 = bidi.algorithm.get_display(t7)
                    t8 = arabic_reshaper.reshape(" المواد الكيماوية")
                    bidi8 = bidi.algorithm.get_display(t8)
                    t9 = arabic_reshaper.reshape("مكان التخزين")
                    bidi9 = bidi.algorithm.get_display(t9)
                    # Add text to the header cells
                    hdr_cells[0].text = bidi0
                    hdr_cells[1].text = bidi2
                    hdr_cells[2].text = bidi3
                    hdr_cells[3].text = bidi4
                    hdr_cells[4].text = bidi5
                    hdr_cells[5].text = bidi6
                    hdr_cells[6].text = bidi7
                    hdr_cells[7].text = bidi8
                    hdr_cells[8].text = bidi9
                    # Set the width of the header cells
                    hdr_cells[0].width = Inches(2)  # set the width of the first column to 1.5 inches
                    hdr_cells[1].width = Inches(2.4)
                    hdr_cells[2].width = Inches(1.8)
                    hdr_cells[3].width = Inches(1.4)
                    hdr_cells[4].width = Inches(1.4)
                    hdr_cells[5].width = Inches(1.4)
                    hdr_cells[6].width = Inches(1.5)
                    hdr_cells[7].width = Inches(1.2)
                    hdr_cells[8].width = Inches(2)
                    for cell in table.rows[0].cells:
                        cell.paragraphs[0].runs[0].font.name = 'Calibri'
                        cell.paragraphs[0].runs[0].font.size = Pt(11)
                        cell.paragraphs[0].runs[0].font.bold = True
                    # Add the data to the table
                    if rows:  # Check if the 'rows' list is not empty
                        table_cols = len(rows[0])
                        for row in range(len(rows)):
                            table.add_row()
                            for col in range(table_cols):
                                cell = table.cell(row + 1, col)
                                cell.text = str(rows[row][col])
                                # Set font properties
                                cell.paragraphs[0].runs[0].font.name = 'Calibri'
                                cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                                cell.paragraphs[0].runs[0].font.bold = False
                                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    else:
                        # Handle the case when there are no rows in the 'rows' list
                        table.add_row()
                        cell = table.cell(1, 0)
                        cell.text = "No data available."
                    for row in table.rows:
                        row.height = Inches(0.6)
                    # Save and open the document for printing
                    path, _ = QFileDialog.getSaveFileName(None, "Enregistrer la fiche", "", "Fichiers DOCX (*.docx)")
                    if path:
                        self.doc.save(path)
                    conn.close()
                    msgbox = QtWidgets.QMessageBox()
                    msgbox.setWindowTitle('تاكيد العملية')
                    msgbox.setText('تم حفظ الملف بنجاح')
                    msgbox.exec()

                elif selected_date and self.produitphytofiltre.currentText():
                    produitphyto = self.produitphytofiltre.currentText()
                    selected_month = self.date_edit.date().month()
                    # Create a list of Arabic month names
                    arabic_month_names = ["جانفي", "فيفري", "مارس", "أفريل", "ماي", "جوان", "جويلية", "أوت", "سبتمبر",
                                          "أكتوبر", "نوفمبر", "ديسمبر"]
                    # Create a custom formatted Arabic month string
                    formatted_month = arabic_month_names[selected_month - 1]
                    mois = formatted_month

                    conn = mysql.connector.connect(host="localhost", user=user, password=password,database="datta")
                    cursor = conn.cursor()
                    cursor.execute("""SELECT
                                      REPLACE(group_concat(LIEU_DE_TRAITEMENT), ',', '                                     ') AS LIEU_DE_TRAITEMENT,
                                      REPLACE(group_concat(NATURE_ET_QUANTITE_DE_PRODUIT_TRAITE), ',', '  ') AS NATURE_ET_QUANTITE_DE_PRODUIT_TRAITE,
                                      REPLACE(group_concat(DATE_DE_PEREMPTION), ',', ',') AS DATE_DE_PEREMPTION,
                                      STOCK_FINAL,
                                      SUM(QUANTITE_UTILISEE) AS QUANTITE_UTILISEE,
                                      QUANTITE_ACHETE,
                                      STOCK_INITIAL,
                                      PRODUITS_PHYTOSANITAIRES,
                                      LIEU_DE_STOCKAGE
                                      FROM phytotablebidi  WHERE DATE_FORMAT('%Y-%m', le)=%s  and PRODUITS_PHYTOSANITAIRES=%s GROUP BY PRODUITS_PHYTOSANITAIRES ; """,
                                   (selected_date, produitphyto,))
                    rows = cursor.fetchall()

                    ########################################################################
                    # Create a new document and add a table
                    self.doc = docx.Document()
                    section = self.doc.sections[0]
                    section.orientation = WD_ORIENTATION.LANDSCAPE
                    section.page_width = docx.shared.Cm(29.7)
                    section.page_height = docx.shared.Cm(21.0)
                    section.top_margin = docx.shared.Cm(0.5)
                    section.bottom_margin = docx.shared.Cm(1.5)
                    section.left_margin = docx.shared.Cm(1)
                    section.right_margin = docx.shared.Cm(1)
                    # add second heading

                    reshap = arabic_reshaper.reshape('حالة مخزون مواد الصحة النباتية  \n ')
                    bidi1 = bidi.algorithm.get_display(reshap)
                    heading1 = self.doc.add_heading(bidi1, level=1)
                    heading1.style.font.name = 'Times New Roman'
                    heading1.style.font.size = Pt(20)
                    heading1.style.font.bold = True
                    heading1.style.font.color.rgb = RGBColor(0, 0, 0)
                    heading1.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

                    paragraph = self.doc.add_paragraph()
                    paragraph.style.font.name = 'Times New Roman'
                    paragraph.style.font.size = Pt(14)
                    paragraph.style.font.bold = True
                    paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                    reshapl = arabic_reshaper.reshape('شهر:' + mois)
                    bidil = bidi.algorithm.get_display(reshapl)
                    left_run = paragraph.add_run(bidil)

                    reshap1 = arabic_reshaper.reshape('تعاونية الحبوب والبقول الجافة غليزان\t\t\t\t\t\t\t\t\t\t\t\t\n')
                    bidi2 = bidi.algorithm.get_display(reshap1)
                    left_run1 = paragraph.add_run(bidi2)
                    table = self.doc.add_table(rows=1, cols=9)
                    table.style = "Table Grid"  # set the table style
                    hdr_cells = table.rows[0].cells
                    t1 = arabic_reshaper.reshape("مكان المعالجة")
                    bidi0 = bidi.algorithm.get_display(t1)
                    t2 = arabic_reshaper.reshape("نوع وكمية المادة المعالجة")
                    bidi2 = bidi.algorithm.get_display(t2)
                    t3 = arabic_reshaper.reshape(" تاريخ نهاية الصلاحية")
                    bidi3 = bidi.algorithm.get_display(t3)
                    t4 = arabic_reshaper.reshape(" المخزون النهائي")
                    bidi4 = bidi.algorithm.get_display(t4)
                    t5 = arabic_reshaper.reshape("الكمية المستعملة")
                    bidi5 = bidi.algorithm.get_display(t5)
                    t6 = arabic_reshaper.reshape("الكمية المشترات")
                    bidi6 = bidi.algorithm.get_display(t6)
                    t7 = arabic_reshaper.reshape("المخزون الابتدائي")
                    bidi7 = bidi.algorithm.get_display(t7)
                    t8 = arabic_reshaper.reshape("المواد الكيماوية")
                    bidi8 = bidi.algorithm.get_display(t8)
                    t9 = arabic_reshaper.reshape("مكان التخزين")
                    bidi9 = bidi.algorithm.get_display(t9)
                    # Add text to the header cells
                    hdr_cells[0].text = bidi0
                    hdr_cells[1].text = bidi2
                    hdr_cells[2].text = bidi3
                    hdr_cells[3].text = bidi4
                    hdr_cells[4].text = bidi5
                    hdr_cells[5].text = bidi6
                    hdr_cells[6].text = bidi7
                    hdr_cells[7].text = bidi8
                    hdr_cells[8].text = bidi9
                    # Set the width of the header cells
                    hdr_cells[0].width = Inches(2)  # set the width of the first column to 1.5 inches
                    hdr_cells[1].width = Inches(2.4)
                    hdr_cells[2].width = Inches(1.8)
                    hdr_cells[3].width = Inches(1.4)
                    hdr_cells[4].width = Inches(1.4)
                    hdr_cells[5].width = Inches(1.4)
                    hdr_cells[6].width = Inches(1.5)
                    hdr_cells[7].width = Inches(1.2)
                    hdr_cells[8].width = Inches(2)

                    for cell in table.rows[0].cells:
                        cell.paragraphs[0].runs[0].font.name = 'Calibri'
                        cell.paragraphs[0].runs[0].font.size = Pt(11)
                        cell.paragraphs[0].runs[0].font.bold = True
                    # Add the data to the table
                    if rows:  # Check if the 'rows' list is not empty
                        table_cols = len(rows[0])
                        for row in range(len(rows)):
                            table.add_row()
                            for col in range(table_cols):
                                cell = table.cell(row + 1, col)
                                cell.text = str(rows[row][col])
                                # Set font properties
                                cell.paragraphs[0].runs[0].font.name = 'Calibri'
                                cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                                cell.paragraphs[0].runs[0].font.bold = False
                                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                            for row in table.rows:
                                for cell in row.cells:
                                    for paragraph in cell.paragraphs:
                                        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT

                    else:
                        # Handle the case when there are no rows in the 'rows' list
                        table.add_row()
                        cell = table.cell(1, 0)
                        cell.text = "No data available."
                    for row in table.rows:
                        row.height = Inches(0.6)
                    # Save and open the document for printing
                    path, _ = QFileDialog.getSaveFileName(None, "Enregistrer la fiche", "", "Fichiers DOCX (*.docx)")
                    if path:
                        self.doc.save(path)
                    conn.close()
                    msgbox = QtWidgets.QMessageBox()
                    msgbox.setWindowTitle('تاكيد العملية')
                    msgbox.setText('تم حفظ الملف بنجاح')
                    msgbox.exec()

            except Exception as e:
                print(e)

        def print_docx(self):
            self.dialog = QtWidgets.QDialog()
            self.dialog.setStyleSheet(""" QWidget
                            {
                                color: #000000;
                                background-color: #ffffff;
                                border-width: 1px;
                                border-color: #1e1e1e;
                                border-style: solid;
                                border-radius: 6;
                                padding: 0px;
                                font-size: 18px;
                                padding-left: 1px;
                                padding-right: 1px
                            }
                            QWidget:item:hover
                            {
                                background-color: #3daee9;
                                color: #eff0f1;
                            }
                            QWidget:item:selected
                            {
                                background-color: #3daee9;
                            }
                            QWidget:disabled
                            {
                                color: #454545;
                                background-color: #31363b;
                            }
                            QPushButton
                            {
                                color: #000000;
                                background-color:#ade3e7;
                                border-width: 1px;
                                border-color: #1e1e1e;
                                border-style: solid;
                                border-radius: 6;
                                padding: 3px;
                                font-size: 12px;
                                padding-left: 5px;
                                padding-right: 5px;
                                min-width: 40px
                            }
                            QPushButton:disabled
                            {
                                background-color: #31363b;
                                border-width: 1px;
                                border-color: #454545;
                                border-style: solid;
                                padding-top: 5px;
                                padding-bottom: 5px;
                                padding-left: 10px;
                                padding-right: 10px;
                                border-radius: 2px;
                                color: #454545;
                            }

                            QPushButton:pressed
                            {
                                background-color: #3daee9;
                                padding-top: -15px;
                                padding-bottom: -17px;
                            }
                            QPushButton:hover
                            {
                                border: 1px solid #ff8c00;
                                color: #000000;
                            }
                             QLabel
                            {
                                font-size: 18px;
                                border: 0px solid orange;
                            }

                        """)
            self.dialog.setWindowTitle("اختر التاريخ والمنتج")
            self.dialog.setGeometry(200, 200, 475, 200)
            label0 = QtWidgets.QLabel("اختر التاريخ والمنتج", self.dialog)
            label0.setGeometry(130, 10, 180, 30)
            label1 = QtWidgets.QLabel("التاريخ", self.dialog)
            label1.setGeometry(75, 50, 80, 30)
            label2 = QtWidgets.QLabel("المنتج", self.dialog)
            label2.setGeometry(290, 50, 80, 30)
            self.date_edit = QtWidgets.QDateEdit(self.dialog)
            self.date_edit.setDisplayFormat("yyyy-MM-dd")
            self.date_edit.setDate(self.dateday)
            self.date_edit.setAlignment(QtCore.Qt.AlignmentFlag.AlignRight)

            self.date_edit.resize(180, 30)
            self.date_edit.move(40, 82)
            self.produitphytofiltre = QtWidgets.QComboBox(self.dialog)
            self.produitphytofiltre.setGeometry(250, 82, 180, 30)
            self.produitphytofiltre.addItem(" ")
            self.produitphytofiltre.addItem("PHOSTOXIN(PH3)")
            self.produitphytofiltre.addItem("ACTELLIC")
            self.produitphytofiltre.addItem("CIRATHRINE")
            self.produitphytofiltre.addItem("RATICIDE")
            self.produitphytofiltre.addItem("TEXTO")
            self.ok_button = QtWidgets.QPushButton("طباعة", self.dialog, clicked=lambda: self.print_docx_situation())
            self.ok_button.setGeometry(310, 140, 80, 40)
            cancel_button = QtWidgets.QPushButton("خروج", self.dialog, clicked=lambda: self.dialog.close())
            cancel_button.setGeometry(90, 140, 80, 40)
            self.print_button = QtWidgets.QPushButton("حفظ", self.dialog, clicked=lambda: self.SITUATION_phyto())
            self.print_button.setGeometry(200, 140, 80, 40)
            self.dialog.exec()

        def print_docx_situation(self):
            try:
                selected_date = self.date_edit.date().toString("yyyy-MM")
                if self.produitphytofiltre.currentText() not in ['PHOSTOXIN(PH3)', 'ACTELLIC', 'CIRATHRINE', 'RATICIDE',
                                                                 'TEXTO'] and selected_date == "2000-01":
                    conn = mysql.connector.connect(host="localhost", user=user, password=password,database="datta")
                    cursor = conn.cursor()
                    cursor.execute(
                        "SELECT   LIEU_DE_TRAITEMENT, NATURE_ET_QUANTITE_DE_PRODUIT_TRAITE, DATE_DE_PEREMPTION, STOCK_FINAL, QUANTITE_UTILISEE, QUANTITE_ACHETE, STOCK_INITIAL, PRODUITS_PHYTOSANITAIRES, LIEU_DE_STOCKAGE,le FROM phytotablebidi")
                    rows = cursor.fetchall()
                    ########################################################################
                    # Create a new document and add a table
                    self.doc = docx.Document()
                    section = self.doc.sections[0]
                    section.orientation = WD_ORIENTATION.LANDSCAPE
                    section.page_width = docx.shared.Cm(29.7)
                    section.page_height = docx.shared.Cm(21.0)
                    section.top_margin = docx.shared.Cm(1.5)
                    section.bottom_margin = docx.shared.Cm(1.5)
                    section.left_margin = docx.shared.Cm(1)
                    section.right_margin = docx.shared.Cm(1)
                    # add second heading
                    paragraph = self.doc.add_paragraph()
                    paragraph.style.font.name = 'Times New Roman'
                    paragraph.style.font.size = Pt(20)
                    paragraph.style.font.bold = True
                    paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                    reshap1 = arabic_reshaper.reshape('حالة مخزون مواد الصحة النباتية')
                    bidi2 = bidi.algorithm.get_display(reshap1)
                    left_run1 = paragraph.add_run(bidi2)
                    reshap = arabic_reshaper.reshape('تعاونية الحبوب والبقول الجافة غليزان\n\n')
                    bidi1 = bidi.algorithm.get_display(reshap)
                    heading1 = self.doc.add_heading(bidi1, level=1)
                    heading1.style.font.name = 'Times New Roman'
                    heading1.style.font.size = Pt(14)
                    heading1.style.font.bold = True
                    heading1.style.font.color.rgb = RGBColor(0, 0, 0)
                    heading1.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT

                    table = self.doc.add_table(rows=1, cols=10)
                    table.style = "Table Grid"  # set the table style
                    hdr_cells = table.rows[0].cells
                    t1 = arabic_reshaper.reshape("مكان المعالجة")
                    bidi0 = bidi.algorithm.get_display(t1)
                    t2 = arabic_reshaper.reshape("نوع وكمية المادة المعالجة")
                    bidi2 = bidi.algorithm.get_display(t2)
                    t3 = arabic_reshaper.reshape("تاريخ نهاية الصلاحية")
                    bidi3 = bidi.algorithm.get_display(t3)
                    t4 = arabic_reshaper.reshape("المخزون النهائي")
                    bidi4 = bidi.algorithm.get_display(t4)
                    t5 = arabic_reshaper.reshape("الكمية المستعملة")
                    bidi5 = bidi.algorithm.get_display(t5)
                    t6 = arabic_reshaper.reshape("الكمية المشترات")
                    bidi6 = bidi.algorithm.get_display(t6)
                    t7 = arabic_reshaper.reshape("المخزون الابتدائي")
                    bidi7 = bidi.algorithm.get_display(t7)
                    t8 = arabic_reshaper.reshape(" المواد الكيماوية")
                    bidi8 = bidi.algorithm.get_display(t8)
                    t9 = arabic_reshaper.reshape("مكان التخزين")
                    bidi9 = bidi.algorithm.get_display(t9)
                    t10 = arabic_reshaper.reshape("تاريخ المعالجة")
                    bidi10 = bidi.algorithm.get_display(t10)
                    # Add text to the header cells
                    hdr_cells[0].text = bidi0
                    hdr_cells[1].text = bidi2
                    hdr_cells[2].text = bidi3
                    hdr_cells[3].text = bidi4
                    hdr_cells[4].text = bidi5
                    hdr_cells[5].text = bidi6
                    hdr_cells[6].text = bidi7
                    hdr_cells[7].text = bidi8
                    hdr_cells[8].text = bidi9
                    hdr_cells[9].text = bidi10
                    # Set the width of the header cells
                    hdr_cells[0].width = Inches(2)  # set the width of the first column to 1.5 inches
                    hdr_cells[1].width = Inches(2.4)
                    hdr_cells[2].width = Inches(1.8)
                    hdr_cells[3].width = Inches(1.4)
                    hdr_cells[4].width = Inches(1.4)
                    hdr_cells[5].width = Inches(1.4)
                    hdr_cells[6].width = Inches(1.5)
                    hdr_cells[7].width = Inches(1.2)
                    hdr_cells[8].width = Inches(2)
                    hdr_cells[9].width = Inches(1.5)

                    for cell in table.rows[0].cells:
                        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                        cell.paragraphs[0].runs[0].font.size = Pt(10)
                        cell.paragraphs[0].runs[0].font.bold = True

                    # Add the data to the table
                    if rows:  # Check if the 'rows' list is not empty
                        table_cols = len(rows[0])
                        for row in range(len(rows)):
                            table.add_row()
                            for col in range(table_cols):
                                cell = table.cell(row + 1, col)
                                cell_text = str(rows[row][col]).replace(" ", "\n")
                                cell.text = str(rows[row][col])

                                # Set font properties
                                cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                                cell.paragraphs[0].runs[0].font.size = Pt(9)
                                cell.paragraphs[0].runs[0].font.bold = True
                                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                            for row in table.rows:
                                for cell in row.cells:
                                    for paragraph in cell.paragraphs:
                                        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        # Handle the case when there are no rows in the 'rows' list
                        table.add_row()
                        cell = table.cell(1, 0)
                        celltext = "لا توجد بيانات"
                        err = arabic_reshaper.reshape(celltext)
                        errbidi = bidi.algorithm.get_display(err)
                        cell.text = errbidi

                    for row in table.rows:
                        row.height = Inches(0.6)

                    # Save and open the document for printing
                    temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                    doc_pdf = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False).name
                    self.doc.save(temp_file)
                    if temp_file:
                        a = self.progress_bar()
                        os.startfile(temp_file, "open")
                        # docx2pdf.convert(temp_file,doc_pdf)
                        # app_path = 'C:\\Program Files\\ONLYOFFICE\\DesktopEditors\\DesktopEditors.exe'
                        # subprocess.Popen([app_path,doc_pdf])
                        # word = win32a.gencache.EnsureDispatch("Word.Application")

                        # doc = word.Documents.Open(temp_file)
                        # doc.PrintOut()
                        # doc.Close()
                        # word.Quit()
                    conn.close()

                    # subprocess.Popen([app_path,temp_file])
                    # os.startfile(temp_file,"print")
                elif selected_date and self.produitphytofiltre.currentText() not in ['PHOSTOXIN(PH3)', 'ACTELLIC',
                                                                                     'CIRATHRINE', 'RATICIDE', 'TEXTO']:
                    selected_month = self.date_edit.date().month()
                    # Create a list of Arabic month names
                    arabic_month_names = ["جانفي", "فيفري", "مارس", "أفريل", "ماي", "جوان", "جويلية", "أوت", "سبتمبر",
                                          "أكتوبر", "نوفمبر", "ديسمبر"]
                    # Create a custom formatted Arabic month string
                    formatted_month = arabic_month_names[selected_month - 1]
                    mois = formatted_month

                    conn = mysql.connector.connect(host="localhost", user=user, password=password,database="datta")
                    cursor = conn.cursor()
                    cursor.execute(
                        """SELECT
                                      REPLACE(group_concat(LIEU_DE_TRAITEMENT), ',','\n') AS LIEU_DE_TRAITEMENT,
                                      REPLACE(group_concat(NATURE_ET_QUANTITE_DE_PRODUIT_TRAITE), ',', '\n') AS NATURE_ET_QUANTITE_DE_PRODUIT_TRAITE,
                                      DATE_DE_PEREMPTION,
                                      STOCK_FINAL,
                                      SUM(QUANTITE_UTILISEE) AS QUANTITE_UTILISEE,
                                      QUANTITE_ACHETE,
                                      STOCK_INITIAL,
                                      PRODUITS_PHYTOSANITAIRES,   
                                      LIEU_DE_STOCKAGE
                                      FROM phytotablebidi WHERE DATE_FORMAT('%Y-%m', le) = %s GROUP BY PRODUITS_PHYTOSANITAIRES;""",
                        (selected_date,))
                    rows = cursor.fetchall()

                    ########################################################################
                    # Create a new document and add a table
                    self.doc = docx.Document()
                    section = self.doc.sections[0]
                    section.orientation = WD_ORIENTATION.LANDSCAPE
                    section.page_width = docx.shared.Cm(29.7)
                    section.page_height = docx.shared.Cm(21.0)
                    section.top_margin = docx.shared.Cm(0.5)
                    section.bottom_margin = docx.shared.Cm(1.5)
                    section.left_margin = docx.shared.Cm(1)
                    section.right_margin = docx.shared.Cm(1)
                    # add second heading

                    reshap = arabic_reshaper.reshape('حالة مخزون مواد الصحة النباتية  \n ')
                    bidi1 = bidi.algorithm.get_display(reshap)
                    heading1 = self.doc.add_heading(bidi1, level=1)
                    heading1.style.font.name = 'Times New Roman'
                    heading1.style.font.size = Pt(20)
                    heading1.style.font.bold = True
                    heading1.style.font.color.rgb = RGBColor(0, 0, 0)
                    heading1.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

                    paragraph = self.doc.add_paragraph()
                    paragraph.style.font.name = 'Times New Roman'
                    paragraph.style.font.size = Pt(14)
                    paragraph.style.font.bold = True
                    paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                    reshapl = arabic_reshaper.reshape('شهر:' + mois)
                    bidil = bidi.algorithm.get_display(reshapl)
                    left_run = paragraph.add_run(bidil)

                    reshap1 = arabic_reshaper.reshape('تعاونية الحبوب والبقول الجافة غليزان\t\t\t\t\t\t\t\t\t\t\t\t\n')
                    bidi2 = bidi.algorithm.get_display(reshap1)
                    left_run1 = paragraph.add_run(bidi2)

                    table = self.doc.add_table(rows=1, cols=9)
                    table.style = "Table Grid"  # set the table style
                    hdr_cells = table.rows[0].cells
                    t1 = arabic_reshaper.reshape("مكان المعالجة")
                    bidi0 = bidi.algorithm.get_display(t1)
                    t2 = arabic_reshaper.reshape("نوع وكمية المادة المعالجة")
                    bidi2 = bidi.algorithm.get_display(t2)
                    t3 = arabic_reshaper.reshape(" تاريخ نهاية الصلاحية")
                    bidi3 = bidi.algorithm.get_display(t3)
                    t4 = arabic_reshaper.reshape(" المخزون النهائي")
                    bidi4 = bidi.algorithm.get_display(t4)
                    t5 = arabic_reshaper.reshape("الكمية المستعملة")
                    bidi5 = bidi.algorithm.get_display(t5)
                    t6 = arabic_reshaper.reshape("الكمية المشترات")
                    bidi6 = bidi.algorithm.get_display(t6)
                    t7 = arabic_reshaper.reshape("المخزون الابتدائي")
                    bidi7 = bidi.algorithm.get_display(t7)
                    t8 = arabic_reshaper.reshape("المواد الكيماوية")
                    bidi8 = bidi.algorithm.get_display(t8)
                    t9 = arabic_reshaper.reshape("مكان التخزين")
                    bidi9 = bidi.algorithm.get_display(t9)
                    # Add text to the header cells
                    hdr_cells[0].text = bidi0
                    hdr_cells[1].text = bidi2
                    hdr_cells[2].text = bidi3
                    hdr_cells[3].text = bidi4
                    hdr_cells[4].text = bidi5
                    hdr_cells[5].text = bidi6
                    hdr_cells[6].text = bidi7
                    hdr_cells[7].text = bidi8
                    hdr_cells[8].text = bidi9
                    # Set the width of the header cells
                    hdr_cells[0].width = Inches(1.9)  # set the width of the first column to 1.5 inches
                    hdr_cells[1].width = Inches(2.6)
                    hdr_cells[2].width = Inches(1.8)
                    hdr_cells[3].width = Inches(1.4)
                    hdr_cells[4].width = Inches(1.4)
                    hdr_cells[5].width = Inches(1.4)
                    hdr_cells[6].width = Inches(1.5)
                    hdr_cells[7].width = Inches(1.2)
                    hdr_cells[8].width = Inches(2)

                    for cell in table.rows[0].cells:
                        cell.paragraphs[0].runs[0].font.name = 'Calibri'
                        cell.paragraphs[0].runs[0].font.size = Pt(10)
                        cell.paragraphs[0].runs[0].font.bold = True

                    # Add the data to the table
                    if rows:  # Check if the 'rows' list is not empty
                        table_cols = len(rows[0])
                        for row in range(len(rows)):
                            table.add_row()
                            for col in range(table_cols):
                                cell = table.cell(row + 1, col)
                                cell.text = str(rows[row][col])
                                # Set font properties
                                cell.paragraphs[0].runs[0].font.name = 'Calibri'
                                cell.paragraphs[0].runs[0].font.size = Pt(9)
                                cell.paragraphs[0].runs[0].font.bold = True
                                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                            for row in table.rows:
                                for cell in row.cells:
                                    for paragraph in cell.paragraphs:
                                        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        # Handle the case when there are no rows in the 'rows' list
                        table.add_row()
                        cell = table.cell(1, 0)
                        cell.text = "No data available."
                    for row in table.rows:
                        row.height = Inches(0.6)
                    # Save and open the document for printing
                    temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                    self.doc.save(temp_file)
                    if temp_file:
                        a = self.progress_bar()
                        a = self.progress_bar()
                        os.startfile(temp_file, "open")
                        # word = win32a.gencache.EnsureDispatch("Word.Application")

                        # doc = word.Documents.Open(temp_file)
                        # doc.PrintOut()
                        # doc.Close()
                        # word.Quit()
                    conn.close()


                elif self.produitphytofiltre.currentText() and selected_date == "2000-01":
                    produitphyto = self.produitphytofiltre.currentText()
                    selected_date = self.date_edit.date().toString("yyyy-MM")
                    locale.setlocale(locale.LC_ALL, 'fr_FR.utf8')
                    QtCore.QLocale.setDefault(QtCore.QLocale(QtCore.QLocale.Language.French))
                    selected_date1 = self.date_edit.date().toPyDate()
                    mois = selected_date1.DATE_FORMAT("%B").lower()

                    conn = mysql.connector.connect(host="localhost", user=user, password=password,database="datta")
                    cursor = conn.cursor()
                    cursor.execute(
                        """SELECT
                                      REPLACE(group_concat(LIEU_DE_TRAITEMENT), ',', '\n') AS LIEU_DE_TRAITEMENT,
                                      REPLACE(group_concat(NATURE_ET_QUANTITE_DE_PRODUIT_TRAITE), ',', '\n') AS NATURE_ET_QUANTITE_DE_PRODUIT_TRAITE,
                                      REPLACE(group_concat(DATE_DE_PEREMPTION), ',', ',') AS DATE_DE_PEREMPTION,
                                      STOCK_FINAL,
                                      SUM(QUANTITE_UTILISEE) AS QUANTITE_UTILISEE,
                                      QUANTITE_ACHETE,
                                      STOCK_INITIAL,
                                      PRODUITS_PHYTOSANITAIRES,
                                      LIEU_DE_STOCKAGE
                                      FROM phytotablebidi  WHERE PRODUITS_PHYTOSANITAIRES= %s GROUP BY PRODUITS_PHYTOSANITAIRES  ;""",
                        (produitphyto,))
                    rows = cursor.fetchall()

                    ########################################################################
                    # Create a new document and add a table
                    self.doc = docx.Document()
                    section = self.doc.sections[0]
                    section.orientation = WD_ORIENTATION.LANDSCAPE
                    section.page_width = docx.shared.Cm(29.7)
                    section.page_height = docx.shared.Cm(21.0)
                    section.top_margin = docx.shared.Cm(0.5)
                    section.bottom_margin = docx.shared.Cm(1.5)
                    section.left_margin = docx.shared.Cm(1)
                    section.right_margin = docx.shared.Cm(1)
                    # add second heading

                    reshap = arabic_reshaper.reshape('حالة مخزون مواد الصحة النباتية  \n ')
                    bidi1 = bidi.algorithm.get_display(reshap)
                    heading1 = self.doc.add_heading(bidi1, level=1)
                    heading1.style.font.name = 'Times New Roman'
                    heading1.style.font.size = Pt(20)
                    heading1.style.font.bold = True
                    heading1.style.font.color.rgb = RGBColor(0, 0, 0)
                    heading1.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

                    paragraph = self.doc.add_paragraph()
                    paragraph.style.font.name = 'Times New Roman'
                    paragraph.style.font.size = Pt(14)
                    paragraph.style.font.bold = True
                    paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                    reshapl = arabic_reshaper.reshape('المنتج:' + produitphyto)
                    bidil = bidi.algorithm.get_display(reshapl)
                    left_run = paragraph.add_run(bidil)

                    reshap1 = arabic_reshaper.reshape('تعاونية الحبوب والبقول الجافة غليزان\t\t\t\t\t\t\t\t\t\t\t\t\n')
                    bidi2 = bidi.algorithm.get_display(reshap1)
                    left_run1 = paragraph.add_run(bidi2)

                    table = self.doc.add_table(rows=1, cols=9)
                    table.style = "Table Grid"  # set the table style
                    hdr_cells = table.rows[0].cells
                    t1 = arabic_reshaper.reshape("مكان المعالجة")
                    bidi0 = bidi.algorithm.get_display(t1)
                    t2 = arabic_reshaper.reshape("نوع وكمية المادة المعالجة")
                    bidi2 = bidi.algorithm.get_display(t2)
                    t3 = arabic_reshaper.reshape("تاريخ نهاية الصلاحية")
                    bidi3 = bidi.algorithm.get_display(t3)
                    t4 = arabic_reshaper.reshape("المخزون النهائي")
                    bidi4 = bidi.algorithm.get_display(t4)
                    t5 = arabic_reshaper.reshape("الكمية المستعملة")
                    bidi5 = bidi.algorithm.get_display(t5)
                    t6 = arabic_reshaper.reshape("الكمية المشترات")
                    bidi6 = bidi.algorithm.get_display(t6)
                    t7 = arabic_reshaper.reshape("المخزون الابتدائي")
                    bidi7 = bidi.algorithm.get_display(t7)
                    t8 = arabic_reshaper.reshape(" المواد الكيماوية")
                    bidi8 = bidi.algorithm.get_display(t8)
                    t9 = arabic_reshaper.reshape("مكان التخزين")
                    bidi9 = bidi.algorithm.get_display(t9)
                    # Add text to the header cells
                    hdr_cells[0].text = bidi0
                    hdr_cells[1].text = bidi2
                    hdr_cells[2].text = bidi3
                    hdr_cells[3].text = bidi4
                    hdr_cells[4].text = bidi5
                    hdr_cells[5].text = bidi6
                    hdr_cells[6].text = bidi7
                    hdr_cells[7].text = bidi8
                    hdr_cells[8].text = bidi9
                    # Set the width of the header cells
                    hdr_cells[0].width = Inches(1.9)  # set the width of the first column to 1.5 inches
                    hdr_cells[1].width = Inches(2.6)
                    hdr_cells[2].width = Inches(1.8)
                    hdr_cells[3].width = Inches(1.4)
                    hdr_cells[4].width = Inches(1.4)
                    hdr_cells[5].width = Inches(1.4)
                    hdr_cells[6].width = Inches(1.5)
                    hdr_cells[7].width = Inches(1.2)
                    hdr_cells[8].width = Inches(2)
                    for cell in table.rows[0].cells:
                        cell.paragraphs[0].runs[0].font.name = 'Calibri'
                        cell.paragraphs[0].runs[0].font.size = Pt(11)
                        cell.paragraphs[0].runs[0].font.bold = True
                    # Add the data to the table
                    if rows:  # Check if the 'rows' list is not empty
                        table_cols = len(rows[0])
                        for row in range(len(rows)):
                            table.add_row()
                            for col in range(table_cols):
                                cell = table.cell(row + 1, col)
                                cell.text = str(rows[row][col])
                                # Set font properties
                                cell.paragraphs[0].runs[0].font.name = 'Calibri'
                                cell.paragraphs[0].runs[0].font.size = Pt(9)
                                cell.paragraphs[0].runs[0].font.bold = True
                                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                            for row in table.rows:
                                for cell in row.cells:
                                    for paragraph in cell.paragraphs:
                                        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        # Handle the case when there are no rows in the 'rows' list
                        table.add_row()
                        cell = table.cell(1, 0)
                        cell.text = "No data available."
                    for row in table.rows:
                        row.height = Inches(0.6)
                    # Save and open the document for printing
                    temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                    self.doc.save(temp_file)
                    if temp_file:
                        a = self.progress_bar()
                        a = self.progress_bar()
                        os.startfile(temp_file, "open")
                        # word = win32a.gencache.EnsureDispatch("Word.Application")
                        # a = self.progress_bar()
                        # doc = word.Documents.Open(temp_file)

                        # doc.PrintOut()
                        # doc.Close()
                        # word.Quit()
                    conn.close()


                elif selected_date and self.produitphytofiltre.currentText():
                    produitphyto = self.produitphytofiltre.currentText()
                    selected_month = self.date_edit.date().month()
                    # Create a list of Arabic month names
                    arabic_month_names = ["جانفي", "فيفري", "مارس", "أفريل", "ماي", "جوان", "جويلية", "أوت", "سبتمبر",
                                          "أكتوبر", "نوفمبر", "ديسمبر"]
                    # Create a custom formatted Arabic month string
                    formatted_month = arabic_month_names[selected_month - 1]
                    mois = formatted_month

                    conn = mysql.connector.connect(host="localhost", user=user, password=password,database="datta")
                    cursor = conn.cursor()
                    cursor.execute(
                        """SELECT
                                      REPLACE(group_concat(LIEU_DE_TRAITEMENT), ',', '\n') AS LIEU_DE_TRAITEMENT,
                                      REPLACE(group_concat(NATURE_ET_QUANTITE_DE_PRODUIT_TRAITE), ',', '\n') AS NATURE_ET_QUANTITE_DE_PRODUIT_TRAITE,
                                      REPLACE(group_concat(DATE_DE_PEREMPTION), ',', ',') AS DATE_DE_PEREMPTION,
                                      STOCK_FINAL,
                                      SUM(QUANTITE_UTILISEE) AS QUANTITE_UTILISEE,
                                      QUANTITE_ACHETE,
                                      STOCK_INITIAL,
                                      PRODUITS_PHYTOSANITAIRES,
                                      LIEU_DE_STOCKAGE
                                      FROM phytotablebidi  WHERE DATE_FORMAT('%Y-%m', le)=%s  and PRODUITS_PHYTOSANITAIRES=%s GROUP BY PRODUITS_PHYTOSANITAIRES  ;""",
                        (selected_date, produitphyto,))
                    rows = cursor.fetchall()

                    ########################################################################
                    # Create a new document and add a table
                    self.doc = docx.Document()
                    section = self.doc.sections[0]
                    section.orientation = WD_ORIENTATION.LANDSCAPE
                    section.page_width = docx.shared.Cm(29.7)
                    section.page_height = docx.shared.Cm(21.0)
                    section.top_margin = docx.shared.Cm(0.5)
                    section.bottom_margin = docx.shared.Cm(1.5)
                    section.left_margin = docx.shared.Cm(1)
                    section.right_margin = docx.shared.Cm(1)
                    # add second heading

                    reshap = arabic_reshaper.reshape('حالة مخزون مواد الصحة النباتية  \n ')
                    bidi1 = bidi.algorithm.get_display(reshap)
                    heading1 = self.doc.add_heading(bidi1, level=1)
                    heading1.style.font.name = 'Times New Roman'
                    heading1.style.font.size = Pt(20)
                    heading1.style.font.bold = True
                    heading1.style.font.color.rgb = RGBColor(0, 0, 0)
                    heading1.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

                    paragraph = self.doc.add_paragraph()
                    paragraph.style.font.name = 'Times New Roman'
                    paragraph.style.font.size = Pt(14)
                    paragraph.style.font.bold = True
                    paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                    reshapl = arabic_reshaper.reshape('شهر:' + mois)
                    bidil = bidi.algorithm.get_display(reshapl)
                    left_run = paragraph.add_run(bidil)

                    reshap1 = arabic_reshaper.reshape('تعاونية الحبوب والبقول الجافة غليزان\t\t\t\t\t\t\t\t\t\t\t\t\n')
                    bidi2 = bidi.algorithm.get_display(reshap1)
                    left_run1 = paragraph.add_run(bidi2)
                    table = self.doc.add_table(rows=1, cols=9)
                    table.style = "Table Grid"  # set the table style
                    hdr_cells = table.rows[0].cells
                    t1 = arabic_reshaper.reshape("مكان المعالجة")
                    bidi0 = bidi.algorithm.get_display(t1)
                    t2 = arabic_reshaper.reshape("نوع وكمية المادة المعالجة")
                    bidi2 = bidi.algorithm.get_display(t2)
                    t3 = arabic_reshaper.reshape(" تاريخ نهاية الصلاحية")
                    bidi3 = bidi.algorithm.get_display(t3)
                    t4 = arabic_reshaper.reshape(" المخزون النهائي")
                    bidi4 = bidi.algorithm.get_display(t4)
                    t5 = arabic_reshaper.reshape("الكمية المستعملة")
                    bidi5 = bidi.algorithm.get_display(t5)
                    t6 = arabic_reshaper.reshape("الكمية المشترات")
                    bidi6 = bidi.algorithm.get_display(t6)
                    t7 = arabic_reshaper.reshape("المخزون الابتدائي")
                    bidi7 = bidi.algorithm.get_display(t7)
                    t8 = arabic_reshaper.reshape("المواد الكيماوية")
                    bidi8 = bidi.algorithm.get_display(t8)
                    t9 = arabic_reshaper.reshape("مكان التخزين")
                    bidi9 = bidi.algorithm.get_display(t9)
                    # Add text to the header cells
                    hdr_cells[0].text = bidi0
                    hdr_cells[1].text = bidi2
                    hdr_cells[2].text = bidi3
                    hdr_cells[3].text = bidi4
                    hdr_cells[4].text = bidi5
                    hdr_cells[5].text = bidi6
                    hdr_cells[6].text = bidi7
                    hdr_cells[7].text = bidi8
                    hdr_cells[8].text = bidi9
                    # Set the width of the header cells
                    hdr_cells[0].width = Inches(1.9)  # set the width of the first column to 1.5 inches
                    hdr_cells[1].width = Inches(2.6)
                    hdr_cells[2].width = Inches(1.8)
                    hdr_cells[3].width = Inches(1.4)
                    hdr_cells[4].width = Inches(1.4)
                    hdr_cells[5].width = Inches(1.4)
                    hdr_cells[6].width = Inches(1.5)
                    hdr_cells[7].width = Inches(1.2)
                    hdr_cells[8].width = Inches(2)

                    for cell in table.rows[0].cells:
                        cell.paragraphs[0].runs[0].font.name = 'Calibri'
                        cell.paragraphs[0].runs[0].font.size = Pt(11)
                        cell.paragraphs[0].runs[0].font.bold = True
                    # Add the data to the table
                    if rows:  # Check if the 'rows' list is not empty
                        table_cols = len(rows[0])
                        for row in range(len(rows)):
                            table.add_row()
                            for col in range(table_cols):
                                cell = table.cell(row + 1, col)
                                cell.text = str(rows[row][col])
                                # Set font properties
                                cell.paragraphs[0].runs[0].font.name = 'Calibri'
                                cell.paragraphs[0].runs[0].font.size = Pt(9)
                                cell.paragraphs[0].runs[0].font.bold = True
                                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                            for row in table.rows:
                                for cell in row.cells:
                                    for paragraph in cell.paragraphs:
                                        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        # Handle the case when there are no rows in the 'rows' list
                        table.add_row()
                        cell = table.cell(1, 0)
                        cell.text = "No data available."
                    for row in table.rows:
                        row.height = Inches(0.6)
                    # Save and open the document for printing
                    temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                    self.doc.save(temp_file)
                    if temp_file:
                        a = self.progress_bar()
                        a = self.progress_bar()
                        os.startfile(temp_file, "open")
                        # word = win32a.gencache.EnsureDispatch("Word.Application")
                        # a = self.progress_bar()
                        # doc = word.Documents.Open(temp_file)
                        # doc.PrintOut()
                        # doc.Close()
                        # word.Quit()
                    conn.close()
            except Exception as e:
                print(e)

        def progress_bar(self):
            self.widgetprogress = QtWidgets.QDialog()
            self.widgetprogress.setStyleSheet(""" QWidget
                            {
                                color: #000000;
                                background-color: #ffffff;
                                border-width: 1px;
                                border-color: #1e1e1e;
                                border-style: solid;
                                border-radius: 6;
                                padding: 0px;
                                font-size: 18px;
                                padding-left: 1px;
                                padding-right: 1px
                            }
                            QWidget:item:hover
                            {
                                background-color: #3daee9;
                                color: #eff0f1;
                            }
                            QWidget:item:selected
                            {
                                background-color: #3daee9;
                            }
                            QWidget:disabled
                            {
                                color: #454545;
                                background-color: #31363b;
                            }
                            QPushButton
                            {
                                color: #000000;
                                background-color:#ade3e7;
                                border-width: 1px;
                                border-color: #1e1e1e;
                                border-style: solid;
                                border-radius: 6;
                                padding: 3px;
                                font-size: 12px;
                                padding-left: 5px;
                                padding-right: 5px;
                                min-width: 40px
                            }
                            QPushButton:disabled
                            {
                                background-color: #31363b;
                                border-width: 1px;
                                border-color: #454545;
                                border-style: solid;
                                padding-top: 5px;
                                padding-bottom: 5px;
                                padding-left: 10px;
                                padding-right: 10px;
                                border-radius: 2px;
                                color: #454545;
                            }

                            QPushButton:pressed
                            {
                                background-color: #3daee9;
                                padding-top: -15px;
                                padding-bottom: -17px;
                            }
                            QPushButton:hover
                            {
                                border: 1px solid #ff8c00;
                                color: #000000;
                            }
                             QLabel
                            {
                                font-size: 18px;
                                border: 0px solid orange;
                            }

                        """)
            self.widgetprogress.setWindowTitle("جاري طباعة الملف يرجى الانتظار ")
            self.widgetprogress.setGeometry(550, 450, 250, 20)
            self.progressBar = QtWidgets.QProgressBar(self.widgetprogress)
            self.progressBar.setGeometry(10, 10, 200, 10)
            self.progressBar.setMinimum(0)
            self.progressBar.setMaximum(100)
            self.progressBar.setStyleSheet("""QProgressBar
    {
    border: solid grey;
    border-radius: 15px;
    color: black;
    }
    QProgressBar::chunk 
    {
    background-color: #05B8CC;
    border-radius :15px;
    }      """)
            self.progressBar.setAlignment(Qt.AlignmentFlag.AlignCenter)
            self.vbox = QVBoxLayout(self.widgetprogress)
            self.vbox.addWidget(self.progressBar)
            self.timer = QtCore.QTimer()
            self.timer.timeout.connect(self.update_progress)
            self.timer.start(100)  # Update progress every
            self.widgetprogress.show()

        def update_progress(self):
            # Simulate file download progress
            current_value = self.progressBar.value()
            if current_value < 100:
                new_value = current_value + 10
                self.progressBar.setValue(new_value)
                if current_value == 99:
                    self.timer.stop()
                    self.progressBar.close()
                    self.widgetprogress.close()

        def impot_filter(self):
            try:
                datefl = self.datefilter.date().toString("yyyy-MM")
                produitphytofl = self.produitphytofl.currentText()
                conn = mysql.connector.connect(host="localhost", user=user, password=password,database="datta")
                cursor = conn.cursor()
                cursordate = conn.cursor()
                if datefl == "2000-01":

                    cursor.execute(
                        "SELECT LIEU_DE_TRAITEMENT, NATURE_ET_QUANTITE_DE_PRODUIT_TRAITE, DATE_DE_PEREMPTION, STOCK_FINAL, QUANTITE_UTILISEE, QUANTITE_ACHETE, STOCK_INITIAL, PRODUITS_PHYTOSANITAIRES, LIEU_DE_STOCKAGE, le FROM phytotable WHERE  PRODUITS_PHYTOSANITAIRES = %s ",
                        (produitphytofl,))
                    result = cursor.fetchall()
                    self.textEdit.setRowCount(0)
                    for row, row_datta in enumerate(result):
                        self.textEdit.insertRow(row)
                        for colum, datta in enumerate(row_datta):
                            self.textEdit.setItem(row, colum, QTableWidgetItem(str(datta)))
                elif produitphytofl == '--------------':
                    cursordate.execute(
                        "SELECT LIEU_DE_TRAITEMENT, NATURE_ET_QUANTITE_DE_PRODUIT_TRAITE, DATE_DE_PEREMPTION, STOCK_FINAL, QUANTITE_UTILISEE, QUANTITE_ACHETE, STOCK_INITIAL, PRODUITS_PHYTOSANITAIRES, LIEU_DE_STOCKAGE, le FROM phytotable WHERE DATE_FORMAT(le, '%Y-%m') = %s ",
                        (datefl,))
                    resultdate = cursordate.fetchall()
                    self.textEdit.setRowCount(0)
                    for rowdate, row_dattadate in enumerate(resultdate):
                        self.textEdit.insertRow(rowdate)
                        for columdate, dattadate in enumerate(row_dattadate):
                            self.textEdit.setItem(rowdate, columdate, QTableWidgetItem(str(dattadate)))
                else:
                    cursor.execute(
                        "SELECT LIEU_DE_TRAITEMENT, NATURE_ET_QUANTITE_DE_PRODUIT_TRAITE, DATE_DE_PEREMPTION, STOCK_FINAL, QUANTITE_UTILISEE, QUANTITE_ACHETE, STOCK_INITIAL, PRODUITS_PHYTOSANITAIRES, LIEU_DE_STOCKAGE, le FROM phytotable WHERE  PRODUITS_PHYTOSANITAIRES = %s AND DATE_FORMAT('%Y-%m',le ) = %s ",
                        (produitphytofl, datefl,))
                    result = cursor.fetchall()
                    self.textEdit.setRowCount(0)
                    for row, row_datta in enumerate(result):
                        self.textEdit.insertRow(row)
                        for colum, datta in enumerate(row_datta):
                            self.textEdit.setItem(row, colum, QTableWidgetItem(str(datta)))




            except sqlite3.Error as e:
                print(e)

        def add_datta_save(self):
            try:
                datte = self.dateedite.text()
                lieudestock = self.lieucombo.currentText()
                lieudestockt = arabic_reshaper.reshape(lieudestock)
                lieudestockbidi = bidi.algorithm.get_display(lieudestockt)
                produitpyhto = self.produitphyto.currentText()
                stockinitial = self.STOCKINITIAL.value()
                quantiteutilise = self.quantitéutilsé.value()
                stockfinall = (stockinitial - quantiteutilise)
                datedeprom = self.dateprempo.text()
                lieudetraitement = self.lieudetraitementcombo.currentText()
                lieudetraitementt = arabic_reshaper.reshape(lieudetraitement)
                lieudetraitementbidi = bidi.algorithm.get_display(lieudetraitementt)
                lieudetraitementbidi.replace(" ", "")
                natureetquantite = self.naturelieu.toPlainText()
                natureetquantitet = arabic_reshaper.reshape(natureetquantite)
                natureetquantitebidi = bidi.algorithm.get_display(natureetquantitet)
                quantiteachete = self.STOCKINITIAL.value()
                if produitpyhto:
                    dattabase = mysql.connector.connect(host="localhost", user=user, password=password,
                                                        database="datta")
                    curs = dattabase.cursor()
                    cursachat = dattabase.cursor()
                    cursbidi = dattabase.cursor()

                    if produitpyhto == "PHOSTOXIN(PH3)":
                        curs.execute(
                            "INSERT INTO phytotable (PRODUITS_PHYTOSANITAIRES,le,LIEU_DE_STOCKAGE,QUANTITE_ACHETE,QUANTITE_UTILISEE,STOCK_FINAL,DATE_DE_PEREMPTION,NATURE_ET_QUANTITE_DE_PRODUIT_TRAITE,LIEU_DE_TRAITEMENT,STOCK_INITIAL_PHOSTOXIN_PH3,QUANTITE_ACHETE_PHOSTOXIN_PH3,QUANTITE_UTILISEE_PHOSTOXIN_PH3) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                            (produitpyhto, datte, lieudestock, quantiteachete, quantiteutilise, stockfinall,
                             datedeprom,
                             natureetquantite, lieudetraitement, stockinitial, quantiteachete, quantiteutilise))
                        curs.execute(
                            "UPDATE phytotable SET STOCK_FINAL = (SELECT SUM(STOCK_INITIAL_PHOSTOXIN_PH3 - QUANTITE_UTILISEE_PHOSTOXIN_PH3)FROM phytotable WHERE PRODUITS_PHYTOSANITAIRES = 'PHOSTOXIN(PH3)')WHERE PRODUITS_PHYTOSANITAIRES = 'PHOSTOXIN(PH3)'")
                        curs.execute(
                            "UPDATE phytotable SET STOCK_INITIAL = (SELECT SUM(QUANTITE_ACHETE_PHOSTOXIN_PH3 )FROM phytotable WHERE PRODUITS_PHYTOSANITAIRES = 'PHOSTOXIN(PH3)')")

                        # curspif.execute("UPDATE phytotable SET STOCK_INITIAL_PHOSTOXIN_PH3 = IFNULL( STOCK_INITIAL_PHOSTOXIN_PH3 , 0) WHERE STOCK_FINAL = 0 AND PRODUITS_PHYTOSANITAIRES = 'PHOSTOXIN(PH3)'")
                        ##################################bidi############
                        cursbidi.execute(
                            "INSERT INTO phytotablebidi (PRODUITS_PHYTOSANITAIRES,le,LIEU_DE_STOCKAGE,QUANTITE_ACHETE,QUANTITE_UTILISEE,STOCK_FINAL,DATE_DE_PEREMPTION,NATURE_ET_QUANTITE_DE_PRODUIT_TRAITE,LIEU_DE_TRAITEMENT,STOCK_INITIAL_PHOSTOXIN_PH3 ,QUANTITE_ACHETE_PHOSTOXIN_PH3,QUANTITE_UTILISEE_PHOSTOXIN_PH3) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                            (produitpyhto, datte, lieudestockbidi, quantiteachete, quantiteutilise, stockfinall,
                             datedeprom,
                             natureetquantitebidi, lieudetraitementbidi, stockinitial, quantiteachete,
                             quantiteutilise))
                        cursbidi.execute(
                            "UPDATE phytotablebidi SET STOCK_FINAL = (SELECT SUM(STOCK_INITIAL_PHOSTOXIN_PH3 - QUANTITE_UTILISEE_PHOSTOXIN_PH3)FROM phytotable WHERE PRODUITS_PHYTOSANITAIRES = 'PHOSTOXIN(PH3)')WHERE PRODUITS_PHYTOSANITAIRES = 'PHOSTOXIN(PH3)'")
                        cursbidi.execute(
                            "UPDATE phytotablebidi SET STOCK_INITIAL = (SELECT SUM(QUANTITE_ACHETE_PHOSTOXIN_PH3 )FROM phytotable WHERE PRODUITS_PHYTOSANITAIRES = 'PHOSTOXIN(PH3)')")
                    if produitpyhto == "ACTELLIC":
                        curs.execute(
                            "INSERT INTO phytotable (PRODUITS_PHYTOSANITAIRES,le,LIEU_DE_STOCKAGE,QUANTITE_ACHETE,QUANTITE_UTILISEE,STOCK_FINAL,DATE_DE_PEREMPTION,NATURE_ET_QUANTITE_DE_PRODUIT_TRAITE,LIEU_DE_TRAITEMENT,STOCK_INITIAL_ACTELLIC,QUANTITE_ACHETE_ACTELLIC,QUANTITE_UTILISEE_ACTELLIC) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                            (produitpyhto, datte, lieudestock, quantiteachete, quantiteutilise, stockfinall,
                             datedeprom,
                             natureetquantite, lieudetraitement, stockinitial, quantiteachete, quantiteutilise))
                        curs.execute(
                            "UPDATE phytotable SET STOCK_FINAL = (SELECT SUM(STOCK_INITIAL_ACTELLIC - QUANTITE_UTILISEE_ACTELLIC)FROM phytotable WHERE PRODUITS_PHYTOSANITAIRES = 'ACTELLIC')WHERE PRODUITS_PHYTOSANITAIRES = 'ACTELLIC'")
                        cursachat.execute(
                            "UPDATE phytotable SET STOCK_INITIAL = (SELECT SUM(QUANTITE_ACHETE_ACTELLIC )FROM phytotable WHERE PRODUITS_PHYTOSANITAIRES = 'ACTELLIC')WHERE PRODUITS_PHYTOSANITAIRES = 'ACTELLIC'")
                        #############################bidiactilic
                        cursbidi.execute(
                            "INSERT INTO phytotablebidi (PRODUITS_PHYTOSANITAIRES,le,LIEU_DE_STOCKAGE,QUANTITE_ACHETE,QUANTITE_UTILISEE,STOCK_FINAL,DATE_DE_PEREMPTION,NATURE_ET_QUANTITE_DE_PRODUIT_TRAITE,LIEU_DE_TRAITEMENT,STOCK_INITIAL_ACTELLIC,QUANTITE_ACHETE_ACTELLIC,QUANTITE_UTILISEE_ACTELLIC) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s) ",
                            (produitpyhto, datte, lieudestockbidi, quantiteachete, quantiteutilise, stockfinall,
                             datedeprom,
                             natureetquantitebidi, lieudetraitementbidi, stockinitial, quantiteachete,
                             quantiteutilise))
                        cursbidi.execute(
                            "UPDATE phytotablebidi SET STOCK_FINAL = (SELECT SUM(STOCK_INITIAL_ACTELLIC - QUANTITE_UTILISEE_ACTELLIC)FROM phytotable WHERE PRODUITS_PHYTOSANITAIRES = 'ACTELLIC')WHERE PRODUITS_PHYTOSANITAIRES = 'ACTELLIC'")
                        cursbidi.execute(
                            "UPDATE phytotablebidi SET STOCK_INITIAL = (SELECT SUM(QUANTITE_ACHETE_ACTELLIC )FROM phytotable WHERE PRODUITS_PHYTOSANITAIRES = 'ACTELLIC')WHERE PRODUITS_PHYTOSANITAIRES = 'ACTELLIC'")
                    if produitpyhto == "CIRATHRINE":
                        curs.execute(
                            "INSERT INTO phytotable (PRODUITS_PHYTOSANITAIRES,le,LIEU_DE_STOCKAGE,STOCK_INITIAL,QUANTITE_ACHETE,QUANTITE_UTILISEE,STOCK_FINAL,DATE_DE_PEREMPTION,NATURE_ET_QUANTITE_DE_PRODUIT_TRAITE,LIEU_DE_TRAITEMENT,STOCK_INITIAL_CIRATHRINE,QUANTITE_ACHETE_CIRATHRINE,QUANTITE_UTILISEE_CIRATHRINE) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                            (produitpyhto, datte, lieudestock, stockinitial, quantiteachete, quantiteutilise,
                             stockfinall,
                             datedeprom, natureetquantite, lieudetraitement, stockinitial, quantiteachete,
                             quantiteutilise))
                        curs.execute(
                            "UPDATE phytotable SET STOCK_FINAL = (SELECT SUM(STOCK_INITIAL_CIRATHRINE - QUANTITE_UTILISEE_CIRATHRINE)FROM phytotable WHERE PRODUITS_PHYTOSANITAIRES = 'CIRATHRINE')WHERE PRODUITS_PHYTOSANITAIRES = 'CIRATHRINE'")
                        curs.execute(
                            "UPDATE phytotable SET STOCK_INITIAL = (SELECT SUM(QUANTITE_ACHETE_CIRATHRINE)FROM phytotable WHERE PRODUITS_PHYTOSANITAIRES = 'CIRATHRINE')WHERE PRODUITS_PHYTOSANITAIRES = 'CIRATHRINE'")
                        ###############################bidicirathrine
                        cursbidi.execute(
                            "INSERT INTO phytotablebidi (PRODUITS_PHYTOSANITAIRES,le,LIEU_DE_STOCKAGE,STOCK_INITIAL,QUANTITE_ACHETE,QUANTITE_UTILISEE,STOCK_FINAL,DATE_DE_PEREMPTION,NATURE_ET_QUANTITE_DE_PRODUIT_TRAITE,LIEU_DE_TRAITEMENT,STOCK_INITIAL_CIRATHRINE,QUANTITE_ACHETE_CIRATHRINE,QUANTITE_UTILISEE_CIRATHRINE) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                            (produitpyhto, datte, lieudestockbidi, stockinitial, quantiteachete, quantiteutilise,
                             stockfinall,
                             datedeprom, natureetquantitebidi, lieudetraitementbidi, stockinitial, quantiteachete,
                             quantiteutilise))
                        cursbidi.execute(
                            "UPDATE phytotablebidi SET STOCK_FINAL = (SELECT SUM(STOCK_INITIAL_CIRATHRINE - QUANTITE_UTILISEE_CIRATHRINE)FROM phytotable WHERE PRODUITS_PHYTOSANITAIRES = 'CIRATHRINE')WHERE PRODUITS_PHYTOSANITAIRES = 'CIRATHRINE'")
                        cursbidi.execute(
                            "UPDATE phytotablebidi SET STOCK_INITIAL = (SELECT SUM(QUANTITE_ACHETE_CIRATHRINE)FROM phytotable WHERE PRODUITS_PHYTOSANITAIRES = 'CIRATHRINE')WHERE PRODUITS_PHYTOSANITAIRES = 'CIRATHRINE'")
                    if produitpyhto == "DEKATRINE":
                        curs.execute(
                            "INSERT INTO phytotable (PRODUITS_PHYTOSANITAIRES,le,LIEU_DE_STOCKAGE,STOCK_INITIAL,QUANTITE_ACHETE,QUANTITE_UTILISEE,STOCK_FINAL,DATE_DE_PEREMPTION,NATURE_ET_QUANTITE_DE_PRODUIT_TRAITE,LIEU_DE_TRAITEMENT,STOCK_INITIAL_DEKATRINE,QUANTITE_ACHETE_DEKATRINE,QUANTITE_UTILISEE_DEKATRINE) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                            (produitpyhto, datte, lieudestock, stockinitial, quantiteachete, quantiteutilise,
                             stockfinall,
                             datedeprom, natureetquantite, lieudetraitement, stockinitial, quantiteachete,
                             quantiteutilise))
                        curs.execute(
                            "UPDATE phytotable SET STOCK_FINAL = (SELECT SUM(STOCK_INITIAL_DEKATRINE - QUANTITE_UTILISEE_DEKATRINE)FROM phytotable WHERE PRODUITS_PHYTOSANITAIRES = 'DEKATRINE')WHERE PRODUITS_PHYTOSANITAIRES = 'DEKATRINE'")
                        curs.execute(
                            "UPDATE phytotable SET STOCK_INITIAL = (SELECT SUM(QUANTITE_ACHETE_DEKATRINE)FROM phytotable WHERE PRODUITS_PHYTOSANITAIRES = 'DEKATRINE')WHERE PRODUITS_PHYTOSANITAIRES = 'DEKATRINE'")
                        ###############################bididekatrine
                        cursbidi.execute(
                            "INSERT INTO phytotablebidi (PRODUITS_PHYTOSANITAIRES,le,LIEU_DE_STOCKAGE,STOCK_INITIAL,QUANTITE_ACHETE,QUANTITE_UTILISEE,STOCK_FINAL,DATE_DE_PEREMPTION,NATURE_ET_QUANTITE_DE_PRODUIT_TRAITE,LIEU_DE_TRAITEMENT,STOCK_INITIAL_DEKATRINE,QUANTITE_ACHETE_DEKATRINE,QUANTITE_UTILISEE_DEKATRINE) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                            (produitpyhto, datte, lieudestockbidi, stockinitial, quantiteachete, quantiteutilise,
                             stockfinall,
                             datedeprom, natureetquantitebidi, lieudetraitementbidi, stockinitial, quantiteachete,
                             quantiteutilise))
                        cursbidi.execute(
                            "UPDATE phytotablebidi SET STOCK_FINAL = (SELECT SUM(STOCK_INITIAL_DEKATRINE - QUANTITE_UTILISEE_DEKATRINE)FROM phytotable WHERE PRODUITS_PHYTOSANITAIRES = 'DEKATRINE')WHERE PRODUITS_PHYTOSANITAIRES = 'DEKATRINE'")
                        cursbidi.execute(
                            "UPDATE phytotablebidi SET STOCK_INITIAL = (SELECT SUM(QUANTITE_ACHETE_DEKATRINE)FROM phytotable WHERE PRODUITS_PHYTOSANITAIRES = 'DEKATRINE')WHERE PRODUITS_PHYTOSANITAIRES = 'DEKATRINE'")
                    if produitpyhto == "RATICIDE":
                        curs.execute(
                            "INSERT INTO phytotable (PRODUITS_PHYTOSANITAIRES,le,LIEU_DE_STOCKAGE,STOCK_INITIAL,QUANTITE_ACHETE,QUANTITE_UTILISEE,STOCK_FINAL,DATE_DE_PEREMPTION,NATURE_ET_QUANTITE_DE_PRODUIT_TRAITE,LIEU_DE_TRAITEMENT,STOCK_INITIAL_RATICIDE,QUANTITE_ACHETE_RATICIDE,QUANTITE_UTILISEE_RATICIDE) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                            (produitpyhto, datte, lieudestock, stockinitial, quantiteachete, quantiteutilise,
                             stockfinall,
                             datedeprom, natureetquantite, lieudetraitement, stockinitial, quantiteachete,
                             quantiteutilise))
                        curs.execute(
                            "UPDATE phytotable SET STOCK_FINAL = (SELECT SUM(STOCK_INITIAL_RATICIDE - QUANTITE_UTILISEE_RATICIDE)FROM phytotable WHERE PRODUITS_PHYTOSANITAIRES = 'RATICIDE')WHERE PRODUITS_PHYTOSANITAIRES = 'RATICIDE'")
                        curs.execute(
                            "UPDATE phytotable SET STOCK_INITIAL = (SELECT SUM(QUANTITE_ACHETE_RATICIDE)FROM phytotable WHERE PRODUITS_PHYTOSANITAIRES ='RATICIDE')WHERE PRODUITS_PHYTOSANITAIRES = 'RATICIDE'")
                        ############################bidiraticid
                        cursbidi.execute(
                            "INSERT INTO phytotablebidi (PRODUITS_PHYTOSANITAIRES,le,LIEU_DE_STOCKAGE,STOCK_INITIAL,QUANTITE_ACHETE,QUANTITE_UTILISEE,STOCK_FINAL,DATE_DE_PEREMPTION,NATURE_ET_QUANTITE_DE_PRODUIT_TRAITE,LIEU_DE_TRAITEMENT,STOCK_INITIAL_RATICIDE,QUANTITE_ACHETE_RATICIDE,QUANTITE_UTILISEE_RATICIDE) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                            (produitpyhto, datte, lieudestockbidi, stockinitial, quantiteachete, quantiteutilise,
                             stockfinall,
                             datedeprom, natureetquantitebidi, lieudetraitementbidi, stockinitial, quantiteachete,
                             quantiteutilise))
                        cursbidi.execute(
                            "UPDATE phytotablebidi SET STOCK_FINAL = (SELECT SUM(STOCK_INITIAL_RATICIDE - QUANTITE_UTILISEE_RATICIDE)FROM phytotable WHERE PRODUITS_PHYTOSANITAIRES = 'RATICIDE')WHERE PRODUITS_PHYTOSANITAIRES = 'RATICIDE'")
                        cursbidi.execute(
                            "UPDATE phytotablebidi SET STOCK_INITIAL = (SELECT SUM(QUANTITE_ACHETE_RATICIDE)FROM phytotable WHERE PRODUITS_PHYTOSANITAIRES = 'RATICIDE')WHERE PRODUITS_PHYTOSANITAIRES = 'RATICIDE'")
                    if produitpyhto == "TEXTO":
                        curs.execute(
                            "INSERT INTO phytotable (PRODUITS_PHYTOSANITAIRES,le,LIEU_DE_STOCKAGE,STOCK_INITIAL,QUANTITE_ACHETE,QUANTITE_UTILISEE,STOCK_FINAL,DATE_DE_PEREMPTION,NATURE_ET_QUANTITE_DE_PRODUIT_TRAITE,LIEU_DE_TRAITEMENT,STOCK_INITIAL_TEXTO,QUANTITE_ACHETE_TEXTO,QUANTITE_UTILISEE_TEXTO) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                            (produitpyhto, datte, lieudestock, stockinitial, quantiteachete, quantiteutilise,
                             stockfinall,
                             datedeprom, natureetquantite, lieudetraitement, stockinitial, quantiteachete,
                             quantiteutilise))
                        curs.execute(
                            "UPDATE phytotable SET STOCK_FINAL = (SELECT SUM(STOCK_INITIAL_TEXTO - QUANTITE_UTILISEE_TEXTO)FROM phytotable WHERE PRODUITS_PHYTOSANITAIRES ='TEXTO')WHERE PRODUITS_PHYTOSANITAIRES = 'TEXTO'")
                        curs.execute(
                            "UPDATE phytotable SET STOCK_INITIAL = (SELECT SUM(QUANTITE_ACHETE_TEXTO)FROM phytotable WHERE PRODUITS_PHYTOSANITAIRES ='TEXTO')WHERE PRODUITS_PHYTOSANITAIRES = 'TEXTO'")

                        ###########################biditexto
                        cursbidi.execute(
                            "INSERT INTO phytotablebidi (PRODUITS_PHYTOSANITAIRES,le,LIEU_DE_STOCKAGE,STOCK_INITIAL,QUANTITE_ACHETE,QUANTITE_UTILISEE,STOCK_FINAL,DATE_DE_PEREMPTION,NATURE_ET_QUANTITE_DE_PRODUIT_TRAITE,LIEU_DE_TRAITEMENT,STOCK_INITIAL_TEXTO,QUANTITE_ACHETE_TEXTO,QUANTITE_UTILISEE_TEXTO) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                            (produitpyhto, datte, lieudestockbidi, stockinitial, quantiteachete, quantiteutilise,
                             stockfinall,
                             datedeprom, natureetquantitebidi, lieudetraitementbidi, stockinitial, quantiteachete,
                             quantiteutilise))
                        cursbidi.execute(
                            "UPDATE phytotablebidi SET STOCK_FINAL = (SELECT SUM(STOCK_INITIAL_TEXTO - QUANTITE_UTILISEE_TEXTO)FROM phytotable WHERE PRODUITS_PHYTOSANITAIRES ='TEXTO')WHERE PRODUITS_PHYTOSANITAIRES = 'TEXTO'")
                        cursbidi.execute(
                            "UPDATE phytotablebidi SET STOCK_INITIAL = (SELECT SUM(QUANTITE_ACHETE_TEXTO)FROM phytotable WHERE PRODUITS_PHYTOSANITAIRES ='TEXTO')WHERE PRODUITS_PHYTOSANITAIRES = 'TEXTO'")

                    dattabase.commit()
                    dattabase.close()
                    self.impot_all()
                    self.importph()
                    self.importtexto()
                    self.importraticide()
                    self.importactilic()
                    self.importdekatrine()
                    self.importcirathrine()
                    msgbox = QtWidgets.QMessageBox()
                    msgbox.setWindowTitle('تاكيد العملية')
                    msgbox.setText('تمت الاضافة بنجاح')
                    msgbox.exec()
                else:
                    msgbox = QtWidgets.QMessageBox()
                    msgbox.setWindowTitle("خطأ")
                    msgbox.setText("خطأ: اختر المنتج ")
                    msgbox.exec()

            except mysql.connector.Error as e:
                print(e)


        def impot_all(self):
            try:
                dattabase = mysql.connector.connect(host="localhost", user=user, password=password,database="datta")
                cur = dattabase.cursor()
                cur.execute("SELECT LIEU_DE_TRAITEMENT,NATURE_ET_QUANTITE_DE_PRODUIT_TRAITE ,DATE_DE_PEREMPTION ,STOCK_FINAL ,QUANTITE_UTILISEE ,QUANTITE_ACHETE ,STOCK_INITIAL , PRODUITS_PHYTOSANITAIRES ,LIEU_DE_STOCKAGE ,le,id  FROM phytotable  WHERE LIEU_DE_TRAITEMENT !='' AND le !='' ")
                result = cur.fetchall()
                self.textEdit.setRowCount(0)
                for row, row_datta in enumerate(result):

                    self.textEdit.insertRow(row)
                    for colum, datta in enumerate(row_datta):
                        self.textEdit.setItem(row, colum, QTableWidgetItem(str(datta)))
                dattabase.commit()
            except mysql.connector.Error as e:
                print(e)

        def importph(self):
            """Imports the stock of PHOSTOXIN(PH3) from the database."""
            try:
                with mysql.connector.connect(host="localhost", user=user, password=password,database="datta") as conn:
                    cur = conn.cursor()

                    # Get the total stock of PHOSTOXIN(PH3).
                    cur.execute(
                        "SELECT SUM(STOCK_INITIAL_PHOSTOXIN_PH3)-SUM(QUANTITE_UTILISEE_PHOSTOXIN_PH3) as STOCK_FINAL_PHOSTOXIN_PH3 FROM phytotable WHERE PRODUITS_PHYTOSANITAIRES='PHOSTOXIN(PH3)'"
                    )
                    stock_final_ph = cur.fetchone()[0]

                    # Set the value of the `phfinal` field.
                    self.phfinal.setValue(stock_final_ph)

                    if stock_final_ph == 0:
                        cur.execute(
                            "UPDATE phytotable SET QUANTITE_ACHETE_PHOSTOXIN_PH3=0 WHERE PRODUITS_PHYTOSANITAIRES ='PHOSTOXIN(PH3)'")
                    # Commit the changes.
                    conn.commit()
            except Exception as e:
                print(e)

        def importactilic(self):
            """Imports the stock of ACTELLIC from the database."""
            try:
                with mysql.connector.connect(host="localhost", user=user, password=password,database="datta") as conn:
                    cur = conn.cursor()

                    # Get the total stock of ACTELLIC.
                    cur.execute(
                        "SELECT SUM(STOCK_INITIAL_ACTELLIC)-SUM(QUANTITE_UTILISEE_ACTELLIC) as STOCK_FINAL_ACTELLIC FROM phytotable WHERE PRODUITS_PHYTOSANITAIRES='ACTELLIC'"
                    )
                    stock_final_actellic = cur.fetchone()[0]

                    # Set the value of the `ACTELLICfinal` field.
                    self.ACTELLICfinal.setValue(stock_final_actellic)

                    if stock_final_actellic == 0:
                        cur.execute(
                            "UPDATE phytotable SET   QUANTITE_ACHETE_ACTELLIC = 0 WHERE PRODUITS_PHYTOSANITAIRES ='ACTELLIC'")

                    # Commit the changes.
                    conn.commit()
            except Exception as e:
                print(e)

        def importcirathrine(self):
            """Imports the stock of CIRATHRINE from the database."""
            try:
                with mysql.connector.connect(host="localhost", user=user, password=password,database="datta") as conn:
                    cur = conn.cursor()

                    # Get the total stock of CIRATHRINE.
                    cur.execute(
                        "SELECT SUM(STOCK_INITIAL_CIRATHRINE)-SUM(QUANTITE_UTILISEE_CIRATHRINE) as STOCK_FINAL_CIRATHRINE FROM phytotable WHERE PRODUITS_PHYTOSANITAIRES='CIRATHRINE'"
                    )
                    stock_final_cirathrine = cur.fetchone()[0]

                    # Set the value of the `CIRATHRINEfinal` field.
                    self.CIRATHRINEfinal.setValue(stock_final_cirathrine)

                    if stock_final_cirathrine == 0:
                        cur.execute(
                            "UPDATE phytotable SET  QUANTITE_ACHETE_CIRATHRINE = 0 WHERE PRODUITS_PHYTOSANITAIRES ='CIRATHRINE'")

                    # Commit the changes.
                    conn.commit()
            except Exception as e:
                print(e)

        def importdekatrine(self):
            """Imports the stock of DEKATRINE from the database."""
            try:
                with mysql.connector.connect(host="localhost", user=user, password=password,database="datta") as conn:
                    cur = conn.cursor()

                    # Get the total stock of DEKATRINE.
                    cur.execute(
                        "SELECT SUM(STOCK_INITIAL_DEKATRINE)-SUM(QUANTITE_UTILISEE_DEKATRINE) as STOCK_FINAL_DEKATRINE FROM phytotable WHERE PRODUITS_PHYTOSANITAIRES='DEKATRINE'"
                    )
                    stock_final_dekatrine = cur.fetchone()[0]

                    # Set the value of the `DEKATRINEfinal` field.
                    self.DEKATRINEfinal.setValue(stock_final_dekatrine)

                    if stock_final_dekatrine == 0:
                        cur.execute(
                            "UPDATE phytotable SET  QUANTITE_ACHETE_DEKATRINE = 0 WHERE PRODUITS_PHYTOSANITAIRES ='DEKATRINE'")

                    # Commit the changes.
                    conn.commit()
            except Exception as e:
                print(e)

        def importraticide(self):
            """Imports the stock of RATICIDE from the database."""
            try:
                with mysql.connector.connect(host="localhost", user=user, password=password,database="datta") as conn:
                    cur = conn.cursor()

                    # Get the total stock of RATICIDE.
                    cur.execute(
                        "SELECT SUM(STOCK_INITIAL_RATICIDE)-SUM(QUANTITE_UTILISEE_RATICIDE) as STOCK_FINAL_RATICIDE FROM phytotable WHERE PRODUITS_PHYTOSANITAIRES='RATICIDE'"
                    )
                    stock_final_raticide = cur.fetchone()[0]

                    # Set the value of the `RATICIDEfinal` field.
                    self.RATICIDEfinal.setValue(stock_final_raticide)

                    if stock_final_raticide == 0:
                        cur.execute(
                            "UPDATE phytotable SET   QUANTITE_ACHETE_RATICIDE = 0 WHERE PRODUITS_PHYTOSANITAIRES ='RATICIDE'")
                    # Commit the changes.
                    conn.commit()
            except Exception as e:
                print(e)

        def importtexto(self):
            """Imports the stock of TEXTO from the database.
            If the stock of TEXTO is 0, the initial stock is also set to 0.
            """
            try:
                with mysql.connector.connect(host="localhost", user=user, password=password,database="datta") as conn:
                    cur = conn.cursor()

                    # Get the total stock of TEXTO.
                    cur.execute(
                        "SELECT SUM(STOCK_INITIAL_TEXTO)-SUM(QUANTITE_UTILISEE_TEXTO) as STOCK_FINAL_TEXTO FROM phytotable WHERE PRODUITS_PHYTOSANITAIRES='TEXTO'"
                    )
                    stock_final_texto = cur.fetchone()[0]

                    # Set the value of the `TEXTOfinal` field.
                    self.TEXTOfinal.setValue(stock_final_texto)

                    # If the stock of TEXTO is 0, set the initial stock to 0 as well.
                    if stock_final_texto == 0:
                        cur.execute(
                            "UPDATE phytotable SET   QUANTITE_ACHETE_TEXTO = 0 WHERE PRODUITS_PHYTOSANITAIRES ='TEXTO'")
                        conn.commit()

            except Exception as e:
                print(e)

        def delete_item(self):
            try:
                msgbox = QMessageBox()
                msgbox.setWindowTitle("تحذير")
                msgbox.setText("هل ترغب في حذف العنصر !")
                yesbutton = QtWidgets.QPushButton("نعم")
                nobuttons = QtWidgets.QPushButton("لا")
                # msgbox.setStandardButtons(QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
                msgbox.addButton(yesbutton, QMessageBox.ButtonRole.ActionRole)
                msgbox.addButton(nobuttons, QMessageBox.ButtonRole.ActionRole)
                push = msgbox.exec()
                if msgbox.clickedButton() == nobuttons:
                    return nobuttons

                elif msgbox.clickedButton() == yesbutton:



                    curentrow = self.textEdit.currentRow()
                    id_ = self.textEdit.item(curentrow, 10).text()
                    database = mysql.connector.connect(host="localhost", user=user, password=password,database="datta")
                    curs = database.cursor()
                    curs.execute("DELETE FROM phytotable WHERE id=%s", (id_,))
                    curs.execute("DELETE FROM phytotablebidi WHERE idbidi=%s", (id_,))
                    database.commit()
                    dattabase.close()
                    msgbox = QtWidgets.QMessageBox()
                    msgbox.setWindowTitle('تأكيد')
                    msgbox.setText('تم حذف العنصر')
                    msgbox.exec()
            except Exception as e:
                print(e)

        def fiche_traitement(self):
            try:
                self.msgbox = QtWidgets.QDialog()
                self.msgbox.setGeometry(300, 100, 935, 560)
                self.msgbox.setStyleSheet(""" QWidget
                            {
                                color: #000000;
                                background-color: #ffffff;
                                border-width: 1px;
                                border-color: #1e1e1e;
                                border-style: solid;
                                border-radius: 6;
                                padding: 0px;
                                font-size: 18px;

                                padding-left: 1px;
                                padding-right: 1px
                            }
                            QWidget:item:hover
                            {
                                background-color: #3daee9;
                                color: #eff0f1;
                            }
                            QWidget:item:selected
                            {
                                background-color: #00000;
                            }
                            QWidget:disabled
                            {
                                color: #454545;
                                background-color: #31363b;
                            }
                            QPushButton
                            {
                                color: #000000;
                                background-color:#ade3e7;
                                border-width: 1px;
                                border-color: #1e1e1e;
                                border-style: solid;
                                border-radius: 6;
                                padding: 3px;
                                font-size: 12px;
                                padding-left: 5px;
                                padding-right: 5px;
                                min-width: 40px
                            }
                            QPushButton:disabled
                            {
                                background-color: #31363b;
                                border-width: 1px;
                                border-color: #454545;
                                border-style: solid;
                                padding-top: 5px;
                                padding-bottom: 5px;
                                padding-left: 10px;
                                padding-right: 10px;
                                border-radius: 2px;
                                color: #454545;
                            }

                            QPushButton:pressed
                            {
                                background-color: #3daee9;
                                padding-top: -15px;
                                padding-bottom: -17px;
                            }
                            QPushButton:hover
                            {
                                border: 1px solid #ff8c00;
                                color: #000000;
                            }
                             QLabel
                            {
                                font-size: 18px;
                                border: 0px solid orange;
                            }

                        """)
                self.msgbox.setWindowTitle("ورقة المعالجة")

                self.ok_button = QtWidgets.QPushButton("طباعة", self.msgbox,
                                                       clicked=lambda: self.docx_fichetraitement())
                self.ok_button.setGeometry(580, 400, 200, 40)
                self.cancel_button = QtWidgets.QPushButton("خروج", self.msgbox, clicked=lambda: self.msgbox.close())
                self.cancel_button.setGeometry(170, 400, 200, 40)

                self.save_button = QtWidgets.QPushButton("حفظ", self.msgbox, clicked=lambda: self.save_fichetraitemen())
                self.save_button.setGeometry(375, 400, 200, 40)

                # msgbox.setText("\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\n\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t")

                # msgbox.addButton(ok_button,QMessageBox.ButtonRole.ActionRole)
                # msgbox.addButton(cancel_button,QMessageBox.ButtonRole.ActionRole)
                fichetraitementxt = QtWidgets.QLabel("<h2>ورقة المعالجة<h2/>", self.msgbox)
                fichetraitementxt.setGeometry(150, 5, 400, 60)
                dateDeTraitementtx = QtWidgets.QLabel("التاريخ :", self.msgbox)
                dateDeTraitementtx.setGeometry(795, 70, 100, 30)
                self.dateDeTraitement = QtWidgets.QDateEdit(self.msgbox)
                self.dateDeTraitement.setDate(self.dateday)
                self.dateDeTraitement.setGeometry(510, 60, 200, 50)
                lieuDeTraitementtxt = QtWidgets.QLabel("مكان المالجة:", self.msgbox)
                lieuDeTraitementtxt.setGeometry(220, 70, 230, 30)
                self.lieuDeTraitement = QtWidgets.QComboBox(self.msgbox)
                self.lieuDeTraitement.setGeometry(50, 60, 200, 50)
                self.lieuDeTraitement.addItem("")
                self.lieuDeTraitement.addItem("المخزن الرئيسي غليزان")
                self.lieuDeTraitement.addItem("مخزن ماسرة")
                self.lieuDeTraitement.addItem("مخزن الكهف الازرق")
                self.lieuDeTraitement.addItem("مخزن زمورة")
                self.lieuDeTraitement.addItem("المحطة الجديدة منداس")
                self.lieuDeTraitement.addItem("محطة منداس")
                self.lieuDeTraitement.addItem("مخزن اوفلا ")
                self.lieuDeTraitement.addItem("مخزن بلعسل")

                self.Naturedeproduittraitétxt = QtWidgets.QLabel("طبيعة المنتج المعالج:", self.msgbox)
                self.Naturedeproduittraitétxt.setGeometry(665, 125, 230, 30)

                self.Naturedeproduittraité = QtWidgets.QTextEdit(self.msgbox)
                self.Naturedeproduittraité.setGeometry(510, 115, 200, 50)
                self.Naturedeproduittraité.setAlignment(Qt.AlignmentFlag.AlignRight)

                self.quantitedeproduittraitétxt = QtWidgets.QLabel("كمية المنتج المعالج:", self.msgbox)
                self.quantitedeproduittraitétxt.setGeometry(220, 125, 230, 30)

                self.quantitedeproduittraité = QtWidgets.QLineEdit(self.msgbox)

                self.quantitedeproduittraité.setGeometry(50, 115, 200, 50)
                self.quantitedeproduittraité.setAlignment(Qt.AlignmentFlag.AlignRight)

                self.dateDebutDeTraitementtxt = QtWidgets.QLabel("تاريخ بداية المعالجة:", self.msgbox)
                self.dateDebutDeTraitementtxt.setGeometry(650, 185, 242, 30)

                self.dateDebutDeTraitement = QtWidgets.QDateEdit(self.msgbox)
                self.dateDebutDeTraitement.setGeometry(510, 175, 200, 50)
                self.dateDebutDeTraitement.setDate(self.dateday)

                self.dateDefinDeTraitementtxt = QtWidgets.QLabel("تاريخ نهاية المعالجة:", self.msgbox)
                self.dateDefinDeTraitementtxt.setGeometry(220, 185, 230, 30)

                self.dateDefinDeTraitement = QtWidgets.QDateEdit(self.msgbox)
                self.dateDefinDeTraitement.setGeometry(50, 175, 200, 50)
                self.dateDefinDeTraitement.setDate(self.dateday)

                self.natureDuProduitUtilisétxt = QtWidgets.QLabel("نوع المبيد المستخدم:", self.msgbox)
                self.natureDuProduitUtilisétxt.setGeometry(660, 245, 230, 30)

                self.natureDuProduitUtilisé = QtWidgets.QComboBox(self.msgbox, editable=True)
                self.natureDuProduitUtilisé.setGeometry(510, 235, 200, 50)
                self.natureDuProduitUtilisé.addItem("")
                self.natureDuProduitUtilisé.addItem("PHOSTOXIN(PH3)")
                self.natureDuProduitUtilisé.addItem("ACTELLIC")
                self.natureDuProduitUtilisé.addItem("CIRATHRINE")
                self.natureDuProduitUtilisé.addItem("RATICIDE")
                self.natureDuProduitUtilisé.addItem("TEXTO")

                self.quantiteDuProduitUtilisétxt = QtWidgets.QLabel("كمية المبيد المستخدمة:", self.msgbox)
                self.quantiteDuProduitUtilisétxt.setGeometry(220, 245, 230, 30)

                self.quantiteDuProduitUtilisé = QtWidgets.QLineEdit(self.msgbox)
                self.quantiteDuProduitUtilisé.setGeometry(50, 235, 200, 50)
                self.quantiteDuProduitUtilisé.setAlignment(Qt.AlignmentFlag.AlignRight)

                self.doseDeProduittxt = QtWidgets.QLabel("تركيز المبيد :", self.msgbox)
                self.doseDeProduittxt.setGeometry(730, 305, 160, 30)

                self.doseDeProduit = QtWidgets.QLineEdit(self.msgbox)
                self.doseDeProduit.setGeometry(510, 295, 200, 50)
                self.doseDeProduit.setAlignment(Qt.AlignmentFlag.AlignRight)
                # self.doseDeProduit.setInputMask("999/aaa 999/aaa")

                self.volumdelottxt = QtWidgets.QLabel("حجم المساحة المعاجة:", self.msgbox)
                self.volumdelottxt.setGeometry(250, 305, 200, 30)
                self.volumdelot = QtWidgets.QLineEdit(self.msgbox)
                self.volumdelot.setGeometry(50, 295, 200, 50)
                self.volumdelot.setAlignment(Qt.AlignmentFlag.AlignRight)
                self.msgbox.show()
                self.msgbox.exec()
            except Exception as e:
                print(e)

        def docx_fichetraitement(self):
            try:
                datte = self.dateDeTraitement.text()
                tproduitphyto = self.natureDuProduitUtilisé.currentText()
                produitphytot1 = arabic_reshaper.reshape(tproduitphyto)
                produitphyto = bidi.algorithm.get_display(produitphytot1)
                tquantiteutilisee = self.quantiteDuProduitUtilisé.text()
                tquantiteutilisee1 = arabic_reshaper.reshape(tquantiteutilisee)
                quantiteutilisee = bidi.algorithm.get_display(tquantiteutilisee1)
                datedebutdetraitemen = self.dateDebutDeTraitement.text()

                datedeprom = self.dateDefinDeTraitement.text()
                lieudetraitement = self.lieuDeTraitement.currentText()
                tnatureetquantite = self.Naturedeproduittraité.toPlainText()
                tnatureetquantite1 = arabic_reshaper.reshape(tnatureetquantite)
                natureetquantite = bidi.algorithm.get_display(tnatureetquantite1)
                tqantitedeprduittraite = self.quantitedeproduittraité.text()
                tqantitedeprduittraite1 = arabic_reshaper.reshape(tqantitedeprduittraite)
                qantitedeprduittraite = bidi.algorithm.get_display(tqantitedeprduittraite1)
                tvolumdelot = self.volumdelot.text()
                tvolumdelot1 = arabic_reshaper.reshape(tvolumdelot)
                volumdelot = bidi.algorithm.get_display(tvolumdelot1)
                # nature=self.natur.text()
                tdosage = self.doseDeProduit.text()
                tdosage1 = arabic_reshaper.reshape(tdosage)
                dosage = bidi.algorithm.get_display(tdosage1)
                doctraitemen = DocxTemplate("Docxfiles/fiche de traitemen/fichedetraitementemplate.docx")
                doctraitemen.render(
                    {
                        "vdl": volumdelot,
                        "dt": datte,
                        "mg": lieudetraitement,
                        "pd": natureetquantite,
                        "pde": qantitedeprduittraite,
                        "dtd": datedebutdetraitemen,
                        "dtf": datedeprom,
                        "npd": produitphyto,
                        "dsg": dosage,
                        "pdq": quantiteutilisee,
                    }
                )
                self.tempfilebd = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                doctraitemen.save(self.tempfilebd)
                # imp = os.startfile(tempfilebd, "print")
                if self.tempfilebd:
                    os.startfile(self.tempfilebd)

            except Exception as e:
                print(e)

        def save_fichetraitemen(self):
            try:
                datte = self.dateDeTraitement.text()
                tproduitphyto = self.natureDuProduitUtilisé.currentText()
                produitphytot1 = arabic_reshaper.reshape(tproduitphyto)
                produitphyto = bidi.algorithm.get_display(produitphytot1)
                tquantiteutilisee = self.quantiteDuProduitUtilisé.text()
                tquantiteutilisee1 = arabic_reshaper.reshape(tquantiteutilisee)
                quantiteutilisee = bidi.algorithm.get_display(tquantiteutilisee1)
                datedebutdetraitemen = self.dateDebutDeTraitement.text()

                datedeprom = self.dateDefinDeTraitement.text()
                lieudetraitement = self.lieuDeTraitement.currentText()
                tnatureetquantite = self.Naturedeproduittraité.toPlainText()
                tnatureetquantite1 = arabic_reshaper.reshape(tnatureetquantite)
                natureetquantite = bidi.algorithm.get_display(tnatureetquantite1)
                tqantitedeprduittraite = self.quantitedeproduittraité.text()
                tqantitedeprduittraite1 = arabic_reshaper.reshape(tqantitedeprduittraite)
                qantitedeprduittraite = bidi.algorithm.get_display(tqantitedeprduittraite1)
                tvolumdelot = self.volumdelot.text()
                tvolumdelot1 = arabic_reshaper.reshape(tvolumdelot)
                volumdelot = bidi.algorithm.get_display(tvolumdelot1)
                # nature=self.natur.text()
                tdosage = self.doseDeProduit.text()
                tdosage1 = arabic_reshaper.reshape(tdosage)
                dosage = bidi.algorithm.get_display(tdosage1)
                doctraitemen = DocxTemplate("Docxfiles/fiche de traitemen/fichedetraitementemplate.docx")
                # Render the template with the data
                doctraitemen.render(
                    {
                        "vdl": volumdelot,
                        "dt": datte,
                        "mg": lieudetraitement,
                        "pd": natureetquantite,
                        "pde": qantitedeprduittraite,
                        "dtd": datedebutdetraitemen,
                        "dtf": datedeprom,
                        "npd": produitphyto,
                        "dsg": dosage,
                        "pdq": quantiteutilisee,
                    }
                )
                # Get the path to save the file
                path, _ = QFileDialog.getSaveFileName(None, "Enregistrer la fiche", "", "Fichiers DOCX (*.docx)")

                # Save the file
                if path:
                    doctraitemen.save(path)
                    # Display a message box confirming that the file was saved
                    msgbox = QtWidgets.QMessageBox()
                    msgbox.setWindowTitle('تاكيد')
                    msgbox.setText("تم حفظ الملف بنجاح")
                    msgbox.exec()

            except Exception as e:
                print(e)


    if __name__ == "__main__":
        import sys

        app = QtWidgets.QApplication(sys.argv)
        MainWindow = QtWidgets.QMainWindow()
        ui = Phyto_Window()
        ui.phyoto_produit(MainWindow)
        MainWindow.show()

        sys.exit(app.exec())
except Exception as e:
    print(e)
