import locale
import tempfile
import docx
from PyQt6 import QtCore, QtGui, QtWidgets
from PyQt6.QtCore import QTimer, QDate
from PyQt6.QtWidgets import *
import csv
import sqlite3
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches, RGBColor
import os
import mysql.connector
from docxtpl import DocxTemplate
userslist = ["nadir", "qualite"]
passwordlist = ["Nadir206@", "qualite48"]

for user ,password in zip(userslist,passwordlist):
    try:
        # Connection for entrytable

        database = mysql.connector.connect(
            host="localhost",
            user=user,
            password=password
        )

        curs = database.cursor()

        # Create the database if it doesn't exist
        curs.execute("CREATE DATABASE IF NOT EXISTS datta_legumesec_entry")
        curs.execute("USE datta_legumesec_entry")

        curs.execute("""
            CREATE TABLE IF NOT EXISTS entrytable (
                id INT AUTO_INCREMENT PRIMARY KEY,
                nlot TEXT,
                nlotsch TEXT,
                le TEXT,
                produit TEXT,
                origine TEXT,
                quantité FLOAT,
                unité TEXT,
                sacherie TEXT,
                fournisseur TEXT,
                imatricule TEXT,
                date_recolte TEXT,
                date_CND TEXT,
                DLUM TEXT,
                quantitérizentrée FLOAT,
                quantitérizetvntrée FLOAT,
                quantitépchentrée FLOAT,
                quantitépchplentrée FLOAT,
                quantitélenentrée FLOAT,
                quantitélenrougeentrée FLOAT,
                quantitélenplntrée FLOAT,
                quantitéharntrée FLOAT,
                quantitéharlsbntrée FLOAT,
                quantitépchneufentrée FLOAT,
                quantitépchhuitntrée FLOAT,
                quantitépchsixntrée FLOAT,
                quantitélenvertntrée FLOAT,
                quantitérizetotale FLOAT,
                quantitérizetvtotale FLOAT,
                quantitépchtotale FLOAT,
                quantitépchpltotale FLOAT,
                quantitélentotale FLOAT,
                quantitélenrougetotale FLOAT,
                quantitélenpltotale FLOAT,
                quantitéhartotale FLOAT,
                quantitéharlsbtotal FLOAT,
                quantitéentrétotal FLOAT,
                quantitépchneufetotal FLOAT,
                quantitépchhuittotal FLOAT,
                quantitépchsixtotal FLOAT,
                quantitélenverttotal FLOAT
            )
        """)

        database.commit()
        database.close()

        # Connection for outtable
        database1 = mysql.connector.connect(
            host='localhost',
             user=user,
            password=password
        )

        curs1 = database1.cursor()

        # Create the database if it doesn't exist
        curs1.execute("CREATE DATABASE IF NOT EXISTS datta_legumsec_out")
        curs1.execute("USE datta_legumsec_out")

        curs1.execute("""
            CREATE TABLE IF NOT EXISTS outtable (
                id INT AUTO_INCREMENT PRIMARY KEY,
                le TEXT,
                nlot TEXT,
                nlotsch TEXT,
                unité TEXT,
                produit TEXT,
                origine TEXT,
                quantité FLOAT,
                achteur TEXT,
                sacherie TEXT,
                imatricule TEXT,
                date_recoltes TEXT,
                date_CND TEXT,
                DLUM TEXT,
                fournisseur TEXT,
                quantitérizentré FLOAT,
                quantitérizesorté FLOAT,
                quantitérizerest FLOAT,
                quantitérizetotal FLOAT,
                quantitérizetvntré FLOAT,
                quantitérizetvsorté FLOAT,
                quantitérizetvrest FLOAT,
                quantitérizetvtotal FLOAT,
                quantitepchentré FLOAT,
                quantitépchsorté FLOAT,
                quantitépchrest FLOAT,
                quantitépchtotal FLOAT,
                quantitépchplentré FLOAT,
                quantitépchplsorté FLOAT,
                quantitépchplrest FLOAT,
                quantitépchpltotal FLOAT,
                quantitélenentré FLOAT,
                quantitélensorté FLOAT,
                quantitélenrest FLOAT,
                quantitélentotal FLOAT,
                quantitélenrougeentré FLOAT,
                quantitélenrougesorté FLOAT,
                quantitélenrougerest FLOAT,
                quantitélenrougetotal FLOAT,
                quantitélenplentré FLOAT,
                quantitélenplsorté FLOAT,
                quantitélenplrest FLOAT,
                quantitélenpltotal FLOAT,
                quantitéharentré FLOAT,
                quantitéharsorté FLOAT,
                quantitéharrest FLOAT,
                quantitéharlsbentré FLOAT,
                quantitéharlsbsorté FLOAT,
                quantitéharlsbrest FLOAT,
                quantitéhartotal FLOAT,
                quantitéentretotal FLOAT,
                quantitérestotal FLOAT,
                quantitépchneufeentré FLOAT,
                quantitépchneufesorté FLOAT,
                quantitépchneuferest FLOAT,
                quantitépchneufetotal FLOAT,
                quantitépchhuitentré FLOAT,
                quantitépchhuitsorté FLOAT,
                quantitépchhuitrest FLOAT,
                quantitépchhuittotal FLOAT,
                quantitépchsixentré FLOAT,
                quantitépchsixsorté FLOAT,
                quantitépchsixrest FLOAT,
                quantitépchsixtotal FLOAT,
                quantitélenvertentré FLOAT,
                quantitélenvertsorté FLOAT,
                quantitélenvertrest FLOAT,
                quantitélenverttotal FLOAT
            )
        """)

        database1.commit()
        database1.close()

    except mysql.connector.Error as e:
        print("Error:", e)




try:
    class Stock_Legumesec(object):

        def stock_legumesec(self, MainWindow):

            MainWindow.setObjectName("MainWindow")
            MainWindow.resize(1338, 700)
            self.centralwidget = QtWidgets.QWidget(MainWindow)
            self.centralwidget.setObjectName("centralwidget")
            self.verticalLayout = QtWidgets.QVBoxLayout(self.centralwidget)
            self.verticalLayout.setObjectName("verticalLayout")
            self.verticalLayout.setContentsMargins(0,0,0,0)
            self.tabWidget = QtWidgets.QTabWidget(self.centralwidget)
            self.tabWidget.setStyleSheet("""QToolTip
        {
            border: 1px solid #76797C;
            background-color: rgb(90, 102, 117);;
            color: white;
            padding: 5px;
            opacity: 200;
        }

        QWidget
        {
            color: #000000;
            background-color: #ffffff;
            selection-background-color:#3daee9;
            selection-color: #3daee9;
            background-clip: border;
            border-image: none;
            border: 0px transparent black;
            outline: 0;
        }

        QWidget:item:hover
        {
            background-color: #3daee9;
            color: #eff0f1;
        }

        QWidget:item:selected
        {
            background-color: #3daee9;
        }



        QWidget:disabled
        {
            color: #454545;
            background-color: #31363b;
        }

        QAbstractItemView
        {
            alternate-background-color: #31363b;
            color: #eff0f1;
            border: 1px solid 3A3939;
            border-radius: 2px;
        }

        QWidget:focus, QMenuBar:focus
        {
            border: 1px solid #3daee9;
        }

        QTabWidget:focus, QCheckBox:focus, QRadioButton:focus, QSlider:focus
        {
            border: none;
        }

        QLineEdit
        {
            background-color: #FDFEFE;
    border: 1px solid #76797C;
    color:#000000;
    border-radius: 0.25em;
    padding: 0.0em 0.0em;
    font-size: 1.25rem;
    cursor: pointer;
        }
        QDoubleSpinBox
        {
            background-color: #FDFEFE;
            padding: 1px;
            border-style: solid;
            border: 1px solid #76797C;
            border-radius: 0px;
            color:#000000;
            font-size: 11px;
            font-weight: bold;

        }
        QDoubleSpinBox:focus{
            background-color: #FDFEFE;
            border-style: solid;
            border: 2px solid #76797C;
            border-radius: 4px;
            border-color: #ff8c00;
        }
        QDoubleSpinBox::drop-down
        {
            subcontrol-origin: padding;
            subcontrol-position: top right;
            width: 1px;

            border-left-width: 0px;
            border-left-color: #302629;
            border-left-style: solid;
            border-top-right-radius: 1px;
            border-bottom-right-radius: 1px;
        }



        QGroupBox {
            border:1px solid #76797C;
            border-radius: 2px;
            margin-top: 5px;
        }

        QGroupBox::title {
            subcontrol-origin: margin;
            subcontrol-position: top center;
            padding-left: 4px;
            padding-right: 4px;
            padding-top: 4px;
        }

        QAbstractScrollArea
        {
            border-radius: 2px;
            border: 1px solid #76797C;
            background-color: transparent;
        }

        QScrollBar:horizontal
        {
            height: 15px;
            margin: 3px 15px 3px 15px;
            border: 1px transparent #2A2929;
            border-radius: 4px;
            background-color: #2A2929;
        }

        QScrollBar::handle:horizontal
        {
            background-color: #605F5F;
            min-width: 5px;
            border-radius: 4px;
        }

        QScrollBar::add-line:horizontal
        {
            margin: 0px 3px 0px 3px;
            border-image: url(:/qss_icons/Dark_rc/right_arrow_disabled.png);
            width: 10px;
            height: 10px;
            subcontrol-position: right;
            subcontrol-origin: margin;
        }

        QScrollBar::sub-line:horizontal
        {
            margin: 0px 3px 0px 3px;
            border-image: url(:/qss_icons/Dark_rc/left_arrow_disabled.png);
            height: 10px;
            width: 10px;
            subcontrol-position: left;
            subcontrol-origin: margin;
        }

        QScrollBar::add-line:horizontal:hover,QScrollBar::add-line:horizontal:on
        {
            border-image: url(:/qss_icons/Dark_rc/right_arrow.png);
            height: 10px;
            width: 10px;
            subcontrol-position: right;
            subcontrol-origin: margin;
        }


        QScrollBar::sub-line:horizontal:hover, QScrollBar::sub-line:horizontal:on
        {
            border-image: url(:/qss_icons/Dark_rc/left_arrow.png);
            height: 10px;
            width: 10px;
            subcontrol-position: left;
            subcontrol-origin: margin;
        }

        QScrollBar::up-arrow:horizontal, QScrollBar::down-arrow:horizontal
        {
            background: none;
        }


        QScrollBar::add-page:horizontal, QScrollBar::sub-page:horizontal
        {
            background: none;
        }

        QScrollBar:vertical
        {
            background-color: #2A2929;
            width: 15px;
            margin: 15px 3px 15px 3px;
            border: 1px transparent #2A2929;
            border-radius: 4px;
        }

        QScrollBar::handle:vertical
        {
            background-color: #605F5F;
            min-height: 5px;
            border-radius: 4px;
        }

        QScrollBar::sub-line:vertical
        {
            margin: 3px 0px 3px 0px;
            border-image: url(:/qss_icons/Dark_rc/up_arrow_disabled.png);
            height: 10px;
            width: 10px;
            subcontrol-position: top;
            subcontrol-origin: margin;
        }

        QScrollBar::add-line:vertical
        {
            margin: 3px 0px 3px 0px;
            border-image: url(:/qss_icons/Dark_rc/down_arrow_disabled.png);
            height: 10px;
            width: 10px;
            subcontrol-position: bottom;
            subcontrol-origin: margin;
        }

        QScrollBar::sub-line:vertical:hover,QScrollBar::sub-line:vertical:on
        {

            border-image: url(:/qss_icons/Dark_rc/up_arrow.png);
            height: 10px;
            width: 10px;
            subcontrol-position: top;
            subcontrol-origin: margin;
        }


        QScrollBar::add-line:vertical:hover, QScrollBar::add-line:vertical:on
        {
            border-image: url(:/qss_icons/Dark_rc/down_arrow.png);
            height: 10px;
            width: 10px;
            subcontrol-position: bottom;
            subcontrol-origin: margin;
        }

        QScrollBar::up-arrow:vertical, QScrollBar::down-arrow:vertical
        {
            background: none;
        }


        QScrollBar::add-page:vertical, QScrollBar::sub-page:vertical
        {
            background: none;
        }

        QTextEdit
        {
            background-color: #302629;
            color: #eff0f1;
            border: 1px solid #76797C;
        }

        QPlainTextEdit
        {
            background-color: #302629;;
            color: #eff0f1;
            border-radius: 2px;
            border: 1px solid #76797C;
        }

        QHeaderView::section
        {
            background-color: #76797C;
            color: #eff0f1;
            padding: 1px;
            border: 1px solid #76797C;
        }

        QSizeGrip {
            width: 12px;
            height: 12px;
        }


        QMainWindow::separator
        {
            background-color: #31363b;
            color: white;
            padding-left: 4px;
            spacing: 2px;
            border: 1px dashed #76797C;
        }

        QMainWindow::separator:hover
        {

            background-color: #787876;
            color: white;
            padding-left: 4px;
            border: 1px solid #76797C;
            spacing: 2px;
        }


        QMenu::separator
        {
            height: 1px;
            background-color: #76797C;
            color: white;
            padding-left: 4px;
            margin-left: 10px;
            margin-right: 5px;
        }


        QFrame
        {
            border-radius: 2px;
            border: 1px solid #76797C;
        }

        QFrame[frameShape="0"]
        {
            border-radius: 2px;
            border: 1px transparent #76797C;
        }

        QStackedWidget
        {
            border: 1px transparent black;
        }


        QPushButton
        {
            color: #00000;
            background-color:#ade3e7;
            border-width: 1px;
            border-color: #1e1e1e;
            border-style: solid;
            border-radius: 6;
            padding: 3px;
            font-size: 12px;
            padding-left: 5px;
            padding-right: 5px;
            min-width: 40px;

        }

        QPushButton:disabled
        {
            background-color:#03ecff;
            border-width: 1px;
            border-color: #454545;
            border-style: solid;
            padding-top: 5px;
            padding-bottom: 5px;
            padding-left: 10px;
            padding-right: 10px;
            border-radius: 2px;
            color: #454545;
        }
        QPushButton:pressed
        {
            background-color: #3daee9;
            padding-top: -15px;
            padding-bottom: -17px;
        }

        QComboBox {
    background-color: #FDFEFE;
    border: 1px solid #76797C;
    color:#000000;
    border-radius: 0.25em;
    padding: 0.0em 0.0em;
    font-size: 1.25rem;
    cursor: pointer;
}

QComboBox::drop-down {
    subcontrol-origin: padding;
    subcontrol-position: top right;
    width: 1.3em;
    border-left: 0px solid #777;
    border-radius: 0.25em;
}

QComboBox::drop-down::icon {
    image: url('E:/pythonProject_moullin-application.3.5/images/down-arroww.png');
}

        QPushButton:checked{
            background-color: #76797C;
            border-color: #6A6969;
        }
        QComboBox:on
        {
            padding-top: 1px;
            padding-left: 1px;
            selection-background-color: #FDFEFE;
        }
        QComboBox QAbstractItemView
        {
            background-color: #FDFEFE;
            border-radius: 2px;
            border: 1px solid #76797C;
            color:#000000;
            selection-background-color: #000000;
        }
        QComboBox:hover,QDoubleSpinBox:Hover,QPushButton:hover,QAbstractSpinBox:hover,QLineEdit:hover,QTextEdit:hover,QPlainTextEdit:hover,QAbstractView:hover,QTreeView:hover
        {
            border: 1px solid #ff8c00;
            color: #000000;
        }

        


        QLabel
        {
            border: 2px solid black;
        }

        QTabWidget{
            border: 0px transparent black;
        }

        QTabWidget::pane {
            border: 1px solid #76797C;
            padding: 5px;
            margin: 0px;
        }

        QTabBar
        {
            qproperty-drawBase: 0;
            left: 5px; /* move to the right by 5px */
            border-radius: 3px;
        }

        QTabBar:focus
        {
            border: 0px transparent black;
        }

        QTabBar::close-button  {
            image: url(:/qss_icons/Dark_rc/close.png);
            background: transparent;
        }

        QTabBar::close-button:hover
        {
            image: url(:/qss_icons/Dark_rc/close-hover.png);
            background: transparent;
        }

        QTabBar::close-button:pressed {
            image: url(:/qss_icons/Dark_rc/close-pressed.png);
            background: transparent;
        }

        /* TOP TABS */
        QTabBar::tab:top {
            color: #eff0f1;
            border: 1px solid #76797C;
            border-bottom: 1px transparent black;
            background-color: #31363b;
            padding: 5px;
            min-width: 10px;
            border-top-left-radius: 2px;
            border-top-right-radius: 2px;
        }

        QTabBar::tab:top:!selected
        {
            color: #eff0f1;
            background-color: #54575B;
            border: 1px solid #76797C;
            border-bottom: 1px transparent black;
            border-top-left-radius: 2px;
            border-top-right-radius: 2px;    
        }

        QTabBar::tab:top:!selected:hover {
            background-color: #3daee9;
        }

        /* BOTTOM TABS */
        QTabBar::tab:bottom {
            color: #eff0f1;
            border: 1px solid #76797C;
            border-top: 1px transparent black;
            background-color: #31363b;
            padding: 5px;
            border-bottom-left-radius: 2px;
            border-bottom-right-radius: 2px;
            min-width: 10px;
        }

        QTabBar::tab:bottom:!selected
        {
            color: #eff0f1;
            background-color: #54575B;
            border: 1px solid #76797C;
            border-top: 1px transparent black;
            border-bottom-left-radius: 2px;
            border-bottom-right-radius: 2px;
        }

        QTabBar::tab:bottom:!selected:hover {
            background-color: #3daee9;
        }

        /* LEFT TABS */
        QTabBar::tab:left {
            color: #eff0f1;
            border: 1px solid #76797C;
            border-left: 1px transparent black;
            background-color: #31363b;
            padding: 5px;
            border-top-right-radius: 2px;
            border-bottom-right-radius: 2px;
            min-height: 50px;
        }

        QTabBar::tab:left:!selected
        {
            color: #eff0f1;
            background-color: #54575B;
            border: 1px solid #76797C;
            border-left: 1px transparent black;
            border-top-right-radius: 2px;
            border-bottom-right-radius: 2px;
        }

        QTabBar::tab:left:!selected:hover {
            background-color: #3daee9;
        }


        /* RIGHT TABS */
        QTabBar::tab:right {
            color: #eff0f1;
            border: 1px solid #76797C;
            border-right: 1px transparent black;
            background-color: #31363b;
            padding: 5px;
            border-top-left-radius: 2px;
            border-bottom-left-radius: 2px;
            min-height: 50px;
        }

        QTabBar::tab:right:!selected
        {
            color: #eff0f1;
            background-color: #54575B;
            border: 1px solid #76797C;
            border-right: 1px transparent black;
            border-top-left-radius: 2px;
            border-bottom-left-radius: 2px;
        }





        QSlider::groove:horizontal {
            border: 1px solid #565a5e;
            height: 4px;
            background: #565a5e;
            margin: 0px;
            border-radius: 2px;
        }

        QSlider::handle:horizontal {
            background: #302629;
            border: 1px solid #565a5e;
            width: 16px;
            height: 16px;
            margin: -8px 0;
            border-radius: 9px;
        }

        QSlider::groove:vertical {
            border: 1px solid #565a5e;
            width: 4px;
            background: #565a5e;
            margin: 0px;
            border-radius: 3px;
        }

        QSlider::handle:vertical {
            background: #302629;
            border: 1px solid #565a5e;
            width: 16px;
            height: 16px;
            margin: 0 -8px;
            border-radius: 9px;
        }

        QToolButton {
            background-color: transparent;
            border: 1px transparent #76797C;
            border-radius: 2px;
            margin: 3px;
            padding: 5px;
        }

        QToolButton[popupMode="1"] { /* only for MenuButtonPopup */
         padding-right: 20px; /* make way for the popup button */
         border: 1px #76797C;
         border-radius: 5px;
        }

        QToolButton[popupMode="2"] { /* only for InstantPopup */
         padding-right: 10px; /* make way for the popup button */
         border: 1px #76797C;
        }


        QToolButton:hover, QToolButton::menu-button:hover {
            background-color: transparent;
            border: 1px solid #3daee9;
            padding: 5px;
        }

        QToolButton:checked, QToolButton:pressed,
                QToolButton::menu-button:pressed {
            background-color: #3daee9;
            border: 1px solid #3daee9;
            padding: 5px;
        }

        /* the subcontrol below is used only in the InstantPopup or DelayedPopup mode */
        QToolButton::menu-indicator {
            background-color:ff8c00;
            top: -7px; left: -2px; /* shift it a bit */
        }

        /* the subcontrols below are used only in the MenuButtonPopup mode */
        QToolButton::menu-button {
            border: 1px transparent #76797C;
            border-top-right-radius: 6px;
            border-bottom-right-radius: 6px;
            /* 16px width + 4px for border = 20px allocated above */
            width: 16px;
            outline: none;
        }

        QToolButton::menu-arrow {
           background-color:ff8c00;
        }

        QToolButton::menu-arrow:open {
            border: 1px solid #76797C;
        }

        QPushButton::menu-indicator  {
            subcontrol-origin: padding;
            subcontrol-position: bottom right;
            left: 8px;
        }

        QTableView
        {
            border: 1px solid #76797C;
            gridline-color: #31363b;
            background-color: #FDFEFE;
            color:#000000;
        }


        QTableView, QHeaderView
        {
            background-color: #FDFEFE;
            color:#000000;
            border-radius: 0px;
        }

        QTableView::item:pressed, QListView::item:pressed, QTreeView::item:pressed  {
            background: #FDFEFE;
            color: #000000;
        }

        QTableView::item:selected:active, QTreeView::item:selected:active, QListView::item:selected:active  {
            background: #3daee9;
            color: #000000;
        }


        QHeaderView
        {
            background-color: #FDFEFE;
            border: 1px transparent;
            border-radius: 0px;
            margin: 0px;
            padding: 0px;

        }

        QHeaderView::section  {
            background-color:#80f1f9;
            color: #000000;
            padding: 5px;
            border: 1px solid #76797C;
            border-radius: 0px;
            text-align: center;
        }

        QHeaderView::section::vertical::first, QHeaderView::section::vertical::only-one
        {
            border-top: 1px solid #76797C;
        }

        QHeaderView::section::vertical
        {
            border-top: transparent;
        }

        QHeaderView::section::horizontal::first, QHeaderView::section::horizontal::only-one
        {
            border-left: 1px solid #76797C;
        }

        QHeaderView::section::horizontal
        {
            border-left: transparent;
        }


        QHeaderView::section:checked
         {
            color: #000000;
            background-color: #3daee9;
         }

         /* style the sort indicator */
        QHeaderView::down-arrow {

        }

        QHeaderView::up-arrow {

        }


        QTableCornerButton::section {
            background-color: #31363b;
            border: 1px transparent #76797C;
            border-radius: 0px;
        }

        QToolBox  {
            padding: 5px;
            border: 1px transparent black;
        }

        QToolBox::tab {
            color: #eff0f1;
            background-color: #31363b;
            border: 1px solid #76797C;
            border-bottom: 1px transparent #31363b;
            border-top-left-radius: 5px;
            border-top-right-radius: 5px;
        }

        QToolBox::tab:selected { /* italicize selected tabs */
            font: italic;
            background-color: #31363b;
            border-color: #3daee9;
         }

        QStatusBar::item {
            border: 0px transparent dark;
         }


        QFrame[height="3"], QFrame[width="3"] {
            background-color: #76797C;
        }




        QDateEdit
        {
            background-color: #302629;;
            border-style: solid;
            border: 1px solid #76797C;
            border-radius: 2px;
            padding: 1px;
            min-width: 75px;
        }

        QDateEdit:on
        {
            padding-top: 2px;
            padding-left: 2px;
            selection-background-color: #4a4a4a;
        }

        QDateEdit QAbstractItemView
        {
            background-color: #ff8c00;
            border-radius: 2px;
            border: 1px solid #3375A3;
            selection-background-color:ff8c00;
        }

        QDateEdit::drop-down
        {
            background-color: #FDFEFE;
    border: 1px solid #76797C;
    color:#000000;
    border-radius: 0.25em;
    padding: 0.0em 0.0em;
    font-size: 1.25rem;
    cursor: pointer;
        }   
        QDateTimeEdit
        {
            background-color: #302629;;
            border-style: solid;
            border: 1px solid #76797C;
            border-radius: 2px;
            padding: 1px;
            min-width: 75px;

        }    
        """)
            self.tabWidget.setObjectName("tabWidget")
            self.tabENTRER = QtWidgets.QWidget()
            self.tabENTRER.setObjectName("tabENTRER")

            self.addbtn = QtWidgets.QPushButton(self.tabENTRER)
            self.addbtn.setGeometry(QtCore.QRect(20, 275, 100, 40))
            self.addbtn.setObjectName("addbtn")
            self.addbtn.clicked.connect(self.add)

            ##################################################entre###############

            self.deletebtn = QtWidgets.QPushButton(self.tabENTRER,clicked=lambda :self.delete_entry())
            self.deletebtn.setGeometry(QtCore.QRect(190, 275, 100, 40))
            self.deletebtn.setObjectName("deletebtn")

            self.printbtn = QtWidgets.QPushButton(self.tabENTRER,clicked=lambda :self.print_docx())
            self.printbtn.setGeometry(QtCore.QRect(360, 275, 100, 40))
            self.printbtn.setObjectName("printbtn")



            self.cclstxt = QtWidgets.QLabel(self.tabENTRER)
            self.cclstxt.setGeometry(QtCore.QRect(430, 0, 500, 41))
            font = QtGui.QFont()
            font.setPointSize(14)
            font.setBold(True)
            font.setWeight(75)
            self.cclstxt.setFont(font)
            self.cclstxt.setMouseTracking(False)
            self.cclstxt.setLayoutDirection(QtCore.Qt.LayoutDirection.LeftToRight)
            self.cclstxt.setAutoFillBackground(False)
            self.cclstxt.setFrameShape(QtWidgets.QFrame.Shape.WinPanel)
            self.cclstxt.setLineWidth(0)
            self.cclstxt.setMidLineWidth(0)
            self.cclstxt.setTextFormat(QtCore.Qt.TextFormat.AutoText)
            self.cclstxt.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
            self.cclstxt.setObjectName("cclstxt")
            self.entrielabel = QtWidgets.QLabel("<h3>ENTRER<h3/>", self.tabENTRER)
            self.entrielabel.setGeometry(650, 75, 150, 40)
            self.entrielabel.setFont(font)
            self.titletxt = QtWidgets.QLabel(self.tabENTRER)
            self.titletxt.setGeometry(QtCore.QRect(540, 50, 350, 30))
            font = QtGui.QFont()
            font.setPointSize(14)
            font.setBold(True)
            font.setUnderline(True)
            font.setWeight(75)
            self.titletxt.setFont(font)
            self.titletxt.setAlignment(
                QtCore.Qt.AlignmentFlag.AlignLeading | QtCore.Qt.AlignmentFlag.AlignLeft | QtCore.Qt.AlignmentFlag.AlignVCenter)
            self.titletxt.setObjectName("titletxt")
            #############################################
            self.txtdate = QtWidgets.QLabel(self.tabENTRER)
            self.txtdate.setGeometry(QtCore.QRect(20, 80, 120, 16))
            self.txtdate.setObjectName("txtdate")
            ####################################################
            self.dateedite = QtWidgets.QLineEdit(self.tabENTRER)
            self.dateedite.setInputMask("99-99-9999")
            self.dateedite.setGeometry(QtCore.QRect(115, 77, 140, 30))
            self.dateedite.setObjectName("dateedite")
            ########################################################
            self.produittxt = QtWidgets.QLabel(self.tabENTRER)
            self.produittxt.setGeometry(QtCore.QRect(20, 132, 61, 16))
            self.produittxt.setObjectName("produittxt")
            self.produitcombo = QtWidgets.QComboBox(self.tabENTRER)
            self.produitcombo.setGeometry(QtCore.QRect(115, 130, 140, 30))
            self.produitcombo.setObjectName("produitcombo")
            self.produitcombo.addItem("")
            self.produitcombo.addItem("POIS CHICHES PL")
            self.produitcombo.addItem("POIS CHICHES IMP 12 mm")
            self.produitcombo.addItem("POIS CHICHES IMP 09 mm")
            self.produitcombo.addItem("POIS CHICHES IMP 08 mm")
            self.produitcombo.addItem("POIS CHICHES IMP 06 mm")
            self.produitcombo.addItem("HARICOT BLANC IMP")
            self.produitcombo.addItem("HARICOT LSB")
            self.produitcombo.addItem("LENTILLE PL")
            self.produitcombo.addItem("LENTILLE IMP 05-07 mm")
            self.produitcombo.addItem("LENTILLE IMP vert")
            self.produitcombo.addItem("LENTILLE IMP ROUGE")
            self.produitcombo.addItem("RIZ IMP")
            self.produitcombo.addItem("RIZ ETUVE")
            #####################################################################################
            self.origintxt = QtWidgets.QLabel(self.tabENTRER)
            self.origintxt.setGeometry(QtCore.QRect(270, 133, 121, 16))
            self.origintxt.setObjectName("origintxt")
            ###################################################################
            self.originel = QtWidgets.QComboBox(self.tabENTRER)
            all_countries = [" ","Algeria","Argentina",  "Australia",  "Azerbaijan", "Bangladesh", "Brazil", "Canada", "China", "Colombia",
            "Costa Rica",
            "Ecuador",
            "Egypt",
            "India", "Indonesia",
            "Italy", "Japan", "Kazakhstan",
            "Kuwait", "Kyrgyzstan",
            "Macedonia",  "Malaysia",
            "Mauritius", "Mexico",
            "Pakistan", "Paraguay", "Peru", "Philippines",
            "Poland", "Portugal",  "Romania", "Russian ",
            "Saudi Arabia",
            "Singapore",
            "South Africa", "Spain", "Sri Lanka", "Sudan",
            "Swaziland", "Sweden", "Switzerland",  "Taiwan", "Tajikistan",
            "Tanzania", "Thailand", "Tunisia", "Turkey",
            "Turkmenistan",  "Ukraine",
            "United Kingdom", "United States", "Uruguay", "Uzbekistan",
             "Venezuela", "Vietnam", "Yemen", "Zambia", "Zimbabwe"
        ]
            self.originel.setGeometry(QtCore.QRect(397, 130, 140, 30))
            self.originel.addItems(all_countries)

            ###############################################################################
            self.quatitelineEdit = QtWidgets.QDoubleSpinBox(self.tabENTRER)
            self.quatitelineEdit.setSpecialValueText(' ')
            self.quatitelineEdit.setRange(0, 1000000)
            self.quatitelineEdit.setGeometry(QtCore.QRect(695,130, 140, 30))
            self.quatitelineEdit.setObjectName("quatitelineEdit")
            ################################################################################

            self.txtquantite = QtWidgets.QLabel(self.tabENTRER)
            self.txtquantite.setGeometry(QtCore.QRect(550, 133, 51, 16))
            self.txtquantite.setObjectName("txtquantite")
            #############################################################################
            self.labelmagasin = QtWidgets.QLabel(self.tabENTRER)
            self.labelmagasin.setGeometry(QtCore.QRect(850, 133, 116, 16))
            self.magasincombo = QtWidgets.QComboBox(self.tabENTRER)
            self.magasincombo.setGeometry(QtCore.QRect(980, 130, 140, 30))
            self.magasincombo.setObjectName("magasincombo")
            self.magasincombo.addItem("")
            self.magasincombo.addItem("DOCK CENTRAL")
            self.magasincombo.addItem("MAGASIN MESRA")
            self.magasincombo.addItem("MARCHE GROS BELACEL")


            ###########################################################################
            self.fourniseurcombo = QtWidgets.QComboBox(self.tabENTRER)
            self.fourniseurcombo.setGeometry(QtCore.QRect(397, 185, 140, 30))
            self.fourniseurcombo.setObjectName("fourniseurcombo")
            self.fourniseurcombo.addItem("")
            self.fourniseurcombo.addItem("UCA ALGER")
            self.fourniseurcombo.addItem("UCA ORAN")

            self.fornisseurtxt = QtWidgets.QLabel(self.tabENTRER)
            self.fornisseurtxt.setGeometry(QtCore.QRect(270, 188, 81, 16))
            self.fornisseurtxt.setObjectName("fornisseurtxt")
            #########################################################################
            self.CNDTXT = QtWidgets.QLabel(self.tabENTRER)
            self.CNDTXT.setGeometry(QtCore.QRect(20, 188, 91, 16))
            self.CNDTXT.setObjectName("CNDTXT")
            self.cndcombo = QtWidgets.QComboBox(self.tabENTRER)
            self.cndcombo.setGeometry(QtCore.QRect(115, 185, 140, 30))
            self.cndcombo.setObjectName("cndcombo")
            self.cndcombo.addItem("")
            self.cndcombo.addItem("SAC 50KG")
            self.cndcombo.addItem("SAC 25KG")
            self.cndcombo.addItem("BIG BAG")
            ########################################################################
            self.TRANSPRTTXT = QtWidgets.QLabel(self.tabENTRER)
            self.TRANSPRTTXT.setGeometry(QtCore.QRect(550, 188, 131, 16))
            self.TRANSPRTTXT.setObjectName("TRANSPRTTXT")
            self.TRANSPORTcombo = QtWidgets.QComboBox(self.tabENTRER,editable=True)
            self.TRANSPORTcombo.setGeometry(QtCore.QRect(695, 185, 140, 30))
            self.TRANSPORTcombo.setObjectName("TRANSPORTcombo")
            self.TRANSPORTcombo.addItem("")
            ########################################################################
            self.label_2 = QtWidgets.QLabel(self.tabENTRER)
            self.label_2.setGeometry(QtCore.QRect(850, 188, 81, 16))
            self.label_2.setObjectName("label_2")
            self.imatricullineeditr = QtWidgets.QLineEdit(self.tabENTRER)
            self.imatricullineeditr.setInputMask("99999-999-99")
            self.imatricullineeditr.setGeometry(QtCore.QRect(980, 185, 140, 30))

            #########################################################################
            self.DATEDECNDTXT = QtWidgets.QLabel(self.tabENTRER)
            self.DATEDECNDTXT.setGeometry(QtCore.QRect(270, 243, 120, 13))
            self.DATEDECNDTXT.setObjectName("DATEDECNDTXT")
            self.dateeditedcnd = QtWidgets.QLineEdit(self.tabENTRER)
            self.dateeditedcnd.setInputMask("99-99-9999")
            self.dateeditedcnd.setGeometry(QtCore.QRect(397, 235, 140, 30))
            self.dateeditedcnd.setObjectName("dateeditedcnd")

            self.dlum = QtWidgets.QLabel(self.tabENTRER)
            self.dlum.setGeometry(QtCore.QRect(550, 243, 131, 16))
            self.dlum.setObjectName("dlum")
            self.dateEditdlum = QtWidgets.QLineEdit(self.tabENTRER)
            self.dateEditdlum.setInputMask("99-99-9999")
            self.dateEditdlum.setGeometry(QtCore.QRect(695, 235, 140, 30))
            self.dateEditdlum.setObjectName("dateEditdlum")





            self.nemuroDelottxt = QtWidgets.QLabel("N° DE LOT:",self.tabENTRER)
            self.nemuroDelottxt.setGeometry(QtCore.QRect(850,243,120,13))
            self.nemuroDelot=QtWidgets.QLineEdit(self.tabENTRER)
            self.nemuroDelot.setGeometry(QtCore.QRect(980,235,140,30))
            self.nemuroDelotschtxt = QtWidgets.QLabel("N°LOT SCH:", self.tabENTRER)
            self.nemuroDelotschtxt.setGeometry(QtCore.QRect(1130, 243, 120, 13))
            self.nemuroDelotsch = QtWidgets.QLineEdit(self.tabENTRER)
            self.nemuroDelotsch.setGeometry(QtCore.QRect(1200, 235, 140, 30))
            self.daterecolttxt = QtWidgets.QLabel("Date de récolte",self.tabENTRER)
            self.daterecolttxt.setGeometry(QtCore.QRect(20, 243, 81, 16))

            self.daterecolte = QtWidgets.QLineEdit(self.tabENTRER)
            self.daterecolte.setInputMask("99-99-9999")
            self.daterecolte.setGeometry(QtCore.QRect(115, 235, 140, 30))



            self.textEdit = QtWidgets.QTableWidget(self.tabENTRER)
            self.textEdit.setRowCount(0)
            self.textEdit.setColumnCount(14)
            self.textEdit.setColumnWidth(0, 3)
            self.textEdit.setColumnWidth(1, 10)
            self.textEdit.setColumnWidth(2, 20)
            self.textEdit.setColumnWidth(3, 130)
            self.textEdit.setColumnWidth(4, 130)
            self.textEdit.setColumnWidth(5, 130)
            self.textEdit.setColumnWidth(6, 130)
            self.textEdit.setColumnWidth(7, 130)
            self.textEdit.setColumnWidth(8, 130)
            self.textEdit.setColumnWidth(9, 130)
            self.textEdit.setColumnWidth(10, 130)
            self.textEdit.setColumnWidth(11,130)
            self.textEdit.setColumnWidth(12,100)
            self.textEdit.setColumnWidth(13, 100)
            self.textEdit.verticalHeader().setVisible(False)

            self.textEdit.setHorizontalHeaderLabels(("N°","N°LOT","N°LOT SCH","DATE:", "PRODUIT", "ORIGINE", "QUANTITE", "MAGASIN DE STOCK",
                                                     "SASHERIE CND", "FOURNISSEUR", "IMATRICUL","DATE DE RECOLTE", "DATE DE CND",
                                                     "DLUM" ))
            self.textEdit.setGeometry(QtCore.QRect(20, 320, 1310, 305))
            self.textEdit.setStyleSheet("background-color:rgb(255, 255, 255)")
            self.textEdit.setObjectName("textEdit")
            self.textEdit.setStyleSheet(" background-color: #FDFEFE")

            self.totalfont = QtGui.QFont()
            self.totalfont.setPointSize(9)
            self.totalfont.setBold(True)
            self.totalfont.bold()

            self.totaltxte = QtWidgets.QLabel("TOTAL")
            self.totaltxte.setFrameShape(QtWidgets.QFrame.Shape.WinPanel.VLine)
            self.totaltxte.setGeometry(1160, 573, 170, 30)
            self.totaltxte.setFont(self.totalfont)

            self.totalentree = QtWidgets.QDoubleSpinBox( readOnly=True)
            self.totalentree.setPrefix(" ")
            self.totalentree.setGeometry(1205, 578, 120, 20)
            self.totalentree.setRange(0, 9000000)

            self.totalriztxte = QtWidgets.QLabel("RIZ IMP")
            self.totalriztxte.setFrameShape(QtWidgets.QFrame.Shape.WinPanel.VLine)
            self.totalriztxte.setGeometry(20, 573, 162, 30)
            self.totalriztxte.setFont(self.totalfont)

            self.totalrizee = QtWidgets.QDoubleSpinBox( readOnly=True)
            self.totalrizee.setPrefix("TOTAL: ")
            self.totalrizee.setGeometry(54, 578, 140, 20)
            self.totalrizee.setRange(0, 9000000.00)

            self.totalpchtxte = QtWidgets.QLabel("P.CHICHES IMP")
            self.totalpchtxte.setFrameShape(QtWidgets.QFrame.Shape.WinPanel.VLine)
            self.totalpchtxte.setGeometry(187, 573, 193, 30)
            self.totalpchtxte.setFont(self.totalfont)

            self.totalpche = QtWidgets.QDoubleSpinBox( readOnly=True)
            self.totalpche.setPrefix("TOTAL: ")
            self.totalpche.setGeometry(257, 578, 120, 20)
            self.totalpche.setRange(0, 9000000)

            self.totalpchpltxte = QtWidgets.QLabel("P.CHICHES pl")
            self.totalpchpltxte.setFrameShape(QtWidgets.QFrame.Shape.WinPanel.VLine)
            self.totalpchpltxte.setGeometry(383, 573, 205, 30)
            self.totalpchpltxte.setFont(self.totalfont)

            self.totalpchple = QtWidgets.QDoubleSpinBox( readOnly=True)
            self.totalpchple.setPrefix("TOTAL: ")
            self.totalpchple.setGeometry(465, 578, 120, 20)
            self.totalpchple.setRange(0, 9000000)

            self.totallnpltxte = QtWidgets.QLabel("LENTILLE PL")
            self.totallnpltxte.setFrameShape(QtWidgets.QFrame.Shape.WinPanel.VLine)
            self.totallnpltxte.setGeometry(591, 573, 197, 30)
            self.totallnpltxte.setFont(self.totalfont)

            self.totallnple = QtWidgets.QDoubleSpinBox( readOnly=True)
            self.totallnple.setPrefix("TOTAL: ")
            self.totallnple.setGeometry(664, 578, 120, 20)
            self.totallnple.setRange(0, 9000000)

            self.totallntxte = QtWidgets.QLabel("LENTILLE IMP")
            self.totallntxte.setFrameShape(QtWidgets.QFrame.Shape.WinPanel.VLine)
            self.totallntxte.setGeometry(790, 573, 182, 30)
            self.totallntxte.setFont(self.totalfont)

            self.totallne = QtWidgets.QDoubleSpinBox( readOnly=True)
            self.totallne.setPrefix("TOTAL: ")
            self.totallne.setGeometry(850, 578, 120, 20)
            self.totallne.setRange(0, 9000000)

            self.totalhrtxte = QtWidgets.QLabel("HARICOT BLANC IMP")
            self.totalhrtxte.setFrameShape(QtWidgets.QFrame.Shape.WinPanel.VLine)
            self.totalhrtxte.setGeometry(974, 573, 185, 30)
            self.totalhrtxte.setFont(self.totalfont)

            self.totalhre = QtWidgets.QDoubleSpinBox( readOnly=True)
            self.totalhre.setPrefix("TOTAL: ")
            self.totalhre.setGeometry(1035, 578, 120, 20)
            self.totalhre.setRange(0, 9000000)

            self.filtertxt = QtWidgets.QLabel("Sélectionnez la date ", self.tabENTRER)
            self.filtertxt.setFrameShape(QtWidgets.QFrame.Shape.WinPanel.VLine)
            self.filtertxt.setGeometry(QtCore.QRect(730, 272, 590, 45))

            self.datefilter = QtWidgets.QLineEdit(self.tabENTRER)
            self.datefilter.setGeometry(QtCore.QRect(855, 275, 130, 38))
            self.datefilter.setInputMask('99-99-9999')

            self.produitfl = QtWidgets.QComboBox(self.tabENTRER)
            self.produitfl.setGeometry(QtCore.QRect(1060, 275, 130, 38))
            self.produitfl.addItem(" ")
            self.produitfl.addItem("POIS CHICHES PL")
            self.produitfl.addItem("POIS CHICHES IMP 12 mm")
            self.produitfl.addItem("POIS CHICHES IMP 09 mm")
            self.produitfl.addItem("POIS CHICHES IMP 08 mm")
            self.produitfl.addItem("POIS CHICHES IMP 06 mm")
            self.produitfl.addItem("HARICOT BLANC IMP")
            self.produitfl.addItem("HARICOT LSB")
            self.produitfl.addItem("LENTILLE PL")
            self.produitfl.addItem("LENTILLE IMP 05-07 mm")
            self.produitfl.addItem("LENTILLE IMP vert")
            self.produitfl.addItem("LENTILLE IMP ROUGE")
            self.produitfl.addItem("RIZ IMP")
            self.produitfl.addItem("RIZ ETUVE")

            self.btnfilter = QtWidgets.QPushButton("ok", self.tabENTRER, clicked=lambda: self.impot_filter())
            self.btnfilter.setGeometry((QtCore.QRect(1200, 275, 40, 40)))

            self.btnfcncl = QtWidgets.QPushButton("exit", self.tabENTRER, clicked=lambda: self.impot_all())
            self.btnfcncl.setGeometry((QtCore.QRect(1260, 275, 40, 40)))

            self.datafiltertxt = QtWidgets.QLabel("produit:", self.tabENTRER)
            self.datafiltertxt.setGeometry(QtCore.QRect(1000, 285, 60, 20))

            self.filtertxt = QtWidgets.QLabel("Filtrage des données:",self.tabENTRER)
            self.filtertxt.setGeometry(QtCore.QRect(735, 270, 120, 10))



            self.tabWidget.addTab(self.tabENTRER,"")
            ####################################################tab sortie###################################################
            ##################################################################################################################
            ################################################################################################################

            self.tabSORTIE = QtWidgets.QWidget()
            self.tabSORTIE.setObjectName("tabSORTIE")

            font1 = QtGui.QFont()
            font1.setPointSize(14)
            font1.setBold(True)
            font1.bold()
            font1.setWeight(75)

            self.addbtn_1 = QtWidgets.QPushButton(self.tabSORTIE)
            self.addbtn_1.setGeometry(QtCore.QRect(20, 275, 100, 40))
            self.addbtn_1.setObjectName("addbtn")
            self.addbtn_1.clicked.connect(self.add_sortie)

            self.deletebtn_1 = QtWidgets.QPushButton(self.tabSORTIE, clicked=lambda: self.delete_out())
            self.deletebtn_1.setGeometry(QtCore.QRect(190, 275, 100, 40))
            self.deletebtn_1.setObjectName("deletebtn")

            self.printbtn_1 = QtWidgets.QPushButton(self.tabSORTIE,clicked=lambda :self.print_docx_sortie())
            self.printbtn_1.setGeometry(QtCore.QRect(360, 275, 100, 40))
            self.printbtn_1.setObjectName("printbtn")

            self.printbtn_factur = QtWidgets.QPushButton('Imprimer \nfacture',self.tabSORTIE,clicked=lambda :self.print_facture_sortie())
            self.printbtn_factur.setGeometry(QtCore.QRect(530, 275, 100, 40))


            #self.actualisation = QtWidgets.QPushButton("Act", self.tabSORTIE, clicked=lambda: self.impot_all_sortie())
            #self.actualisation.setGeometry(QtCore.QRect(1280, 5, 20, 30))

            self.sortielabel = QtWidgets.QLabel("<h3>SORTIE<h3/>", self.tabSORTIE)
            self.sortielabel.setGeometry(650, 75, 150, 40)
            self.sortielabel.setFont(font1)

            self.cclstxt_1 = QtWidgets.QLabel(self.tabSORTIE)
            self.cclstxt_1.setGeometry(QtCore.QRect(430, 0, 500, 41))
            self.cclstxt_1.setFont(font1)
            self.cclstxt_1.setMouseTracking(False)
            self.cclstxt_1.setLayoutDirection(QtCore.Qt.LayoutDirection.LeftToRight)
            self.cclstxt_1.setAutoFillBackground(False)
            self.cclstxt_1.setFrameShape(QtWidgets.QFrame.Shape.WinPanel)
            self.cclstxt_1.setLineWidth(0)
            self.cclstxt_1.setMidLineWidth(0)
            self.cclstxt_1.setTextFormat(QtCore.Qt.TextFormat.AutoText)
            self.cclstxt_1.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
            self.cclstxt_1.setObjectName("cclstxt")
            self.titletxt_1 = QtWidgets.QLabel(self.tabSORTIE)
            self.titletxt_1.setGeometry(QtCore.QRect(540, 50, 350, 30))

            self.titletxt_1.setFont(font)
            self.titletxt_1.setAlignment(
                QtCore.Qt.AlignmentFlag.AlignLeading | QtCore.Qt.AlignmentFlag.AlignLeft | QtCore.Qt.AlignmentFlag.AlignVCenter)
            self.titletxt_1.setObjectName("titletxt")
            #############################################
            self.txtdate_1 = QtWidgets.QLabel(self.tabSORTIE)
            self.txtdate_1.setGeometry(QtCore.QRect(20, 80, 120, 16))
            self.txtdate_1.setObjectName("txtdate")
            ####################################################
            self.dateedite_1 = QtWidgets.QLineEdit(self.tabSORTIE)
            self.dateedite_1.setInputMask("99-99-9999")
            self.dateedite_1.setGeometry(QtCore.QRect(140, 77, 140, 30))
            self.dateedite_1.setObjectName("dateedite")
            ########################################################
            self.produittxt_1 = QtWidgets.QLabel(self.tabSORTIE)
            self.produittxt_1.setGeometry(QtCore.QRect(20, 132, 61, 16))
            self.produittxt_1.setObjectName("produittxt")
            self.produitcombo_1 = QtWidgets.QComboBox(self.tabSORTIE)
            self.produitcombo_1.setGeometry(QtCore.QRect(140, 130, 140, 30))
            self.produitcombo_1.setObjectName("produitcombo")
            self.produitcombo_1.addItem("")
            self.produitcombo_1.addItem("POIS CHICHES PL")
            self.produitcombo_1.addItem("POIS CHICHES IMP 12 mm")
            self.produitcombo_1.addItem("POIS CHICHES IMP 09 mm")
            self.produitcombo_1.addItem("POIS CHICHES IMP 08 mm")
            self.produitcombo_1.addItem("POIS CHICHES IMP 06 mm")
            self.produitcombo_1.addItem("HARICOT BLANC IMP")
            self.produitcombo_1.addItem("HARICOT LSB")
            self.produitcombo_1.addItem("LENTILLE PL")
            self.produitcombo_1.addItem("LENTILLE IMP 05-07 mm")
            self.produitcombo_1.addItem("LENTILLE IMP vert")
            self.produitcombo_1.addItem("LENTILLE IMP ROUGE")
            self.produitcombo_1.addItem("RIZ IMP")
            self.produitcombo_1.addItem("RIZ ETUVE")
            #####################################################################################
            self.origintxt_1 = QtWidgets.QLabel(self.tabSORTIE)
            self.origintxt_1.setGeometry(QtCore.QRect(300, 133, 121, 16))
            self.origintxt_1.setObjectName("origintxt")
            ###################################################################
            self.originel_1 = QtWidgets.QComboBox(self.tabSORTIE)
            all_countries = [
            " ","Algeria","Argentina",  "Australia",  "Azerbaijan", "Bangladesh", "Brazil", "Canada", "China", "Colombia",
            "Costa Rica",
            "Ecuador",
            "Egypt",
            "India", "Indonesia",
            "Italy", "Japan", "Kazakhstan",
            "Kuwait", "Kyrgyzstan",
            "Macedonia",  "Malaysia",
            "Mauritius", "Mexico",
            "Pakistan", "Paraguay", "Peru", "Philippines",
            "Poland", "Portugal",  "Romania", "Russian ",
            "Saudi Arabia",
            "Singapore",
            "South Africa", "Spain", "Sri Lanka", "Sudan",
            "Swaziland", "Sweden", "Switzerland",  "Taiwan", "Tajikistan",
            "Tanzania", "Thailand", "Tunisia", "Turkey",
            "Turkmenistan",  "Ukraine",
            "United Kingdom", "United States", "Uruguay", "Uzbekistan",
             "Venezuela", "Vietnam", "Yemen", "Zambia", "Zimbabwe"
        ]
            self.originel_1.setGeometry(QtCore.QRect(440, 130, 140, 30))
            self.originel_1.addItems(all_countries)
            ###############################################################################
            self.quatitelineEdit_1 = QtWidgets.QDoubleSpinBox(self.tabSORTIE)
            self.quatitelineEdit_1.setSpecialValueText(' ')
            self.quatitelineEdit_1.setRange(0, 1000000)
            self.quatitelineEdit_1.setGeometry(QtCore.QRect(695, 130, 140, 30))
            self.quatitelineEdit_1.setObjectName("quatitelineEdit")
            ################################################################################

            self.txtquantite_1 = QtWidgets.QLabel(self.tabSORTIE)
            self.txtquantite_1.setGeometry(QtCore.QRect(620, 133, 51, 16))
            self.txtquantite_1.setObjectName("txtquantite")
            #############################################################################
            self.labelmagasin_1 = QtWidgets.QLabel(self.tabSORTIE)
            self.labelmagasin_1.setGeometry(QtCore.QRect(870, 133, 116, 16))
            self.magasincombo_1 = QtWidgets.QComboBox(self.tabSORTIE,editable=True)
            self.magasincombo_1.setGeometry(QtCore.QRect(965, 130, 140, 30))
            self.magasincombo_1.setObjectName("magasincombo")
            self.magasincombo_1.addItem("")
            self.magasincombo_1.addItem("DOCK CENTRAL")
            self.magasincombo_1.addItem("MAGASIN MESRA")
            self.magasincombo_1.addItem("MARCHE GROS BELACEL")
            self.magasincombo_1.addItem("N.S.MENDES")
            self.magasincombo_1.addItem("MAGASIN ZEMOURA")
            self.magasincombo_1.addItem("POINT DE VENTE MOBILE")
            ###########################################################################
            self.labelmagasins = QtWidgets.QLabel("MAGASIN DE STOCK",self.tabSORTIE)
            self.labelmagasins.setGeometry(QtCore.QRect(20, 188, 110, 16))
            self.magasincombos = QtWidgets.QComboBox(self.tabSORTIE)
            self.magasincombos.setGeometry(QtCore.QRect(140, 185, 140, 30))
            self.magasincombos.setObjectName("magasincombo")
            self.magasincombos.addItem("")
            self.magasincombos.addItem("DOCK CENTRAL")
            self.magasincombos.addItem("MAGASIN MESRA")
            self.magasincombos.addItem("N.S.MENDES")

            #########################################################################
            self.CNDTXT_1 = QtWidgets.QLabel(self.tabSORTIE)
            self.CNDTXT_1.setGeometry(QtCore.QRect( 870, 188, 116, 16))
            self.CNDTXT_1.setObjectName("CNDTXT")
            self.cndcombo_1 = QtWidgets.QComboBox(self.tabSORTIE)
            self.cndcombo_1.setGeometry(QtCore.QRect(965, 185, 140, 30))
            self.cndcombo_1.setObjectName("cndcombo")
            self.cndcombo_1.addItem("")
            self.cndcombo_1.addItem("SAC 50KG")
            self.cndcombo_1.addItem("SAC 25KG")
            self.cndcombo_1.addItem("BIG BAG")
            ########################################################################
            self.TRANSPRTTXT_1 = QtWidgets.QLabel(self.tabSORTIE)
            self.TRANSPRTTXT_1.setGeometry(QtCore.QRect(300, 188, 131, 16))
            self.TRANSPRTTXT_1.setObjectName("TRANSPRTTXT")
            self.TRANSPORTcombo_1 = QtWidgets.QComboBox(self.tabSORTIE,editable=True)
            self.TRANSPORTcombo_1.setGeometry(QtCore.QRect(440, 185, 140, 30))
            self.TRANSPORTcombo_1.setObjectName("TRANSPORTcombo")
            self.TRANSPORTcombo_1.addItem("")
            ########################################################################
            self.label_2_1 = QtWidgets.QLabel(self.tabSORTIE)
            self.label_2_1.setGeometry(QtCore.QRect(620, 188, 81, 16))
            self.label_2_1.setObjectName("label_2")
            self.imatricullineeditr_1 = QtWidgets.QLineEdit(self.tabSORTIE)
            self.imatricullineeditr_1.setInputMask("99999-999-99")
            self.imatricullineeditr_1.setGeometry(QtCore.QRect(695, 185, 140, 30))

            #########################################################################
            self.dateeditedcnd_1 = QtWidgets.QLineEdit(self.tabSORTIE)
            self.dateeditedcnd_1.setInputMask("99-99-9999")
            self.dateeditedcnd_1.setGeometry(QtCore.QRect(440, 235, 140, 30))
            self.dateeditedcnd_1.setObjectName("dateeditedcnd")

            self.dateEditdlum_1 = QtWidgets.QLineEdit(self.tabSORTIE)
            self.dateEditdlum_1.setInputMask("99-99-9999")
            self.dateEditdlum_1.setGeometry(QtCore.QRect(695, 235, 140, 30))
            self.dateEditdlum_1.setObjectName("dateEditdlum")

            self.DATEDECNDTXT_1 = QtWidgets.QLabel(self.tabSORTIE)
            self.DATEDECNDTXT_1.setGeometry(QtCore.QRect(300, 243, 131, 16))
            self.DATEDECNDTXT_1.setObjectName("DATEDECNDTXT")

            self.dlum_1 = QtWidgets.QLabel(self.tabSORTIE)
            self.dlum_1.setGeometry(QtCore.QRect(620, 243, 47, 13))
            self.dlum_1.setObjectName("dlum")

            self.nemuroDelottxt = QtWidgets.QLabel("N° DE LOT:", self.tabSORTIE)
            self.nemuroDelottxt.setGeometry(QtCore.QRect(870, 243, 120, 13))
            self.nemuroDelots = QtWidgets.QLineEdit(self.tabSORTIE)
            self.nemuroDelots.setGeometry(QtCore.QRect(965, 235, 140, 30))

            self.nemuroDelotschtxt = QtWidgets.QLabel("N°LOT SCH :", self.tabSORTIE)
            self.nemuroDelotschtxt.setGeometry(QtCore.QRect(1128, 243, 120, 13))
            self.nemuroDelotssch = QtWidgets.QLineEdit(self.tabSORTIE)
            self.nemuroDelotssch.setGeometry(QtCore.QRect(1200, 235, 140, 30))

            self.daterecolttxtso = QtWidgets.QLabel("Date de récolte",self.tabSORTIE)
            self.daterecolttxtso.setGeometry(QtCore.QRect(20, 243, 81, 16))

            self.daterecolteso = QtWidgets.QLineEdit(self.tabSORTIE)
            self.daterecolteso.setInputMask("99-99-9999")
            self.daterecolteso.setGeometry(QtCore.QRect(140, 235, 140, 30))

            self.fourniseurcombostxt=QtWidgets.QLabel("EXPEDITEUR:",self.tabSORTIE)
            self.fourniseurcombostxt.setGeometry(QtCore.QRect(1128, 130, 140, 30))
            self.fourniseurcombos = QtWidgets.QComboBox(self.tabSORTIE)
            self.fourniseurcombos.setGeometry(QtCore.QRect(1200, 130, 140, 30))

            self.fourniseurcombos.addItem("")
            self.fourniseurcombos.addItem("UCA ALGER")
            self.fourniseurcombos.addItem("UCA ORAN")

            self.textEdit_1 = QTableWidget(self.tabSORTIE)
            self.textEdit_1.setRowCount(0)
            self.textEdit_1.setColumnCount(14)

            self.textEdit_1.setColumnWidth(1, 50)
            self.textEdit_1.setColumnWidth(2, 10)
            self.textEdit_1.setColumnWidth(3, 10)
            self.textEdit_1.setColumnWidth(4, 0)
            self.textEdit_1.setColumnWidth(5, 0)
            self.textEdit_1.setColumnWidth(6, 0)
            self.textEdit_1.setColumnWidth(7, 0)
            self.textEdit_1.setColumnWidth(8, 0)
            self.textEdit_1.setColumnWidth(9, 0)
            self.textEdit_1.setColumnWidth(10, 0)
            self.textEdit_1.setColumnWidth(11, 0)
            self.textEdit_1.setColumnWidth(12, 0)
            self.textEdit_1.setColumnWidth(13, 0)
            self.textEdit_1.verticalHeader().setVisible(False)

            self.textEdit_1.setHorizontalHeaderLabels(("N°","N°LOT","N°LOT SCH","DATE","MAGASIN DE STOCK", "PRODUIT", "ORIGINE", "QUANTITE (Qx)", "ACHTEUR",
                                                       "SASHERIE DE CND", "IMATRICULE","DATE DE RECOLTE", "DATE DE CND",
                                                       "DLUM"))
            self.textEdit_1.setGeometry(QtCore.QRect(20, 320, 1310, 305))
            self.textEdit_1.setObjectName("textEdit")
            self.textEdit_1.setStyleSheet(" background-color: #ffffff;")

            self.totaltxt = QtWidgets.QLabel("TOTAL")
            self.totaltxt.setFrameShape(QtWidgets.QFrame.Shape.WinPanel.VLine)
            self.totaltxt.setGeometry(1148, 573, 183, 50)
            self.totaltxt.setFont(self.totalfont)

            self.rest = QtWidgets.QDoubleSpinBox( readOnly=True)
            self.rest.setPrefix("REST: ")
            self.rest.setGeometry(1195, 599, 133, 20)
            self.rest.setRange(0, 900000)



            self.totalriztxt = QtWidgets.QLabel("RIZ IMP")
            self.totalriztxt.setFrameShape(QtWidgets.QFrame.Shape.WinPanel.VLine)
            self.totalriztxt.setGeometry(20, 573, 154, 50)
            self.totalriztxt.setFont(self.totalfont)

            self.resteriz = QtWidgets.QDoubleSpinBox( readOnly=True)
            self.resteriz.setPrefix("REST: ")
            self.resteriz.setGeometry(51, 599, 120, 20)
            self.resteriz.setRange(0, 900000)



            self.totalpchtxt = QtWidgets.QLabel("P.CHICHES IMP")
            self.totalpchtxt.setFrameShape(QtWidgets.QFrame.Shape.WinPanel.VLine)
            self.totalpchtxt.setGeometry(175, 573, 193, 50)
            self.totalpchtxt.setFont(self.totalfont)

            self.restpch = QtWidgets.QDoubleSpinBox( readOnly=True)
            self.restpch.setPrefix("REST: ")
            self.restpch.setGeometry(245, 599, 120, 20)
            self.restpch.setRange(0, 900000)



            self.totalpchpltxt = QtWidgets.QLabel("P.CHICHES pl")
            self.totalpchpltxt.setFrameShape(QtWidgets.QFrame.Shape.WinPanel.VLine)
            self.totalpchpltxt.setGeometry(370, 573, 205, 50)
            self.totalpchpltxt.setFont(self.totalfont)

            self.restpchpl = QtWidgets.QDoubleSpinBox( readOnly=True)
            self.restpchpl.setPrefix("REST: ")
            self.restpchpl.setGeometry(450, 599, 120, 20)
            self.restpchpl.setRange(0, 900000)



            self.totallnpltxt = QtWidgets.QLabel("LENTILLE PL")
            self.totallnpltxt.setFrameShape(QtWidgets.QFrame.Shape.WinPanel.VLine)
            self.totallnpltxt.setGeometry(578, 573, 200, 50)
            self.totalpchpltxt.setFont(self.totalfont)



            self.restlnpl = QtWidgets.QDoubleSpinBox( readOnly=True)
            self.restlnpl.setPrefix("REST: ")
            self.restlnpl.setGeometry(655, 599, 120, 20)
            self.restlnpl.setRange(0, 900000)

            self.totallntxt = QtWidgets.QLabel("LENTILLE IMP")
            self.totallntxt.setFrameShape(QtWidgets.QFrame.Shape.WinPanel.VLine)
            self.totallntxt.setGeometry(780, 573, 188, 50)
            self.totallntxt.setFont(self.totalfont)



            self.restln = QtWidgets.QDoubleSpinBox( readOnly=True)
            self.restln.setPrefix("REST: ")
            self.restln.setGeometry(845, 599, 120, 20)
            self.restln.setRange(0, 900000)

            self.totalhrtxt = QtWidgets.QLabel("HARICOT BLANC IMP")
            self.totalhrtxt.setFrameShape(QtWidgets.QFrame.Shape.WinPanel.VLine)
            self.totalhrtxt.setGeometry(970, 573, 175, 50)
            self.totalhrtxt.setFont(self.totalfont)



            self.resthr = QtWidgets.QDoubleSpinBox( readOnly=True)
            self.resthr.setPrefix("REST: ")
            self.resthr.setGeometry(1030, 599, 110, 20)
            self.resthr.setRange(0, 900000)

            self.filtertxt = QtWidgets.QLabel("Sélectionnez la date ", self.tabSORTIE)
            self.filtertxt.setFrameShape(QtWidgets.QFrame.Shape.WinPanel.VLine)
            self.filtertxt.setGeometry(QtCore.QRect(730, 272, 590, 45))

            self.filtertxts = QtWidgets.QLabel("Sélectionnez la date ")
            self.filtertxts.setFrameShape(QtWidgets.QFrame.Shape.WinPanel.VLine)
            self.filtertxts.setGeometry(QtCore.QRect(730, 268, 590, 51))

            self.datefilters = QtWidgets.QLineEdit(self.tabSORTIE)
            self.datefilters.setGeometry(QtCore.QRect(855, 275, 130, 38))
            self.datefilters.setInputMask('99-99-9999')

            self.produitfls = QtWidgets.QComboBox(self.tabSORTIE)
            self.produitfls.setGeometry(QtCore.QRect(1060, 275, 130, 38))
            self.produitfls.addItem(" ")
            self.produitfls.addItem("POIS CHICHES PL")
            self.produitfls.addItem("POIS CHICHES IMP 12 mm")
            self.produitfls.addItem("POIS CHICHES IMP 09 mm")
            self.produitfls.addItem("POIS CHICHES IMP 08 mm")
            self.produitfls.addItem("POIS CHICHES IMP 06 mm")
            self.produitfls.addItem("HARICOT BLANC IMP")
            self.produitfls.addItem("POIS CHICHES PL")
            self.produitfls.addItem("HARICOT LSB")
            self.produitfls.addItem("LENTILLE PL")
            self.produitfls.addItem("LENTILLE IMP 05-07 mm")
            self.produitfls.addItem("LENTILLE IMP vert")
            self.produitfls.addItem("RIZ IMP")
            self.produitfls.addItem("RIZ ETUVE")



            self.btnfilters = QtWidgets.QPushButton("ok", self.tabSORTIE, clicked=lambda: self.impot_filter_sortie())
            self.btnfilters.setGeometry((QtCore.QRect(1200, 275, 40, 40)))

            self.btnfcncls = QtWidgets.QPushButton("exit", self.tabSORTIE, clicked=lambda: self.impot_all_sortie())
            self.btnfcncls.setGeometry((QtCore.QRect(1260, 275, 40, 40)))

            self.datafiltertxts = QtWidgets.QLabel("produit:",self.tabSORTIE)
            self.datafiltertxts.setGeometry(QtCore.QRect(1000, 285, 60, 20))

            self.filtertxts = QtWidgets.QLabel("Filtrage des données:", self.tabSORTIE)
            self.filtertxts.setGeometry(QtCore.QRect(735, 270, 115, 10))



            self.tabWidget.addTab(self.tabSORTIE, "")
            self.verticalLayout.addWidget(self.tabWidget)
            MainWindow.setCentralWidget(self.centralwidget)
            self.statusbar = QtWidgets.QStatusBar(MainWindow)
            self.statusbar.setObjectName("statusbar")
            MainWindow.setStatusBar(self.statusbar)

            self.retranslateUi(MainWindow)
            self.tabWidget.setCurrentIndex(1)
            QtCore.QMetaObject.connectSlotsByName(MainWindow)

            self.impot_all()
            self.impot_all_sortie()







####################################################################################################################
        ##########################################Stock############################################
            self.tabStock=QtWidgets.QWidget()
            font2 = QtGui.QFont()
            font2.setPointSize(10)
            font2.setBold(False)


            self.addd_buttons = QtWidgets.QPushButton("STOCK FINAL", self.tabStock, clicked=lambda: self.stock_final_magasin() and self.impot_filter_stock())
            self.addd_buttons.setGeometry(680, 250, 160, 40)

            self.printbts = QtWidgets.QPushButton("IMPRIMER SITUATION \nDES STOCK",self.tabStock, clicked=lambda: self.docxStock())
            self.printbts.setGeometry(QtCore.QRect(500, 250, 160, 40))


            fichetraitementxts = QtWidgets.QLabel("ÉTAT DES STOCKS DE LÉGUMES SECS", self.tabStock)
            fichetraitementxts.setGeometry(490, 5, 430, 60)
            fichetraitementxts.setStyleSheet("font: bold 24px;")

            dateDeFacturetxs = QtWidgets.QLabel("Date de stock:")
            dateDeFacturetxs.setGeometry(5, 60, 160, 30)

            self.Unitetxts = QtWidgets.QLabel("UNITÉ DE STOCK :", self.tabStock)
            self.Unitetxts.setGeometry(340, 115, 300, 30)
            self.Unitetxts.setFont(font2)

            self.Unites = QtWidgets.QComboBox(self.tabStock, editable=False)
            self.Unites.addItem("")
            self.Unites.addItem("DOCK CENTRAL")
            self.Unites.addItem("MAGASIN MESRA")
            self.Unites.addItem("MARCHE GROS BELACEL")
            self.Unites.setGeometry(460, 115, 192, 35)


            self.Produittxts = QtWidgets.QLabel("PRODUIT :", self.tabStock)
            self.Produittxts.setGeometry(10, 115, 300, 30)
            self.Produittxts.setFont(font2)

            self.Produits = QtWidgets.QComboBox(self.tabStock)
            self.Produits.setGeometry(90, 115, 192, 35)
            self.Produits.addItem("")

            self.Produits.addItem("POIS CHICHES PL")
            self.Produits.addItem("POIS CHICHES IMP 12 mm")
            self.Produits.addItem("POIS CHICHES IMP 06 mm")
            self.Produits.addItem("POIS CHICHES IMP 08 mm")
            self.Produits.addItem("POIS CHICHES IMP 09 mm")
            self.Produits.addItem("HARICOT BLANC IMP")
            self.Produits.addItem("HARICOT LSB")
            self.Produits.addItem("LENTILLE PL")
            self.Produits.addItem("LENTILLE IMP 05-07 mm")
            self.Produits.addItem("LENTILLE IMP vert")
            self.Produits.addItem("LENTILLE IMP ROUGE")
            self.Produits.addItem("RIZ IMP")
            self.Produits.addItem("RIZ ETUVE")




            self.OrigineDuProduittxts = QtWidgets.QLabel("ORIGINE DE PRODUIT :", self.tabStock)
            self.OrigineDuProduittxts.setGeometry(705, 115, 300, 30)
            self.OrigineDuProduittxts.setFont(font2)
            all_countries =[
            "","Algeria","Argentina","Australia","Azerbaijan","Bangladesh","Brazil","Canada","China","Colombia",
            "Costa Rica",
            "Ecuador",
            "Egypt",
            "India", "Indonesia",
            "Italy", "Japan", "Kazakhstan",
            "Kuwait", "Kyrgyzstan",
            "Macedonia",  "Malaysia",
            "Mauritius", "Mexico",
            "Pakistan", "Paraguay", "Peru", "Philippines",
            "Poland", "Portugal",  "Romania", "Russian ",
            "Saudi Arabia",
            "Singapore",
            "South Africa", "Spain", "Sri Lanka", "Sudan",
            "Swaziland", "Sweden", "Switzerland",  "Taiwan", "Tajikistan",
            "Tanzania", "Thailand", "Tunisia", "Turkey",
            "Turkmenistan",  "Ukraine",
            "United Kingdom", "United States", "Uruguay", "Uzbekistan",
             "Venezuela", "Vietnam", "Yemen", "Zambia", "Zimbabwe"
            ]
            self.OrigineDuProduits = QtWidgets.QComboBox(self.tabStock)
            self.OrigineDuProduits.addItems(all_countries)
            self.OrigineDuProduits.setGeometry(850, 115, 170, 35)





            self.expediteurtxt=QtWidgets.QLabel("EXPÉDITEUR :",self.tabStock)
            self.expediteurtxt.setGeometry(QtCore.QRect(1080,115,150,30))
            self.expediteurtxt.setFont(font2)
            self.expediteur = QtWidgets.QComboBox( self.tabStock)
            self.expediteur.setGeometry(QtCore.QRect(1170, 115, 170, 35))
            self.expediteur.addItem("")
            self.expediteur.addItem("UCA ALGER")
            self.expediteur.addItem("UCA ORAN")
            self.expediteur.addItem("UCA MOSTAGANEM")
            self.expediteur.addItem("")
            self.expediteur.addItem("")
            self.expediteur.addItem("")

            datetext = QLabel("   DATE :", self.tabStock)
            datetext.setFrameShape(QtWidgets.QFrame.Shape.WinPanel.VLine)
            datetext.setGeometry(QtCore.QRect(10, 165, 650, 73))
            datetextss = QLabel("PERIODE :", self.tabStock)
            datetextss.setGeometry(QtCore.QRect(25, 155, 55, 18))
            datetexts = QLabel("au :", self.tabStock)
            datetexts.setGeometry(QtCore.QRect(380, 185, 80, 35))
            self.date_edits = QtWidgets.QLineEdit(self.tabStock)
            self.date_edits.setInputMask("99-99-9999")
            self.date_edits.setGeometry(QtCore.QRect(90,185,192,35))

            self.date_edit2 = QtWidgets.QLineEdit(self.tabStock)
            self.date_edit2.setInputMask("99-99-9999")
            self.date_edit2.setGeometry(QtCore.QRect(460,185,192,35))




            self.txtfacturetxts = QtWidgets.QLabel(" ", self.tabStock)
            self.txtfacturetxts.setGeometry(400, 400, 100, 20)
            self.txtfactures = QTableWidget(self.tabStock)
            self.txtfactures.setGeometry(10, 300, 1330, 250)
            self.txtfactures.setRowCount(0)
            self.txtfactures.setColumnCount(9)
            self.txtfactures.setColumnWidth(0, 140)
            self.txtfactures.setColumnWidth(1, 170)
            self.txtfactures.setColumnWidth(2, 160)
            self.txtfactures.setColumnWidth(3, 175)
            self.txtfactures.setColumnWidth(4, 160)
            self.txtfactures.setColumnWidth(5, 175)
            self.txtfactures.setColumnWidth(6, 175)
            self.txtfactures.setColumnWidth(7, 175)
            self.txtfactures.verticalHeader().setVisible(False)
            self.txtfactures.setHorizontalHeaderLabels(
                ("UNITÉ", "PRODUIT", "QUANTITÉ", "ORIGINE DE PRODUIT", "DLUM","DATE DE CND","N°LOT",'N°LOT SCH',"EXPEDITEUR"))


            self.stockFinaltxt=QtWidgets.QLabel("STOCK FINAL",self.tabStock)
            self.stockFinaltxt.setStyleSheet("font: bold 24px;")
            self.stockFinaltxt.setGeometry(QtCore.QRect(835,570,180,40))
            self.stockFinal=QtWidgets.QLineEdit(self.tabStock)
            self.stockFinal.setGeometry(QtCore.QRect(1007,570,333,40))
           ##########################difinition of value
            self.selected_produit = None
            self.selected_magasin = None
            self.selected_dates = None
            self.selected_datesAU = None



            self.tabWidget.addTab(self.tabStock,"STOCK")









        def retranslateUi(self, MainWindow):
            _translate = QtCore.QCoreApplication.translate
            MainWindow.setWindowTitle(_translate("MainWindow", "MainWindow"))
            self.addbtn.setText(_translate("MainWindow", "Ajouter"))
            self.deletebtn.setText(_translate("MainWindow", "suprime"))
            self.printbtn.setText(_translate("MainWindow", "Imprimer"))
            self.cclstxt.setText(_translate("MainWindow", "<h2>CCLS RELIZANE SERVICE QUALITE<h2/>"))
            self.titletxt.setText(_translate("MainWindow", "<h3>GESTION STOCk LEGUMES SECS<h3/>"))
            self.txtdate.setText(_translate("MainWindow", "DATE DE ENTRER:"))
            self.produittxt.setText(_translate("MainWindow", "PRODUIT:"))
            self.origintxt.setText(_translate("MainWindow", "ORIGINE DE PRODUIT:"))
            self.txtquantite.setText(_translate("MainWindow", "QUATITE:"))
            self.labelmagasin.setText(_translate("MainWindow", "MAGASAIN DE STOCK::"))
            self.fornisseurtxt.setText(_translate("MainWindow", "EXPEDITEUR:"))
            self.CNDTXT.setText(_translate("MainWindow", "SACHERIE CND :"))
            self.TRANSPRTTXT.setText(_translate("MainWindow", "MOIYEN DE TRANSPORT :"))
            self.label_2.setText(_translate("MainWindow", "IMATRICULE:"))
            self.DATEDECNDTXT.setText(_translate("MainWindow", "DATE DE CND :"))
            self.dlum.setText(_translate("MainWindow", "DLUM :"))

            self.tabWidget.setTabText(self.tabWidget.indexOf(self.tabENTRER), _translate("MainWindow", "ENTRER"))
            self.addbtn_1.setText(_translate("MainWindow", "Ajouter"))
            self.deletebtn_1.setText(_translate("MainWindow", "suprime"))
            self.printbtn_1.setText(_translate("MainWindow", "Imprimer"))
            self.cclstxt_1.setText(_translate("MainWindow", "<h2>CCLS RELIZANE SERVICE QUALITE<h2/>"))
            self.titletxt_1.setText(_translate("MainWindow", "<h3>GESTION STOCk LEGUMES SECS<h3/>"))
            self.txtdate_1.setText(_translate("MainWindow", "DATE DE SORTIE:"))
            self.produittxt_1.setText(_translate("MainWindow", "PRODUIT:"))
            self.origintxt_1.setText(_translate("MainWindow", "ORIGINE DE PRODUIT:"))
            self.txtquantite_1.setText(_translate("MainWindow", "QUATITE:"))
            self.labelmagasin_1.setText(_translate("MainWindow", "ACHTEUR"))
            self.CNDTXT_1.setText(_translate("MainWindow", "SACHERIE CND :"))
            self.TRANSPRTTXT_1.setText(_translate("MainWindow", "MOIYEN DE TRANSPORT :"))
            self.label_2_1.setText(_translate("MainWindow", "IMATRICULE:"))
            self.DATEDECNDTXT_1.setText(_translate("MainWindow", "DATE DE CND :"))
            self.dlum_1.setText(_translate("MainWindow", "DLUM :"))
            self.tabWidget.setTabText(self.tabWidget.indexOf(self.tabSORTIE), _translate("MainWindow", "SORTIE"))

        #####################################################################################################################################

        def add(self):
            try:
                dateentre = self.dateedite.text()
                nlot=self.nemuroDelot.text()
                nlotsch=self.nemuroDelotsch.text()
                produitt = self.produitcombo.currentText()
                country = self.originel.currentText()
                quantité = self.quatitelineEdit.value()
                lieu = self.magasincombo.currentText()
                sac = self.cndcombo.currentText()
                fourni = self.fourniseurcombo.currentText()
                matric = self.imatricullineeditr.text()
                datrecolte=self.daterecolte.text()
                cnddate = self.dateeditedcnd.text()
                datedlm = self.dateEditdlum.text()

                if produitt:
                    if produitt == 'RIZ IMP':
                        dattabase = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumsec_out')
                        curs = dattabase.cursor()
                        dattabasei = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumesec_entry')
                        cursi = dattabasei.cursor()
                        curs.execute("INSERT INTO outtable (quantitérizentré,le,produit,unité,origine,fournisseur)VALUES(%s,%s,%s,%s,%s,%s)",(quantité,dateentre,produitt,lieu,country,fourni))
                        cursi.execute(
                            "INSERT INTO entrytable (le,nlot,nlotsch,produit,origine,quantité,unité,sacherie,fournisseur,imatricule,date_recolte,date_CND,DLUM,quantitérizentrée)VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                            (dateentre,nlot,nlotsch,produitt,country,quantité, lieu, sac, fourni, matric,datrecolte, cnddate, datedlm,
                             quantité))
                        dattabase.commit()
                        dattabasei.commit()
                    if produitt == 'RIZ ETUVE':
                        dattabase = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumsec_out')
                        curs = dattabase.cursor()
                        dattabasei = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumesec_entry')
                        cursi = dattabasei.cursor()
                        curs.execute("INSERT INTO outtable (quantitérizetvntré,le,produit,unité,origine,fournisseur)VALUES(%s,%s,%s,%s,%s,%s)",(quantité,dateentre,produitt,lieu,country,fourni))
                        cursi.execute(
                            "INSERT INTO entrytable (le,nlot,nlotsch,produit,origine,quantité,unité,sacherie,fournisseur,imatricule,date_recolte,date_CND,DLUM,quantitérizetvntrée)VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                            (dateentre,nlot,nlotsch, produitt, country, quantité, lieu, sac, fourni, matric,datrecolte, cnddate, datedlm,
                             quantité))
                        dattabase.commit()
                        dattabasei.commit()
                    if produitt == 'LENTILLE IMP vert':
                        dattabase = mysql.connector.connect(
        host="localhost",
         user=user,
        password=password,database='datta_legumsec_out'
    )
                        curs = dattabase.cursor()
                        dattabasei = mysql.connector.connect(
        host="localhost",
         user=user,
        password=password,database='datta_legumesec_entry'
    )
                        cursi = dattabasei.cursor()
                        curs.execute("INSERT INTO outtable (quantitélenvertentré,le,produit,unité,origine,fournisseur)VALUES(%s,%s,%s,%s,%s,%s)",(quantité,dateentre,produitt,lieu,country,fourni))
                        cursi.execute(
                            "INSERT INTO entrytable (le,nlot,nlotsch,produit,origine,quantité,unité,sacherie,fournisseur,imatricule,date_recolte,date_CND,DLUM,quantitélenvertntrée)VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                            (dateentre,nlot,nlotsch, produitt, country, quantité, lieu, sac, fourni, matric,datrecolte, cnddate, datedlm,
                             quantité))
                        dattabase.commit()
                        dattabasei.commit()

                    if produitt == 'LENTILLE IMP 05-07 mm':
                        dattabase = mysql.connector.connect(
        host="localhost",
         user=user,
        password=password,database='datta_legumsec_out'
    )
                        curs = dattabase.cursor()
                        dattabasei = mysql.connector.connect(
        host="localhost",
         user=user,
        password=password,database='datta_legumesec_entry'
    )
                        cursi = dattabasei.cursor()
                        curs.execute("INSERT INTO outtable (quantitélenentré,le,produit,unité,origine,fournisseur)VALUES(%s,%s,%s,%s,%s,%s)",(quantité,dateentre,produitt,lieu,country,fourni))
                        cursi.execute(
                            "INSERT INTO entrytable (le,nlot,nlotsch,produit,origine,quantité,unité,sacherie,fournisseur,imatricule,date_recolte,date_CND,DLUM,quantitélenentrée)VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                            (dateentre,nlot,nlotsch, produitt, country, quantité, lieu, sac, fourni, matric,datrecolte, cnddate, datedlm,
                             quantité))
                        dattabase.commit()
                        dattabasei.commit()

                    if produitt == 'LENTILLE IMP ROUGE':
                        dattabase = mysql.connector.connect(
                            host="localhost",
                             user=user,
                            password=password, database='datta_legumsec_out'
                        )
                        curs = dattabase.cursor()
                        dattabasei = mysql.connector.connect(
                            host="localhost",
                             user=user,
                            password=password, database='datta_legumesec_entry'
                        )
                        cursi = dattabasei.cursor()
                        curs.execute(
                            "INSERT INTO outtable (quantitélenrougeentré,le,produit,unité,origine,fournisseur)VALUES(%s,%s,%s,%s,%s,%s)",
                            (quantité,dateentre, produitt, lieu, country, fourni))
                        cursi.execute(
                            "INSERT INTO entrytable (le,nlot,nlotsch,produit,origine,quantité,unité,sacherie,fournisseur,imatricule,date_recolte,date_CND,DLUM,quantitélenrougeentrée)VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                            (dateentre, nlot,nlotsch,produitt,country, quantité, lieu, sac, fourni, matric, datrecolte,
                             cnddate, datedlm,
                             quantité))
                        dattabase.commit()
                        dattabasei.commit()

                    if produitt == 'LENTILLE PL':
                        dattabase = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumsec_out')
                        dattabasei = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumesec_entry')
                        curs=dattabase.cursor()
                        cursi=dattabasei.cursor()
                        curs.execute("INSERT INTO outtable (quantitélenplentré,le,produit,unité,origine,fournisseur)VALUES(%s,%s,%s,%s,%s,%s)",(quantité,dateentre,produitt,lieu,country,fourni))
                        cursi.execute(
                            "INSERT INTO entrytable (le,nlot,nlotsch,produit,origine,quantité,unité,sacherie,fournisseur,imatricule,date_recolte,date_CND,DLUM,quantitélenplntrée)VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                            (dateentre,nlot,nlotsch, produitt, country, quantité, lieu, sac, fourni, matric,datrecolte, cnddate, datedlm,
                             quantité))
                        dattabase.commit()
                        dattabasei.commit()

                    if produitt == 'HARICOT BLANC IMP':
                        dattabase = mysql.connector.connect(
        host="localhost",
         user=user,
        password=password,database='datta_legumsec_out'
    )
                        curs = dattabase.cursor()
                        dattabasei = mysql.connector.connect(
        host="localhost",
         user=user,
        password=password,database='datta_legumesec_entry'
    )
                        cursi = dattabasei.cursor()
                        curs.execute("INSERT INTO outtable (quantitéharentré,le,produit,unité,origine,fournisseur)VALUES(%s,%s,%s,%s,%s,%s)",(quantité,dateentre,produitt,lieu,country,fourni))
                        cursi.execute(
                            "INSERT INTO entrytable (le,nlot,nlotsch,produit,origine,quantité,unité,sacherie,fournisseur,imatricule,date_recolte,date_CND,DLUM,quantitéharntrée)VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                            (dateentre,nlot,nlotsch, produitt, country, quantité, lieu, sac, fourni, matric,datrecolte ,cnddate, datedlm,
                             quantité))
                        dattabase.commit()
                        dattabasei.commit()

                    if produitt == 'HARICOT LSB':
                        dattabase = mysql.connector.connect(
        host="localhost",
         user=user,
        password=password,database='datta_legumsec_out'
    )
                        curs = dattabase.cursor()
                        dattabasei = mysql.connector.connect(
        host="localhost",
         user=user,
        password=password,database='datta_legumesec_entry'
    )
                        cursi = dattabasei.cursor()
                        curs.execute("INSERT INTO outtable (quantitéharlsbentré,le,produit,unité,origine,fournisseur)VALUES(%s,%s,%s,%s,%s,%s)",(quantité,dateentre,produitt,lieu,country,fourni))
                        cursi.execute(
                            "INSERT INTO entrytable (le,nlot,nlotsch,produit,origine,quantité,unité,sacherie,fournisseur,imatricule,date_recolte,date_CND,DLUM,quantitéharlsbntrée)VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                            (dateentre,nlot,nlotsch,produitt, country, quantité, lieu, sac, fourni, matric,datrecolte ,cnddate, datedlm,
                             quantité))
                        dattabase.commit()
                        dattabasei.commit()

                    if produitt == 'POIS CHICHES IMP 12 mm':
                        dattabase = mysql.connector.connect(
        host="localhost",
         user=user,
        password=password,database='datta_legumsec_out'
    )
                        curs = dattabase.cursor()
                        dattabasei = mysql.connector.connect(
        host="localhost",
         user=user,
        password=password,database='datta_legumesec_entry'
    )
                        cursi = dattabasei.cursor()
                        curs.execute("INSERT INTO outtable (quantitepchentré,le,produit,unité,origine,fournisseur)VALUES(%s,%s,%s,%s,%s,%s)",(quantité,dateentre,produitt,lieu,country,fourni))
                        cursi.execute(
                            "INSERT INTO entrytable (le,nlot,nlotsch,produit,origine,quantité,unité,sacherie,fournisseur,imatricule,date_recolte,date_CND,DLUM,quantitépchentrée)VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                            (dateentre,nlot,nlotsch,produitt,country,quantité, lieu, sac, fourni, matric, datrecolte,cnddate, datedlm,
                             quantité))
                        dattabase.commit()
                        dattabasei.commit()
                    if produitt == 'POIS CHICHES IMP 09 mm':
                        dattabase = mysql.connector.connect(
        host="localhost",
         user=user,
        password=password,database='datta_legumsec_out'
    )
                        curs = dattabase.cursor()
                        dattabasei = mysql.connector.connect(
        host="localhost",
         user=user,
        password=password,database='datta_legumesec_entry'
    )
                        cursi = dattabasei.cursor()
                        curs.execute("INSERT INTO outtable (quantitépchneufeentré,le,produit,unité,origine,fournisseur)VALUES(%s,%s,%s,%s,%s,%s)",(quantité,dateentre,produitt,lieu,country,fourni))
                        cursi.execute(
                            "INSERT INTO entrytable (le,nlot,nlotsch,produit,origine,quantité,unité,sacherie,fournisseur,imatricule,date_recolte,date_CND,DLUM,quantitépchneufentrée)VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                            (dateentre,nlot,nlotsch ,produitt, country, quantité, lieu, sac, fourni, matric,datrecolte, cnddate, datedlm,
                             quantité))
                        dattabase.commit()
                        dattabasei.commit()
                    if produitt == 'POIS CHICHES IMP 08 mm':
                        dattabase = mysql.connector.connect(
        host="localhost",
         user=user,
        password=password,database='datta_legumsec_out'
    )
                        curs = dattabase.cursor()
                        dattabasei = mysql.connector.connect(
        host="localhost",
         user=user,
        password=password,database='datta_legumesec_entry'
    )
                        cursi = dattabasei.cursor()
                        curs.execute("INSERT INTO outtable ( quantitépchhuitentré,le,produit,unité,origine,fournisseur)VALUES(%s,%s,%s,%s,%s,%s)",(quantité,dateentre,produitt,lieu,country,fourni))
                        cursi.execute(
                            "INSERT INTO entrytable (le,nlot,nlotsch,produit,origine,quantité,unité,sacherie,fournisseur,imatricule,date_recolte,date_CND,DLUM,quantitépchhuitntrée)VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                            (dateentre,nlot,nlotsch, produitt, country, quantité, lieu, sac, fourni, matric,datrecolte, cnddate, datedlm,
                             quantité))
                        dattabase.commit()
                        dattabasei.commit()
                    if produitt == 'POIS CHICHES IMP 06 mm':
                        dattabase = mysql.connector.connect(
        host="localhost",
         user=user,
        password=password,database='datta_legumsec_out'
    )
                        curs = dattabase.cursor()
                        dattabasei = mysql.connector.connect(
        host="localhost",
         user=user,
        password=password,database='datta_legumesec_entry'
    )
                        cursi = dattabasei.cursor()
                        curs.execute("INSERT INTO outtable (quantitépchsixentré,le,produit,unité,origine,fournisseur)VALUES(%s,%s,%s,%s,%s,%s)",(quantité,dateentre,produitt,lieu,country,fourni))
                        cursi.execute(
                            "INSERT INTO entrytable (le,nlot,nlotsch,produit,origine,quantité,unité,sacherie,fournisseur,imatricule,date_recolte,date_CND,DLUM,quantitépchsixntrée)VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                            (dateentre,nlot,nlotsch, produitt, country, quantité, lieu, sac, fourni, matric,datrecolte, cnddate, datedlm,
                             quantité))
                        dattabase.commit()
                        dattabasei.commit()
                    if produitt == 'POIS CHICHES PL':
                        dattabase = mysql.connector.connect(
        host="localhost",
         user=user,
        password=password,database='datta_legumsec_out'
    )
                        curs = dattabase.cursor()
                        dattabasei = mysql.connector.connect(
        host="localhost",
         user=user,
        password=password,database='datta_legumesec_entry'
    )
                        cursi = dattabasei.cursor()
                        curs.execute("INSERT INTO outtable (quantitépchplentré,le,produit,unité,origine,fournisseur)VALUES(%s,%s,%s,%s,%s,%s)",(quantité,dateentre,produitt,lieu,country,fourni))
                        cursi.execute(
                            "INSERT INTO entrytable (le,nlot,nlotsch,produit,origine,quantité,unité,sacherie,fournisseur,imatricule,date_recolte,date_CND,DLUM,quantitépchplentrée)VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                            (dateentre,nlot,nlotsch, produitt, country, quantité, lieu, sac, fourni, matric,datrecolte, cnddate, datedlm,
                             quantité))
                        dattabase.commit()
                        dattabasei.commit()
                    msgbox = QtWidgets.QMessageBox()
                    msgbox.setWindowTitle('confirmation')
                    msgbox.setText('Lopération a été ajoutée avec succès')
                    msgbox.exec()
                    self.impot_all()

                else:
                    msgbox = QtWidgets.QMessageBox()
                    msgbox.setWindowTitle('Erreur')
                    msgbox.setText("Erreur: Sélectionne le produit")
                    msgbox.exec()
            except mysql.connector.Error as er:
                print(er)

        def add_sortie(self):
            try:
                dateentre1 = self.dateedite_1.text()
                nlots = self.nemuroDelots.text()
                nlotssch=self.nemuroDelotssch.text()
                magasin_de_stock = self.magasincombos.currentText()
                produitt1 = self.produitcombo_1.currentText()
                country1 = self.originel_1.currentText()
                quantité1 = self.quatitelineEdit_1.value()
                achteur = self.magasincombo_1.currentText()
                sac1 = self.cndcombo_1.currentText()
                matric1 = self.imatricullineeditr_1.text()
                daterecoltso = self.daterecolteso.text()
                cnddate1 = self.dateeditedcnd_1.text()
                datedlm1 = self.dateEditdlum_1.text()
                expediteur = self.fourniseurcombos.currentText()

                try:
                    if produitt1:
                        dattabase1 = mysql.connector.connect(host="localhost",  user=user, password=password,
                                                             database='datta_legumsec_out')
                        curs1 = dattabase1.cursor()
                        if produitt1 == 'RIZ IMP':
                            curs1.execute(
                                "INSERT INTO outtable (le, nlot,nlotsch, unité, produit, origine, quantité, achteur, sacherie, imatricule, date_recoltes, date_CND, DLUM, fournisseur, quantitérizesorté) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                                (dateentre1, nlots,nlotssch, magasin_de_stock, produitt1, country1, quantité1, achteur, sac1,
                                 matric1, daterecoltso, cnddate1, datedlm1, expediteur, quantité1))
                        elif produitt1 == 'RIZ ETUVE':
                            curs1.execute(
                                "INSERT INTO outtable (le, nlot,nlotsch, unité, produit, origine, quantité, achteur, sacherie, imatricule, date_recoltes, date_CND, DLUM, fournisseur, quantitérizetvsorté) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                                (dateentre1, nlots,nlotssch, magasin_de_stock, produitt1, country1, quantité1, achteur, sac1,
                                 matric1, daterecoltso, cnddate1, datedlm1, expediteur, quantité1))
                        elif produitt1 == 'LENTILLE IMP 05-07 mm':
                            curs1.execute(
                                "INSERT INTO outtable (le, nlot,nlotsch, unité, produit, origine, quantité, achteur, sacherie, imatricule, date_recoltes, date_CND, DLUM, fournisseur, quantitélensorté) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                                (dateentre1, nlots,nlotssch, magasin_de_stock, produitt1, country1, quantité1, achteur, sac1,
                                 matric1, daterecoltso, cnddate1, datedlm1, expediteur, quantité1))
                        elif produitt1 == 'LENTILLE IMP vert':
                            curs1.execute(
                                "INSERT INTO outtable (le, nlot, nlotsch,unité, produit, origine, quantité, achteur, sacherie, imatricule, date_recoltes, date_CND, DLUM, fournisseur, quantitélenvertsorté) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                                (dateentre1, nlots,nlotssch, magasin_de_stock, produitt1, country1, quantité1, achteur, sac1,
                                 matric1, daterecoltso, cnddate1, datedlm1, expediteur, quantité1))
                        elif produitt1 == 'LENTILLE IMP ROUGE':
                            curs1.execute(
                                "INSERT INTO outtable (le, nlot,nlotsch, unité, produit, origine, quantité, achteur, sacherie, imatricule, date_recoltes, date_CND, DLUM, fournisseur, quantitélenrougesorté) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                                (dateentre1, nlots,nlotssch, magasin_de_stock, produitt1, country1, quantité1, achteur, sac1,
                                 matric1, daterecoltso, cnddate1, datedlm1, expediteur, quantité1))
                        # Add other conditions for different products here
                        elif produitt1 == 'POIS CHICHES IMP 12 mm':
                            curs1.execute(
                                "INSERT INTO outtable (le, nlot,nlotsch, unité, produit, origine, quantité, achteur, sacherie, imatricule, date_recoltes, date_CND, DLUM, fournisseur, quantitépchsorté) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                                (dateentre1, nlots,nlotssch, magasin_de_stock, produitt1, country1, quantité1, achteur, sac1,
                                 matric1, daterecoltso, cnddate1, datedlm1, expediteur, quantité1))
                        elif produitt1 == 'POIS CHICHES IMP 09 mm':
                            curs1.execute(
                                "INSERT INTO outtable (le, nlot,nlotsch, unité, produit, origine, quantité, achteur, sacherie, imatricule, date_recoltes, date_CND, DLUM, fournisseur, quantitépchneufesorté) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                                (dateentre1, nlots,nlotssch, magasin_de_stock, produitt1, country1, quantité1, achteur, sac1,
                                 matric1, daterecoltso, cnddate1, datedlm1, expediteur, quantité1))
                        elif produitt1 == 'POIS CHICHES IMP 08 mm':
                            curs1.execute(
                                "INSERT INTO outtable (le, nlot,nlotsch, unité, produit, origine, quantité, achteur, sacherie, imatricule, date_recoltes, date_CND, DLUM, fournisseur, quantitépchhuitsorté) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                                (dateentre1, nlots,nlotssch, magasin_de_stock, produitt1, country1, quantité1, achteur, sac1,
                                 matric1, daterecoltso, cnddate1, datedlm1, expediteur, quantité1))
                        elif produitt1 == 'POIS CHICHES IMP 06 mm':
                            curs1.execute(
                                "INSERT INTO outtable (le, nlot,nlotsch, unité, produit, origine, quantité, achteur, sacherie, imatricule, date_recoltes, date_CND, DLUM, fournisseur, quantitépchsixsorté) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                                (dateentre1, nlots,nlotssch, magasin_de_stock, produitt1, country1, quantité1, achteur, sac1,
                                 matric1, daterecoltso, cnddate1, datedlm1, expediteur, quantité1))
                        elif produitt1 == 'POIS CHICHES PL':
                            curs1.execute(
                                "INSERT INTO outtable (le, nlot,nlotsch, unité, produit, origine, quantité, achteur, sacherie, imatricule, date_recoltes, date_CND, DLUM, fournisseur, quantitépchplsorté) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                                (dateentre1, nlots,nlotssch, magasin_de_stock, produitt1, country1, quantité1, achteur, sac1,
                                 matric1, daterecoltso, cnddate1, datedlm1, expediteur, quantité1))
                        elif produitt1 == 'HARICOT BLANC IMP':
                            curs1.execute(
                                "INSERT INTO outtable (le, nlot,nlotsch, unité, produit, origine, quantité, achteur, sacherie, imatricule, date_recoltes, date_CND, DLUM, fournisseur, quantitéharsorté) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                                (dateentre1, nlots,nlotssch, magasin_de_stock, produitt1, country1, quantité1, achteur, sac1,
                                 matric1, daterecoltso, cnddate1, datedlm1, expediteur, quantité1))
                        elif produitt1 == 'HARICOT LSB':
                            curs1.execute(
                                "INSERT INTO outtable (le, nlot,nlotsch, unité, produit, origine, quantité, achteur, sacherie, imatricule, date_recoltes, date_CND, DLUM, fournisseur, quantitéharlsbsorté) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                                (dateentre1, nlots,nlotssch, magasin_de_stock, produitt1, country1, quantité1, achteur, sac1,
                                 matric1, daterecoltso, cnddate1, datedlm1, expediteur, quantité1))
                        elif produitt1 == 'LENTILLE PL':
                            curs1.execute(
                                "INSERT INTO outtable (le, nlot,nlotsch, unité, produit, origine, quantité, achteur, sacherie, imatricule, date_recoltes, date_CND, DLUM, fournisseur, quantitélenplsorté) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                                (dateentre1, nlots,nlotssch, magasin_de_stock, produitt1, country1, quantité1, achteur, sac1,
                                 matric1, daterecoltso, cnddate1, datedlm1, expediteur, quantité1))

                        dattabase1.commit()
                        dattabase1.close()
                        msgbox = QtWidgets.QMessageBox()
                        msgbox.setWindowTitle('confirmation')
                        msgbox.setText('L\'opération a été ajoutée avec succès')
                        msgbox.exec()
                        self.impot_all_sortie()
                    else:
                        msgbox = QtWidgets.QMessageBox()
                        msgbox.setWindowTitle('Erreur')
                        msgbox.setText("Erreur: Sélectionnez le produit")
                        msgbox.exec()
                except Exception as e:
                    print(e)
            except mysql.connector.Error as e:
                print(e)

        def impot_filter(self):
            try:
                datefl = self.datefilter.text()
                produitfl = self.produitfl.currentText()
                dattabase = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumesec_entry')

                curs = dattabase.cursor()
                curs.execute(
                    "SELECT nlot,nlotsch,le ,produit ,origine,quantité ,unité,sacherie,fournisseur ,imatricule  ,date_CND ,DLUM  FROM entrytable WHERE le = %s or produit= %s",
                    (datefl, produitfl,))
                result = curs.fetchall()
                self.textEdit.setRowCount(0)
                for row, row_datta in enumerate(result):
                    self.textEdit.insertRow(row)
                    for colum, datta in enumerate(row_datta):
                        self.textEdit.setItem(row, colum, QTableWidgetItem(str(datta)))
                dattabase.commit()
                dattabase.close()
            except Exception as e :
                print(e)

        def impot_filter_sortie(self):
            try:
                datefls = self.datefilters.text()
                produitfls = self.produitfls.currentText()
                dattabase = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumsec_out')
                curs = dattabase.cursor()
                curs.execute(
                    "SELECT nlot,nlotsch,le,unité,produit ,origine ,quantité ,achteur ,sacherie ,imatricule,date_CND ,DLUM  FROM outtable WHERE le = %s or produit= %s",
                    (datefls, produitfls,))
                result = curs.fetchall()
                self.textEdit_1.setRowCount(0)
                for row, row_datta in enumerate(result):
                    self.textEdit_1.insertRow(row)
                    for colum, datta in enumerate(row_datta):
                        self.textEdit_1.setItem(row, colum, QTableWidgetItem(str(datta)))
                dattabase.commit()
                dattabase.close()
            except Exception as e:
                print(e)

        def print_docxStockDatabase(self):
            self.poischicheDouze()
            self.poischicheNeuf()
            self.poischicheOuit()
            self.poichichSIxSept()
            self.operation_dattabase()
            self.haricot()
            self.haricotLsb()
            self.lentilleCinqSept()
            self.lenVert()
            self.lenRouge()
            self.lenpl()
            self.poichichepl()
            self.rizeEtuve()
            self.selected_produit = self.Produits.currentText()
            self.selected_magasin = self.Unites.currentText()
            self.selected_origine=self.OrigineDuProduits.currentText()
            self.selected_expediteur=self.expediteur.currentText()
            self.selected_dates = self.date_edits.text()
            self.selected_datesAU = self.date_edit2.text()
            try:
                if self.selected_produit and self.selected_magasin and self.selected_origine and self.selected_dates=="--" and self.selected_datesAU=="--":
                    conn = mysql.connector.connect(host="localhost",  user=user, password=password,
                                                   database="datta_legumesec_entry")
                    cursor = conn.cursor()
                    cursor.execute(
                        "SELECT le,produit,origine,quantité,unité,sacherie,fournisseur, imatricule,date_recolte, date_CND, DLUM FROM entrytable WHERE produit = %s AND unité = %s AND fournisseur=%s  ",
                        (self.selected_produit, self.selected_magasin,self.selected_expediteur,)
                    )
                    rows = cursor.fetchall()

                    ########################################################################
                    # Create a new document and add a table
                    self.doc = docx.Document()
                    section = self.doc.sections[0]
                    section.orientation = WD_ORIENTATION.LANDSCAPE
                    section.page_width = docx.shared.Cm(29.7)
                    section.page_height = docx.shared.Cm(21.0)
                    section.top_margin = docx.shared.Cm(0.5)
                    section.bottom_margin = docx.shared.Cm(1.5)
                    section.right_margin = docx.shared.Cm(1.5)
                    section.left_margin = docx.shared.Cm(1.5)
                    heading1 = self.doc.add_heading(
                        '\t\t\t\t\t\tSITUATION DES STOCKS DES LEGUMES SECS\n\t\t\t\t\t\t\t    PRODUITS RÉCEPTIONNER', level=1)
                    heading1.style.font.name = 'Times New Roman'
                    heading1.style.font.size = Pt(14)
                    heading1.style.font.bold = True
                    heading1.style.font.color.rgb = RGBColor(0, 0, 0)
                    # add second heading
                    paragraph0 = self.doc.add_paragraph()
                    paragraph0.style.font.name = 'Times New Roman'
                    left_run0 = paragraph0.add_run('Période du :' + self.selected_dates + "  au:" + self.selected_datesAU)
                    paragraph0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    paragraph1 = self.doc.add_paragraph()
                    paragraph1.style.font.name = 'Times New Roman'
                    current_date = QDate.currentDate().getDate()
                    left_run1 = paragraph1.add_run("\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t Date d'édition : {}/{}/{}".format(current_date[2], current_date[1], current_date[0]))
                    left_run1.font.size = Pt(9)

                    paragraph = self.doc.add_paragraph()
                    paragraph.style.font.name = 'Times New Roman'
                    paragraph.style.font.size = Pt(12)
                    left_run = paragraph.add_run(self.selected_magasin + '\t\t\t\t\t\t\t\t\t\t\t\t\t')
                    left_run.bold = True

                    table = self.doc.add_table(rows=1, cols=11)
                    table.style = "Table Grid"  # set the table style
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'DATE'
                    hdr_cells[1].text = 'PRODUIT'
                    hdr_cells[2].text = 'ORIGINE'
                    hdr_cells[3].text = 'QUANTITE'
                    hdr_cells[4].text = 'MAGASIN DE STOCK'
                    hdr_cells[5].text = 'SASHERIE  CND'
                    hdr_cells[6].text = 'EXPIDITEUR'
                    hdr_cells[7].text = 'IMATRICUL'
                    hdr_cells[8].text = 'Date de récolte'
                    hdr_cells[9].text = 'Date de CND'
                    hdr_cells[10].text = 'Date de DLUM'
                    # Set the width of the header cells
                    hdr_cells[0].width = Inches(0.5)  # set the width of the first column to 1.5 inches
                    hdr_cells[1].width = Inches(2)
                    hdr_cells[2].width = Inches(1)
                    hdr_cells[3].width = Inches(0.4)
                    hdr_cells[4].width = Inches(1.3)
                    hdr_cells[5].width = Inches(1)
                    hdr_cells[6].width = Inches(1)
                    hdr_cells[7].width = Inches(1)
                    hdr_cells[8].width = Inches(0.5)
                    hdr_cells[9].width = Inches(0.5)
                    hdr_cells[10].width = Inches(0.5)

                    for cell in table.rows[0].cells:
                        cell.paragraphs[0].runs[0].font.name = 'Calibri'
                        cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                        cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                        cell.paragraphs[0].runs[0].font.bold = True
                        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    # Add the data to the table
                    table_rows = len(rows)
                    table_cols = len(rows[0])
                    for row in range(table_rows):
                        table.add_row()
                        for col in range(table_cols):
                            cell = table.cell(row + 1, col)
                            cell.text = str(rows[row][col])
                            # Set font properties
                            cell.paragraphs[0].runs[0].font.name = 'Calibri'
                            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
                            cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                            cell.paragraphs[0].runs[0].font.bold = False
                            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    for row in table.rows:
                        row.height = Inches(0.3)
                    # Create the document
                    # Create the table
                    self.doc.add_paragraph().add_run().add_break()
                    table_b = self.doc.add_table(rows=14, cols=2)
                    table_b.style = "Table Grid"
                    # Add table headers
                    hdr_colum = table_b.rows[0].cells
                    hdr_colum[0].text = "PRODUIT"
                    hdr_colum[1].text = "STOCK"
                    for hdr_colum in table_b.rows[0].cells:
                        hdr_colum.paragraphs[0].runs[0].font.name = 'Calibri'
                        hdr_colum.paragraphs[0].runs[0].font.size = Pt(10.5)
                        hdr_colum.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                        hdr_colum.paragraphs[0].runs[0].font.bold = True
                        hdr_colum.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    hdr_colum = table_b.rows[1].cells
                    hdr_colum[0].text = "POIS CHICHES PL"
                    hdr_colum[1].text = ""
                    hdr_colum = table_b.rows[2].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 12 mm"
                    hdr_colum[1].text = self.restpoichichedouze + "  QX"
                    hdr_colum = table_b.rows[3].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 09 mm"
                    hdr_colum[1].text = self.restpoichicheneuf + " QX"
                    hdr_colum = table_b.rows[4].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 08 mm"
                    hdr_colum[1].text = self.restpoichicheouit + " QX"
                    hdr_colum = table_b.rows[5].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 06 mm"
                    hdr_colum[1].text = self.restpoichiche + " QX"
                    hdr_colum = table_b.rows[6].cells
                    hdr_colum[0].text = "HARICOT BLANC IMP"
                    hdr_colum[1].text = self.restharicotimp + " QX"
                    hdr_colum = table_b.rows[7].cells
                    hdr_colum[0].text = "HARICOT LSB"
                    hdr_colum[1].text = self.restharicotlsb + " QX"
                    hdr_colum = table_b.rows[8].cells
                    hdr_colum[0].text = "LENTILLE PL"
                    hdr_colum[1].text = self.restlenpl + " QX"
                    hdr_colum = table_b.rows[9].cells
                    hdr_colum[0].text = "LENTILLE IMP 05-07 mm"
                    hdr_colum[1].text = self.restlentsinq + " QX"
                    hdr_colum = table_b.rows[10].cells
                    hdr_colum[0].text = "LENTILLE IMP vert"
                    hdr_colum[1].text = self.restlentvert + " QX"
                    hdr_colum = table_b.rows[11].cells
                    hdr_colum[0].text = "LENTILLE IMP ROUGE"
                    hdr_colum[1].text = self.restlentrouge + " QX"
                    hdr_colum = table_b.rows[12].cells
                    hdr_colum[0].text = "RIZ IMP"
                    hdr_colum[1].text = self.resulxtt + " QX"
                    hdr_colum = table_b.rows[13].cells
                    hdr_colum[0].text = "RIZ ETUVE"
                    hdr_colum[1].text = self.resulxttv + " QX"
                    for row in table_b.rows[1:13]:
                        for cell in row.cells:
                            cell.paragraphs[0].runs[0].font.name = 'Calibri'
                            cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                            cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                            cell.paragraphs[0].runs[0].font.bold = False

                    # Save and open the document for printing
                    temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                    self.doc.save(temp_file)
                    os.startfile(temp_file, "open")
                    conn.close()
                elif self.selected_produit and self.selected_magasin and self.selected_origine and self.selected_dates and self.selected_datesAU:

                    conn = mysql.connector.connect(host="localhost",  user=user, password=password,
                                                   database="datta_legumesec_entry")
                    cursor = conn.cursor()
                    cursor.execute(
                        "SELECT le,produit,origine,quantité,unité,sacherie,fournisseur, imatricule,date_recolte, date_CND, DLUM FROM entrytable WHERE produit = %s AND unité = %s AND fournisseur=%s AND le BETWEEN %s AND %s",
                        (self.selected_produit, self.selected_magasin, self.selected_expediteur,self.selected_dates,self.selected_datesAU))
                    rows = cursor.fetchall()

                    ########################################################################
                    # Create a new document and add a table
                    self.doc = docx.Document()
                    section = self.doc.sections[0]
                    section.orientation = WD_ORIENTATION.LANDSCAPE
                    section.page_width = docx.shared.Cm(29.7)
                    section.page_height = docx.shared.Cm(21.0)
                    section.top_margin = docx.shared.Cm(0.5)
                    section.bottom_margin = docx.shared.Cm(1.5)
                    section.right_margin = docx.shared.Cm(1.5)
                    section.left_margin = docx.shared.Cm(1.5)
                    heading1 = self.doc.add_heading(
                        '\t\t\t\t\t\tSITUATION DES STOCKS DES LEGUMES SECS\n\t\t\t\t\t\t\t    PRODUITS RÉCEPTIONNER', level=1)
                    heading1.style.font.name = 'Times New Roman'
                    heading1.style.font.size = Pt(14)
                    heading1.style.font.bold = True
                    heading1.style.font.color.rgb = RGBColor(0, 0, 0)
                    # add second heading
                    # add second heading
                    paragraph0 = self.doc.add_paragraph()
                    paragraph0.style.font.name = 'Times New Roman'
                    left_run0 = paragraph0.add_run(
                        'Période du :' + self.selected_dates + "  au:" + self.selected_datesAU)
                    paragraph0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    paragraph1 = self.doc.add_paragraph()
                    paragraph1.style.font.name = 'Times New Roman'
                    current_date = QDate.currentDate().getDate()
                    left_run1 = paragraph1.add_run(
                        "\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t Date d'édition : {}/{}/{}".format(current_date[2],
                                                                                                current_date[1],
                                                                                                current_date[0]))
                    left_run1.font.size = Pt(9)

                    paragraph = self.doc.add_paragraph()
                    paragraph.style.font.name = 'Times New Roman'
                    paragraph.style.font.size = Pt(12)
                    left_run = paragraph.add_run(self.selected_magasin + '\t\t\t\t\t\t\t\t\t\t\t\t\t')
                    left_run.bold = True

                    table = self.doc.add_table(rows=1, cols=11)
                    table.style = "Table Grid"  # set the table style
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'DATE'
                    hdr_cells[1].text = 'PRODUIT'
                    hdr_cells[2].text = 'ORIGINE'
                    hdr_cells[3].text = 'QUANTITE'
                    hdr_cells[4].text = 'MAGASIN DE STOCK'
                    hdr_cells[5].text = 'SASHERIE  CND'
                    hdr_cells[6].text = 'EXPIDITEUR'
                    hdr_cells[7].text = 'IMATRICUL'
                    hdr_cells[8].text = 'Date de récolte'
                    hdr_cells[9].text = 'Date de CND'
                    hdr_cells[10].text = 'Date de DLUM'
                    # Set the width of the header cells
                    hdr_cells[0].width = Inches(0.5)  # set the width of the first column to 1.5 inches
                    hdr_cells[1].width = Inches(2)
                    hdr_cells[2].width = Inches(1)
                    hdr_cells[3].width = Inches(0.4)
                    hdr_cells[4].width = Inches(1.3)
                    hdr_cells[5].width = Inches(1)
                    hdr_cells[6].width = Inches(1)
                    hdr_cells[7].width = Inches(1)
                    hdr_cells[8].width = Inches(0.5)
                    hdr_cells[9].width = Inches(0.5)
                    hdr_cells[10].width = Inches(0.5)

                    for cell in table.rows[0].cells:
                        cell.paragraphs[0].runs[0].font.name = 'Calibri'
                        cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                        cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                        cell.paragraphs[0].runs[0].font.bold = True
                        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    # Add the data to the table
                    table_rows = len(rows)
                    table_cols = len(rows[0])
                    for row in range(table_rows):
                        table.add_row()
                        for col in range(table_cols):
                            cell = table.cell(row + 1, col)
                            cell.text = str(rows[row][col])
                            # Set font properties
                            cell.paragraphs[0].runs[0].font.name = 'Calibri'
                            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
                            cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                            cell.paragraphs[0].runs[0].font.bold = False
                            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    for row in table.rows:
                        row.height = Inches(0.3)
                    # Create the document
                    # Create the table
                    self.doc.add_paragraph().add_run().add_break()
                    table_b = self.doc.add_table(rows=14, cols=2)
                    table_b.style = "Table Grid"
                    # Add table headers
                    hdr_colum = table_b.rows[0].cells
                    hdr_colum[0].text = "PRODUIT"
                    hdr_colum[1].text = "STOCK"
                    for hdr_colum in table_b.rows[0].cells:
                        hdr_colum.paragraphs[0].runs[0].font.name = 'Calibri'
                        hdr_colum.paragraphs[0].runs[0].font.size = Pt(10.5)
                        hdr_colum.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                        hdr_colum.paragraphs[0].runs[0].font.bold = True
                        hdr_colum.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    hdr_colum = table_b.rows[1].cells
                    hdr_colum[0].text = "POIS CHICHES PL"
                    hdr_colum[1].text = ""
                    hdr_colum = table_b.rows[2].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 12 mm"
                    hdr_colum[1].text = self.restpoichichedouze + "  QX"
                    hdr_colum = table_b.rows[3].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 09 mm"
                    hdr_colum[1].text = self.restpoichicheneuf + " QX"
                    hdr_colum = table_b.rows[4].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 08 mm"
                    hdr_colum[1].text = self.restpoichicheouit + " QX"
                    hdr_colum = table_b.rows[5].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 06 mm"
                    hdr_colum[1].text = self.restpoichiche + " QX"
                    hdr_colum = table_b.rows[6].cells
                    hdr_colum[0].text = "HARICOT BLANC IMP"
                    hdr_colum[1].text = self.restharicotimp + " QX"
                    hdr_colum = table_b.rows[7].cells
                    hdr_colum[0].text = "HARICOT LSB"
                    hdr_colum[1].text = self.restharicotlsb + " QX"
                    hdr_colum = table_b.rows[8].cells
                    hdr_colum[0].text = "LENTILLE PL"
                    hdr_colum[1].text = self.restlenpl + " QX"
                    hdr_colum = table_b.rows[9].cells
                    hdr_colum[0].text = "LENTILLE IMP 05-07 mm"
                    hdr_colum[1].text = self.restlentsinq + " QX"
                    hdr_colum = table_b.rows[10].cells
                    hdr_colum[0].text = "LENTILLE IMP vert"
                    hdr_colum[1].text = self.restlentvert + " QX"
                    hdr_colum = table_b.rows[11].cells
                    hdr_colum[0].text = "LENTILLE IMP ROUGE"
                    hdr_colum[1].text = self.restlentrouge + " QX"
                    hdr_colum = table_b.rows[12].cells
                    hdr_colum[0].text = "RIZ IMP"
                    hdr_colum[1].text = self.resulxtt + " QX"
                    hdr_colum = table_b.rows[13].cells
                    hdr_colum[0].text = "RIZ ETUVE"
                    hdr_colum[1].text = self.resulxtt + " QX"
                    for row in table_b.rows[1:14]:
                        for cell in row.cells:
                            cell.paragraphs[0].runs[0].font.name = 'Calibri'
                            cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                            cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                            cell.paragraphs[0].runs[0].font.bold = False

                    # Save and open the document for printing
                    temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                    self.doc.save(temp_file)
                    os.startfile(temp_file, "open")
                    conn.close()
                elif self.selected_produit == "" and self.selected_magasin == "" and self.selected_origine == "" and self.selected_expediteur == "" and self.selected_dates == "--" and self.selected_datesAU == "--":

                    conn = mysql.connector.connect(host="localhost",  user=user, password=password,
                                                   database="datta_legumesec_entry")
                    cursor = conn.cursor()
                    cursor.execute(
                        "SELECT le,produit,origine,quantité,unité,sacherie,fournisseur,imatricule,date_recolte,date_CND,DLUM FROM entrytable")
                    rows = cursor.fetchall()

                    ########################################################################

                    # Create a new document and add a table
                    self.doc = docx.Document()
                    section = self.doc.sections[0]
                    section.orientation = WD_ORIENTATION.LANDSCAPE
                    section.page_width = docx.shared.Cm(29.7)
                    section.page_height = docx.shared.Cm(21.0)
                    section.top_margin = docx.shared.Cm(0.5)
                    section.bottom_margin = docx.shared.Cm(1.5)
                    section.right_margin = docx.shared.Cm(1.5)
                    section.left_margin = docx.shared.Cm(1.5)
                    heading1 = self.doc.add_heading(
                        '\t\t\t\t\t\tSITUATION DES STOCKS DES LEGUMES SECS\n\t\t\t\t\t\t\t    PRODUITS RÉCEPTIONNER',
                        level=1)
                    heading1.style.font.name = 'Times New Roman'
                    heading1.style.font.size = Pt(14)
                    heading1.style.font.bold = True
                    heading1.style.font.color.rgb = RGBColor(0, 0, 0)
                    # add second heading
                    paragraph = self.doc.add_paragraph()
                    paragraph.style.font.name = 'Times New Roman'
                    paragraph.style.font.size = Pt(9)
                    left_run = paragraph.add_run('\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t')
                    left_run.bold = True
                    current_date = QDate.currentDate().getDate()
                    right_run = paragraph.add_run(
                        "Date d'édition : {}/{}/{}".format(current_date[2], current_date[1], current_date[0]))
                    right_run.bold = False
                    table = self.doc.add_table(rows=1, cols=11)
                    table.style = "Table Grid"  # set the table style
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'DATE'
                    hdr_cells[1].text = 'PRODUIT'
                    hdr_cells[2].text = 'ORIGINE'
                    hdr_cells[3].text = 'QUANTITE'
                    hdr_cells[4].text = 'MAGASIN DE STOCK'
                    hdr_cells[5].text = 'SASHERIE  CND'
                    hdr_cells[6].text = 'FOURNISSEUR'
                    hdr_cells[7].text = 'IMATRICUL'
                    hdr_cells[8].text = 'Date de récolte'
                    hdr_cells[9].text = 'Date de CND'
                    hdr_cells[10].text = 'Date de DLUM'
                    # Set the width of the header cells
                    hdr_cells[0].width = Inches(0.5)  # set the width of the first column to 1.5 inches
                    hdr_cells[1].width = Inches(2)
                    hdr_cells[2].width = Inches(1)
                    hdr_cells[3].width = Inches(0.4)
                    hdr_cells[4].width = Inches(1.3)
                    hdr_cells[5].width = Inches(1)
                    hdr_cells[6].width = Inches(1)
                    hdr_cells[7].width = Inches(1)
                    hdr_cells[8].width = Inches(0.5)
                    hdr_cells[9].width = Inches(0.5)
                    hdr_cells[10].width = Inches(0.5)

                    for cell in table.rows[0].cells:
                        cell.paragraphs[0].runs[0].font.name = 'Calibri'
                        cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                        cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                        cell.paragraphs[0].runs[0].font.bold = True
                        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    # Add the data to the table
                    table_rows = len(rows)
                    table_cols = len(rows[0])
                    for row in range(table_rows):
                        table.add_row()
                        for col in range(table_cols):
                            cell = table.cell(row + 1, col)
                            cell.text = str(rows[row][col])
                            # Set font properties
                            cell.paragraphs[0].runs[0].font.name = 'Calibri'
                            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
                            cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                            cell.paragraphs[0].runs[0].font.bold = False
                            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    for row in table.rows:
                        row.height = Inches(0.3)
                    # Create the document
                        # Create the table
                    self.doc.add_paragraph().add_run().add_break()
                    table_b = self.doc.add_table(rows=14, cols=2)
                    table_b.style = "Table Grid"
                    # Add table headers
                    hdr_colum = table_b.rows[0].cells
                    hdr_colum[0].text = "PRODUIT"
                    hdr_colum[1].text = "STOCK"
                    for hdr_colum in table_b.rows[0].cells:
                        hdr_colum.paragraphs[0].runs[0].font.name = 'Calibri'
                        hdr_colum.paragraphs[0].runs[0].font.size = Pt(10.5)
                        hdr_colum.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                        hdr_colum.paragraphs[0].runs[0].font.bold = True
                        hdr_colum.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    hdr_colum = table_b.rows[1].cells
                    hdr_colum[0].text = "POIS CHICHES PL"
                    hdr_colum[1].text = self.restpoischichepltt +"  QX"
                    hdr_colum = table_b.rows[2].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 12 mm"
                    hdr_colum[1].text = self.restpoichichedouzett + "  QX"
                    hdr_colum = table_b.rows[3].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 09 mm"
                    hdr_colum[1].text = self.restpoichicheneuftt + " QX"
                    hdr_colum = table_b.rows[4].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 08 mm"
                    hdr_colum[1].text = self.restpoichicheouittt + " QX"
                    hdr_colum = table_b.rows[5].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 06 mm"
                    hdr_colum[1].text = self.restpoichichett + " QX"
                    hdr_colum = table_b.rows[6].cells
                    hdr_colum[0].text = "HARICOT BLANC IMP"
                    hdr_colum[1].text = self.restharicotimptt + " QX"
                    hdr_colum = table_b.rows[7].cells
                    hdr_colum[0].text = "HARICOT LSB"
                    hdr_colum[1].text = self.restharicotlsbtt + " QX"
                    hdr_colum = table_b.rows[8].cells
                    hdr_colum[0].text = "LENTILLE PL"
                    hdr_colum[1].text = self.restlenpltt + " QX"
                    hdr_colum = table_b.rows[9].cells
                    hdr_colum[0].text = "LENTILLE IMP 05-07 mm"
                    hdr_colum[1].text = self.restlentsinqtt + " QX"
                    hdr_colum = table_b.rows[10].cells
                    hdr_colum[0].text = "LENTILLE IMP vert"
                    hdr_colum[1].text = self.restlentvertt + " QX"
                    hdr_colum = table_b.rows[11].cells
                    hdr_colum[0].text = "LENTILLE IMP ROUGE"
                    hdr_colum[1].text = self.restlentrougett + " QX"
                    hdr_colum = table_b.rows[12].cells
                    hdr_colum[0].text = "RIZ IMP"
                    hdr_colum[1].text = self.resultt + " QX"
                    hdr_colum = table_b.rows[13].cells
                    hdr_colum[0].text = "RIZ ETUVE"
                    hdr_colum[1].text = self.resuletv + " QX"
                    for row in table_b.rows[1:14]:
                        for cell in row.cells:
                            cell.paragraphs[0].runs[0].font.name = 'Calibri'
                            cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                            cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                            cell.paragraphs[0].runs[0].font.bold = False
                    # Save and open the document for printing
                    temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                    self.doc.save(temp_file)
                    os.startfile(temp_file, "open")
                elif self.selected_produit=="" and self.selected_magasin=="" and self.selected_origine=="" and self.selected_dates and self.selected_datesAU:

                    conn = mysql.connector.connect(host="localhost",  user=user, password=password,
                                                   database="datta_legumesec_entry")
                    dattabase1 = mysql.connector.connect(host="localhost",  user=user, password=password,
                                                         database='datta_legumsec_out')
                    cursor = conn.cursor()
                    cursor.execute(
                        "SELECT le,produit,origine,quantité,unité,sacherie,fournisseur, imatricule, date_CND, DLUM FROM entrytable WHERE le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    rows = cursor.fetchall()
                    curs1x = dattabase1.cursor()
                    curs1x.execute(
                        "SELECT SUM(quantitérizentré) - SUM(quantitérizesorté) as quantitérizerest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    self.resulxs = curs1x.fetchone()
                    self.resulxttm = str(self.resulxs[0])
                    curs1xtv = dattabase1.cursor()
                    curs1xtv.execute(
                        "SELECT SUM(quantitérizetvntré) - SUM(quantitérizetvsorté) as quantitérizetvrest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    self.resulxstv = curs1xtv.fetchone()
                    self.resulxtttvm = str(self.resulxs[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitepchentré) - sum(quantitépchsorté) as quantitépchrest  FROM outtable   WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resultpch_att = curstt.fetchone()
                    self.restpoichichedouzettsd = str(resultpch_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT SUM(quantitépchsixentré) - SUM(quantitépchsixsorté) as quantitépchsixrest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resultpch_att = curstt.fetchone()
                    self.restpoichichettm = str(resultpch_att[0])

                    curs5tt = dattabase1.cursor()
                    curs5tt.execute(
                        "SELECT sum(quantitépchplentré) - sum(quantitépchplsorté) as quantitépchplhrest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resultpchpl_att = curs5tt.fetchone()
                    self.restpoischicheplttm = str(resultpchpl_att[0])
                    curs7tt = dattabase1.cursor()
                    curs7tt.execute(
                        "SELECT sum(quantitélenplentré) - sum(quantitélenplsorté) as quantitélenplrest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resultlenpl_att = curs7tt.fetchone()
                    self.restlenplttm = str(resultlenpl_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenvertentré) - sum(quantitélenvertsorté) as quantitélenrest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resultlenverttt = curs9tt.fetchone()
                    self.restlentvertttm = str(resultlenverttt[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenrougeentré) - sum(quantitélenrougesorté) as quantitélenrougerest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resultlenrougett = curs9tt.fetchone()
                    self.restlentrougettm = str(resultlenrougett[0])
                    curs11tt = dattabase1.cursor()
                    curs11tt.execute(
                        "SELECT sum(quantitéharentré) - sum(quantitéharsorté) as quantitéharrest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resulthar_att = curs11tt.fetchone()
                    self.restharicotimpttm = str(resulthar_att[0])
                    curs11tt.execute(
                        "SELECT SUM(quantitéharlsbentré) - SUM(quantitéharlsbsorté) as quantitéharlsbrest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resultharlsb_att = curs11tt.fetchone()
                    self.restharicotlsbttm = str(resultharlsb_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchneufeentré) - sum(quantitépchneufesorté) as quantitépchneuferest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheneufttm = str(resultpch_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchhuitentré) - sum(quantitépchhuitsorté) as quantitépchhuitrest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheouitttm = str(resultpch_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenentré) - sum(quantitélensorté) as quantitélenrest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resultlensinqtt = curs9tt.fetchone()
                    self.restlentsinqttm = str(resultlensinqtt[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT SUM(quantitépchsixentré) - SUM(quantitépchsixsorté) as quantitépchsixrest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resultpch_attm = curstt.fetchone()
                    self.restpoichichett = str(resultpch_att[0])

                    curs5tt = dattabase1.cursor()
                    curs5tt.execute(
                        "SELECT sum(quantitépchplentré) - sum(quantitépchplsorté) as quantitépchplhrest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resultpchpl_att = curs5tt.fetchone()
                    self.restpoischicheplttm = str(resultpchpl_att[0])
                    curs7tt = dattabase1.cursor()
                    curs7tt.execute(
                        "SELECT sum(quantitélenplentré) - sum(quantitélenplsorté) as quantitélenplrest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resultlenpl_att = curs7tt.fetchone()
                    self.restlenplttm = str(resultlenpl_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenvertentré) - sum(quantitélenvertsorté) as quantitélenrest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resultlenverttt = curs9tt.fetchone()
                    self.restlentvertttm = str(resultlenverttt[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenrougeentré) - sum(quantitélenrougesorté) as quantitélenrougerest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resultlenrougett = curs9tt.fetchone()
                    self.restlentrougettm = str(resultlenrougett[0])
                    curs11tt = dattabase1.cursor()
                    curs11tt.execute(
                        "SELECT sum(quantitéharentré) - sum(quantitéharsorté) as quantitéharrest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resulthar_att = curs11tt.fetchone()
                    self.restharicotimptt = str(resulthar_att[0])
                    curs11tt.execute(
                        "SELECT SUM(quantitéharlsbentré) - SUM(quantitéharlsbsorté) as quantitéharlsbrest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resultharlsb_att = curs11tt.fetchone()
                    self.restharicotlsbttm = str(resultharlsb_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchneufeentré) - sum(quantitépchneufesorté) as quantitépchneuferest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheneufttm = str(resultpch_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchhuitentré) - sum(quantitépchhuitsorté) as quantitépchhuitrest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheouitttm = str(resultpch_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenentré) - sum(quantitélensorté) as quantitélenrest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resultlensinqtt = curs9tt.fetchone()
                    self.restlentsinqttm = str(resultlensinqtt[0])

                    ########################################################################
                    ########################################################################
                    # Create a new document and add a table
                    self.doc = docx.Document()
                    section = self.doc.sections[0]
                    section.orientation = WD_ORIENTATION.LANDSCAPE
                    section.page_width = docx.shared.Cm(29.7)
                    section.page_height = docx.shared.Cm(21.0)
                    section.top_margin = docx.shared.Cm(0.5)
                    section.bottom_margin = docx.shared.Cm(1.5)
                    section.right_margin = docx.shared.Cm(1.5)
                    section.left_margin = docx.shared.Cm(1.5)
                    heading1 = self.doc.add_heading(
                        '\t\t\t\t\t\tSITUATION DES STOCKS DES LEGUMES SECS\n\t\t\t\t\t\t\t    PRODUITS RÉCEPTIONNER', level=1)
                    heading1.style.font.name = 'Times New Roman'
                    heading1.style.font.size = Pt(14)
                    heading1.style.font.bold = True
                    heading1.style.font.color.rgb = RGBColor(0, 0, 0)
                    # add second heading
                    paragraph0 = self.doc.add_paragraph()
                    paragraph0.style.font.name = 'Times New Roman'
                    left_run0 = paragraph0.add_run(
                        'Période du :' + self.selected_dates + "  au:" + self.selected_datesAU)
                    paragraph0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    paragraph1 = self.doc.add_paragraph()
                    paragraph1.style.font.name = 'Times New Roman'
                    current_date = QDate.currentDate().getDate()
                    left_run1 = paragraph1.add_run(
                        "\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t Date d'édition : {}/{}/{}".format(current_date[2],
                                                                                                current_date[1],
                                                                                                current_date[0]))
                    left_run1.font.size = Pt(9)

                    paragraph = self.doc.add_paragraph()
                    paragraph.style.font.name = 'Times New Roman'
                    paragraph.style.font.size = Pt(12)
                    left_run = paragraph.add_run(self.selected_magasin + '\t\t\t\t\t\t\t\t\t\t\t\t\t')
                    left_run.bold = True

                    table = self.doc.add_table(rows=1, cols=11)
                    table.style = "Table Grid"  # set the table style
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'DATE'
                    hdr_cells[1].text = 'PRODUIT'
                    hdr_cells[2].text = 'ORIGINE'
                    hdr_cells[3].text = 'QUANTITE'
                    hdr_cells[4].text = 'MAGASIN DE STOCK'
                    hdr_cells[5].text = 'SASHERIE  CND'
                    hdr_cells[6].text = 'EXPIDITEUR'
                    hdr_cells[7].text = 'IMATRICUL'
                    hdr_cells[8].text = 'Date de récolte'
                    hdr_cells[9].text = 'Date de CND'
                    hdr_cells[10].text = 'Date de DLUM'
                    # Set the width of the header cells
                    hdr_cells[0].width = Inches(0.5)  # set the width of the first column to 1.5 inches
                    hdr_cells[1].width = Inches(2)
                    hdr_cells[2].width = Inches(1)
                    hdr_cells[3].width = Inches(0.4)
                    hdr_cells[4].width = Inches(1.3)
                    hdr_cells[5].width = Inches(1)
                    hdr_cells[6].width = Inches(1)
                    hdr_cells[7].width = Inches(1)
                    hdr_cells[8].width = Inches(0.5)
                    hdr_cells[9].width = Inches(0.5)
                    hdr_cells[10].width = Inches(0.5)

                    for cell in table.rows[0].cells:
                        cell.paragraphs[0].runs[0].font.name = 'Calibri'
                        cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                        cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                        cell.paragraphs[0].runs[0].font.bold = True
                        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    # Add the data to the table
                    table_rows = len(rows)
                    table_cols = len(rows[0])
                    for row in range(table_rows):
                        table.add_row()
                        for col in range(table_cols):
                            cell = table.cell(row + 1, col)
                            cell.text = str(rows[row][col])
                            # Set font properties
                            cell.paragraphs[0].runs[0].font.name = 'Calibri'
                            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
                            cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                            cell.paragraphs[0].runs[0].font.bold = False
                            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    for row in table.rows:
                        row.height = Inches(0.3)
                    # Create the document
                    # Create the table
                    self.doc.add_paragraph().add_run().add_break()
                    table_b = self.doc.add_table(rows=14, cols=2)
                    table_b.style = "Table Grid"
                    # Add table headers
                    hdr_colum = table_b.rows[0].cells
                    hdr_colum[0].text = "PRODUIT"
                    hdr_colum[1].text = "STOCK"
                    for hdr_colum in table_b.rows[0].cells:
                        hdr_colum.paragraphs[0].runs[0].font.name = 'Calibri'
                        hdr_colum.paragraphs[0].runs[0].font.size = Pt(10.5)
                        hdr_colum.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                        hdr_colum.paragraphs[0].runs[0].font.bold = True
                        hdr_colum.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    hdr_colum = table_b.rows[1].cells
                    hdr_colum[0].text = "POIS CHICHES PL"
                    hdr_colum[1].text = ""
                    hdr_colum = table_b.rows[2].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 12 mm"
                    hdr_colum[1].text = self.restpoichichedouzettsd  + "  QX"
                    hdr_colum = table_b.rows[3].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 09 mm"
                    hdr_colum[1].text = self.restpoichicheneufttm + " QX"
                    hdr_colum = table_b.rows[4].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 08 mm"
                    hdr_colum[1].text = self.restpoichicheouitttm + " QX"
                    hdr_colum = table_b.rows[5].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 06 mm"
                    hdr_colum[1].text = self.restpoichichettm + " QX"
                    hdr_colum = table_b.rows[6].cells
                    hdr_colum[0].text = "HARICOT BLANC IMP"
                    hdr_colum[1].text = self.restharicotimpttm + " QX"
                    hdr_colum = table_b.rows[7].cells
                    hdr_colum[0].text = "HARICOT LSB"
                    hdr_colum[1].text = self.restharicotlsbttm + " QX"
                    hdr_colum = table_b.rows[8].cells
                    hdr_colum[0].text = "LENTILLE PL"
                    hdr_colum[1].text = self.restlenplttm + " QX"
                    hdr_colum = table_b.rows[9].cells
                    hdr_colum[0].text = "LENTILLE IMP 05-07 mm"
                    hdr_colum[1].text = self.restlentsinqttm + " QX"
                    hdr_colum = table_b.rows[10].cells
                    hdr_colum[0].text = "LENTILLE IMP vert"
                    hdr_colum[1].text = self.restlentvertttm + " QX"
                    hdr_colum = table_b.rows[11].cells
                    hdr_colum[0].text = "LENTILLE IMP ROUGE"
                    hdr_colum[1].text = self.restlentrougettm + " QX"
                    hdr_colum = table_b.rows[12].cells
                    hdr_colum[0].text = "RIZ IMP"
                    hdr_colum[1].text = self.resulxttm + " QX"
                    hdr_colum = table_b.rows[13].cells
                    hdr_colum[0].text = "RIZ ETUVE"
                    hdr_colum[1].text = self.resulxttm + " QX"
                    for row in table_b.rows[1:14]:
                        for cell in row.cells:
                            cell.paragraphs[0].runs[0].font.name = 'Calibri'
                            cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                            cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                            cell.paragraphs[0].runs[0].font.bold = False
                    # Save and open the document for printing
                    temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                    self.doc.save(temp_file)
                    os.startfile(temp_file, "open")
                    conn.close()
                elif self.selected_produit=="" and self.selected_magasin and self.selected_origine=="" and self.selected_dates=="--" and self.selected_datesAU=="--"  :

                    conn = mysql.connector.connect(host="localhost",  user=user, password=password,
                                                   database="datta_legumesec_entry")
                    dattabase1 = mysql.connector.connect(host="localhost",  user=user, password=password,
                                                         database='datta_legumsec_out')
                    cursor = conn.cursor()
                    cursor.execute(
                        "SELECT le,produit,origine,quantité,unité,sacherie,fournisseur, imatricule,date_recolte,date_CND,DLUM FROM entrytable WHERE unité = %s",
                        (self.selected_magasin,))
                    rows = cursor.fetchall()
                    curs1x = dattabase1.cursor()
                    curs1x.execute("SELECT SUM(quantitérizentré) - SUM(quantitérizesorté) as quantitérizerest FROM outtable WHERE  unité=%s",
                        (self.selected_magasin,))
                    self.resulxs = curs1x.fetchone()
                    self.resulxttm = str(self.resulxs[0])
                    curs1xtv = dattabase1.cursor()
                    curs1xtv.execute("SELECT SUM(quantitérizetvntré) - SUM(quantitérizetvsorté) as quantitérizetvrest FROM outtable WHERE  unité=%s",
                        (self.selected_magasin,))
                    self.resulxstv = curs1xtv.fetchone()
                    self.resulxtttvm = str(self.resulxstv[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitepchentré) - sum(quantitépchsorté) as quantitépchrest  FROM outtable   WHERE   unité=%s",
                        (self.selected_magasin,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichichedouzettm = str(resultpch_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT SUM(quantitépchsixentré) - SUM(quantitépchsixsorté) as quantitépchsixrest FROM outtable WHERE   unité=%s",
                        (self.selected_magasin,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichichettm = str(resultpch_att[0])

                    curs5tt = dattabase1.cursor()
                    curs5tt.execute(
                        "SELECT sum(quantitépchplentré) - sum(quantitépchplsorté) as quantitépchplhrest FROM outtable WHERE   unité=%s",
                        (self.selected_magasin,))
                    resultpchpl_att = curs5tt.fetchone()
                    self.restpoischicheplttm = str(resultpchpl_att[0])
                    curs7tt = dattabase1.cursor()
                    curs7tt.execute(
                        "SELECT sum(quantitélenplentré) - sum(quantitélenplsorté) as quantitélenplrest FROM outtable WHERE   unité=%s",
                        (self.selected_magasin,))
                    resultlenpl_att = curs7tt.fetchone()
                    self.restlenplttm = str(resultlenpl_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenvertentré) - sum(quantitélenvertsorté) as quantitélenrest FROM outtable WHERE   unité=%s",
                        (self.selected_magasin,))
                    resultlenverttt = curs9tt.fetchone()
                    self.restlentvertttm = str(resultlenverttt[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenrougeentré) - sum(quantitélenrougesorté) as quantitélenrougerest FROM outtable WHERE   unité=%s",
                        (self.selected_magasin,))
                    resultlenrougett = curs9tt.fetchone()
                    self.restlentrougettm = str(resultlenrougett[0])
                    curs11tt = dattabase1.cursor()
                    curs11tt.execute(
                        "SELECT sum(quantitéharentré) - sum(quantitéharsorté) as quantitéharrest FROM outtable WHERE   unité=%s",
                        (self.selected_magasin,))
                    resulthar_att = curs11tt.fetchone()
                    self.restharicotimpttm = str(resulthar_att[0])
                    curs11tt.execute(
                        "SELECT SUM(quantitéharlsbentré) - SUM(quantitéharlsbsorté) as quantitéharlsbrest FROM outtable WHERE   unité=%s",
                        (self.selected_magasin,))
                    resultharlsb_att = curs11tt.fetchone()
                    self.restharicotlsbttm = str(resultharlsb_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchneufeentré) - sum(quantitépchneufesorté) as quantitépchneuferest FROM outtable   WHERE unité=%s",
                        (self.selected_magasin,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheneufttm = str(resultpch_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchhuitentré) - sum(quantitépchhuitsorté) as quantitépchhuitrest FROM outtable WHERE   unité=%s",
                        (self.selected_magasin,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheouitttm = str(resultpch_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenentré) - sum(quantitélensorté) as quantitélenrest FROM outtable WHERE   unité=%s",
                        (self.selected_magasin,))
                    resultlensinqtt = curs9tt.fetchone()
                    self.restlentsinqttm = str(resultlensinqtt[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT SUM(quantitépchsixentré) - SUM(quantitépchsixsorté) as quantitépchsixrest FROM outtable  WHERE unité=%s",
                        (self.selected_magasin,))
                    resultpch_attm = curstt.fetchone()
                    self.restpoichichett = str(resultpch_att[0])

                    curs5tt = dattabase1.cursor()
                    curs5tt.execute(
                        "SELECT sum(quantitépchplentré) - sum(quantitépchplsorté) as quantitépchplhrest FROM outtable   WHERE unité=%s",
                        (self.selected_magasin,))
                    resultpchpl_att = curs5tt.fetchone()
                    self.restpoischicheplttm = str(resultpchpl_att[0])
                    curs7tt = dattabase1.cursor()
                    curs7tt.execute(
                        "SELECT sum(quantitélenplentré) - sum(quantitélenplsorté) as quantitélenplrest FROM outtable  WHERE unité=%s",
                        (self.selected_magasin,))
                    resultlenpl_att = curs7tt.fetchone()
                    self.restlenplttm = str(resultlenpl_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenvertentré) - sum(quantitélenvertsorté) as quantitélenrest FROM outtable WHERE unité=%s",
                        (self.selected_magasin,))
                    resultlenverttt = curs9tt.fetchone()
                    self.restlentvertttm = str(resultlenverttt[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenrougeentré) - sum(quantitélenrougesorté) as quantitélenrougerest FROM outtable WHERE unité=%s",
                        (self.selected_magasin,))
                    resultlenrougett = curs9tt.fetchone()
                    self.restlentrougettm = str(resultlenrougett[0])
                    curs11tt = dattabase1.cursor()
                    curs11tt.execute(
                        "SELECT sum(quantitéharentré) - sum(quantitéharsorté) as quantitéharrest FROM outtable WHERE  unité=%s",
                        (self.selected_magasin,))
                    resulthar_att = curs11tt.fetchone()
                    self.restharicotimptt = str(resulthar_att[0])
                    curs11tt.execute(
                        "SELECT SUM(quantitéharlsbentré) - SUM(quantitéharlsbsorté) as quantitéharlsbrest FROM outtable   WHERE unité=%s",
                        (self.selected_magasin,))
                    resultharlsb_att = curs11tt.fetchone()
                    self.restharicotlsbttm = str(resultharlsb_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchneufeentré) - sum(quantitépchneufesorté) as quantitépchneuferest FROM outtable  WHERE unité=%s",
                        (self.selected_magasin,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheneufttm = str(resultpch_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchhuitentré) - sum(quantitépchhuitsorté) as quantitépchhuitrest FROM outtable WHERE unité=%s",
                        (self.selected_magasin,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheouitttm = str(resultpch_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenentré) - sum(quantitélensorté) as quantitélenrest FROM outtable WHERE  unité=%s",
                        (self.selected_magasin,))
                    resultlensinqtt = curs9tt.fetchone()
                    self.restlentsinqttm = str(resultlensinqtt[0])

                    ########################################################################
                    # Create a new document and add a table
                    self.doc = docx.Document()
                    section = self.doc.sections[0]
                    section.orientation = WD_ORIENTATION.LANDSCAPE
                    section.page_width = docx.shared.Cm(29.7)
                    section.page_height = docx.shared.Cm(21.0)
                    section.top_margin = docx.shared.Cm(0.5)
                    section.bottom_margin = docx.shared.Cm(1.5)
                    section.right_margin = docx.shared.Cm(1.5)
                    section.left_margin = docx.shared.Cm(1.5)
                    heading1 = self.doc.add_heading(
                        '\t\t\t\t\t\tSITUATION DES STOCKS DES LEGUMES SECS\n\t\t\t\t\t\t\t    PRODUITS RÉCEPTIONNER', level=1)
                    heading1.style.font.name = 'Times New Roman'
                    heading1.style.font.size = Pt(14)
                    heading1.style.font.bold = True
                    heading1.style.font.color.rgb = RGBColor(0, 0, 0)
                    # add second heading

                    paragraph1 = self.doc.add_paragraph()
                    paragraph1.style.font.name = 'Times New Roman'
                    current_date = QDate.currentDate().getDate()
                    left_run1 = paragraph1.add_run(
                        "\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t Date d'édition : {}/{}/{}".format(current_date[2],
                                                                                                current_date[1],
                                                                                                current_date[0]))
                    left_run1.font.size = Pt(9)

                    paragraph = self.doc.add_paragraph()
                    paragraph.style.font.name = 'Times New Roman'
                    paragraph.style.font.size = Pt(12)
                    left_run = paragraph.add_run(self.selected_magasin + '\t\t\t\t\t\t\t\t\t\t\t\t\t')
                    left_run.bold = True

                    table = self.doc.add_table(rows=1, cols=11)
                    table.style = "Table Grid"  # set the table style
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'DATE'
                    hdr_cells[1].text = 'PRODUIT'
                    hdr_cells[2].text = 'ORIGINE'
                    hdr_cells[3].text = 'QUANTITE'
                    hdr_cells[4].text = 'MAGASIN DE STOCK'
                    hdr_cells[5].text = 'SASHERIE  CND'
                    hdr_cells[6].text = 'EXPIDITEUR'
                    hdr_cells[7].text = 'IMATRICUL'
                    hdr_cells[8].text = 'Date de récolte'
                    hdr_cells[9].text = 'Date de CND'
                    hdr_cells[10].text = 'Date de DLUM'
                    # Set the width of the header cells
                    hdr_cells[0].width = Inches(0.5)  # set the width of the first column to 1.5 inches
                    hdr_cells[1].width = Inches(2)
                    hdr_cells[2].width = Inches(1)
                    hdr_cells[3].width = Inches(0.4)
                    hdr_cells[4].width = Inches(1.3)
                    hdr_cells[5].width = Inches(1)
                    hdr_cells[6].width = Inches(1)
                    hdr_cells[7].width = Inches(1)
                    hdr_cells[8].width = Inches(0.5)
                    hdr_cells[9].width = Inches(0.5)
                    hdr_cells[10].width = Inches(0.5)

                    for cell in table.rows[0].cells:
                        cell.paragraphs[0].runs[0].font.name = 'Calibri'
                        cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                        cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                        cell.paragraphs[0].runs[0].font.bold = True
                        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    # Add the data to the table
                    table_rows = len(rows)
                    table_cols = len(rows[0])
                    for row in range(table_rows):
                        table.add_row()
                        for col in range(table_cols):
                            cell = table.cell(row + 1, col)
                            cell.text = str(rows[row][col])
                            # Set font properties
                            cell.paragraphs[0].runs[0].font.name = 'Calibri'
                            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
                            cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                            cell.paragraphs[0].runs[0].font.bold = False
                            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    for row in table.rows:
                        row.height = Inches(0.3)
                    # Create the document
                    # Create the table
                    self.doc.add_paragraph().add_run().add_break()
                    table_b = self.doc.add_table(rows=14, cols=2)
                    table_b.style = "Table Grid"
                    # Add table headers
                    hdr_colum = table_b.rows[0].cells
                    hdr_colum[0].text = "PRODUIT"
                    hdr_colum[1].text = "STOCK"
                    for hdr_colum in table_b.rows[0].cells:
                        hdr_colum.paragraphs[0].runs[0].font.name = 'Calibri'
                        hdr_colum.paragraphs[0].runs[0].font.size = Pt(10.5)
                        hdr_colum.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                        hdr_colum.paragraphs[0].runs[0].font.bold = True
                        hdr_colum.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    hdr_colum = table_b.rows[1].cells
                    hdr_colum[0].text = "POIS CHICHES PL"
                    hdr_colum[1].text = ""
                    hdr_colum = table_b.rows[2].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 12 mm"
                    hdr_colum[1].text = self.restpoichichedouzettm + "  QX"
                    hdr_colum = table_b.rows[3].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 09 mm"
                    hdr_colum[1].text = self.restpoichicheneufttm + " QX"
                    hdr_colum = table_b.rows[4].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 08 mm"
                    hdr_colum[1].text = self.restpoichicheouitttm + " QX"
                    hdr_colum = table_b.rows[5].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 06 mm"
                    hdr_colum[1].text = self.restpoichichettm + " QX"
                    hdr_colum = table_b.rows[6].cells
                    hdr_colum[0].text = "HARICOT BLANC IMP"
                    hdr_colum[1].text = self.restharicotimpttm + " QX"
                    hdr_colum = table_b.rows[7].cells
                    hdr_colum[0].text = "HARICOT LSB"
                    hdr_colum[1].text = self.restharicotlsbttm + " QX"
                    hdr_colum = table_b.rows[8].cells
                    hdr_colum[0].text = "LENTILLE PL"
                    hdr_colum[1].text = self.restlenplttm + " QX"
                    hdr_colum = table_b.rows[9].cells
                    hdr_colum[0].text = "LENTILLE IMP 05-07 mm"
                    hdr_colum[1].text = self.restlentsinqttm + " QX"
                    hdr_colum = table_b.rows[10].cells
                    hdr_colum[0].text = "LENTILLE IMP vert"
                    hdr_colum[1].text = self.restlentvertttm + " QX"
                    hdr_colum = table_b.rows[11].cells
                    hdr_colum[0].text = "LENTILLE IMP ROUGE"
                    hdr_colum[1].text = self.restlentrougettm + " QX"
                    hdr_colum = table_b.rows[12].cells
                    hdr_colum[0].text = "RIZ IMP"
                    hdr_colum[1].text = self.resulxttm + " QX"
                    hdr_colum = table_b.rows[13].cells
                    hdr_colum[0].text = "RIZ ETUVE"
                    hdr_colum[1].text = self.resulxttm + " QX"
                    for row in table_b.rows[1:14]:
                        for cell in row.cells:
                            cell.paragraphs[0].runs[0].font.name = 'Calibri'
                            cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                            cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                            cell.paragraphs[0].runs[0].font.bold = False
                    # Save and open the document for printing
                    temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                    self.doc.save(temp_file)
                    os.startfile(temp_file, "open")
                    conn.close()
                elif self.selected_produit and self.selected_magasin=="" and self.selected_origine=="" and self.selected_dates=="--" and self.selected_datesAU=="--":
                    conn = mysql.connector.connect(host="localhost",  user=user, password=password,
                                                   database="datta_legumesec_entry")
                    dattabase1 = mysql.connector.connect(host="localhost",  user=user, password=password,
                                                         database='datta_legumsec_out')
                    cursor = conn.cursor()
                    cursor.execute(
                        "SELECT le, produit, origine, quantité, unité, sacherie, fournisseur, imatricule,date_recolte, date_CND, DLUM FROM entrytable WHERE produit = %s",
                        (self.selected_produit,))
                    rows = cursor.fetchall()
                    curs1x = dattabase1.cursor()
                    curs1x.execute(
                        "SELECT SUM(quantitérizentré) - SUM(quantitérizesorté) as quantitérizerest FROM outtable WHERE  produit= %s",
                        (self.selected_produit,))
                    self.resulxs = curs1x.fetchone()
                    self.resulxttp = str(self.resulxs[0])
                    curs1xtv = dattabase1.cursor()
                    curs1xtv.execute(
                        "SELECT SUM(quantitérizetvntré) - SUM(quantitérizetvsorté) as quantitérizetvrest FROM outtable WHERE  produit= %s",
                        (self.selected_produit,))
                    self.resulxstv = curs1xtv.fetchone()
                    self.resulxtttvp = str(self.resulxstv[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitepchentré) - sum(quantitépchsorté) as quantitépchrest  FROM outtable   WHERE   produit=%s",
                        (self.selected_produit,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichichedouzettp = str(resultpch_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT SUM(quantitépchsixentré) - SUM(quantitépchsixsorté) as quantitépchsixrest FROM outtable WHERE   produit=%s",
                        (self.selected_produit,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichichettp = str(resultpch_att[0])

                    curs5tt = dattabase1.cursor()
                    curs5tt.execute(
                        "SELECT sum(quantitépchplentré) - sum(quantitépchplsorté) as quantitépchplhrest FROM outtable WHERE   produit=%s",
                        (self.selected_produit,))
                    resultpchpl_att = curs5tt.fetchone()
                    self.restpoischicheplttp = str(resultpchpl_att[0])
                    curs7tt = dattabase1.cursor()
                    curs7tt.execute(
                        "SELECT sum(quantitélenplentré) - sum(quantitélenplsorté) as quantitélenplrest FROM outtable WHERE   produit=%s",
                        (self.selected_produit,))
                    resultlenpl_att = curs7tt.fetchone()
                    self.restlenplttp = str(resultlenpl_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenvertentré) - sum(quantitélenvertsorté) as quantitélenrest FROM outtable WHERE   produit=%s",
                        (self.selected_produit,))
                    resultlenverttt = curs9tt.fetchone()
                    self.restlentvertttp = str(resultlenverttt[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenrougeentré) - sum(quantitélenrougesorté) as quantitélenrougerest FROM outtable WHERE   produit=%s",
                        (self.selected_produit,))
                    resultlenrougett = curs9tt.fetchone()
                    self.restlentrougettp = str(resultlenrougett[0])
                    curs11tt = dattabase1.cursor()
                    curs11tt.execute(
                        "SELECT sum(quantitéharentré) - sum(quantitéharsorté) as quantitéharrest FROM outtable WHERE   produit=%s",
                        (self.selected_produit,))
                    resulthar_att = curs11tt.fetchone()
                    self.restharicotimpttp = str(resulthar_att[0])
                    curs11tt.execute(
                        "SELECT SUM(quantitéharlsbentré) - SUM(quantitéharlsbsorté) as quantitéharlsbrest FROM outtable WHERE   produit=%s",
                        (self.selected_produit,))
                    resultharlsb_att = curs11tt.fetchone()
                    self.restharicotlsbttp = str(resultharlsb_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchneufeentré) - sum(quantitépchneufesorté) as quantitépchneuferest FROM outtable   WHERE produit=%s",
                        (self.selected_produit,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheneufttp = str(resultpch_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchhuitentré) - sum(quantitépchhuitsorté) as quantitépchhuitrest FROM outtable WHERE   produit=%s",
                        (self.selected_produit,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheouitttp = str(resultpch_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenentré) - sum(quantitélensorté) as quantitélenrest FROM outtable WHERE   produit=%s",
                        (self.selected_produit,))
                    resultlensinqtt = curs9tt.fetchone()
                    self.restlentsinqttp = str(resultlensinqtt[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT SUM(quantitépchsixentré) - SUM(quantitépchsixsorté) as quantitépchsixrest FROM outtable  WHERE produit=%s",
                        (self.selected_produit,))
                    resultpch_attp = curstt.fetchone()
                    self.restpoichichettp = str(resultpch_attp[0])

                    curs5tt = dattabase1.cursor()
                    curs5tt.execute(
                        "SELECT sum(quantitépchplentré) - sum(quantitépchplsorté) as quantitépchplhrest FROM outtable   WHERE produit=%s",
                        (self.selected_produit,))
                    resultpchpl_att = curs5tt.fetchone()
                    self.restpoischicheplttp = str(resultpchpl_att[0])
                    curs7tt = dattabase1.cursor()
                    curs7tt.execute(
                        "SELECT sum(quantitélenplentré) - sum(quantitélenplsorté) as quantitélenplrest FROM outtable  WHERE produit=%s",
                        (self.selected_produit,))
                    resultlenpl_att = curs7tt.fetchone()
                    self.restlenplttp = str(resultlenpl_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenvertentré) - sum(quantitélenvertsorté) as quantitélenrest FROM outtable WHERE produit=%s",
                        (self.selected_produit,))
                    resultlenverttt = curs9tt.fetchone()
                    self.restlentvertttp = str(resultlenverttt[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenrougeentré) - sum(quantitélenrougesorté) as quantitélenrougerest FROM outtable WHERE produit=%s",
                        (self.selected_produit,))
                    resultlenrougett = curs9tt.fetchone()
                    self.restlentrougettp = str(resultlenrougett[0])
                    curs11tt = dattabase1.cursor()
                    curs11tt.execute(
                        "SELECT sum(quantitéharentré) - sum(quantitéharsorté) as quantitéharrest FROM outtable WHERE  produit=%s",
                        (self.selected_produit,))
                    resulthar_att = curs11tt.fetchone()
                    self.restharicotimpttp = str(resulthar_att[0])
                    curs11tt.execute(
                        "SELECT SUM(quantitéharlsbentré) - SUM(quantitéharlsbsorté) as quantitéharlsbrest FROM outtable   WHERE produit=%s",
                        (self.selected_produit,))
                    resultharlsb_att = curs11tt.fetchone()
                    self.restharicotlsbttp = str(resultharlsb_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchneufeentré) - sum(quantitépchneufesorté) as quantitépchneuferest FROM outtable  WHERE produit=%s",
                        (self.selected_produit,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheneufttp = str(resultpch_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchhuitentré) - sum(quantitépchhuitsorté) as quantitépchhuitrest FROM outtable WHERE produit=%s",
                        (self.selected_produit,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheouitttp = str(resultpch_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenentré) - sum(quantitélensorté) as quantitélenrest FROM outtable WHERE  produit=%s",
                        (self.selected_produit,))
                    resultlensinqtt = curs9tt.fetchone()
                    self.restlentsinqttp = str(resultlensinqtt[0])
                    ########################################################################
                    # Create a new document and add a table
                    self.doc = docx.Document()
                    section = self.doc.sections[0]
                    section.orientation = WD_ORIENTATION.LANDSCAPE
                    section.page_width = docx.shared.Cm(29.7)
                    section.page_height = docx.shared.Cm(21.0)
                    section.top_margin = docx.shared.Cm(0.5)
                    section.bottom_margin = docx.shared.Cm(1.5)
                    section.right_margin = docx.shared.Cm(1.5)
                    section.left_margin = docx.shared.Cm(1.5)
                    heading1 = self.doc.add_heading(
                        '\t\t\t\t\t\tSITUATION DES STOCKS DES LEGUMES SECS\n\t\t\t\t\t\t\t    PRODUITS RÉCEPTIONNER', level=1)
                    heading1.style.font.name = 'Times New Roman'
                    heading1.style.font.size = Pt(14)
                    heading1.style.font.bold = True
                    heading1.style.font.color.rgb = RGBColor(0, 0, 0)
                    # add second heading


                    paragraph1 = self.doc.add_paragraph()
                    paragraph1.style.font.name = 'Times New Roman'
                    current_date = QDate.currentDate().getDate()
                    left_run1 = paragraph1.add_run(
                        "\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t Date d'édition : {}/{}/{}".format(current_date[2],
                                                                                                current_date[1],
                                                                                                current_date[0]))
                    left_run1.font.size = Pt(9)

                    paragraph = self.doc.add_paragraph()
                    paragraph.style.font.name = 'Times New Roman'
                    paragraph.style.font.size = Pt(12)
                    left_run = paragraph.add_run("PRODUIT :"+self.selected_produit + '\t\t\t\t\t\t\t\t\t\t\t\t\t')
                    left_run.bold = True

                    table = self.doc.add_table(rows=1, cols=11)
                    table.style = "Table Grid"  # set the table style
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'DATE'
                    hdr_cells[1].text = 'PRODUIT'
                    hdr_cells[2].text = 'ORIGINE'
                    hdr_cells[3].text = 'QUANTITE'
                    hdr_cells[4].text = 'MAGASIN DE STOCK'
                    hdr_cells[5].text = 'SASHERIE  CND'
                    hdr_cells[6].text = 'EXPIDITEUR'
                    hdr_cells[7].text = 'IMATRICUL'
                    hdr_cells[8].text = 'Date de récolte'
                    hdr_cells[9].text = 'Date de CND'
                    hdr_cells[10].text = 'Date de DLUM'
                    # Set the width of the header cells
                    hdr_cells[0].width = Inches(0.5)  # set the width of the first column to 1.5 inches
                    hdr_cells[1].width = Inches(2)
                    hdr_cells[2].width = Inches(1)
                    hdr_cells[3].width = Inches(0.4)
                    hdr_cells[4].width = Inches(1.3)
                    hdr_cells[5].width = Inches(1)
                    hdr_cells[6].width = Inches(1)
                    hdr_cells[7].width = Inches(1)
                    hdr_cells[8].width = Inches(0.5)
                    hdr_cells[9].width = Inches(0.5)
                    hdr_cells[10].width = Inches(0.5)

                    for cell in table.rows[0].cells:
                        cell.paragraphs[0].runs[0].font.name = 'Calibri'
                        cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                        cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                        cell.paragraphs[0].runs[0].font.bold = True
                        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    # Add the data to the table
                    table_rows = len(rows)
                    table_cols = len(rows[0])
                    for row in range(table_rows):
                        table.add_row()
                        for col in range(table_cols):
                            cell = table.cell(row + 1, col)
                            cell.text = str(rows[row][col])
                            # Set font properties
                            cell.paragraphs[0].runs[0].font.name = 'Calibri'
                            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
                            cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                            cell.paragraphs[0].runs[0].font.bold = False
                            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    for row in table.rows:
                        row.height = Inches(0.3)
                    # Create the document
                    # Create the table
                    self.doc.add_paragraph().add_run().add_break()
                    table_b = self.doc.add_table(rows=2, cols=2)
                    table_b.style = "Table Grid"
                    # Add table headers
                    hdr_colum = table_b.rows[0].cells
                    hdr_colum[0].text = "PRODUIT"
                    hdr_colum[1].text = "STOCK"
                    for hdr_colum in table_b.rows[0].cells:
                        hdr_colum.paragraphs[0].runs[0].font.name = 'Calibri'
                        hdr_colum.paragraphs[0].runs[0].font.size = Pt(10.5)
                        hdr_colum.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                        hdr_colum.paragraphs[0].runs[0].font.bold = True
                        hdr_colum.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    hdr_colum = table_b.rows[1].cells

                    if self.selected_produit=="POIS CHICHES IMP 12 mm" :
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text =self.restpoichichedouzettp + "  QX"

                    if self.selected_produit== "POIS CHICHES IMP 12 mm":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restpoichichedouzettp + "  QX"
                    if self.selected_produit== "POIS CHICHES IMP 09 mm":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restpoichicheneufttp + " QX"
                    if self.selected_produit == "POIS CHICHES IMP 08 mm":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restpoichicheouitttp + " QX"
                    if self.selected_produit== "POIS CHICHES IMP 06 mm":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restpoichichettp + " QX"
                    if self.selected_produit== "HARICOT BLANC IMP":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restharicotimpttp + " QX"
                    if self.selected_produit== "HARICOT LSB":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restharicotlsbttp + " QX"
                    if self.selected_produit== "LENTILLE PL":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restlenplttp + " QX"
                    if self.selected_produit== "LENTILLE IMP 05-07 mm":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restlentsinqttp + " QX"
                    if self.selected_produit== "LENTILLE IMP vert":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restlentvertttp + " QX"
                    if self.selected_produit== "LENTILLE IMP ROUGE":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restlentrougettp + " QX"
                    if self.selected_produit == "RIZ IMP":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.resulxttp + " QX"
                    if self.selected_produit== "RIZ ETUVE":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.resulxtttvp + " QX"
                    for row in table_b.rows[1:2]:
                        for cell in row.cells:
                            cell.paragraphs[0].runs[0].font.name = 'Calibri'
                            cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                            cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                            cell.paragraphs[0].runs[0].font.bold = False
                    # Save and open the document for printing
                    temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                    self.doc.save(temp_file)
                    os.startfile(temp_file, "open")
                    conn.close()
                elif self.selected_produit and self.selected_magasin and self.selected_origine == "" and self.selected_dates == "--" and self.selected_datesAU == "--":
                    conn = mysql.connector.connect(host="localhost",  user=user, password=password,
                                                   database="datta_legumesec_entry")
                    dattabase1 = mysql.connector.connect(host="localhost",  user=user, password=password,
                                                         database='datta_legumsec_out')
                    cursor = conn.cursor()
                    cursor.execute(
                        "SELECT le,produit,origine,quantité,unité,sacherie,fournisseur, imatricule,date_recolte,date_CND, DLUM FROM entrytable WHERE produit = %s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    rows = cursor.fetchall()
                    curs1x = dattabase1.cursor()
                    curs1x.execute(
                        "SELECT SUM(quantitérizentré) - SUM(quantitérizesorté) as quantitérizerest FROM outtable WHERE  produit= %s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    self.resulxs = curs1x.fetchone()
                    self.resulxttp = str(self.resulxs[0])
                    curs1xtv = dattabase1.cursor()
                    curs1xtv.execute(
                        "SELECT SUM(quantitérizetvntré) - SUM(quantitérizetvsorté) as quantitérizetvrest FROM outtable WHERE  produit= %s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    self.resulxstv = curs1xtv.fetchone()
                    self.resulxtttvp = str(self.resulxstv[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitepchentré) - sum(quantitépchsorté) as quantitépchrest  FROM outtable   WHERE   produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichichedouzettp = str(resultpch_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT SUM(quantitépchsixentré) - SUM(quantitépchsixsorté) as quantitépchsixrest FROM outtable WHERE   produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichichettp = str(resultpch_att[0])

                    curs5tt = dattabase1.cursor()
                    curs5tt.execute(
                        "SELECT sum(quantitépchplentré) - sum(quantitépchplsorté) as quantitépchplhrest FROM outtable WHERE   produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resultpchpl_att = curs5tt.fetchone()
                    self.restpoischicheplttp = str(resultpchpl_att[0])
                    curs7tt = dattabase1.cursor()
                    curs7tt.execute(
                        "SELECT sum(quantitélenplentré) - sum(quantitélenplsorté) as quantitélenplrest FROM outtable WHERE   produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resultlenpl_att = curs7tt.fetchone()
                    self.restlenplttp = str(resultlenpl_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenvertentré) - sum(quantitélenvertsorté) as quantitélenrest FROM outtable WHERE   produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resultlenverttt = curs9tt.fetchone()
                    self.restlentvertttp = str(resultlenverttt[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenrougeentré) - sum(quantitélenrougesorté) as quantitélenrougerest FROM outtable WHERE   produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resultlenrougett = curs9tt.fetchone()
                    self.restlentrougettp = str(resultlenrougett[0])
                    curs11tt = dattabase1.cursor()
                    curs11tt.execute(
                        "SELECT sum(quantitéharentré) - sum(quantitéharsorté) as quantitéharrest FROM outtable WHERE   produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resulthar_att = curs11tt.fetchone()
                    self.restharicotimpttp = str(resulthar_att[0])
                    curs11tt.execute(
                        "SELECT SUM(quantitéharlsbentré) - SUM(quantitéharlsbsorté) as quantitéharlsbrest FROM outtable WHERE   produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resultharlsb_att = curs11tt.fetchone()
                    self.restharicotlsbttp = str(resultharlsb_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchneufeentré) - sum(quantitépchneufesorté) as quantitépchneuferest FROM outtable   WHERE produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheneufttp = str(resultpch_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchhuitentré) - sum(quantitépchhuitsorté) as quantitépchhuitrest FROM outtable WHERE   produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheouitttp = str(resultpch_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenentré) - sum(quantitélensorté) as quantitélenrest FROM outtable WHERE   produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resultlensinqtt = curs9tt.fetchone()
                    self.restlentsinqttp = str(resultlensinqtt[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT SUM(quantitépchsixentré) - SUM(quantitépchsixsorté) as quantitépchsixrest FROM outtable  WHERE produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resultpch_attp = curstt.fetchone()
                    self.restpoichichettp = str(resultpch_attp[0])

                    curs5tt = dattabase1.cursor()
                    curs5tt.execute(
                        "SELECT sum(quantitépchplentré) - sum(quantitépchplsorté) as quantitépchplhrest FROM outtable   WHERE produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resultpchpl_att = curs5tt.fetchone()
                    self.restpoischicheplttp = str(resultpchpl_att[0])
                    curs7tt = dattabase1.cursor()
                    curs7tt.execute(
                        "SELECT sum(quantitélenplentré) - sum(quantitélenplsorté) as quantitélenplrest FROM outtable  WHERE produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resultlenpl_att = curs7tt.fetchone()
                    self.restlenplttp = str(resultlenpl_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenvertentré) - sum(quantitélenvertsorté) as quantitélenrest FROM outtable WHERE produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resultlenverttt = curs9tt.fetchone()
                    self.restlentvertttp = str(resultlenverttt[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenrougeentré) - sum(quantitélenrougesorté) as quantitélenrougerest FROM outtable WHERE produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resultlenrougett = curs9tt.fetchone()
                    self.restlentrougettp = str(resultlenrougett[0])
                    curs11tt = dattabase1.cursor()
                    curs11tt.execute(
                        "SELECT sum(quantitéharentré) - sum(quantitéharsorté) as quantitéharrest FROM outtable WHERE  produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resulthar_att = curs11tt.fetchone()
                    self.restharicotimpttp = str(resulthar_att[0])
                    curs11tt.execute(
                        "SELECT SUM(quantitéharlsbentré) - SUM(quantitéharlsbsorté) as quantitéharlsbrest FROM outtable   WHERE produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resultharlsb_att = curs11tt.fetchone()
                    self.restharicotlsbttp = str(resultharlsb_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchneufeentré) - sum(quantitépchneufesorté) as quantitépchneuferest FROM outtable  WHERE produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheneufttp = str(resultpch_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchhuitentré) - sum(quantitépchhuitsorté) as quantitépchhuitrest FROM outtable WHERE produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheouitttp = str(resultpch_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenentré) - sum(quantitélensorté) as quantitélenrest FROM outtable WHERE  produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resultlensinqtt = curs9tt.fetchone()
                    self.restlentsinqttp = str(resultlensinqtt[0])
                    ########################################################################
                    # Create a new document and add a table
                    self.doc = docx.Document()
                    section = self.doc.sections[0]
                    section.orientation = WD_ORIENTATION.LANDSCAPE
                    section.page_width = docx.shared.Cm(29.7)
                    section.page_height = docx.shared.Cm(21.0)
                    section.top_margin = docx.shared.Cm(0.5)
                    section.bottom_margin = docx.shared.Cm(1.5)
                    section.right_margin = docx.shared.Cm(1.5)
                    section.left_margin = docx.shared.Cm(1.5)
                    heading1 = self.doc.add_heading(
                        '\t\t\t\t\t\tSITUATION DES STOCKS DES LEGUMES SECS\n\t\t\t\t\t\t\t    PRODUITS RÉCEPTIONNER',
                        level=1)
                    heading1.style.font.name = 'Times New Roman'
                    heading1.style.font.size = Pt(14)
                    heading1.style.font.bold = True
                    heading1.style.font.color.rgb = RGBColor(0, 0, 0)
                    # add second heading

                    paragraph1 = self.doc.add_paragraph()
                    paragraph1.style.font.name = 'Times New Roman'
                    current_date = QDate.currentDate().getDate()
                    left_run1 = paragraph1.add_run(
                        "\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t Date d'édition : {}/{}/{}".format(current_date[2],
                                                                                                current_date[1],
                                                                                                current_date[0]))
                    left_run1.font.size = Pt(9)

                    paragraph = self.doc.add_paragraph()
                    paragraph.style.font.name = 'Times New Roman'
                    paragraph.style.font.size = Pt(12)
                    left_run = paragraph.add_run(self.selected_magasin+"\nPRODUIT :" + self.selected_produit + '\t\t\t\t\t\t\t\t\t\t\t\t\t')
                    left_run.bold = True

                    table = self.doc.add_table(rows=1, cols=11)
                    table.style = "Table Grid"  # set the table style
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'DATE'
                    hdr_cells[1].text = 'PRODUIT'
                    hdr_cells[2].text = 'ORIGINE'
                    hdr_cells[3].text = 'QUANTITE'
                    hdr_cells[4].text = 'MAGASIN DE STOCK'
                    hdr_cells[5].text = 'SASHERIE  CND'
                    hdr_cells[6].text = 'EXPIDITEUR'
                    hdr_cells[7].text = 'IMATRICUL'
                    hdr_cells[8].text = 'Date de récolte'
                    hdr_cells[9].text = 'Date de CND'
                    hdr_cells[10].text = 'Date de DLUM'
                    # Set the width of the header cells
                    hdr_cells[0].width = Inches(0.5)  # set the width of the first column to 1.5 inches
                    hdr_cells[1].width = Inches(2)
                    hdr_cells[2].width = Inches(1)
                    hdr_cells[3].width = Inches(0.4)
                    hdr_cells[4].width = Inches(1.3)
                    hdr_cells[5].width = Inches(1)
                    hdr_cells[6].width = Inches(1)
                    hdr_cells[7].width = Inches(1)
                    hdr_cells[8].width = Inches(0.5)
                    hdr_cells[9].width = Inches(0.5)
                    hdr_cells[10].width = Inches(0.5)

                    for cell in table.rows[0].cells:
                        cell.paragraphs[0].runs[0].font.name = 'Calibri'
                        cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                        cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                        cell.paragraphs[0].runs[0].font.bold = True
                        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    # Add the data to the table
                    table_rows = len(rows)
                    table_cols = len(rows[0])
                    for row in range(table_rows):
                        table.add_row()
                        for col in range(table_cols):
                            cell = table.cell(row + 1, col)
                            cell.text = str(rows[row][col])
                            # Set font properties
                            cell.paragraphs[0].runs[0].font.name = 'Calibri'
                            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
                            cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                            cell.paragraphs[0].runs[0].font.bold = False
                            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    for row in table.rows:
                        row.height = Inches(0.3)
                    # Create the document
                    # Create the table
                    self.doc.add_paragraph().add_run().add_break()
                    table_b = self.doc.add_table(rows=2, cols=2)
                    table_b.style = "Table Grid"
                    # Add table headers
                    hdr_colum = table_b.rows[0].cells
                    hdr_colum[0].text = "PRODUIT"
                    hdr_colum[1].text = "STOCK"
                    for hdr_colum in table_b.rows[0].cells:
                        hdr_colum.paragraphs[0].runs[0].font.name = 'Calibri'
                        hdr_colum.paragraphs[0].runs[0].font.size = Pt(10.5)
                        hdr_colum.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                        hdr_colum.paragraphs[0].runs[0].font.bold = True
                        hdr_colum.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    hdr_colum = table_b.rows[1].cells

                    if self.selected_produit == "POIS CHICHES IMP 12 mm":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restpoichichedouzettp + "  QX"

                    if self.selected_produit == "POIS CHICHES IMP 12 mm":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restpoichichedouzettp + "  QX"
                    if self.selected_produit == "POIS CHICHES IMP 09 mm":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restpoichicheneufttp + " QX"
                    if self.selected_produit == "POIS CHICHES IMP 08 mm":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restpoichicheouitttp + " QX"
                    if self.selected_produit == "POIS CHICHES IMP 06 mm":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restpoichichettp + " QX"
                    if self.selected_produit == "HARICOT BLANC IMP":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restharicotimpttp + " QX"
                    if self.selected_produit == "HARICOT LSB":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restharicotlsbttp + " QX"
                    if self.selected_produit == "LENTILLE PL":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restlenplttp + " QX"
                    if self.selected_produit == "LENTILLE IMP 05-07 mm":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restlentsinqttp + " QX"
                    if self.selected_produit == "LENTILLE IMP vert":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restlentvertttp + " QX"
                    if self.selected_produit == "LENTILLE IMP ROUGE":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restlentrougettp + " QX"
                    if self.selected_produit == "RIZ IMP":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.resulxttp + " QX"
                    if self.selected_produit == "RIZ ETUVE":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.resulxtttvp + " QX"
                    for row in table_b.rows[1:2]:
                        for cell in row.cells:
                            cell.paragraphs[0].runs[0].font.name = 'Calibri'
                            cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                            cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                            cell.paragraphs[0].runs[0].font.bold = False
                    # Save and open the document for printing
                    temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                    self.doc.save(temp_file)
                    os.startfile(temp_file, "open")
                    conn.close()
                elif self.selected_produit=="" and self.selected_magasin=="" and self.selected_origine and self.selected_dates=="--" and self.selected_datesAU=="--":
                    conn = mysql.connector.connect(host="localhost",  user=user, password=password,
                                                   database="datta_legumesec_entry")
                    dattabase1 = mysql.connector.connect(host="localhost",  user=user, password=password,
                                                         database='datta_legumsec_out')
                    cursor = conn.cursor()
                    cursor.execute(
                        "SELECT le,produit,origine,quantité,produit,sacherie,fournisseur, imatricule,date_recolte,date_CND, DLUM FROM entrytable WHERE origine=%s",
                        (self.selected_origine,))
                    rows = cursor.fetchall()
                    curs1x = dattabase1.cursor()
                    curs1x.execute(
                        "SELECT SUM(quantitérizentré) - SUM(quantitérizesorté) as quantitérizerest FROM outtable WHERE origine=%s",
                        (self.selected_origine,))
                    self.resulxs = curs1x.fetchone()
                    self.resulxtto = str(self.resulxs[0])
                    curs1xtv = dattabase1.cursor()
                    curs1xtv.execute(
                        "SELECT SUM(quantitérizetvntré) - SUM(quantitérizetvsorté) as quantitérizetvrest FROM outtable WHERE origine=%s",
                        (self.selected_origine,))
                    self.resulxstv = curs1xtv.fetchone()
                    self.resulxtttvo = str(self.resulxstv[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitepchentré) - sum(quantitépchsorté) as quantitépchrest  FROM outtable   WHERE  origine=%s",
                        (self.selected_origine,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichichedouzetto = str(resultpch_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT SUM(quantitépchsixentré) - SUM(quantitépchsixsorté) as quantitépchsixrest FROM outtable WHERE  origine=%s",
                        (self.selected_origine,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichichetto = str(resultpch_att[0])

                    curs5tt = dattabase1.cursor()
                    curs5tt.execute(
                        "SELECT sum(quantitépchplentré) - sum(quantitépchplsorté) as quantitépchplhrest FROM outtable WHERE  origine=%s",
                        (self.selected_origine,))
                    resultpchpl_att = curs5tt.fetchone()
                    self.restpoischichepltto = str(resultpchpl_att[0])
                    curs7tt = dattabase1.cursor()
                    curs7tt.execute(
                        "SELECT sum(quantitélenplentré) - sum(quantitélenplsorté) as quantitélenplrest FROM outtable WHERE  origine=%s",
                        (self.selected_origine,))
                    resultlenpl_att = curs7tt.fetchone()
                    self.restlenpltto = str(resultlenpl_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenvertentré) - sum(quantitélenvertsorté) as quantitélenrest FROM outtable WHERE  origine=%s",
                        (self.selected_origine,))
                    resultlenverttt = curs9tt.fetchone()
                    self.restlentverttto = str(resultlenverttt[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenrougeentré) - sum(quantitélenrougesorté) as quantitélenrougerest FROM outtable WHERE  origine=%s",
                        (self.selected_origine,))
                    resultlenrougett = curs9tt.fetchone()
                    self.restlentrougetto = str(resultlenrougett[0])
                    curs11tt = dattabase1.cursor()
                    curs11tt.execute(
                        "SELECT sum(quantitéharentré) - sum(quantitéharsorté) as quantitéharrest FROM outtable WHERE  origine=%s",
                        (self.selected_origine,))
                    resulthar_att = curs11tt.fetchone()
                    self.restharicotimptto = str(resulthar_att[0])
                    curs11tt.execute(
                        "SELECT SUM(quantitéharlsbentré) - SUM(quantitéharlsbsorté) as quantitéharlsbrest FROM outtable WHERE  origine=%s",
                        (self.selected_origine,))
                    resultharlsb_att = curs11tt.fetchone()
                    self.restharicotlsbtto = str(resultharlsb_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchneufeentré) - sum(quantitépchneufesorté) as quantitépchneuferest FROM outtable  WHERE origine=%s",
                        (self.selected_origine,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheneuftto = str(resultpch_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchhuitentré) - sum(quantitépchhuitsorté) as quantitépchhuitrest FROM outtable WHERE  origine=%s",
                        (self.selected_origine,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheouittto = str(resultpch_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenentré) - sum(quantitélensorté) as quantitélenrest FROM outtable WHERE  origine=%s",
                        (self.selected_origine,))
                    resultlensinqtt = curs9tt.fetchone()
                    self.restlentsinqtto = str(resultlensinqtt[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT SUM(quantitépchsixentré) - SUM(quantitépchsixsorté) as quantitépchsixrest FROM outtable  WHERE origine=%s",
                        (self.selected_origine,))
                    resultpch_atto = curstt.fetchone()
                    self.restpoichichett = str(resultpch_att[0])

                    curs5tt = dattabase1.cursor()
                    curs5tt.execute(
                        "SELECT sum(quantitépchplentré) - sum(quantitépchplsorté) as quantitépchplhrest FROM outtable   WHERE origine=%s",
                        (self.selected_origine,))
                    resultpchpl_att = curs5tt.fetchone()
                    self.restpoischichepltto = str(resultpchpl_att[0])
                    curs7tt = dattabase1.cursor()
                    curs7tt.execute(
                        "SELECT sum(quantitélenplentré) - sum(quantitélenplsorté) as quantitélenplrest FROM outtable  WHERE origine=%s",
                        (self.selected_origine,))
                    resultlenpl_att = curs7tt.fetchone()
                    self.restlenpltto = str(resultlenpl_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenvertentré) - sum(quantitélenvertsorté) as quantitélenrest FROM outtable WHERE origine=%s",
                        (self.selected_origine,))
                    resultlenverttt = curs9tt.fetchone()
                    self.restlentverttto = str(resultlenverttt[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenrougeentré) - sum(quantitélenrougesorté) as quantitélenrougerest FROM outtable WHERE origine=%s",
                        (self.selected_origine,))
                    resultlenrougett = curs9tt.fetchone()
                    self.restlentrougetto = str(resultlenrougett[0])
                    curs11tt = dattabase1.cursor()
                    curs11tt.execute(
                        "SELECT sum(quantitéharentré) - sum(quantitéharsorté) as quantitéharrest FROM outtable WHERE origine=%s",
                        (self.selected_origine,))
                    resulthar_att = curs11tt.fetchone()
                    self.restharicotimptt = str(resulthar_att[0])
                    curs11tt.execute(
                        "SELECT SUM(quantitéharlsbentré) - SUM(quantitéharlsbsorté) as quantitéharlsbrest FROM outtable   WHERE origine=%s",
                        (self.selected_origine,))
                    resultharlsb_att = curs11tt.fetchone()
                    self.restharicotlsbtto = str(resultharlsb_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchneufeentré) - sum(quantitépchneufesorté) as quantitépchneuferest FROM outtable  WHERE origine=%s",
                        (self.selected_origine,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheneuftto = str(resultpch_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchhuitentré) - sum(quantitépchhuitsorté) as quantitépchhuitrest FROM outtable WHERE origine=%s",
                        (self.selected_origine,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheouittto = str(resultpch_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenentré) - sum(quantitélensorté) as quantitélenrest FROM outtable WHERE origine=%s",
                        (self.selected_origine,))
                    resultlensinqtt = curs9tt.fetchone()
                    self.restlentsinqtto = str(resultlensinqtt[0])
                    ########################################################################
                    # Create a new document and add a table
                    self.doc = docx.Document()
                    section = self.doc.sections[0]
                    section.orientation = WD_ORIENTATION.LANDSCAPE
                    section.page_width = docx.shared.Cm(29.7)
                    section.page_height = docx.shared.Cm(21.0)
                    section.top_margin = docx.shared.Cm(0.5)
                    section.bottom_margin = docx.shared.Cm(1.5)
                    section.right_margin = docx.shared.Cm(1.5)
                    section.left_margin = docx.shared.Cm(1.5)
                    heading1 = self.doc.add_heading(
                        '\t\t\t\t\t\tSITUATION DES STOCKS DES LEGUMES SECS\n\t\t\t\t\t\t\t    PRODUITS RÉCEPTIONNER', level=1)
                    heading1.style.font.name = 'Times New Roman'
                    heading1.style.font.size = Pt(14)
                    heading1.style.font.bold = True
                    heading1.style.font.color.rgb = RGBColor(0, 0, 0)
                    # add second heading


                    paragraph1 = self.doc.add_paragraph()
                    paragraph1.style.font.name = 'Times New Roman'
                    current_date = QDate.currentDate().getDate()
                    left_run1 = paragraph1.add_run(
                        "\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t Date d'édition : {}/{}/{}".format(current_date[2],
                                                                                                current_date[1],
                                                                                                current_date[0]))
                    left_run1.font.size = Pt(9)

                    paragraph = self.doc.add_paragraph()
                    paragraph.style.font.name = 'Times New Roman'
                    paragraph.style.font.size = Pt(12)
                    left_run = paragraph.add_run(self.selected_magasin + '\t\t\t\t\t\t\t\t\t\t\t\t\t')
                    left_run.bold = True

                    table = self.doc.add_table(rows=1, cols=11)
                    table.style = "Table Grid"  # set the table style
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'DATE'
                    hdr_cells[1].text = 'PRODUIT'
                    hdr_cells[2].text = 'ORIGINE'
                    hdr_cells[3].text = 'QUANTITE'
                    hdr_cells[4].text = 'MAGASIN DE STOCK'
                    hdr_cells[5].text = 'SASHERIE  CND'
                    hdr_cells[6].text = 'EXPIDITEUR'
                    hdr_cells[7].text = 'IMATRICUL'
                    hdr_cells[8].text = 'Date de récolte'
                    hdr_cells[9].text = 'Date de CND'
                    hdr_cells[10].text = 'Date de DLUM'
                    # Set the width of the header cells
                    hdr_cells[0].width = Inches(0.5)  # set the width of the first column to 1.5 inches
                    hdr_cells[1].width = Inches(2)
                    hdr_cells[2].width = Inches(1)
                    hdr_cells[3].width = Inches(0.4)
                    hdr_cells[4].width = Inches(1.3)
                    hdr_cells[5].width = Inches(1)
                    hdr_cells[6].width = Inches(1)
                    hdr_cells[7].width = Inches(1)
                    hdr_cells[8].width = Inches(0.5)
                    hdr_cells[9].width = Inches(0.5)
                    hdr_cells[10].width = Inches(0.5)

                    for cell in table.rows[0].cells:
                        cell.paragraphs[0].runs[0].font.name = 'Calibri'
                        cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                        cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                        cell.paragraphs[0].runs[0].font.bold = True
                        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    # Add the data to the table
                    table_rows = len(rows)
                    table_cols = len(rows[0])
                    for row in range(table_rows):
                        table.add_row()
                        for col in range(table_cols):
                            cell = table.cell(row + 1, col)
                            cell.text = str(rows[row][col])
                            # Set font properties
                            cell.paragraphs[0].runs[0].font.name = 'Calibri'
                            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
                            cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                            cell.paragraphs[0].runs[0].font.bold = False
                            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    for row in table.rows:
                        row.height = Inches(0.3)
                    # Create the document
                    # Create the table
                    self.doc.add_paragraph().add_run().add_break()
                    table_b = self.doc.add_table(rows=14, cols=2)
                    table_b.style = "Table Grid"
                    # Add table headers
                    hdr_colum = table_b.rows[0].cells
                    hdr_colum[0].text = "PRODUIT"
                    hdr_colum[1].text = "STOCK"
                    for hdr_colum in table_b.rows[0].cells:
                        hdr_colum.paragraphs[0].runs[0].font.name = 'Calibri'
                        hdr_colum.paragraphs[0].runs[0].font.size = Pt(10.5)
                        hdr_colum.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                        hdr_colum.paragraphs[0].runs[0].font.bold = True
                        hdr_colum.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    hdr_colum = table_b.rows[1].cells
                    hdr_colum[0].text = "POIS CHICHES PL"
                    hdr_colum[1].text = ""
                    hdr_colum = table_b.rows[2].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 12 mm"
                    hdr_colum[1].text = self.restpoichichedouzetto + "  QX"
                    hdr_colum = table_b.rows[3].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 09 mm"
                    hdr_colum[1].text = self.restpoichicheneuftto + " QX"
                    hdr_colum = table_b.rows[4].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 08 mm"
                    hdr_colum[1].text = self.restpoichicheouittto + " QX"
                    hdr_colum = table_b.rows[5].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 06 mm"
                    hdr_colum[1].text = self.restpoichichetto + " QX"
                    hdr_colum = table_b.rows[6].cells
                    hdr_colum[0].text = "HARICOT BLANC IMP"
                    hdr_colum[1].text = self.restharicotimptto + " QX"
                    hdr_colum = table_b.rows[7].cells
                    hdr_colum[0].text = "HARICOT LSB"
                    hdr_colum[1].text = self.restharicotlsbtto + " QX"
                    hdr_colum = table_b.rows[8].cells
                    hdr_colum[0].text = "LENTILLE PL"
                    hdr_colum[1].text = self.restlenpltto + " QX"
                    hdr_colum = table_b.rows[9].cells
                    hdr_colum[0].text = "LENTILLE IMP 05-07 mm"
                    hdr_colum[1].text = self.restlentsinqtto + " QX"
                    hdr_colum = table_b.rows[10].cells
                    hdr_colum[0].text = "LENTILLE IMP vert"
                    hdr_colum[1].text = self.restlentverttto + " QX"
                    hdr_colum = table_b.rows[11].cells
                    hdr_colum[0].text = "LENTILLE IMP ROUGE"
                    hdr_colum[1].text = self.restlentrougetto + " QX"
                    hdr_colum = table_b.rows[12].cells
                    hdr_colum[0].text = "RIZ IMP"
                    hdr_colum[1].text = self.resulxtto + " QX"
                    hdr_colum = table_b.rows[13].cells
                    hdr_colum[0].text = "RIZ ETUVE"
                    hdr_colum[1].text = self.resulxtto + " QX"
                    for row in table_b.rows[1:14]:
                        for cell in row.cells:
                            cell.paragraphs[0].runs[0].font.name = 'Calibri'
                            cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                            cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                            cell.paragraphs[0].runs[0].font.bold = False
                    # Save and open the document for printing
                    temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                    self.doc.save(temp_file)
                    os.startfile(temp_file, "open")
                    conn.close()
                else:
                    print("hello world")
                    conn = mysql.connector.connect(host="localhost",  user=user, password=password,
                                                   database="datta_legumesec_entry")
                    dattabase1 = mysql.connector.connect(host="localhost",  user=user, password=password,
                                                         database='datta_legumsec_out')
                    cursor = conn.cursor()
                    cursor.execute(
                        "SELECT le,produit,origine,quantité,produit,sacherie,fournisseur, imatricule, date_CND, DLUM FROM entrytable WHERE fournisseur=%s",
                        (self.selected_expediteur,))
                    rows = cursor.fetchall()
                    curs1x = dattabase1.cursor()
                    curs1x.execute(
                        "SELECT SUM(quantitérizentré) - SUM(quantitérizesorté) as quantitérizerest FROM outtable WHERE fournisseur=%s",
                        (self.selected_expediteur,))
                    self.resulxs = curs1x.fetchone()
                    self.resulxttex = str(self.resulxs[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitepchentré) - sum(quantitépchsorté) as quantitépchrest  FROM outtable   WHERE  fournisseur=%s",
                        (self.selected_expediteur,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichichedouzettex = str(resultpch_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT SUM(quantitépchsixentré) - SUM(quantitépchsixsorté) as quantitépchsixrest FROM outtable WHERE  fournisseur=%s",
                        (self.selected_expediteur,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichichettex = str(resultpch_att[0])

                    curs5tt = dattabase1.cursor()
                    curs5tt.execute(
                        "SELECT sum(quantitépchplentré) - sum(quantitépchplsorté) as quantitépchplhrest FROM outtable WHERE  fournisseur=%s",
                        (self.selected_expediteur,))
                    resultpchpl_att = curs5tt.fetchone()
                    self.restpoischicheplttex = str(resultpchpl_att[0])
                    curs7tt = dattabase1.cursor()
                    curs7tt.execute(
                        "SELECT sum(quantitélenplentré) - sum(quantitélenplsorté) as quantitélenplrest FROM outtable WHERE  fournisseur=%s",
                        (self.selected_expediteur,))
                    resultlenpl_att = curs7tt.fetchone()
                    self.restlenplttex = str(resultlenpl_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenvertentré) - sum(quantitélenvertsorté) as quantitélenrest FROM outtable WHERE  fournisseur=%s",
                        (self.selected_expediteur,))
                    resultlenverttt = curs9tt.fetchone()
                    self.restlentvertttex = str(resultlenverttt[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenrougeentré) - sum(quantitélenrougesorté) as quantitélenrougerest FROM outtable WHERE  fournisseur=%s",
                        (self.selected_expediteur,))
                    resultlenrougett = curs9tt.fetchone()
                    self.restlentrougettex = str(resultlenrougett[0])
                    curs11tt = dattabase1.cursor()
                    curs11tt.execute(
                        "SELECT sum(quantitéharentré) - sum(quantitéharsorté) as quantitéharrest FROM outtable WHERE  fournisseur=%s",
                        (self.selected_expediteur,))
                    resulthar_att = curs11tt.fetchone()
                    self.restharicotimpttex = str(resulthar_att[0])
                    curs11tt.execute(
                        "SELECT SUM(quantitéharlsbentré) - SUM(quantitéharlsbsorté) as quantitéharlsbrest FROM outtable WHERE  fournisseur=%s",
                        (self.selected_expediteur,))
                    resultharlsb_att = curs11tt.fetchone()
                    self.restharicotlsbttex = str(resultharlsb_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchneufeentré) - sum(quantitépchneufesorté) as quantitépchneuferest FROM outtable  WHERE fournisseur=%s",
                        (self.selected_expediteur,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheneufttex = str(resultpch_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchhuitentré) - sum(quantitépchhuitsorté) as quantitépchhuitrest FROM outtable WHERE  fournisseur=%s",
                        (self.selected_expediteur,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheouitttex = str(resultpch_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenentré) - sum(quantitélensorté) as quantitélenrest FROM outtable WHERE  fournisseur=%s",
                        (self.selected_expediteur,))
                    resultlensinqtt = curs9tt.fetchone()
                    self.restlentsinqttex = str(resultlensinqtt[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT SUM(quantitépchsixentré) - SUM(quantitépchsixsorté) as quantitépchsixrest FROM outtable  WHERE fournisseur=%s",
                        (self.selected_expediteur,))
                    resultpch_attex = curstt.fetchone()
                    self.restpoichichett = str(resultpch_att[0])

                    curs5tt = dattabase1.cursor()
                    curs5tt.execute(
                        "SELECT sum(quantitépchplentré) - sum(quantitépchplsorté) as quantitépchplhrest FROM outtable   WHERE fournisseur=%s",
                        (self.selected_expediteur,))
                    resultpchpl_att = curs5tt.fetchone()
                    self.restpoischicheplttex = str(resultpchpl_att[0])
                    curs7tt = dattabase1.cursor()
                    curs7tt.execute(
                        "SELECT sum(quantitélenplentré) - sum(quantitélenplsorté) as quantitélenplrest FROM outtable  WHERE fournisseur=%s",
                        (self.selected_expediteur,))
                    resultlenpl_att = curs7tt.fetchone()
                    self.restlenplttex = str(resultlenpl_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenvertentré) - sum(quantitélenvertsorté) as quantitélenrest FROM outtable WHERE fournisseur=%s",
                        (self.selected_expediteur,))
                    resultlenverttt = curs9tt.fetchone()
                    self.restlentvertttex = str(resultlenverttt[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenrougeentré) - sum(quantitélenrougesorté) as quantitélenrougerest FROM outtable WHERE fournisseur=%s",
                        (self.selected_expediteur,))
                    resultlenrougett = curs9tt.fetchone()
                    self.restlentrougettex = str(resultlenrougett[0])
                    curs11tt = dattabase1.cursor()
                    curs11tt.execute(
                        "SELECT sum(quantitéharentré) - sum(quantitéharsorté) as quantitéharrest FROM outtable WHERE fournisseur=%s",
                        (self.selected_expediteur,))
                    resulthar_att = curs11tt.fetchone()
                    self.restharicotimptt = str(resulthar_att[0])
                    curs11tt.execute(
                        "SELECT SUM(quantitéharlsbentré) - SUM(quantitéharlsbsorté) as quantitéharlsbrest FROM outtable   WHERE fournisseur=%s",
                        (self.selected_expediteur,))
                    resultharlsb_att = curs11tt.fetchone()
                    self.restharicotlsbttex = str(resultharlsb_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchneufeentré) - sum(quantitépchneufesorté) as quantitépchneuferest FROM outtable  WHERE fournisseur=%s",
                        (self.selected_expediteur,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheneufttex = str(resultpch_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchhuitentré) - sum(quantitépchhuitsorté) as quantitépchhuitrest FROM outtable WHERE fournisseur=%s",
                        (self.selected_expediteur,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheouitttex = str(resultpch_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenentré) - sum(quantitélensorté) as quantitélenrest FROM outtable WHERE fournisseur=%s",
                        (self.selected_expediteur,))
                    resultlensinqtt = curs9tt.fetchone()
                    self.restlentsinqttex = str(resultlensinqtt[0])
                    ########################################################################
                    # Create a new document and add a table
                    self.doc = docx.Document()
                    section = self.doc.sections[0]
                    section.orientation = WD_ORIENTATION.LANDSCAPE
                    section.page_width = docx.shared.Cm(29.7)
                    section.page_height = docx.shared.Cm(21.0)
                    section.top_margin = docx.shared.Cm(0.5)
                    section.bottom_margin = docx.shared.Cm(1.5)
                    section.right_margin = docx.shared.Cm(1.5)
                    section.left_margin = docx.shared.Cm(1.5)
                    heading1 = self.doc.add_heading(
                        '\t\t\t\t\t\tSITUATION DES STOCKS DES LEGUMES SECS', level=1)
                    heading1.style.font.name = 'Times New Roman'
                    heading1.style.font.size = Pt(18)
                    heading1.style.font.bold = True
                    heading1.style.font.color.rgb = RGBColor(0, 0, 0)
                    # add second heading
                    paragraph0 = self.doc.add_paragraph()
                    paragraph0.style.font.name = 'Times New Roman'
                    left_run0 = paragraph0.add_run(
                        'Période du :' + self.selected_dates + "  au:" + self.selected_datesAU)
                    paragraph0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    paragraph1 = self.doc.add_paragraph()
                    paragraph1.style.font.name = 'Times New Roman'
                    current_date = QDate.currentDate().getDate()
                    left_run1 = paragraph1.add_run(
                        "\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t Date d'édition : {}/{}/{}".format(current_date[2],
                                                                                                current_date[1],
                                                                                                current_date[0]))
                    left_run1.font.size = Pt(9)

                    paragraph = self.doc.add_paragraph()
                    paragraph.style.font.name = 'Times New Roman'
                    paragraph.style.font.size = Pt(12)
                    left_run = paragraph.add_run(self.selected_magasin + '\t\t\t\t\t\t\t\t\t\t\t\t\t')
                    left_run.bold = True

                    table = self.doc.add_table(rows=1, cols=11)
                    table.style = "Table Grid"  # set the table style
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'DATE'
                    hdr_cells[1].text = 'PRODUIT'
                    hdr_cells[2].text = 'ORIGINE'
                    hdr_cells[3].text = 'QUANTITE'
                    hdr_cells[4].text = 'MAGASIN DE STOCK'
                    hdr_cells[5].text = 'SASHERIE  CND'
                    hdr_cells[6].text = 'EXPIDITEUR'
                    hdr_cells[7].text = 'IMATRICUL'
                    hdr_cells[8].text = 'Date de récolte'
                    hdr_cells[9].text = 'Date de CND'
                    hdr_cells[10].text = 'Date de DLUM'
                    # Set the width of the header cells
                    hdr_cells[0].width = Inches(0.5)  # set the width of the first column to 1.5 inches
                    hdr_cells[1].width = Inches(2)
                    hdr_cells[2].width = Inches(1)
                    hdr_cells[3].width = Inches(0.4)
                    hdr_cells[4].width = Inches(1.3)
                    hdr_cells[5].width = Inches(1)
                    hdr_cells[6].width = Inches(1)
                    hdr_cells[7].width = Inches(1)
                    hdr_cells[8].width = Inches(0.5)
                    hdr_cells[9].width = Inches(0.5)
                    hdr_cells[10].width = Inches(0.5)

                    for cell in table.rows[0].cells:
                        cell.paragraphs[0].runs[0].font.name = 'Calibri'
                        cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                        cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                        cell.paragraphs[0].runs[0].font.bold = True
                        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    # Add the data to the table
                    table_rows = len(rows)
                    table_cols = len(rows[0])
                    for row in range(table_rows):
                        table.add_row()
                        for col in range(table_cols):
                            cell = table.cell(row + 1, col)
                            cell.text = str(rows[row][col])
                            # Set font properties
                            cell.paragraphs[0].runs[0].font.name = 'Calibri'
                            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
                            cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                            cell.paragraphs[0].runs[0].font.bold = False
                            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    for row in table.rows:
                        row.height = Inches(0.3)
                    # Create the document
                    # Create the table
                    self.doc.add_paragraph().add_run().add_break()
                    table_b = self.doc.add_table(rows=14, cols=2)
                    table_b.style = "Table Grid"
                    # Add table headers
                    hdr_colum = table_b.rows[0].cells
                    hdr_colum[0].text = "PRODUIT"
                    hdr_colum[1].text = "STOCK"
                    for hdr_colum in table_b.rows[0].cells:
                        hdr_colum.paragraphs[0].runs[0].font.name = 'Calibri'
                        hdr_colum.paragraphs[0].runs[0].font.size = Pt(10.5)
                        hdr_colum.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                        hdr_colum.paragraphs[0].runs[0].font.bold = True
                        hdr_colum.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    hdr_colum = table_b.rows[1].cells
                    hdr_colum[0].text = "POIS CHICHES PL"
                    hdr_colum[1].text = ""
                    hdr_colum = table_b.rows[2].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 12 mm"
                    hdr_colum[1].text = self.restpoichichedouzetto + "  QX"
                    hdr_colum = table_b.rows[3].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 09 mm"
                    hdr_colum[1].text = self.restpoichicheneuftto + " QX"
                    hdr_colum = table_b.rows[4].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 08 mm"
                    hdr_colum[1].text = self.restpoichicheouittto + " QX"
                    hdr_colum = table_b.rows[5].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 06 mm"
                    hdr_colum[1].text = self.restpoichichetto + " QX"
                    hdr_colum = table_b.rows[6].cells
                    hdr_colum[0].text = "HARICOT BLANC IMP"
                    hdr_colum[1].text = self.restharicotimptto + " QX"
                    hdr_colum = table_b.rows[7].cells
                    hdr_colum[0].text = "HARICOT LSB"
                    hdr_colum[1].text = self.restharicotlsbtto + " QX"
                    hdr_colum = table_b.rows[8].cells
                    hdr_colum[0].text = "LENTILLE PL"
                    hdr_colum[1].text = self.restlenpltto + " QX"
                    hdr_colum = table_b.rows[9].cells
                    hdr_colum[0].text = "LENTILLE IMP 05-07 mm"
                    hdr_colum[1].text = self.restlentsinqtto + " QX"
                    hdr_colum = table_b.rows[10].cells
                    hdr_colum[0].text = "LENTILLE IMP vert"
                    hdr_colum[1].text = self.restlentverttto + " QX"
                    hdr_colum = table_b.rows[11].cells
                    hdr_colum[0].text = "LENTILLE IMP ROUGE"
                    hdr_colum[1].text = self.restlentrougetto + " QX"
                    hdr_colum = table_b.rows[12].cells
                    hdr_colum[0].text = "RIZ IMP"
                    hdr_colum[1].text = self.resulxtto + " QX"
                    hdr_colum = table_b.rows[13].cells
                    hdr_colum[0].text = "RIZ IMP"
                    hdr_colum[1].text = self.resulxtto + " QX"
                    for row in table_b.rows[1:14]:
                        for cell in row.cells:
                            cell.paragraphs[0].runs[0].font.name = 'Calibri'
                            cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                            cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                            cell.paragraphs[0].runs[0].font.bold = False
                    # Save and open the document for printing
                    temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                    self.doc.save(temp_file)
                    os.startfile(temp_file, "open")
                    conn.close()

            except Exception as e:
                print(e)

        def print_docxStockDatabaseSortie(self):
            self.poischicheDouze()
            self.poischicheNeuf()
            self.poischicheOuit()
            self.poichichSIxSept()
            self.operation_dattabase()
            self.haricot()
            self.haricotLsb()
            self.lentilleCinqSept()
            self.lenVert()
            self.lenRouge()
            self.lenpl()
            self.poichichepl()
            self.rizeEtuve()
            self.selected_produit = self.Produits.currentText()
            self.selected_magasin = self.Unites.currentText()
            self.selected_origine=self.OrigineDuProduits.currentText()
            self.selected_expediteur=self.expediteur.currentText()
            self.selected_dates = self.date_edits.text()
            self.selected_datesAU = self.date_edit2.text()
            try:
                if self.selected_produit and self.selected_magasin and self.selected_origine and self.selected_dates=="--" and self.selected_datesAU=="--":
                    conn = mysql.connector.connect(host="localhost",  user=user, password=password,
                                                   database="datta_legumsec_out")
                    cursor = conn.cursor()
                    cursor.execute(
                        "SELECT le,produit,origine,quantité,unité,sacherie,fournisseur, imatricule,date_recoltes, date_CND, DLUM FROM outtable WHERE produit = %s AND unité = %s AND fournisseur=%s  AND date_recoltes IS NOT NULL AND date_CND IS NOT NULL AND DLUM IS NOT NULL ",
                        (self.selected_produit, self.selected_magasin,self.selected_expediteur,)
                    )
                    rows = cursor.fetchall()

                    ########################################################################
                    # Create a new document and add a table
                    self.doc = docx.Document()
                    section = self.doc.sections[0]
                    section.orientation = WD_ORIENTATION.LANDSCAPE
                    section.page_width = docx.shared.Cm(29.7)
                    section.page_height = docx.shared.Cm(21.0)
                    section.top_margin = docx.shared.Cm(0.5)
                    section.bottom_margin = docx.shared.Cm(1.5)
                    section.right_margin = docx.shared.Cm(1.5)
                    section.left_margin = docx.shared.Cm(1.5)
                    heading1 = self.doc.add_heading(
                        '\t\t\t\t\t\tSITUATION DES STOCKS DES LEGUMES SECS\n\t\t\t\t\t\t\t\tPRODUITS LIVRÉE',
                        level=1)
                    heading1.style.font.name = 'Times New Roman'
                    heading1.style.font.size = Pt(14)
                    heading1.style.font.bold = True
                    heading1.style.font.color.rgb = RGBColor(0, 0, 0)
                    # add second heading
                    paragraph0 = self.doc.add_paragraph()
                    paragraph0.style.font.name = 'Times New Roman'
                    left_run0 = paragraph0.add_run('Période du :' + self.selected_dates + "  au:" + self.selected_datesAU)
                    paragraph0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    paragraph1 = self.doc.add_paragraph()
                    paragraph1.style.font.name = 'Times New Roman'
                    current_date = QDate.currentDate().getDate()
                    left_run1 = paragraph1.add_run("\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t Date d'édition : {}/{}/{}".format(current_date[2], current_date[1], current_date[0]))
                    left_run1.font.size = Pt(9)

                    paragraph = self.doc.add_paragraph()
                    paragraph.style.font.name = 'Times New Roman'
                    paragraph.style.font.size = Pt(12)
                    left_run = paragraph.add_run(self.selected_magasin + '\t\t\t\t\t\t\t\t\t\t\t\t\t')
                    left_run.bold = True

                    table = self.doc.add_table(rows=1, cols=11)
                    table.style = "Table Grid"  # set the table style
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'DATE'
                    hdr_cells[1].text = 'PRODUIT'
                    hdr_cells[2].text = 'ORIGINE'
                    hdr_cells[3].text = 'QUANTITE'
                    hdr_cells[4].text = 'MAGASIN DE STOCK'
                    hdr_cells[5].text = 'SASHERIE  CND'
                    hdr_cells[6].text = 'EXPIDITEUR'
                    hdr_cells[7].text = 'IMATRICUL'
                    hdr_cells[8].text = 'Date de récolte'
                    hdr_cells[9].text = 'Date de CND'
                    hdr_cells[10].text = 'Date de DLUM'
                    # Set the width of the header cells
                    hdr_cells[0].width = Inches(0.5)  # set the width of the first column to 1.5 inches
                    hdr_cells[1].width = Inches(2)
                    hdr_cells[2].width = Inches(1)
                    hdr_cells[3].width = Inches(0.4)
                    hdr_cells[4].width = Inches(1.3)
                    hdr_cells[5].width = Inches(1)
                    hdr_cells[6].width = Inches(1)
                    hdr_cells[7].width = Inches(1)
                    hdr_cells[8].width = Inches(0.5)
                    hdr_cells[9].width = Inches(0.5)
                    hdr_cells[10].width = Inches(0.5)

                    for cell in table.rows[0].cells:
                        cell.paragraphs[0].runs[0].font.name = 'Calibri'
                        cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                        cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                        cell.paragraphs[0].runs[0].font.bold = True
                        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    # Add the data to the table
                    table_rows = len(rows)
                    table_cols = len(rows[0])
                    for row in range(table_rows):
                        table.add_row()
                        for col in range(table_cols):
                            cell = table.cell(row + 1, col)
                            cell.text = str(rows[row][col])
                            # Set font properties
                            cell.paragraphs[0].runs[0].font.name = 'Calibri'
                            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
                            cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                            cell.paragraphs[0].runs[0].font.bold = False
                            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    for row in table.rows:
                        row.height = Inches(0.3)
                    # Create the document
                    # Create the table
                    self.doc.add_paragraph().add_run().add_break()
                    table_b = self.doc.add_table(rows=14, cols=2)
                    table_b.style = "Table Grid"
                    # Add table headers
                    hdr_colum = table_b.rows[0].cells
                    hdr_colum[0].text = "PRODUIT"
                    hdr_colum[1].text = "STOCK"
                    for hdr_colum in table_b.rows[0].cells:
                        hdr_colum.paragraphs[0].runs[0].font.name = 'Calibri'
                        hdr_colum.paragraphs[0].runs[0].font.size = Pt(10.5)
                        hdr_colum.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                        hdr_colum.paragraphs[0].runs[0].font.bold = True
                        hdr_colum.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    hdr_colum = table_b.rows[1].cells
                    hdr_colum[0].text = "POIS CHICHES PL"
                    hdr_colum[1].text = ""
                    hdr_colum = table_b.rows[2].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 12 mm"
                    hdr_colum[1].text = self.restpoichichedouze + "  QX"
                    hdr_colum = table_b.rows[3].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 09 mm"
                    hdr_colum[1].text = self.restpoichicheneuf + " QX"
                    hdr_colum = table_b.rows[4].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 08 mm"
                    hdr_colum[1].text = self.restpoichicheouit + " QX"
                    hdr_colum = table_b.rows[5].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 06 mm"
                    hdr_colum[1].text = self.restpoichiche + " QX"
                    hdr_colum = table_b.rows[6].cells
                    hdr_colum[0].text = "HARICOT BLANC IMP"
                    hdr_colum[1].text = self.restharicotimp + " QX"
                    hdr_colum = table_b.rows[7].cells
                    hdr_colum[0].text = "HARICOT LSB"
                    hdr_colum[1].text = self.restharicotlsb + " QX"
                    hdr_colum = table_b.rows[8].cells
                    hdr_colum[0].text = "LENTILLE PL"
                    hdr_colum[1].text = self.restlenpl + " QX"
                    hdr_colum = table_b.rows[9].cells
                    hdr_colum[0].text = "LENTILLE IMP 05-07 mm"
                    hdr_colum[1].text = self.restlentsinq + " QX"
                    hdr_colum = table_b.rows[10].cells
                    hdr_colum[0].text = "LENTILLE IMP vert"
                    hdr_colum[1].text = self.restlentvert + " QX"
                    hdr_colum = table_b.rows[11].cells
                    hdr_colum[0].text = "LENTILLE IMP ROUGE"
                    hdr_colum[1].text = self.restlentrouge + " QX"
                    hdr_colum = table_b.rows[12].cells
                    hdr_colum[0].text = "RIZ IMP"
                    hdr_colum[1].text = self.resulxtt + " QX"
                    hdr_colum = table_b.rows[13].cells
                    hdr_colum[0].text = "RIZ ETUVE"
                    hdr_colum[1].text = self.resulxttv + " QX"
                    for row in table_b.rows[1:13]:
                        for cell in row.cells:
                            cell.paragraphs[0].runs[0].font.name = 'Calibri'
                            cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                            cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                            cell.paragraphs[0].runs[0].font.bold = False

                    # Save and open the document for printing
                    temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                    self.doc.save(temp_file)
                    os.startfile(temp_file, "open")
                    conn.close()
                elif self.selected_produit and self.selected_magasin and self.selected_origine and self.selected_dates and self.selected_datesAU:

                    conn = mysql.connector.connect(host="localhost",  user=user, password=password,
                                                   database="datta_legumsec_out")
                    cursor = conn.cursor()
                    cursor.execute(
                        "SELECT le,produit,origine,quantité,unité,sacherie,fournisseur, imatricule,date_recoltes,date_CND, DLUM FROM outtable WHERE produit = %s AND unité = %s AND fournisseur=%s AND le BETWEEN %s AND %s AND date_recoltes IS NOT NULL AND date_CND IS NOT NULL AND DLUM IS NOT NULL",
                        (self.selected_produit, self.selected_magasin, self.selected_expediteur,self.selected_dates,self.selected_datesAU))
                    rows = cursor.fetchall()

                    ########################################################################
                    # Create a new document and add a table
                    self.doc = docx.Document()
                    section = self.doc.sections[0]
                    section.orientation = WD_ORIENTATION.LANDSCAPE
                    section.page_width = docx.shared.Cm(29.7)
                    section.page_height = docx.shared.Cm(21.0)
                    section.top_margin = docx.shared.Cm(0.5)
                    section.bottom_margin = docx.shared.Cm(1.5)
                    section.right_margin = docx.shared.Cm(1.5)
                    section.left_margin = docx.shared.Cm(1.5)
                    heading1 = self.doc.add_heading(
                        '\t\t\t\t\t\tSITUATION DES STOCKS DES LEGUMES SECS', level=1)
                    heading1.style.font.name = 'Times New Roman'
                    heading1.style.font.size = Pt(18)
                    heading1.style.font.bold = True
                    heading1.style.font.color.rgb = RGBColor(0, 0, 0)
                    # add second heading
                    # add second heading
                    paragraph0 = self.doc.add_paragraph()
                    paragraph0.style.font.name = 'Times New Roman'
                    left_run0 = paragraph0.add_run(
                        'Période du :' + self.selected_dates + "  au:" + self.selected_datesAU)
                    paragraph0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    paragraph1 = self.doc.add_paragraph()
                    paragraph1.style.font.name = 'Times New Roman'
                    current_date = QDate.currentDate().getDate()
                    left_run1 = paragraph1.add_run(
                        "\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t Date d'édition : {}/{}/{}".format(current_date[2],
                                                                                                current_date[1],
                                                                                                current_date[0]))
                    left_run1.font.size = Pt(9)

                    paragraph = self.doc.add_paragraph()
                    paragraph.style.font.name = 'Times New Roman'
                    paragraph.style.font.size = Pt(12)
                    left_run = paragraph.add_run(self.selected_magasin + '\t\t\t\t\t\t\t\t\t\t\t\t\t')
                    left_run.bold = True

                    table = self.doc.add_table(rows=1, cols=11)
                    table.style = "Table Grid"  # set the table style
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'DATE'
                    hdr_cells[1].text = 'PRODUIT'
                    hdr_cells[2].text = 'ORIGINE'
                    hdr_cells[3].text = 'QUANTITE'
                    hdr_cells[4].text = 'MAGASIN DE STOCK'
                    hdr_cells[5].text = 'SASHERIE  CND'
                    hdr_cells[6].text = 'EXPIDITEUR'
                    hdr_cells[7].text = 'IMATRICUL'
                    hdr_cells[8].text = 'Date de récolte'
                    hdr_cells[9].text = 'Date de CND'
                    hdr_cells[10].text = 'Date de DLUM'
                    # Set the width of the header cells
                    hdr_cells[0].width = Inches(0.5)  # set the width of the first column to 1.5 inches
                    hdr_cells[1].width = Inches(2)
                    hdr_cells[2].width = Inches(1)
                    hdr_cells[3].width = Inches(0.4)
                    hdr_cells[4].width = Inches(1.3)
                    hdr_cells[5].width = Inches(1)
                    hdr_cells[6].width = Inches(1)
                    hdr_cells[7].width = Inches(1)
                    hdr_cells[8].width = Inches(0.5)
                    hdr_cells[9].width = Inches(0.5)
                    hdr_cells[10].width = Inches(0.5)

                    for cell in table.rows[0].cells:
                        cell.paragraphs[0].runs[0].font.name = 'Calibri'
                        cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                        cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                        cell.paragraphs[0].runs[0].font.bold = True
                        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    # Add the data to the table
                    table_rows = len(rows)
                    table_cols = len(rows[0])
                    for row in range(table_rows):
                        table.add_row()
                        for col in range(table_cols):
                            cell = table.cell(row + 1, col)
                            cell.text = str(rows[row][col])
                            # Set font properties
                            cell.paragraphs[0].runs[0].font.name = 'Calibri'
                            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
                            cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                            cell.paragraphs[0].runs[0].font.bold = False
                            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    for row in table.rows:
                        row.height = Inches(0.3)
                    # Create the document
                    # Create the table
                    self.doc.add_paragraph().add_run().add_break()
                    table_b = self.doc.add_table(rows=14, cols=2)
                    table_b.style = "Table Grid"
                    # Add table headers
                    hdr_colum = table_b.rows[0].cells
                    hdr_colum[0].text = "PRODUIT"
                    hdr_colum[1].text = "STOCK"
                    for hdr_colum in table_b.rows[0].cells:
                        hdr_colum.paragraphs[0].runs[0].font.name = 'Calibri'
                        hdr_colum.paragraphs[0].runs[0].font.size = Pt(10.5)
                        hdr_colum.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                        hdr_colum.paragraphs[0].runs[0].font.bold = True
                        hdr_colum.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    hdr_colum = table_b.rows[1].cells
                    hdr_colum[0].text = "POIS CHICHES PL"
                    hdr_colum[1].text = ""
                    hdr_colum = table_b.rows[2].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 12 mm"
                    hdr_colum[1].text = self.restpoichichedouze + "  QX"
                    hdr_colum = table_b.rows[3].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 09 mm"
                    hdr_colum[1].text = self.restpoichicheneuf + " QX"
                    hdr_colum = table_b.rows[4].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 08 mm"
                    hdr_colum[1].text = self.restpoichicheouit + " QX"
                    hdr_colum = table_b.rows[5].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 06 mm"
                    hdr_colum[1].text = self.restpoichiche + " QX"
                    hdr_colum = table_b.rows[6].cells
                    hdr_colum[0].text = "HARICOT BLANC IMP"
                    hdr_colum[1].text = self.restharicotimp + " QX"
                    hdr_colum = table_b.rows[7].cells
                    hdr_colum[0].text = "HARICOT LSB"
                    hdr_colum[1].text = self.restharicotlsb + " QX"
                    hdr_colum = table_b.rows[8].cells
                    hdr_colum[0].text = "LENTILLE PL"
                    hdr_colum[1].text = self.restlenpl + " QX"
                    hdr_colum = table_b.rows[9].cells
                    hdr_colum[0].text = "LENTILLE IMP 05-07 mm"
                    hdr_colum[1].text = self.restlentsinq + " QX"
                    hdr_colum = table_b.rows[10].cells
                    hdr_colum[0].text = "LENTILLE IMP vert"
                    hdr_colum[1].text = self.restlentvert + " QX"
                    hdr_colum = table_b.rows[11].cells
                    hdr_colum[0].text = "LENTILLE IMP ROUGE"
                    hdr_colum[1].text = self.restlentrouge + " QX"
                    hdr_colum = table_b.rows[12].cells
                    hdr_colum[0].text = "RIZ IMP"
                    hdr_colum[1].text = self.resulxtt + " QX"
                    hdr_colum = table_b.rows[13].cells
                    hdr_colum[0].text = "RIZ ETUVE"
                    hdr_colum[1].text = self.resulxtt + " QX"
                    for row in table_b.rows[1:14]:
                        for cell in row.cells:
                            cell.paragraphs[0].runs[0].font.name = 'Calibri'
                            cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                            cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                            cell.paragraphs[0].runs[0].font.bold = False

                    # Save and open the document for printing
                    temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                    self.doc.save(temp_file)
                    os.startfile(temp_file, "open")
                    conn.close()
                elif self.selected_produit == "" and self.selected_magasin == "" and self.selected_origine == "" and self.selected_expediteur == "" and self.selected_dates == "--" and self.selected_datesAU == "--":

                    conn = mysql.connector.connect(host="localhost",  user=user, password=password,
                                                   database="datta_legumsec_out")
                    cursor = conn.cursor()
                    cursor.execute(
                        "SELECT le,produit,origine,quantité,unité,sacherie,fournisseur,imatricule,date_recoltes,date_CND,DLUM FROM outtable WHERE  date_recoltes IS NOT NULL AND date_CND IS NOT NULL AND DLUM IS NOT NULL")
                    rows = cursor.fetchall()

                    ########################################################################

                    # Create a new document and add a table
                    self.doc = docx.Document()
                    section = self.doc.sections[0]
                    section.orientation = WD_ORIENTATION.LANDSCAPE
                    section.page_width = docx.shared.Cm(29.7)
                    section.page_height = docx.shared.Cm(21.0)
                    section.top_margin = docx.shared.Cm(0.5)
                    section.bottom_margin = docx.shared.Cm(1.5)
                    section.right_margin = docx.shared.Cm(1.5)
                    section.left_margin = docx.shared.Cm(1.5)
                    heading1 = self.doc.add_heading(
                        '\t\t\t\t\t\tSITUATION DES STOCKS DES LEGUMES SECS\n\t\t\t\t\t\t\t\tPRODUITS LIVRÉE',
                        level=1)
                    heading1.style.font.name = 'Times New Roman'
                    heading1.style.font.size = Pt(14)
                    heading1.style.font.bold = True
                    heading1.style.font.color.rgb = RGBColor(0, 0, 0)
                    # add second heading
                    paragraph = self.doc.add_paragraph()
                    paragraph.style.font.name = 'Times New Roman'
                    paragraph.style.font.size = Pt(9)
                    left_run = paragraph.add_run('\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t')
                    left_run.bold = True
                    current_date = QDate.currentDate().getDate()
                    right_run = paragraph.add_run(
                        "Date d'édition : {}/{}/{}".format(current_date[2], current_date[1], current_date[0]))
                    right_run.bold = False
                    table = self.doc.add_table(rows=1, cols=11)
                    table.style = "Table Grid"  # set the table style
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'DATE'
                    hdr_cells[1].text = 'PRODUIT'
                    hdr_cells[2].text = 'ORIGINE'
                    hdr_cells[3].text = 'QUANTITE'
                    hdr_cells[4].text = 'MAGASIN DE STOCK'
                    hdr_cells[5].text = 'SASHERIE  CND'
                    hdr_cells[6].text = 'FOURNISSEUR'
                    hdr_cells[7].text = 'IMATRICUL'
                    hdr_cells[8].text = 'Date de récolte'
                    hdr_cells[9].text = 'Date de CND'
                    hdr_cells[10].text = 'Date de DLUM'
                    # Set the width of the header cells
                    hdr_cells[0].width = Inches(0.5)  # set the width of the first column to 1.5 inches
                    hdr_cells[1].width = Inches(2)
                    hdr_cells[2].width = Inches(1)
                    hdr_cells[3].width = Inches(0.4)
                    hdr_cells[4].width = Inches(1.3)
                    hdr_cells[5].width = Inches(1)
                    hdr_cells[6].width = Inches(1)
                    hdr_cells[7].width = Inches(1)
                    hdr_cells[8].width = Inches(0.5)
                    hdr_cells[9].width = Inches(0.5)
                    hdr_cells[10].width = Inches(0.5)

                    for cell in table.rows[0].cells:
                        cell.paragraphs[0].runs[0].font.name = 'Calibri'
                        cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                        cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                        cell.paragraphs[0].runs[0].font.bold = True
                        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    # Add the data to the table
                    table_rows = len(rows)
                    table_cols = len(rows[0])
                    for row in range(table_rows):
                        table.add_row()
                        for col in range(table_cols):
                            cell = table.cell(row + 1, col)
                            cell.text = str(rows[row][col])
                            # Set font properties
                            cell.paragraphs[0].runs[0].font.name = 'Calibri'
                            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
                            cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                            cell.paragraphs[0].runs[0].font.bold = False
                            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    for row in table.rows:
                        row.height = Inches(0.3)
                    # Create the document
                    # Create the table
                    self.doc.add_paragraph().add_run().add_break()
                    table_b = self.doc.add_table(rows=14, cols=2)
                    table_b.style = "Table Grid"
                    # Add table headers
                    hdr_colum = table_b.rows[0].cells
                    hdr_colum[0].text = "PRODUIT"
                    hdr_colum[1].text = "STOCK"
                    for hdr_colum in table_b.rows[0].cells:
                        hdr_colum.paragraphs[0].runs[0].font.name = 'Calibri'
                        hdr_colum.paragraphs[0].runs[0].font.size = Pt(10.5)
                        hdr_colum.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                        hdr_colum.paragraphs[0].runs[0].font.bold = True
                        hdr_colum.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    hdr_colum = table_b.rows[1].cells
                    hdr_colum[0].text = "POIS CHICHES PL"
                    hdr_colum[1].text = self.restpoischichepltt +"  QX"
                    hdr_colum = table_b.rows[2].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 12 mm"
                    hdr_colum[1].text = self.restpoichichedouzett + "  QX"
                    hdr_colum = table_b.rows[3].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 09 mm"
                    hdr_colum[1].text = self.restpoichicheneuftt + " QX"
                    hdr_colum = table_b.rows[4].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 08 mm"
                    hdr_colum[1].text = self.restpoichicheouittt + " QX"
                    hdr_colum = table_b.rows[5].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 06 mm"
                    hdr_colum[1].text = self.restpoichichett + " QX"
                    hdr_colum = table_b.rows[6].cells
                    hdr_colum[0].text = "HARICOT BLANC IMP"
                    hdr_colum[1].text = self.restharicotimptt + " QX"
                    hdr_colum = table_b.rows[7].cells
                    hdr_colum[0].text = "HARICOT LSB"
                    hdr_colum[1].text = self.restharicotlsbtt + " QX"
                    hdr_colum = table_b.rows[8].cells
                    hdr_colum[0].text = "LENTILLE PL"
                    hdr_colum[1].text = self.restlenpltt + " QX"
                    hdr_colum = table_b.rows[9].cells
                    hdr_colum[0].text = "LENTILLE IMP 05-07 mm"
                    hdr_colum[1].text = self.restlentsinqtt + " QX"
                    hdr_colum = table_b.rows[10].cells
                    hdr_colum[0].text = "LENTILLE IMP vert"
                    hdr_colum[1].text =  self.restlentvertt + " QX"
                    hdr_colum = table_b.rows[11].cells
                    hdr_colum[0].text = "LENTILLE IMP ROUGE"
                    hdr_colum[1].text = self.restlentrougett + " QX"
                    hdr_colum = table_b.rows[12].cells
                    hdr_colum[0].text = "RIZ IMP"
                    hdr_colum[1].text = self.resultt + " QX"
                    hdr_colum = table_b.rows[13].cells
                    hdr_colum[0].text = "RIZ ETUVE"
                    hdr_colum[1].text = self.resuletv + " QX"
                    for row in table_b.rows[1:14]:
                        for cell in row.cells:
                            cell.paragraphs[0].runs[0].font.name = 'Calibri'
                            cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                            cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                            cell.paragraphs[0].runs[0].font.bold = False
                    # Save and open the document for printing
                    temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                    self.doc.save(temp_file)
                    os.startfile(temp_file, "open")
                elif self.selected_produit=="" and self.selected_magasin=="" and self.selected_origine=="" and self.selected_dates and self.selected_datesAU:

                    dattabase1 = mysql.connector.connect(host="localhost",  user=user, password=password,
                                                         database='datta_legumsec_out')
                    cursor = dattabase1.cursor()
                    cursor.execute(
                        "SELECT le,produit,origine,quantité,unité,sacherie,fournisseur, imatricule,date_recoltes, date_CND, DLUM FROM outtable WHERE le BETWEEN %s AND %s  AND date_recoltes IS NOT NULL AND date_CND IS NOT NULL AND DLUM IS NOT NULL",
                        (self.selected_dates, self.selected_datesAU))
                    rows = cursor.fetchall()
                    curs1x = dattabase1.cursor()
                    curs1x.execute(
                        "SELECT SUM(quantitérizentré) - SUM(quantitérizesorté) as quantitérizerest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    self.resulxs = curs1x.fetchone()
                    self.resulxttm = str(self.resulxs[0])
                    curs1xtv = dattabase1.cursor()
                    curs1xtv.execute(
                        "SELECT SUM(quantitérizetvntré) - SUM(quantitérizetvsorté) as quantitérizetvrest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    self.resulxstv = curs1xtv.fetchone()
                    self.resulxtttvm = str(self.resulxs[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitepchentré) - sum(quantitépchsorté) as quantitépchrest  FROM outtable   WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resultpch_att = curstt.fetchone()
                    self.restpoichichedouzettsd = str(resultpch_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT SUM(quantitépchsixentré) - SUM(quantitépchsixsorté) as quantitépchsixrest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resultpch_att = curstt.fetchone()
                    self.restpoichichettm = str(resultpch_att[0])

                    curs5tt = dattabase1.cursor()
                    curs5tt.execute(
                        "SELECT sum(quantitépchplentré) - sum(quantitépchplsorté) as quantitépchplhrest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resultpchpl_att = curs5tt.fetchone()
                    self.restpoischicheplttm = str(resultpchpl_att[0])
                    curs7tt = dattabase1.cursor()
                    curs7tt.execute(
                        "SELECT sum(quantitélenplentré) - sum(quantitélenplsorté) as quantitélenplrest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resultlenpl_att = curs7tt.fetchone()
                    self.restlenplttm = str(resultlenpl_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenvertentré) - sum(quantitélenvertsorté) as quantitélenrest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resultlenverttt = curs9tt.fetchone()
                    self.restlentvertttm = str(resultlenverttt[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenrougeentré) - sum(quantitélenrougesorté) as quantitélenrougerest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resultlenrougett = curs9tt.fetchone()
                    self.restlentrougettm = str(resultlenrougett[0])
                    curs11tt = dattabase1.cursor()
                    curs11tt.execute(
                        "SELECT sum(quantitéharentré) - sum(quantitéharsorté) as quantitéharrest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resulthar_att = curs11tt.fetchone()
                    self.restharicotimpttm = str(resulthar_att[0])
                    curs11tt.execute(
                        "SELECT SUM(quantitéharlsbentré) - SUM(quantitéharlsbsorté) as quantitéharlsbrest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resultharlsb_att = curs11tt.fetchone()
                    self.restharicotlsbttm = str(resultharlsb_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchneufeentré) - sum(quantitépchneufesorté) as quantitépchneuferest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheneufttm = str(resultpch_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchhuitentré) - sum(quantitépchhuitsorté) as quantitépchhuitrest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheouitttm = str(resultpch_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenentré) - sum(quantitélensorté) as quantitélenrest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resultlensinqtt = curs9tt.fetchone()
                    self.restlentsinqttm = str(resultlensinqtt[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT SUM(quantitépchsixentré) - SUM(quantitépchsixsorté) as quantitépchsixrest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resultpch_attm = curstt.fetchone()
                    self.restpoichichett = str(resultpch_att[0])

                    curs5tt = dattabase1.cursor()
                    curs5tt.execute(
                        "SELECT sum(quantitépchplentré) - sum(quantitépchplsorté) as quantitépchplhrest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resultpchpl_att = curs5tt.fetchone()
                    self.restpoischicheplttm = str(resultpchpl_att[0])
                    curs7tt = dattabase1.cursor()
                    curs7tt.execute(
                        "SELECT sum(quantitélenplentré) - sum(quantitélenplsorté) as quantitélenplrest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resultlenpl_att = curs7tt.fetchone()
                    self.restlenplttm = str(resultlenpl_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenvertentré) - sum(quantitélenvertsorté) as quantitélenrest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resultlenverttt = curs9tt.fetchone()
                    self.restlentvertttm = str(resultlenverttt[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenrougeentré) - sum(quantitélenrougesorté) as quantitélenrougerest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resultlenrougett = curs9tt.fetchone()
                    self.restlentrougettm = str(resultlenrougett[0])
                    curs11tt = dattabase1.cursor()
                    curs11tt.execute(
                        "SELECT sum(quantitéharentré) - sum(quantitéharsorté) as quantitéharrest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resulthar_att = curs11tt.fetchone()
                    self.restharicotimptt = str(resulthar_att[0])
                    curs11tt.execute(
                        "SELECT SUM(quantitéharlsbentré) - SUM(quantitéharlsbsorté) as quantitéharlsbrest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resultharlsb_att = curs11tt.fetchone()
                    self.restharicotlsbttm = str(resultharlsb_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchneufeentré) - sum(quantitépchneufesorté) as quantitépchneuferest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheneufttm = str(resultpch_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchhuitentré) - sum(quantitépchhuitsorté) as quantitépchhuitrest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheouitttm = str(resultpch_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenentré) - sum(quantitélensorté) as quantitélenrest FROM outtable WHERE  le BETWEEN %s AND %s",
                        (self.selected_dates, self.selected_datesAU))
                    resultlensinqtt = curs9tt.fetchone()
                    self.restlentsinqttm = str(resultlensinqtt[0])

                    ########################################################################
                    ########################################################################
                    # Create a new document and add a table
                    self.doc = docx.Document()
                    section = self.doc.sections[0]
                    section.orientation = WD_ORIENTATION.LANDSCAPE
                    section.page_width = docx.shared.Cm(29.7)
                    section.page_height = docx.shared.Cm(21.0)
                    section.top_margin = docx.shared.Cm(0.5)
                    section.bottom_margin = docx.shared.Cm(1.5)
                    section.right_margin = docx.shared.Cm(1.5)
                    section.left_margin = docx.shared.Cm(1.5)
                    heading1 = self.doc.add_heading(
                        '\t\t\t\t\t\tSITUATION DES STOCKS DES LEGUMES SECS\n\t\t\t\t\t\t\t\tPRODUITS LIVRÉE',
                        level=1)
                    heading1.style.font.name = 'Times New Roman'
                    heading1.style.font.size = Pt(14)
                    heading1.style.font.bold = True
                    heading1.style.font.color.rgb = RGBColor(0, 0, 0)
                    # add second heading
                    paragraph0 = self.doc.add_paragraph()
                    paragraph0.style.font.name = 'Times New Roman'
                    left_run0 = paragraph0.add_run(
                        'Période du :' + self.selected_dates + "  au:" + self.selected_datesAU)
                    paragraph0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    paragraph1 = self.doc.add_paragraph()
                    paragraph1.style.font.name = 'Times New Roman'
                    current_date = QDate.currentDate().getDate()
                    left_run1 = paragraph1.add_run(
                        "\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t Date d'édition : {}/{}/{}".format(current_date[2],
                                                                                                current_date[1],
                                                                                                current_date[0]))
                    left_run1.font.size = Pt(9)

                    paragraph = self.doc.add_paragraph()
                    paragraph.style.font.name = 'Times New Roman'
                    paragraph.style.font.size = Pt(12)
                    left_run = paragraph.add_run(self.selected_magasin + '\t\t\t\t\t\t\t\t\t\t\t\t\t')
                    left_run.bold = True

                    table = self.doc.add_table(rows=1, cols=11)
                    table.style = "Table Grid"  # set the table style
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'DATE'
                    hdr_cells[1].text = 'PRODUIT'
                    hdr_cells[2].text = 'ORIGINE'
                    hdr_cells[3].text = 'QUANTITE'
                    hdr_cells[4].text = 'MAGASIN DE STOCK'
                    hdr_cells[5].text = 'SASHERIE  CND'
                    hdr_cells[6].text = 'EXPIDITEUR'
                    hdr_cells[7].text = 'IMATRICUL'
                    hdr_cells[8].text = 'Date de récolte'
                    hdr_cells[9].text = 'Date de CND'
                    hdr_cells[10].text = 'Date de DLUM'
                    # Set the width of the header cells
                    hdr_cells[0].width = Inches(0.5)  # set the width of the first column to 1.5 inches
                    hdr_cells[1].width = Inches(2)
                    hdr_cells[2].width = Inches(1)
                    hdr_cells[3].width = Inches(0.4)
                    hdr_cells[4].width = Inches(1.3)
                    hdr_cells[5].width = Inches(1)
                    hdr_cells[6].width = Inches(1)
                    hdr_cells[7].width = Inches(1)
                    hdr_cells[8].width = Inches(0.5)
                    hdr_cells[9].width = Inches(0.5)
                    hdr_cells[10].width = Inches(0.5)

                    for cell in table.rows[0].cells:
                        cell.paragraphs[0].runs[0].font.name = 'Calibri'
                        cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                        cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                        cell.paragraphs[0].runs[0].font.bold = True
                        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    # Add the data to the table
                    table_rows = len(rows)
                    table_cols = len(rows[0])
                    for row in range(table_rows):
                        table.add_row()
                        for col in range(table_cols):
                            cell = table.cell(row + 1, col)
                            cell.text = str(rows[row][col])
                            # Set font properties
                            cell.paragraphs[0].runs[0].font.name = 'Calibri'
                            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
                            cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                            cell.paragraphs[0].runs[0].font.bold = False
                            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    for row in table.rows:
                        row.height = Inches(0.3)
                    # Create the document
                    # Create the table
                    self.doc.add_paragraph().add_run().add_break()
                    table_b = self.doc.add_table(rows=14, cols=2)
                    table_b.style = "Table Grid"
                    # Add table headers
                    hdr_colum = table_b.rows[0].cells
                    hdr_colum[0].text = "PRODUIT"
                    hdr_colum[1].text = "STOCK"
                    for hdr_colum in table_b.rows[0].cells:
                        hdr_colum.paragraphs[0].runs[0].font.name = 'Calibri'
                        hdr_colum.paragraphs[0].runs[0].font.size = Pt(10.5)
                        hdr_colum.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                        hdr_colum.paragraphs[0].runs[0].font.bold = True
                        hdr_colum.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    hdr_colum = table_b.rows[1].cells
                    hdr_colum[0].text = "POIS CHICHES PL"
                    hdr_colum[1].text = ""
                    hdr_colum = table_b.rows[2].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 12 mm"
                    hdr_colum[1].text = self.restpoichichedouzettsd  + "  QX"
                    hdr_colum = table_b.rows[3].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 09 mm"
                    hdr_colum[1].text = self.restpoichicheneufttm + " QX"
                    hdr_colum = table_b.rows[4].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 08 mm"
                    hdr_colum[1].text = self.restpoichicheouitttm + " QX"
                    hdr_colum = table_b.rows[5].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 06 mm"
                    hdr_colum[1].text = self.restpoichichettm + " QX"
                    hdr_colum = table_b.rows[6].cells
                    hdr_colum[0].text = "HARICOT BLANC IMP"
                    hdr_colum[1].text = self.restharicotimpttm + " QX"
                    hdr_colum = table_b.rows[7].cells
                    hdr_colum[0].text = "HARICOT LSB"
                    hdr_colum[1].text = self.restharicotlsbttm + " QX"
                    hdr_colum = table_b.rows[8].cells
                    hdr_colum[0].text = "LENTILLE PL"
                    hdr_colum[1].text = self.restlenplttm + " QX"
                    hdr_colum = table_b.rows[9].cells
                    hdr_colum[0].text = "LENTILLE IMP 05-07 mm"
                    hdr_colum[1].text = self.restlentsinqttm + " QX"
                    hdr_colum = table_b.rows[10].cells
                    hdr_colum[0].text = "LENTILLE IMP vert"
                    hdr_colum[1].text = self.restlentvertttm + " QX"
                    hdr_colum = table_b.rows[11].cells
                    hdr_colum[0].text = "LENTILLE IMP ROUGE"
                    hdr_colum[1].text = self.restlentrougettm + " QX"
                    hdr_colum = table_b.rows[12].cells
                    hdr_colum[0].text = "RIZ IMP"
                    hdr_colum[1].text = self.resulxttm + " QX"
                    hdr_colum = table_b.rows[13].cells
                    hdr_colum[0].text = "RIZ ETUVE"
                    hdr_colum[1].text = self.resulxttm + " QX"
                    for row in table_b.rows[1:14]:
                        for cell in row.cells:
                            cell.paragraphs[0].runs[0].font.name = 'Calibri'
                            cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                            cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                            cell.paragraphs[0].runs[0].font.bold = False
                    # Save and open the document for printing
                    temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                    self.doc.save(temp_file)
                    os.startfile(temp_file, "open")
                    dattabase1.close()
                elif self.selected_produit=="" and self.selected_magasin and self.selected_origine=="" and self.selected_dates=="--" and self.selected_datesAU=="--"  :


                    dattabase1 = mysql.connector.connect(host="localhost",  user=user, password=password,
                                                         database='datta_legumsec_out')
                    cursor = dattabase1.cursor()
                    cursor.execute(
                        "SELECT le,produit,origine,quantité,unité,sacherie,fournisseur, imatricule,date_recoltes, date_CND, DLUM FROM outtable WHERE unité = %s AND date_recoltes IS NOT NULL AND date_CND IS NOT NULL AND DLUM IS NOT NULL",
                        (self.selected_magasin,))
                    rows = cursor.fetchall()
                    curs1x = dattabase1.cursor()
                    curs1x.execute("SELECT SUM(quantitérizentré) - SUM(quantitérizesorté) as quantitérizerest FROM outtable WHERE  unité=%s",
                        (self.selected_magasin,))
                    self.resulxs = curs1x.fetchone()
                    self.resulxttm = str(self.resulxs[0])
                    curs1xtv = dattabase1.cursor()
                    curs1xtv.execute("SELECT SUM(quantitérizetvntré) - SUM(quantitérizetvsorté) as quantitérizetvrest FROM outtable WHERE  unité=%s",
                        (self.selected_magasin,))
                    self.resulxstv = curs1xtv.fetchone()
                    self.resulxtttvm = str(self.resulxstv[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitepchentré) - sum(quantitépchsorté) as quantitépchrest  FROM outtable   WHERE   unité=%s",
                        (self.selected_magasin,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichichedouzettm = str(resultpch_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT SUM(quantitépchsixentré) - SUM(quantitépchsixsorté) as quantitépchsixrest FROM outtable WHERE   unité=%s",
                        (self.selected_magasin,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichichettm = str(resultpch_att[0])

                    curs5tt = dattabase1.cursor()
                    curs5tt.execute(
                        "SELECT sum(quantitépchplentré) - sum(quantitépchplsorté) as quantitépchplhrest FROM outtable WHERE   unité=%s",
                        (self.selected_magasin,))
                    resultpchpl_att = curs5tt.fetchone()
                    self.restpoischicheplttm = str(resultpchpl_att[0])
                    curs7tt = dattabase1.cursor()
                    curs7tt.execute(
                        "SELECT sum(quantitélenplentré) - sum(quantitélenplsorté) as quantitélenplrest FROM outtable WHERE   unité=%s",
                        (self.selected_magasin,))
                    resultlenpl_att = curs7tt.fetchone()
                    self.restlenplttm = str(resultlenpl_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenvertentré) - sum(quantitélenvertsorté) as quantitélenrest FROM outtable WHERE   unité=%s",
                        (self.selected_magasin,))
                    resultlenverttt = curs9tt.fetchone()
                    self.restlentvertttm = str(resultlenverttt[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenrougeentré) - sum(quantitélenrougesorté) as quantitélenrougerest FROM outtable WHERE   unité=%s",
                        (self.selected_magasin,))
                    resultlenrougett = curs9tt.fetchone()
                    self.restlentrougettm = str(resultlenrougett[0])
                    curs11tt = dattabase1.cursor()
                    curs11tt.execute(
                        "SELECT sum(quantitéharentré) - sum(quantitéharsorté) as quantitéharrest FROM outtable WHERE   unité=%s",
                        (self.selected_magasin,))
                    resulthar_att = curs11tt.fetchone()
                    self.restharicotimpttm = str(resulthar_att[0])
                    curs11tt.execute(
                        "SELECT SUM(quantitéharlsbentré) - SUM(quantitéharlsbsorté) as quantitéharlsbrest FROM outtable WHERE   unité=%s",
                        (self.selected_magasin,))
                    resultharlsb_att = curs11tt.fetchone()
                    self.restharicotlsbttm = str(resultharlsb_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchneufeentré) - sum(quantitépchneufesorté) as quantitépchneuferest FROM outtable   WHERE unité=%s",
                        (self.selected_magasin,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheneufttm = str(resultpch_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchhuitentré) - sum(quantitépchhuitsorté) as quantitépchhuitrest FROM outtable WHERE   unité=%s",
                        (self.selected_magasin,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheouitttm = str(resultpch_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenentré) - sum(quantitélensorté) as quantitélenrest FROM outtable WHERE   unité=%s",
                        (self.selected_magasin,))
                    resultlensinqtt = curs9tt.fetchone()
                    self.restlentsinqttm = str(resultlensinqtt[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT SUM(quantitépchsixentré) - SUM(quantitépchsixsorté) as quantitépchsixrest FROM outtable  WHERE unité=%s",
                        (self.selected_magasin,))
                    resultpch_attm = curstt.fetchone()
                    self.restpoichichett = str(resultpch_att[0])

                    curs5tt = dattabase1.cursor()
                    curs5tt.execute(
                        "SELECT sum(quantitépchplentré) - sum(quantitépchplsorté) as quantitépchplhrest FROM outtable   WHERE unité=%s",
                        (self.selected_magasin,))
                    resultpchpl_att = curs5tt.fetchone()
                    self.restpoischicheplttm = str(resultpchpl_att[0])
                    curs7tt = dattabase1.cursor()
                    curs7tt.execute(
                        "SELECT sum(quantitélenplentré) - sum(quantitélenplsorté) as quantitélenplrest FROM outtable  WHERE unité=%s",
                        (self.selected_magasin,))
                    resultlenpl_att = curs7tt.fetchone()
                    self.restlenplttm = str(resultlenpl_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenvertentré) - sum(quantitélenvertsorté) as quantitélenrest FROM outtable WHERE unité=%s",
                        (self.selected_magasin,))
                    resultlenverttt = curs9tt.fetchone()
                    self.restlentvertttm = str(resultlenverttt[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenrougeentré) - sum(quantitélenrougesorté) as quantitélenrougerest FROM outtable WHERE unité=%s",
                        (self.selected_magasin,))
                    resultlenrougett = curs9tt.fetchone()
                    self.restlentrougettm = str(resultlenrougett[0])
                    curs11tt = dattabase1.cursor()
                    curs11tt.execute(
                        "SELECT sum(quantitéharentré) - sum(quantitéharsorté) as quantitéharrest FROM outtable WHERE  unité=%s",
                        (self.selected_magasin,))
                    resulthar_att = curs11tt.fetchone()
                    self.restharicotimptt = str(resulthar_att[0])
                    curs11tt.execute(
                        "SELECT SUM(quantitéharlsbentré) - SUM(quantitéharlsbsorté) as quantitéharlsbrest FROM outtable   WHERE unité=%s",
                        (self.selected_magasin,))
                    resultharlsb_att = curs11tt.fetchone()
                    self.restharicotlsbttm = str(resultharlsb_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchneufeentré) - sum(quantitépchneufesorté) as quantitépchneuferest FROM outtable  WHERE unité=%s",
                        (self.selected_magasin,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheneufttm = str(resultpch_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchhuitentré) - sum(quantitépchhuitsorté) as quantitépchhuitrest FROM outtable WHERE unité=%s",
                        (self.selected_magasin,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheouitttm = str(resultpch_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenentré) - sum(quantitélensorté) as quantitélenrest FROM outtable WHERE  unité=%s",
                        (self.selected_magasin,))
                    resultlensinqtt = curs9tt.fetchone()
                    self.restlentsinqttm = str(resultlensinqtt[0])

                    ########################################################################
                    # Create a new document and add a table
                    self.doc = docx.Document()
                    section = self.doc.sections[0]
                    section.orientation = WD_ORIENTATION.LANDSCAPE
                    section.page_width = docx.shared.Cm(29.7)
                    section.page_height = docx.shared.Cm(21.0)
                    section.top_margin = docx.shared.Cm(0.5)
                    section.bottom_margin = docx.shared.Cm(1.5)
                    section.right_margin = docx.shared.Cm(1.5)
                    section.left_margin = docx.shared.Cm(1.5)
                    heading1 = self.doc.add_heading(
                        '\t\t\t\t\t\tSITUATION DES STOCKS DES LEGUMES SECS\n\t\t\t\t\t\t\t\tPRODUITS LIVRÉE',
                        level=1)
                    heading1.style.font.name = 'Times New Roman'
                    heading1.style.font.size = Pt(14)
                    heading1.style.font.bold = True
                    heading1.style.font.color.rgb = RGBColor(0, 0, 0)
                    # add second heading

                    paragraph1 = self.doc.add_paragraph()
                    paragraph1.style.font.name = 'Times New Roman'
                    current_date = QDate.currentDate().getDate()
                    left_run1 = paragraph1.add_run(
                        "\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t Date d'édition : {}/{}/{}".format(current_date[2],
                                                                                                current_date[1],
                                                                                                current_date[0]))
                    left_run1.font.size = Pt(9)

                    paragraph = self.doc.add_paragraph()
                    paragraph.style.font.name = 'Times New Roman'
                    paragraph.style.font.size = Pt(12)
                    left_run = paragraph.add_run(self.selected_magasin + '\t\t\t\t\t\t\t\t\t\t\t\t\t')
                    left_run.bold = True

                    table = self.doc.add_table(rows=1, cols=11)
                    table.style = "Table Grid"  # set the table style
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'DATE'
                    hdr_cells[1].text = 'PRODUIT'
                    hdr_cells[2].text = 'ORIGINE'
                    hdr_cells[3].text = 'QUANTITE'
                    hdr_cells[4].text = 'MAGASIN DE STOCK'
                    hdr_cells[5].text = 'SASHERIE  CND'
                    hdr_cells[6].text = 'EXPIDITEUR'
                    hdr_cells[7].text = 'IMATRICUL'
                    hdr_cells[8].text = 'Date de récolte'
                    hdr_cells[9].text = 'Date de CND'
                    hdr_cells[10].text = 'Date de DLUM'
                    # Set the width of the header cells
                    hdr_cells[0].width = Inches(0.5)  # set the width of the first column to 1.5 inches
                    hdr_cells[1].width = Inches(2)
                    hdr_cells[2].width = Inches(1)
                    hdr_cells[3].width = Inches(0.4)
                    hdr_cells[4].width = Inches(1.3)
                    hdr_cells[5].width = Inches(1)
                    hdr_cells[6].width = Inches(1)
                    hdr_cells[7].width = Inches(1)
                    hdr_cells[8].width = Inches(0.5)
                    hdr_cells[9].width = Inches(0.5)
                    hdr_cells[10].width = Inches(0.5)

                    for cell in table.rows[0].cells:
                        cell.paragraphs[0].runs[0].font.name = 'Calibri'
                        cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                        cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                        cell.paragraphs[0].runs[0].font.bold = True
                        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    # Add the data to the table
                    table_rows = len(rows)
                    table_cols = len(rows[0])
                    for row in range(table_rows):
                        table.add_row()
                        for col in range(table_cols):
                            cell = table.cell(row + 1, col)
                            cell.text = str(rows[row][col])
                            # Set font properties
                            cell.paragraphs[0].runs[0].font.name = 'Calibri'
                            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
                            cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                            cell.paragraphs[0].runs[0].font.bold = False
                            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    for row in table.rows:
                        row.height = Inches(0.3)
                    # Create the document
                    # Create the table
                    self.doc.add_paragraph().add_run().add_break()
                    table_b = self.doc.add_table(rows=14, cols=2)
                    table_b.style = "Table Grid"
                    # Add table headers
                    hdr_colum = table_b.rows[0].cells
                    hdr_colum[0].text = "PRODUIT"
                    hdr_colum[1].text = "STOCK"
                    for hdr_colum in table_b.rows[0].cells:
                        hdr_colum.paragraphs[0].runs[0].font.name = 'Calibri'
                        hdr_colum.paragraphs[0].runs[0].font.size = Pt(10.5)
                        hdr_colum.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                        hdr_colum.paragraphs[0].runs[0].font.bold = True
                        hdr_colum.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    hdr_colum = table_b.rows[1].cells
                    hdr_colum[0].text = "POIS CHICHES PL"
                    hdr_colum[1].text = ""
                    hdr_colum = table_b.rows[2].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 12 mm"
                    hdr_colum[1].text = self.restpoichichedouzettm + "  QX"
                    hdr_colum = table_b.rows[3].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 09 mm"
                    hdr_colum[1].text = self.restpoichicheneufttm + " QX"
                    hdr_colum = table_b.rows[4].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 08 mm"
                    hdr_colum[1].text = self.restpoichicheouitttm + " QX"
                    hdr_colum = table_b.rows[5].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 06 mm"
                    hdr_colum[1].text = self.restpoichichettm + " QX"
                    hdr_colum = table_b.rows[6].cells
                    hdr_colum[0].text = "HARICOT BLANC IMP"
                    hdr_colum[1].text = self.restharicotimpttm + " QX"
                    hdr_colum = table_b.rows[7].cells
                    hdr_colum[0].text = "HARICOT LSB"
                    hdr_colum[1].text = self.restharicotlsbttm + " QX"
                    hdr_colum = table_b.rows[8].cells
                    hdr_colum[0].text = "LENTILLE PL"
                    hdr_colum[1].text = self.restlenplttm + " QX"
                    hdr_colum = table_b.rows[9].cells
                    hdr_colum[0].text = "LENTILLE IMP 05-07 mm"
                    hdr_colum[1].text = self.restlentsinqttm + " QX"
                    hdr_colum = table_b.rows[10].cells
                    hdr_colum[0].text = "LENTILLE IMP vert"
                    hdr_colum[1].text = self.restlentvertttm + " QX"
                    hdr_colum = table_b.rows[11].cells
                    hdr_colum[0].text = "LENTILLE IMP ROUGE"
                    hdr_colum[1].text = self.restlentrougettm + " QX"
                    hdr_colum = table_b.rows[12].cells
                    hdr_colum[0].text = "RIZ IMP"
                    hdr_colum[1].text = self.resulxttm + " QX"
                    hdr_colum = table_b.rows[13].cells
                    hdr_colum[0].text = "RIZ ETUVE"
                    hdr_colum[1].text = self.resulxttm + " QX"
                    for row in table_b.rows[1:14]:
                        for cell in row.cells:
                            cell.paragraphs[0].runs[0].font.name = 'Calibri'
                            cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                            cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                            cell.paragraphs[0].runs[0].font.bold = False
                    # Save and open the document for printing
                    temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                    self.doc.save(temp_file)
                    os.startfile(temp_file, "open")
                    dattabase1.close()
                elif self.selected_produit and self.selected_magasin=="" and self.selected_origine=="" and self.selected_dates=="--" and self.selected_datesAU=="--":

                    dattabase1 = mysql.connector.connect(host="localhost",  user=user, password=password,
                                                         database='datta_legumsec_out')
                    cursor = dattabase1.cursor()
                    cursor.execute(
                        "SELECT le,produit,origine,quantité,unité,sacherie,fournisseur, imatricule,date_recoltes ,date_CND, DLUM FROM outtable WHERE produit = %s  and date_recoltes IS NOT NULL AND date_CND IS NOT NULL AND DLUM IS NOT NULL",
                        (self.selected_produit,))
                    rows = cursor.fetchall()
                    curs1x = dattabase1.cursor()
                    curs1x.execute(
                        "SELECT SUM(quantitérizentré) - SUM(quantitérizesorté) as quantitérizerest FROM outtable WHERE  produit= %s",
                        (self.selected_produit,))
                    self.resulxs = curs1x.fetchone()
                    self.resulxttp = str(self.resulxs[0])
                    curs1xtv = dattabase1.cursor()
                    curs1xtv.execute(
                        "SELECT SUM(quantitérizetvntré) - SUM(quantitérizetvsorté) as quantitérizetvrest FROM outtable WHERE  produit= %s",
                        (self.selected_produit,))
                    self.resulxstv = curs1xtv.fetchone()
                    self.resulxtttvp = str(self.resulxstv[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitepchentré) - sum(quantitépchsorté) as quantitépchrest  FROM outtable   WHERE   produit=%s",
                        (self.selected_produit,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichichedouzettp = str(resultpch_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT SUM(quantitépchsixentré) - SUM(quantitépchsixsorté) as quantitépchsixrest FROM outtable WHERE   produit=%s",
                        (self.selected_produit,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichichettp = str(resultpch_att[0])

                    curs5tt = dattabase1.cursor()
                    curs5tt.execute(
                        "SELECT sum(quantitépchplentré) - sum(quantitépchplsorté) as quantitépchplhrest FROM outtable WHERE   produit=%s",
                        (self.selected_produit,))
                    resultpchpl_att = curs5tt.fetchone()
                    self.restpoischicheplttp = str(resultpchpl_att[0])
                    curs7tt = dattabase1.cursor()
                    curs7tt.execute(
                        "SELECT sum(quantitélenplentré) - sum(quantitélenplsorté) as quantitélenplrest FROM outtable WHERE   produit=%s",
                        (self.selected_produit,))
                    resultlenpl_att = curs7tt.fetchone()
                    self.restlenplttp = str(resultlenpl_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenvertentré) - sum(quantitélenvertsorté) as quantitélenrest FROM outtable WHERE   produit=%s",
                        (self.selected_produit,))
                    resultlenverttt = curs9tt.fetchone()
                    self.restlentvertttp = str(resultlenverttt[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenrougeentré) - sum(quantitélenrougesorté) as quantitélenrougerest FROM outtable WHERE   produit=%s",
                        (self.selected_produit,))
                    resultlenrougett = curs9tt.fetchone()
                    self.restlentrougettp = str(resultlenrougett[0])
                    curs11tt = dattabase1.cursor()
                    curs11tt.execute(
                        "SELECT sum(quantitéharentré) - sum(quantitéharsorté) as quantitéharrest FROM outtable WHERE   produit=%s",
                        (self.selected_produit,))
                    resulthar_att = curs11tt.fetchone()
                    self.restharicotimpttp = str(resulthar_att[0])
                    curs11tt.execute(
                        "SELECT SUM(quantitéharlsbentré) - SUM(quantitéharlsbsorté) as quantitéharlsbrest FROM outtable WHERE   produit=%s",
                        (self.selected_produit,))
                    resultharlsb_att = curs11tt.fetchone()
                    self.restharicotlsbttp = str(resultharlsb_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchneufeentré) - sum(quantitépchneufesorté) as quantitépchneuferest FROM outtable   WHERE produit=%s",
                        (self.selected_produit,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheneufttp = str(resultpch_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchhuitentré) - sum(quantitépchhuitsorté) as quantitépchhuitrest FROM outtable WHERE   produit=%s",
                        (self.selected_produit,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheouitttp = str(resultpch_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenentré) - sum(quantitélensorté) as quantitélenrest FROM outtable WHERE   produit=%s",
                        (self.selected_produit,))
                    resultlensinqtt = curs9tt.fetchone()
                    self.restlentsinqttp = str(resultlensinqtt[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT SUM(quantitépchsixentré) - SUM(quantitépchsixsorté) as quantitépchsixrest FROM outtable  WHERE produit=%s",
                        (self.selected_produit,))
                    resultpch_attp = curstt.fetchone()
                    self.restpoichichettp = str(resultpch_attp[0])

                    curs5tt = dattabase1.cursor()
                    curs5tt.execute(
                        "SELECT sum(quantitépchplentré) - sum(quantitépchplsorté) as quantitépchplhrest FROM outtable   WHERE produit=%s",
                        (self.selected_produit,))
                    resultpchpl_att = curs5tt.fetchone()
                    self.restpoischicheplttp = str(resultpchpl_att[0])
                    curs7tt = dattabase1.cursor()
                    curs7tt.execute(
                        "SELECT sum(quantitélenplentré) - sum(quantitélenplsorté) as quantitélenplrest FROM outtable  WHERE produit=%s",
                        (self.selected_produit,))
                    resultlenpl_att = curs7tt.fetchone()
                    self.restlenplttp = str(resultlenpl_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenvertentré) - sum(quantitélenvertsorté) as quantitélenrest FROM outtable WHERE produit=%s",
                        (self.selected_produit,))
                    resultlenverttt = curs9tt.fetchone()
                    self.restlentvertttp = str(resultlenverttt[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenrougeentré) - sum(quantitélenrougesorté) as quantitélenrougerest FROM outtable WHERE produit=%s",
                        (self.selected_produit,))
                    resultlenrougett = curs9tt.fetchone()
                    self.restlentrougettp = str(resultlenrougett[0])
                    curs11tt = dattabase1.cursor()
                    curs11tt.execute(
                        "SELECT sum(quantitéharentré) - sum(quantitéharsorté) as quantitéharrest FROM outtable WHERE  produit=%s",
                        (self.selected_produit,))
                    resulthar_att = curs11tt.fetchone()
                    self.restharicotimpttp = str(resulthar_att[0])
                    curs11tt.execute(
                        "SELECT SUM(quantitéharlsbentré) - SUM(quantitéharlsbsorté) as quantitéharlsbrest FROM outtable   WHERE produit=%s",
                        (self.selected_produit,))
                    resultharlsb_att = curs11tt.fetchone()
                    self.restharicotlsbttp = str(resultharlsb_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchneufeentré) - sum(quantitépchneufesorté) as quantitépchneuferest FROM outtable  WHERE produit=%s",
                        (self.selected_produit,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheneufttp = str(resultpch_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchhuitentré) - sum(quantitépchhuitsorté) as quantitépchhuitrest FROM outtable WHERE produit=%s",
                        (self.selected_produit,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheouitttp = str(resultpch_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenentré) - sum(quantitélensorté) as quantitélenrest FROM outtable WHERE  produit=%s",
                        (self.selected_produit,))
                    resultlensinqtt = curs9tt.fetchone()
                    self.restlentsinqttp = str(resultlensinqtt[0])
                    ########################################################################
                    # Create a new document and add a table
                    self.doc = docx.Document()
                    section = self.doc.sections[0]
                    section.orientation = WD_ORIENTATION.LANDSCAPE
                    section.page_width = docx.shared.Cm(29.7)
                    section.page_height = docx.shared.Cm(21.0)
                    section.top_margin = docx.shared.Cm(0.5)
                    section.bottom_margin = docx.shared.Cm(1.5)
                    section.right_margin = docx.shared.Cm(1.5)
                    section.left_margin = docx.shared.Cm(1.5)
                    heading1 = self.doc.add_heading(
                        '\t\t\t\t\t\tSITUATION DES STOCKS DES LEGUMES SECS\n\t\t\t\t\t\t\t\tPRODUITS LIVRÉE',
                        level=1)
                    heading1.style.font.name = 'Times New Roman'
                    heading1.style.font.size = Pt(14)
                    heading1.style.font.bold = True
                    heading1.style.font.color.rgb = RGBColor(0, 0, 0)
                    # add second heading


                    paragraph1 = self.doc.add_paragraph()
                    paragraph1.style.font.name = 'Times New Roman'
                    current_date = QDate.currentDate().getDate()
                    left_run1 = paragraph1.add_run(
                        "\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t Date d'édition : {}/{}/{}".format(current_date[2],
                                                                                                current_date[1],
                                                                                                current_date[0]))
                    left_run1.font.size = Pt(9)

                    paragraph = self.doc.add_paragraph()
                    paragraph.style.font.name = 'Times New Roman'
                    paragraph.style.font.size = Pt(12)
                    left_run = paragraph.add_run("PRODUIT :"+self.selected_produit + '\t\t\t\t\t\t\t\t\t\t\t\t\t')
                    left_run.bold = True

                    table = self.doc.add_table(rows=1, cols=11)
                    table.style = "Table Grid"  # set the table style
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'DATE'
                    hdr_cells[1].text = 'PRODUIT'
                    hdr_cells[2].text = 'ORIGINE'
                    hdr_cells[3].text = 'QUANTITE'
                    hdr_cells[4].text = 'MAGASIN DE STOCK'
                    hdr_cells[5].text = 'SASHERIE  CND'
                    hdr_cells[6].text = 'EXPIDITEUR'
                    hdr_cells[7].text = 'IMATRICUL'
                    hdr_cells[8].text = 'Date de récolte'
                    hdr_cells[9].text = 'Date de CND'
                    hdr_cells[10].text = 'Date de DLUM'
                    # Set the width of the header cells
                    hdr_cells[0].width = Inches(0.5)  # set the width of the first column to 1.5 inches
                    hdr_cells[1].width = Inches(2)
                    hdr_cells[2].width = Inches(1)
                    hdr_cells[3].width = Inches(0.4)
                    hdr_cells[4].width = Inches(1.3)
                    hdr_cells[5].width = Inches(1)
                    hdr_cells[6].width = Inches(1)
                    hdr_cells[7].width = Inches(1)
                    hdr_cells[8].width = Inches(0.5)
                    hdr_cells[9].width = Inches(0.5)
                    hdr_cells[10].width = Inches(0.5)

                    for cell in table.rows[0].cells:
                        cell.paragraphs[0].runs[0].font.name = 'Calibri'
                        cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                        cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                        cell.paragraphs[0].runs[0].font.bold = True
                        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    # Add the data to the table
                    table_rows = len(rows)
                    table_cols = len(rows[0])
                    for row in range(table_rows):
                        table.add_row()
                        for col in range(table_cols):
                            cell = table.cell(row + 1, col)
                            cell.text = str(rows[row][col])
                            # Set font properties
                            cell.paragraphs[0].runs[0].font.name = 'Calibri'
                            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
                            cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                            cell.paragraphs[0].runs[0].font.bold = False
                            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    for row in table.rows:
                        row.height = Inches(0.3)
                    # Create the document
                    # Create the table
                    self.doc.add_paragraph().add_run().add_break()
                    table_b = self.doc.add_table(rows=2, cols=2)
                    table_b.style = "Table Grid"
                    # Add table headers
                    hdr_colum = table_b.rows[0].cells
                    hdr_colum[0].text = "PRODUIT"
                    hdr_colum[1].text = "STOCK"
                    for hdr_colum in table_b.rows[0].cells:
                        hdr_colum.paragraphs[0].runs[0].font.name = 'Calibri'
                        hdr_colum.paragraphs[0].runs[0].font.size = Pt(10.5)
                        hdr_colum.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                        hdr_colum.paragraphs[0].runs[0].font.bold = True
                        hdr_colum.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    hdr_colum = table_b.rows[1].cells

                    if self.selected_produit == "POIS CHICHES IMP 12 mm":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restpoichichedouzettp + "  QX"

                    if self.selected_produit == "POIS CHICHES IMP 12 mm":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restpoichichedouzettp + "  QX"
                    if self.selected_produit == "POIS CHICHES IMP 09 mm":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restpoichicheneufttp + " QX"
                    if self.selected_produit == "POIS CHICHES IMP 08 mm":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restpoichicheouitttp + " QX"
                    if self.selected_produit == "POIS CHICHES IMP 06 mm":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restpoichichettp + " QX"
                    if self.selected_produit == "HARICOT BLANC IMP":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restharicotimpttp + " QX"
                    if self.selected_produit == "HARICOT LSB":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restharicotlsbttp + " QX"
                    if self.selected_produit == "LENTILLE PL":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restlenplttp + " QX"
                    if self.selected_produit == "LENTILLE IMP 05-07 mm":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restlentsinqttp + " QX"
                    if self.selected_produit == "LENTILLE IMP vert":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restlentvertttp + " QX"
                    if self.selected_produit == "LENTILLE IMP ROUGE":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restlentrougettp + " QX"
                    if self.selected_produit == "RIZ IMP":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.resulxttp + " QX"
                    if self.selected_produit == "RIZ ETUVE":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.resulxtttvp + " QX"
                    for row in table_b.rows[1:2]:
                        for cell in row.cells:
                            cell.paragraphs[0].runs[0].font.name = 'Calibri'
                            cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                            cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                            cell.paragraphs[0].runs[0].font.bold = False
                    # Save and open the document for printing
                    temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                    self.doc.save(temp_file)
                    os.startfile(temp_file, "open")
                    dattabase1.close()
                elif self.selected_produit and self.selected_magasin and self.selected_origine == "" and self.selected_dates == "--" and self.selected_datesAU == "--":
                    conn = mysql.connector.connect(host="localhost",  user=user, password=password,
                                                   database="datta_legumesec_entry")
                    dattabase1 = mysql.connector.connect(host="localhost",  user=user, password=password,
                                                         database='datta_legumsec_out')
                    cursor = conn.cursor()
                    cursor.execute(
                        "SELECT le,produit,origine,quantité,unité,sacherie,fournisseur, imatricule,date_recoltes,date_CND,DLUM FROM entrytable WHERE produit = %s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    rows = cursor.fetchall()
                    curs1x = dattabase1.cursor()
                    curs1x.execute(
                        "SELECT SUM(quantitérizentré) - SUM(quantitérizesorté) as quantitérizerest FROM outtable WHERE  produit= %s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    self.resulxs = curs1x.fetchone()
                    self.resulxttp = str(self.resulxs[0])
                    curs1xtv = dattabase1.cursor()
                    curs1xtv.execute(
                        "SELECT SUM(quantitérizetvntré) - SUM(quantitérizetvsorté) as quantitérizetvrest FROM outtable WHERE  produit= %s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    self.resulxstv = curs1xtv.fetchone()
                    self.resulxtttvp = str(self.resulxstv[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitepchentré) - sum(quantitépchsorté) as quantitépchrest  FROM outtable   WHERE   produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichichedouzettp = str(resultpch_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT SUM(quantitépchsixentré) - SUM(quantitépchsixsorté) as quantitépchsixrest FROM outtable WHERE   produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichichettp = str(resultpch_att[0])

                    curs5tt = dattabase1.cursor()
                    curs5tt.execute(
                        "SELECT sum(quantitépchplentré) - sum(quantitépchplsorté) as quantitépchplhrest FROM outtable WHERE   produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resultpchpl_att = curs5tt.fetchone()
                    self.restpoischicheplttp = str(resultpchpl_att[0])
                    curs7tt = dattabase1.cursor()
                    curs7tt.execute(
                        "SELECT sum(quantitélenplentré) - sum(quantitélenplsorté) as quantitélenplrest FROM outtable WHERE   produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resultlenpl_att = curs7tt.fetchone()
                    self.restlenplttp = str(resultlenpl_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenvertentré) - sum(quantitélenvertsorté) as quantitélenrest FROM outtable WHERE   produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resultlenverttt = curs9tt.fetchone()
                    self.restlentvertttp = str(resultlenverttt[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenrougeentré) - sum(quantitélenrougesorté) as quantitélenrougerest FROM outtable WHERE   produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resultlenrougett = curs9tt.fetchone()
                    self.restlentrougettp = str(resultlenrougett[0])
                    curs11tt = dattabase1.cursor()
                    curs11tt.execute(
                        "SELECT sum(quantitéharentré) - sum(quantitéharsorté) as quantitéharrest FROM outtable WHERE   produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resulthar_att = curs11tt.fetchone()
                    self.restharicotimpttp = str(resulthar_att[0])
                    curs11tt.execute(
                        "SELECT SUM(quantitéharlsbentré) - SUM(quantitéharlsbsorté) as quantitéharlsbrest FROM outtable WHERE   produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resultharlsb_att = curs11tt.fetchone()
                    self.restharicotlsbttp = str(resultharlsb_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchneufeentré) - sum(quantitépchneufesorté) as quantitépchneuferest FROM outtable   WHERE produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheneufttp = str(resultpch_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchhuitentré) - sum(quantitépchhuitsorté) as quantitépchhuitrest FROM outtable WHERE   produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheouitttp = str(resultpch_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenentré) - sum(quantitélensorté) as quantitélenrest FROM outtable WHERE   produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resultlensinqtt = curs9tt.fetchone()
                    self.restlentsinqttp = str(resultlensinqtt[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT SUM(quantitépchsixentré) - SUM(quantitépchsixsorté) as quantitépchsixrest FROM outtable  WHERE produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resultpch_attp = curstt.fetchone()
                    self.restpoichichettp = str(resultpch_attp[0])

                    curs5tt = dattabase1.cursor()
                    curs5tt.execute(
                        "SELECT sum(quantitépchplentré) - sum(quantitépchplsorté) as quantitépchplhrest FROM outtable   WHERE produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resultpchpl_att = curs5tt.fetchone()
                    self.restpoischicheplttp = str(resultpchpl_att[0])
                    curs7tt = dattabase1.cursor()
                    curs7tt.execute(
                        "SELECT sum(quantitélenplentré) - sum(quantitélenplsorté) as quantitélenplrest FROM outtable  WHERE produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resultlenpl_att = curs7tt.fetchone()
                    self.restlenplttp = str(resultlenpl_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenvertentré) - sum(quantitélenvertsorté) as quantitélenrest FROM outtable WHERE produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resultlenverttt = curs9tt.fetchone()
                    self.restlentvertttp = str(resultlenverttt[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenrougeentré) - sum(quantitélenrougesorté) as quantitélenrougerest FROM outtable WHERE produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resultlenrougett = curs9tt.fetchone()
                    self.restlentrougettp = str(resultlenrougett[0])
                    curs11tt = dattabase1.cursor()
                    curs11tt.execute(
                        "SELECT sum(quantitéharentré) - sum(quantitéharsorté) as quantitéharrest FROM outtable WHERE  produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resulthar_att = curs11tt.fetchone()
                    self.restharicotimpttp = str(resulthar_att[0])
                    curs11tt.execute(
                        "SELECT SUM(quantitéharlsbentré) - SUM(quantitéharlsbsorté) as quantitéharlsbrest FROM outtable   WHERE produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resultharlsb_att = curs11tt.fetchone()
                    self.restharicotlsbttp = str(resultharlsb_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchneufeentré) - sum(quantitépchneufesorté) as quantitépchneuferest FROM outtable  WHERE produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheneufttp = str(resultpch_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchhuitentré) - sum(quantitépchhuitsorté) as quantitépchhuitrest FROM outtable WHERE produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheouitttp = str(resultpch_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenentré) - sum(quantitélensorté) as quantitélenrest FROM outtable WHERE  produit=%s AND unité = %s",
                        (self.selected_produit,self.selected_magasin,))
                    resultlensinqtt = curs9tt.fetchone()
                    self.restlentsinqttp = str(resultlensinqtt[0])
                    ########################################################################
                    # Create a new document and add a table
                    self.doc = docx.Document()
                    section = self.doc.sections[0]
                    section.orientation = WD_ORIENTATION.LANDSCAPE
                    section.page_width = docx.shared.Cm(29.7)
                    section.page_height = docx.shared.Cm(21.0)
                    section.top_margin = docx.shared.Cm(0.5)
                    section.bottom_margin = docx.shared.Cm(1.5)
                    section.right_margin = docx.shared.Cm(1.5)
                    section.left_margin = docx.shared.Cm(1.5)
                    heading1 = self.doc.add_heading(
                        '\t\t\t\t\t\tSITUATION DES STOCKS DES LEGUMES SECS\n\t\t\t\t\t\t\t\tPRODUITS LIVRÉE',
                        level=1)
                    heading1.style.font.name = 'Times New Roman'
                    heading1.style.font.size = Pt(14)
                    heading1.style.font.bold = True
                    heading1.style.font.color.rgb = RGBColor(0, 0, 0)
                    # add second heading

                    paragraph1 = self.doc.add_paragraph()
                    paragraph1.style.font.name = 'Times New Roman'
                    current_date = QDate.currentDate().getDate()
                    left_run1 = paragraph1.add_run(
                        "\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t Date d'édition : {}/{}/{}".format(current_date[2],
                                                                                                current_date[1],
                                                                                                current_date[0]))
                    left_run1.font.size = Pt(9)

                    paragraph = self.doc.add_paragraph()
                    paragraph.style.font.name = 'Times New Roman'
                    paragraph.style.font.size = Pt(12)
                    left_run = paragraph.add_run(self.selected_magasin+"\nPRODUIT :" + self.selected_produit + '\t\t\t\t\t\t\t\t\t\t\t\t\t')
                    left_run.bold = True

                    table = self.doc.add_table(rows=1, cols=11)
                    table.style = "Table Grid"  # set the table style
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'DATE'
                    hdr_cells[1].text = 'PRODUIT'
                    hdr_cells[2].text = 'ORIGINE'
                    hdr_cells[3].text = 'QUANTITE'
                    hdr_cells[4].text = 'MAGASIN DE STOCK'
                    hdr_cells[5].text = 'SASHERIE  CND'
                    hdr_cells[6].text = 'EXPIDITEUR'
                    hdr_cells[7].text = 'IMATRICUL'
                    hdr_cells[8].text = 'Date de récolte'
                    hdr_cells[9].text = 'Date de CND'
                    hdr_cells[10].text = 'Date de DLUM'
                    # Set the width of the header cells
                    hdr_cells[0].width = Inches(0.5)  # set the width of the first column to 1.5 inches
                    hdr_cells[1].width = Inches(2)
                    hdr_cells[2].width = Inches(1)
                    hdr_cells[3].width = Inches(0.4)
                    hdr_cells[4].width = Inches(1.3)
                    hdr_cells[5].width = Inches(1)
                    hdr_cells[6].width = Inches(1)
                    hdr_cells[7].width = Inches(1)
                    hdr_cells[8].width = Inches(0.5)
                    hdr_cells[9].width = Inches(0.5)
                    hdr_cells[10].width = Inches(0.5)

                    for cell in table.rows[0].cells:
                        cell.paragraphs[0].runs[0].font.name = 'Calibri'
                        cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                        cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                        cell.paragraphs[0].runs[0].font.bold = True
                        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    # Add the data to the table
                    table_rows = len(rows)
                    table_cols = len(rows[0])
                    for row in range(table_rows):
                        table.add_row()
                        for col in range(table_cols):
                            cell = table.cell(row + 1, col)
                            cell.text = str(rows[row][col])
                            # Set font properties
                            cell.paragraphs[0].runs[0].font.name = 'Calibri'
                            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
                            cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                            cell.paragraphs[0].runs[0].font.bold = False
                            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    for row in table.rows:
                        row.height = Inches(0.3)
                    # Create the document
                    # Create the table
                    self.doc.add_paragraph().add_run().add_break()
                    table_b = self.doc.add_table(rows=2, cols=2)
                    table_b.style = "Table Grid"
                    # Add table headers
                    hdr_colum = table_b.rows[0].cells
                    hdr_colum[0].text = "PRODUIT"
                    hdr_colum[1].text = "STOCK"
                    for hdr_colum in table_b.rows[0].cells:
                        hdr_colum.paragraphs[0].runs[0].font.name = 'Calibri'
                        hdr_colum.paragraphs[0].runs[0].font.size = Pt(10.5)
                        hdr_colum.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                        hdr_colum.paragraphs[0].runs[0].font.bold = True
                        hdr_colum.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    hdr_colum = table_b.rows[1].cells

                    if self.selected_produit == "POIS CHICHES IMP 12 mm":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restpoichichedouzettp + "  QX"

                    if self.selected_produit == "POIS CHICHES IMP 12 mm":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restpoichichedouzettp + "  QX"
                    if self.selected_produit == "POIS CHICHES IMP 09 mm":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restpoichicheneufttp + " QX"
                    if self.selected_produit == "POIS CHICHES IMP 08 mm":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restpoichicheouitttp + " QX"
                    if self.selected_produit == "POIS CHICHES IMP 06 mm":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restpoichichettp + " QX"
                    if self.selected_produit == "HARICOT BLANC IMP":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restharicotimpttp + " QX"
                    if self.selected_produit == "HARICOT LSB":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restharicotlsbttp + " QX"
                    if self.selected_produit == "LENTILLE PL":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restlenplttp + " QX"
                    if self.selected_produit == "LENTILLE IMP 05-07 mm":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restlentsinqttp + " QX"
                    if self.selected_produit == "LENTILLE IMP vert":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restlentvertttp + " QX"
                    if self.selected_produit == "LENTILLE IMP ROUGE":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.restlentrougettp + " QX"
                    if self.selected_produit == "RIZ IMP":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.resulxttp + " QX"
                    if self.selected_produit == "RIZ ETUVE":
                        hdr_colum[0].text = self.selected_produit
                        hdr_colum[1].text = self.resulxtttvp + " QX"
                    for row in table_b.rows[1:2]:
                        for cell in row.cells:
                            cell.paragraphs[0].runs[0].font.name = 'Calibri'
                            cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                            cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                            cell.paragraphs[0].runs[0].font.bold = False
                    # Save and open the document for printing
                    temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                    self.doc.save(temp_file)
                    os.startfile(temp_file, "open")
                    conn.close()
                elif self.selected_produit=="" and self.selected_magasin=="" and self.selected_origine and self.selected_dates=="--" and self.selected_datesAU=="--":
                    dattabase1 = mysql.connector.connect(host="localhost",  user=user, password=password,
                                                         database='datta_legumsec_out')
                    cursor = dattabase1.cursor()
                    cursor.execute(
                        "SELECT le,produit,origine,quantité,produit,sacherie,fournisseur, imatricule,date_recoltes,date_CND, DLUM FROM outtable WHERE origine=%s AND date_recoltes IS NOT NULL AND date_CND IS NOT NULL AND DLUM IS NOT NULL",
                        (self.selected_origine,))
                    rows = cursor.fetchall()
                    curs1x = dattabase1.cursor()
                    curs1x.execute(
                        "SELECT SUM(quantitérizentré) - SUM(quantitérizesorté) as quantitérizerest FROM outtable WHERE origine=%s",
                        (self.selected_origine,))
                    self.resulxs = curs1x.fetchone()
                    self.resulxtto = str(self.resulxs[0])
                    curs1xtv = dattabase1.cursor()
                    curs1xtv.execute(
                        "SELECT SUM(quantitérizetvntré) - SUM(quantitérizetvsorté) as quantitérizetvrest FROM outtable WHERE origine=%s",
                        (self.selected_origine,))
                    self.resulxstv = curs1xtv.fetchone()
                    self.resulxtttvo = str(self.resulxstv[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitepchentré) - sum(quantitépchsorté) as quantitépchrest  FROM outtable   WHERE  origine=%s",
                        (self.selected_origine,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichichedouzetto = str(resultpch_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT SUM(quantitépchsixentré) - SUM(quantitépchsixsorté) as quantitépchsixrest FROM outtable WHERE  origine=%s",
                        (self.selected_origine,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichichetto = str(resultpch_att[0])

                    curs5tt = dattabase1.cursor()
                    curs5tt.execute(
                        "SELECT sum(quantitépchplentré) - sum(quantitépchplsorté) as quantitépchplhrest FROM outtable WHERE  origine=%s",
                        (self.selected_origine,))
                    resultpchpl_att = curs5tt.fetchone()
                    self.restpoischichepltto = str(resultpchpl_att[0])
                    curs7tt = dattabase1.cursor()
                    curs7tt.execute(
                        "SELECT sum(quantitélenplentré) - sum(quantitélenplsorté) as quantitélenplrest FROM outtable WHERE  origine=%s",
                        (self.selected_origine,))
                    resultlenpl_att = curs7tt.fetchone()
                    self.restlenpltto = str(resultlenpl_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenvertentré) - sum(quantitélenvertsorté) as quantitélenrest FROM outtable WHERE  origine=%s",
                        (self.selected_origine,))
                    resultlenverttt = curs9tt.fetchone()
                    self.restlentverttto = str(resultlenverttt[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenrougeentré) - sum(quantitélenrougesorté) as quantitélenrougerest FROM outtable WHERE  origine=%s",
                        (self.selected_origine,))
                    resultlenrougett = curs9tt.fetchone()
                    self.restlentrougetto = str(resultlenrougett[0])
                    curs11tt = dattabase1.cursor()
                    curs11tt.execute(
                        "SELECT sum(quantitéharentré) - sum(quantitéharsorté) as quantitéharrest FROM outtable WHERE  origine=%s",
                        (self.selected_origine,))
                    resulthar_att = curs11tt.fetchone()
                    self.restharicotimptto = str(resulthar_att[0])
                    curs11tt.execute(
                        "SELECT SUM(quantitéharlsbentré) - SUM(quantitéharlsbsorté) as quantitéharlsbrest FROM outtable WHERE  origine=%s",
                        (self.selected_origine,))
                    resultharlsb_att = curs11tt.fetchone()
                    self.restharicotlsbtto = str(resultharlsb_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchneufeentré) - sum(quantitépchneufesorté) as quantitépchneuferest FROM outtable  WHERE origine=%s",
                        (self.selected_origine,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheneuftto = str(resultpch_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchhuitentré) - sum(quantitépchhuitsorté) as quantitépchhuitrest FROM outtable WHERE  origine=%s",
                        (self.selected_origine,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheouittto = str(resultpch_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenentré) - sum(quantitélensorté) as quantitélenrest FROM outtable WHERE  origine=%s",
                        (self.selected_origine,))
                    resultlensinqtt = curs9tt.fetchone()
                    self.restlentsinqtto = str(resultlensinqtt[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT SUM(quantitépchsixentré) - SUM(quantitépchsixsorté) as quantitépchsixrest FROM outtable  WHERE origine=%s",
                        (self.selected_origine,))
                    resultpch_atto = curstt.fetchone()
                    self.restpoichichett = str(resultpch_att[0])

                    curs5tt = dattabase1.cursor()
                    curs5tt.execute(
                        "SELECT sum(quantitépchplentré) - sum(quantitépchplsorté) as quantitépchplhrest FROM outtable   WHERE origine=%s",
                        (self.selected_origine,))
                    resultpchpl_att = curs5tt.fetchone()
                    self.restpoischichepltto = str(resultpchpl_att[0])
                    curs7tt = dattabase1.cursor()
                    curs7tt.execute(
                        "SELECT sum(quantitélenplentré) - sum(quantitélenplsorté) as quantitélenplrest FROM outtable  WHERE origine=%s",
                        (self.selected_origine,))
                    resultlenpl_att = curs7tt.fetchone()
                    self.restlenpltto = str(resultlenpl_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenvertentré) - sum(quantitélenvertsorté) as quantitélenrest FROM outtable WHERE origine=%s",
                        (self.selected_origine,))
                    resultlenverttt = curs9tt.fetchone()
                    self.restlentverttto = str(resultlenverttt[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenrougeentré) - sum(quantitélenrougesorté) as quantitélenrougerest FROM outtable WHERE origine=%s",
                        (self.selected_origine,))
                    resultlenrougett = curs9tt.fetchone()
                    self.restlentrougetto = str(resultlenrougett[0])
                    curs11tt = dattabase1.cursor()
                    curs11tt.execute(
                        "SELECT sum(quantitéharentré) - sum(quantitéharsorté) as quantitéharrest FROM outtable WHERE origine=%s",
                        (self.selected_origine,))
                    resulthar_att = curs11tt.fetchone()
                    self.restharicotimptt = str(resulthar_att[0])
                    curs11tt.execute(
                        "SELECT SUM(quantitéharlsbentré) - SUM(quantitéharlsbsorté) as quantitéharlsbrest FROM outtable   WHERE origine=%s",
                        (self.selected_origine,))
                    resultharlsb_att = curs11tt.fetchone()
                    self.restharicotlsbtto = str(resultharlsb_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchneufeentré) - sum(quantitépchneufesorté) as quantitépchneuferest FROM outtable  WHERE origine=%s",
                        (self.selected_origine,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheneuftto = str(resultpch_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchhuitentré) - sum(quantitépchhuitsorté) as quantitépchhuitrest FROM outtable WHERE origine=%s",
                        (self.selected_origine,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheouittto = str(resultpch_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenentré) - sum(quantitélensorté) as quantitélenrest FROM outtable WHERE origine=%s",
                        (self.selected_origine,))
                    resultlensinqtt = curs9tt.fetchone()
                    self.restlentsinqtto = str(resultlensinqtt[0])
                    ########################################################################
                    # Create a new document and add a table
                    self.doc = docx.Document()
                    section = self.doc.sections[0]
                    section.orientation = WD_ORIENTATION.LANDSCAPE
                    section.page_width = docx.shared.Cm(29.7)
                    section.page_height = docx.shared.Cm(21.0)
                    section.top_margin = docx.shared.Cm(0.5)
                    section.bottom_margin = docx.shared.Cm(1.5)
                    section.right_margin = docx.shared.Cm(1.5)
                    section.left_margin = docx.shared.Cm(1.5)
                    heading1 = self.doc.add_heading(
                        '\t\t\t\t\t\tSITUATION DES STOCKS DES LEGUMES SECS\n\t\t\t\t\t\t\t\tPRODUITS LIVRÉE',
                        level=1)
                    heading1.style.font.name = 'Times New Roman'
                    heading1.style.font.size = Pt(14)
                    heading1.style.font.bold = True
                    heading1.style.font.color.rgb = RGBColor(0, 0, 0)
                    # add second heading


                    paragraph1 = self.doc.add_paragraph()
                    paragraph1.style.font.name = 'Times New Roman'
                    current_date = QDate.currentDate().getDate()
                    left_run1 = paragraph1.add_run(
                        "\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t Date d'édition : {}/{}/{}".format(current_date[2],
                                                                                                current_date[1],
                                                                                                current_date[0]))
                    left_run1.font.size = Pt(9)

                    paragraph = self.doc.add_paragraph()
                    paragraph.style.font.name = 'Times New Roman'
                    paragraph.style.font.size = Pt(12)
                    left_run = paragraph.add_run(self.selected_magasin + '\t\t\t\t\t\t\t\t\t\t\t\t\t')
                    left_run.bold = True

                    table = self.doc.add_table(rows=1, cols=11)
                    table.style = "Table Grid"  # set the table style
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'DATE'
                    hdr_cells[1].text = 'PRODUIT'
                    hdr_cells[2].text = 'ORIGINE'
                    hdr_cells[3].text = 'QUANTITE'
                    hdr_cells[4].text = 'MAGASIN DE STOCK'
                    hdr_cells[5].text = 'SASHERIE  CND'
                    hdr_cells[6].text = 'EXPIDITEUR'
                    hdr_cells[7].text = 'IMATRICUL'
                    hdr_cells[8].text = 'Date de récolte'
                    hdr_cells[9].text = 'Date de CND'
                    hdr_cells[10].text = 'Date de DLUM'
                    # Set the width of the header cells
                    hdr_cells[0].width = Inches(0.5)  # set the width of the first column to 1.5 inches
                    hdr_cells[1].width = Inches(2)
                    hdr_cells[2].width = Inches(1)
                    hdr_cells[3].width = Inches(0.4)
                    hdr_cells[4].width = Inches(1.3)
                    hdr_cells[5].width = Inches(1)
                    hdr_cells[6].width = Inches(1)
                    hdr_cells[7].width = Inches(1)
                    hdr_cells[8].width = Inches(0.5)
                    hdr_cells[9].width = Inches(0.5)
                    hdr_cells[10].width = Inches(0.5)

                    for cell in table.rows[0].cells:
                        cell.paragraphs[0].runs[0].font.name = 'Calibri'
                        cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                        cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                        cell.paragraphs[0].runs[0].font.bold = True
                        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    # Add the data to the table
                    table_rows = len(rows)
                    table_cols = len(rows[0])
                    for row in range(table_rows):
                        table.add_row()
                        for col in range(table_cols):
                            cell = table.cell(row + 1, col)
                            cell.text = str(rows[row][col])
                            # Set font properties
                            cell.paragraphs[0].runs[0].font.name = 'Calibri'
                            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
                            cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                            cell.paragraphs[0].runs[0].font.bold = False
                            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    for row in table.rows:
                        row.height = Inches(0.3)
                    # Create the document
                    # Create the table
                    self.doc.add_paragraph().add_run().add_break()
                    table_b = self.doc.add_table(rows=14, cols=2)
                    table_b.style = "Table Grid"
                    # Add table headers
                    hdr_colum = table_b.rows[0].cells
                    hdr_colum[0].text = "PRODUIT"
                    hdr_colum[1].text = "STOCK"
                    for hdr_colum in table_b.rows[0].cells:
                        hdr_colum.paragraphs[0].runs[0].font.name = 'Calibri'
                        hdr_colum.paragraphs[0].runs[0].font.size = Pt(10.5)
                        hdr_colum.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                        hdr_colum.paragraphs[0].runs[0].font.bold = True
                        hdr_colum.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    hdr_colum = table_b.rows[1].cells
                    hdr_colum[0].text = "POIS CHICHES PL"
                    hdr_colum[1].text = ""
                    hdr_colum = table_b.rows[2].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 12 mm"
                    hdr_colum[1].text = self.restpoichichedouzetto + "  QX"
                    hdr_colum = table_b.rows[3].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 09 mm"
                    hdr_colum[1].text = self.restpoichicheneuftto + " QX"
                    hdr_colum = table_b.rows[4].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 08 mm"
                    hdr_colum[1].text = self.restpoichicheouittto + " QX"
                    hdr_colum = table_b.rows[5].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 06 mm"
                    hdr_colum[1].text = self.restpoichichetto + " QX"
                    hdr_colum = table_b.rows[6].cells
                    hdr_colum[0].text = "HARICOT BLANC IMP"
                    hdr_colum[1].text = self.restharicotimptto + " QX"
                    hdr_colum = table_b.rows[7].cells
                    hdr_colum[0].text = "HARICOT LSB"
                    hdr_colum[1].text = self.restharicotlsbtto + " QX"
                    hdr_colum = table_b.rows[8].cells
                    hdr_colum[0].text = "LENTILLE PL"
                    hdr_colum[1].text = self.restlenpltto + " QX"
                    hdr_colum = table_b.rows[9].cells
                    hdr_colum[0].text = "LENTILLE IMP 05-07 mm"
                    hdr_colum[1].text = self.restlentsinqtto + " QX"
                    hdr_colum = table_b.rows[10].cells
                    hdr_colum[0].text = "LENTILLE IMP vert"
                    hdr_colum[1].text = self.restlentverttto + " QX"
                    hdr_colum = table_b.rows[11].cells
                    hdr_colum[0].text = "LENTILLE IMP ROUGE"
                    hdr_colum[1].text = self.restlentrougetto + " QX"
                    hdr_colum = table_b.rows[12].cells
                    hdr_colum[0].text = "RIZ IMP"
                    hdr_colum[1].text = self.resulxtto + " QX"
                    hdr_colum = table_b.rows[13].cells
                    hdr_colum[0].text = "RIZ ETUVE"
                    hdr_colum[1].text = self.resulxtto + " QX"
                    for row in table_b.rows[1:14]:
                        for cell in row.cells:
                            cell.paragraphs[0].runs[0].font.name = 'Calibri'
                            cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                            cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                            cell.paragraphs[0].runs[0].font.bold = False
                    # Save and open the document for printing
                    temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                    self.doc.save(temp_file)
                    os.startfile(temp_file, "open")
                    dattabase1.close()
                else:
                    print("hello world")
                    conn = mysql.connector.connect(host="localhost",  user=user, password=password,
                                                   database="datta_legumesec_entry")
                    dattabase1 = mysql.connector.connect(host="localhost",  user=user, password=password,
                                                         database='datta_legumsec_out')
                    cursor = conn.cursor()
                    cursor.execute(
                        "SELECT le,produit,origine,quantité,produit,sacherie,fournisseur, imatricule, date_CND, DLUM FROM entrytable WHERE fournisseur=%s",
                        (self.selected_expediteur,))
                    rows = cursor.fetchall()
                    curs1x = dattabase1.cursor()
                    curs1x.execute(
                        "SELECT SUM(quantitérizentré) - SUM(quantitérizesorté) as quantitérizerest FROM outtable WHERE fournisseur=%s",
                        (self.selected_expediteur,))
                    self.resulxs = curs1x.fetchone()
                    self.resulxttex = str(self.resulxs[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitepchentré) - sum(quantitépchsorté) as quantitépchrest  FROM outtable   WHERE  fournisseur=%s",
                        (self.selected_expediteur,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichichedouzettex = str(resultpch_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT SUM(quantitépchsixentré) - SUM(quantitépchsixsorté) as quantitépchsixrest FROM outtable WHERE  fournisseur=%s",
                        (self.selected_expediteur,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichichettex = str(resultpch_att[0])

                    curs5tt = dattabase1.cursor()
                    curs5tt.execute(
                        "SELECT sum(quantitépchplentré) - sum(quantitépchplsorté) as quantitépchplhrest FROM outtable WHERE  fournisseur=%s",
                        (self.selected_expediteur,))
                    resultpchpl_att = curs5tt.fetchone()
                    self.restpoischicheplttex = str(resultpchpl_att[0])
                    curs7tt = dattabase1.cursor()
                    curs7tt.execute(
                        "SELECT sum(quantitélenplentré) - sum(quantitélenplsorté) as quantitélenplrest FROM outtable WHERE  fournisseur=%s",
                        (self.selected_expediteur,))
                    resultlenpl_att = curs7tt.fetchone()
                    self.restlenplttex = str(resultlenpl_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenvertentré) - sum(quantitélenvertsorté) as quantitélenrest FROM outtable WHERE  fournisseur=%s",
                        (self.selected_expediteur,))
                    resultlenverttt = curs9tt.fetchone()
                    self.restlentvertttex = str(resultlenverttt[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenrougeentré) - sum(quantitélenrougesorté) as quantitélenrougerest FROM outtable WHERE  fournisseur=%s",
                        (self.selected_expediteur,))
                    resultlenrougett = curs9tt.fetchone()
                    self.restlentrougettex = str(resultlenrougett[0])
                    curs11tt = dattabase1.cursor()
                    curs11tt.execute(
                        "SELECT sum(quantitéharentré) - sum(quantitéharsorté) as quantitéharrest FROM outtable WHERE  fournisseur=%s",
                        (self.selected_expediteur,))
                    resulthar_att = curs11tt.fetchone()
                    self.restharicotimpttex = str(resulthar_att[0])
                    curs11tt.execute(
                        "SELECT SUM(quantitéharlsbentré) - SUM(quantitéharlsbsorté) as quantitéharlsbrest FROM outtable WHERE  fournisseur=%s",
                        (self.selected_expediteur,))
                    resultharlsb_att = curs11tt.fetchone()
                    self.restharicotlsbttex = str(resultharlsb_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchneufeentré) - sum(quantitépchneufesorté) as quantitépchneuferest FROM outtable  WHERE fournisseur=%s",
                        (self.selected_expediteur,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheneufttex = str(resultpch_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchhuitentré) - sum(quantitépchhuitsorté) as quantitépchhuitrest FROM outtable WHERE  fournisseur=%s",
                        (self.selected_expediteur,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheouitttex = str(resultpch_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenentré) - sum(quantitélensorté) as quantitélenrest FROM outtable WHERE  fournisseur=%s",
                        (self.selected_expediteur,))
                    resultlensinqtt = curs9tt.fetchone()
                    self.restlentsinqttex = str(resultlensinqtt[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT SUM(quantitépchsixentré) - SUM(quantitépchsixsorté) as quantitépchsixrest FROM outtable  WHERE fournisseur=%s",
                        (self.selected_expediteur,))
                    resultpch_attex = curstt.fetchone()
                    self.restpoichichett = str(resultpch_att[0])

                    curs5tt = dattabase1.cursor()
                    curs5tt.execute(
                        "SELECT sum(quantitépchplentré) - sum(quantitépchplsorté) as quantitépchplhrest FROM outtable   WHERE fournisseur=%s",
                        (self.selected_expediteur,))
                    resultpchpl_att = curs5tt.fetchone()
                    self.restpoischicheplttex = str(resultpchpl_att[0])
                    curs7tt = dattabase1.cursor()
                    curs7tt.execute(
                        "SELECT sum(quantitélenplentré) - sum(quantitélenplsorté) as quantitélenplrest FROM outtable  WHERE fournisseur=%s",
                        (self.selected_expediteur,))
                    resultlenpl_att = curs7tt.fetchone()
                    self.restlenplttex = str(resultlenpl_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenvertentré) - sum(quantitélenvertsorté) as quantitélenrest FROM outtable WHERE fournisseur=%s",
                        (self.selected_expediteur,))
                    resultlenverttt = curs9tt.fetchone()
                    self.restlentvertttex = str(resultlenverttt[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenrougeentré) - sum(quantitélenrougesorté) as quantitélenrougerest FROM outtable WHERE fournisseur=%s",
                        (self.selected_expediteur,))
                    resultlenrougett = curs9tt.fetchone()
                    self.restlentrougettex = str(resultlenrougett[0])
                    curs11tt = dattabase1.cursor()
                    curs11tt.execute(
                        "SELECT sum(quantitéharentré) - sum(quantitéharsorté) as quantitéharrest FROM outtable WHERE fournisseur=%s",
                        (self.selected_expediteur,))
                    resulthar_att = curs11tt.fetchone()
                    self.restharicotimptt = str(resulthar_att[0])
                    curs11tt.execute(
                        "SELECT SUM(quantitéharlsbentré) - SUM(quantitéharlsbsorté) as quantitéharlsbrest FROM outtable   WHERE fournisseur=%s",
                        (self.selected_expediteur,))
                    resultharlsb_att = curs11tt.fetchone()
                    self.restharicotlsbttex = str(resultharlsb_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchneufeentré) - sum(quantitépchneufesorté) as quantitépchneuferest FROM outtable  WHERE fournisseur=%s",
                        (self.selected_expediteur,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheneufttex = str(resultpch_att[0])
                    curstt = dattabase1.cursor()
                    curstt.execute(
                        "SELECT sum(quantitépchhuitentré) - sum(quantitépchhuitsorté) as quantitépchhuitrest FROM outtable WHERE fournisseur=%s",
                        (self.selected_expediteur,))
                    resultpch_att = curstt.fetchone()
                    self.restpoichicheouitttex = str(resultpch_att[0])
                    curs9tt = dattabase1.cursor()
                    curs9tt.execute(
                        "SELECT sum(quantitélenentré) - sum(quantitélensorté) as quantitélenrest FROM outtable WHERE fournisseur=%s",
                        (self.selected_expediteur,))
                    resultlensinqtt = curs9tt.fetchone()
                    self.restlentsinqttex = str(resultlensinqtt[0])
                    ########################################################################
                    # Create a new document and add a table
                    self.doc = docx.Document()
                    section = self.doc.sections[0]
                    section.orientation = WD_ORIENTATION.LANDSCAPE
                    section.page_width = docx.shared.Cm(29.7)
                    section.page_height = docx.shared.Cm(21.0)
                    section.top_margin = docx.shared.Cm(0.5)
                    section.bottom_margin = docx.shared.Cm(1.5)
                    section.right_margin = docx.shared.Cm(1.5)
                    section.left_margin = docx.shared.Cm(1.5)
                    heading1 = self.doc.add_heading(
                        '\t\t\t\t\t\tSITUATION DES STOCKS DES LEGUMES SECS', level=1)
                    heading1.style.font.name = 'Times New Roman'
                    heading1.style.font.size = Pt(18)
                    heading1.style.font.bold = True
                    heading1.style.font.color.rgb = RGBColor(0, 0, 0)
                    # add second heading
                    paragraph0 = self.doc.add_paragraph()
                    paragraph0.style.font.name = 'Times New Roman'
                    left_run0 = paragraph0.add_run(
                        'Période du :' + self.selected_dates + "  au:" + self.selected_datesAU)
                    paragraph0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    paragraph1 = self.doc.add_paragraph()
                    paragraph1.style.font.name = 'Times New Roman'
                    current_date = QDate.currentDate().getDate()
                    left_run1 = paragraph1.add_run(
                        "\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t Date d'édition : {}/{}/{}".format(current_date[2],
                                                                                                current_date[1],
                                                                                                current_date[0]))
                    left_run1.font.size = Pt(9)

                    paragraph = self.doc.add_paragraph()
                    paragraph.style.font.name = 'Times New Roman'
                    paragraph.style.font.size = Pt(12)
                    left_run = paragraph.add_run(self.selected_magasin + '\t\t\t\t\t\t\t\t\t\t\t\t\t')
                    left_run.bold = True

                    table = self.doc.add_table(rows=1, cols=11)
                    table.style = "Table Grid"  # set the table style
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'DATE'
                    hdr_cells[1].text = 'PRODUIT'
                    hdr_cells[2].text = 'ORIGINE'
                    hdr_cells[3].text = 'QUANTITE'
                    hdr_cells[4].text = 'MAGASIN DE STOCK'
                    hdr_cells[5].text = 'SASHERIE  CND'
                    hdr_cells[6].text = 'EXPIDITEUR'
                    hdr_cells[7].text = 'IMATRICUL'
                    hdr_cells[8].text = 'Date de récolte'
                    hdr_cells[9].text = 'Date de CND'
                    hdr_cells[10].text = 'Date de DLUM'
                    # Set the width of the header cells
                    hdr_cells[0].width = Inches(0.5)  # set the width of the first column to 1.5 inches
                    hdr_cells[1].width = Inches(2)
                    hdr_cells[2].width = Inches(1)
                    hdr_cells[3].width = Inches(0.4)
                    hdr_cells[4].width = Inches(1.3)
                    hdr_cells[5].width = Inches(1)
                    hdr_cells[6].width = Inches(1)
                    hdr_cells[7].width = Inches(1)
                    hdr_cells[8].width = Inches(0.5)
                    hdr_cells[9].width = Inches(0.5)
                    hdr_cells[10].width = Inches(0.5)

                    for cell in table.rows[0].cells:
                        cell.paragraphs[0].runs[0].font.name = 'Calibri'
                        cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                        cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                        cell.paragraphs[0].runs[0].font.bold = True
                        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    # Add the data to the table
                    table_rows = len(rows)
                    table_cols = len(rows[0])
                    for row in range(table_rows):
                        table.add_row()
                        for col in range(table_cols):
                            cell = table.cell(row + 1, col)
                            cell.text = str(rows[row][col])
                            # Set font properties
                            cell.paragraphs[0].runs[0].font.name = 'Calibri'
                            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
                            cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                            cell.paragraphs[0].runs[0].font.bold = False
                            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    for row in table.rows:
                        row.height = Inches(0.3)
                    # Create the document
                    # Create the table
                    self.doc.add_paragraph().add_run().add_break()
                    table_b = self.doc.add_table(rows=14, cols=2)
                    table_b.style = "Table Grid"
                    # Add table headers
                    hdr_colum = table_b.rows[0].cells
                    hdr_colum[0].text = "PRODUIT"
                    hdr_colum[1].text = "STOCK"
                    for hdr_colum in table_b.rows[0].cells:
                        hdr_colum.paragraphs[0].runs[0].font.name = 'Calibri'
                        hdr_colum.paragraphs[0].runs[0].font.size = Pt(10.5)
                        hdr_colum.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                        hdr_colum.paragraphs[0].runs[0].font.bold = True
                        hdr_colum.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    hdr_colum = table_b.rows[1].cells
                    hdr_colum[0].text = "POIS CHICHES PL"
                    hdr_colum[1].text = ""
                    hdr_colum = table_b.rows[2].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 12 mm"
                    hdr_colum[1].text = self.restpoichichedouzetto + "  QX"
                    hdr_colum = table_b.rows[3].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 09 mm"
                    hdr_colum[1].text = self.restpoichicheneuftto + " QX"
                    hdr_colum = table_b.rows[4].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 08 mm"
                    hdr_colum[1].text = self.restpoichicheouittto + " QX"
                    hdr_colum = table_b.rows[5].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 06 mm"
                    hdr_colum[1].text = self.restpoichichetto + " QX"
                    hdr_colum = table_b.rows[6].cells
                    hdr_colum[0].text = "HARICOT BLANC IMP"
                    hdr_colum[1].text = self.restharicotimptto + " QX"
                    hdr_colum = table_b.rows[7].cells
                    hdr_colum[0].text = "HARICOT LSB"
                    hdr_colum[1].text = self.restharicotlsbtto + " QX"
                    hdr_colum = table_b.rows[8].cells
                    hdr_colum[0].text = "LENTILLE PL"
                    hdr_colum[1].text = self.restlenpltto + " QX"
                    hdr_colum = table_b.rows[9].cells
                    hdr_colum[0].text = "LENTILLE IMP 05-07 mm"
                    hdr_colum[1].text = self.restlentsinqtto + " QX"
                    hdr_colum = table_b.rows[10].cells
                    hdr_colum[0].text = "LENTILLE IMP vert"
                    hdr_colum[1].text = self.restlentverttto + " QX"
                    hdr_colum = table_b.rows[11].cells
                    hdr_colum[0].text = "LENTILLE IMP ROUGE"
                    hdr_colum[1].text = self.restlentrougetto + " QX"
                    hdr_colum = table_b.rows[12].cells
                    hdr_colum[0].text = "RIZ IMP"
                    hdr_colum[1].text = self.resulxtto + " QX"
                    hdr_colum = table_b.rows[13].cells
                    hdr_colum[0].text = "RIZ IMP"
                    hdr_colum[1].text = self.resulxtto + " QX"
                    for row in table_b.rows[1:14]:
                        for cell in row.cells:
                            cell.paragraphs[0].runs[0].font.name = 'Calibri'
                            cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                            cell.paragraphs[0].runs[0].font.spacing = Pt(0.1)
                            cell.paragraphs[0].runs[0].font.bold = False
                    # Save and open the document for printing
                    temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                    self.doc.save(temp_file)
                    os.startfile(temp_file, "open")
                    conn.close()

            except Exception as e:
                print(e)

        def operation_dattabase(self):
            try:
                # Riz operation
                with mysql.connector.connect(host="localhost",  user=user, password=password,
                                             database='datta_legumesec_entry') as dattabase:
                    with dattabase.cursor() as curs:
                        curs.execute("SELECT SUM(quantitérizentrée) FROM entrytable")
                        result = curs.fetchone()

                with mysql.connector.connect(host="localhost",  user=user, password=password,
                                             database='datta_legumsec_out') as dattabase1:
                    with dattabase1.cursor() as curs1:
                        curs1.execute(
                            "SELECT SUM(quantitérizentré) - SUM(quantitérizesorté) as quantitérizerest FROM outtable")
                        self.resul = curs1.fetchone()
                        self.resultt =str(self.resul[0])
                    with dattabase1.cursor() as curs:
                        curs.execute(
                            "SELECT SUM(quantitérizentré) - SUM(quantitérizesorté) as quantitérizerest FROM outtable WHERE produit=%s AND unité=%s OR origine=%s OR fournisseur=%s",
                            (self.Produits.currentText(), self.Unites.currentText(),
                             self.OrigineDuProduits.currentText(), self.expediteur.currentText()))
                        resultpch_a = curs.fetchone()
                        self.resultstock = str(resultpch_a[0])
                        self.resulttext = str(self.resul[0])
                    with dattabase1.cursor() as curstt:
                        curstt.execute(
                            "SELECT SUM(quantitérizentré) - SUM(quantitérizesorté) as quantitérizerest FROM outtable WHERE produit=%s AND unité=%s OR origine=%s OR fournisseur=%s",
                            (self.Produits.currentText(), self.Unites.currentText(),
                             self.OrigineDuProduits.currentText(), self.expediteur.currentText()))
                        resultpch_ax = curstt.fetchone()
                        self.resulxtt = str(resultpch_ax[0])
                    ########################for docx out
                    with mysql.connector.connect(host="localhost",  user=user, password=password,
                                                 database='datta_legumsec_out') as dattabase1:
                        with dattabase1.cursor() as curs1x:
                            curs1x.execute(
                                "SELECT SUM(quantitérizentré) - SUM(quantitérizesorté) as quantitérizerest FROM outtable WHERE produit=%s AND unité=%s AND le BETWEEN %s AND %s", (self.selected_produit, self.selected_magasin,self.selected_dates,self.selected_datesAU))
                            self.resulxs = curs1x.fetchone()
                            self.resulx=str(self.resulxs[0])

            except mysql.connector.Error as e:
                # Handle the error more gracefully, for example, log it or raise a custom exception
                print(f"An error occurred: {e}")

        def rizeEtuve(self):
            try:
                # Riz operation
                with mysql.connector.connect(host="localhost",  user=user, password=password,
                                             database='datta_legumesec_entry') as dattabase:
                    with dattabase.cursor() as curs:
                        curs.execute("SELECT SUM(quantitérizetvntrée) FROM entrytable")
                        result = curs.fetchone()

                with mysql.connector.connect(host="localhost",  user=user, password=password,
                                             database='datta_legumsec_out') as dattabase1:
                    with dattabase1.cursor() as curs1:
                        curs1.execute(
                            "SELECT SUM(quantitérizetvntré) - SUM(quantitérizetvsorté) as quantitérizetvrest FROM outtable")
                        resulterzetv = curs1.fetchone()
                        self.resuletv=str(resulterzetv[0])
                    with dattabase1.cursor() as curs:
                        curs.execute(
                            "SELECT SUM(quantitérizetvntré) - SUM(quantitérizetvsorté) as quantitérizetvrest FROM outtable WHERE produit=%s AND unité=%s OR origine=%s OR fournisseur=%s",
                            (self.Produits.currentText(), self.Unites.currentText(),
                             self.OrigineDuProduits.currentText(), self.expediteur.currentText()))
                        resultpch_a = curs.fetchone()
                        self.resultstocktv = str(resultpch_a[0])
                        self.resulttextv = str(self.resul[0])
                    with dattabase1.cursor() as curstt:
                        curstt.execute(
                            "SELECT SUM(quantitérizetvntré) - SUM(quantitérizetvsorté) as quantitérizetvrest FROM outtable WHERE produit=%s AND unité=%s OR origine=%s OR fournisseur=%s",
                            (self.Produits.currentText(), self.Unites.currentText(),
                             self.OrigineDuProduits.currentText(), self.expediteur.currentText()))
                        resultpch_axtv = curstt.fetchone()
                        self.resulxttv = str(resultpch_axtv[0])
                    ########################for docx out
                    with mysql.connector.connect(host="localhost",  user=user, password=password,
                                                 database='datta_legumsec_out') as dattabase1:
                        with dattabase1.cursor() as curs1x:
                            curs1x.execute(
                                "SELECT SUM(quantitérizetvntré) - SUM(quantitérizetvsorté) as quantitérizetvrest FROM outtable WHERE produit=%s AND unité=%s AND le BETWEEN %s AND %s", (self.selected_produit, self.selected_magasin,self.selected_dates,self.selected_datesAU))
                            self.resulxstv = curs1x.fetchone()
                            self.resulxtv=str(self.resulxstv[0])

            except mysql.connector.Error as e:
                # Handle the error more gracefully, for example, log it or raise a custom exception
                print(f"An error occurred: {e}")

        def poichichSIxSept(self):
            try:
                # Pois chiches operation
                dattabase = mysql.connector.connect(host="localhost",  user=user, password=password,database='datta_legumesec_entry')
                dattabase1 = mysql.connector.connect(host="localhost",  user=user, password=password,database='datta_legumsec_out')
                curs=dattabase.cursor()
                curs.execute("SELECT SUM(quantitépchsixntrée) AS quantitépchsixtotal FROM entrytable")
                resultpch = curs.fetchone()
                curs=dattabase1.cursor()
                curs.execute("SELECT SUM(quantitépchsixentré) - SUM(quantitépchsixsorté) as quantitépchsixrest FROM outtable WHERE produit=%s AND unité=%s OR origine=%s OR fournisseur=%s",(self.Produits.currentText(), self.Unites.currentText(),self.OrigineDuProduits.currentText(), self.expediteur.currentText()))
                resultpch_a = curs.fetchone()
                self.restpoichiche = str(resultpch_a[0])
                curs = dattabase1.cursor()
                curs.execute(
                    "SELECT SUM(quantitépchsixentré) - SUM(quantitépchsixsorté) as quantitépchsixrest FROM outtable ")
                resultpch_at = curs.fetchone()
                self.restpoichichett = str(resultpch_at[0])
                ##################for docx out
                curstt= dattabase1.cursor()
                curstt.execute("SELECT SUM(quantitépchsixentré) - SUM(quantitépchsixsorté) as quantitépchsixrest FROM outtable WHERE produit=%s AND unité=%s AND le BETWEEN %s AND %s", (self.selected_produit, self.selected_magasin,self.selected_dates,self.selected_datesAU))
                resultpch_att = curstt.fetchone()
                self.restpoichichettx = str(resultpch_att[0])
            except Exception as e:
                # Handle the error more gracefully, for example, log it or raise a custom exception
                print(f"An error occurred: {e}")

        def poichichepl(self):
            try:
                ####################pois chichz pl operation
                dattabase = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumesec_entry')
                dattabase1=mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumsec_out')
                curs4 = dattabase.cursor()
                curs4.execute("SELECT SUM(quantitépchplentrée) AS quantitépchpltotale FROM entrytable ")
                resultpchpl = curs4.fetchone()

                dattabase.commit()
                curs5 =dattabase1.cursor()
                curs5.execute(
                    "SELECT sum(quantitépchplentré) - sum(quantitépchplsorté) as quantitépchplhrest FROM outtable WHERE produit=%s AND unité=%s OR origine=%s OR fournisseur=%s", (self.Produits.currentText(), self.Unites.currentText(),self.OrigineDuProduits.currentText(),self.expediteur.currentText()))
                resultpchpl_a = curs5.fetchone()
                self.restpoischichepl=str(resultpchpl_a[0])
                dattabase.commit()
                curs5 = dattabase1.cursor()
                curs5.execute(
                    "SELECT sum(quantitépchplentré) - sum(quantitépchplsorté) as quantitépchplhrest FROM outtable ")
                resultpchpl_at = curs5.fetchone()
                self.restpoischichepltt = str(resultpchpl_at[0])
                dattabase.commit()
                #####for docx out
                curs5tt = dattabase1.cursor()
                curs5tt.execute(
                    "SELECT sum(quantitépchplentré) - sum(quantitépchplsorté) as quantitépchplhrest FROM outtable WHERE produit=%s AND unité=%s AND le BETWEEN %s AND %s", (self.selected_produit, self.selected_magasin,self.selected_dates,self.selected_datesAU))
                resultpchpl_att = curs5tt.fetchone()
                self.restpoischicheplttx = str(resultpchpl_att[0])
                dattabase.commit()
                dattabase.close()

            except Exception as e:
                print(e)

        def lenpl(self):
            try:
                ###############lentille pl operation
                dattabase = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumesec_entry')
                dattabase1 = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumsec_out')
                curs6 = dattabase.cursor()
                curs6.execute("SELECT SUM(quantitélenplntrée) AS quantitélenpltotale FROM entrytable ")
                resultlenpl = curs6.fetchone()

                dattabase.commit()

                curs7 = dattabase1.cursor()
                curs7.execute(
                    "SELECT sum(quantitélenplentré) - sum(quantitélenplsorté) as quantitélenplrest FROM outtable WHERE produit=%s AND unité=%s OR origine=%s OR fournisseur=%s", (self.Produits.currentText(), self.Unites.currentText(),self.OrigineDuProduits.currentText(),self.expediteur.currentText()))
                resultlenpl_a = curs7.fetchone()
                self.restlenpl=str(resultlenpl_a[0])
                dattabase.commit()
                curs7 = dattabase1.cursor()
                curs7.execute( "SELECT sum(quantitélenplentré) - sum(quantitélenplsorté) as quantitélenplrest FROM outtable")
                resultlenpl_att = curs7.fetchone()
                self.restlenpltt = str(resultlenpl_att[0])
                dattabase.commit()
                ###############for docx out
                curs7tt = dattabase1.cursor()
                curs7tt.execute(
                    "SELECT sum(quantitélenplentré) - sum(quantitélenplsorté) as quantitélenplrest FROM outtable WHERE produit=%s AND unité=%s AND le BETWEEN %s AND %s", (self.selected_produit, self.selected_magasin,self.selected_dates,self.selected_datesAU))
                resultlenpl_at = curs7tt.fetchone()
                self.restlenplttx = str(resultlenpl_at[0])
                dattabase.commit()
                dattabase.close()

            except Exception as e:
                print(e)

        def lenVert(self):
            try:
                ###########lentille
                dattabase = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumesec_entry')
                dattabase1 = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumsec_out')
                curs8 = dattabase.cursor()
                curs8.execute("SELECT SUM(quantitélenvertntrée) AS quantitélentotale FROM entrytable ")
                resultlen = curs8.fetchone()

                dattabase.commit()
                curs9 = dattabase1.cursor()
                curs9.execute("SELECT sum(quantitélenvertentré) - sum(quantitélenvertsorté) as quantitélenrest FROM outtable WHERE produit=%s AND unité=%s OR origine=%s OR fournisseur=%s", (self.Produits.currentText(), self.Unites.currentText(),self.OrigineDuProduits.currentText(),self.expediteur.currentText()))
                resultlenvert = curs9.fetchone()
                self.restlentvert=str(resultlenvert[0])
                dattabase.commit()
                dattabase.commit()
                curs9 = dattabase1.cursor()
                curs9.execute(
                    "SELECT sum(quantitélenvertentré) - sum(quantitélenvertsorté) as quantitélenrest FROM outtable ")
                resultlenvertt = curs9.fetchone()
                self.restlentvertt = str(resultlenvertt[0])
                dattabase.commit()
                #################for docx out
                curs9tt = dattabase1.cursor()
                curs9tt.execute(
                    "SELECT sum(quantitélenvertentré) - sum(quantitélenvertsorté) as quantitélenrest FROM outtable WHERE produit=%s AND unité=%s AND le BETWEEN %s AND %s", (self.selected_produit, self.selected_magasin,self.selected_dates,self.selected_datesAU))
                resultlenverttt = curs9tt.fetchone()
                self.restlentverttx = str(resultlenverttt[0])
                dattabase.commit()
                dattabase.close()

            except Exception as e:
                print(e)


        def lenRouge(self):
            try:
                ###########lentilleRouge
                dattabase = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumesec_entry')
                dattabase1 = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumsec_out')
                curs8 = dattabase.cursor()
                curs8.execute("SELECT SUM(quantitélenrougeentrée) AS quantitélenrougetotale FROM entrytable ")
                resultlen = curs8.fetchone()

                dattabase.commit()
                curs9 = dattabase1.cursor()
                curs9.execute("SELECT sum(quantitélenrougeentré) - sum(quantitélenrougesorté) as quantitélenrougerest FROM outtable WHERE produit=%s AND unité=%s OR origine=%s OR fournisseur=%s", (self.Produits.currentText(), self.Unites.currentText(),self.OrigineDuProduits.currentText(),self.expediteur.currentText()))
                resultlenrouge = curs9.fetchone()
                self.restlentrouge=str(resultlenrouge[0])
                dattabase.commit()
                dattabase.commit()
                curs9 = dattabase1.cursor()
                curs9.execute(
                    "SELECT sum(quantitélenrougeentré) - sum(quantitélenrougesorté) as quantitélenrougerest FROM outtable ")
                resultlenrouget = curs9.fetchone()
                self.restlentrougett = str(resultlenrouget[0])
                dattabase.commit()
                #################for docx out
                curs9tt = dattabase1.cursor()
                curs9tt.execute(
                    "SELECT sum(quantitélenrougeentré) - sum(quantitélenrougesorté) as quantitélenrougerest FROM outtable WHERE produit=%s AND unité=%s AND le BETWEEN %s AND %s", (self.selected_produit, self.selected_magasin,self.selected_dates,self.selected_datesAU))
                resultlenrougettx = curs9tt.fetchone()
                self.restlentrougettx = str(resultlenrougettx[0])
                dattabase.commit()
                dattabase.close()

            except Exception as e:
                print(e)

        def haricot(self):
            try:
                ####################haricot operation
                dattabase = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumesec_entry')
                dattabase1 = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumsec_out')
                curs10 = dattabase.cursor()
                curs10.execute("SELECT SUM(quantitéharntrée) AS quantitéhartotale FROM entrytable ")
                resulthar = curs10.fetchone()

                dattabase.commit()
                curs11 = dattabase1.cursor()
                curs11.execute("SELECT sum(quantitéharentré) - sum(quantitéharsorté) as quantitéharrest FROM outtable WHERE produit=%s AND unité=%s OR origine=%s OR fournisseur=%s", (self.Produits.currentText(), self.Unites.currentText(),self.OrigineDuProduits.currentText(),self.expediteur.currentText()))
                resulthar_a = curs11.fetchone()
                self.restharicotimp=str(resulthar_a[0])
                dattabase.commit()
                dattabase.commit()
                curs11 = dattabase1.cursor()
                curs11.execute("SELECT sum(quantitéharentré) - sum(quantitéharsorté) as quantitéharrest FROM outtable")
                resulthar_at = curs11.fetchone()
                self.restharicotimptt = str(resulthar_at[0])
                dattabase.commit()
                ############for docx out
                curs11tt = dattabase1.cursor()
                curs11tt.execute(
                    "SELECT sum(quantitéharentré) - sum(quantitéharsorté) as quantitéharrest FROM outtable WHERE produit=%s AND unité=%s AND le BETWEEN %s AND %s", (self.selected_produit, self.selected_magasin,self.selected_dates,self.selected_datesAU))
                resulthar_att = curs11tt.fetchone()
                self.restharicotimpttx = str(resulthar_att[0])
                dattabase.commit()
                dattabase.close()
            except Exception as e:
                print(e)

        def haricotLsb(self):
            try:
                # Connect to the first database
                with mysql.connector.connect(host="localhost",  user=user, password=password,
                                             database='datta_legumesec_entry') as dattabase:
                    with dattabase.cursor() as curs10:
                        # Query for entrytable
                        curs10.execute("SELECT SUM(quantitéharlsbntrée) AS quantitéharlsbtotal FROM entrytable")
                        resulthar = curs10.fetchone()
                        self.quantiteharlsbtotal = str(resulthar[0])

                # Connect to the second database
                with mysql.connector.connect(host="localhost",  user=user, password=password,
                                             database='datta_legumsec_out') as dattabase1:
                    with dattabase1.cursor() as curs11:
                        # Query for outtable with specific conditions
                        curs11.execute(
                            "SELECT SUM(quantitéharlsbentré) - SUM(quantitéharlsbsorté) as quantitéharlsbrest FROM outtable WHERE produit=%s AND unité=%s AND origine=%s AND fournisseur=%s",
                            (self.Produits.currentText(), self.Unites.currentText(),
                             self.OrigineDuProduits.currentText(), self.expediteur.currentText()))
                        resultharlsb_a = curs11.fetchone()
                        self.restharicotlsb = str(resultharlsb_a[0]) if resultharlsb_a is not None else "No data found"
                        print(self.restharicotlsb)

                        # Connect to the second database
                        with mysql.connector.connect(host="localhost",  user=user, password=password,
                                                     database='datta_legumsec_out') as dattabase1:
                            with dattabase1.cursor() as curs11:
                                # Query for outtable with specific conditions
                                curs11.execute(
                                    "SELECT SUM(quantitéharlsbentré) - SUM(quantitéharlsbsorté) as quantitéharlsbrest FROM outtable WHERE produit=%s AND unité=%s AND origine=%s AND fournisseur=%s",
                                    (self.Produits.currentText(), self.Unites.currentText(),
                                     self.OrigineDuProduits.currentText(), self.expediteur.currentText()))
                                resultharlsb_a = curs11.fetchone()
                                self.restharicotlsb = str(
                                    resultharlsb_a[0]) if resultharlsb_a is not None else "No data found"
                                print(self.restharicotlsb)
                            with dattabase1.cursor() as curs11:
                                # Query for outtable with specific conditions
                                curs11.execute("SELECT SUM(quantitéharlsbentré) - SUM(quantitéharlsbsorté) as quantitéharlsbrest FROM outtable ")
                                resultharlsb_at = curs11.fetchone()
                                self.restharicotlsbtt = str(resultharlsb_at[0])
                    # Additional query (docx)
                    with dattabase1.cursor() as curs11tt:
                        curs11tt.execute(
                            "SELECT SUM(quantitéharlsbentré) - SUM(quantitéharlsbsorté) as quantitéharlsbrest FROM outtable WHERE produit=%s AND unité=%s AND le BETWEEN %s AND %s",
                            (self.selected_produit, self.selected_magasin, self.selected_dates, self.selected_datesAU))
                        resultharlsb_att = curs11tt.fetchone()
                        self.restharicotlsbtt = str(
                            resultharlsb_att[0]) if resultharlsb_att is not None else "No data found"
                        print("result=" + self.restharicotlsbtt)

            except mysql.connector.Error as e:
                print(f"Error: {e}")
                # Handle the exception appropriately

        def poischicheDouze(self):
            try:
                produitstock = self.Produits.currentText()
                expediteurstock = self.expediteur.currentText()
                unitestock = self.Unites.currentText()
                originestock = self.OrigineDuProduits.currentText()

                # Connect to the first database
                with mysql.connector.connect(host="localhost",  user=user, password=password,
                                             database='datta_legumesec_entry') as dattabase:
                    with dattabase.cursor() as curs:
                        curs.execute("SELECT SUM(quantitépchentrée) AS quantitépchtotale FROM entrytable ")
                        resultpch = curs.fetchone()
                        self.poischichdouzeS = str(resultpch[0])
                        dattabase.commit()

                # Connect to the second database
                with mysql.connector.connect(host="localhost",  user=user, password=password,
                                             database='datta_legumsec_out') as dattabase1:
                    with dattabase1.cursor() as curs:
                        curs.execute(
                            "SELECT sum(quantitepchentré) - sum(quantitépchsorté) as quantitépchrest FROM outtable WHERE produit=%s AND unité=%s AND origine=%s AND fournisseur=%s", (produitstock,unitestock,originestock,expediteurstock))

                        resultpch_a = curs.fetchone()
                        if resultpch_a is not None:
                            self.restpoichichedouze = str(resultpch_a[0])
                        else:
                            # Handle the case where no results are returned
                            self.restpoichichedouze = "No data found"

                    # Additional query (docx)
                    with dattabase1.cursor() as curstt:
                            query = """
                                SELECT sum(quantitepchentré) - sum(quantitépchsorté) as quantitépchrest 
                                FROM outtable 
                                WHERE produit=%s AND unité=%s AND le BETWEEN %s AND %s
                            """
                            curstt.execute(query, (
                            self.selected_produit, self.selected_magasin, self.selected_dates, self.selected_datesAU))

                            resultpch_att = curstt.fetchone()

                            if resultpch_att is not None:
                                self.restpoichichedouzetts = str(resultpch_att[0])
                            else:
                                self.restpoichichedouzetts = "No result found"





                    with dattabase1.cursor() as cursstt:
                        cursstt.execute(
                            "SELECT sum(quantitepchentré) - sum(quantitépchsorté) as quantitépchrest FROM outtable  "

                        )

                        resultpch_astt = cursstt.fetchone()

                        self.restpoichichedouzett = str(resultpch_astt[0])
            except mysql.connector.Error as e:
                print(f"Error: {e}")
                # Handle the exception appropriately
                dattabase.commit()


        def poischicheNeuf(self):
            try:
                ########pois chiches operatio
                dattabase = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumesec_entry')
                dattabase1 = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumsec_out')
                curs = dattabase.cursor()
                curs.execute("SELECT SUM(quantitépchneufentrée) AS quantitépchneufetotal FROM entrytable ")
                resultpch = curs.fetchone()
                dattabase.commit()

                curs = dattabase1.cursor()
                curs.execute("SELECT sum(quantitépchneufeentré) - sum(quantitépchneufesorté) as quantitépchneuferest FROM outtable WHERE produit=%s AND unité=%s OR origine=%s OR fournisseur=%s", (self.Produits.currentText(), self.Unites.currentText(),self.OrigineDuProduits.currentText(),self.expediteur.currentText()))
                resultpch_a = curs.fetchone()
                self.restpoichicheneuf=str(resultpch_a[0])
                curs = dattabase1.cursor()
                curs.execute("SELECT sum(quantitépchneufeentré) - sum(quantitépchneufesorté) as quantitépchneuferest FROM outtable ")
                resultpch_at = curs.fetchone()
                self.restpoichicheneuftt = str(resultpch_at[0])
                ########################for docx out
                curstt = dattabase1.cursor()
                curstt.execute("SELECT sum(quantitépchneufeentré) - sum(quantitépchneufesorté) as quantitépchneuferest FROM outtable WHERE produit=%s AND unité=%s AND le BETWEEN %s AND %s", (self.selected_produit, self.selected_magasin,self.selected_dates,self.selected_datesAU))
                resultpch_att = curstt.fetchone()
                self.restpoichicheneufttx = str(resultpch_att[0])
                dattabase.commit()
            except Exception as e:
                print(e)
            pass

        def poischicheOuit(self):
            try:
                ########pois chiches operatio
                dattabase = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumesec_entry')
                dattabase1 = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumsec_out')
                curs = dattabase.cursor()
                curs.execute("SELECT SUM(quantitépchhuitntrée) AS quantitépchhuittotal FROM entrytable ")
                resultpch = curs.fetchone()

                dattabase.commit()

                curs = dattabase1.cursor()
                curs.execute("SELECT sum(quantitépchhuitentré) - sum(quantitépchhuitsorté) as quantitépchhuitrest FROM outtable WHERE produit=%s AND unité=%s OR origine=%s OR fournisseur=%s", (self.Produits.currentText(), self.Unites.currentText(),self.OrigineDuProduits.currentText(),self.expediteur.currentText()))
                resultpch_a = curs.fetchone()
                self.restpoichicheouit=str(resultpch_a[0])
                curs = dattabase1.cursor()
                curs.execute(
                    "SELECT sum(quantitépchhuitentré) - sum(quantitépchhuitsorté) as quantitépchhuitrest FROM outtable")
                resultpch_at = curs.fetchone()
                self.restpoichicheouittt = str(resultpch_at[0])
                ##############for docx out
                curstt = dattabase1.cursor()
                curstt.execute("SELECT sum(quantitépchhuitentré) - sum(quantitépchhuitsorté) as quantitépchhuitrest FROM outtable WHERE produit=%s AND unité=%s AND le BETWEEN %s AND %s", (self.selected_produit, self.selected_magasin,self.selected_dates,self.selected_datesAU))
                resultpch_att = curstt.fetchone()
                self.restpoichicheouitttx = str(resultpch_att[0])
                dattabase.commit()
                dattabase.close()
            except Exception as e:
                print(e)
            pass

        def lentilleCinqSept(self):
            try:
                ###########lentille
                dattabase = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumesec_entry')
                dattabase1 = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumsec_out')
                curs8 = dattabase.cursor()
                curs8.execute("SELECT SUM(quantitélenentrée) AS quantitélentotale FROM entrytable ")
                resultlen = curs8.fetchone()

                dattabase.commit()

                curs9 = dattabase1.cursor()
                curs9.execute("SELECT sum(quantitélenentré) - sum(quantitélensorté) as quantitélenrest FROM outtable WHERE produit=%s AND unité=%s OR origine=%s OR fournisseur=%s", (self.Produits.currentText(), self.Unites.currentText(),self.OrigineDuProduits.currentText(),self.expediteur.currentText()))
                resultlensinq = curs9.fetchone()
                self.restlentsinq=str(resultlensinq[0])
                curs9 = dattabase1.cursor()
                curs9.execute("SELECT sum(quantitélenentré) - sum(quantitélensorté) as quantitélenrest FROM outtable ")
                resultlensinqt = curs9.fetchone()
                self.restlentsinqtt = str(resultlensinqt[0])
                ################for docx out
                curs9tt = dattabase1.cursor()
                curs9tt.execute("SELECT sum(quantitélenentré) - sum(quantitélensorté) as quantitélenrest FROM outtable WHERE produit=%s AND unité=%s AND le BETWEEN %s AND %s", (self.selected_produit, self.selected_magasin,self.selected_dates,self.selected_datesAU))
                resultlensinqtt = curs9tt.fetchone()
                self.restlentsinqttx = str(resultlensinqtt[0])
                dattabase.commit()

            except Exception as e:
                print(e)
            pass

        def totalfunction(self):
            try:
                dattabase = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumesec_entry')
                dattabase1 = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumsec_out')
                curstotal = dattabase.cursor()
                curstotal.execute(
                    "SELECT (SUM(quantitérizentrée)+SUM(quantitépchentrée)+SUM(quantitépchplentrée)+SUM(quantitélenplntrée)+SUM(quantitélenentrée)+SUM(quantitéharntrée) ) as quantitéentrétotal FROM entrytable")

                totatentre = curstotal.fetchone()

                dattabase.commit()
                curstotalrest = dattabase1.cursor()
                curstotalrest.execute(
                    "SELECT ( SUM(quantitérizentré) + SUM(quantitepchentré)+ SUM(quantitépchplentré) + SUM(quantitélenplentré)+ SUM(quantitélenentré) + SUM(quantitéharentré))-(SUM(quantitérizesorté)+SUM(quantitépchsorté)+SUM(quantitépchplsorté)+SUM(quantitélenplsorté)+SUM(quantitélensorté)+SUM(quantitéharsorté))  as quantitérestotal  FROM outtable")
                totalrest = curstotalrest.fetchone()
                self.rest.setValue(totalrest[0])
                dattabase1.commit()
            except Exception as e:
                print(e)

        def impot_all(self):
            try:
                dattabase = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumesec_entry')
                curs = dattabase.cursor()
                curs.execute("SELECT id,nlot,nlotsch,le,produit,origine,quantité,unité,sacherie,fournisseur,imatricule ,date_recolte,date_CND ,DLUM  FROM entrytable")
                result = curs.fetchall()
                self.textEdit.setRowCount(0)
                for row, row_datta in enumerate(result):
                    self.textEdit.insertRow(row)
                    for colum, datta in enumerate(row_datta):
                        self.textEdit.setItem(row, colum, QTableWidgetItem(str(datta)))
                dattabase.commit()
                dattabase.close()
            except Exception as e :
                print(e)


        def impot_all_sortie(self):
            try:
                dattabase1 = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumsec_out')
                curs1 = dattabase1.cursor()
                curs1.execute("SELECT id, nlot,nlotsch, le, unité, produit, origine, quantité, achteur, sacherie, imatricule, date_recoltes, date_CND, DLUM FROM outtable WHERE  date_recoltes IS NOT NULL AND date_CND IS NOT NULL AND DLUM IS NOT NULL")
                result1 = curs1.fetchall()
                self.textEdit_1.setRowCount(0)
                for row1, row_datta1 in enumerate(result1):
                    print(row1)
                    self.textEdit_1.insertRow(row1)
                    for colum1, datta1 in enumerate(row_datta1):
                        self.textEdit_1.setItem(row1, colum1, QTableWidgetItem(str(datta1)))
                dattabase1.commit()
                dattabase1.close()
            except Exception as e:
                print(e)

        def delete_entry(self):
            try:
                msgbox = QMessageBox()
                msgbox.setWindowTitle("Alerte")
                msgbox.setText("Voulez-vous supprimer !")
                yesbutton = QtWidgets.QPushButton("OUI")
                nobuttons = QtWidgets.QPushButton("NON")
                # msgbox.setStandardButtons(QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
                msgbox.addButton(yesbutton, QMessageBox.ButtonRole.ActionRole)
                msgbox.addButton(nobuttons, QMessageBox.ButtonRole.ActionRole)
                push = msgbox.exec()
                clicked_button = msgbox.clickedButton()
                if clicked_button == nobuttons:
                    print("NO")
                elif clicked_button == yesbutton:
                    print("yes clicked")
                    curentrow = self.textEdit.currentRow()
                    id_ = self.textEdit.item(curentrow, 0).text()
                    database = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumesec_entry')
                    curs = database.cursor()
                    curs.execute("DELETE FROM entrytable WHERE id=%s", (id_,))
                    database.commit()
                    curs.close()
                    database.close()
                    database = mysql.connector.connect(host="localhost", user=user,password=password,database="datta_legumsec_out")
                    curs = database.cursor()
                    curs.execute("DELETE FROM outtable WHERE id=%s", (id_,))
                    database.commit()
                    database.close()
                    self.impot_all()
            except mysql.connector.Error as e:
                print(e)

        def delete_out(self):
            try:
                msgbox = QMessageBox()
                msgbox.setStyleSheet(""" QWidget
                    {
                        color: #eff0f1;
                        background-color: #31363b;
                        border-width: 1px;
                        border-color: #1e1e1e;
                        border-style: solid;
                        border-radius: 6;
                        padding: 3px;
                        font-size: 12px;
                        padding-left: 5px;
                        padding-right: 5px
                    }
                    QWidget:item:hover
                    {
                        background-color: #3daee9;
                        color: #eff0f1;
                    }
                    QWidget:item:selected
                    {
                        background-color: #3daee9;
                    }
                    QWidget:disabled
                    {
                        color: #454545;
                        background-color: #31363b;
                    }
                    QPushButton
                    {
                        color: #b1b1b1;
                        background-color: QLinearGradient( x1: 0, y1: 0, x2: 0, y2: 1, stop: 0 #565656, stop: 0.1 #525252, stop: 0.5 #4e4e4e, stop: 0.9 #4a4a4a, stop: 1 #464646);
                        border-width: 1px;
                        border-color: #1e1e1e;
                        border-style: solid;
                        border-radius: 6;
                        padding: 3px;
                        font-size: 12px;
                        padding-left: 5px;
                        padding-right: 5px;
                        min-width: 40px;
                    }
                    QPushButton:disabled
                    {
                        background-color: #31363b;
                        border-width: 1px;
                        border-color: #454545;
                        border-style: solid;
                        padding-top: 5px;
                        padding-bottom: 5px;
                        padding-left: 10px;
                        padding-right: 10px;
                        border-radius: 2px;
                        color: #454545;
                    }

                    QPushButton:pressed
                    {
                        background-color: #3daee9;
                        padding-top: -15px;
                        padding-bottom: -17px;
                    }
                    QPushButton:hover
                    {
                        border: 1px solid #ff8c00;
                        color: #eff0f1;
                    }
                     QLabel
                    {
                        font-size: 18px;
                        border: 1px solid orange;
                    }

                """)
                msgbox.setWindowTitle("Alerte")
                msgbox.setText("Voulez-vous supprimer !")
                yesbutton = QtWidgets.QPushButton("OUI")
                nobuttons = QtWidgets.QPushButton("NON")
                # msgbox.setStandardButtons(QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
                msgbox.addButton(yesbutton, QMessageBox.ButtonRole.ActionRole)
                msgbox.addButton(nobuttons, QMessageBox.ButtonRole.ActionRole)
                push = msgbox.exec()
                if msgbox.clickedButton() == nobuttons:
                    print("no")
                elif msgbox.clickedButton() == yesbutton:
                    curentrow = self.textEdit_1.currentRow()
                    id_ = self.textEdit_1.item(curentrow, 0).text()
                    database = mysql.connector.connect(
        host="localhost",
         user=user,
        password=password,
        database="datta_legumsec_out"
    )
                    curs = database.cursor()
                    curs.execute("DELETE FROM outtable WHERE id=%s", (id_,))
                    database.commit()
                    database.close()
                    self.impot_all_sortie()
            except:
                print(' ')

        def print_docx(self):
            self.poischicheDouze()
            self.poischicheNeuf()
            self.poischicheOuit()
            self.poichichSIxSept()
            self.operation_dattabase()
            self.haricot()
            self.haricotLsb()
            self.lentilleCinqSept()
            self.lenVert()
            self.lenpl()
            self.poichichepl()
            dialog = QMessageBox()
            dialog.setStyleSheet(""" QWidget
                    {
                        color: #000000;
                        background-color: #ffffff;
                        border-width: 1px;
                        border-color: #1e1e1e;
                        border-style: solid;
                        border-radius: 6;
                        padding: 3px;
                        font-size: 18px;
                        padding-left: 5px;
                        padding-right: 5px
                    }
                    QWidget:item:hover
                    {
                        background-color: #3daee9;
                        color: #eff0f1;
                    }
                    QWidget:item:selected
                    {
                        background-color: #3daee9;
                    }
                    QWidget:disabled
                    {
                        color: #454545;
                        background-color: #31363b;
                    }
                    QPushButton
                    {
                        color: #b1b1b1;
                        background-color: QLinearGradient( x1: 0, y1: 0, x2: 0, y2: 1, stop: 0 #565656, stop: 0.1 #525252, stop: 0.5 #4e4e4e, stop: 0.9 #4a4a4a, stop: 1 #464646);
                        border-width: 1px;
                        border-color: #1e1e1e;
                        border-style: solid;
                        border-radius: 6;
                        padding: 3px;
                        font-size: 12px;
                        padding-left: 5px;
                        padding-right: 5px;
                        min-width: 40px;
                    }
                    QPushButton:disabled
                    {
                        background-color: #31363b;
                        border-width: 1px;
                        border-color: #454545;
                        border-style: solid;
                        padding-top: 5px;
                        padding-bottom: 5px;
                        padding-left: 10px;
                        padding-right: 10px;
                        border-radius: 2px;
                        color: #454545;
                    }

                    QPushButton:pressed
                    {
                        background-color: #3daee9;
                        padding-top: -15px;
                        padding-bottom: -17px;
                    }
                    QPushButton:hover
                    {
                        border: 1px solid #ff8c00;
                        color: #eff0f1;
                    }
                     QLabel
                    {
                        font-size: 18px;
                        border: 0px solid orange;
                    }

                """)
            dialog.setWindowTitle("Select a Date")
            dialog.setText("Sélectionnez la date et produit souhaitée\n")
            self.date_edit = QtWidgets.QLineEdit(dialog)
            self.date_edit.setInputMask("99-99-9999")
            self.date_edit.resize(180, 30)
            self.date_edit.move(40, 82)
            self.produitphytofiltre = QtWidgets.QComboBox(dialog)
            self.produitphytofiltre.setGeometry(40, 42, 180, 30)
            self.produitphytofiltre.addItem("")
            self.produitphytofiltre.addItem("POIS CHICHES")
            self.produitphytofiltre.addItem("LENTILLE")
            self.produitphytofiltre.addItem("RIZ")
            self.produitphytofiltre.addItem("HARICOT")
            self.produitphytofiltre.addItem("LENTILLE PL")
            self.produitphytofiltre.addItem("POIS CHICHES PL")
            ok_button = QtWidgets.QPushButton("OK", dialog)
            cancel_button = QtWidgets.QPushButton("Cancel", dialog)
            dialog.addButton(ok_button, QMessageBox.ButtonRole.ActionRole)
            dialog.addButton(cancel_button, QMessageBox.ButtonRole.ActionRole)
            dialog.exec()
            try:
                self.produitphyto = self.produitphytofiltre.currentText()
                self.selected_date = self.date_edit.text()
                if dialog.clickedButton() == ok_button and  self.produitphytofiltre.currentText() not in ['POIS CHICHES','LENTILLE','RIZ','HARICOT','LENTILLE PL','POIS CHICHES PL'] and self.selected_date=="00-00-0000":
                    conn = mysql.connector.connect(host="localhost", user=user,password=password,database="datta_legumesec_entry")
                    cursor = conn.cursor()
                    cursor.execute("SELECT le,produit,origine,quantité,unité,sacherie,fournisseur,imatricule,date_recolte,date_CND,DLUM FROM entrytable")
                    rows = cursor.fetchall()
                    ########################################################################
                    # Create a new document and add a table
                    self.doc = docx.Document()
                    section = self.doc.sections[0]
                    section.orientation = WD_ORIENTATION.LANDSCAPE
                    section.page_width = docx.shared.Cm(29.7)
                    section.page_height = docx.shared.Cm(21.0)
                    section.top_margin = docx.shared.Cm(1.5)
                    section.bottom_margin = docx.shared.Cm(1.5)
                    section.right_margin = docx.shared.Cm(1.5)
                    section.left_margin = docx.shared.Cm(1.5)
                    heading1 = self.doc.add_heading('\t\t\tSITUATION DES STOCKS DES LEGUMES SECS\n\t\t\t\t\t\t\t\t\t\t\tENTRE', level=1)
                    heading1.style.font.name = 'Times New Roman'
                    heading1.style.font.size = Pt(30)
                    heading1.style.font.bold = True
                    heading1.style.font.color.rgb = RGBColor(0, 0, 0)
                    # add second heading
                    paragraph = self.doc.add_paragraph()
                    paragraph.style.font.name = 'Times New Roman'
                    paragraph.style.font.size = Pt(14)
                    left_run = paragraph.add_run('CCLS RELIZANE\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t')
                    left_run.bold = True
                    right_run = paragraph.add_run('Global')
                    right_run.bold = True

                    table = self.doc.add_table(rows=1, cols=11)
                    table.style = "Table Grid"  # set the table style
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'DATE'
                    hdr_cells[1].text = 'PRODUIT'
                    hdr_cells[2].text = 'ORIGINE'
                    hdr_cells[3].text = 'QUANTITE'
                    hdr_cells[4].text = 'MAGASIN DE STOCK'
                    hdr_cells[5].text = 'SASHERIE  CND'
                    hdr_cells[6].text = 'FOURNISSEUR'
                    hdr_cells[7].text = 'IMATRICUL'
                    hdr_cells[8].text = 'Date de récolte'
                    hdr_cells[9].text = 'Date de CND'
                    hdr_cells[10].text = 'Date de DLUM'
                    # Set the width of the header cells
                    hdr_cells[0].width = Inches(1)  # set the width of the first column to 1.5 inches
                    hdr_cells[1].width = Inches(1)
                    hdr_cells[2].width = Inches(1)
                    hdr_cells[3].width = Inches(1)
                    hdr_cells[4].width = Inches(1)
                    hdr_cells[5].width = Inches(1)
                    hdr_cells[6].width = Inches(1)
                    hdr_cells[7].width = Inches(1)
                    hdr_cells[8].width = Inches(1)
                    hdr_cells[9].width = Inches(1)
                    hdr_cells[10].width = Inches(1)
                    for cell in table.rows[0].cells:
                        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                        cell.paragraphs[0].runs[0].font.size = Pt(11.5)
                        cell.paragraphs[0].runs[0].font.bold = True
                    # Add the data to the table
                    table_rows = len(rows)
                    table_cols = len(rows[0])
                    for row in range(table_rows):
                        table.add_row()
                        for col in range(table_cols):
                            cell = table.cell(row + 1, col)
                            cell.text = str(rows[row][col])
                            # Set font properties
                            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                            cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    for row in table.rows:
                        row.height = Inches(0.6)
                    # Create the document
                    # Create the table
                    table_b = self.doc.add_table(rows=12, cols=2)
                    table_b.style = "Table Grid"
                    # Add table headers
                    hdr_colum = table_b.rows[0].cells
                    hdr_colum[0].text = "PRODUIT"
                    hdr_colum[1].text = "STOCK"
                    hdr_colum = table_b.rows[1].cells
                    hdr_colum[0].text = "POIS CHICHES PL"
                    hdr_colum[1].text =""
                    hdr_colum=table_b.rows[2].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 12 mm"
                    hdr_colum[1].text =self.restpoichichedouze +  "  QX"
                    hdr_colum = table_b.rows[3].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 09 mm"
                    hdr_colum[1].text = self.restpoichicheneufs+" QX"
                    hdr_colum = table_b.rows[4].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 08 mm"
                    hdr_colum[1].text =self.restpoichicheouits +" QX"
                    hdr_colum = table_b.rows[5].cells
                    hdr_colum[0].text = "POIS CHICHES IMP 06 mm"
                    hdr_colum[1].text = self.restpoichiches+" QX"
                    hdr_colum = table_b.rows[6].cells
                    hdr_colum[0].text = "HARICOT BLANC IMP"
                    hdr_colum[1].text = self.restharicotimps+" QX"
                    hdr_colum = table_b.rows[7].cells
                    hdr_colum[0].text = "HARICOT LSB"
                    hdr_colum[1].text =self.restharicotlsbs+ " QX"
                    hdr_colum = table_b.rows[8].cells
                    hdr_colum[0].text = "LENTILLE PL"
                    hdr_colum[1].text =self.restlenpls+ " QX"
                    hdr_colum = table_b.rows[9].cells
                    hdr_colum[0].text = "LENTILLE IMP 05-07 mm"
                    hdr_colum[1].text =self.restlentsinqs+ " QX"
                    hdr_colum = table_b.rows[10].cells
                    hdr_colum[0].text = "LENTILLE IMP vert"
                    hdr_colum[1].text =self.restlentverts+ " QX"
                    hdr_colum = table_b.rows[11].cells
                    hdr_colum[0].text = "RIZ IMP"
                    hdr_colum[1].text =self.resulxs+ " QX"



                    # Save and open the document for printing
                    temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                    self.doc.save(temp_file)
                    os.startfile(temp_file, "open")
                elif dialog.clickedButton() == ok_button and  self.produitphytofiltre.currentText() not in ['POIS CHICHES','LENTILLE','RIZ','HARICOT','LENTILLE PL','POIS CHICHES PL'] and  self.date_edit.text()==self.selected_date:
                    print(self.selected_date)
                    conn =  mysql.connector.connect(host="localhost", user=user,password=password,database="datta_legumesec_entry")
                    cursor = conn.cursor()
                    cursor.execute("SELECT le,produit,origine,quantité,unité,sacherie,fournisseur, imatricule, date_CND, DLUM FROM entrytable WHERE le = %s",(self.selected_date,))
                    rows = cursor.fetchall()
                    ########################################################################
                    # Create a new document and add a table
                    self.doc = docx.Document()
                    section = self.doc.sections[0]
                    section.orientation = WD_ORIENTATION.LANDSCAPE
                    section.page_width = docx.shared.Cm(29.7)
                    section.page_height = docx.shared.Cm(21.0)
                    section.top_margin = docx.shared.Cm(1.5)
                    section.bottom_margin = docx.shared.Cm(1.5)
                    section.right_margin = docx.shared.Cm(1.5)
                    section.left_margin = docx.shared.Cm(1.5)
                    heading1 = self.doc.add_heading('\t\t\tSITUATION DES STOCKS DES LEGUMES SECS\n\t\t\t\t\t\t\t\t\t\t\tENTRE', level=1)
                    heading1.style.font.name = 'Times New Roman'
                    heading1.style.font.size = Pt(30)
                    heading1.style.font.bold = True
                    heading1.style.font.color.rgb = RGBColor(0, 0, 0)
                    # add second heading
                    paragraph = self.doc.add_paragraph()
                    paragraph.style.font.name = 'Times New Roman'
                    paragraph.style.font.size = Pt(14)
                    left_run = paragraph.add_run('CCLS RELIZANE\t\t\t\t\t\t\t\t\t\t\t\t\t\t')
                    left_run.bold = True
                    right_run = paragraph.add_run('Mois de:' + self.selected_date)
                    right_run.bold = True
                    table = self.doc.add_table(rows=1, cols=11)
                    table.style = "Table Grid"  # set the table style
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'DATE'
                    hdr_cells[1].text = 'PRODUIT'
                    hdr_cells[2].text = 'ORIGINE'
                    hdr_cells[3].text = 'QUANTITE'
                    hdr_cells[4].text = 'MAGASIN DE STOCK'
                    hdr_cells[5].text = 'SASHERIE  CND'
                    hdr_cells[6].text = 'FOURNISSEUR'
                    hdr_cells[7].text = 'IMATRICUL'
                    hdr_cells[8].text = 'Date de récolte'
                    hdr_cells[9].text = 'Date de CND'
                    hdr_cells[10].text = 'Date de DLUM'
                    # Set the width of the header cells
                    hdr_cells[0].width = Inches(1)  # set the width of the first column to 1.5 inches
                    hdr_cells[1].width = Inches(1)
                    hdr_cells[2].width = Inches(1)
                    hdr_cells[3].width = Inches(1)
                    hdr_cells[4].width = Inches(1)
                    hdr_cells[5].width = Inches(1)
                    hdr_cells[6].width = Inches(1)
                    hdr_cells[7].width = Inches(1)
                    hdr_cells[8].width = Inches(1)
                    hdr_cells[9].width = Inches(1)
                    hdr_cells[10].width = Inches(1)

                    for cell in table.rows[0].cells:
                        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                        cell.paragraphs[0].runs[0].font.size = Pt(11.5)
                        cell.paragraphs[0].runs[0].font.bold = True
                    # Add the data to the table
                    table_rows = len(rows)
                    table_cols = len(rows[0])
                    for row in range(table_rows):
                        table.add_row()
                        for col in range(table_cols):
                            cell = table.cell(row + 1, col)
                            cell.text = str(rows[row][col])
                            # Set font properties
                            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                            cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    for row in table.rows:
                        row.height = Inches(0.6)
                    # Save and open the document for printing
                    temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                    self.doc.save(temp_file)
                    os.startfile(temp_file, "print")
                    conn.close()

                elif dialog.clickedButton() == ok_button and  self.produitphytofiltre.currentText() ==self.produitphyto :

                    conn =  mysql.connector.connect(host="localhost", user=user,password=password,database="datta_legumesec_entry")
                    cursor = conn.cursor()
                    cursor.execute("SELECT le,produit,origine,quantité,unité,sacherie,fournisseur, imatricule, date_CND, DLUM FROM entrytable WHERE produit= %s",( self.produitphyto,))
                    rows = cursor.fetchall()
                    ########################################################################
                    # Create a new document and add a table
                    self.doc = docx.Document()
                    section = self.doc.sections[0]
                    section.orientation = WD_ORIENTATION.LANDSCAPE
                    section.page_width = docx.shared.Cm(29.7)
                    section.page_height = docx.shared.Cm(21.0)
                    section.top_margin = docx.shared.Cm(1.5)
                    section.bottom_margin = docx.shared.Cm(1.5)
                    section.right_margin = docx.shared.Cm(1.5)
                    section.left_margin = docx.shared.Cm(1.5)
                    heading1 = self.doc.add_heading('\t\t\tSITUATION DES STOCKS DES LEGUMES SECS\n\t\t\t\t\t\t\t\t\t\t\tENTRE ', level=1)
                    heading1.style.font.name = 'Times New Roman'
                    heading1.style.font.size = Pt(30)
                    heading1.style.font.bold = True
                    heading1.style.font.color.rgb = RGBColor(0, 0, 0)
                    # add second heading
                    paragraph = self.doc.add_paragraph()
                    paragraph.style.font.name = 'Times New Roman'
                    paragraph.style.font.size = Pt(14)
                    left_run = paragraph.add_run('CCLS RELIZANE\t\t\t\t\t\t\t\t\t\t\t\t\t')
                    left_run.bold = True
                    right_run = paragraph.add_run('Produit:'+self.produitphyto)
                    right_run.bold = True
                    table = self.doc.add_table(rows=1, cols=11)
                    table.style = "Table Grid"  # set the table style
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'DATE'
                    hdr_cells[1].text = 'PRODUIT'
                    hdr_cells[2].text = 'ORIGINE'
                    hdr_cells[3].text = 'QUANTITE'
                    hdr_cells[4].text = 'MAGASIN DE STOCK'
                    hdr_cells[5].text = 'SASHERIE  CND'
                    hdr_cells[6].text = 'FOURNISSEUR'
                    hdr_cells[7].text = 'IMATRICUL'
                    hdr_cells[8].text = 'Date de récolte'
                    hdr_cells[9].text = 'Date de CND'
                    hdr_cells[10].text = 'Date de DLUM'
                    # Set the width of the header cells
                    hdr_cells[0].width = Inches(1)  # set the width of the first column to 1.5 inches
                    hdr_cells[1].width = Inches(1)
                    hdr_cells[2].width = Inches(1)
                    hdr_cells[3].width = Inches(1)
                    hdr_cells[4].width = Inches(1)
                    hdr_cells[5].width = Inches(1)
                    hdr_cells[6].width = Inches(1)
                    hdr_cells[7].width = Inches(1)
                    hdr_cells[8].width = Inches(1)
                    hdr_cells[9].width = Inches(1)
                    hdr_cells[10].width = Inches(1)
                    for cell in table.rows[0].cells:
                        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                        cell.paragraphs[0].runs[0].font.size = Pt(11.5)
                        cell.paragraphs[0].runs[0].font.bold = True
                    # Add the data to the table
                    table_rows = len(rows)
                    table_cols = len(rows[0])
                    for row in range(table_rows):
                        table.add_row()
                        for col in range(table_cols):
                            cell = table.cell(row + 1, col)
                            cell.text = str(rows[row][col])
                            # Set font properties
                            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                            cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    for row in table.rows:
                        row.height = Inches(0.6)
                    # Save and open the document for printing
                    temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                    self.doc.save(temp_file)
                    os.startfile(temp_file, "print")
                    conn.close()
            except Exception as e:
                print(e)

        def print_facture_sortie(self):
            try:
                self.msgbox = QtWidgets.QDialog()
                self.msgbox.setGeometry(300, 40, 810, 670)
                self.msgbox.setStyleSheet(""" QWidget
                                       {
                                           color: #eff0f1;
                                           background-color: #31363b;
                                           border-width: 1px;
                                           border-color: #1e1e1e;
                                           border-style: solid;
                                           border-radius: 6;
                                           padding: 0px;
                                           font-size: 18px;
                                           padding-left: 5px;
                                           padding-right: 5px
                                       }
                                       QWidget:item:hover
                                       {
                                           background-color: #3daee9;
                                           color: #eff0f1;
                                       }
                                       QWidget:item:selected
                                       {
                                           background-color: #3daee9;
                                       }
                                       QWidget:disabled
                                       {
                                           color: #454545;
                                           background-color: #31363b;
                                       }
                                       QPushButton
                                       {
                                           color: #b1b1b1;
                                           background-color: QLinearGradient( x1: 0, y1: 0, x2: 0, y2: 1, stop: 0 #565656, stop: 0.1 #525252, stop: 0.5 #4e4e4e, stop: 0.9 #4a4a4a, stop: 1 #464646);
                                           border-width: 1px;
                                           border-color: #1e1e1e;
                                           border-style: solid;
                                           border-radius: 6;
                                           padding: 3px;
                                           font-size: 16px;
                                           padding-left: 5px;
                                           padding-right: 5px;
                                           min-width: 40px;
                                       }
                                       QPushButton:disabled
                                       {
                                           background-color: #31363b;
                                           border-width: 1px;
                                           border-color: #454545;
                                           border-style: solid;
                                           padding-top: 5px;
                                           padding-bottom: 5px;
                                           padding-left: 10px;
                                           padding-right: 10px;
                                           border-radius: 2px;
                                           color: #454545;
                                       }

                                       QPushButton:pressed
                                       {
                                           background-color: #3daee9;
                                           padding-top: -15px;
                                           padding-bottom: -17px;
                                       }
                                       QPushButton:hover
                                       {
                                           border: 1px solid #ff8c00;
                                           color: #eff0f1;
                                       }
                                        QLabel
                                       {
                                           font-size: 16px;
                                           border: 0px solid orange;
                                       }
                                       
                                       QTableView
        {
            border: 1px solid #76797C;
            gridline-color: #31363b;
            background-color: #302629;
            
        }


        QTableView, QHeaderView
        {
            border-radius: 0px;
        }

        QTableView::item:pressed, QListView::item:pressed, QTreeView::item:pressed  {
            background: #3daee9;
            color: #eff0f1;
            font-size: 16px;
        }

        QTableView::item:selected:active, QTreeView::item:selected:active, QListView::item:selected:active  {
            background: #3daee9;
            color: #eff0f1;
            font-size: 16px;
        }


        QHeaderView
        {
            background-color: #31363b;
            border: 1px transparent;
            border-radius: 0px;
            margin: 0px;
            padding: 0px;
            font-size: 16px;

        }

        QHeaderView::section  {
            background-color: #31363b;
            color: #eff0f1;
            padding: 1px;
            border: 1px solid #76797C;
            border-radius: 0px;
            text-align: center;
            font-size: 16px;
        }

        QHeaderView::section::vertical::first, QHeaderView::section::vertical::only-one
        {
            border-top: 1px solid #76797C;
        }

        QHeaderView::section::vertical
        {
            border-top: transparent;
        }

        QHeaderView::section::horizontal::first, QHeaderView::section::horizontal::only-one
        {
            border-left: 1px solid #76797C;
        }

        QHeaderView::section::horizontal
        {
            border-left: transparent;
        }


        QHeaderView::section:checked
         {
            color: white;
            background-color: #334e5e;
            font-size: 16px;
         }

         /* style the sort indicator */
        QHeaderView::down-arrow {

        }

        QHeaderView::up-arrow {

        }


        QTableCornerButton::section {
            background-color: #31363b;
            border: 1px transparent #76797C;
            border-radius: 0px;
        }

                                   """)
                self.msgbox.setWindowTitle("Fiche de traitement")

                self.ok_button = QtWidgets.QPushButton("Imprimer", self.msgbox,
                                                       clicked=lambda: self.docx_facture_sortie())
                self.ok_button.setGeometry(100, 580, 200, 40)
                self.cancel_button = QtWidgets.QPushButton("Annuler", self.msgbox, clicked=lambda: self.msgbox.close())
                self.cancel_button.setGeometry(500, 580, 200, 40)

                self.save_button = QtWidgets.QPushButton("Enregistrer", self.msgbox,
                                                         clicked=lambda: self.save_facture_sortie())
                self.save_button.setGeometry(300, 580, 200, 40)

                self.addd_button=QtWidgets.QPushButton("Insérer",self.msgbox,clicked=lambda:self.add_table())
                self.addd_button.setGeometry(400,295,390,40)

                # msgbox.setText("\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\n\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t")

                # msgbox.addButton(ok_button,QMessageBox.ButtonRole.ActionRole)
                # msgbox.addButton(cancel_button,QMessageBox.ButtonRole.ActionRole)
                fichetraitementxt = QtWidgets.QLabel("LA FACTURE", self.msgbox)
                fichetraitementxt.setGeometry(300, 5, 200, 60)
                dateDeFacturetx = QtWidgets.QLabel("Date de facture:", self.msgbox)
                dateDeFacturetx.setGeometry(5, 60, 160, 30)
                self.dateDeFacture = QtWidgets.QLineEdit("Date", self.msgbox)
                self.dateDeFacture.setInputMask("99/99/9999")
                self.dateDeFacture.setGeometry(190, 60, 170, 30)
                NmDeFacturetxt = QtWidgets.QLabel("Numero de facture :", self.msgbox)
                NmDeFacturetxt.setGeometry(400, 60, 300, 30)
                self.NmDeFacture = QtWidgets.QLineEdit(self.msgbox,readOnly=True)
                self.NmDeFacture.setGeometry(620, 60, 170, 30)
                self.NmDeFacture.setText(str(self.textEdit_1.rowCount()))


                self.Nomtxt = QtWidgets.QLabel("Client:", self.msgbox)
                self.Nomtxt.setGeometry(5, 115, 300, 30)

                self.Nom = QtWidgets.QLineEdit(self.msgbox)
                self.Nom.setGeometry(190, 115, 170, 30)



                self. IMATRICULEtxt = QtWidgets.QLabel("Imatricule:", self.msgbox)
                self.IMATRICULEtxt.setGeometry(5, 175, 242, 30)

                self.IMATRICULE = QtWidgets.QLineEdit(self.msgbox)
                self.IMATRICULE.setGeometry(190, 175, 170, 30)
                self.IMATRICULE.setInputMask("99999-999-99")

                self.Unitetxt = QtWidgets.QLabel("Unité de stock:", self.msgbox)
                self.Unitetxt.setGeometry(5, 305, 300, 30)

                self.Unite = QtWidgets.QComboBox(self.msgbox,editable=True)
                self.Unite.addItem(" ")
                self.Unite.addItem("Dock Central")
                self.Unite.addItem("Mesra")
                self.Unite.addItem("Kef lazreg")
                self.Unite.addItem("Zemmoura")
                self.Unite.addItem("N.S.Mendes")
                self.Unite.addItem(" Mendes")
                self.Unite.addItem("OFLA ")
                self.Unite.addItem("Magasin Belhacel ")
                self.Unite.setGeometry(190, 305, 170, 30)


                self.Produittxt = QtWidgets.QLabel(" Produit :", self.msgbox)
                self.Produittxt.setGeometry(400, 115, 300, 30)

                self.Produit = QtWidgets.QComboBox(self.msgbox, editable=True)
                self.Produit.setGeometry(620, 115, 170, 30)
                self.Produit.addItem("")
                self.Produit.addItem("POIS CHICHES")
                self.Produit.addItem("LENTILLE")
                self.Produit.addItem("RIZ")
                self.Produit.addItem("HARICOT")
                self.Produit.addItem("LENTILLE PL")
                self.Produit.addItem("POIS CHICHES PL")



                self.OrigineDuProduittxt = QtWidgets.QLabel("Origine de produit:", self.msgbox)
                self.OrigineDuProduittxt.setGeometry(400, 305, 300, 30)

                self.OrigineDuProduit = QtWidgets.QLineEdit(self.msgbox)
                self.OrigineDuProduit.setGeometry(620, 305, 170, 30)

                self.DLUMtxt = QtWidgets.QLabel("DLUM:", self.msgbox)
                self.DLUMtxt.setGeometry(5, 295, 160, 30)

                self.DLUM = QtWidgets.QLineEdit(self.msgbox)
                self.DLUM.setGeometry(190, 295, 170, 30)
                self.DLUM.setInputMask("99/99/9999")

                self.quantitetxt = QtWidgets.QLabel("Quantité:", self.msgbox)
                self.quantitetxt.setGeometry(400, 175, 160, 30)
                self.quantite = QtWidgets.QLineEdit(self.msgbox)
                self.quantite.setGeometry(620, 175, 170, 30)

                self.txtfacturetxt=QtWidgets.QLabel(" ",self.msgbox)
                self.txtfacturetxt.setGeometry(400,400,100,20)
                self.txtfacture = QtWidgets.QTableWidget(self.msgbox)
                self.txtfacture.setGeometry(10, 350, 780, 200)
                self.txtfacture.setStyleSheet(" background-color: #302629 ")
                self.txtfacture.setRowCount(0)
                self.txtfacture.setColumnCount(5)
                self.txtfacture.setColumnWidth(0, 140)
                self.txtfacture.setColumnWidth(1, 140)
                self.txtfacture.setColumnWidth(2, 145)
                self.txtfacture.setColumnWidth(3, 175)
                self.txtfacture.setColumnWidth(4, 145)
                self.txtfacture.setHorizontalHeaderLabels(
                    ("UNITÉ", "PRODUIT", "QUANTITÉ", "ORIGINE DE PRODUIT", "DLUM"))

                for row in range(self.txtfacture.rowCount()):
                    for col in range(self.txtfacture.columnCount()):
                        item = self.txtfacture.item(row, col)
                        if item is not None:
                            global cell_text
                            cell_text= item.text()
                            print(f"Row {row}, Column {col}: {self.cell_text}")
                self.msgbox.show()
                self.msgbox.exec()
            except Exception as e:
                print(e)


        def docx_facture_sortie(self):
            try:
                # Retrieve data from various widgets
                dateFactur = self.dateDeFacture.text()
                nom = self.Nom.text()
                imatricule = self.IMATRICULE.text()
                unité = self.Unite.currentText()
                produit = self.Produit.currentText()
                origineProduit = self.OrigineDuProduit.text()
                DLUM = self.DLUM.text()
                quantite = self.quantite.text()

                # Loop through the table to get cell data
                table_data = []
                for row in range(self.txtfacture.rowCount()):
                    row_data = []
                    for col in range(self.txtfacture.columnCount()):
                        item = self.txtfacture.item(row, col)
                        if item is not None:
                            cell_text = item.text()
                            row_data.append(cell_text)
                    if row_data:
                        table_data.append(row_data)
                # Accessing the database (you may need to uncomment this part)
                # database = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumsec_out')
                # cursor = database.cursor()
                # cursor.execute("SELECT id FROM outtable WHERE  quantité  IS  NOT NULL ORDER BY id DESC LIMIT 1")
                # result = cursor.fetchone()
                # if result is not None:
                #     last_id = result[0]
                #     numeroDeFacture = str(last_id)
                # Calculate row count (you may need to uncomment this part)
                # rowcount = self.textEdit_1.rowCount()
                # self.numeroDeFacture = rowcount
                # Render the template with the data
                self.docFacture = DocxTemplate("fiche_legume-sec/legum_sec_templfactur.docx")
                context = {
                    "dtf": dateFactur,
                    # "nmf": self.numeroDeFacture,
                    "nom": nom,
                    "imt": imatricule,
                    "msn": unité,
                    "prd": produit,
                    "org": origineProduit,
                    "dlm": DLUM,
                    "qtt": quantite,
                    "msn":  table_data[0][0],
                    "msna": table_data[1][0],
                    "msnb": table_data[2][0],
                    "msnc": table_data[3][0],
                    "msnd": table_data[4][0],
                    "msne": table_data[0][1],
                    "prda": table_data[5][0],
                    "prd":  table_data[1][1],
                    "prdb": table_data[2][1],
                    "prdc": table_data[3][1],
                    "prdd": table_data[4][1],
                    "prde": table_data[5][1]
                }
                self.docFacture.render(context)
                # Save the template to a temporary file
                tempfilebd = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                self.docFacture.save(tempfilebd)
                # Open the temporary file
                os.startfile(tempfilebd, "open")
            except Exception as e:
                print(e)


        def save_facture_sortie(self):
            try:
                dateFatur = self.dateedite_1.text()
                client = self.magasincombo_1.currentText()
                imatricule = self.imatricullineeditr_1.text()
                unité = self.magasincombos.currentText()
                produit = self.produitcombo_1.currentText()
                origineProduit = self.originel_1.text()
                DLUM = self.dateEditdlum_1.text()
                qauntite = self.quatitelineEdit_1.text()
                try:
                    # database = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumsec_out')
                    # cursor = database.cursor()
                    # cursor.execute("SELECT id FROM outtable WHERE  quantité  IS  NOT NULL ORDER BY id DESC LIMIT 1")
                    # result = cursor.fetchone()
                    # if result is not None:
                    # last_id = result[0]
                    # numeroDeFacture=(str(last_id))
                    rowcount = self.textEdit_1.rowCount()
                    numeroDeFacture = rowcount
                except sqlite3.Error as e:
                    print("Error accessing the database:", e)
                self.docFacture = DocxTemplate("fiche_legume-sec/legum_sec_templfactur.docx")
                self.docFacture.render({"dtf": dateFatur,
                                        "nmf": numeroDeFacture,
                                        "prn": client,
                                        "imt": imatricule,
                                        "msn": unité,
                                        "prd": produit,
                                        "org": origineProduit,
                                        "dlm": DLUM,
                                        "qtt": qauntite, })

                path, _ = QFileDialog.getSaveFileName(None, "Enregistrer la fiche", "", "Fichiers DOCX (*.docx)")
                if path:
                    self.docFacture.save(path)
                    msgbox = QtWidgets.QMessageBox()
                    msgbox.setWindowTitle('confirmation')
                    msgbox.setText('fiche de traitemen a été enregistrée')
                    msgbox.exec()

            except Exception as e:
                print(e)

        def print_docx_sortie(self):
            dialog = QMessageBox()
            dialog.setStyleSheet(""" QWidget
                    {
                        color: #eff0f1;
                        background-color: #31363b;
                        border-width: 1px;
                        border-color: #1e1e1e;
                        border-style: solid;
                        border-radius: 6;
                        padding: 3px;
                        font-size: 18px;
                        padding-left: 5px;
                        padding-right: 5px
                    }
                    QWidget:item:hover
                    {
                        background-color: #3daee9;
                        color: #eff0f1;
                    }
                    QWidget:item:selected
                    {
                        background-color: #3daee9;
                    }
                    QWidget:disabled
                    {
                        color: #454545;
                        background-color: #31363b;
                    }
                    QPushButton
                    {
                        color: #b1b1b1;
                        background-color: QLinearGradient( x1: 0, y1: 0, x2: 0, y2: 1, stop: 0 #565656, stop: 0.1 #525252, stop: 0.5 #4e4e4e, stop: 0.9 #4a4a4a, stop: 1 #464646);
                        border-width: 1px;
                        border-color: #1e1e1e;
                        border-style: solid;
                        border-radius: 6;
                        padding: 3px;
                        font-size: 12px;
                        padding-left: 5px;
                        padding-right: 5px;
                        min-width: 40px;
                    }
                    QPushButton:disabled
                    {
                        background-color: #31363b;
                        border-width: 1px;
                        border-color: #454545;
                        border-style: solid;
                        padding-top: 5px;
                        padding-bottom: 5px;
                        padding-left: 10px;
                        padding-right: 10px;
                        border-radius: 2px;
                        color: #454545;
                    }

                    QPushButton:pressed
                    {
                        background-color: #3daee9;
                        padding-top: -15px;
                        padding-bottom: -17px;
                    }
                    QPushButton:hover
                    {
                        border: 1px solid #ff8c00;
                        color: #eff0f1;
                    }
                     QLabel
                    {
                        font-size: 18px;
                        border: 0px solid orange;
                    }

                """)
            dialog.setWindowTitle("Select a Date")
            dialog.setText("Sélectionnez la date et produit souhaitée\n")
            self.date_edit = QtWidgets.QLineEdit(dialog)
            self.date_edit.setInputMask("99-99-9999")
            self.date_edit.resize(180, 30)
            self.date_edit.move(40, 82)
            self.produitphytofiltre = QtWidgets.QComboBox(dialog)
            self.produitphytofiltre.setGeometry(40, 42, 180, 30)
            self.produitphytofiltre.addItem("")
            self.produitphytofiltre.addItem("POIS CHICHES")
            self.produitphytofiltre.addItem("LENTILLE")
            self.produitphytofiltre.addItem("RIZ")
            self.produitphytofiltre.addItem("HARICOT")
            self.produitphytofiltre.addItem("LENTILLE PL")
            self.produitphytofiltre.addItem("POIS CHICHES PL")
            ok_button = QtWidgets.QPushButton("OK", dialog)
            cancel_button = QtWidgets.QPushButton("Cancel", dialog)
            dialog.addButton(ok_button, QMessageBox.ButtonRole.ActionRole)
            dialog.addButton(cancel_button, QMessageBox.ButtonRole.ActionRole)
            dialog.exec()
            try:
                self.produitphyto = self.produitphytofiltre.currentText()
                self.selected_date = self.date_edit.text()
                if dialog.clickedButton() == ok_button and  self.produitphytofiltre.currentText() not in ['POIS CHICHES','LENTILLE','RIZ','HARICOT','LENTILLE PL','POIS CHICHES PL'] and self.selected_date=="00-00-0000":
                    conn = sqlite3.connect('datta_legumsec_out.db')
                    cursor = conn.cursor()
                    cursor.execute("SELECT le,unité,produit ,origine ,quantité ,achteur ,sacherie ,imatricule,date_recoltes,date_CND ,DLUM FROM outtable  WHERE  quantité  IS  NOT NULL")
                    rows = cursor.fetchall()
                    print(rows)

                    ########################################################################
                    # Create a new document and add a table
                    self.doc = docx.Document()
                    section = self.doc.sections[0]
                    section.orientation = WD_ORIENTATION.LANDSCAPE
                    section.page_width = docx.shared.Cm(29.7)
                    section.page_height = docx.shared.Cm(21.0)
                    section.top_margin = docx.shared.Cm(1.5)
                    section.bottom_margin = docx.shared.Cm(1.5)
                    section.right_margin = docx.shared.Cm(1.5)
                    section.left_margin = docx.shared.Cm(1.5)
                    heading1 = self.doc.add_heading('\t\t\t\t\t\t\tSITUATION DES STOCKS DES LEGUMES SECS\n\t\t\t\t\t\t\t\t\t\t\tSORTIE', level=1)
                    heading1.style.font.name = 'Times New Roman'
                    heading1.style.font.size = Pt(30)
                    heading1.style.font.bold = True
                    heading1.style.font.color.rgb = RGBColor(0, 0, 0)
                    # add second heading
                    paragraph = self.doc.add_paragraph()
                    paragraph.style.font.name = 'Times New Roman'
                    paragraph.style.font.size = Pt(14)
                    left_run = paragraph.add_run('CCLS RELIZANE\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t')
                    left_run.bold = True
                    right_run = paragraph.add_run('Global')
                    right_run.bold = True
                    table = self.doc.add_table(rows=1, cols=11)
                    table.style = "Table Grid"  # set the table style
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'DATE'
                    hdr_cells[1].text = 'MAGASIN DE STOCK'
                    hdr_cells[2].text = 'PRODUIT'
                    hdr_cells[3].text = 'ORIGINE'
                    hdr_cells[4].text = 'QUANTITE'
                    hdr_cells[5].text = 'ACHTEUR'
                    hdr_cells[6].text = 'SACHERIE DE CND'
                    hdr_cells[7].text = 'IMATRICUL'
                    hdr_cells[8].text = 'Date de récolte'
                    hdr_cells[9].text = 'Date de CND'
                    hdr_cells[10].text = 'Date de DLUM'
                    # Set the width of the header cells
                    hdr_cells[0].width = Inches(2)  # set the width of the first column to 1.5 inches
                    hdr_cells[1].width = Inches(2)
                    hdr_cells[2].width = Inches(2)
                    hdr_cells[3].width = Inches(1.5)
                    hdr_cells[4].width = Inches(1.5)
                    hdr_cells[5].width = Inches(3)
                    hdr_cells[6].width = Inches(1.5)
                    hdr_cells[7].width = Inches(1.5)
                    hdr_cells[8].width = Inches(2)
                    hdr_cells[9].width = Inches(2)
                    hdr_cells[10].width = Inches(2)
                    for cell in table.rows[0].cells:
                        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                        cell.paragraphs[0].runs[0].font.size = Pt(11.5)
                        cell.paragraphs[0].runs[0].font.bold = True
                    # Add the data to the table
                    table_rows = len(rows)
                    table_cols = len(rows[0])
                    for row in range(table_rows):
                        table.add_row()
                        for col in range(table_cols):
                            cell = table.cell(row + 1, col)
                            cell.text = str(rows[row][col])
                            # Set font properties
                            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                            cell.paragraphs[0].runs[0].font.size = Pt(10)
                            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    for row in table.rows:
                        row.height = Inches(0.6)
                    database = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumesec_entry')
                    database1 = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumsec_out')
                    curs = database.cursor()
                    curs1 = database1.cursor()
                    # Execute the queries
                    curs.execute("SELECT SUM(quantitérizentrée) FROM entrytable")
                    sumresultriz = curs.fetchone()
                    curs1.execute("SELECT SUM(quantitérizentré)-SUM(quantitérizesorté) as quantitérizerest FROM outtable")
                    resultriz = curs1.fetchone()


                    dattabase = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumesec_entry')
                    dattabase1 = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumsec_out')
                    curs = dattabase.cursor()
                    curs.execute("SELECT SUM(quantitépchentrée) AS quantitépchtotale FROM entrytable ")
                    resultpch = curs.fetchone()
                    dattabase.commit()
                    curs = dattabase1.cursor()
                    curs.execute("SELECT sum(quantitepchentré) - sum(quantitépchsorté) as quantitépchrest FROM outtable")
                    resultpch_a = curs.fetchone()
                    dattabase.commit()

                    ####################pois chichz pl operation
                    dattabase = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumesec_entry')
                    dattabase1 = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumsec_out')
                    curs4 = dattabase.cursor()
                    curs4.execute("SELECT SUM(quantitépchplentrée) AS quantitépchpltotale FROM entrytable ")
                    resultpchpl = curs4.fetchone()
                    dattabase.commit()

                    curs5 = dattabase1.cursor()
                    curs5.execute("SELECT sum(quantitépchplentré) - sum(quantitépchplsorté) as quantitépchplhrest FROM outtable")
                    resultpchpl_a = curs5.fetchone()
                    dattabase.commit()

                    ###############lentille pl operation
                    dattabase = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumesec_entry')
                    dattabase1 = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumsec_out')
                    curs6 = dattabase.cursor()
                    curs6.execute("SELECT SUM(quantitélenplntrée) AS quantitélenpltotale FROM entrytable ")
                    resultlenpl = curs6.fetchone()
                    dattabase.commit()
                    curs7 = dattabase1.cursor()
                    curs7.execute("SELECT sum(quantitélenplentré) - sum(quantitélenplsorté) as quantitélenplrest FROM outtable")
                    resultlenpl_a = curs7.fetchone()
                    dattabase.commit()

                    ###########lentille
                    dattabase = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumesec_entry')
                    dattabase1 = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumsec_out')
                    curs8 = dattabase.cursor()
                    curs8.execute("SELECT SUM(quantitélenentrée) AS quantitélentotale FROM entrytable ")
                    resultlen = curs8.fetchone()

                    dattabase.commit()
                    curs9 = dattabase1.cursor()
                    curs9.execute(
                        "SELECT sum(quantitélenentré) - sum(quantitélensorté) as quantitélenrest FROM outtable")
                    resultlen_a = curs9.fetchone()
                    dattabase.commit()
                    ####################haricot operation
                    dattabase = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumesec_entry')
                    dattabase1 = mysql.connector.connect(host="localhost", user=user,password=password,database='datta_legumsec_out')
                    curs10 = dattabase.cursor()
                    curs10.execute("SELECT SUM(quantitéharntrée) AS quantitéhartotale FROM entrytable ")
                    resulthar = curs10.fetchone()
                    dattabase.commit()
                    curs11 = dattabase1.cursor()
                    curs11.execute("SELECT sum(quantitéharentré) - sum(quantitéharsorté) as quantitéharrest FROM outtable")
                    resulthar_a = curs11.fetchone()
                    dattabase.commit()
                    dattabase.close()

                # Create the document
                    # Create the table
                    table_b = self.doc.add_table(rows=3, cols=7)
                    table_b.style = "Table Grid"
                    # Add table headers
                    hdr_colum = table_b.rows[0].cells
                    hdr_colum[0].text = "PRODUIT"
                    hdr_colum[1].text = "POIS CHICHES"
                    hdr_colum[2].text = "POIS CHICHES PL"
                    hdr_colum[3].text = "RIZ"
                    hdr_colum[4].text = "HARICOT"
                    hdr_colum[5].text = "LENTILLE"
                    hdr_colum[6].text = "LENTILLE PL"
                    hdr_colum = table_b.rows[1].cells
                    hdr_colum[0].text = "TOTAL"
                    hdr_colum[1].text=str(resultpch[0])+ " QX"
                    hdr_colum[3].text= str(sumresultriz[0])+ " QX"
                    hdr_colum[2].text = str(resultpchpl[0])+ " QX"
                    hdr_colum[4].text = str(resulthar[0]) + " QX"
                    hdr_colum[5].text = str(resultlen[0])+" QX"
                    hdr_colum[6].text = str(resultlenpl[0])+ " QX"
                    hdr_colum = table_b.rows[2].cells
                    hdr_colum[0].text = "REST"
                    hdr_colum[3].text=str(resultriz[0])+ " QX"
                    hdr_colum[1].text = str(resultpch_a[0])+ " QX"
                    hdr_colum[2].text = str(resultpchpl_a[0])+ " QX"
                    hdr_colum[5].text = str(resultlen_a[0])+ " QX"
                    hdr_colum[6].text = str(resultlenpl_a[0])+ " QX"
                    hdr_colum[4].text = str(resulthar_a[0]) + " QX"
                    # Save and open the document for printing
                    temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                    self.doc.save(temp_file)
                    os.startfile(temp_file, "print")

                elif dialog.clickedButton() == ok_button and  self.produitphytofiltre.currentText() not in ['POIS CHICHES','LENTILLE','RIZ','HARICOT','LENTILLE PL','POIS CHICHES PL'] and  self.date_edit.text()==self.selected_date:
                    print(self.selected_date)
                    conn = sqlite3.connect('datta_legumsec_out.db')
                    cursor = conn.cursor()
                    cursor.execute("SELECT le ,unité,produit,origine,quantité,achteur,sacherie,imatricule ,date_CND,DLUM FROM outtable WHERE le = %s",(self.selected_date,))
                    rows = cursor.fetchall()
                    ########################################################################
                    # Create a new document and add a table
                    self.doc = docx.Document()
                    section = self.doc.sections[0]
                    section.orientation = WD_ORIENTATION.LANDSCAPE
                    section.page_width = docx.shared.Cm(29.7)
                    section.page_height = docx.shared.Cm(21.0)
                    section.top_margin = docx.shared.Cm(1.5)
                    section.bottom_margin = docx.shared.Cm(1.5)
                    section.right_margin = docx.shared.Cm(1.5)
                    section.left_margin = docx.shared.Cm(1.5)
                    heading1 = self.doc.add_heading('\t\t\t\t\t\t\tSITUATION DES STOCKS DES LEGUMES SECS\n\t\t\t\t\t\t\t\t\t\t\tSORTIE', level=1)
                    heading1.style.font.name = 'Times New Roman'
                    heading1.style.font.size = Pt(30)
                    heading1.style.font.bold = True
                    heading1.style.font.color.rgb = RGBColor(0, 0, 0)
                    # add second heading
                    paragraph = self.doc.add_paragraph()
                    paragraph.style.font.name = 'Times New Roman'
                    paragraph.style.font.size = Pt(14)
                    left_run = paragraph.add_run('CCLS RELIZANE\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t')
                    left_run.bold = True
                    right_run = paragraph.add_run('Date:'+self.selected_date)
                    right_run.bold = True
                    table = self.doc.add_table(rows=1, cols=11)
                    table.style = "Table Grid"  # set the table style
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'DATE'
                    hdr_cells[1].text = 'MAGASIN DE STOCK'
                    hdr_cells[2].text = 'PRODUIT'
                    hdr_cells[3].text = 'ORIGINE'
                    hdr_cells[4].text = 'QUANTITE'
                    hdr_cells[5].text = 'ACHTEUR'
                    hdr_cells[6].text = 'SACHERIE DE CND'
                    hdr_cells[7].text = 'IMATRICUL'
                    hdr_cells[8].text = 'Date de récolte'
                    hdr_cells[9].text = 'Date de CND'
                    hdr_cells[10].text = 'Date de DLUM'
                    # Set the width of the header cells
                    hdr_cells[0].width = Inches(2)  # set the width of the first column to 1.5 inches
                    hdr_cells[1].width = Inches(2)
                    hdr_cells[2].width = Inches(2)
                    hdr_cells[3].width = Inches(1.5)
                    hdr_cells[4].width = Inches(1.5)
                    hdr_cells[5].width = Inches(3)
                    hdr_cells[6].width = Inches(1.5)
                    hdr_cells[7].width = Inches(1.5)
                    hdr_cells[8].width = Inches(2)
                    hdr_cells[9].width = Inches(2)
                    hdr_cells[10].width = Inches(2)
                    for cell in table.rows[0].cells:
                        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                        cell.paragraphs[0].runs[0].font.size = Pt(11.5)
                        cell.paragraphs[0].runs[0].font.bold = True
                    # Add the data to the table
                    table_rows = len(rows)
                    table_cols = len(rows[0])
                    for row in range(table_rows):
                        table.add_row()
                        for col in range(table_cols):
                            cell = table.cell(row + 1, col)
                            cell.text = str(rows[row][col])
                            # Set font properties
                            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                            cell.paragraphs[0].runs[0].font.size = Pt(10)
                            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    for row in table.rows:
                        row.height = Inches(0.6)
                    # Save and open the document for printing
                    temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                    self.doc.save(temp_file)
                    os.startfile(temp_file, "print")
                    conn.close()

                elif dialog.clickedButton() == ok_button and  self.produitphytofiltre.currentText() ==self.produitphyto :
                    print("nooooooooo" + self.selected_date, self.produitphyto)
                    conn = sqlite3.connect('datta_legumsec_out.db')
                    cursor = conn.cursor()
                    cursor.execute("SELECT le ,unité,produit,origine,quantité,achteur,sacherie,imatricule ,date_CND,DLUM FROM outtable WHERE produit= %s",( self.produitphyto,))
                    rows = cursor.fetchall()
                    ########################################################################
                    # Create a new document and add a table
                    self.doc = docx.Document()
                    section = self.doc.sections[0]
                    section.orientation = WD_ORIENTATION.LANDSCAPE
                    section.page_width = docx.shared.Cm(29.7)
                    section.page_height = docx.shared.Cm(21.0)
                    section.top_margin = docx.shared.Cm(1.5)
                    section.bottom_margin = docx.shared.Cm(1.5)
                    section.right_margin=docx.shared.Cm(1.5)
                    section.left_margin = docx.shared.Cm(1.5)
                    heading1 = self.doc.add_heading('\t\t\t\t\t\t\tSITUATION DES STOCKS DES LEGUMES SECS\n\t\t\t\t\t\t\t\t\t\t\tSORTIE', level=1)
                    heading1.style.font.name = 'Times New Roman'
                    heading1.style.font.size = Pt(30)
                    heading1.style.font.bold = True
                    heading1.style.font.color.rgb = RGBColor(0, 0, 0)
                    # add second heading
                    paragraph = self.doc.add_paragraph()
                    paragraph.style.font.name = 'Times New Roman'
                    paragraph.style.font.size = Pt(14)
                    left_run = paragraph.add_run('CCLS RELIZANE\t\t\t\t\t\t\t\t\t\t\t\t\t')
                    left_run.bold = True
                    right_run = paragraph.add_run('Produit:' + self.produitphyto)
                    right_run.bold = True
                    table = self.doc.add_table(rows=1, cols=11)
                    table.style = "Table Grid"  # set the table style
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'DATE'
                    hdr_cells[1].text = 'MAGASIN DE STOCK'
                    hdr_cells[2].text = 'PRODUIT'
                    hdr_cells[3].text = 'ORIGINE'
                    hdr_cells[4].text = 'QUANTITE'
                    hdr_cells[5].text = 'ACHTEUR'
                    hdr_cells[6].text = 'SACHERIE DE CND'
                    hdr_cells[7].text = 'IMATRICUL'
                    hdr_cells[8].text = 'Date de récolte'
                    hdr_cells[9].text = 'Date de CND'
                    hdr_cells[10].text = 'Date de DLUM'
                    # Set the width of the header cells
                    hdr_cells[0].width = Inches(2)  # set the width of the first column to 1.5 inches
                    hdr_cells[1].width = Inches(2)
                    hdr_cells[2].width = Inches(2)
                    hdr_cells[3].width = Inches(1.5)
                    hdr_cells[4].width = Inches(1.5)
                    hdr_cells[5].width = Inches(3)
                    hdr_cells[6].width = Inches(1.5)
                    hdr_cells[7].width = Inches(1.5)
                    hdr_cells[8].width = Inches(2)
                    hdr_cells[9].width = Inches(2)
                    hdr_cells[10].width = Inches(2)
                    for cell in table.rows[0].cells:
                        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                        cell.paragraphs[0].runs[0].font.size = Pt(11.5)
                        cell.paragraphs[0].runs[0].font.bold = True
                    # Add the data to the table
                    table_rows = len(rows)
                    table_cols = len(rows[0])
                    for row in range(table_rows):
                        table.add_row()
                        for col in range(table_cols):
                            cell = table.cell(row + 1, col)
                            cell.text = str(rows[row][col])
                            # Set font properties
                            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                            cell.paragraphs[0].runs[0].font.size = Pt(10)
                            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    for row in table.rows:
                        row.height = Inches(0.6)
                    # Save and open the document for printing
                    temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                    self.doc.save(temp_file)
                    os.startfile(temp_file, "print")
                    conn.close()
            except Exception as e:
                print('e')


        def add_table(self):
            try:

                unité = self.Unite.currentText()
                produit = self.Produit.currentText()
                origineProduit = self.OrigineDuProduit.text()
                DLUM = self.DLUM.text()
                qauntite = self.quantite.text()
                rows=self.txtfacture.rowCount()
                self.txtfacture.insertRow(rows)
                self.txtfacture.setItem(rows,0, QTableWidgetItem(unité))
                self.txtfacture.setItem(rows,1,QTableWidgetItem(produit))
                self.txtfacture.setItem(rows,2,QTableWidgetItem(qauntite))
                self.txtfacture.setItem(rows,3,QTableWidgetItem(origineProduit))
                self.txtfacture.setItem(rows,4,QTableWidgetItem(DLUM))
            except Exception as e:
                print(e)

        def stock_final_magasin(self):
            self.selected_produit = self.Produits.currentText()
            self.selected_magasin = self.Unites.currentText()
            self.selected_origine = self.OrigineDuProduits.currentText()
            self.selected_expediteur = self.expediteur.currentText()
            self.selected_dates = self.date_edits.text()
            self.selected_datesAU = self.date_edit2.text()
            if self.selected_produit and self.selected_magasin and self.selected_origine and  self.selected_expediteur and self.selected_dates == "--" and self.selected_datesAU == "--":
                if self.Produits.currentText() == "RIZ IMP":
                    self.operation_dattabase()
                    self.impot_filter_stock()
                    self.style = self.stockFinal.setStyleSheet(
                        "font: bold 16px;background-color:#88ffaa;color:#000000;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                    self.stockFinal.setText(self.Produits.currentText() + "=" + self.resultstock + " QX")
                elif self.Produits.currentText() == "POIS CHICHES IMP 12 mm":
                    self.poischicheDouze()
                    self.impot_filter_stock()
                    self.style = self.stockFinal.setStyleSheet(
                        "font: bold 16px;background-color:#88ffaa;color:#000000;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                    self.stockFinal.setText(self.Produits.currentText() + "=" + self.restpoichichedouze + " QX")
                elif self.Produits.currentText() == "POIS CHICHES IMP 09 mm":
                    self.poischicheNeuf()
                    self.impot_filter_stock()
                    self.style = self.stockFinal.setStyleSheet(
                        "font: bold 16px;background-color:#88ffaa;color:#000000;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                    self.stockFinal.setText(self.Produits.currentText() + "=" + self.restpoichicheneuf + " QX")
                elif self.Produits.currentText() == "POIS CHICHES IMP 08 mm":
                    self.poischicheOuit()
                    self.impot_filter_stock()
                    self.style = self.stockFinal.setStyleSheet(
                        "font: bold 16px;background-color:#88ffaa;color:#000000;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                    self.stockFinal.setText(self.Produits.currentText() + "=" + self.restpoichicheouit + " QX")
                elif self.Produits.currentText() == "POIS CHICHES IMP 06 mm":
                    self.poichichSIxSept()
                    self.impot_filter_stock()
                    self.style = self.stockFinal.setStyleSheet(
                        "font: bold 16px;background-color:#88ffaa;color:#000000;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                    self.stockFinal.setText(self.Produits.currentText() + "=" + self.restpoichiche + " QX")
                elif self.Produits.currentText() == "LENTILLE PL":
                    self.lenpl()
                    self.impot_filter_stock()
                    self.style = self.stockFinal.setStyleSheet(
                        "font: bold 16px;background-color:#88ffaa;color:#000000;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                    self.stockFinal.setText(self.Produits.currentText() + "=" + self.restlenpl + " QX")
                elif self.Produits.currentText() == "LENTILLE IMP 05-07 mm":
                    self.lentilleCinqSept()
                    self.impot_filter_stock()
                    self.style = self.stockFinal.setStyleSheet(
                        "font: bold 16px;background-color:#88ffaa;color:#000000;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                    self.stockFinal.setText(self.Produits.currentText() + "=" + self.restlentsinq + " QX")
                elif self.Produits.currentText() == "LENTILLE IMP vert":
                    self.lenVert()
                    self.impot_filter_stock()
                    self.style = self.stockFinal.setStyleSheet(
                        "font: bold 16px;background-color:#88ffaa;color:#000000;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                    self.stockFinal.setText(self.Produits.currentText() + "=" + self.restlentvert + " QX")

                elif self.Produits.currentText() == "LENTILLE IMP ROUGE":
                    self.lenRouge()
                    self.impot_filter_stock()
                    self.style = self.stockFinal.setStyleSheet(
                        "font: bold 16px;background-color:#88ffaa;color:#000000;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                    self.stockFinal.setText(self.Produits.currentText() + "=" + self.restlentrouge + " QX")
                elif self.Produits.currentText() == "HARICOT BLANC IMP":
                    self.haricot()
                    self.impot_filter_stock()
                    self.style = self.stockFinal.setStyleSheet(
                        "font: bold 16px;background-color:#88ffaa;color:#000000;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                    self.stockFinal.setText(self.Produits.currentText() + "=" + self.restharicotimp + " QX")
                elif self.Produits.currentText() == "HARICOT LSB":
                    self.haricotLsb()
                    self.impot_filter_stock()
                    self.style = self.stockFinal.setStyleSheet(
                        "font: bold 16px;background-color:#88ffaa;color:#000000;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                    self.stockFinal.setText(self.Produits.currentText() + "=" + self.restharicotlsb + " QX")
            elif self.selected_produit and self.selected_magasin and self.selected_origine == "" and self.selected_expediteur=="" and self.selected_dates == "--" and self.selected_datesAU == "--":
                conn = mysql.connector.connect(host="localhost",  user=user, password=password,
                                               database="datta_legumesec_entry")
                dattabase1 = mysql.connector.connect(host="localhost",  user=user, password=password,
                                                     database='datta_legumsec_out')
                cursor = conn.cursor()
                cursor.execute(
                    "SELECT le,produit,origine,quantité,unité,sacherie,fournisseur, imatricule,date_recolte,date_CND, DLUM FROM entrytable WHERE produit = %s AND unité = %s",
                    (self.selected_produit, self.selected_magasin,))
                rows = cursor.fetchall()
                curs1x = dattabase1.cursor()
                curs1x.execute(
                    "SELECT SUM(quantitérizentré) - SUM(quantitérizesorté) as quantitérizerest FROM outtable WHERE  produit= %s AND unité = %s",
                    (self.selected_produit, self.selected_magasin,))
                self.resulxs = curs1x.fetchone()
                self.resulxttp = str(self.resulxs[0])
                curs1xtv = dattabase1.cursor()
                curs1xtv.execute(
                    "SELECT SUM(quantitérizetvntré) - SUM(quantitérizetvsorté) as quantitérizetvrest FROM outtable WHERE  produit= %s AND unité = %s",
                    (self.selected_produit, self.selected_magasin,))
                self.resulxstv = curs1xtv.fetchone()
                self.resulxtttvp = str(self.resulxstv[0])
                curstt = dattabase1.cursor()
                curstt.execute(
                    "SELECT sum(quantitepchentré) - sum(quantitépchsorté) as quantitépchrest  FROM outtable   WHERE   produit=%s AND unité = %s",
                    (self.selected_produit, self.selected_magasin,))
                resultpch_att = curstt.fetchone()
                self.restpoichichedouzettp = str(resultpch_att[0])
                curstt = dattabase1.cursor()
                curstt.execute(
                    "SELECT SUM(quantitépchsixentré) - SUM(quantitépchsixsorté) as quantitépchsixrest FROM outtable WHERE   produit=%s AND unité = %s",
                    (self.selected_produit, self.selected_magasin,))
                resultpch_att = curstt.fetchone()
                self.restpoichichettp = str(resultpch_att[0])

                curs5tt = dattabase1.cursor()
                curs5tt.execute(
                    "SELECT sum(quantitépchplentré) - sum(quantitépchplsorté) as quantitépchplhrest FROM outtable WHERE   produit=%s AND unité = %s",
                    (self.selected_produit, self.selected_magasin,))
                resultpchpl_att = curs5tt.fetchone()
                self.restpoischicheplttp = str(resultpchpl_att[0])
                curs7tt = dattabase1.cursor()
                curs7tt.execute(
                    "SELECT sum(quantitélenplentré) - sum(quantitélenplsorté) as quantitélenplrest FROM outtable WHERE   produit=%s AND unité = %s",
                    (self.selected_produit, self.selected_magasin,))
                resultlenpl_att = curs7tt.fetchone()
                self.restlenplttp = str(resultlenpl_att[0])
                curs9tt = dattabase1.cursor()
                curs9tt.execute(
                    "SELECT sum(quantitélenvertentré) - sum(quantitélenvertsorté) as quantitélenrest FROM outtable WHERE   produit=%s AND unité = %s",
                    (self.selected_produit, self.selected_magasin,))
                resultlenverttt = curs9tt.fetchone()
                self.restlentvertttp = str(resultlenverttt[0])
                curs9tt = dattabase1.cursor()
                curs9tt.execute(
                    "SELECT sum(quantitélenrougeentré) - sum(quantitélenrougesorté) as quantitélenrougerest FROM outtable WHERE   produit=%s AND unité = %s",
                    (self.selected_produit, self.selected_magasin,))
                resultlenrougett = curs9tt.fetchone()
                self.restlentrougettp = str(resultlenrougett[0])
                curs11tt = dattabase1.cursor()
                curs11tt.execute(
                    "SELECT sum(quantitéharentré) - sum(quantitéharsorté) as quantitéharrest FROM outtable WHERE   produit=%s AND unité = %s",
                    (self.selected_produit, self.selected_magasin,))
                resulthar_att = curs11tt.fetchone()
                self.restharicotimpttp = str(resulthar_att[0])
                curs11tt.execute(
                    "SELECT SUM(quantitéharlsbentré) - SUM(quantitéharlsbsorté) as quantitéharlsbrest FROM outtable WHERE   produit=%s AND unité = %s",
                    (self.selected_produit, self.selected_magasin,))
                resultharlsb_att = curs11tt.fetchone()
                self.restharicotlsbttp = str(resultharlsb_att[0])
                curstt = dattabase1.cursor()
                curstt.execute(
                    "SELECT sum(quantitépchneufeentré) - sum(quantitépchneufesorté) as quantitépchneuferest FROM outtable   WHERE produit=%s AND unité = %s",
                    (self.selected_produit, self.selected_magasin,))
                resultpch_att = curstt.fetchone()
                self.restpoichicheneufttp = str(resultpch_att[0])
                curstt = dattabase1.cursor()
                curstt.execute(
                    "SELECT sum(quantitépchhuitentré) - sum(quantitépchhuitsorté) as quantitépchhuitrest FROM outtable WHERE   produit=%s AND unité = %s",
                    (self.selected_produit, self.selected_magasin,))
                resultpch_att = curstt.fetchone()
                self.restpoichicheouitttp = str(resultpch_att[0])
                curs9tt = dattabase1.cursor()
                curs9tt.execute(
                    "SELECT sum(quantitélenentré) - sum(quantitélensorté) as quantitélenrest FROM outtable WHERE   produit=%s AND unité = %s",
                    (self.selected_produit, self.selected_magasin,))
                resultlensinqtt = curs9tt.fetchone()
                self.restlentsinqttp = str(resultlensinqtt[0])
                curstt = dattabase1.cursor()
                curstt.execute(
                    "SELECT SUM(quantitépchsixentré) - SUM(quantitépchsixsorté) as quantitépchsixrest FROM outtable  WHERE produit=%s AND unité = %s",
                    (self.selected_produit, self.selected_magasin,))
                resultpch_attp = curstt.fetchone()
                self.restpoichichettp = str(resultpch_attp[0])

                curs5tt = dattabase1.cursor()
                curs5tt.execute(
                    "SELECT sum(quantitépchplentré) - sum(quantitépchplsorté) as quantitépchplhrest FROM outtable   WHERE produit=%s AND unité = %s",
                    (self.selected_produit, self.selected_magasin,))
                resultpchpl_att = curs5tt.fetchone()
                self.restpoischicheplttp = str(resultpchpl_att[0])
                curs7tt = dattabase1.cursor()
                curs7tt.execute(
                    "SELECT sum(quantitélenplentré) - sum(quantitélenplsorté) as quantitélenplrest FROM outtable  WHERE produit=%s AND unité = %s",
                    (self.selected_produit, self.selected_magasin,))
                resultlenpl_att = curs7tt.fetchone()
                self.restlenplttp = str(resultlenpl_att[0])
                curs9tt = dattabase1.cursor()
                curs9tt.execute(
                    "SELECT sum(quantitélenvertentré) - sum(quantitélenvertsorté) as quantitélenrest FROM outtable WHERE produit=%s AND unité = %s",
                    (self.selected_produit, self.selected_magasin,))
                resultlenverttt = curs9tt.fetchone()
                self.restlentvertttp = str(resultlenverttt[0])
                curs9tt = dattabase1.cursor()
                curs9tt.execute(
                    "SELECT sum(quantitélenrougeentré) - sum(quantitélenrougesorté) as quantitélenrougerest FROM outtable WHERE produit=%s AND unité = %s",
                    (self.selected_produit, self.selected_magasin,))
                resultlenrougett = curs9tt.fetchone()
                self.restlentrougettp = str(resultlenrougett[0])
                curs11tt = dattabase1.cursor()
                curs11tt.execute(
                    "SELECT sum(quantitéharentré) - sum(quantitéharsorté) as quantitéharrest FROM outtable WHERE  produit=%s AND unité = %s",
                    (self.selected_produit, self.selected_magasin,))
                resulthar_att = curs11tt.fetchone()
                self.restharicotimpttp = str(resulthar_att[0])
                curs11tt.execute(
                    "SELECT SUM(quantitéharlsbentré) - SUM(quantitéharlsbsorté) as quantitéharlsbrest FROM outtable   WHERE produit=%s AND unité = %s",
                    (self.selected_produit, self.selected_magasin,))
                resultharlsb_att = curs11tt.fetchone()
                self.restharicotlsbttp = str(resultharlsb_att[0])
                curstt = dattabase1.cursor()
                curstt.execute(
                    "SELECT sum(quantitépchneufeentré) - sum(quantitépchneufesorté) as quantitépchneuferest FROM outtable  WHERE produit=%s AND unité = %s",
                    (self.selected_produit, self.selected_magasin,))
                resultpch_att = curstt.fetchone()
                self.restpoichicheneufttp = str(resultpch_att[0])
                curstt = dattabase1.cursor()
                curstt.execute(
                    "SELECT sum(quantitépchhuitentré) - sum(quantitépchhuitsorté) as quantitépchhuitrest FROM outtable WHERE produit=%s AND unité = %s",
                    (self.selected_produit, self.selected_magasin,))
                resultpch_att = curstt.fetchone()
                self.restpoichicheouitttp = str(resultpch_att[0])
                curs9tt = dattabase1.cursor()
                curs9tt.execute(
                    "SELECT sum(quantitélenentré) - sum(quantitélensorté) as quantitélenrest FROM outtable WHERE  produit=%s AND unité = %s",
                    (self.selected_produit, self.selected_magasin,))
                resultlensinqtt = curs9tt.fetchone()
                self.restlentsinqttp = str(resultlensinqtt[0])
                if self.selected_produit == "POIS CHICHES IMP 12 mm":
                    self.impot_filter_stock()
                    self.style = self.stockFinal.setStyleSheet("font: bold 16px;background-color:#88ffaa;color:#000000;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                    self.stockFinal.setText(self.Produits.currentText() + "=" + self.restpoichichedouzettp + " QX")
                if self.selected_produit == "POIS CHICHES IMP 09 mm":
                    self.impot_filter_stock()
                    self.style = self.stockFinal.setStyleSheet("font: bold 16px;background-color:#88ffaa;color:#000000;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                    self.stockFinal.setText(self.Produits.currentText() + "=" +self.restpoichicheneufttpp + " QX")
                if self.selected_produit == "POIS CHICHES IMP 08 mm":
                    self.impot_filter_stock()
                    self.style = self.stockFinal.setStyleSheet("font: bold 16px;background-color:#88ffaa;color:#000000;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                    self.stockFinal.setText(self.Produits.currentText() + "=" + self.restpoichicheouitttp + " QX")
                if self.selected_produit == "POIS CHICHES IMP 06 mm":
                    self.impot_filter_stock()
                    self.style = self.stockFinal.setStyleSheet("font: bold 16px;background-color:#88ffaa;color:#000000;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                    self.stockFinal.setText(self.Produits.currentText() + "=" + self.restpoichichettp + " QX")
                if self.selected_produit == "HARICOT BLANC IMP":
                    self.impot_filter_stock()
                    self.style = self.stockFinal.setStyleSheet("font: bold 16px;background-color:#88ffaa;color:#000000;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                    self.stockFinal.setText(self.Produits.currentText() + "=" + self.restharicotimpttp + " QX")
                if self.selected_produit == "HARICOT LSB":
                    self.impot_filter_stock()
                    self.style = self.stockFinal.setStyleSheet("font: bold 16px;background-color:#88ffaa;color:#000000;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                    self.stockFinal.setText(self.Produits.currentText() + "=" + self.restharicotlsbttp + " QX")
                if self.selected_produit == "LENTILLE PL":
                    self.impot_filter_stock()
                    self.style = self.stockFinal.setStyleSheet("font: bold 16px;background-color:#88ffaa;color:#000000;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                    self.stockFinal.setText(self.Produits.currentText() + "=" + self.restlenplttp + " QX")
                if self.selected_produit == "LENTILLE IMP 05-07 mm":
                    self.impot_filter_stock()
                    self.style = self.stockFinal.setStyleSheet("font: bold 16px;background-color:#88ffaa;color:#000000;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                    self.stockFinal.setText(self.Produits.currentText() + "=" + self.restlentsinqttp + " QX")
                if self.selected_produit == "LENTILLE IMP vert":
                    self.impot_filter_stock()
                    self.style = self.stockFinal.setStyleSheet("font: bold 16px;background-color:#88ffaa;color:#000000;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                    self.stockFinal.setText(self.Produits.currentText() + "=" + self.restlentvertttp + " QX")
                if self.selected_produit == "LENTILLE IMP ROUGE":
                    self.impot_filter_stock()
                    self.style = self.stockFinal.setStyleSheet("font: bold 16px;background-color:#88ffaa;color:#000000;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                    self.stockFinal.setText(self.Produits.currentText() + "=" + self.restlentrougettp + " QX")
                if self.selected_produit == "RIZ IMP":
                    self.impot_filter_stock()
                    self.style = self.stockFinal.setStyleSheet("font: bold 16px;background-color:#88ffaa;color:#000000;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                    self.stockFinal.setText(self.Produits.currentText() + "=" + self.resulxttp + " QX")
                if self.selected_produit == "RIZ ETUVE":
                    self.impot_filter_stock()
                    self.style = self.stockFinal.setStyleSheet("font: bold 16px;background-color:#88ffaa;color:#000000;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                    self.stockFinal.setText(self.Produits.currentText() + "=" + self.resulxtttvp + " QX")


            else:
                self.stockFinal.setStyleSheet("font: bold 16px;background-color:#ffffff;color:#000000;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                self.stockFinal.setText("0 ")

        def impot_filter_stock(self):
            try:

                produitstock = self.Produits.currentText()
                expediteurstock = self.expediteur.currentText()
                unitestock=self.Unites.currentText()
                originestock=self.OrigineDuProduits.currentText()
                if produitstock and expediteurstock and unitestock and originestock:
                    dattabase = mysql.connector.connect(host="localhost", user=user, password=password,
                                                        database='datta_legumesec_entry')
                    curs = dattabase.cursor()
                    curs.execute(
                        "SELECT unité,produit ,quantité ,origine,date_CND,DLUM,nlot,nlotsch,fournisseur  FROM entrytable WHERE produit = %s AND  unité= %s  AND fournisseur= %s  AND origine= %s",
                        (produitstock, unitestock, expediteurstock, originestock,))

                    result = curs.fetchall()
                    self.txtfactures.setRowCount(0)
                    for row, row_datta in enumerate(result):
                        self.txtfactures.insertRow(row)
                        for colum, datta in enumerate(row_datta):
                            self.txtfactures.setItem(row, colum, QTableWidgetItem(str(datta)))
                    dattabase.commit()
                    dattabase.close()
                if produitstock and unitestock and originestock=='' and expediteurstock=='':
                    dattabase = mysql.connector.connect(host="localhost", user=user, password=password,
                                                        database='datta_legumesec_entry')
                    curs = dattabase.cursor()
                    curs.execute("SELECT unité,produit ,quantité ,origine,date_CND,DLUM,nlot,nlotsch,fournisseur  FROM entrytable WHERE produit = %s AND  unité= %s ",(produitstock, unitestock,))

                    result = curs.fetchall()
                    self.txtfactures.setRowCount(0)
                    for row, row_datta in enumerate(result):
                        self.txtfactures.insertRow(row)
                        for colum, datta in enumerate(row_datta):
                            self.txtfactures.setItem(row, colum, QTableWidgetItem(str(datta)))
                    dattabase.commit()
                    dattabase.close()
            except mysql.connector.Error as e :
                print(e)


        def docxStock(self):
            try:
                dialog = QDialog()
                dialog.setStyleSheet(""" 
                                    QPushButton
        {
            color: #00000;
            background-color:#ade3e7;
            border-width: 1px;
            border-color: #1e1e1e;
            border-style: solid;
            border-radius: 6;
            padding: 3px;
            font-size: 12px;
            padding-left: 5px;
            padding-right: 5px;
            min-width: 40px;

        }

        QPushButton:disabled
        {
            background-color:#03ecff;
            border-width: 1px;
            border-color: #454545;
            border-style: solid;
            padding-top: 5px;
            padding-bottom: 5px;
            padding-left: 10px;
            padding-right: 10px;
            border-radius: 2px;
            color: #454545;
        }
        QPushButton:pressed
        {
            background-color: #3daee9;
            padding-top: -15px;
            padding-bottom: -17px;
        }
                                     QLabel
                                    {
                                        font-size: 18px;
                                        border: 0px solid orange;
                                    }

                                """)
                dialog.setGeometry(QtCore.QRect(500,200,300,100))
                dialog.setWindowTitle("Sélectionnez  ")
                #dialog.setText("Sélectionnez la date et produit souhaitée\n")
                ok_button = QtWidgets.QPushButton("ENTRE", dialog,clicked=lambda :self.print_docxStockDatabase())
                ok_button.setGeometry(QtCore.QRect(40,30,100,40))
                cancel_button = QtWidgets.QPushButton("SORTIE", dialog,clicked=lambda :self.print_docxStockDatabaseSortie())
                cancel_button.setGeometry(QtCore.QRect(160,30,100,40))
                #dialog.addButton(ok_button, QMessageBox.ButtonRole.ActionRole)
                #dialog.addButton(cancel_button, QMessageBox.ButtonRole.ActionRole)
                dialog.exec()
            except Exception as e:
                print(e)




    if __name__ == "__main__":
        import sys

        app = QtWidgets.QApplication(sys.argv)
        MainWindow = QtWidgets.QMainWindow()
        ui = Stock_Legumesec()
        ui.stock_legumesec(MainWindow)
        MainWindow.show()
        sys.exit(app.exec())

except Exception as e:
    print(e)






