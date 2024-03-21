import locale
import os
import tempfile
import docx
import docxtpl
from PyQt6 import QtCore, QtGui, QtWidgets
from PyQt6.QtCore import QDate, Qt, QLocale, QTimer
from PyQt6.QtGui import QIcon
from PyQt6.QtWidgets import *
import datetime
import csv
import sqlite3
from docx.enum.section import WD_ORIENTATION
from docx.shared import Inches, Pt, RGBColor
from docxtpl import DocxTemplate
from PyQt6.QtPrintSupport import QPrinter, QPrintDialog
import mysql.connector

userslist = ["nadir", "qualite"]
passwordlist = ["Nadir206@", "qualite48"]

for user ,password in zip(userslist,passwordlist):
    database = mysql.connector.connect(host="localhost", user=user, password=password)
    curs = database.cursor()
    curs.execute("CREATE DATABASE IF NOT EXISTS datta_refus")
    curs.execute("USE datta_refus")
    curs.execute("""CREATE TABLE IF NOT EXISTS refutable
                    (
                      id INT AUTO_INCREMENT PRIMARY KEY,
                      point_de_collecte TEXT,
                      date TEXT, 
                      nom_et_prénom TEXT,
                      fils_de TEXT,    
                      matricule TEXT,
                      produit TEXT, 
                      cause_de_refus TEXT

                    )
                    """)
    database.commit()
    database.close()
try:
    class Refus_Window():
        def refus_produit(self, MainWindow):
            MainWindow.setObjectName("Window")
            MainWindow.resize(1338, 700)
            MainWindow.setStyleSheet("""QToolTip
            {
                border: 1px solid #76797C;
                background-color: rgb(90, 102, 117);;
                color: white;
                padding: 5px;
                opacity: 200;
            }

            QWidget
            {
                color: #000000;
                background-color: #ffffff;
                selection-background-color:#3daee9;
                selection-color: #3daee9;
                background-clip: border;
                border-image: none;
                border: 0px transparent black;
                outline: 0;
            }

            QWidget:item:hover
            {
                background-color: #3daee9;
                color: #eff0f1;
            }

            QWidget:item:selected
            {
                background-color: #3daee9;
            }



            QWidget:disabled
            {
                color: #454545;
                background-color: #31363b;
            }

            QAbstractItemView
            {
                alternate-background-color: #31363b;
                color: #eff0f1;
                border: 1px solid 3A3939;
                border-radius: 2px;
            }

            QWidget:focus, QMenuBar:focus
            {
                border: 1px solid #3daee9;
            }

            QTabWidget:focus, QCheckBox:focus, QRadioButton:focus, QSlider:focus
            {
                border: none;
            }

            QLineEdit
            {
                background-color: #FDFEFE;
                padding: 1px;
                border-style: solid;
                border: 1px solid #76797C;
                border-radius: 0px;
                color: #000000;
            }
            QDoubleSpinBox
            {
                background-color: #FDFEFE;
                padding: 1px;
                border-style: solid;
                border: 1px solid #76797C;
                border-radius: 0px;
                color:#000000;
                font-size: 11px;
                font-weight: bold;

            }
            QDoubleSpinBox:focus{
                background-color: #FDFEFE;
                border-style: solid;
                border: 2px solid #76797C;
                border-radius: 4px;
                border-color: #ff8c00;
            }
            QDoubleSpinBox::drop-down
            {
                subcontrol-origin: padding;
                subcontrol-position: top right;
                width: 1px;

                border-left-width: 0px;
                border-left-color: #232629;
                border-left-style: solid;
                border-top-right-radius: 1px;
                border-bottom-right-radius: 1px;
            }



            QGroupBox {
                border:1px solid #76797C;
                border-radius: 2px;
                margin-top: 5px;
            }

            QGroupBox::title {
                subcontrol-origin: margin;
                subcontrol-position: top center;
                padding-left: 4px;
                padding-right: 4px;
                padding-top: 4px;
            }

            QAbstractScrollArea
            {
                border-radius: 2px;
                border: 1px solid #76797C;
                background-color: transparent;
            }

            QScrollBar:horizontal
            {
                height: 15px;
                margin: 3px 15px 3px 15px;
                border: 1px transparent #2A2929;
                border-radius: 4px;
                background-color: #2A2929;
            }

            QScrollBar::handle:horizontal
            {
                background-color: #605F5F;
                min-width: 5px;
                border-radius: 4px;
            }

            QScrollBar::add-line:horizontal
            {
                margin: 0px 3px 0px 3px;
                border-image: url(:/qss_icons/Dark_rc/right_arrow_disabled.png);
                width: 10px;
                height: 10px;
                subcontrol-position: right;
                subcontrol-origin: margin;
            }

            QScrollBar::sub-line:horizontal
            {
                margin: 0px 3px 0px 3px;
                border-image: url(:/qss_icons/Dark_rc/left_arrow_disabled.png);
                height: 10px;
                width: 10px;
                subcontrol-position: left;
                subcontrol-origin: margin;
            }

            QScrollBar::add-line:horizontal:hover,QScrollBar::add-line:horizontal:on
            {
                border-image: url(:/qss_icons/Dark_rc/right_arrow.png);
                height: 10px;
                width: 10px;
                subcontrol-position: right;
                subcontrol-origin: margin;
            }


            QScrollBar::sub-line:horizontal:hover, QScrollBar::sub-line:horizontal:on
            {
                border-image: url(:/qss_icons/Dark_rc/left_arrow.png);
                height: 10px;
                width: 10px;
                subcontrol-position: left;
                subcontrol-origin: margin;
            }

            QScrollBar::up-arrow:horizontal, QScrollBar::down-arrow:horizontal
            {
                background: none;
            }


            QScrollBar::add-page:horizontal, QScrollBar::sub-page:horizontal
            {
                background: none;
            }

            QScrollBar:vertical
            {
                background-color: #2A2929;
                width: 15px;
                margin: 15px 3px 15px 3px;
                border: 1px transparent #2A2929;
                border-radius: 4px;
            }

            QScrollBar::handle:vertical
            {
                background-color: #605F5F;
                min-height: 5px;
                border-radius: 4px;
            }

            QScrollBar::sub-line:vertical
            {
                margin: 3px 0px 3px 0px;
                border-image: url(:/qss_icons/Dark_rc/up_arrow_disabled.png);
                height: 10px;
                width: 10px;
                subcontrol-position: top;
                subcontrol-origin: margin;
            }

            QScrollBar::add-line:vertical
            {
                margin: 3px 0px 3px 0px;
                border-image: url(:/qss_icons/Dark_rc/down_arrow_disabled.png);
                height: 10px;
                width: 10px;
                subcontrol-position: bottom;
                subcontrol-origin: margin;
            }

            QScrollBar::sub-line:vertical:hover,QScrollBar::sub-line:vertical:on
            {

                border-image: url(:/qss_icons/Dark_rc/up_arrow.png);
                height: 10px;
                width: 10px;
                subcontrol-position: top;
                subcontrol-origin: margin;
            }


            QScrollBar::add-line:vertical:hover, QScrollBar::add-line:vertical:on
            {
                border-image: url(:/qss_icons/Dark_rc/down_arrow.png);
                height: 10px;
                width: 10px;
                subcontrol-position: bottom;
                subcontrol-origin: margin;
            }

            QScrollBar::up-arrow:vertical, QScrollBar::down-arrow:vertical
            {
                background: none;
            }


            QScrollBar::add-page:vertical, QScrollBar::sub-page:vertical
            {
                background: none;
            }

            QTextEdit
            {
                background-color: #FDFEFE;
                color: #000000;
                border: 1px solid #76797C;
            }

            QPlainTextEdit
            {
                background-color: #232629;;
                color: #000000;
                border-radius: 2px;
                border: 1px solid #76797C;
            }

            QHeaderView::section
            {
                background-color: #76797C;
                color: #eff0f1;
                padding: 1px;
                border: 1px solid #76797C;
            }

            QSizeGrip {
                width: 12px;
                height: 12px;
            }


            QMainWindow::separator
            {
                background-color: #31363b;
                color: white;
                padding-left: 4px;
                spacing: 2px;
                border: 1px dashed #76797C;
            }

            QMainWindow::separator:hover
            {

                background-color: #787876;
                color: white;
                padding-left: 4px;
                border: 1px solid #76797C;
                spacing: 2px;
            }


            QMenu::separator
            {
                height: 1px;
                background-color: #76797C;
                color: white;
                padding-left: 4px;
                margin-left: 10px;
                margin-right: 5px;
            }


            QFrame
            {
                border-radius: 2px;
                border: 1px solid #76797C;
            }

            QFrame[frameShape="0"]
            {
                border-radius: 2px;
                border: 1px transparent #76797C;
            }

            QStackedWidget
            {
                border: 1px transparent black;
            }


            QPushButton
            {
                color: #00000;
                background-color:#ade3e7;
                border-width: 1px;
                border-color: #1e1e1e;
                border-style: solid;
                border-radius: 6;
                padding: 3px;
                font-size: 12px;
                padding-left: 5px;
                padding-right: 5px;
                min-width: 40px;

            }

            QPushButton:disabled
            {
                background-color:#03ecff;
                border-width: 1px;
                border-color: #454545;
                border-style: solid;
                padding-top: 5px;
                padding-bottom: 5px;
                padding-left: 10px;
                padding-right: 10px;
                border-radius: 2px;
                color: #454545;
            }
            QPushButton:pressed
            {
                background-color: #3daee9;
                padding-top: -15px;
                padding-bottom: -17px;
            }

            QComboBox
            {
               background-color: #FDFEFE;
                border-style: solid;
                border: 1px solid #76797C;
                border-radius: 2px;
                min-width: 40px;
            }

            QPushButton:checked{
                background-color: #76797C;
                border-color: #6A6969;
            }

            QComboBox:hover,QDoubleSpinBox:Hover,QPushButton:hover,QAbstractSpinBox:hover,QLineEdit:hover,QTextEdit:hover,QPlainTextEdit:hover,QAbstractView:hover,QTreeView:hover
            {
                border: 1px solid #ff8c00;
                color: #000000;
            }

            QComboBox:on
            {
                padding-top: 1px;
                padding-left: 1px;
                selection-background-color: #FDFEFE;
            }

            QComboBox QAbstractItemView
            {
                background-color: #FDFEFE;
                border-radius: 2px;
                border: 1px solid #76797C;
                color:#000000;
                selection-background-color: #000000;
            }

            QComboBox::drop-down
            {
                subcontrol-origin: padding;
                subcontrol-position: top right;
                width: 15px;

                border-left-width: 0px;
                border-left-color: ff8c00;
                border-left-style: solid;
                border-top-right-radius: 1px;
                border-bottom-right-radius: 1px;
            }


            QLabel
            {
                border: 2px solid black;
            }

            QTabWidget{
                border: 0px transparent black;
            }

            QTabWidget::pane {
                border: 1px solid #76797C;
                padding: 5px;
                margin: 0px;
            }

            QTabBar
            {
                qproperty-drawBase: 0;
                left: 5px; /* move to the right by 5px */
                border-radius: 3px;
            }

            QTabBar:focus
            {
                border: 0px transparent black;
            }

            QTabBar::close-button  {
                image: url(:/qss_icons/Dark_rc/close.png);
                background: transparent;
            }

            QTabBar::close-button:hover
            {
                image: url(:/qss_icons/Dark_rc/close-hover.png);
                background: transparent;
            }

            QTabBar::close-button:pressed {
                image: url(:/qss_icons/Dark_rc/close-pressed.png);
                background: transparent;
            }

            /* TOP TABS */
            QTabBar::tab:top {
                color: #eff0f1;
                border: 1px solid #76797C;
                border-bottom: 1px transparent black;
                background-color: #31363b;
                padding: 5px;
                min-width: 10px;
                border-top-left-radius: 2px;
                border-top-right-radius: 2px;
            }

            QTabBar::tab:top:!selected
            {
                color: #eff0f1;
                background-color: #54575B;
                border: 1px solid #76797C;
                border-bottom: 1px transparent black;
                border-top-left-radius: 2px;
                border-top-right-radius: 2px;    
            }

            QTabBar::tab:top:!selected:hover {
                background-color: #3daee9;
            }

            /* BOTTOM TABS */
            QTabBar::tab:bottom {
                color: #eff0f1;
                border: 1px solid #76797C;
                border-top: 1px transparent black;
                background-color: #31363b;
                padding: 5px;
                border-bottom-left-radius: 2px;
                border-bottom-right-radius: 2px;
                min-width: 10px;
            }

            QTabBar::tab:bottom:!selected
            {
                color: #eff0f1;
                background-color: #54575B;
                border: 1px solid #76797C;
                border-top: 1px transparent black;
                border-bottom-left-radius: 2px;
                border-bottom-right-radius: 2px;
            }

            QTabBar::tab:bottom:!selected:hover {
                background-color: #3daee9;
            }

            /* LEFT TABS */
            QTabBar::tab:left {
                color: #eff0f1;
                border: 1px solid #76797C;
                border-left: 1px transparent black;
                background-color: #31363b;
                padding: 5px;
                border-top-right-radius: 2px;
                border-bottom-right-radius: 2px;
                min-height: 50px;
            }

            QTabBar::tab:left:!selected
            {
                color: #eff0f1;
                background-color: #54575B;
                border: 1px solid #76797C;
                border-left: 1px transparent black;
                border-top-right-radius: 2px;
                border-bottom-right-radius: 2px;
            }

            QTabBar::tab:left:!selected:hover {
                background-color: #3daee9;
            }


            /* RIGHT TABS */
            QTabBar::tab:right {
                color: #eff0f1;
                border: 1px solid #76797C;
                border-right: 1px transparent black;
                background-color: #31363b;
                padding: 5px;
                border-top-left-radius: 2px;
                border-bottom-left-radius: 2px;
                min-height: 50px;
            }

            QTabBar::tab:right:!selected
            {
                color: #eff0f1;
                background-color: #54575B;
                border: 1px solid #76797C;
                border-right: 1px transparent black;
                border-top-left-radius: 2px;
                border-bottom-left-radius: 2px;
            }





            QSlider::groove:horizontal {
                border: 1px solid #565a5e;
                height: 4px;
                background: #565a5e;
                margin: 0px;
                border-radius: 2px;
            }

            QSlider::handle:horizontal {
                background: #232629;
                border: 1px solid #565a5e;
                width: 16px;
                height: 16px;
                margin: -8px 0;
                border-radius: 9px;
            }

            QSlider::groove:vertical {
                border: 1px solid #565a5e;
                width: 4px;
                background: #565a5e;
                margin: 0px;
                border-radius: 3px;
            }

            QSlider::handle:vertical {
                background: #232629;
                border: 1px solid #565a5e;
                width: 16px;
                height: 16px;
                margin: 0 -8px;
                border-radius: 9px;
            }

            QToolButton {
                background-color: transparent;
                border: 1px transparent #76797C;
                border-radius: 2px;
                margin: 3px;
                padding: 5px;
            }

            QToolButton[popupMode="1"] { /* only for MenuButtonPopup */
             padding-right: 20px; /* make way for the popup button */
             border: 1px #76797C;
             border-radius: 5px;
            }

            QToolButton[popupMode="2"] { /* only for InstantPopup */
             padding-right: 10px; /* make way for the popup button */
             border: 1px #76797C;
            }


            QToolButton:hover, QToolButton::menu-button:hover {
                background-color: transparent;
                border: 1px solid #3daee9;
                padding: 5px;
            }

            QToolButton:checked, QToolButton:pressed,
                    QToolButton::menu-button:pressed {
                background-color: #3daee9;
                border: 1px solid #3daee9;
                padding: 5px;
            }

            /* the subcontrol below is used only in the InstantPopup or DelayedPopup mode */
            QToolButton::menu-indicator {
                background-color:ff8c00;
                top: -7px; left: -2px; /* shift it a bit */
            }

            /* the subcontrols below are used only in the MenuButtonPopup mode */
            QToolButton::menu-button {
                border: 1px transparent #76797C;
                border-top-right-radius: 6px;
                border-bottom-right-radius: 6px;
                /* 16px width + 4px for border = 20px allocated above */
                width: 16px;
                outline: none;
            }

            QToolButton::menu-arrow {
               background-color:ff8c00;
            }

            QToolButton::menu-arrow:open {
                border: 1px solid #76797C;
            }

            QPushButton::menu-indicator  {
                subcontrol-origin: padding;
                subcontrol-position: bottom right;
                left: 8px;
            }

            QTableView
            {
                border: 1px solid #76797C;
                gridline-color: #31363b;
                background-color: #FDFEFE;
                color:#000000;
            }


            QTableView, QHeaderView
            {
                background-color: #FDFEFE;
                color:#000000;
                border-radius: 0px;
            }

            QTableView::item:pressed, QListView::item:pressed, QTreeView::item:pressed  {
                background: #FDFEFE;
                color: #000000;
            }

            QTableView::item:selected:active, QTreeView::item:selected:active, QListView::item:selected:active  {
                background: #3daee9;
                color: #000000;
            }


            QHeaderView
            {
                background-color: #FDFEFE;
                border: 1px transparent;
                border-radius: 0px;
                margin: 0px;
                padding: 0px;

            }

            QHeaderView::section  {
                background-color:#80f1f9;
                color: #000000;
                padding: 5px;
                border: 1px solid #76797C;
                border-radius: 0px;
                text-align: center;
            }

            QHeaderView::section::vertical::first, QHeaderView::section::vertical::only-one
            {
                border-top: 1px solid #76797C;
            }

            QHeaderView::section::vertical
            {
                border-top: transparent;
            }

            QHeaderView::section::horizontal::first, QHeaderView::section::horizontal::only-one
            {
                border-left: 1px solid #76797C;
            }

            QHeaderView::section::horizontal
            {
                border-left: transparent;
            }


            QHeaderView::section:checked
             {
                color: #000000;
                background-color: #3daee9;
             }

             /* style the sort indicator */
            QHeaderView::down-arrow {

            }

            QHeaderView::up-arrow {

            }


            QTableCornerButton::section {
                background-color: #31363b;
                border: 1px transparent #76797C;
                border-radius: 0px;
            }

            QToolBox  {
                padding: 5px;
                border: 1px transparent black;
            }

            QToolBox::tab {
                color: #eff0f1;
                background-color: #31363b;
                border: 1px solid #76797C;
                border-bottom: 1px transparent #31363b;
                border-top-left-radius: 5px;
                border-top-right-radius: 5px;
            }

            QToolBox::tab:selected { /* italicize selected tabs */
                font: italic;
                background-color: #31363b;
                border-color: #3daee9;
             }

            QStatusBar::item {
                border: 0px transparent dark;
             }


            QFrame[height="3"], QFrame[width="3"] {
                background-color: #76797C;
            }




            QDateEdit
            {
                background-color: #232629;;
                border-style: solid;
                border: 1px solid #76797C;
                border-radius: 2px;
                padding: 1px;
                min-width: 75px;
            }

            QDateEdit:on
            {
                padding-top: 2px;
                padding-left: 2px;
                selection-background-color: #4a4a4a;
            }

            QDateEdit QAbstractItemView
            {
                background-color: #ff8c00;
                border-radius: 2px;
                border: 1px solid #3375A3;
                selection-background-color:ff8c00;
            }

            QDateEdit::drop-down
            {
                subcontrol-origin: padding;
                subcontrol-position: top right;
                width: 15px;
                border-left-width: 0px;
                border-left-color: darkgray;
                border-left-style: solid;
                border-top-right-radius: 3px;
                border-bottom-right-radius: 3px;
            }   
            QDateTimeEdit
            {
                background-color: #232629;;
                border-style: solid;
                border: 1px solid #76797C;
                border-radius: 2px;
                padding: 1px;
                min-width: 75px;

            }    
            """)

            self.addbtn = QtWidgets.QPushButton("Ajouter", MainWindow, clicked=lambda: self.add_datta_save())
            self.addbtn.setGeometry(QtCore.QRect(20, 275, 100, 40))

            self.modifierbtn = QtWidgets.QPushButton("Enregistrer", MainWindow, clicked=lambda: self.SITUATION_phyto())
            self.modifierbtn.setGeometry(QtCore.QRect(260, 275, 100, 40))
            ##################################################sortie###############
            self.deletebtn = QtWidgets.QPushButton("Suprime", MainWindow, clicked=lambda: self.delete_item())
            self.deletebtn.setGeometry(QtCore.QRect(140, 275, 100, 40))

            self.printbtn = QtWidgets.QPushButton("Imprimer", MainWindow, clicked=lambda: self.print_docx())
            self.printbtn.setGeometry(QtCore.QRect(380, 275, 100, 40))

            self.datafiltertxt = QtWidgets.QLabel("Nom et Prénom: ", MainWindow)
            self.datafiltertxt.setFrameShape(QtWidgets.QFrame.Shape.WinPanel.VLine)
            self.datafiltertxt.setGeometry(QtCore.QRect(777, 268, 565, 51))

            self.nomefiltrage = QtWidgets.QLineEdit(MainWindow)
            self.nomefiltrage.setGeometry(QtCore.QRect(882, 275, 160, 40))

            self.oledtxt = QtWidgets.QLabel("Fils de: ", MainWindow)
            self.oledtxt.setGeometry(QtCore.QRect(1057, 285, 100, 20))

            self.oled = QtWidgets.QLineEdit(MainWindow)
            self.oled.setGeometry(QtCore.QRect(1100, 275, 100, 40))

            self.btnfilter = QtWidgets.QPushButton("ok", MainWindow, clicked=lambda: self.impot_filter())
            self.btnfilter.setGeometry((QtCore.QRect(1217, 275, 40, 40)))

            self.btnfcncl = QtWidgets.QPushButton("exit", MainWindow, clicked=lambda: self.impot_all())
            self.btnfcncl.setGeometry((QtCore.QRect(1282, 275, 40, 40)))

            self.filtertxt = QtWidgets.QLabel("Filtrage des données:", MainWindow)
            self.filtertxt.setGeometry(QtCore.QRect(792, 254, 120, 20))

            self.cclstxt = QtWidgets.QLabel("<h1>CCLS RELIZANE SERVICE QUALITE<h1/>", MainWindow)
            self.cclstxt.setGeometry(QtCore.QRect(430, 0, 600, 70))
            font = QtGui.QFont()
            font.setPointSize(14)
            font.setBold(True)
            font.setWeight(75)
            self.cclstxt.setFont(font)
            self.cclstxt.setMouseTracking(False)
            self.cclstxt.setLayoutDirection(QtCore.Qt.LayoutDirection.LeftToRight)
            self.cclstxt.setAutoFillBackground(False)
            self.cclstxt.setFrameShape(QtWidgets.QFrame.Shape.WinPanel)
            self.cclstxt.setLineWidth(0)
            self.cclstxt.setMidLineWidth(0)
            self.cclstxt.setTextFormat(QtCore.Qt.TextFormat.AutoText)
            self.cclstxt.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
            self.entrielabel = QtWidgets.QLabel("<h2>REGISTRE DE REFUS<h2/>", MainWindow)
            self.entrielabel.setGeometry(370, 75, 680, 30)
            self.entrielabel.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
            self.entrielabel.setFont(font)

            font = QtGui.QFont()
            font.setPointSize(14)
            font.setBold(True)
            font.setUnderline(True)
            font.setWeight(75)

            #############################################
            self.txtlieu = QtWidgets.QLabel("Point de collecte :", MainWindow)
            self.txtlieu.setGeometry(QtCore.QRect(20, 90, 110, 16))
            ####################################################
            self.lieu = QtWidgets.QComboBox(MainWindow)
            self.lieu.addItem('')
            self.lieu.addItem("DOCK CENTRAL")
            self.lieu.addItem("MAGASIN MESRA")
            self.lieu.addItem("MAGASIN KEF LAZREG")
            self.lieu.addItem('N.S.MENDES')
            self.lieu.addItem('S.MENDES')
            self.lieu.setGeometry(QtCore.QRect(150, 77, 150, 40))

            self.txtdate = QtWidgets.QLabel("DATE :", MainWindow)
            self.txtdate.setGeometry(QtCore.QRect(20, 40, 110, 16))
            ####################################################
            self.dateedite = QtWidgets.QLineEdit(MainWindow)
            self.dateedite.setInputMask("99/99/9999")
            self.dateedite.setGeometry(QtCore.QRect(150, 22, 150, 40))

            ########################################################
            self.nometprénomtxt = QtWidgets.QLabel("NOM ET PRENOM", MainWindow)
            self.nometprénomtxt.setGeometry(QtCore.QRect(20, 140, 100, 20))

            self.nomeetprénom = QtWidgets.QLineEdit(MainWindow)
            self.nomeetprénom.setGeometry(QtCore.QRect(150, 130, 150, 40))

            self.fissedetxt = QtWidgets.QLabel("Père : ", MainWindow)
            self.fissedetxt.setGeometry(QtCore.QRect(330, 140, 100, 20))

            self.fissede = QtWidgets.QLineEdit(MainWindow)
            self.fissede.setGeometry(QtCore.QRect(390, 130, 150, 40))
            #####################################################################################
            self.matriculetxt = QtWidgets.QLabel("MATRICULE", MainWindow)
            self.matriculetxt.setGeometry(QtCore.QRect(570, 130, 150, 40))
            ###################################################################
            self.matricule = QtWidgets.QLineEdit(MainWindow)
            self.matricule.setInputMask("99999-999-99")
            self.matricule.setGeometry(QtCore.QRect(660, 130, 150, 40))

            ###############################################################################
            self.produittxt = QtWidgets.QLabel("PRODUIT", MainWindow)
            self.produittxt.setGeometry(QtCore.QRect(840, 130, 120, 40))

            self.produit = QtWidgets.QComboBox(MainWindow)
            self.produit.addItem(' ')
            self.produit.addItem('Blé dur')
            self.produit.addItem('Blé tendre')
            self.produit.addItem('Orge')
            self.produit.addItem('Avoine')
            self.produit.addItem('Tritical')
            self.produit.setGeometry(QtCore.QRect(920, 130, 150, 40))
            ############################################################################
            self.causerefustxt = QtWidgets.QLabel("CAUSE DE REFUS", MainWindow)
            self.causerefustxt.setGeometry(QtCore.QRect(1100, 140, 100, 16))

            self.causerefus = QtWidgets.QTextEdit(MainWindow)
            self.causerefus.setGeometry(QtCore.QRect(1220, 130, 120, 40))

            ###########################################################################

            ########################################################################
            ####################################################################
            self.totalfont = QtGui.QFont()
            self.totalfont.setPointSize(10)
            self.totalfont.setBold(True)
            self.totalfont.bold()
            self.textEdit = QtWidgets.QTableWidget(MainWindow)
            self.textEdit.setRowCount(0)
            self.textEdit.setColumnCount(8)
            self.textEdit.setColumnWidth(0, 5)
            self.textEdit.setColumnWidth(1, 178)
            self.textEdit.setColumnWidth(2, 178)
            self.textEdit.setColumnWidth(3, 178)
            self.textEdit.setColumnWidth(4, 178)
            self.textEdit.setColumnWidth(5, 178)
            self.textEdit.setColumnWidth(6, 178)
            self.textEdit.setColumnWidth(7, 187)
            # Set the height of the row
            header_labels = ["ID", "Point de collecte", "Date", "Nom et Prénom", "fils_de", "Matricule", "Produit",
                             "Cause de refus"]
            self.textEdit.setHorizontalHeaderLabels(header_labels)
            font = self.textEdit.horizontalHeader().font()
            font.setBold(True)
            font.setPointSize(12)
            self.textEdit.horizontalHeader().setFont(font)
            self.textEdit.setGeometry(QtCore.QRect(20, 320, 1323, 340))
            # self.textEdit.setStyleSheet("background-color:rgb(255, 255, 255)")
            self.textEdit.setObjectName("textEdit")
            # self.textEdit.setStyleSheet(" background-color: #232629")
            self.textEdit.setFont(self.totalfont)
            self.impot_all()

        ####################################################################################################################################
        def SITUATION_phyto(self):
            dialog = QMessageBox()
            dialog.setStyleSheet("""
                QWidget {
                    color: #eff0f1;
                    background-color: #31363b;
                    border-width: 1px;
                    border-color: #1e1e1e;
                    border-style: solid;
                    border-radius: 6;
                    padding: 3px;
                    font-size: 18px;
                    padding-left: 5px;
                    padding-right: 5px;
                }
                QWidget:item:hover {
                    background-color: #3daee9;
                    color: #eff0f1;
                }
                QWidget:item:selected {
                    background-color: #3daee9;
                }
                QWidget:disabled {
                    color: #454545;
                    background-color: #31363b;
                }
                QPushButton {
                    color: #b1b1b1;
                    background-color: linear-gradient(x1: 0, y1: 0, x2: 0, y2: 1, stop: 0 #565656, stop: 0.1 #525252, stop: 0.5 #4e4e4e, stop: 0.9 #4a4a4a, stop: 1 #464646);
                    border-width: 1px;
                    border-color: #1e1e1e;
                    border-style: solid;
                    border-radius: 6;
                    padding: 3px;
                    font-size: 12px;
                    padding-left: 5px;
                    padding-right: 5px;
                    min-width: 40px;
                }
                QPushButton:disabled {
                    background-color: #31363b;
                    border-width: 1px;
                    border-color: #454545;
                    border-style: solid;
                    padding-top: 5px;
                    padding-bottom: 5px;
                    padding-left: 10px;
                    padding-right: 10px;
                    border-radius: 2px;
                    color: #454545;
                }
                QPushButton:pressed {
                    background-color: #3daee9;
                    padding-top: -15px;
                    padding-bottom: -17px;
                }
                QPushButton:hover {
                    border: 1px solid #ff8c00;
                    color: #eff0f1;
                }
                QLabel {
                    font-size: 18px;
                    border: 0px solid orange;
                }
            """)
            dialog.setWindowTitle("Select a Date")
            dialog.setText("Sélectionnez la date et produit souhaitée\n")
            self.date_edit = QtWidgets.QDateEdit()
            self.date = QtWidgets.QLineEdit(dialog)
            self.date.setInputMask("99-99-9999")
            self.date.resize(180, 30)
            self.date.move(40, 82)
            self.produitphytofiltre = QtWidgets.QComboBox(dialog)
            self.produitphytofiltre.setGeometry(40, 42, 180, 30)
            self.produitphytofiltre.addItem(" ")
            self.produitphytofiltre.addItem("Blé dur")
            self.produitphytofiltre.addItem("Blé tendre")
            self.produitphytofiltre.addItem("Orge")
            self.produitphytofiltre.addItem("Avoine")
            self.produitphytofiltre.addItem("Tritical")
            # self.date_edit.setDate(QDate.currentDate())
            ok_button = QtWidgets.QPushButton("OK", dialog)
            cancel_button = QtWidgets.QPushButton("Cancel", dialog)
            dialog.addButton(ok_button, QMessageBox.ButtonRole.ActionRole)
            dialog.addButton(cancel_button, QMessageBox.ButtonRole.ActionRole)
            dialog.exec()
            try:
                if dialog.clickedButton() == ok_button:
                    if self.produitphytofiltre.currentText() not in ['Blé dur', 'Blé tendre', 'Orge', 'Avoine',
                                                                     'Tritical'] and self.date.text() == '':
                        selected_date1 = self.date_edit.date().toString("yyyy")
                        selected_date = self.date.text()
                        conn = mysql.connector.connect(host="localhost", user=user, password=password,database="datta_refus")
                        cursor = conn.cursor()
                        cursor.execute(
                            "SELECT point_de_collecte,date,nom_et_prénom,fils_de,matricule,produit,cause_de_refus FROM refutable")
                        rows = cursor.fetchall()
                        ########################################################################
                        # Create a new document and add a table
                        self.doc = docx.Document()
                        section = self.doc.sections[0]
                        section.page_width = docx.shared.Cm(29.7)
                        section.page_height = docx.shared.Cm(21.0)
                        section.top_margin = docx.shared.Cm(1.5)
                        section.bottom_margin = docx.shared.Cm(1.5)
                        heading1 = self.doc.add_heading('\t\tLES CAS DE REFUS DES PRODUIT DE LA COMPANGNE ', level=1)
                        heading1.style.font.name = 'Times New Roman'
                        heading1.style.font.size = Pt(22)
                        heading1.style.font.bold = True
                        heading1.style.font.color.rgb = RGBColor(0, 0, 0)
                        # add second heading
                        paragraph = self.doc.add_paragraph()
                        paragraph.style.font.name = 'Times New Roman'
                        paragraph.style.font.size = Pt(14)
                        left_run = paragraph.add_run('CCLS RELIZANE\t\t\t\t\t\t\t\t\t\t\t\t\t')
                        left_run.bold = True
                        right_run = paragraph.add_run('Année:' + selected_date1)
                        right_run.bold = True

                        table = self.doc.add_table(rows=1, cols=7)
                        table.style = "Table Grid"  # set the table style
                        hdr_cells = table.rows[0].cells
                        hdr_cells[0].text = 'Point de collecte'
                        hdr_cells[1].text = 'Date'
                        hdr_cells[2].text = 'Nom et Prénom'
                        hdr_cells[3].text = 'fils_de'
                        hdr_cells[4].text = 'Matricule'
                        hdr_cells[5].text = 'Produit'
                        hdr_cells[6].text = 'Cause de refus'
                        # Set the width of the header cells
                        hdr_cells[0].width = Inches(2)  # set the width of the first column to 1.5 inches
                        hdr_cells[1].width = Inches(1)
                        hdr_cells[2].width = Inches(2)
                        hdr_cells[3].width = Inches(1)
                        hdr_cells[4].width = Inches(1)
                        hdr_cells[5].width = Inches(1)
                        hdr_cells[6].width = Inches(1.5)
                        # set hight of the column
                        hdr_cells[0].height = Inches(1)  # set the width of the first column to 1.5 inches
                        hdr_cells[1].height = Inches(1)
                        hdr_cells[2].height = Inches(1)
                        hdr_cells[3].height = Inches(1)
                        hdr_cells[4].height = Inches(1)
                        hdr_cells[5].height = Inches(1)
                        hdr_cells[6].height = Inches(1)

                        for cell in table.rows[0].cells:
                            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                            cell.paragraphs[0].runs[0].font.size = Pt(12)
                            cell.paragraphs[0].runs[0].font.bold = True
                        # Add the data to the table
                        table_rows = len(rows)
                        table_cols = len(rows[0])
                        for row in range(table_rows):
                            table.add_row()
                            for col in range(table_cols):
                                cell = table.cell(row + 1, col)
                                cell.text = str(rows[row][col])
                                # Set font properties
                                cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                                cell.paragraphs[0].runs[0].font.size = Pt(12)
                                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                        for row in table.rows:
                            row.height = Inches(0.5)
                        # Save and open the document for printing
                        temp_file = "les cas de refus" + selected_date1 + ".docx"
                        path, _ = QFileDialog.getSaveFileName(None, "Enregistrer la fiche", "",
                                                              "Fichiers DOCX (*.docx)")
                        if path:
                            self.doc.save(path)
                try:
                    if dialog.clickedButton() == ok_button:
                        selected_date = self.date.text()
                        if self.produitphytofiltre.currentText() not in ['Blé dur', 'Blé tendre', 'Orge', 'Avoine',
                                                                         'Tritical'] and self.date.text() == selected_date:
                            # selected_date = self.date_edit.date().toString("dd-MM-yyyy")
                            # locale.setlocale(locale.LC_ALL, 'fr_FR.utf8')
                            # QtCore.QLocale.setDefault(QtCore.QLocale(QtCore.QLocale.Language.French))
                            # selected_date1 = self.date_edit.date().toPyDate()
                            # mois = selected_date1.strftime("%B").lower()
                            print(selected_date)
                            # print(mois)
                            conn = mysql.connector.connect(host="localhost", user=user, password=password,database="datta_refus")
                            cursor = conn.cursor()
                            cursor.execute(
                                "SELECT point_de_collecte,date,nom_et_prénom,fils_de,matricule,produit,cause_de_refus FROM refutable WHERE date = %s",
                                (selected_date,))
                            rows = cursor.fetchall()
                            ########################################################################
                            print("Number of rows fetched:", len(rows))
                            print("Fetched rows:", rows)

                            # Create a new document and add a table
                            self.doc = docx.Document()
                            section = self.doc.sections[0]
                            section.orientation = WD_ORIENTATION.LANDSCAPE
                            section.page_width = docx.shared.Cm(29.7)
                            section.page_height = docx.shared.Cm(21.0)
                            section.top_margin = docx.shared.Cm(1.5)
                            section.bottom_margin = docx.shared.Cm(1.5)
                            heading1 = self.doc.add_heading('\t\tLES CAS DE REFUS DES PRODUIT DE LA COMPANGNE', level=1)
                            heading1.style.font.name = 'Times New Roman'
                            heading1.style.font.size = Pt(22)
                            heading1.style.font.bold = True
                            heading1.style.font.color.rgb = RGBColor(0, 0, 0)
                            # add second heading
                            paragraph = self.doc.add_paragraph()
                            paragraph.style.font.name = 'Times New Roman'
                            paragraph.style.font.size = Pt(14)
                            left_run = paragraph.add_run('CCLS RELIZANE\t\t\t\t\t\t\t\t\t\t\t\t')
                            left_run.bold = True
                            right_run = paragraph.add_run('Mois de:' + selected_date)
                            right_run.bold = True
                            table = self.doc.add_table(rows=1, cols=7)
                            table.style = "Table Grid"  # set the table style
                            hdr_cells = table.rows[0].cells
                            hdr_cells[0].text = 'Point de collecte'
                            hdr_cells[1].text = 'Date'
                            hdr_cells[2].text = 'Nom et Prénom'
                            hdr_cells[3].text = 'fils_de '
                            hdr_cells[4].text = 'Matricule'
                            hdr_cells[5].text = 'Produit'
                            hdr_cells[6].text = 'Cause de refus'
                            # Set the width of the header cells
                            hdr_cells[0].width = Inches(2)  # set the width of the first column to 1.5 inches
                            hdr_cells[1].width = Inches(1)
                            hdr_cells[2].width = Inches(2)
                            hdr_cells[3].width = Inches(1)
                            hdr_cells[4].width = Inches(1)
                            hdr_cells[5].width = Inches(1)
                            hdr_cells[6].width = Inches(1.5)
                            # set hight of the column
                            hdr_cells[0].height = Inches(1)  # set the width of the first column to 1.5 inches
                            hdr_cells[1].height = Inches(1)
                            hdr_cells[2].height = Inches(1)
                            hdr_cells[3].height = Inches(1)
                            hdr_cells[4].height = Inches(1)
                            hdr_cells[5].height = Inches(1)
                            hdr_cells[6].height = Inches(1)

                            for cell in table.rows[0].cells:
                                cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                                cell.paragraphs[0].runs[0].font.size = Pt(12)
                                cell.paragraphs[0].runs[0].font.bold = True
                            # Add the data to the table
                            table_rows = len(rows)
                            table_cols = len(rows[0])
                            for row in range(table_rows):
                                table.add_row()
                                for col in range(table_cols):
                                    cell = table.cell(row + 1, col)
                                    cell.text = str(rows[row][col])
                                    # Set font properties
                                    cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                                    cell.paragraphs[0].runs[0].font.size = Pt(12)
                                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                            for row in table.rows:
                                row.height = Inches(0.5)
                            # Save and open the document for printing
                            path, _ = QFileDialog.getSaveFileName(None, "Enregistrer la fiche", "",
                                                                  "Fichiers DOCX (*.docx)")
                            if path:
                                self.doc.save(path)
                            conn.close()
                except Exception as e:
                    QMessageBox.critical(MainWindow, "erreur", "erreur: Il n'y a pas des données à cette date")

                try:
                    if dialog.clickedButton() == ok_button:
                        if self.produitphytofiltre.currentText() in ['Blé dur', 'Blé tendre', 'Orge', 'Avoine',
                                                                     'Tritical'] and self.date.text() == "":
                            produitfiltre = self.produitphytofiltre.currentText()
                            print(produitfiltre)
                            conn = mysql.connector.connect(host="localhost", user=user, password=password,database="datta_refus")
                            cursor = conn.cursor()
                            cursor.execute(
                                "SELECT point_de_collecte,date,nom_et_prénom,fils_de,matricule,produit,cause_de_refus FROM refutable WHERE produit= %s ",
                                (produitfiltre,))
                            rows = cursor.fetchall()
                            ########################################################################
                            print("Number of rows fetched:", len(rows))
                            print("Fetched rows:", rows)
                            # Create a new document and add a table
                            self.doc = docx.Document()
                            section = self.doc.sections[0]
                            section.orientation = WD_ORIENTATION.LANDSCAPE
                            section.page_width = docx.shared.Cm(29.7)
                            section.page_height = docx.shared.Cm(21.0)
                            section.top_margin = docx.shared.Cm(1.5)
                            section.bottom_margin = docx.shared.Cm(1.5)
                            heading1 = self.doc.add_heading('\t\tLES CAS DE REFUS DES PRODUIT DE LA COMPANGNE', level=1)
                            heading1.style.font.name = 'Times New Roman'
                            heading1.style.font.size = Pt(22)
                            heading1.style.font.bold = True
                            heading1.style.font.color.rgb = RGBColor(0, 0, 0)
                            # add second heading
                            paragraph = self.doc.add_paragraph()
                            paragraph.style.font.name = 'Times New Roman'
                            paragraph.style.font.size = Pt(14)
                            left_run = paragraph.add_run('CCLS RELIZANE\t\t\t\t\t\t\t\t\t\t\t\t\t')
                            left_run.bold = True
                            right_run = paragraph.add_run('Espece:' + produitfiltre)
                            right_run.bold = True
                            table = self.doc.add_table(rows=1, cols=7)
                            table.style = "Table Grid"  # set the table style
                            hdr_cells = table.rows[0].cells
                            hdr_cells[0].text = 'Point de collecte'
                            hdr_cells[1].text = 'Date'
                            hdr_cells[2].text = 'Nom et Prénom'
                            hdr_cells[3].text = 'fils_de'
                            hdr_cells[4].text = 'Matricule'
                            hdr_cells[5].text = 'Produit'
                            hdr_cells[6].text = 'Cause de refus'
                            # Set the width of the header cells
                            hdr_cells[0].width = Inches(2)  # set the width of the first column to 1.5 inches
                            hdr_cells[1].width = Inches(1)
                            hdr_cells[2].width = Inches(2)
                            hdr_cells[3].width = Inches(1)
                            hdr_cells[4].width = Inches(1)
                            hdr_cells[5].width = Inches(1)
                            hdr_cells[6].width = Inches(1.5)
                            # set hight of the column
                            hdr_cells[0].height = Inches(1)  # set the width of the first column to 1.5 inches
                            hdr_cells[1].height = Inches(1)
                            hdr_cells[2].height = Inches(1)
                            hdr_cells[3].height = Inches(1)
                            hdr_cells[4].height = Inches(1)
                            hdr_cells[5].height = Inches(1)
                            hdr_cells[6].height = Inches(1)

                            for cell in table.rows[0].cells:
                                cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                                cell.paragraphs[0].runs[0].font.size = Pt(12)
                                cell.paragraphs[0].runs[0].font.bold = True
                            # Add the data to the table
                            table_rows = len(rows)
                            table_cols = len(rows[0])
                            for row in range(table_rows):
                                table.add_row()
                                for col in range(table_cols):
                                    cell = table.cell(row + 1, col)
                                    cell.text = str(rows[row][col])
                                    # Set font properties
                                    cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                                    cell.paragraphs[0].runs[0].font.size = Pt(12)
                                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                            for row in table.rows:
                                row.height = Inches(0.5)
                            # Save and open the document for printing
                            path, _ = QFileDialog.getSaveFileName(None, "Enregistrer la fiche", "",
                                                                  "Fichiers DOCX (*.docx)")
                            if path:
                                self.doc.save(path)
                            conn.close()
                except Exception as e:
                    msg = QMessageBox.critical(MainWindow, "Erreur", "Erreur")
            except Exception as e:
                print(e)
        def print_docx(self):
            dialog = QMessageBox()
            dialog.setStyleSheet("""
                       QWidget {
                           color: #eff0f1;
                           background-color: #31363b;
                           border-width: 1px;
                           border-color: #1e1e1e;
                           border-style: solid;
                           border-radius: 6;
                           padding: 3px;
                           font-size: 18px;
                           padding-left: 5px;
                           padding-right: 5px;
                       }
                       QWidget:item:hover {
                           background-color: #3daee9;
                           color: #eff0f1;
                       }
                       QWidget:item:selected {
                           background-color: #3daee9;
                       }
                       QWidget:disabled {
                           color: #454545;
                           background-color: #31363b;
                       }
                       QPushButton {
                           color: #b1b1b1;
                           background-color: linear-gradient(x1: 0, y1: 0, x2: 0, y2: 1, stop: 0 #565656, stop: 0.1 #525252, stop: 0.5 #4e4e4e, stop: 0.9 #4a4a4a, stop: 1 #464646);
                           border-width: 1px;
                           border-color: #1e1e1e;
                           border-style: solid;
                           border-radius: 6;
                           padding: 3px;
                           font-size: 12px;
                           padding-left: 5px;
                           padding-right: 5px;
                           min-width: 40px;
                       }
                       QPushButton:disabled {
                           background-color: #31363b;
                           border-width: 1px;
                           border-color: #454545;
                           border-style: solid;
                           padding-top: 5px;
                           padding-bottom: 5px;
                           padding-left: 10px;
                           padding-right: 10px;
                           border-radius: 2px;
                           color: #454545;
                       }
                       QPushButton:pressed {
                           background-color: #3daee9;
                           padding-top: -15px;
                           padding-bottom: -17px;
                       }
                       QPushButton:hover {
                           border: 1px solid #ff8c00;
                           color: #eff0f1;
                       }
                       QLabel {
                           font-size: 18px;
                           border: 0px solid orange;
                       }
                   """)
            dialog.setWindowTitle("Select a Date")
            dialog.setText("Sélectionnez la date et produit souhaitée\n")
            self.date_edit = QtWidgets.QDateEdit()
            self.date = QtWidgets.QLineEdit(dialog)
            self.date.setInputMask('99-99-9999')
            self.date.resize(180, 30)
            self.date.move(40, 82)
            self.produitphytofiltre = QtWidgets.QComboBox(dialog)
            self.produitphytofiltre.setGeometry(40, 42, 180, 30)
            self.produitphytofiltre.addItem(" ")
            self.produitphytofiltre.addItem("Blé dur")
            self.produitphytofiltre.addItem("Blé tendre")
            self.produitphytofiltre.addItem("Orge")
            self.produitphytofiltre.addItem("Avoine")
            self.produitphytofiltre.addItem("Tritical")
            # self.date_edit.setDate(QDate.currentDate())
            ok_button = QtWidgets.QPushButton("OK", dialog)
            cancel_button = QtWidgets.QPushButton("Cancel", dialog)
            dialog.addButton(ok_button, QMessageBox.ButtonRole.ActionRole)
            dialog.addButton(cancel_button, QMessageBox.ButtonRole.ActionRole)
            dialog.exec()
            try:
                if dialog.clickedButton() == ok_button:
                    if self.produitphytofiltre.currentText() not in ['Blé dur', 'Blé tendre', 'Orge', 'Avoine',
                                                                     'Tritical'] and self.date.text() == "--":
                        selected_date1 = self.date_edit.date().toString("yyyy")
                        selected_date = self.date.text()
                        conn = mysql.connector.connect(host="localhost", user=user, password=password,database="datta_refus")
                        cursor = conn.cursor()
                        cursor.execute(
                            "SELECT point_de_collecte,date,nom_et_prénom,fils_de,matricule,produit,cause_de_refus FROM refutable")
                        rows = cursor.fetchall()
                        ########################################################################
                        # Create a new document and add a table
                        self.doc = docx.Document()
                        section = self.doc.sections[0]
                        section.page_width = docx.shared.Cm(29.7)
                        section.page_height = docx.shared.Cm(21.0)
                        section.top_margin = docx.shared.Cm(1.5)
                        section.bottom_margin = docx.shared.Cm(1.5)
                        heading1 = self.doc.add_heading('\t\tLES CAS DE REFUS DES PRODUIT DE LA COMPANGNE ', level=1)
                        heading1.style.font.name = 'Times New Roman'
                        heading1.style.font.size = Pt(22)
                        heading1.style.font.bold = True
                        heading1.style.font.color.rgb = RGBColor(0, 0, 0)
                        # add second heading
                        paragraph = self.doc.add_paragraph()
                        paragraph.style.font.name = 'Times New Roman'
                        paragraph.style.font.size = Pt(14)
                        left_run = paragraph.add_run('CCLS RELIZANE\t\t\t\t\t\t\t\t\t\t\t\t\t')
                        left_run.bold = True
                        right_run = paragraph.add_run('Année:' + selected_date1)
                        right_run.bold = True

                        table = self.doc.add_table(rows=1, cols=7)
                        table.style = "Table Grid"  # set the table style
                        hdr_cells = table.rows[0].cells
                        hdr_cells[0].text = 'Point de collecte'
                        hdr_cells[1].text = 'Date'
                        hdr_cells[2].text = 'Nom et Prénom'
                        hdr_cells[3].text = 'fils_de'
                        hdr_cells[4].text = 'Matricule'
                        hdr_cells[5].text = 'Produit'
                        hdr_cells[6].text = 'Cause de refus'
                        # Set the width of the header cells
                        hdr_cells[0].width = Inches(2)  # set the width of the first column to 1.5 inches
                        hdr_cells[1].width = Inches(1)
                        hdr_cells[2].width = Inches(2)
                        hdr_cells[3].width = Inches(1)
                        hdr_cells[4].width = Inches(1)
                        hdr_cells[5].width = Inches(1)
                        hdr_cells[6].width = Inches(1.5)
                        # set hight of the column
                        hdr_cells[0].height = Inches(1)  # set the width of the first column to 1.5 inches
                        hdr_cells[1].height = Inches(1)
                        hdr_cells[2].height = Inches(1)
                        hdr_cells[3].height = Inches(1)
                        hdr_cells[4].height = Inches(1)
                        hdr_cells[5].height = Inches(1)
                        hdr_cells[6].height = Inches(1)

                        for cell in table.rows[0].cells:
                            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                            cell.paragraphs[0].runs[0].font.size = Pt(12)
                            cell.paragraphs[0].runs[0].font.bold = True
                        # Add the data to the table
                        table_rows = len(rows)
                        table_cols = len(rows[0])
                        for row in range(table_rows):
                            table.add_row()
                            for col in range(table_cols):
                                cell = table.cell(row + 1, col)
                                cell.text = str(rows[row][col])
                                # Set font properties
                                cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                                cell.paragraphs[0].runs[0].font.size = Pt(12)
                                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                        for row in table.rows:
                            row.height = Inches(0.5)
                        # Save and open the document for printing
                        temp_file = "les cas de refus" + selected_date + ".docx"
                        self.doc.save(temp_file)
                        os.startfile(temp_file, "print")
                try:
                    if dialog.clickedButton() == ok_button:
                        selected_date = self.date.text()
                        if self.produitphytofiltre.currentText() not in ['Blé dur', 'Blé tendre', 'Orge', 'Avoine',
                                                                         'Tritical'] and self.date.text() == selected_date:
                            # selected_date = self.date_edit.date().toString("dd-MM-yyyy")
                            # locale.setlocale(locale.LC_ALL, 'fr_FR.utf8')
                            # QtCore.QLocale.setDefault(QtCore.QLocale(QtCore.QLocale.Language.French))
                            # selected_date1 = self.date_edit.date().toPyDate()
                            # mois = selected_date1.strftime("%B").lower()

                            # print(mois)
                            conn = mysql.connector.connect(host="localhost", user=user, password=password,database="datta_refus")
                            cursor = conn.cursor()
                            cursor.execute(
                                "SELECT point_de_collecte,date,nom_et_prénom,fils_de,matricule,produit,cause_de_refus FROM refutable WHERE date = %s",
                                (selected_date,))
                            rows = cursor.fetchall()
                            ########################################################################
                            print("Number of rows fetched:", len(rows))
                            # Create a new document and add a table
                            self.doc = docx.Document()
                            section = self.doc.sections[0]
                            section.orientation = WD_ORIENTATION.LANDSCAPE
                            section.page_width = docx.shared.Cm(29.7)
                            section.page_height = docx.shared.Cm(21.0)
                            section.top_margin = docx.shared.Cm(1.5)
                            section.bottom_margin = docx.shared.Cm(1.5)
                            heading1 = self.doc.add_heading('\t\tLES CAS DE REFUS DES PRODUIT DE LA COMPANGNE', level=1)
                            heading1.style.font.name = 'Times New Roman'
                            heading1.style.font.size = Pt(22)
                            heading1.style.font.bold = True
                            heading1.style.font.color.rgb = RGBColor(0, 0, 0)
                            # add second heading
                            paragraph = self.doc.add_paragraph()
                            paragraph.style.font.name = 'Times New Roman'
                            paragraph.style.font.size = Pt(14)
                            left_run = paragraph.add_run('CCLS RELIZANE\t\t\t\t\t\t\t\t\t\t\t\t')
                            left_run.bold = True
                            right_run = paragraph.add_run('Le:' + selected_date)
                            right_run.bold = True
                            table = self.doc.add_table(rows=1, cols=7)
                            table.style = "Table Grid"  # set the table style
                            hdr_cells = table.rows[0].cells
                            hdr_cells[0].text = 'Point de collecte'
                            hdr_cells[1].text = 'Date'
                            hdr_cells[2].text = 'Nom et Prénom'
                            hdr_cells[3].text = 'fils_de'
                            hdr_cells[4].text = 'Matricule'
                            hdr_cells[5].text = 'Produit'
                            hdr_cells[6].text = 'Cause de refus'
                            # Set the width of the header cells
                            hdr_cells[0].width = Inches(2)  # set the width of the first column to 1.5 inches
                            hdr_cells[1].width = Inches(1)
                            hdr_cells[2].width = Inches(2)
                            hdr_cells[3].width = Inches(1)
                            hdr_cells[4].width = Inches(1)
                            hdr_cells[5].width = Inches(1)
                            hdr_cells[6].width = Inches(1.5)
                            # set hight of the column
                            hdr_cells[0].height = Inches(1)  # set the width of the first column to 1.5 inches
                            hdr_cells[1].height = Inches(1)
                            hdr_cells[2].height = Inches(1)
                            hdr_cells[3].height = Inches(1)
                            hdr_cells[4].height = Inches(1)
                            hdr_cells[5].height = Inches(1)
                            hdr_cells[6].height = Inches(1)
                            for cell in table.rows[0].cells:
                                cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                                cell.paragraphs[0].runs[0].font.size = Pt(12)
                                cell.paragraphs[0].runs[0].font.bold = True
                            # Add the data to the table
                            table_rows = len(rows)
                            table_cols = len(rows[0])
                            for row in range(table_rows):
                                table.add_row()
                                for col in range(table_cols):
                                    cell = table.cell(row + 1, col)
                                    cell.text = str(rows[row][col])
                                    # Set font properties
                                    cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                                    cell.paragraphs[0].runs[0].font.size = Pt(12)
                                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                            for row in table.rows:
                                row.height = Inches(0.5)
                            # Save and open the document for printing
                            # temp_file = "les cas de refus" + selected_date + ".docx"
                            temp_file = "les cas de refus" + selected_date + ".docx"
                            self.doc.save(temp_file)
                            os.startfile(temp_file, "print")
                            conn.close()
                except:
                    msg = QMessageBox.critical(MainWindow, 'Erreur', "Erreur: Il n'y a pas des données à cette date")
                try:
                    if dialog.clickedButton() == ok_button:
                        if self.produitphytofiltre.currentText() in ['Blé dur', 'Blé tendre', 'Orge', 'Avoine',
                                                                     'Tritical'] and self.date.text() == "":
                            produitfiltre = self.produitphytofiltre.currentText()
                            conn = mysql.connector.connect(host="localhost", user=user, password=password,database="datta_refus")
                            cursor = conn.cursor()
                            cursor.execute(
                                "SELECT point_de_collecte,date,nom_et_prénom,fils_de,matricule,produit,cause_de_refus FROM refutable WHERE produit= %s ",
                                (produitfiltre,))
                            rows = cursor.fetchall()
                            ########################################################################
                            print("Number of rows fetched:", len(rows))
                            print("Fetched rows:", rows)
                            # Create a new document and add a table
                            self.doc = docx.Document()
                            section = self.doc.sections[0]
                            section.orientation = WD_ORIENTATION.LANDSCAPE
                            section.page_width = docx.shared.Cm(29.7)
                            section.page_height = docx.shared.Cm(21.0)
                            section.top_margin = docx.shared.Cm(1.5)
                            section.bottom_margin = docx.shared.Cm(1.5)
                            heading1 = self.doc.add_heading('\t\tLES CAS DE REFUS DES PRODUIT DE LA COMPANGNE', level=1)
                            heading1.style.font.name = 'Times New Roman'
                            heading1.style.font.size = Pt(22)
                            heading1.style.font.bold = True
                            heading1.style.font.color.rgb = RGBColor(0, 0, 0)
                            # add second heading
                            paragraph = self.doc.add_paragraph()
                            paragraph.style.font.name = 'Times New Roman'
                            paragraph.style.font.size = Pt(14)
                            left_run = paragraph.add_run('CCLS RELIZANE\t\t\t\t\t\t\t\t\t\t\t\t\t')
                            left_run.bold = True
                            right_run = paragraph.add_run('Espece:' + produitfiltre)
                            right_run.bold = True
                            table = self.doc.add_table(rows=1, cols=7)
                            table.style = "Table Grid"  # set the table style
                            hdr_cells = table.rows[0].cells
                            hdr_cells[0].text = 'Point de collecte'
                            hdr_cells[1].text = 'Date'
                            hdr_cells[2].text = 'Nom et Prénom'
                            hdr_cells[3].text = 'fils_de'
                            hdr_cells[4].text = 'Matricule'
                            hdr_cells[5].text = 'Produit'
                            hdr_cells[6].text = 'Cause de refus'
                            # Set the width of the header cells
                            hdr_cells[0].width = Inches(2)  # set the width of the first column to 1.5 inches
                            hdr_cells[1].width = Inches(1)
                            hdr_cells[2].width = Inches(2)
                            hdr_cells[3].width = Inches(1)
                            hdr_cells[4].width = Inches(1)
                            hdr_cells[5].width = Inches(1)
                            hdr_cells[6].width = Inches(1.5)
                            # set hight of the column
                            hdr_cells[0].height = Inches(1)  # set the width of the first column to 1.5 inches
                            hdr_cells[1].height = Inches(1)
                            hdr_cells[2].height = Inches(1)
                            hdr_cells[3].height = Inches(1)
                            hdr_cells[4].height = Inches(1)
                            hdr_cells[5].height = Inches(1)
                            hdr_cells[6].height = Inches(1)

                            for cell in table.rows[0].cells:
                                cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                                cell.paragraphs[0].runs[0].font.size = Pt(12)
                                cell.paragraphs[0].runs[0].font.bold = True
                            # Add the data to the table
                            table_rows = len(rows)
                            table_cols = len(rows[0])
                            for row in range(table_rows):
                                table.add_row()
                                for col in range(table_cols):
                                    cell = table.cell(row + 1, col)
                                    cell.text = str(rows[row][col])
                                    # Set font properties
                                    cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                                    cell.paragraphs[0].runs[0].font.size = Pt(12)
                                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                            for row in table.rows:
                                row.height = Inches(0.5)
                            # Save and open the document for printing
                            path, _ = QFileDialog.getSaveFileName(None, "Enregistrer la fiche", "",
                                                                  "Fichiers DOCX (*.docx)")
                            if path:
                                self.doc.save(path)
                            conn.close()
                except Exception as e:
                    msg = QMessageBox.critical(MainWindow, "Erreur", "Erreur")


            except Exception as e:
                print(e)
        def impot_filter(self):
            nomefiltrage = self.nomefiltrage.text()
            oled = self.oled.text()
            print(nomefiltrage, oled)
            try:
                conn = mysql.connector.connect(host="localhost", user=user, password=password,database="datta_refus")
                cursor = conn.cursor()
                cursor.execute(
                    "SELECT id,point_de_collecte,date,nom_et_prénom,fils_de, matricule, produit,cause_de_refus FROM refutable WHERE nom_et_prénom =%s OR fils_de=%s",
                    (nomefiltrage, oled,))
                result = cursor.fetchall()
                self.textEdit.setRowCount(0)
                for row, row_datta in enumerate(result):
                    self.textEdit.insertRow(row)
                    for colum, datta in enumerate(row_datta):
                        self.textEdit.setItem(row, colum, QTableWidgetItem(str(datta)))
                conn.commit()
                conn.close()
            except sqlite3.Error as error:
                print(error)
        def add_datta_save(self):
            datte = self.dateedite.text()
            point_de_collecte = self.lieu.currentText()
            nome_et_prénom = self.nomeetprénom.text()
            matricule = self.matricule.text()
            produit = self.produit.currentText()
            causederefus = self.causerefus.toPlainText()
            fils_de = self.fissede.text()
            try:
                if nome_et_prénom:
                    dattabase = mysql.connector.connect(host="localhost",user=user,password=password,database="datta_refus")
                    cursor = dattabase.cursor()
                    cursor.execute(
                        "INSERT INTO refutable (point_de_collecte,date,nom_et_prénom,fils_de,matricule,produit,cause_de_refus) VALUES (%s,%s,%s,%s,%s,%s,%s)",
                        (point_de_collecte, datte, nome_et_prénom, fils_de, matricule, produit, causederefus))

                    dattabase.commit()
                    dattabase.close()
                    msgbox = QtWidgets.QMessageBox()
                    msgbox.setWindowTitle('confirmation')
                    msgbox.setText('Lopération a été ajoutée avec succès')
                    msgbox.exec()
                    self.impot_all()
                else:
                    msgbox = QtWidgets.QMessageBox()
                    msgbox.setWindowTitle('Erreur')
                    msgbox.setText("Erreur: entrer le nom et prénom")
                    msgbox.exec()


            except mysql.connector.Error as error:
                print("Failed to insert data into sqlite table", error)
        def impot_all(self):
            self.oled.clear()
            self.nomefiltrage.clear()
            try:
                dattabase = mysql.connector.connect(host="localhost", user=user, password=password,database="datta_refus")
                cur = dattabase.cursor()
                cur.execute("SELECT * FROM refutable ")
                result = cur.fetchall()
                self.textEdit.setRowCount(0)
                for row, row_datta in enumerate(result):
                    print(row)
                    self.textEdit.insertRow(row)
                    for colum, datta in enumerate(row_datta):
                        self.textEdit.setItem(row, colum, QTableWidgetItem(str(datta)))
                dattabase.commit()
                dattabase.close()
            except:
                print(' ')
        def delete_item(self):
            try:
                msgbox = QMessageBox()

                msgbox.setWindowTitle("Alerte")
                msgbox.setText("Voulez-vous supprimer !")
                yesbutton = QtWidgets.QPushButton("OUI")
                nobuttons = QtWidgets.QPushButton("NON")
                # msgbox.setStandardButtons(QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
                msgbox.addButton(yesbutton, QMessageBox.ButtonRole.ActionRole)
                msgbox.addButton(nobuttons, QMessageBox.ButtonRole.ActionRole)
                push = msgbox.exec()
                if msgbox.clickedButton() == nobuttons:
                    print("NO")
                elif msgbox.clickedButton() == yesbutton:
                    
                    curentrow = self.textEdit.currentRow()
                    id_ = self.textEdit.item(curentrow, 0).text()
                    database = mysql.connector.connect(host="localhost",user=user,password=password,database="datta_refus")
                    curs = database.cursor()
                    curs.execute("DELETE FROM refutable WHERE id=%s", (id_,))
                    database.commit()
                    database.close()
                    self.impot_all()
            except:
                print(' ')
    if __name__ == "__main__":
        import sys
        app = QtWidgets.QApplication(sys.argv)
        MainWindow = QtWidgets.QMainWindow()
        ui = Refus_Window()
        ui.refus_produit(MainWindow)
        MainWindow.show()
        sys.exit(app.exec())
except Exception as e:
    print(e)

