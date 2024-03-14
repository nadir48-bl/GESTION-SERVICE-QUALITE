import locale
import os
import pickle
import subprocess
import tempfile
import docx
import docxtpl
from PyQt6 import QtCore, QtGui, QtWidgets
from PyQt6.QtCore import QDate, Qt, QLocale, QTimer, QSize, QTime
from PyQt6.QtGui import QIcon, QIntValidator, QDoubleValidator
from PyQt6.QtPrintSupport import QPrinter
from PyQt6.QtWidgets import QMessageBox
from PyQt6.QtWidgets import *
import datetime
import sqlite3
from docx.enum.section import WD_ORIENTATION
from docx.shared import Inches, Pt, RGBColor
from docx2pdf import convert
from docxtpl import DocxTemplate
from PyQt6.QtPrintSupport import QPrinter, QPrintDialog
from win32com import client as win32a
import sys



database=sqlite3.connect("data_orge.db")
cur=database.cursor()
cur.execute("""create table if not exists orgetable
                       (
                        num_bulletin INTEGER PRIMARY KEY,
                        date TEXT,
                        nom_producteur TEXT,
                        pere TEXT,
                        matricule INTEGER,
                        point_collect TEXT,
                        nom_agreeur TEXT,
                        quantite INTEGER,
                        ps INTEGER,
                        total INTEGER,
                        path TEXT
                        )
                        """)
database.commit()
database.close()

database=sqlite3.connect("data_orge.db")
cur=database.cursor()
cur.execute("""create table if not exists orgetablebulletinsortie
                       (
                        num_bulletinor INTEGER PRIMARY KEY,
                        dateor TEXT,
                        nom_producteuror TEXT,
                        pereor TEXT,
                        quantiteor  INTEGER,
                        pathor TEXT,
                        date_sortieor,
                        ncartesor INTEGER
                        )
                        """)
database.commit()
database.close()


databasebd=sqlite3.connect("data_bd.db")
cur=databasebd.cursor()
cur.execute("""create table if not exists bdtable
                       (
                        num_bulletinbd INTEGER PRIMARY KEY,
                        datebd TEXT,
                        nom_producteurbd TEXT,
                        perebd TEXT,
                        matriculebd INTEGER,
                        point_collectbd TEXT,
                        nom_agreeurbd TEXT,
                        quantitebd INTEGER,
                        psbd INTEGER,
                        totalbd INTEGER,
                        pathbd TEXT,
                        date_sortie TEXT,
                        num_cart TEXT
                        )
                        """)
databasebd.commit()
databasebd.close()
databasebd=sqlite3.connect("data_bd.db")
cur=databasebd.cursor()
cur.execute("""create table if not exists bdtable_bulletin_sortie
                       (
                        num_bulletinbd INTEGER PRIMARY KEY,
                        datebd TEXT,
                        nom_producteurbd TEXT,
                        perebd TEXT,
                        quantite INTEGER,
                        pathbd TEXT,
                        date_sortie TEXT,
                        num_cart TEXT
                        )
                        """)
databasebd.commit()
databasebd.close()

databasebt=sqlite3.connect("data_bt.db")
cur=databasebt.cursor()
cur.execute("""create table if not exists bttable
                       (
                        num_bulletinbt INTEGER PRIMARY KEY,
                        datebt TEXT,
                        nom_producteurbt TEXT,
                        perebt TEXT,
                        matriculebt INTEGER,
                        point_collectbt TEXT,
                        nom_agreeurbt TEXT,
                        quantitebt INTEGER,
                        psbt INTEGER,
                        totalbt INTEGER,
                        pathbt TEXT,
                        date_sortie TEXT,
                        num_cart TEXT
                        )
                        """)
databasebt.commit()
databasebt.close()
databasebt=sqlite3.connect("data_bt.db")
cur=databasebt.cursor()
cur.execute("""create table if not exists bttable_bulletin_sortie
                       (
                        num_bulletinbts INTEGER PRIMARY KEY,
                        datebts TEXT,
                        nom_producteurbts TEXT,
                        perebts TEXT,
                        quantites INTEGER,
                        pathbts TEXT,
                        date_sortie TEXT,
                        num_carts TEXT
                        )
                        """)
databasebt.commit()
databasebt.close()

try:
    class Agréage_Window(object):
        ################################ORGE########################
        ############################################################
        def bonification_ps_orge(self):
            self.bpsor.clear()
            self.bpsor.setStyleSheet("background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding: 0px")
            while self.vpsor.value() >= 62.01:
                self.bpsor.setValue(0.24)
                self.bpsor.setStyleSheet(
                    "background-color:#88ffaa;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                if self.vpsor.value() <= 62.50:
                    self.bpsor.setValue(0.24)
                elif self.vpsor.value() >= 62.51 and self.vpsor.value() <= 63.00:
                    self.bpsor.setValue(0.48)
                elif self.vpsor.value() >= 63.01 and self.vpsor.value() <= 63.50:
                    self.bpsor.setValue(0.72)
                elif self.vpsor.value() >= 63.51 and self.vpsor.value() <= 64.00:
                    self.bpsor.setValue(0.96)
                elif self.vpsor.value() >= 64.01 and self.vpsor.value() <= 64.50:
                    self.bpsor.setValue(1.20)
                elif self.vpsor.value() >= 64.51 and self.vpsor.value() <= 65.00:
                    self.bpsor.setValue(1.44)
                elif self.vpsor.value() >= 65.01 and self.vpsor.value() <= 65.50:
                    self.bpsor.setValue(1.68)
                elif self.vpsor.value() >= 65.51 and self.vpsor.value() <= 66.00:
                    self.bpsor.setValue(1.92)
                elif self.vpsor.value() >= 66.01 and self.vpsor.value() <= 66.50:
                    self.bpsor.setValue(2.16)
                elif self.vpsor.value() >= 66.51 and self.vpsor.value() <= 67.00:
                    self.bpsor.setValue(2.40)
                elif self.vpsor.value() >= 67.01 and self.vpsor.value() <= 67.50:
                    self.bpsor.setValue(2.64)
                elif self.vpsor.value() >= 67.51 and self.vpsor.value() <= 68.00:
                    self.bpsor.setValue(2.88)
                elif self.vpsor.value() >= 68.01 and self.vpsor.value() <= 68.50:
                    self.bpsor.setValue(3.12)
                elif self.vpsor.value() >= 68.51 and self.vpsor.value() <= 69.00:
                    self.bpsor.setValue(3.36)
                elif self.vpsor.value() >= 69.01 and self.vpsor.value() <= 69.50:
                    self.bpsor.setValue(3.60)
                elif self.vpsor.value() >= 69.51 and self.vpsor.value() <= 70.00:
                    self.bpsor.setValue(3.84)
                elif self.vpsor.value() >= 70.01 and self.vpsor.value() <= 70.50:
                    self.bpsor.setValue(4.08)
                elif self.vpsor.value() >= 70.51 and self.vpsor.value() <= 71.00:
                    self.bpsor.setValue(4.32)
                elif self.vpsor.value() >= 61.01 and self.vpsor.value() <= 71.50:
                    self.bpsor.setValue(4.56)
                elif self.vpsor.value() >= 71.51 and self.vpsor.value() <= 72.00:
                    self.bpsor.setValue(4.70)
                else:
                    self.bpsor.setStyleSheet(
                        "background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                    self.bpsor.setValue(0)

                break

        def réfaction_ps_orge(self):
            self.rpsor.clear()
            self.rpsor.setStyleSheet("background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding: 0px")

            while self.vpsor.value() >= 52.00:
                self.rpsor.setValue(1.44)
                self.rpsor.setStyleSheet(
                    "background-color:#e22630;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                if self.vpsor.value() <= 52.49:
                    self.rpsor.setValue(1.44)
                elif self.vpsor.value() >= 52.50 and self.vpsor.value() <= 52.99:
                    self.rpsor.setValue(1.32)
                elif self.vpsor.value() >= 53.00 and self.vpsor.value() <= 53.49:
                    self.rpsor.setValue(1.20)
                elif self.vpsor.value() >= 53.50 and self.vpsor.value() <= 53.99:
                    self.rpsor.setValue(1.08)
                elif self.vpsor.value() >= 54.00 and self.vpsor.value() <= 54.49:
                    self.rpsor.setValue(0.96)
                elif self.vpsor.value() >= 54.50 and self.vpsor.value() <= 54.99:
                    self.rpsor.setValue(0.84)
                elif self.vpsor.value() >= 55.00 and self.vpsor.value() <= 55.49:
                    self.rpsor.setValue(0.72)
                elif self.vpsor.value() >= 55.50 and self.vpsor.value() <= 55.99:
                    self.rpsor.setValue(0.60)
                elif self.vpsor.value() >= 56.00 and self.vpsor.value() <= 56.49:
                    self.rpsor.setValue(0.48)
                elif self.vpsor.value() >= 56.50 and self.vpsor.value() <= 56.99:
                    self.rpsor.setValue(0.36)
                elif self.vpsor.value() >= 57.00 and self.vpsor.value() <= 57.49:
                    self.rpsor.setValue(0.24)
                elif self.vpsor.value() >= 57.50 and self.vpsor.value() <= 57.99:
                    self.rpsor.setValue(0.12)
                else:
                    self.rpsor.setStyleSheet(
                        "background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding:0px")
                    self.rpsor.setValue(0)

                break

        def refaction_impurté_orge(self):
            self.rtotalor.clear()
            self.rtotalor.setStyleSheet(
                "background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding: 0px")
            while self.vtotalor.value() >= 2.01:
                self.rtotalor.setValue(0.12)
                self.rtotalor.setStyleSheet(
                    "background-color:#e22630;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                if self.vtotalor.value() <= 2.50:
                    self.rtotalor.setValue(0.12)
                elif self.vtotalor.value() >= 2.51 and self.vtotalor.value() <= 3.00:
                    self.rtotalor.setValue(0.24)
                elif self.vtotalor.value() >= 3.01 and self.vtotalor.value() <= 3.50:
                    self.rtotalor.setValue(0.36)
                elif self.vtotalor.value() >= 3.51 and self.vtotalor.value() <= 4.00:
                    self.rtotalor.setValue(0.48)
                elif self.vtotalor.value() >= 4.01 and self.vtotalor.value() <= 4.50:
                    self.rtotalor.setValue(0.60)
                elif self.vtotalor.value() >= 4.51 and self.vtotalor.value() <= 5.00:
                    self.rtotalor.setValue(0.72)
                elif self.vtotalor.value() >= 5.01 and self.vtotalor.value() <= 5.50:
                    self.rtotalor.setValue(0.84)
                elif self.vtotalor.value() >= 5.51 and self.vtotalor.value() <= 6.00:
                    self.rtotalor.setValue(0.96)
                elif self.vtotalor.value() >= 6.01 and self.vtotalor.value() <= 6.50:
                    self.rtotalor.setValue(1.08)
                elif self.vtotalor.value() >= 6.51 and self.vtotalor.value() <= 7.00:
                    self.rtotalor.setValue(1.20)
                elif self.vtotalor.value() >= 7.01 and self.vtotalor.value() <= 7.50:
                    self.rtotalor.setValue(1.32)
                elif self.vtotalor.value() >= 7.51 and self.vtotalor.value() <= 8.00:
                    self.rtotalor.setValue(1.44)
                elif self.vtotalor.value() >= 8.01 and self.vtotalor.value() <= 8.50:
                    self.rtotalor.setValue(1.56)
                elif self.vtotalor.value() >= 8.51 and self.vtotalor.value() <= 9.00:
                    self.rtotalor.setValue(1.68)
                elif self.vtotalor.value() >= 9.01 and self.vtotalor.value() <= 9.50:
                    self.rtotalor.setValue(1.80)
                elif self.vtotalor.value() >= 9.51 and self.vtotalor.value() <= 10.00:
                    self.rtotalor.setValue(1.92)
                elif self.vtotalor.value() >= 10.01 and self.vtotalor.value() <= 20.00:
                    self.rtotalor.setValue(0)
                    self.rtotalor.setSuffix("P/D")
                    self.observationor.setText("PRIX A DEBATTRE")
                else:
                    self.rtotalor.setStyleSheet(
                        "background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                    self.rtotalor.setValue(0)
                break

        def plus_orge(self):
            a = self.bpsor.value()
            b = self.btotalor.value()
            c = self.rpsor.value()
            d = self.rtotalor.value()
            e = self.vGrainsanvaleuror.value()
            f = self.vmatierinertor.value()
            self.vtotalor.setValue(e + f)
            self.bglobaltotalor.setValue(a + b)
            self.rglobaltotalor.setValue(c + d)
            if self.bpsor.text() == "":
                self.bglobaltotalor.clear()

        def clear_orge(self):
            self.vpsor.clear()
            self.vpsor.setValue(0)
            self.vergotor.clear()
            self.vGrainsanvaleuror.clear()
            self.vGrainsanvaleuror.setValue(0)
            self.vmatierinertor.clear()
            self.vmatierinertor.setValue(0)
            self.vtotalor.clear()
            self.vtotalor.setValue(0)
            self.bpsor.clear()
            self.bpsor.setValue(0)
            self.btotalor.clear()
            self.btotalor.setValue(0)
            self.bglobaltotalor.clear()
            self.bglobaltotalor.setValue(0)
            self.rpsor.clear()
            self.rpsor.setValue(0)
            self.rtotalor.clear()
            self.rtotalor.setValue(0)
            self.rtotalor.setSuffix("")
            self.rglobaltotalor.clear()
            self.rglobaltotalor.setValue(0)
            self.observationor.clear()

        def calcul_orge(self):
            self.réfaction_ps_orge()
            self.bonification_ps_orge()
            self.refaction_impurté_orge()
            self.plus_orge()

        def bulletin_orge_print(self):
            try:
                nbulltin = self.n_bultin.text()
                nomproducteur = self.nome_du_producteur.text()
                pere = self.pére.text()
                matricul = self.adresse.text()
                poindecollecte = self.pointdecollect.currentText()
                dateorge = self.dattereceptiont.text()
                agreeur = self.agréeeurcomboorge.currentText()
                quantite = self.quantiteorge.value()
                vpsorge = self.vpsor.value()
                bpsorge = self.bpsor.value()
                rpsorge = self.rpsor.value()
                vergot = self.vergotor.value()
                grainsanvaleur = self.vGrainsanvaleuror.value()
                matiéreinerte = self.vmatierinertor.value()
                vtotal = self.vtotalor.value()
                btotalorge = self.btotalor.value()
                rtotalorge = self.rtotalor.value()
                ndcarteorge = self.n_carte.text()
                observation = self.observationor.toPlainText()
                date = QDate.currentDate()
                current_date = QDate.currentDate()
                current_day = current_date.day()
                current_month = current_date.month()
                current_year = current_date.year()
                current_time = QTime.currentTime()
                current_hour = current_time.hour()
                current_minute = current_time.minute()
                # Format the date as "dd-mm-yyyy"
                self.current_days = f"{current_day:02d}-{current_month:02d}-{current_year}-{current_hour}-{current_minute}"
                self.docorge = DocxTemplate("bulletin agreage/bulletin_orge/bulletin_orge.docx")
                context = {
                    "nmp": nomproducteur,
                    "per": pere,
                    "mat": matricul,
                    "ptt": poindecollecte,
                    "dt": dateorge,
                    "agr": agreeur,
                    "qtt": quantite,
                    'vps': vpsorge,
                    'bps': bpsorge,
                    'rps': rpsorge,
                    'vtt': vtotal,
                    'btt': btotalorge,
                    'rtt': rtotalorge,
                    'ncn': ndcarteorge,
                    "vrg": vergot,
                    "vmt": matiéreinerte,
                    "vgs": grainsanvaleur,
                    "oo": observation,
                    "nm": nbulltin
                }
                self.docorge.render(context)
                data = sqlite3.connect("data_orge.db")
                cursbd = data.cursor()
                cursbd.execute(
                    "SELECT * FROM orgetablebulletinsortie WHERE num_bulletinor AND dateor=? AND nom_producteuror = ? AND pereor = ?  AND quantiteor=?",
                    (dateorge, nomproducteur, pere, quantite,))
                existing_data = cursbd.fetchone()
                if existing_data:
                    self.msgbox = QtWidgets.QDialog()
                    self.msgbox.setWindowTitle('Confirmation de sortie')
                    self.msgbox.setGeometry(550, 200, 305, 100)
                    self.msgbox.setStyleSheet(""" QWidget
                            {
                                color: #000000;
                                background-color: #ffffff;
                                border-width: 1px;
                                border-color: #1e1e1e;
                                border-style: solid;
                                border-radius: 6;
                                padding: 0px;
                                font-size: 18px;
                                padding-left: 1px;
                                padding-right: 1px
                            }
                            QWidget:item:hover
                            {
                                background-color: #3daee9;
                                color: #eff0f1;
                            }
                            QWidget:item:selected
                            {
                                background-color: #3daee9;
                            }
                            QWidget:disabled
                            {
                                color: #454545;
                                background-color: #31363b;
                            }
                            QPushButton
                            {
                                color: #000000;
                                background-color:#ade3e7;
                                border-width: 1px;
                                border-color: #1e1e1e;
                                border-style: solid;
                                border-radius: 6;
                                padding: 3px;
                                font-size: 12px;
                                padding-left: 5px;
                                padding-right: 5px;
                                min-width: 40px
                            }
                            QPushButton:disabled
                            {
                                background-color: #31363b;
                                border-width: 1px;
                                border-color: #454545;
                                border-style: solid;
                                padding-top: 5px;
                                padding-bottom: 5px;
                                padding-left: 10px;
                                padding-right: 10px;
                                border-radius: 2px;
                                color: #454545;
                            }

                            QPushButton:pressed
                            {
                                background-color: #3daee9;
                                padding-top: -15px;
                                padding-bottom: -17px;
                            }
                            QPushButton:hover
                            {
                                border: 1px solid #ff8c00;
                                color: #000000;
                            }
                             QLabel
                            {
                                font-size: 18px;
                                border: 0px solid orange;
                            }

                        """)
                    self.oui_button = QtWidgets.QPushButton("Oui", self.msgbox,
                                                            clicked=lambda: self.print_other_docx_orge())
                    self.oui_button.setGeometry(90, 60, 60, 30)
                    self.non_button = QtWidgets.QPushButton("Non", self.msgbox, clicked=lambda: self.msgbox.close())
                    self.non_button.setGeometry(160, 60, 60, 30)
                    labelmgboxorge = QtWidgets.QLabel("Veuillez , imprimer autre fichier ?", self.msgbox)
                    labelmgboxorge.setGeometry(25, 15, 260, 30)
                    self.msgbox.exec()
                else:
                    doc_names = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                    doc_pdf = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False).name
                    self.docorge.save(doc_names)
                    try:
                        if doc_names:
                            a = self.progress_bar()
                            sys.stderr = open("consoleoutput.log", "w")
                            convert(doc_names, doc_pdf)
                            # Open the resulting .pdf file using the default associated application
                            # os.startfile(doc_pdf, 'open')
                            app_path = 'C:\\Program Files\\Okular\\bin\\okular.exe'
                            subprocess.Popen([app_path, doc_pdf])
                    except Exception as e:
                        print(e)

                    nameorge = nomproducteur + '-' + self.current_days + '-' + str(quantite) + ".docx"
                    self.docorge.save("bulletin agreage/bulletin_orge/" + nameorge)
                    pathorge = os.path.abspath("bulletin agreage/bulletin_orge/" + nameorge)
                    databasebd = sqlite3.connect("data_orge.db")
                    curs = databasebd.cursor()
                    curs.execute(
                        "INSERT INTO orgetablebulletinsortie (dateor, nom_producteuror, pereor,quantiteor, pathor,date_sortieor) VALUES (?,?,?, ?, ?, ?)",
                        (dateorge, nomproducteur, pere, quantite, pathorge, current_day,))
                    databasebd.commit()
                    databasebd.close()
            except Exception as e:
                print(e)

        def print_other_docx_orge(self):
            self.msgbox.close()
            quantite = self.quantiteorge.value()
            nomproducteur = self.nome_du_producteur.text()
            doc_names = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
            doc_pdf = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False).name
            self.docorge.save(doc_names)
            try:
                if doc_names:
                    a = self.progress_bar()
                    sys.stderr = open("consoleoutput.log", "w")
                    convert(doc_names, doc_pdf)
                    # Open the resulting .pdf file using the default associated application
                    # os.startfile(doc_pdf, 'open')
                    app_path = 'C:\\Program Files\\Okular\\bin\\okular.exe'
                    subprocess.Popen([app_path, doc_pdf])
            except Exception as e:
                print(e)
            nameorge = nomproducteur + '-' + self.current_days + '-' + str(quantite) + ".docx"
            self.docorge.save("bulletin agreage/bulletin_orge/" + nameorge)
            pathorge = os.path.abspath("bulletin agreage/bulletin_orge/" + nameorge)

        def add_datta_orge(self):
            try:
                nbulltin = self.n_bultin.text()
                nomproducteur = self.nome_du_producteur.text()
                pere = self.pére.text()
                matricul = self.adresse.text()
                poindecollecte = self.pointdecollect.currentText()
                dateorge = self.dattereceptiont.text()
                agreeur = self.agréeeurcomboorge.currentText()
                quantite = self.quantiteorge.value()
                vpsorge = self.vpsor.value()
                bpsorge = self.bpsor.value()
                rpsorge = self.rpsor.value()
                vergot = self.vergotor.value()
                grainsanvaleur = self.vGrainsanvaleuror.value()
                matiéreinerte = self.vmatierinertor.value()
                vtotal = self.vtotalor.value()
                btotalorge = self.btotalor.value()
                rtotalorge = self.rtotalor.value()
                ndcarteorge = self.n_carte.text()
                observation = self.observationor.toPlainText()
                date = QDate.currentDate()
                current_date = QDate.currentDate()
                current_day = current_date.day()
                current_month = current_date.month()
                current_year = current_date.year()
                current_time = QTime.currentTime()
                current_hour = current_time.hour()
                current_minute = current_time.minute()
                # Format the date as "dd-mm-yyyy"
                current_days = f"{current_day:02d}-{current_month:02d}-{current_year}-{current_hour}-{current_minute}"
                docorge = DocxTemplate("bulletin agreage/bulletin_orge/bulletin_orge.docx")
                context = {
                    "nmp": nomproducteur,
                    "per": pere,
                    "mat": matricul,
                    "ptt": poindecollecte,
                    "dt": dateorge,
                    "agr": agreeur,
                    "qtt": quantite,
                    'vps': vpsorge,
                    'bps': bpsorge,
                    'rps': rpsorge,
                    'vtt': vtotal,
                    'btt': btotalorge,
                    'rtt': rtotalorge,
                    'ncn': ndcarteorge,
                    "vrg": vergot,
                    "vmt": matiéreinerte,
                    "vgs": grainsanvaleur,
                    "oo": observation,
                    "nm": nbulltin
                }
                docorge.render(context)
                nameorge = nomproducteur + '-' + current_days + '-' + str(quantite) + ".docx"
                docorge.save("bulletin agreage/bulletin_orge/" + nameorge)
                pathorge = os.path.abspath("bulletin agreage/bulletin_orge/" + nameorge)
                database = sqlite3.connect('data_orge.db')
                curs = database.cursor()

                # Check if the data already exists
                curs.execute(
                    "SELECT * FROM orgetable WHERE date=? AND nom_producteur=? AND pere=? AND matricule=? AND point_collect=? AND nom_agreeur=? AND quantite=? AND ps=? AND total=? AND path=?",
                    (dateorge, nomproducteur, pere, matricul, poindecollecte, agreeur, quantite, vpsorge,
                     vtotal,
                     pathorge,))
                existing_data = curs.fetchone()

                if existing_data:
                    msgbox = QtWidgets.QMessageBox()
                    msgbox.setWindowTitle('confirmation')
                    msgbox.setText('Les données existent déjà.')
                    msgbox.exec()

                else:
                    curs.execute(
                        "INSERT INTO orgetable (date ,nom_producteur ,pere ,matricule ,point_collect ,nom_agreeur ,quantite ,ps ,total,path )  values(?,?,?,?,?,?,?,?,?,?)",
                        (dateorge, nomproducteur, pere, matricul, poindecollecte, agreeur, quantite, vpsorge, vtotal,
                         pathorge,))
                    database.commit()
                    database.close()
                    msgbox = QtWidgets.QMessageBox()
                    msgbox.setWindowTitle('confirmation')
                    msgbox.setText('Les données sont sauvegardées')
                    msgbox.exec()
            except sqlite3.Error as e:
                print("Error accessing the database:", e)

        def number_bulletin_orge(self):
            try:
                database = sqlite3.connect("data_orge.db")
                cursor = database.cursor()
                cursor.execute("SELECT num_bulletin FROM orgetable ORDER BY num_bulletin DESC LIMIT 1")
                result = cursor.fetchone()
                if result is not None:
                    last_id = result[0]
                    self.n_bultin.setText("DC-" + str(last_id))
                else:
                    # Handle the case when there are no records in the orgetable
                    self.n_bultin.setText("")
            except sqlite3.Error as e:
                print("Error accessing the database:", e)

        def print_day(self):
            try:
                dialog = QtWidgets.QMessageBox()
                dialog.setStyleSheet("""
                                          QWidget {
                                              color: #000000;
                                              background-color: #feffbd;
                                              border-width: 1px;
                                              border-color: #1e1e1e;
                                              border-style: solid;
                                              border-radius: 6;
                                              padding: 3px;
                                              font-size: 18px;
                                              padding-left: 5px;
                                              padding-right: 5px;
                                          }
                                          QWidget:item:hover {
                                              background-color: #3daee9;
                                              color: #eff0f1;
                                          }
                                          QWidget:item:selected {
                                              background-color: #3daee9;
                                          }
                                          QWidget:disabled {
                                              color: #454545;
                                              background-color: #31363b;
                                          }
                                          QPushButton {
                                              color: #000000;
                                              background-color:#84dbc8;
                                              border-width: 1px;
                                              border-color: #1e1e1e;
                                              border-style: solid;
                                              border-radius: 6;
                                              padding: 3px;
                                              font-size: 12px;
                                              padding-left: 5px;
                                              padding-right: 5px;
                                              min-width: 40px;
                                          }
                                          QPushButton:disabled {
                                              background-color: #31363b;
                                              border-width: 1px;
                                              border-color: #454545;
                                              border-style: solid;
                                              padding-top: 5px;
                                              padding-bottom: 5px;
                                              padding-left: 10px;
                                              padding-right: 10px;
                                              border-radius: 2px;
                                              color: #454545;
                                          }
                                          QPushButton:pressed {
                                              background-color: #3daee9;
                                              padding-top: -15px;
                                              padding-bottom: -17px;
                                          }
                                          QPushButton:hover {
                                              border: 1px solid #ff8c00;
                                              color: #eff0f1;
                                          }
                                          QLabel {
                                              font-size: 18px;
                                              border: 0px solid orange;
                                          }
                                      """)
                dialog.setWindowTitle("Sélectionnez la date")
                dialog.setText("Sélectionnez la date  \t\n")
                self.date_edit = QtWidgets.QDateEdit()
                self.dateor = QtWidgets.QDateEdit(dialog)
                self.dateor.setDisplayFormat("dd-MM-yyyy")
                self.dateor.setDate(self.datedaytime)
                self.dateor.resize(180, 30)
                self.dateor.move(40, 50)

                self.date_edit.setDate(QDate.currentDate())
                ok_button = QtWidgets.QPushButton("OK", dialog)
                cancel_button = QtWidgets.QPushButton("Cancel", dialog)
                dialog.addButton(ok_button, QMessageBox.ButtonRole.ActionRole)
                dialog.addButton(cancel_button, QMessageBox.ButtonRole.ActionRole)
                dialog.exec()

                if dialog.clickedButton() == ok_button:
                    self.selected_date = self.dateor.text()
                    if self.dateor.text() == self.selected_date:
                        # selected_date = self.date_edit.date().toString("dd-MM-yyyy")
                        # locale.setlocale(locale.LC_ALL, 'fr_FR.utf8')
                        # QtCore.QLocale.setDefault(QtCore.QLocale(QtCore.QLocale.Language.French))
                        # selected_date1 = self.date_edit.date().toPyDate()
                        # mois = selected_date1.strftime("%B").lower()


                        selected_date1 = self.date_edit.date().toString("dd-MM-yyyy")
                        selected_date = self.dateor.text()

                        conn = sqlite3.connect('data_orge.db')
                        cursor = conn.cursor()
                        cursor.execute(
                            "SELECT num_bulletin, date, nom_producteur, pere, matricule, point_collect, nom_agreeur, quantite FROM orgetable WHERE DATE(substr(date, 7, 4) || '-' || substr(date, 4, 2) || '-' || substr(date, 1, 2)) =?;",
                            (selected_date,))
                        rows = cursor.fetchall()
                        ########################################################################
                        # Create a new document and add a table
                        self.doc = docx.Document()
                        section = self.doc.sections[0]
                        section.page_width = docx.shared.Cm(29.7)
                        section.page_height = docx.shared.Cm(21.0)
                        section.top_margin = docx.shared.Cm(1.5)
                        section.bottom_margin = docx.shared.Cm(1.5)
                        heading1 = self.doc.add_heading("\t\t\t\t\t\t les entre de l'orge ", level=1)
                        heading1.style.font.name = 'Times New Roman'
                        heading1.style.font.size = Pt(22)
                        heading1.style.font.bold = True
                        heading1.style.font.color.rgb = RGBColor(0, 0, 0)
                        # add second heading
                        paragraph = self.doc.add_paragraph()
                        paragraph.style.font.name = 'Times New Roman'
                        paragraph.style.font.size = Pt(14)
                        left_run = paragraph.add_run('CCLS RELIZANE\t\t\t\t\t\t\t\t\t\t\t\t')
                        left_run.bold = True
                        right_run = paragraph.add_run('Date:' + self.selected_date)
                        right_run.bold = True

                        table = self.doc.add_table(rows=1, cols=9)
                        table.style = "Table Grid"  # set the table style
                        hdr_cells = table.rows[0].cells
                        hdr_cells[0].text = 'N° Bulletin \nDC'
                        hdr_cells[1].text = 'Date'
                        hdr_cells[2].text = 'Nom et Prénom'
                        hdr_cells[3].text = 'Pére'
                        hdr_cells[4].text = 'Matricule'
                        hdr_cells[5].text = 'point de collect'
                        hdr_cells[6].text = 'Quantité'
                        hdr_cells[7].text = 'Date de sortie '
                        hdr_cells[8].text = 'N° de la carte '

                        # Set the width of the header cells
                        hdr_cells[0].width = Inches(1.5)  # set the width of the first column to 1.5 inches
                        hdr_cells[1].width = Inches(1.5)
                        hdr_cells[2].width = Inches(2)
                        hdr_cells[3].width = Inches(1)
                        hdr_cells[4].width = Inches(1.5)
                        hdr_cells[5].width = Inches(1.5)
                        hdr_cells[6].width = Inches(1)
                        hdr_cells[7].width = Inches(1.5)
                        hdr_cells[8].width = Inches(1.5)

                        # set hight of the column
                        hdr_cells[0].height = Inches(1)  # set the width of the first column to 1.5 inches
                        hdr_cells[1].height = Inches(1)
                        hdr_cells[2].height = Inches(1)
                        hdr_cells[3].height = Inches(1)
                        hdr_cells[4].height = Inches(1)
                        hdr_cells[5].height = Inches(1)
                        hdr_cells[6].height = Inches(1)
                        hdr_cells[7].width = Inches(1)
                        hdr_cells[8].width = Inches(1)

                        for cell in table.rows[0].cells:
                            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                            cell.paragraphs[0].runs[0].font.size = Pt(12)
                            cell.paragraphs[0].runs[0].font.bold = True
                        # Add the data to the table
                        table_rows = len(rows)
                        table_cols = len(rows[0])
                        for row in range(table_rows):
                            table.add_row()
                            for col in range(table_cols):
                                cell = table.cell(row + 1, col)
                                cell.text = str(rows[row][col])
                                # Set font properties
                                cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                                cell.paragraphs[0].runs[0].font.size = Pt(12)
                                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                        for row in table.rows:
                            row.height = Inches(0.3)
                        # Save and open the document for printing
                        doc_names = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                        doc_pdf = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False).name
                        self.doc.save(doc_names)
                        try:
                            if doc_names:
                                a = self.progress_bar()
                                sys.stderr = open("consoleoutput.log", "w")
                                convert(doc_names, doc_pdf)
                                # Open the resulting .pdf file using the default associated application
                                # os.startfile(doc_pdf, 'open')
                                app_path = 'C:\\Program Files\\Okular\\bin\\okular.exe'
                                subprocess.Popen([app_path, doc_pdf])
                        except Exception as e:
                            print(e)
                    else:
                        msgbox = QtWidgets.QMessageBox()
                        msgbox.setWindowTitle('Erreur')
                        msgbox.setText("Il n'ya pas des données à cette date.")
                        msgbox.exec()
                if dialog.clickedButton() == cancel_button:
                    dialog.close()
            except Exception as e:
                print(e)

        def printBulltinProducteurOrge(self):
            dialogor = QMessageBox()
            dialogor.setFixedSize(300, 600)
            dialogor.setStyleSheet("""
                                          QWidget {
                                              color: #000000;
                                              background-color: #feffbd;
                                              border-width: 1px;
                                              border-color: #1e1e1e;
                                              border-style: solid;
                                              border-radius: 6;
                                              padding: 3px;
                                              font-size: 18px;
                                              padding-left: 5px;
                                              padding-right: 5px;
                                          }
                                          QWidget:item:hover {
                                              background-color: #3daee9;
                                              color: #eff0f1;
                                          }
                                          QWidget:item:selected {
                                              background-color: #3daee9;
                                          }
                                          QWidget:disabled {
                                              color: #454545;
                                              background-color: #31363b;
                                          }
                                          QPushButton {
                                              color: #000000;
                                              background-color:#84dbc8;
                                              border-width: 1px;
                                              border-color: #1e1e1e;
                                              border-style: solid;
                                              border-radius: 6;
                                              padding: 3px;
                                              font-size: 12px;
                                              padding-left: 5px;
                                              padding-right: 5px;
                                              min-width: 40px;
                                          }
                                          QPushButton:disabled {
                                              background-color: #31363b;
                                              border-width: 1px;
                                              border-color: #454545;
                                              border-style: solid;
                                              padding-top: 5px;
                                              padding-bottom: 5px;
                                              padding-left: 10px;
                                              padding-right: 10px;
                                              border-radius: 2px;
                                              color: #454545;
                                          }
                                          QPushButton:pressed {
                                              background-color: #3daee9;
                                              padding-top: -15px;
                                              padding-bottom: -17px;
                                          }
                                          QPushButton:hover {
                                              border: 1px solid #ff8c00;
                                              color: #eff0f1;
                                          }
                                          QLabel {
                                              font-size: 18px;
                                              border: 0px solid orange;
                                          }
                                      """)
            dialogor.setWindowTitle("Sélectionnez le Nom et Prénom et pére")
            dialogor.setText("Sélectionnez le Nom et Prénom et pére  \t\t\t\t\t\t\n\t\t\t\t\t\t\t\t\t\t\n\t\t\t\t")
            self.nomproductureor = QtWidgets.QLineEdit(dialogor)
            self.nomproductureor.setPlaceholderText("Nom et Prénom")
            self.nomproductureor.resize(180, 30)
            self.nomproductureor.move(40, 40)
            self.péreor = QtWidgets.QLineEdit(dialogor)
            self.péreor.setPlaceholderText("Pére")
            self.péreor.resize(180, 30)
            self.péreor.move(240, 40)
            self.datetxt = QtWidgets.QLabel("Date entré", dialogor)
            self.datetxt.setGeometry(QtCore.QRect(30, 80, 100, 20))
            self.dateor = QDateTimeEdit(dialogor)
            self.dateor.setDisplayFormat("dd-MM-yyyy hh:mm")
            self.dateor.resize(180, 30)
            self.dateor.move(40, 105)
            self.quantiteptext = QtWidgets.QLabel("Quantité", dialogor)
            self.quantiteptext.setGeometry(QtCore.QRect(230, 80, 100, 20))
            self.quantiteor = QtWidgets.QDoubleSpinBox(dialogor)
            self.quantiteor.setRange(1, 10000)
            self.quantiteor.setSuffix('  QX')
            self.quantiteor.setSpecialValueText(" ")
            self.quantiteor.setValue(0.0)
            self.quantiteor.resize(180, 30)
            self.quantiteor.move(240, 105)
            ok_button = QtWidgets.QPushButton("OK", dialogor)
            cancel_button = QtWidgets.QPushButton("Cancel", dialogor)
            dialogor.addButton(ok_button, QMessageBox.ButtonRole.ActionRole)
            dialogor.addButton(cancel_button, QMessageBox.ButtonRole.ActionRole)
            dialogor.exec()
            date = QDate.currentDate()
            current_date = QDate.currentDate()
            current_day = current_date.day()
            current_month = current_date.month()
            current_year = current_date.year()

            # Format the date as "dd-mm-yyyy"
            current_day = f"{current_day:02d}-{current_month:02d}-{current_year}"

            if dialogor.clickedButton() == ok_button:
                producteur = self.nomproductureor.text()
                pere = self.péreor.text()
                dateentré = self.dateor.text()
                quantite = self.quantiteor.value()

                try:
                    data = sqlite3.connect("data_orge.db")
                    cursbd = data.cursor()
                    cursbd.execute(
                        "SELECT * FROM orgetablebulletinsortie WHERE num_bulletinor AND dateor=? AND nom_producteuror = ? AND pereor = ?  AND quantiteor=?",
                        (dateentré, producteur, pere, quantite))
                    existing_data = cursbd.fetchone()
                    if existing_data:
                        msgbox = QtWidgets.QMessageBox()
                        msgbox.setWindowTitle('Confirmation')
                        msgbox.setText('Le Bulletin sortie  déjà.')
                        msgbox.exec()
                    else:
                        curs = data.cursor()
                        curs.execute(
                            "SELECT path FROM orgetable WHERE date=? AND nom_producteur=? AND pere=? AND quantite=?",
                            (dateentré, producteur, pere, quantite,))
                        result = curs.fetchall()
                        if result:
                            # Iterate over the paths and open each file
                            for row in result:
                                file_path = row[0]
                                if file_path:
                                    doc_pdf = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False).name
                                    try:
                                        a = self.progress_bar()
                                        sys.stderr = open("consoleoutput.log", "w")
                                        convert(file_path, doc_pdf)
                                        # Open the resulting .pdf file using the default associated application
                                        # os.startfile(doc_pdf, 'open')
                                        app_path = 'C:\\Program Files\\Okular\\bin\\okular.exe'
                                        subprocess.Popen([app_path, doc_pdf])
                                    except Exception as e:
                                        print(e)
                                    # os.startfile(file_path, "open")
                                    curs.execute(
                                        "INSERT INTO orgetablebulletinsortie (dateor, nom_producteuror, pereor, quantiteor, pathor, date_sortieor) VALUES (?,?,?,?,?,?)",
                                        (dateentré, producteur, pere, quantite, file_path, current_day))
                                    data.commit()
                            data.close()
                        else:
                            msgbox = QtWidgets.QMessageBox()
                            msgbox.setWindowTitle('Confirmation')
                            msgbox.setText('Le Bulletin ne existe pas.')
                            msgbox.exec()

                except Exception as e:
                    print(e)

        def eticket_orge(self):
            try:
                dialogor = QtWidgets.QMessageBox()
                dialogor.setFixedSize(300, 600)
                dialogor.setStyleSheet("""
                                          QWidget {
                                              color: #000000;
                                              background-color: #feffbd;
                                              border-width: 1px;
                                              border-color: #1e1e1e;
                                              border-style: solid;
                                              border-radius: 6;
                                              padding: 3px;
                                              font-size: 18px;
                                              padding-left: 5px;
                                              padding-right: 5px;
                                          }
                                          QWidget:item:hover {
                                              background-color: #3daee9;
                                              color: #eff0f1;
                                          }
                                          QWidget:item:selected {
                                              background-color: #3daee9;
                                          }
                                          QWidget:disabled {
                                              color: #454545;
                                              background-color: #31363b;
                                          }
                                          QPushButton {
                                              color: #000000;
                                              background-color:#84dbc8;
                                              border-width: 1px;
                                              border-color: #1e1e1e;
                                              border-style: solid;
                                              border-radius: 6;
                                              padding: 3px;
                                              font-size: 12px;
                                              padding-left: 5px;
                                              padding-right: 5px;
                                              min-width: 40px;
                                          }
                                          QPushButton:disabled {
                                              background-color: #31363b;
                                              border-width: 1px;
                                              border-color: #454545;
                                              border-style: solid;
                                              padding-top: 5px;
                                              padding-bottom: 5px;
                                              padding-left: 10px;
                                              padding-right: 10px;
                                              border-radius: 2px;
                                              color: #454545;
                                          }
                                          QPushButton:pressed {
                                              background-color: #3daee9;
                                              padding-top: -15px;
                                              padding-bottom: -17px;
                                          }
                                          QPushButton:hover {
                                              border: 1px solid #ff8c00;
                                              color: #eff0f1;
                                          }
                                          QLabel {
                                              font-size: 18px;
                                              border: 0px solid orange;
                                          }
                                          QLineEdit
            {
                background-color: #ffffff;
                padding: 1px;
                border-style: solid;
                border: 1px solid #76797C;
                border-radius: 0px;
                color: #000000;
            }
            QDoubleSpinBox
            {
                background-color: #ffffff;
                padding: 1px;
                border-style: solid;
                border: 1px solid #76797C;
                border-radius: 0px;
                color:#000000;
                font-size: 11px;
                font-weight: bold;

            }
            QComboBox
            {
               background-color: #ffffff;
                border-style: solid;
                border: 1px solid #76797C;
                border-radius: 2px;
                min-width: 40px;
            }
                                      """)
                dialogor.setWindowTitle("Sélectionnez le Nom et Prénom et pére")
                dialogor.setText("ORGE \t\t\t\n\t\t\t\t\t\t\n\t\t\t\t\t\t\t\n\t\t\t\t\n\t\t\t\t\t\t\t\t\t\t\n\t\t\t\t")
                nomproductureor = QtWidgets.QLineEdit(dialogor)
                nomproductureor.setPlaceholderText("Nom et Prénom")
                nomproductureor.resize(180, 30)
                nomproductureor.move(40, 125)
                pereor = QtWidgets.QLineEdit(dialogor)
                pereor.setPlaceholderText("Pére")
                pereor.resize(180, 30)
                pereor.move(240, 125)
                matricultxt = QtWidgets.QLabel("Matricul", dialogor)
                matricultxt.setGeometry(QtCore.QRect(30, 165, 100, 20))
                mator = QtWidgets.QLineEdit(dialogor)
                mator.setInputMask("99999-999-99")
                mator.resize(180, 30)
                mator.move(40, 190)
                pointdecollector = QtWidgets.QComboBox(dialogor)
                pointdecollector.setPlaceholderText("Point de collect")
                pointdecollector.addItem('')
                pointdecollector.addItem('Dock central')
                pointdecollector.addItem('Magasin Zemmoura')
                pointdecollector.addItem('Magasin Kef-lazreg')
                pointdecollector.addItem('Magasin Messra')
                pointdecollector.addItem('Magasin Marche gros belacel')
                pointdecollector.resize(180, 30)
                pointdecollector.move(40, 65)

                psor = QtWidgets.QLineEdit(dialogor)
                validator = QDoubleValidator(10, 100, 2)
                psor.setValidator(validator)
                psor.setPlaceholderText("PS")
                psor.resize(180, 30)
                psor.move(240, 65)
                quantiteptext = QtWidgets.QLabel("Quantité", dialogor)
                quantiteptext.setGeometry(QtCore.QRect(230, 165, 100, 20))
                quantiteor = QtWidgets.QDoubleSpinBox(dialogor)
                quantiteor.setRange(1, 10000)
                quantiteor.setSuffix('  QX')
                quantiteor.setSpecialValueText(" ")
                quantiteor.setValue(0.0)
                quantiteor.resize(180, 30)
                quantiteor.move(240, 190)
                ok_button = QtWidgets.QPushButton("OK", dialogor)
                cancel_button = QtWidgets.QPushButton("Cancel", dialogor)
                dialogor.addButton(ok_button, QtWidgets.QMessageBox.ButtonRole.ActionRole)
                dialogor.addButton(cancel_button, QtWidgets.QMessageBox.ButtonRole.ActionRole)
                dialogor.exec()
                current_date = QtCore.QDate.currentDate().toString(QtCore.QDate.currentDate().toString("yyyy-MM-dd"))
                current_day = current_date.split("-")[2]
                current_month = current_date.split("-")[1]
                current_year = current_date.split("-")[0]
                current_daya = f"{current_day}-{current_month}-{current_year}"

                if dialogor.clickedButton() == ok_button:
                    producteur = nomproductureor.text()
                    pere = pereor.text()
                    quantite = quantiteor.text()
                    pointdecolect = pointdecollector.currentText()
                    mator = mator.text()
                    psor = psor.text()

                    eticketorge = DocxTemplate("eticket/Eticket_orge/Eticket_orge.docx")
                    eticketorge.render(
                        {"dt": current_date, "ptt": pointdecolect, "nmp": producteur, "mat": mator, "qtt": quantite,
                         "ps": psor, "per": pere})
                    doc_names = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                    doc_pdf = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False).name
                    eticketorge.save(doc_names)
                    try:
                        if doc_names:
                            a = self.progress_bar()
                            sys.stderr = open("consoleoutput.log", "w")
                            convert(doc_names, doc_pdf)
                            # Open the resulting .pdf file using the default associated application
                            # os.startfile(doc_pdf, 'open')
                            app_path = 'C:\\Program Files\\Okular\\bin\\okular.exe'
                            subprocess.Popen([app_path, doc_pdf])
                    except Exception as e:
                        print(e)
                    name = f"ORGE_{producteur}_{current_daya}.docx"
                    eticketorge.save("eticket/Eticket_orge/" + name)
            except Exception as e:
                print(e)

        ######################################################function blé dur############################
        ##################################################################################################
        ############################################################
        def bonification_ps_bd(self):
            self.bpsbd.clear()
            self.bpsbd.setStyleSheet("background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding: 0px")
            while self.vpsbd.value() >= 80.01:
                self.bpsbd.setValue(0.15)
                self.bpsbd.setStyleSheet(
                    "background-color:#88ffaa;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                if self.vpsbd.value() <= 80.25:
                    self.bpsbd.setValue(0.15)
                elif self.vpsbd.value() >= 80.26 and self.vpsbd.value() <= 80.50:
                    self.bpsbd.setValue(0.30)
                elif self.vpsbd.value() >= 80.51 and self.vpsbd.value() <= 80.75:
                    self.bpsbd.setValue(0.45)
                elif self.vpsbd.value() >= 80.76 and self.vpsbd.value() <= 81.00:
                    self.bpsbd.setValue(0.60)
                elif self.vpsbd.value() >= 81.01 and self.vpsbd.value() <= 81.25:
                    self.bpsbd.setValue(0.75)
                elif self.vpsbd.value() >= 81.26 and self.vpsbd.value() <= 81.50:
                    self.bpsbd.setValue(0.90)
                elif self.vpsbd.value() >= 81.51 and self.vpsbd.value() <= 81.75:
                    self.bpsbd.setValue(1.05)
                elif self.vpsbd.value() >= 81.76 and self.vpsbd.value() <= 82.00:
                    self.bpsbd.setValue(1.20)
                elif self.vpsbd.value() >= 82.01 and self.vpsbd.value() <= 82.25:
                    self.bpsbd.setValue(1.30)
                elif self.vpsbd.value() >= 82.26 and self.vpsbd.value() <= 82.50:
                    self.bpsbd.setValue(1.40)
                elif self.vpsbd.value() >= 82.51 and self.vpsbd.value() <= 82.75:
                    self.bpsbd.setValue(1.50)
                elif self.vpsbd.value() >= 82.76 and self.vpsbd.value() <= 83.00:
                    self.bpsbd.setValue(1.60)
                elif self.vpsbd.value() >= 83.01 and self.vpsbd.value() <= 83.25:
                    self.bpsbd.setValue(1.65)
                elif self.vpsbd.value() >= 83.26 and self.vpsbd.value() <= 83.50:
                    self.bpsbd.setValue(1.70)
                elif self.vpsbd.value() >= 83.51 and self.vpsbd.value() <= 83.75:
                    self.bpsbd.setValue(1.75)
                elif self.vpsbd.value() >= 83.76 and self.vpsbd.value() <= 84.00:
                    self.bpsbd.setValue(1.80)
                else:
                    self.bpsbd.setValue(0)
                    self.bpsbd.setStyleSheet(
                        "background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                break

        def réfaction_ps_bd(self):
            self.rpsbd.clear()
            self.rpsbd.setStyleSheet("background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding: 0px")
            if self.vpsbd.value() < 72 and self.vpsbd.value() > 60:
                self.obesrvationbd.setText("REFUS: POIDS SPECIFIQUE INFERIEUR DE 72 Kg/hl ")
                self.obesrvationbd.setStyleSheet(
                    "background-color:#e22630;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                self.bpsbd.clear()
                self.rpsbd.clear()
            else:
                self.obesrvationbd.clear()
                self.obesrvationbd.setStyleSheet(
                    "background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding: 0px")
            while self.vpsbd.value() >= 72.00:
                self.rpsbd.setValue(3.60)
                self.rpsbd.setStyleSheet(
                    "background-color:#e22630;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                self.bpsbd.clear()
                if self.vpsbd.value() <= 72.24:
                    self.rpsbd.setValue(3.60)
                elif self.vpsbd.value() >= 72.25 and self.vpsbd.value() <= 72.49:
                    self.rpsbd.setValue(3.30)
                elif self.vpsbd.value() >= 72.50 and self.vpsbd.value() <= 72.74:
                    self.rpsbd.setValue(3.00)
                elif self.vpsbd.value() >= 72.75 and self.vpsbd.value() <= 72.99:
                    self.rpsbd.setValue(2.70)
                elif self.vpsbd.value() >= 73.00 and self.vpsbd.value() <= 73.24:
                    self.rpsbd.setValue(2.40)
                elif self.vpsbd.value() >= 73.25 and self.vpsbd.value() <= 73.49:
                    self.rpsbd.setValue(2.10)
                elif self.vpsbd.value() >= 73.50 and self.vpsbd.value() <= 73.74:
                    self.rpsbd.setValue(1.80)
                elif self.vpsbd.value() >= 73.75 and self.vpsbd.value() <= 73.99:
                    self.rpsbd.setValue(1.50)
                elif self.vpsbd.value() >= 74.00 and self.vpsbd.value() <= 74.24:
                    self.rpsbd.setValue(1.20)
                elif self.vpsbd.value() >= 74.25 and self.vpsbd.value() <= 74.49:
                    self.rpsbd.setValue(1.00)
                elif self.vpsbd.value() >= 74.50 and self.vpsbd.value() <= 74.74:
                    self.rpsbd.setValue(0.80)
                elif self.vpsbd.value() >= 74.75 and self.vpsbd.value() <= 74.99:
                    self.rpsbd.setValue(0.60)
                elif self.vpsbd.value() >= 75.00 and self.vpsbd.value() <= 75.24:
                    self.rpsbd.setValue(0.40)
                elif self.vpsbd.value() >= 75.25 and self.vpsbd.value() <= 75.49:
                    self.rpsbd.setValue(0.30)
                elif self.vpsbd.value() >= 75.50 and self.vpsbd.value() <= 75.74:
                    self.rpsbd.setValue(0.20)
                elif self.vpsbd.value() >= 75.75 and self.vpsbd.value() <= 75.99:
                    self.rpsbd.setValue(0.10)
                else:
                    self.rpsbd.setValue(0)
                    self.rpsbd.setStyleSheet(
                        "background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                break

        def bonification_pimpurte_er_bd(self):
            self.btotalprembd.clear()
            self.btotalprembd.setStyleSheet(
                "background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding: 0px")
            while self.vtotalprembd.value() >= 0.01:
                self.btotalprembd.setValue(0.125)
                self.btotalprembd.setStyleSheet(
                    "background-color:#88ffaa;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                if self.vtotalprembd.value() <= 0.25:
                    self.btotalprembd.setValue(0.125)
                elif self.vtotalprembd.value() >= 0.26 and self.vtotalprembd.value() <= 0.50:
                    self.btotalprembd.setValue(0.250)
                elif self.vtotalprembd.value() >= 0.51 and self.vtotalprembd.value() <= 0.75:
                    self.btotalprembd.setValue(0.375)
                elif self.vtotalprembd.value() >= 0.76 and self.vtotalprembd.value() <= 1.00:
                    self.btotalprembd.setValue(0.500)
                else:
                    self.btotalprembd.setValue(0)
                    self.btotalprembd.clear()
                    self.btotalprembd.setStyleSheet(
                        "background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                break

        def refaction_impurté_er_bd(self):
            self.rtotalprembd.clear()
            self.rtotalprembd.setStyleSheet(
                "background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding: 0px")
            while self.vtotalprembd.value() >= 3.01:
                self.rtotalprembd.setValue(0.125)
                self.rtotalprembd.setStyleSheet(
                    "background-color:#e22630;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                self.btotalprembd.clear()
                if self.vtotalprembd.value() <= 3.25:
                    self.rtotalprembd.setValue(0.125)
                elif self.vtotalprembd.value() >= 3.26 and self.vtotalprembd.value() <= 3.50:
                    self.rtotalprembd.setValue(0.250)
                elif self.vtotalprembd.value() >= 3.51 and self.vtotalprembd.value() <= 3.75:
                    self.rtotalprembd.setValue(0.375)
                elif self.vtotalprembd.value() >= 3.76 and self.vtotalprembd.value() <= 4.00:
                    self.rtotalprembd.setValue(0.500)
                elif self.vtotalprembd.value() >= 4.01 and self.vtotalprembd.value() <= 4.25:
                    self.rtotalprembd.setValue(0.625)
                elif self.vtotalprembd.value() >= 4.26 and self.vtotalprembd.value() <= 4.50:
                    self.rtotalprembd.setValue(0.750)
                elif self.vtotalprembd.value() >= 4.51 and self.vtotalprembd.value() <= 4.75:
                    self.rtotalprembd.setValue(0.875)
                elif self.vtotalprembd.value() >= 4.76 and self.vtotalprembd.value() <= 5.00:
                    self.rtotalprembd.setValue(1.000)
                elif self.vtotalprembd.value() >= 5.01 and self.vtotalprembd.value() <= 5.25:
                    self.rtotalprembd.setValue(1.56)
                elif self.vtotalprembd.value() >= 8.51 and self.vtotalprembd.value() <= 9.00:
                    self.rtotalprembd.setValue(1.125)
                elif self.vtotalprembd.value() >= 5.26 and self.vtotalprembd.value() <= 5.50:
                    self.rtotalprembd.setValue(1.250)
                elif self.vtotalprembd.value() >= 5.51 and self.vtotalprembd.value() <= 5.75:
                    self.rtotalprembd.setValue(1.375)
                elif self.vtotalprembd.value() >= 5.76 and self.vtotalprembd.value() <= 6.00:
                    self.rtotalprembd.setValue(1.500)
                elif self.vtotalprembd.value() >= 6.01 and self.vtotalprembd.value() <= 30:
                    self.rtotalprembd.setValue(0)
                    self.obesrvationbd.setText("PRIX A DEBATTRE")

                else:
                    self.rtotalprembd.setValue(0)
                    self.rtotalprembd.clear
                    self.rtotalprembd.setStyleSheet(
                        "background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                break

        def refaction_impurte_eme_bd(self):
            self.rtotaldembd.clear()
            self.rtotaldembd.setStyleSheet(
                "background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding: 0px")
            while self.vtotaldembd.value() >= 10.01:
                self.rtotaldembd.setValue(0.50)
                self.rtotaldembd.setStyleSheet(
                    "background-color:#e22630;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                self.btotaldembd.clear()
                if self.vtotaldembd.value() <= 11.00:
                    self.rtotaldembd.setValue(0.50)
                elif self.vtotaldembd.value() >= 11.01 and self.vtotaldembd.value() <= 12.00:
                    self.rtotaldembd.setValue(1.00)
                elif self.vtotaldembd.value() >= 12.01 and self.vtotaldembd.value() <= 13.00:
                    self.rtotaldembd.setValue(1.50)
                elif self.vtotaldembd.value() >= 13.01 and self.vtotaldembd.value() <= 14.00:
                    self.rtotaldembd.setValue(2.00)
                elif self.vtotaldembd.value() >= 14.01 and self.vtotaldembd.value() <= 15.00:
                    self.rtotaldembd.setValue(2.50)
                elif self.vtotaldembd.value() >= 15.01 and self.vtotaldembd.value() <= 16.00:
                    self.rtotaldembd.setValue(3.00)
                elif self.vtotaldembd.value() >= 16.01 and self.vtotaldembd.value() <= 17.00:
                    self.rtotaldembd.setValue(3.50)
                elif self.vtotaldembd.value() >= 17.01 and self.vtotaldembd.value() <= 18.00:
                    self.rtotaldembd.setValue(4.00)
                elif self.vtotaldembd.value() >= 18.01 and self.vtotaldembd.value() <= 19.00:
                    self.rtotaldembd.setValue(4.50)
                elif self.vtotaldembd.value() >= 19.01 and self.vtotaldembd.value() <= 20.00:
                    self.rtotaldembd.setValue(5.00)
                elif self.vtotaldembd.value() >= 20.01 and self.vtotaldembd.value() <= 50.00:
                    self.rtotaldembd.setValue(0)
                    self.obesrvationbd.setText("PRIX A DEBATTRE")
                else:
                    self.rtotaldembd.setValue(0)
                    self.rtotaldembd.clear()
                    self.rtotaldembd.setStyleSheet(
                        "background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                break

        def rGrains_casse(self):
            self.rgrainscassébd.clear()
            self.rgrainscassébd.setStyleSheet(
                "background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding: 0px")
            while self.vgrainscassébd.value() >= 5.01:
                self.rgrainscassébd.setValue(0.075)
                self.rgrainscassébd.setStyleSheet(
                    "background-color:#e22630;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                if self.vgrainscassébd.value() <= 5.25:
                    self.rgrainscassébd.setValue(0.075)
                elif self.vgrainscassébd.value() >= 5.26 and self.vgrainscassébd.value() <= 5.50:
                    self.rgrainscassébd.setValue(0.150)
                elif self.vgrainscassébd.value() >= 5.51 and self.vgrainscassébd.value() <= 5.75:
                    self.rgrainscassébd.setValue(0.225)
                elif self.vgrainscassébd.value() >= 5.76 and self.vgrainscassébd.value() <= 6.00:
                    self.rgrainscassébd.setValue(0.300)
                elif self.vgrainscassébd.value() >= 6.01 and self.vgrainscassébd.value() <= 6.25:
                    self.rgrainscassébd.setValue(0.375)
                elif self.vgrainscassébd.value() >= 6.26 and self.vgrainscassébd.value() <= 6.50:
                    self.rgrainscassébd.setValue(0.450)
                elif self.vgrainscassébd.value() >= 6.51 and self.vgrainscassébd.value() <= 6.75:
                    self.rgrainscassébd.setValue(0.525)
                elif self.vgrainscassébd.value() >= 6.76 and self.vgrainscassébd.value() <= 7.00:
                    self.rgrainscassébd.setValue(0.600)
                elif self.vgrainscassébd.value() >= 7.01 and self.vgrainscassébd.value() <= 7.25:
                    self.rgrainscassébd.setValue(0.675)
                elif self.vgrainscassébd.value() >= 7.26 and self.vgrainscassébd.value() <= 7.50:
                    self.rgrainscassébd.setValue(0.750)
                elif self.vgrainscassébd.value() >= 7.51 and self.vgrainscassébd.value() <= 7.75:
                    self.rgrainscassébd.setValue(0.825)
                elif self.vgrainscassébd.value() >= 7.76 and self.vgrainscassébd.value() <= 8.00:
                    self.rgrainscassébd.setValue(0.900)
                elif self.vgrainscassébd.value() >= 8.01 and self.vgrainscassébd.value() <= 8.25:
                    self.rgrainscassébd.setValue(0.975)
                elif self.vgrainscassébd.value() >= 8.26 and self.vgrainscassébd.value() <= 8.50:
                    self.rgrainscassébd.setValue(1.050)
                elif self.vgrainscassébd.value() >= 8.51 and self.vgrainscassébd.value() <= 8.75:
                    self.rgrainscassébd.setValue(1.125)
                elif self.vgrainscassébd.value() >= 8.76 and self.vgrainscassébd.value() <= 9.00:
                    self.rgrainscassébd.setValue(1.200)
                elif self.vgrainscassébd.value() >= 9.01 and self.vgrainscassébd.value() <= 30:
                    self.rgrainscassébd.setValue(0)
                    self.obesrvationbd.setText("PRIX A DEBATTRE")
                else:
                    self.rgrainscassébd.setValue(0)
                    self.rgrainscassébd.clear()
                    self.rgrainscassébd.setStyleSheet(
                        "background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                break

        def bonification_mitadinage(self):
            self.bgrainmitadinésbd.setStyleSheet(
                "background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding: 0px")
            self.bgrainmitadinésbd.clear()
            self.bgrainmitadinésbd.setValue(0)
            while self.vgrainmitadinésbd.value() >= 0.01:
                self.bgrainmitadinésbd.setValue(0.25)
                if self.vgrainmitadinésbd.value() <= 20.00:
                    self.bgrainmitadinésbd.setStyleSheet(
                        "background-color:#88ffaa;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                    self.bgrainmitadinésbd.setValue(0.25)
                else:
                    self.bgrainmitadinésbd.setStyleSheet(
                        "background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                    self.bgrainmitadinésbd.clear()
                    self.bgrainmitadinésbd.setValue(0)

                break

        def mitadinage_bd(self):
            self.rgrainmitadinésbd.clear()
            self.bgrainmitadinésbd.clear()
            self.bgrainmitadinésbd.setStyleSheet(
                "background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding: 0px")
            self.rgrainmitadinésbd.setStyleSheet(
                "background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding: 0px")
            while self.vgrainmitadinésbd.value() >= 20.01:
                self.rgrainmitadinésbd.setValue(0.05)
                self.rgrainmitadinésbd.setStyleSheet(
                    "background-color:#e22630;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                if self.vgrainmitadinésbd.value() <= 21.00:
                    self.rgrainmitadinésbd.setValue(0.05)
                elif self.vgrainmitadinésbd.value() >= 21.01 and self.vgrainmitadinésbd.value() <= 22.00:
                    self.rgrainmitadinésbd.setValue(0.10)
                elif self.vgrainmitadinésbd.value() >= 22.01 and self.vgrainmitadinésbd.value() <= 23.00:
                    self.rgrainmitadinésbd.setValue(0.15)
                elif self.vgrainmitadinésbd.value() >= 23.01 and self.vgrainmitadinésbd.value() <= 24.00:
                    self.rgrainmitadinésbd.setValue(0.20)
                elif self.vgrainmitadinésbd.value() >= 24.01 and self.vgrainmitadinésbd.value() <= 25.00:
                    self.rgrainmitadinésbd.setValue(0.25)
                elif self.vgrainmitadinésbd.value() >= 25.01 and self.vgrainmitadinésbd.value() <= 26.00:
                    self.rgrainmitadinésbd.setValue(0.30)
                elif self.vgrainmitadinésbd.value() >= 26.01 and self.vgrainmitadinésbd.value() <= 27.00:
                    self.rgrainmitadinésbd.setValue(0.35)
                elif self.vgrainmitadinésbd.value() >= 27.01 and self.vgrainmitadinésbd.value() <= 28.00:
                    self.rgrainmitadinésbd.setValue(0.40)
                elif self.vgrainmitadinésbd.value() >= 28.01 and self.vgrainmitadinésbd.value() <= 29.00:
                    self.rgrainmitadinésbd.setValue(0.45)
                elif self.vgrainmitadinésbd.value() >= 29.01 and self.vgrainmitadinésbd.value() <= 30.00:
                    self.rgrainmitadinésbd.setValue(0.50)
                elif self.vgrainmitadinésbd.value() >= 30.01 and self.vgrainmitadinésbd.value() <= 31.00:
                    self.rgrainmitadinésbd.setValue(0.55)
                elif self.vgrainmitadinésbd.value() >= 31.01 and self.vgrainmitadinésbd.value() <= 32.00:
                    self.rgrainmitadinésbd.setValue(0.60)
                elif self.vgrainmitadinésbd.value() >= 32.01 and self.vgrainmitadinésbd.value() <= 33.00:
                    self.rgrainmitadinésbd.setValue(0.65)
                elif self.vgrainmitadinésbd.value() >= 33.01 and self.vgrainmitadinésbd.value() <= 34.00:
                    self.rgrainmitadinésbd.setValue(0.70)
                elif self.vgrainmitadinésbd.value() >= 34.01 and self.vgrainmitadinésbd.value() <= 35.00:
                    self.rgrainmitadinésbd.setValue(0.75)
                elif self.vgrainmitadinésbd.value() >= 35.01 and self.vgrainmitadinésbd.value() <= 36.00:
                    self.rgrainmitadinésbd.setValue(0.80)
                elif self.vgrainmitadinésbd.value() >= 36.01 and self.vgrainmitadinésbd.value() <= 37.00:
                    self.rgrainmitadinésbd.setValue(0.85)
                elif self.vgrainmitadinésbd.value() >= 37.01 and self.vgrainmitadinésbd.value() <= 38.00:
                    self.rgrainmitadinésbd.setValue(0.90)
                elif self.vgrainmitadinésbd.value() >= 38.01 and self.vgrainmitadinésbd.value() <= 39.00:
                    self.rgrainmitadinésbd.setValue(0.95)
                elif self.vgrainmitadinésbd.value() >= 39.01 and self.vgrainmitadinésbd.value() <= 40.00:
                    self.rgrainmitadinésbd.setValue(1.00)
                elif self.vgrainmitadinésbd.value() >= 40.01 and self.vgrainmitadinésbd.value() <= 41.00:
                    self.rgrainmitadinésbd.setValue(1.05)
                elif self.vgrainmitadinésbd.value() >= 41.01 and self.vgrainmitadinésbd.value() <= 42.00:
                    self.rgrainmitadinésbd.setValue(1.10)
                elif self.vgrainmitadinésbd.value() >= 42.01 and self.vgrainmitadinésbd.value() <= 43.00:
                    self.rgrainmitadinésbd.setValue(1.15)
                elif self.vgrainmitadinésbd.value() >= 43.01 and self.vgrainmitadinésbd.value() <= 44.00:
                    self.rgrainmitadinésbd.setValue(1.20)
                elif self.vgrainmitadinésbd.value() >= 44.01 and self.vgrainmitadinésbd.value() <= 45.00:
                    self.rgrainmitadinésbd.setValue(1.25)
                elif self.vgrainmitadinésbd.value() >= 45.01 and self.vgrainmitadinésbd.value() <= 46.00:
                    self.rgrainmitadinésbd.setValue(1.30)
                elif self.vgrainmitadinésbd.value() >= 46.01 and self.vgrainmitadinésbd.value() <= 47.00:
                    self.rgrainmitadinésbd.setValue(1.35)
                elif self.vgrainmitadinésbd.value() >= 47.01 and self.vgrainmitadinésbd.value() <= 48.00:
                    self.rgrainmitadinésbd.setValue(1.40)
                elif self.vgrainmitadinésbd.value() >= 48.01 and self.vgrainmitadinésbd.value() <= 49.00:
                    self.rgrainmitadinésbd.setValue(1.45)
                elif self.vgrainmitadinésbd.value() >= 49.01 and self.vgrainmitadinésbd.value() <= 50.00:
                    self.rgrainmitadinésbd.setValue(1.50)
                elif self.vgrainmitadinésbd.value() >= 50.01 and self.vgrainmitadinésbd.value() <= 51.00:
                    self.rgrainmitadinésbd.setValue(1.55)
                elif self.vgrainmitadinésbd.value() >= 51.01 and self.vgrainmitadinésbd.value() <= 52.00:
                    self.rgrainmitadinésbd.setValue(1.60)
                elif self.vgrainmitadinésbd.value() >= 52.01 and self.vgrainmitadinésbd.value() <= 53.00:
                    self.rgrainmitadinésbd.setValue(1.65)
                elif self.vgrainmitadinésbd.value() >= 53.01 and self.vgrainmitadinésbd.value() <= 54.00:
                    self.rgrainmitadinésbd.setValue(1.70)
                elif self.vgrainmitadinésbd.value() >= 54.01 and self.vgrainmitadinésbd.value() <= 55.00:
                    self.rgrainmitadinésbd.setValue(1.75)
                elif self.vgrainmitadinésbd.value() >= 55.01 and self.vgrainmitadinésbd.value() <= 56.00:
                    self.rgrainmitadinésbd.setValue(1.80)
                elif self.vgrainmitadinésbd.value() >= 56.01 and self.vgrainmitadinésbd.value() <= 57.00:
                    self.rgrainmitadinésbd.setValue(1.85)
                elif self.vgrainmitadinésbd.value() >= 57.01 and self.vgrainmitadinésbd.value() <= 58.00:
                    self.rgrainmitadinésbd.setValue(1.90)
                elif self.vgrainmitadinésbd.value() >= 58.01 and self.vgrainmitadinésbd.value() <= 59.00:
                    self.rgrainmitadinésbd.setValue(1.95)
                elif self.vgrainmitadinésbd.value() >= 59.01 and self.vgrainmitadinésbd.value() <= 60.00:
                    self.rgrainmitadinésbd.setValue(2.00)
                elif self.vgrainmitadinésbd.value() >= 60.01 and self.vgrainmitadinésbd.value() <= 61.00:
                    self.rgrainmitadinésbd.setValue(2.05)
                elif self.vgrainmitadinésbd.value() >= 61.01 and self.vgrainmitadinésbd.value() <= 62.00:
                    self.rgrainmitadinésbd.setValue(2.10)
                elif self.vgrainmitadinésbd.value() >= 62.01 and self.vgrainmitadinésbd.value() <= 63.00:
                    self.rgrainmitadinésbd.setValue(2.15)
                elif self.vgrainmitadinésbd.value() >= 63.01 and self.vgrainmitadinésbd.value() <= 64.00:
                    self.rgrainmitadinésbd.setValue(2.20)
                elif self.vgrainmitadinésbd.value() >= 64.01 and self.vgrainmitadinésbd.value() <= 65.00:
                    self.rgrainmitadinésbd.setValue(2.25)
                elif self.vgrainmitadinésbd.value() >= 65.01 and self.vgrainmitadinésbd.value() <= 66.00:
                    self.rgrainmitadinésbd.setValue(2.30)
                elif self.vgrainmitadinésbd.value() >= 66.01 and self.vgrainmitadinésbd.value() <= 67.00:
                    self.rgrainmitadinésbd.setValue(2.35)
                elif self.vgrainmitadinésbd.value() >= 67.01 and self.vgrainmitadinésbd.value() <= 68.00:
                    self.rgrainmitadinésbd.setValue(2.40)
                elif self.vgrainmitadinésbd.value() >= 68.01 and self.vgrainmitadinésbd.value() <= 69.00:
                    self.rgrainmitadinésbd.setValue(2.45)
                elif self.vgrainmitadinésbd.value() >= 69.01 and self.vgrainmitadinésbd.value() <= 70.00:
                    self.rgrainmitadinésbd.setValue(2.50)
                elif self.vgrainmitadinésbd.value() >= 70.01 and self.vgrainmitadinésbd.value() <= 100:
                    self.rgrainmitadinésbd.setValue(0)
                    self.obesrvationbd.setText("PRIX A DEBATTRE")
                    self.rgrainmitadinésbd.setStyleSheet(
                        "background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                else:
                    self.rgrainmitadinésbd.clear()
                    self.vgrainmitadinésbd.clear()
                    self.bgrainmitadinésbd.clear()
                    self.rgrainmitadinésbd.setStyleSheet(
                        "background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                break

        def grain_pounaisée(self):
            self.rgrainepunaisésbd.setValue(0)
            self.rgrainepunaisésbd.clear()
            self.rgrainepunaisésbd.setStyleSheet(
                "background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding: 0px")
            while self.vgrainepunaisésbd.value() >= 2.01:
                self.rgrainepunaisésbd.setValue(0.08)
                self.rgrainepunaisésbd.setStyleSheet(
                    "background-color:#e22630;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                if self.vgrainepunaisésbd.value() <= 2.25:
                    self.rgrainepunaisésbd.setValue(0.08)
                elif self.vgrainepunaisésbd.value() >= 2.26 and self.vgrainepunaisésbd.value() <= 2.50:
                    self.rgrainepunaisésbd.setValue(0.16)
                elif self.vgrainepunaisésbd.value() >= 2.51 and self.vgrainepunaisésbd.value() <= 2.75:
                    self.rgrainepunaisésbd.setValue(0.24)
                elif self.vgrainepunaisésbd.value() >= 2.76 and self.vgrainepunaisésbd.value() <= 3.00:
                    self.rgrainepunaisésbd.setValue(0.32)
                elif self.vgrainepunaisésbd.value() >= 3.01 and self.vgrainepunaisésbd.value() <= 3.25:
                    self.rgrainepunaisésbd.setValue(0.40)
                elif self.vgrainepunaisésbd.value() >= 3.26 and self.vgrainepunaisésbd.value() <= 3.50:
                    self.rgrainepunaisésbd.setValue(0.48)
                elif self.vgrainepunaisésbd.value() >= 3.51 and self.vgrainepunaisésbd.value() <= 3.75:
                    self.rgrainepunaisésbd.setValue(0.56)
                elif self.vgrainepunaisésbd.value() >= 3.76 and self.vgrainepunaisésbd.value() <= 4.00:
                    self.rgrainepunaisésbd.setValue(0.64)
                elif self.vgrainepunaisésbd.value() >= 4.01 and self.vgrainepunaisésbd.value() <= 4.25:
                    self.rgrainepunaisésbd.setValue(0.72)
                elif self.vgrainepunaisésbd.value() >= 4.26 and self.vgrainepunaisésbd.value() <= 4.50:
                    self.rgrainepunaisésbd.setValue(0.80)
                elif self.vgrainepunaisésbd.value() >= 4.51 and self.vgrainepunaisésbd.value() <= 4.75:
                    self.rgrainepunaisésbd.setValue(0.88)
                elif self.vgrainepunaisésbd.value() >= 4.76 and self.vgrainepunaisésbd.value() <= 5.00:
                    self.rgrainepunaisésbd.setValue(0.96)
                elif self.vgrainepunaisésbd.value() >= 5.01 and self.vgrainepunaisésbd.value() <= 5.25:
                    self.rgrainepunaisésbd.setValue(1.04)
                elif self.vgrainepunaisésbd.value() >= 5.26 and self.vgrainepunaisésbd.value() <= 5.50:
                    self.rgrainepunaisésbd.setValue(1.12)
                elif self.vgrainepunaisésbd.value() >= 5.51 and self.vgrainepunaisésbd.value() <= 5.75:
                    self.rgrainepunaisésbd.setValue(1.20)
                elif self.vgrainepunaisésbd.value() >= 5.76 and self.vgrainepunaisésbd.value() <= 6.00:
                    self.rgrainepunaisésbd.setValue(1.28)
                elif self.vgrainepunaisésbd.value() >= 6.01 and self.vgrainepunaisésbd.value() <= 6.25:
                    self.rgrainepunaisésbd.setValue(1.36)
                elif self.vgrainepunaisésbd.value() >= 6.26 and self.vgrainepunaisésbd.value() <= 6.50:
                    self.rgrainepunaisésbd.setValue(1.44)
                elif self.vgrainepunaisésbd.value() >= 6.51 and self.vgrainepunaisésbd.value() <= 6.75:
                    self.rgrainepunaisésbd.setValue(1.52)
                elif self.vgrainepunaisésbd.value() >= 6.76 and self.vgrainepunaisésbd.value() <= 7.00:
                    self.rgrainepunaisésbd.setValue(1.60)
                elif self.vgrainepunaisésbd.value() >= 7.01 and self.vgrainepunaisésbd.value() <= 7.25:
                    self.rgrainepunaisésbd.setValue(1.68)
                elif self.vgrainepunaisésbd.value() >= 7.26 and self.vgrainepunaisésbd.value() <= 7.50:
                    self.rgrainepunaisésbd.setValue(1.76)
                elif self.vgrainepunaisésbd.value() >= 7.51 and self.vgrainepunaisésbd.value() <= 7.75:
                    self.rgrainepunaisésbd.setValue(1.84)
                elif self.vgrainepunaisésbd.value() >= 7.76 and self.vgrainepunaisésbd.value() <= 8.00:
                    self.rgrainepunaisésbd.setValue(1.92)
                elif self.vgrainepunaisésbd.value() >= 8.01 and self.vgrainepunaisésbd.value() <= 8.25:
                    self.rgrainepunaisésbd.setValue(2.00)
                elif self.vgrainepunaisésbd.value() >= 8.26 and self.vgrainepunaisésbd.value() <= 8.50:
                    self.rgrainepunaisésbd.setValue(2.08)
                elif self.vgrainepunaisésbd.value() >= 8.51 and self.vgrainepunaisésbd.value() <= 8.75:
                    self.rgrainepunaisésbd.setValue(2.16)
                elif self.vgrainepunaisésbd.value() >= 8.76 and self.vgrainepunaisésbd.value() <= 9.00:
                    self.rgrainepunaisésbd.setValue(2.24)
                elif self.vgrainepunaisésbd.value() >= 9.01 and self.vgrainepunaisésbd.value() <= 9.25:
                    self.rgrainepunaisésbd.setValue(2.32)
                elif self.vgrainepunaisésbd.value() >= 9.26 and self.vgrainepunaisésbd.value() <= 9.50:
                    self.rgrainepunaisésbd.setValue(2.40)
                elif self.vgrainepunaisésbd.value() >= 9.51 and self.vgrainepunaisésbd.value() <= 9.75:
                    self.rgrainepunaisésbd.setValue(2.48)
                elif self.vgrainepunaisésbd.value() >= 9.76 and self.vgrainepunaisésbd.value() <= 10.00:
                    self.rgrainepunaisésbd.setValue(2.56)
                elif self.vgrainepunaisésbd.value() >= 10.01 and self.vgrainepunaisésbd.value() <= 40:
                    self.rgrainepunaisésbd.setValue(0)
                    self.obesrvationbd.setText("PRIX A DEBATTRE")
                else:
                    self.rgrainepunaisésbd.setValue(0)
                    self.rgrainepunaisésbd.clear()
                    self.rgrainepunaisésbd.setStyleSheet(
                        "background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                break

        def clear_btotalprembd(self):
            if self.vtotalprembd == 0.00:
                self.btotalprembd.setValue(0)
                self.rtotalprembd.setValue(0)
                self.btotalprembd.clear()
                self.rtotalprembd.clear

        def plus_bd(self):
            abd = self.vmatiére20mmbd.value()
            bbd = self.vdébrisvébd.value()
            cbd = self.vgrainnuisiblebd.value()
            dbd = self.vgrainscassébd.value()
            ebd = self.vgrainsboutésbd.value()
            fbd = self.vgrainsrouxbd.value()
            gbd = self.vgrainfortementmouchtébd.value()
            hbd = self.vgrainepunaisésbd.value()
            ibd = self.vgrainpiquebd.value()
            jbd = self.vgrainmitadinésbd.value()
            kbd = self.vbletendredansbledurbd.value()
            totalprembd = abd + bbd + cbd
            totaldem = dbd + ebd + fbd + gbd + hbd + ibd
            totalmitadine = jbd + kbd
            btotalBonietRefacbd = self.btotalprembd.value() + self.btotaldembd.value() + self.btotalmitadinésbd.value() + self.bpsbd.value() + self.bgrainmitadinésbd.value()
            rtotalBonietRefac = self.rtotalprembd.value() + self.rtotaldembd.value() + self.rtotalmitadinésbd.value() + self.rpsbd.value() + self.rgrainmitadinésbd.value() + self.rgrainscassébd.value() + self.rgrainepunaisésbd.value()
            self.btotalbiniEtrefaction.setValue(btotalBonietRefacbd)
            self.rtotalbiniEtrefaction.setValue(rtotalBonietRefac)
            self.vtotalmitadinésbd.setValue(totalmitadine)
            self.vtotalprembd.setValue(totalprembd)
            if self.vgrainscassébd.value() < 5.00:
                self.vtotaldembd.setValue(totaldem)
            elif self.vgrainscassébd.value() >= 5.01:
                self.vtotaldembd.setValue(ebd + fbd + gbd + hbd + ibd)

        def clear_bd(self):
            self.vpsbd.clear()
            self.vpsbd.setValue(0)
            self.vhumiditebd.clear()
            self.vergotbd.clear()
            self.vmatiére20mmbd.clear()
            self.vmatiére20mmbd.setValue(0)
            self.vdébrisvébd.clear()
            self.vdébrisvébd.setValue(0)
            self.vgrainnuisiblebd.clear()
            self.vgrainnuisiblebd.setValue(0)
            self.vtotalprembd.clear()
            self.vtotalprembd.setValue(0)
            self.vgrainscassébd.clear()
            self.vgrainscassébd.setValue(0)
            self.vgrainsboutésbd.clear()
            self.vgrainsboutésbd.setValue(0)
            self.vgrainsrouxbd.clear()
            self.vgrainsrouxbd.setValue(0)
            self.vgrainfortementmouchtébd.clear()
            self.vgrainfortementmouchtébd.setValue(0)
            self.vgrainepunaisésbd.clear()
            self.vgrainepunaisésbd.setValue(0)
            self.vgrainpiquebd.clear()
            self.vgrainpiquebd.setValue(0)
            self.vtotaldembd.clear()
            self.vtotaldembd.setValue(0)
            self.vgrainmitadinésbd.clear()
            self.vgrainmitadinésbd.setValue(0)
            self.vbletendredansbledurbd.clear()
            self.vbletendredansbledurbd.setValue(0)
            self.vtotalmitadinésbd.clear()
            self.vtotalmitadinésbd.setValue(0)
            self.rpsbd.clear()
            self.rpsbd.setValue(0)
            self.rergotbd.clear()
            self.rmatiére20mmbd.clear()
            self.rmatiére20mmbd.setValue(0)
            self.rdébrisvébd.clear()
            self.rdébrisvébd.setValue(0)
            self.rgrainnuisiblebd.clear()
            self.rgrainnuisiblebd.setValue(0)
            self.rtotalprembd.clear()
            self.rtotalprembd.setValue(0)
            self.rgrainscassébd.clear()
            self.rgrainscassébd.setValue(0)
            self.rgrainsboutésbd.clear()
            self.rgrainsboutésbd.setValue(0)
            self.rgrainsrouxbd.clear()
            self.rgrainsrouxbd.setValue(0)
            self.rgrainfortementmouchtébd.clear()
            self.rgrainfortementmouchtébd.setValue(0)
            self.rgrainepunaisésbd.clear()
            self.rgrainepunaisésbd.setValue(0)
            self.rgrainpiquebd.clear()
            self.rgrainpiquebd.setValue(0)
            self.rtotaldembd.clear()
            self.rtotaldembd.setValue(0)
            self.rgrainmitadinésbd.clear()
            self.rgrainmitadinésbd.setValue(0)
            self.rbletendredansbledurbd.clear()
            self.rbletendredansbledurbd.setValue(0)
            self.rtotalmitadinésbd.clear()
            self.rtotalmitadinésbd.setValue(0)
            self.bpsbd.clear()
            self.bpsbd.setValue(0)
            self.bergotbd.clear()
            self.bmatiére20mmbd.clear()
            self.bmatiére20mmbd.setValue(0)
            self.bdébrisvébd.clear()
            self.bdébrisvébd.setValue(0)
            self.bgrainnuisiblebd.clear()
            self.bgrainnuisiblebd.setValue(0)
            self.btotalprembd.clear()
            self.btotalprembd.setValue(0)
            self.bgrainscassébd.clear()
            self.bgrainscassébd.setValue(0)
            self.bgrainsboutésbd.clear()
            self.bgrainsboutésbd.setValue(0)
            self.bgrainsrouxbd.clear()
            self.bgrainsrouxbd.setValue(0)
            self.bgrainfortementmouchtébd.clear()
            self.bgrainfortementmouchtébd.setValue(0)
            self.bgrainepunaisésbd.clear()
            self.bgrainepunaisésbd.setValue(0)
            self.bgrainpiquebd.clear()
            self.bgrainpiquebd.setValue(0)
            self.btotaldembd.clear()
            self.btotaldembd.setValue(0)
            self.bgrainmitadinésbd.clear()
            self.bgrainmitadinésbd.setValue(0)
            self.bbletendredansbledurbd.clear()
            self.bbletendredansbledurbd.setValue(0)
            self.btotalmitadinésbd.clear()
            self.btotalmitadinésbd.setValue(0)

        def calcul_bd(self):
            self.réfaction_ps_bd()
            self.bonification_ps_bd()
            self.refaction_impurté_er_bd()
            self.bonification_pimpurte_er_bd()
            self.refaction_impurte_eme_bd()
            self.rGrains_casse()
            self.mitadinage_bd()
            self.bonification_mitadinage()
            self.grain_pounaisée()
            self.plus_bd()

        def bulletin_bd(self):
            try:
                nbulltinbd = self.n_bultinbd.text()
                nomproducteurbd = self.nome_du_producteurbd.text()
                perebd = self.pérebd.text()
                matriculbd = self.adressebd.text()
                poindecollectebd = self.pointdecollectbd.currentText()
                datebd = self.dattereceptiontbd.dateTime().toString("dd-MM-yyyy hh:mm")
                agreeur = self.agréeeurcombobd.currentText()
                quantitebd = self.quantitebd.value()
                vpsbd = self.vpsbd.value()
                bpsbd = self.bpsbd.text()
                rpsbd = self.rpsbd.text()
                vhumidite = self.vhumiditebd.value()
                vergot = self.vergotbd.value()
                vmatier20mm = self.vmatiére20mmbd.value()
                vdebrit = self.vdébrisvébd.value()
                vgrainnuisible = self.vgrainnuisiblebd.value()
                graincasse = self.vgrainscassébd.value()
                rgraicasse = self.rgrainscassébd.text()
                grainboute = self.vgrainsboutésbd.value()
                grainroux = self.vgrainsrouxbd.value()
                grainfortementmouchte = self.vgrainfortementmouchtébd.value()
                grainpunaise = self.vgrainepunaisésbd.value()
                rgrainpunaise = self.rgrainepunaisésbd.text()
                grainpique = self.vgrainpiquebd.value()
                grainmitadine = self.vgrainmitadinésbd.value()
                bletendedbledur = self.vbletendredansbledurbd.value()
                vtotalbdp = self.vtotalprembd.value()
                vtotalbdd = self.vtotaldembd.value()
                btotalbdp = self.btotalprembd.text()
                btotalbdd = self.btotaldembd.text()
                rtotalbdp = self.rtotalprembd.text()
                rtotalbdd = self.rtotaldembd.text()
                bmitadine = self.bgrainmitadinésbd.text()
                rmitadine = self.rgrainmitadinésbd.text()
                vtotalmitadineetbletendre = self.vtotalmitadinésbd.value()
                btotaldb = self.btotalbiniEtrefaction.text()
                rtotalbd = self.rtotalbiniEtrefaction.text()
                ndcartebd = self.n_cartebd.text()
                observationbd = self.obesrvationbd.toPlainText()
                date = QDate.currentDate()
                current_date = QDate.currentDate()
                current_day = current_date.day()
                current_month = current_date.month()
                current_year = current_date.year()
                current_time = QTime.currentTime()
                current_hour = current_time.hour()
                current_minute = current_time.minute()
                # Format the date as "dd-mm-yyyy"
                current_days = f"{current_day:02d}-{current_month:02d}-{current_year}-{current_hour}-{current_minute}"
                self.docbd = DocxTemplate("bulletin agreage/bulletin_blédur/bulletin_blédur.docx")
                self.docbd.render({
                    "nmp": nomproducteurbd,
                    "per": perebd,
                    "mat": matriculbd,
                    "ptt": poindecollectebd,
                    "dt": datebd,
                    "agr": agreeur,
                    "qtt": quantitebd,
                    "vps": vpsbd,
                    "bps": bpsbd,
                    "rps": rpsbd,
                    "vmh": vhumidite,
                    "vrg": vergot,
                    "vgs": vmatier20mm,
                    "vdb": vdebrit,
                    "vgn": vgrainnuisible,
                    "vgc": graincasse,
                    "rgc": rgraicasse,
                    "vgfb": grainboute,
                    "vgr": grainroux,
                    "vgfm": grainfortementmouchte,
                    "vgpn": grainpunaise,
                    "rgpn": rgrainpunaise,
                    "vgrp": grainpique,
                    "vmt": grainmitadine,
                    "vbtb": bletendedbledur,
                    "vttp": vtotalbdp,
                    "bttp": btotalbdp,
                    "rtt": rtotalbdp,
                    "vtd": vtotalbdd,
                    "btd": btotalbdd,
                    "rtd": rtotalbdd,
                    "btb": btotaldb,
                    "rtr": rtotalbd,
                    "bmt": bmitadine,
                    "rmt": rmitadine,
                    "vttm": vtotalmitadineetbletendre,
                    "ncn": ndcartebd,
                    "oo": observationbd,
                    "nm": nbulltinbd
                })
                datadb = sqlite3.connect("data_bd.db")
                cursbd = datadb.cursor()
                cursbd.execute(
                    "SELECT * FROM bdtable_bulletin_sortie WHERE num_bulletinbd AND datebd=? AND nom_producteurbd = ? AND perebd = ?  AND quantite=?",
                    (datebd, nomproducteurbd, perebd, quantitebd))
                existing_data = cursbd.fetchone()
                if existing_data:
                    msgbox = QtWidgets.QMessageBox()
                    msgbox.setWindowTitle('Confirmation')
                    msgbox.setText('Le Bulletin sortie  déjà.')
                    msgbox.exec()
                    # tempfilebdd = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                    # self.docbd.save( tempfilebdd)
                    # os.startfile(tempfilebdd, "print")
                    # tempfilebd = nomproducteurbd + "-" +current_days+ "B.D" + ".docx"
                    # self.docbd.save("bulletin agreage/bulletin_blédur/" + tempfilebd)
                    # pathbd = os.path.abspath("bulletin agreage/bulletin_blédur/" + tempfilebd)
                else:
                    doc_names = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                    doc_pdf = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False).name
                    self.docbd.save(doc_names)
                    try:
                        if doc_names:
                            a = self.progress_bar()
                            sys.stderr = open("consoleoutput.log", "w")
                            convert(doc_names, doc_pdf)
                            # Open the resulting .pdf file using the default associated application
                            # os.startfile(doc_pdf, 'open')
                            app_path = 'C:\\Program Files\\Okular\\bin\\okular.exe'
                            subprocess.Popen([app_path, doc_pdf])
                    except Exception as e:
                        print(e)
                    tempfilebd = nomproducteurbd + "-" + current_days + "B.D" + ".docx"
                    self.docbd.save("bulletin agreage/bulletin_blédur/" + tempfilebd)
                    pathbd = os.path.abspath("bulletin agreage/bulletin_blédur/" + tempfilebd)
                    databasebd = sqlite3.connect("data_bd.db")
                    curs = databasebd.cursor()
                    curs.execute(
                        "INSERT INTO bdtable_bulletin_sortie (datebd, nom_producteurbd, perebd,quantite, pathbd,date_sortie) VALUES (?,?,?, ?, ?, ?)",
                        (datebd, nomproducteurbd, perebd, quantitebd, pathbd, current_day,))
                    databasebd.commit()
                    databasebd.close()
            except Exception as e:
                print(e)

        def add_datta_bd(self):
            try:
                nbulltinbd = self.n_bultinbd.text()
                nomproducteurbd = self.nome_du_producteurbd.text()
                perebd = self.pérebd.text()
                matriculbd = self.adressebd.text()
                poindecollectebd = self.pointdecollectbd.currentText()
                datebd = self.dattereceptiontbd.dateTime().toString("MM-dd-yyyy-hh:mm")
                agreeurbd = self.agréeeurcombobd.currentText()
                quantitebd = self.quantitebd.value()
                vpsbd = self.vpsbd.value()
                bpsbd = self.bpsbd.text()
                rpsbd = self.rpsbd.text()
                vhumidite = self.vhumiditebd.value()
                vergot = self.vergotbd.value()
                vmatier20mm = self.vmatiére20mmbd.value()
                vdebrit = self.vdébrisvébd.value()
                vgrainnuisible = self.vgrainnuisiblebd.value()
                graincasse = self.vgrainscassébd.value()
                rgraicasse = self.rgrainscassébd.text()
                grainboute = self.vgrainsboutésbd.value()
                grainroux = self.vgrainsrouxbd.value()
                grainfortementmouchte = self.vgrainfortementmouchtébd.value()
                grainpunaise = self.vgrainepunaisésbd.value()
                rgrainpunaise = self.rgrainepunaisésbd.text()
                grainpique = self.vgrainpiquebd.value()
                grainmitadine = self.vgrainmitadinésbd.value()
                bletendedbledur = self.vbletendredansbledurbd.value()
                vtotalbdp = self.vtotalprembd.value()
                vtotalbdd = self.vtotaldembd.value()
                btotalbdp = self.btotalprembd.text()
                btotalbdd = self.btotaldembd.text()
                rtotalbdp = self.rtotalprembd.text()
                rtotalbdd = self.rtotaldembd.text()
                bmitadine = self.bgrainmitadinésbd.text()
                rmitadine = self.rgrainmitadinésbd.text()
                vtotalmitadineetbletendre = self.vtotalmitadinésbd.value()
                btotaldb = self.btotalbiniEtrefaction.text()
                rtotalbd = self.rtotalbiniEtrefaction.text()
                ndcartebd = self.n_cartebd.text()
                observationbd = self.obesrvationbd.toPlainText()
                date = QDate.currentDate()
                current_date = QDate.currentDate()
                current_day = current_date.day()
                current_month = current_date.month()
                current_year = current_date.year()
                current_time = QTime.currentTime()
                current_hour = current_time.hour()
                current_minute = current_time.minute()
                # Format the date as "dd-mm-yyyy"
                current_days = f"{current_day:02d}-{current_month:02d}-{current_year}-{current_hour}-{current_minute}"
                self.docbd = DocxTemplate("bulletin agreage/bulletin_blédur/bulletin_blédur.docx")
                self.docbd.render({
                    "nmp": nomproducteurbd,
                    "per": perebd,
                    "mat": matriculbd,
                    "ptt": poindecollectebd,
                    "dt": datebd,
                    "agr": agreeurbd,
                    "qtt": quantitebd,
                    "vps": vpsbd,
                    "bps": bpsbd,
                    "rps": rpsbd,
                    "vmh": vhumidite,
                    "vrg": vergot,
                    "vgs": vmatier20mm,
                    "vdb": vdebrit,
                    "vgn": vgrainnuisible,
                    "vgc": graincasse,
                    "rgc": rgraicasse,
                    "vgfb": grainboute,
                    "vgr": grainroux,
                    "vgfm": grainfortementmouchte,
                    "vgpn": grainpunaise,
                    "rgpn": rgrainpunaise,
                    "vgrp": grainpique,
                    "vmt": grainmitadine,
                    "vbtb": bletendedbledur,
                    "vttp": vtotalbdp,
                    "bttp": btotalbdp,
                    "rtt": rtotalbdp,
                    "vtd": vtotalbdd,
                    "btd": btotalbdd,
                    "rtd": rtotalbdd,
                    "btb": btotaldb,
                    "rtr": rtotalbd,
                    "bmt": bmitadine,
                    "rmt": rmitadine,
                    "vttm": vtotalmitadineetbletendre,
                    "ncn": ndcartebd,
                    "oo": observationbd,
                    "nm": nbulltinbd
                })
                tempfilebd = nomproducteurbd + "-" + current_days + "-" + "B.D" + ".docx"
                self.docbd.save("bulletin agreage/bulletin_blédur/" + tempfilebd)
                pathbd = os.path.abspath("bulletin agreage/bulletin_blédur/" + tempfilebd)

                # Check if the data already exists
                datadb = sqlite3.connect("data_bd.db")
                cursbd = datadb.cursor()
                cursbd.execute(
                    "SELECT * FROM bdtable WHERE datebd=? AND nom_producteurbd=? AND perebd=? AND matriculebd=? AND point_collectbd=? AND nom_agreeurbd=? AND quantitebd=? AND psbd=? AND totalbd=? AND pathbd=? ",
                    (datebd, nomproducteurbd, perebd, matriculbd, poindecollectebd, agreeurbd, quantitebd, vpsbd,
                     vtotalbdp, pathbd,))
                existing_data = cursbd.fetchone()

                if existing_data:
                    msgbox = QtWidgets.QMessageBox()
                    msgbox.setWindowTitle('confirmation')
                    msgbox.setText('Les données existent déjà.')
                    msgbox.exec()
                else:
                    cursbd.execute(
                        "INSERT INTO bdtable (datebd ,nom_producteurbd ,perebd ,matriculebd ,point_collectbd ,nom_agreeurbd ,quantitebd ,psbd ,totalbd,pathbd  )  values(?,?,?,?,?,?,?,?,?,?)",
                        (datebd, nomproducteurbd, perebd, matriculbd, poindecollectebd, agreeurbd, quantitebd, vpsbd,
                         vtotalbdp, pathbd,))
                    datadb.commit()
                    datadb.close()

                    msgbox = QtWidgets.QMessageBox()
                    msgbox.setWindowTitle('confirmation')
                    msgbox.setText('Les données sont sauvegardées')
                    msgbox.exec()

            except sqlite3.Error as e:
                print("Error accessing the database:", e)

        def number_bulletin_bd(self):
            try:
                databasebd = sqlite3.connect("data_bd.db")
                cursor = databasebd.cursor()
                cursor.execute("SELECT num_bulletinbd FROM bdtable ORDER BY num_bulletinbd DESC LIMIT 1")
                resultbd = cursor.fetchone()
                if resultbd is not None:
                    last_id = resultbd[0]
                    self.n_bultinbd.setText("DC-" + str(last_id))
                else:
                    # Handle the case when there are no records in the orgetable
                    self.n_bultinbd.setText("DC-")
            except sqlite3.Error as e:
                print("Error accessing the database:", e)

        def print_day_bd(self):
            try:
                dialog = QtWidgets.QMessageBox()
                dialog.setStyleSheet("""
                                          QWidget {
                                              color: #000000;
                                              background-color: #ffc6c9;
                                              border-width: 1px;
                                              border-color: #1e1e1e;
                                              border-style: solid;
                                              border-radius: 6;
                                              padding: 3px;
                                              font-size: 18px;
                                              padding-left: 5px;
                                              padding-right: 5px;
                                          }
                                          QWidget:item:hover {
                                              background-color: #3daee9;
                                              color: #eff0f1;
                                          }
                                          QWidget:item:selected {
                                              background-color: #3daee9;
                                          }
                                          QWidget:disabled {
                                              color: #454545;
                                              background-color: #31363b;
                                          }
                                          QPushButton {
                                              color: #000000;
                                              background-color:#84dbc8;
                                              border-width: 1px;
                                              border-color: #1e1e1e;
                                              border-style: solid;
                                              border-radius: 6;
                                              padding: 3px;
                                              font-size: 12px;
                                              padding-left: 5px;
                                              padding-right: 5px;
                                              min-width: 40px;
                                          }
                                          QPushButton:disabled {
                                              background-color: #31363b;
                                              border-width: 1px;
                                              border-color: #454545;
                                              border-style: solid;
                                              padding-top: 5px;
                                              padding-bottom: 5px;
                                              padding-left: 10px;
                                              padding-right: 10px;
                                              border-radius: 2px;
                                              color: #454545;
                                          }
                                          QPushButton:pressed {
                                              background-color: #3daee9;
                                              padding-top: -15px;
                                              padding-bottom: -17px;
                                          }
                                          QPushButton:hover {
                                              border: 1px solid #ff8c00;
                                              color: #eff0f1;
                                          }
                                          QLabel {
                                              font-size: 18px;
                                              border: 0px solid orange;
                                          }
                                      """)
                dialog.setWindowTitle("Sélectionnez la date")
                dialog.setText("Sélectionnez la date  \t\n")
                self.datebd = QtWidgets.QDateEdit(dialog)
                self.datebd.setDate(self.datedaytime)
                self.datebd.resize(180, 30)
                self.datebd.move(40, 50)
                ok_button = QtWidgets.QPushButton("OK", dialog)
                cancel_button = QtWidgets.QPushButton("Cancel", dialog)
                dialog.addButton(ok_button, QMessageBox.ButtonRole.ActionRole)
                dialog.addButton(cancel_button, QMessageBox.ButtonRole.ActionRole)
                dialog.exec()

                if dialog.clickedButton() == ok_button:
                    self.selected_datebd = self.datebd.text()
                    if self.datebd.text() == self.selected_datebd:

                        conn = sqlite3.connect('data_bd.db')
                        cursor = conn.cursor()
                        cursor.execute(
                            "SELECT num_bulletinbd,datebd ,nom_producteurbd ,perebd ,matriculebd ,point_collectbd  ,quantitebd  FROM bdtable WHERE DATE(substr(datebd, 7, 4) || '-' || substr(datebd, 4, 2) || '-' || substr(datebd, 1, 2)) =? ;",
                            (self.selected_datebd,))
                        rows = cursor.fetchall()

                        ########################################################################
                        # Create a new document and add a table
                        self.doc = docx.Document()
                        section = self.doc.sections[0]
                        section.page_width = docx.shared.Cm(29.7)
                        section.page_height = docx.shared.Cm(21.0)
                        section.top_margin = docx.shared.Cm(1.5)
                        section.bottom_margin = docx.shared.Cm(1.5)
                        heading1 = self.doc.add_heading("\t\t\t\t\t\t les entre de Blé Dur ", level=1)
                        heading1.style.font.name = 'Times New Roman'
                        heading1.style.font.size = Pt(22)
                        heading1.style.font.bold = True
                        heading1.style.font.color.rgb = RGBColor(0, 0, 0)
                        # add second heading
                        paragraph = self.doc.add_paragraph()
                        paragraph.style.font.name = 'Times New Roman'
                        paragraph.style.font.size = Pt(14)
                        left_run = paragraph.add_run('CCLS RELIZANE\t\t\t\t\t\t\t\t\t\t\t\t')
                        left_run.bold = True
                        right_run = paragraph.add_run('Date:' + self.selected_datebd)
                        right_run.bold = True

                        table = self.doc.add_table(rows=1, cols=9)
                        table.style = "Table Grid"  # set the table style
                        hdr_cells = table.rows[0].cells
                        hdr_cells[0].text = 'N° Bulletin \nDC'
                        hdr_cells[1].text = 'Date'
                        hdr_cells[2].text = 'Nom et Prénom'
                        hdr_cells[3].text = 'Pére'
                        hdr_cells[4].text = 'Matricule'
                        hdr_cells[5].text = 'point de collect'
                        hdr_cells[6].text = 'Quantité'
                        hdr_cells[7].text = 'Date de sortie '
                        hdr_cells[8].text = 'N° de la carte '

                        # Set the width of the header cells
                        hdr_cells[0].width = Inches(1.5)  # set the width of the first column to 1.5 inches
                        hdr_cells[1].width = Inches(1.5)
                        hdr_cells[2].width = Inches(2)
                        hdr_cells[3].width = Inches(1)
                        hdr_cells[4].width = Inches(1.5)
                        hdr_cells[5].width = Inches(1.5)
                        hdr_cells[6].width = Inches(1)
                        hdr_cells[7].width = Inches(1.5)
                        hdr_cells[8].width = Inches(1.5)

                        # set hight of the column
                        hdr_cells[0].height = Inches(1)  # set the width of the first column to 1.5 inches
                        hdr_cells[1].height = Inches(1)
                        hdr_cells[2].height = Inches(1)
                        hdr_cells[3].height = Inches(1)
                        hdr_cells[4].height = Inches(1)
                        hdr_cells[5].height = Inches(1)
                        hdr_cells[6].height = Inches(1)
                        hdr_cells[7].width = Inches(1)
                        hdr_cells[8].width = Inches(1)

                        for cell in table.rows[0].cells:
                            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                            cell.paragraphs[0].runs[0].font.size = Pt(12)
                            cell.paragraphs[0].runs[0].font.bold = True
                        # Add the data to the table
                        table_rows = len(rows)
                        table_cols = len(rows[0])
                        for row in range(table_rows):
                            table.add_row()
                            for col in range(table_cols):
                                cell = table.cell(row + 1, col)
                                cell.text = str(rows[row][col])
                                # Set font properties
                                cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                                cell.paragraphs[0].runs[0].font.size = Pt(12)
                                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                        for row in table.rows:
                            row.height = Inches(0.3)
                        # Save and open the document for printing
                        doc_names = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                        doc_pdf = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False).name
                        self.doc.save(doc_names)
                        try:
                            if doc_names:
                                a = self.progress_bar()
                                sys.stderr = open("consoleoutput.log", "w")
                                convert(doc_names, doc_pdf)
                                # Open the resulting .pdf file using the default associated application
                                # os.startfile(doc_pdf, 'open')
                                app_path = 'C:\\Program Files\\Okular\\bin\\okular.exe'
                                subprocess.Popen([app_path, doc_pdf])
                        except Exception as e:
                            print(e)
                    else:
                        msgbox = QtWidgets.QMessageBox()
                        msgbox.setWindowTitle('Erreur')
                        msgbox.setText("Il n'ya pas des données à cette date.")
                        msgbox.exec()
                if dialog.clickedButton() == cancel_button:
                    dialog.close()
            except Exception as e:
                print(e)

        def printBulltinProducteur(self):
            dialogor = QMessageBox()
            dialogor.setFixedSize(300, 600)
            dialogor.setStyleSheet("""
                                                  QWidget {
                                                      color: #000000;
                                                      background-color: #ffc6c9;
                                                      border-width: 1px;
                                                      border-color: #1e1e1e;
                                                      border-style: solid;
                                                      border-radius: 6;
                                                      padding: 3px;
                                                      font-size: 18px;
                                                      padding-left: 5px;
                                                      padding-right: 5px;
                                                  }
                                                  QWidget:item:hover {
                                                      background-color: #3daee9;
                                                      color: #eff0f1;
                                                  }
                                                  QWidget:item:selected {
                                                      background-color: #3daee9;
                                                  }
                                                  QWidget:disabled {
                                                      color: #454545;
                                                      background-color: #31363b;
                                                  }
                                                  QPushButton {
                                                      color: #000000;
                                                      background-color:#84dbc8;
                                                      border-width: 1px;
                                                      border-color: #1e1e1e;
                                                      border-style: solid;
                                                      border-radius: 6;
                                                      padding: 3px;
                                                      font-size: 12px;
                                                      padding-left: 5px;
                                                      padding-right: 5px;
                                                      min-width: 40px;
                                                  }
                                                  QPushButton:disabled {
                                                      background-color: #31363b;
                                                      border-width: 1px;
                                                      border-color: #454545;
                                                      border-style: solid;
                                                      padding-top: 5px;
                                                      padding-bottom: 5px;
                                                      padding-left: 10px;
                                                      padding-right: 10px;
                                                      border-radius: 2px;
                                                      color: #454545;
                                                  }
                                                  QPushButton:pressed {
                                                      background-color: #3daee9;
                                                      padding-top: -15px;
                                                      padding-bottom: -17px;
                                                  }
                                                  QPushButton:hover {
                                                      border: 1px solid #ff8c00;
                                                      color: #eff0f1;
                                                  }
                                                  QLabel {
                                                      font-size: 18px;
                                                      border: 0px solid orange;
                                                  }
                                              """)
            dialogor.setWindowTitle("Sélectionnez le Nom et Prénom et pére")
            dialogor.setText("Sélectionnez le Nom et Prénom et pére  \t\t\t\t\t\t\n\t\t\t\t\t\t\t\t\t\t\n\t\t\t\t")
            self.nomproducturep = QtWidgets.QLineEdit(dialogor)
            self.nomproducturep.setPlaceholderText("Nom et Prénom")
            self.nomproducturep.resize(180, 30)
            self.nomproducturep.move(40, 40)
            self.pérep = QtWidgets.QLineEdit(dialogor)
            self.pérep.setPlaceholderText("Pére")
            self.pérep.resize(180, 30)
            self.pérep.move(240, 40)
            self.datetxt = QtWidgets.QLabel("Date entré", dialogor)
            self.datetxt.setGeometry(QtCore.QRect(30, 80, 100, 20))
            datep = QDateTimeEdit(dialogor)
            datep.setDisplayFormat("dd-MM-yyyy hh:mm")
            datep.resize(180, 30)
            datep.move(40, 105)
            self.quantiteptext = QtWidgets.QLabel("Quantité", dialogor)
            self.quantiteptext.setGeometry(QtCore.QRect(230, 80, 100, 20))
            self.quantitep = QtWidgets.QDoubleSpinBox(dialogor)
            self.quantitep.setRange(1, 10000)
            self.quantitep.setSuffix('  QX')
            self.quantitep.setSpecialValueText(" ")
            self.quantitep.setValue(0.0)
            self.quantitep.resize(180, 30)
            self.quantitep.move(240, 105)
            ok_button = QtWidgets.QPushButton("OK", dialogor)
            cancel_button = QtWidgets.QPushButton("Cancel", dialogor)
            dialogor.addButton(ok_button, QMessageBox.ButtonRole.ActionRole)
            dialogor.addButton(cancel_button, QMessageBox.ButtonRole.ActionRole)
            dialogor.exec()
            date = QDate.currentDate()
            current_date = QDate.currentDate()
            current_day = current_date.day()
            current_month = current_date.month()
            current_year = current_date.year()

            # Format the date as "dd-mm-yyyy"
            current_day = f"{current_day:02d}-{current_month:02d}-{current_year}"

            if dialogor.clickedButton() == ok_button:
                producteur = self.nomproducturep.text()
                pere = self.pérep.text()
                dateentré = datep.text()

                quantite = self.quantitep.value()
                try:
                    datadb = sqlite3.connect("data_bd.db")
                    cursbd = datadb.cursor()
                    cursbd.execute(
                        "SELECT * FROM bdtable_bulletin_sortie WHERE datebd=? AND nom_producteurbd=? AND perebd=? AND quantite=?",
                        (dateentré, producteur, pere, quantite))
                    existing_data = cursbd.fetchone()
                    if existing_data:
                        msgbox = QtWidgets.QMessageBox()
                        msgbox.setWindowTitle('Confirmation')
                        msgbox.setText('Le Bulletin sortie existe déjà.')
                        msgbox.exec()
                    else:
                        print("Conditions:", dateentré, producteur, pere, quantite)

                        try:
                            curs = datadb.cursor()
                            curs.execute(
                                "SELECT pathbd FROM bdtable WHERE datebd=? AND  nom_producteurbd=? AND perebd=? AND quantitebd=? ",
                                (dateentré, producteur, pere, quantite,)
                            )
                            result = curs.fetchall()
                            print("Result:", result)
                        except sqlite3.Error as e:
                            print(f"Error retrieving data from bdtable: {e}")

                        if result:
                            # Iterate over the paths and open each file
                            for row in result:
                                file_path = row[0]
                                if file_path:
                                    doc_pdf = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False).name
                                    try:
                                        a = self.progress_bar()
                                        sys.stderr = open("consoleoutput.log", "w")
                                        convert(file_path, doc_pdf)
                                        # Open the resulting .pdf file using the default associated application
                                        # os.startfile(doc_pdf, 'open')
                                        app_path = 'C:\\Program Files\\Okular\\bin\\okular.exe'
                                        subprocess.Popen([app_path, doc_pdf])
                                    except Exception as e:
                                        print(e)
                                    # os.startfile(file_path, "open")
                                    curs.execute(
                                        "INSERT INTO bdtable_bulletin_sortie (datebd, nom_producteurbd, perebd, quantite, pathbd, date_sortie) VALUES (?,?,?,?,?,?)",
                                        (dateentré, producteur, pere, quantite, file_path, current_day))
                                    datadb.commit()
                                    datadb.close()


                        else:
                            msgbox = QtWidgets.QMessageBox()
                            msgbox.setWindowTitle('Confirmation')
                            msgbox.setText('Le Bulletin ne existe pas.')
                            msgbox.exec()
                except Exception as e:
                    print(e)

        def eticket_bledur(self):
            try:
                dialogor = QtWidgets.QMessageBox()
                dialogor.setFixedSize(300, 600)
                dialogor.setStyleSheet("""
                                                  QWidget {
                                                      color: #000000;
                                                      background-color: #ffc6c9;
                                                      border-width: 1px;
                                                      border-color: #1e1e1e;
                                                      border-style: solid;
                                                      border-radius: 6;
                                                      padding: 3px;
                                                      font-size: 18px;
                                                      padding-left: 5px;
                                                      padding-right: 5px;
                                                  }
                                                  QWidget:item:hover {
                                                      background-color: #3daee9;
                                                      color: #eff0f1;
                                                  }
                                                  QWidget:item:selected {
                                                      background-color: #3daee9;
                                                  }
                                                  QWidget:disabled {
                                                      color: #454545;
                                                      background-color: #31363b;
                                                  }
                                                  QPushButton {
                                                      color: #000000;
                                                      background-color:#84dbc8;
                                                      border-width: 1px;
                                                      border-color: #1e1e1e;
                                                      border-style: solid;
                                                      border-radius: 6;
                                                      padding: 3px;
                                                      font-size: 12px;
                                                      padding-left: 5px;
                                                      padding-right: 5px;
                                                      min-width: 40px;
                                                  }
                                                  QPushButton:disabled {
                                                      background-color: #31363b;
                                                      border-width: 1px;
                                                      border-color: #454545;
                                                      border-style: solid;
                                                      padding-top: 5px;
                                                      padding-bottom: 5px;
                                                      padding-left: 10px;
                                                      padding-right: 10px;
                                                      border-radius: 2px;
                                                      color: #454545;
                                                  }
                                                  QPushButton:pressed {
                                                      background-color: #3daee9;
                                                      padding-top: -15px;
                                                      padding-bottom: -17px;
                                                  }
                                                  QPushButton:hover {
                                                      border: 1px solid #ff8c00;
                                                      color: #eff0f1;
                                                  }
                                                  QLabel {
                                                      font-size: 18px;
                                                      border: 0px solid orange;
                                                  }
                                                  QLineEdit
            {
                background-color: #ffffff;
                padding: 1px;
                border-style: solid;
                border: 1px solid #76797C;
                border-radius: 0px;
                color: #000000;
            }
            QDoubleSpinBox
            {
                background-color: #ffffff;
                padding: 1px;
                border-style: solid;
                border: 1px solid #76797C;
                border-radius: 0px;
                color:#000000;
                font-size: 11px;
                font-weight: bold;

            }
            QComboBox
            {
               background-color: #ffffff;
                border-style: solid;
                border: 1px solid #76797C;
                border-radius: 2px;
                min-width: 40px;
            }
                                              """)
                dialogor.setWindowTitle("Sélectionnez le Nom et Prénom et pére")
                dialogor.setText(
                    "BLE DUR \t\t\t\n\t\t\t\t\t\t\n\t\t\t\t\t\t\t\n\t\t\t\t\n\t\t\t\t\t\t\t\t\t\t\n\t\t\t\t")
                nomproductureor = QtWidgets.QLineEdit(dialogor)
                nomproductureor.setPlaceholderText("Nom et Prénom")
                nomproductureor.resize(180, 30)
                nomproductureor.move(40, 125)
                pereor = QtWidgets.QLineEdit(dialogor)
                pereor.setPlaceholderText("Pére")
                pereor.resize(180, 30)
                pereor.move(240, 125)
                matricultxt = QtWidgets.QLabel("Matricul", dialogor)
                matricultxt.setGeometry(QtCore.QRect(30, 165, 100, 20))
                mator = QtWidgets.QLineEdit(dialogor)
                mator.setInputMask("99999-999-99")
                mator.resize(180, 30)
                mator.move(40, 190)
                pointdecollector = QtWidgets.QComboBox(dialogor)
                pointdecollector.setPlaceholderText("Point de collect")
                pointdecollector.addItem('')
                pointdecollector.addItem('Dock central')
                pointdecollector.addItem('Magasin Zemmoura')
                pointdecollector.addItem('Magasin Kef-lazreg')
                pointdecollector.addItem('Magasin Messra')
                pointdecollector.addItem('Magasin Marche gros belacel')
                pointdecollector.resize(180, 30)
                pointdecollector.move(40, 65)

                psor = QtWidgets.QLineEdit(dialogor)
                validator = QDoubleValidator(10, 1000, 2)
                psor.setValidator(validator)
                # psor.setInputMask("99.99")
                psor.setPlaceholderText("PS")
                psor.resize(180, 30)
                psor.move(240, 65)
                quantiteptext = QtWidgets.QLabel("Quantité", dialogor)
                quantiteptext.setGeometry(QtCore.QRect(230, 165, 100, 20))
                quantiteor = QtWidgets.QDoubleSpinBox(dialogor)
                quantiteor.setRange(1, 10000)
                quantiteor.setSuffix('  QX')
                quantiteor.setSpecialValueText(" ")
                quantiteor.setValue(0.0)
                quantiteor.resize(180, 30)
                quantiteor.move(240, 190)
                ok_button = QtWidgets.QPushButton("OK", dialogor)
                cancel_button = QtWidgets.QPushButton("Cancel", dialogor)
                dialogor.addButton(ok_button, QtWidgets.QMessageBox.ButtonRole.ActionRole)
                dialogor.addButton(cancel_button, QtWidgets.QMessageBox.ButtonRole.ActionRole)
                dialogor.exec()
                current_date = QtCore.QDate.currentDate().toString(QtCore.QDate.currentDate().toString("yyyy-MM-dd"))
                current_day = current_date.split("-")[2]
                current_month = current_date.split("-")[1]
                current_year = current_date.split("-")[0]
                current_daya = f"{current_day}-{current_month}-{current_year}"

                if dialogor.clickedButton() == ok_button:
                    producteur = nomproductureor.text()
                    pere = pereor.text()
                    quantite = quantiteor.text()
                    pointdecolect = pointdecollector.currentText()
                    mator = mator.text()
                    psor = psor.text()

                    eticketorge = DocxTemplate("eticket/Eticket_BleDur/Eticket_BleDur.docx")
                    eticketorge.render(
                        {"dt": current_date, "ptt": pointdecolect, "nmp": producteur, "mat": mator, "qtt": quantite,
                         "ps": psor, "per": pere})
                    name = f"B.D_{producteur}_{current_daya}.docx"
                    eticketorge.save("eticket/Eticket_BleDur/" + name)
                    doc_names = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                    doc_pdf = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False).name
                    eticketorge.save(doc_names)
                    try:
                        if doc_names:
                            a = self.progress_bar()
                            sys.stderr = open("consoleoutput.log", "w")
                            convert(doc_names, doc_pdf)
                            # Open the resulting .pdf file using the default associated application
                            # os.startfile(doc_pdf, 'open')
                            app_path = 'C:\\Program Files\\Okular\\bin\\okular.exe'
                            subprocess.Popen([app_path, doc_pdf])
                    except Exception as e:
                        print(e)
            except Exception as e:
                print(e)

        ############################################################################################
        ##############################function blé tendre ##########################################
        ############################################################################################
        def bonification_ps_bt(self):
            self.bpsbt.clear()
            self.bpsbt.setStyleSheet("background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding: 0px")
            while self.vpsbt.value() >= 77.01:
                self.bpsbt.setValue(0.10)
                self.bpsbt.setStyleSheet(
                    "background-color:#88ffaa;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                if self.vpsbt.value() <= 77.25:
                    self.bpsbt.setValue(0.10)
                elif self.vpsbt.value() >= 77.26 and self.vpsbt.value() <= 77.50:
                    self.bpsbt.setValue(0.20)
                elif self.vpsbt.value() >= 77.51 and self.vpsbt.value() <= 77.75:
                    self.bpsbt.setValue(0.30)
                elif self.vpsbt.value() >= 77.76 and self.vpsbt.value() <= 78.00:
                    self.bpsbt.setValue(0.40)
                elif self.vpsbt.value() >= 78.01 and self.vpsbt.value() <= 78.25:
                    self.bpsbt.setValue(0.45)
                elif self.vpsbt.value() >= 78.26 and self.vpsbt.value() <= 78.50:
                    self.bpsbt.setValue(0.50)
                elif self.vpsbt.value() >= 78.51 and self.vpsbt.value() <= 78.75:
                    self.bpsbt.setValue(0.55)
                elif self.vpsbt.value() >= 78.76 and self.vpsbt.value() <= 79.00:
                    self.bpsbt.setValue(0.60)
                elif self.vpsbt.value() >= 79.01 and self.vpsbt.value() <= 79.25:
                    self.bpsbt.setValue(0.65)
                elif self.vpsbt.value() >= 79.26 and self.vpsbt.value() <= 79.50:
                    self.bpsbt.setValue(0.70)
                elif self.vpsbt.value() >= 79.51 and self.vpsbt.value() <= 79.75:
                    self.bpsbt.setValue(0.75)
                elif self.vpsbt.value() >= 79.76 and self.vpsbt.value() <= 80.00:
                    self.bpsbt.setValue(0.80)
                elif self.vpsbt.value() >= 80.01 and self.vpsbt.value() <= 80.25:
                    self.bpsbt.setValue(0.82)
                elif self.vpsbt.value() >= 80.26 and self.vpsbt.value() <= 80.50:
                    self.bpsbt.setValue(0.84)
                elif self.vpsbt.value() >= 80.51 and self.vpsbt.value() <= 80.75:
                    self.bpsbt.setValue(0.86)
                elif self.vpsbt.value() >= 80.76 and self.vpsbt.value() <= 81.00:
                    self.bpsbt.setValue(0.88)
                elif self.vpsbt.value() >= 81.01 and self.vpsbt.value() <= 81.25:
                    self.bpsbt.setValue(0.90)
                elif self.vpsbt.value() >= 81.26 and self.vpsbt.value() <= 81.50:
                    self.bpsbt.setValue(0.92)
                elif self.vpsbt.value() >= 81.51 and self.vpsbt.value() <= 81.75:
                    self.bpsbt.setValue(0.94)
                elif self.vpsbt.value() >= 81.76 and self.vpsbt.value() <= 82.00:
                    self.bpsbt.setValue(0.96)
                elif self.vpsbt.value() >= 82.01 and self.vpsbt.value() <= 82.25:
                    self.bpsbt.setValue(0.98)
                elif self.vpsbt.value() >= 82.26 and self.vpsbt.value() <= 82.50:
                    self.bpsbt.setValue(1.00)
                elif self.vpsbt.value() >= 82.51 and self.vpsbt.value() <= 82.75:
                    self.bpsbt.setValue(1.02)
                elif self.vpsbt.value() >= 82.76 and self.vpsbt.value() <= 83.00:
                    self.bpsbt.setValue(1.04)
                else:
                    self.bpsbt.setValue(0)
                    self.bpsbt.setStyleSheet(
                        "background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                break

        def réfaction_ps_bt(self):
            self.rpsbt.clear()
            self.rpsbt.setStyleSheet("background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding: 0px")
            if self.vpsbt.value() < 69 and self.vpsbt.value() > 50:
                self.obesrvationbt.setText("REFUS: POIDS SPECIFIQUE INFERIEUR DE 72 Kg/hl ")
                self.obesrvationbt.setStyleSheet(
                    "background-color:#e22630;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                self.bpsbt.clear()
                self.rpsbt.clear()
            else:
                self.obesrvationbt.clear()
                self.obesrvationbt.setStyleSheet(
                    "background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding: 0px")
            while self.vpsbt.value() >= 69.00:
                self.rpsbt.setValue(1.36)
                self.rpsbt.setStyleSheet(
                    "background-color:#e22630;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                self.bpsbt.clear()
                if self.vpsbt.value() <= 69.24:
                    self.rpsbt.setValue(1.36)
                elif self.vpsbt.value() >= 69.25 and self.vpsbt.value() <= 69.49:
                    self.rpsbt.setValue(1.16)
                elif self.vpsbt.value() >= 69.50 and self.vpsbt.value() <= 69.74:
                    self.rpsbt.setValue(0.96)
                elif self.vpsbt.value() >= 69.75 and self.vpsbt.value() <= 69.99:
                    self.rpsbt.setValue(0.76)
                elif self.vpsbt.value() >= 70.00 and self.vpsbt.value() <= 71.24:
                    self.rpsbt.setValue(0.56)
                elif self.vpsbt.value() >= 71.25 and self.vpsbt.value() <= 71.49:
                    self.rpsbt.setValue(0.46)
                elif self.vpsbt.value() >= 71.50 and self.vpsbt.value() <= 71.74:
                    self.rpsbt.setValue(0.36)
                elif self.vpsbt.value() >= 71.75 and self.vpsbt.value() <= 71.99:
                    self.rpsbt.setValue(0.26)
                elif self.vpsbt.value() >= 72.00 and self.vpsbt.value() <= 73.24:
                    self.rpsbt.setValue(0.16)
                elif self.vpsbt.value() >= 73.25 and self.vpsbt.value() <= 73.49:
                    self.rpsbt.setValue(0.12)
                elif self.vpsbt.value() >= 73.50 and self.vpsbt.value() <= 73.74:
                    self.rpsbt.setValue(0.08)
                elif self.vpsbt.value() >= 73.75 and self.vpsbt.value() <= 73.99:
                    self.rpsbt.setValue(0.04)
                else:
                    self.rpsbt.setValue(0)
                    self.rpsbt.setStyleSheet(
                        "background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                break

        def bonification_pimpurte_er_bt(self):
            self.btotalprembt.clear()
            self.btotalprembt.setStyleSheet(
                "background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding: 0px")
            while self.vtotalprembt.value() >= 0.01:
                self.btotalprembt.setValue(0.12)
                self.btotalprembt.setStyleSheet(
                    "background-color:#88ffaa;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                if self.vtotalprembt.value() <= 0.25:
                    self.btotalprembt.setValue(0.12)
                elif self.vtotalprembt.value() >= 0.26 and self.vtotalprembt.value() <= 0.50:
                    self.btotalprembt.setValue(0.24)
                elif self.vtotalprembt.value() >= 0.51 and self.vtotalprembt.value() <= 0.75:
                    self.btotalprembt.setValue(0.36)
                elif self.vtotalprembt.value() >= 0.76 and self.vtotalprembt.value() <= 1.00:
                    self.btotalprembt.setValue(0.48)
                else:
                    self.btotalprembt.setValue(0)
                    self.btotalprembt.clear()
                    self.btotalprembt.setStyleSheet(
                        "background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                break

        def refaction_impurté_er_bt(self):
            self.rtotalprembt.clear()
            self.rtotalprembt.setStyleSheet(
                "background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding: 0px")
            while self.vtotalprembt.value() >= 3.01:
                self.rtotalprembt.setValue(0.12)
                self.rtotalprembt.setStyleSheet(
                    "background-color:#e22630;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                self.btotalprembt.clear()
                if self.vtotalprembt.value() <= 3.25:
                    self.rtotalprembt.setValue(0.12)
                elif self.vtotalprembt.value() >= 3.26 and self.vtotalprembt.value() <= 3.50:
                    self.rtotalprembt.setValue(0.24)
                elif self.vtotalprembt.value() >= 3.51 and self.vtotalprembt.value() <= 3.75:
                    self.rtotalprembt.setValue(0.36)
                elif self.vtotalprembt.value() >= 3.76 and self.vtotalprembt.value() <= 4.00:
                    self.rtotalprembt.setValue(0.48)
                elif self.vtotalprembt.value() >= 4.01 and self.vtotalprembt.value() <= 4.25:
                    self.rtotalprembt.setValue(0.60)
                elif self.vtotalprembt.value() >= 4.26 and self.vtotalprembt.value() <= 4.50:
                    self.rtotalprembt.setValue(0.72)
                elif self.vtotalprembt.value() >= 4.51 and self.vtotalprembt.value() <= 4.75:
                    self.rtotalprembt.setValue(0.84)
                elif self.vtotalprembt.value() >= 4.76 and self.vtotalprembt.value() <= 5.00:
                    self.rtotalprembt.setValue(0.96)
                elif self.vtotalprembt.value() >= 5.01 and self.vtotalprembt.value() <= 5.25:
                    self.rtotalprembt.setValue(1.08)
                elif self.vtotalprembt.value() >= 8.51 and self.vtotalprembt.value() <= 9.00:
                    self.rtotalprembt.setValue(1.20)
                elif self.vtotalprembt.value() >= 5.26 and self.vtotalprembt.value() <= 5.50:
                    self.rtotalprembt.setValue(1.32)
                elif self.vtotalprembt.value() >= 5.51 and self.vtotalprembt.value() <= 5.75:
                    self.rtotalprembt.setValue(1.44)
                elif self.vtotalprembt.value() >= 5.76 and self.vtotalprembt.value() <= 6.00:
                    self.rtotalprembt.setValue(1.56)
                elif self.vtotalprembt.value() >= 6.01 and self.vtotalprembt.value() <= 30:
                    self.rtotalprembt.setValue(0)
                    self.obesrvationbt.setText("PRIX A DEBATTRE")

                else:
                    self.rtotalprembt.setValue(0)
                    self.rtotalprembt.clear
                    self.rtotalprembt.setStyleSheet(
                        "background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                break

        def refaction_impurte_eme_bt(self):
            self.rtotaldembt.clear()
            self.rtotaldembt.setStyleSheet(
                "background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding: 0px")
            while self.vtotaldembt.value() >= 6.01:
                self.rtotaldembt.setValue(0.05)
                self.rtotaldembt.setStyleSheet(
                    "background-color:#e22630;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                self.btotaldembt.clear()
                if self.vtotaldembt.value() <= 6.25:
                    self.rtotaldembt.setValue(0.05)
                elif self.vtotaldembt.value() >= 6.26 and self.vtotaldembt.value() <= 6.50:
                    self.rtotaldembt.setValue(0.10)
                elif self.vtotaldembt.value() >= 6.51 and self.vtotaldembt.value() <= 6.75:
                    self.rtotaldembt.setValue(0.15)
                elif self.vtotaldembt.value() >= 6.76 and self.vtotaldembt.value() <= 7.00:
                    self.rtotaldembt.setValue(0.20)
                elif self.vtotaldembt.value() >= 7.01 and self.vtotaldembt.value() <= 7.25:
                    self.rtotaldembt.setValue(0.25)
                elif self.vtotaldembt.value() >= 7.26 and self.vtotaldembt.value() <= 7.50:
                    self.rtotaldembt.setValue(0.30)
                elif self.vtotaldembt.value() >= 7.51 and self.vtotaldembt.value() <= 7.75:
                    self.rtotaldembt.setValue(0.35)
                elif self.vtotaldembt.value() >= 7.76 and self.vtotaldembt.value() <= 8.00:
                    self.rtotaldembt.setValue(0.40)
                elif self.vtotaldembt.value() >= 8.01 and self.vtotaldembt.value() <= 8.25:
                    self.rtotaldembt.setValue(0.45)
                elif self.vtotaldembt.value() >= 8.26 and self.vtotaldembt.value() <= 8.50:
                    self.rtotaldembt.setValue(0.50)
                elif self.vtotaldembt.value() >= 8.51 and self.vtotaldembt.value() <= 8.75:
                    self.rtotaldembt.setValue(0.55)
                elif self.vtotaldembt.value() >= 8.76 and self.vtotaldembt.value() <= 9.00:
                    self.rtotaldembt.setValue(0.60)
                elif self.vtotaldembt.value() >= 9.01 and self.vtotaldembt.value() <= 9.25:
                    self.rtotaldembt.setValue(0.65)
                elif self.vtotaldembt.value() >= 9.26 and self.vtotaldembt.value() <= 9.50:
                    self.rtotaldembt.setValue(0.70)
                elif self.vtotaldembt.value() >= 9.51 and self.vtotaldembt.value() <= 9.75:
                    self.rtotaldembt.setValue(0.75)
                elif self.vtotaldembt.value() >= 9.76 and self.vtotaldembt.value() <= 10.00:
                    self.rtotaldembt.setValue(0.80)
                elif self.vtotaldembt.value() >= 10.01 and self.vtotaldembt.value() <= 10.25:
                    self.rtotaldembt.setValue(0.88)
                elif self.vtotaldembt.value() >= 10.26 and self.vtotaldembt.value() <= 10.50:
                    self.rtotaldembt.setValue(0.96)
                elif self.vtotaldembt.value() >= 10.51 and self.vtotaldembt.value() <= 10.75:
                    self.rtotaldembt.setValue(1.04)
                elif self.vtotaldembt.value() >= 10.76 and self.vtotaldembt.value() <= 11.00:
                    self.rtotaldembt.setValue(1.12)
                elif self.vtotaldembt.value() >= 11.01 and self.vtotaldembt.value() <= 11.25:
                    self.rtotaldembt.setValue(1.20)
                elif self.vtotaldembt.value() >= 11.26 and self.vtotaldembt.value() <= 11.50:
                    self.rtotaldembt.setValue(1.28)
                elif self.vtotaldembt.value() >= 11.51 and self.vtotaldembt.value() <= 11.75:
                    self.rtotaldembt.setValue(1.36)
                elif self.vtotaldembt.value() >= 11.76 and self.vtotaldembt.value() <= 12.00:
                    self.rtotaldembt.setValue(1.44)
                elif self.vtotaldembt.value() >= 12.01 and self.vtotaldembt.value() <= 12.25:
                    self.rtotaldembt.setValue(1.52)
                elif self.vtotaldembt.value() >= 12.26 and self.vtotaldembt.value() <= 12.50:
                    self.rtotaldembt.setValue(1.60)
                elif self.vtotaldembt.value() >= 12.51 and self.vtotaldembt.value() <= 12.75:
                    self.rtotaldembt.setValue(1.68)
                elif self.vtotaldembt.value() >= 12.76 and self.vtotaldembt.value() <= 13.00:
                    self.rtotaldembt.setValue(1.76)
                elif self.vtotaldembt.value() >= 13.01 and self.vtotaldembt.value() <= 13.25:
                    self.rtotaldembt.setValue(1.84)
                elif self.vtotaldembt.value() >= 13.26 and self.vtotaldembt.value() <= 13.50:
                    self.rtotaldembt.setValue(1.92)
                elif self.vtotaldembt.value() >= 13.51 and self.vtotaldembt.value() <= 13.75:
                    self.rtotaldembt.setValue(2.00)
                elif self.vtotaldembt.value() >= 13.76 and self.vtotaldembt.value() <= 14.00:
                    self.rtotaldembt.setValue(2.08)
                elif self.vtotaldembt.value() >= 14.01 and self.vtotaldembt.value() <= 14.25:
                    self.rtotaldembt.setValue(2.16)
                elif self.vtotaldembt.value() >= 14.26 and self.vtotaldembt.value() <= 14.50:
                    self.rtotaldembt.setValue(2.24)
                elif self.vtotaldembt.value() >= 14.51 and self.vtotaldembt.value() <= 14.75:
                    self.rtotaldembt.setValue(2.32)
                elif self.vtotaldembt.value() >= 14.76 and self.vtotaldembt.value() <= 15.00:
                    self.rtotaldembt.setValue(2.40)
                elif self.vtotaldembt.value() >= 15.01 and self.vtotaldembt.value() <= 50.00:
                    self.rtotaldembt.setValue(0)
                    self.obesrvationbt.setText("PRIX A DEBATTRE")
                else:
                    self.rtotaldembt.setValue(0)
                    self.rtotaldembt.clear()
                    self.rtotaldembt.setStyleSheet(
                        "background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                break

        def rGrains_cassebt(self):
            self.rgrainscassébt.clear()
            self.rgrainscassébt.setStyleSheet(
                "background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding: 0px")
            while self.vgrainscassébt.value() >= 4.01:
                self.rgrainscassébt.setValue(0.04)
                self.rgrainscassébt.setStyleSheet(
                    "background-color:#e22630;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                if self.vgrainscassébt.value() <= 4.25:
                    self.rgrainscassébt.setValue(0.04)
                elif self.vgrainscassébt.value() >= 4.26 and self.vgrainscassébt.value() <= 4.50:
                    self.rgrainscassébt.setValue(0.08)
                elif self.vgrainscassébt.value() >= 4.51 and self.vgrainscassébt.value() <= 4.75:
                    self.rgrainscassébt.setValue(0.12)
                elif self.vgrainscassébt.value() >= 4.76 and self.vgrainscassébt.value() <= 5.00:
                    self.rgrainscassébt.setValue(0.16)
                elif self.vgrainscassébt.value() >= 5.01 and self.vgrainscassébt.value() <= 5.25:
                    self.rgrainscassébt.setValue(0.20)
                elif self.vgrainscassébt.value() >= 5.26 and self.vgrainscassébt.value() <= 5.50:
                    self.rgrainscassébt.setValue(0.24)
                elif self.vgrainscassébt.value() >= 5.51 and self.vgrainscassébt.value() <= 5.75:
                    self.rgrainscassébt.setValue(0.28)
                elif self.vgrainscassébt.value() >= 5.76 and self.vgrainscassébt.value() <= 6.00:
                    self.rgrainscassébt.setValue(0.32)
                elif self.vgrainscassébt.value() >= 6.01 and self.vgrainscassébt.value() <= 6.25:
                    self.rgrainscassébt.setValue(0.36)
                elif self.vgrainscassébt.value() >= 6.26 and self.vgrainscassébt.value() <= 6.50:
                    self.rgrainscassébt.setValue(0.40)
                elif self.vgrainscassébt.value() >= 6.51 and self.vgrainscassébt.value() <= 6.75:
                    self.rgrainscassébt.setValue(0.44)
                elif self.vgrainscassébt.value() >= 6.76 and self.vgrainscassébt.value() <= 7.00:
                    self.rgrainscassébt.setValue(0.48)
                elif self.vgrainscassébt.value() >= 7.01 and self.vgrainscassébt.value() <= 7.25:
                    self.rgrainscassébt.setValue(0.52)
                elif self.vgrainscassébt.value() >= 7.26 and self.vgrainscassébt.value() <= 7.50:
                    self.rgrainscassébt.setValue(0.56)
                elif self.vgrainscassébt.value() >= 7.51 and self.vgrainscassébt.value() <= 7.75:
                    self.rgrainscassébt.setValue(0.60)
                elif self.vgrainscassébt.value() >= 7.76 and self.vgrainscassébt.value() <= 8.00:
                    self.rgrainscassébt.setValue(0.64)
                elif self.vgrainscassébt.value() >= 8.01 and self.vgrainscassébt.value() <= 8.25:
                    self.rgrainscassébt.setValue(0.68)
                elif self.vgrainscassébt.value() >= 8.26 and self.vgrainscassébt.value() <= 8.50:
                    self.rgrainscassébt.setValue(0.72)
                elif self.vgrainscassébt.value() >= 8.51 and self.vgrainscassébt.value() <= 8.75:
                    self.rgrainscassébt.setValue(0.76)
                elif self.vgrainscassébt.value() >= 8.76 and self.vgrainscassébt.value() <= 9.00:
                    self.rgrainscassébt.setValue(0.80)
                elif self.vgrainscassébt.value() >= 9.01 and self.vgrainscassébt.value() <= 9.25:
                    self.rgrainscassébt.setValue(0.84)
                elif self.vgrainscassébt.value() >= 9.26 and self.vgrainscassébt.value() <= 9.50:
                    self.rgrainscassébt.setValue(0.88)
                elif self.vgrainscassébt.value() >= 9.51 and self.vgrainscassébt.value() <= 9.75:
                    self.rgrainscassébt.setValue(0.92)
                elif self.vgrainscassébt.value() >= 9.76 and self.vgrainscassébt.value() <= 10.00:
                    self.rgrainscassébt.setValue(0.96)
                elif self.vgrainscassébt.value() >= 10.01 and self.vgrainscassébt.value() <= 10.25:
                    self.rgrainscassébt.setValue(1.00)
                elif self.vgrainscassébt.value() >= 10.26 and self.vgrainscassébt.value() <= 30:
                    self.rgrainscassébt.setValue(0)
                    self.obesrvationbt.setText("PRIX A DEBATTRE")
                else:
                    self.rgrainscassébt.setValue(0)
                    self.rgrainscassébt.clear()
                    self.rgrainscassébt.setStyleSheet(
                        "background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                break

        def grain_pounaiséebt(self):
            self.rgrainpunaiséebt.setValue(0)
            self.rgrainpunaiséebt.clear()
            self.rgrainpunaiséebt.setStyleSheet(
                "background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding: 0px")
            while self.vgrainpunaiséebt.value() >= 2.01:
                self.rgrainpunaiséebt.setValue(0.08)
                self.rgrainpunaiséebt.setStyleSheet(
                    "background-color:#e22630;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                if self.vgrainpunaiséebt.value() <= 2.25:
                    self.rgrainpunaiséebt.setValue(0.08)
                elif self.vgrainpunaiséebt.value() >= 2.26 and self.vgrainpunaiséebt.value() <= 2.50:
                    self.rgrainpunaiséebt.setValue(0.16)
                elif self.vgrainpunaiséebt.value() >= 2.51 and self.vgrainpunaiséebt.value() <= 2.75:
                    self.rgrainpunaiséebt.setValue(0.24)
                elif self.vgrainpunaiséebt.value() >= 2.76 and self.vgrainpunaiséebt.value() <= 3.00:
                    self.rgrainpunaiséebt.setValue(0.32)
                elif self.vgrainpunaiséebt.value() >= 3.01 and self.vgrainpunaiséebt.value() <= 3.25:
                    self.rgrainpunaiséebt.setValue(0.40)
                elif self.vgrainpunaiséebt.value() >= 3.26 and self.vgrainpunaiséebt.value() <= 3.50:
                    self.rgrainpunaiséebt.setValue(0.48)
                elif self.vgrainpunaiséebt.value() >= 3.51 and self.vgrainpunaiséebt.value() <= 3.75:
                    self.rgrainpunaiséebt.setValue(0.56)
                elif self.vgrainpunaiséebt.value() >= 3.76 and self.vgrainpunaiséebt.value() <= 4.00:
                    self.rgrainpunaiséebt.setValue(0.64)
                elif self.vgrainpunaiséebt.value() >= 4.01 and self.vgrainpunaiséebt.value() <= 4.25:
                    self.rgrainpunaiséebt.setValue(0.72)
                elif self.vgrainpunaiséebt.value() >= 4.26 and self.vgrainpunaiséebt.value() <= 4.50:
                    self.rgrainpunaiséebt.setValue(0.80)
                elif self.vgrainpunaiséebt.value() >= 4.51 and self.vgrainpunaiséebt.value() <= 4.75:
                    self.rgrainpunaiséebt.setValue(0.88)
                elif self.vgrainpunaiséebt.value() >= 4.76 and self.vgrainpunaiséebt.value() <= 5.00:
                    self.rgrainpunaiséebt.setValue(0.96)
                elif self.vgrainpunaiséebt.value() >= 5.01 and self.vgrainpunaiséebt.value() <= 5.25:
                    self.rgrainpunaiséebt.setValue(1.04)
                elif self.vgrainpunaiséebt.value() >= 5.26 and self.vgrainpunaiséebt.value() <= 5.50:
                    self.rgrainpunaiséebt.setValue(1.12)
                elif self.vgrainpunaiséebt.value() >= 5.51 and self.vgrainpunaiséebt.value() <= 5.75:
                    self.rgrainpunaiséebt.setValue(1.20)
                elif self.vgrainpunaiséebt.value() >= 5.76 and self.vgrainpunaiséebt.value() <= 6.00:
                    self.rgrainpunaiséebt.setValue(1.28)
                elif self.vgrainpunaiséebt.value() >= 6.01 and self.vgrainpunaiséebt.value() <= 6.25:
                    self.rgrainpunaiséebt.setValue(1.36)
                elif self.vgrainpunaiséebt.value() >= 6.26 and self.vgrainpunaiséebt.value() <= 6.50:
                    self.rgrainpunaiséebt.setValue(1.44)
                elif self.vgrainpunaiséebt.value() >= 6.51 and self.vgrainpunaiséebt.value() <= 6.75:
                    self.rgrainpunaiséebt.setValue(1.52)
                elif self.vgrainpunaiséebt.value() >= 6.76 and self.vgrainpunaiséebt.value() <= 7.00:
                    self.rgrainpunaiséebt.setValue(1.60)
                elif self.vgrainpunaiséebt.value() >= 7.01 and self.vgrainpunaiséebt.value() <= 7.25:
                    self.rgrainpunaiséebt.setValue(1.68)
                elif self.vgrainpunaiséebt.value() >= 7.26 and self.vgrainpunaiséebt.value() <= 7.50:
                    self.rgrainpunaiséebt.setValue(1.76)
                elif self.vgrainpunaiséebt.value() >= 7.51 and self.vgrainpunaiséebt.value() <= 7.75:
                    self.rgrainpunaiséebt.setValue(1.84)
                elif self.vgrainpunaiséebt.value() >= 7.76 and self.vgrainpunaiséebt.value() <= 8.00:
                    self.rgrainpunaiséebt.setValue(1.92)
                elif self.vgrainpunaiséebt.value() >= 8.01 and self.vgrainpunaiséebt.value() <= 8.25:
                    self.rgrainpunaiséebt.setValue(2.00)
                elif self.vgrainpunaiséebt.value() >= 8.26 and self.vgrainpunaiséebt.value() <= 8.50:
                    self.rgrainpunaiséebt.setValue(2.08)
                elif self.vgrainpunaiséebt.value() >= 8.51 and self.vgrainpunaiséebt.value() <= 8.75:
                    self.rgrainpunaiséebt.setValue(2.16)
                elif self.vgrainpunaiséebt.value() >= 8.76 and self.vgrainpunaiséebt.value() <= 9.00:
                    self.rgrainpunaiséebt.setValue(2.24)
                elif self.vgrainpunaiséebt.value() >= 9.01 and self.vgrainpunaiséebt.value() <= 9.25:
                    self.rgrainpunaiséebt.setValue(2.32)
                elif self.vgrainpunaiséebt.value() >= 9.26 and self.vgrainpunaiséebt.value() <= 9.50:
                    self.rgrainpunaiséebt.setValue(2.40)
                elif self.vgrainpunaiséebt.value() >= 9.51 and self.vgrainpunaiséebt.value() <= 9.75:
                    self.rgrainpunaiséebt.setValue(2.48)
                elif self.vgrainpunaiséebt.value() >= 9.76 and self.vgrainpunaiséebt.value() <= 10.00:
                    self.rgrainpunaiséebt.setValue(2.56)
                elif self.vgrainpunaiséebt.value() >= 10.01 and self.vgrainpunaiséebt.value() <= 40:
                    self.rgrainpunaiséebt.setValue(0)
                    self.obesrvationbt.setText("PRIX A DEBATTRE")
                else:
                    self.rgrainpunaiséebt.setValue(0)
                    self.rgrainpunaiséebt.clear()
                    self.rgrainpunaiséebt.setStyleSheet(
                        "background-color:#ffffff;border: 2px solid bleu;border-radius: 4px;padding: 0px")
                break

        def plus_bt(self):
            a = self.vmatiére20mmbt.value()
            b = self.vdébrisvébt.value()
            c = self.vgrainnuisiblebt.value()
            d = self.vgrainscassébt.value()
            e = self.vgrainpunaiséebt.value()
            f = self.vgrainsfortementboutésbt.value()
            g = self.vgrainsfaiblementboutésbt.value()
            h = self.vgrainefortementmouchetesbt.value()
            i = self.vgrainetrangebt.value()
            totalprembt = a + b + c
            totaldem = d + e + f + g + h + i
            btotalBonietRefac = self.btotalprembt.value() + self.btotaldembt.value() + self.bpsbt.value()
            rtotalRefac = self.rtotalprembt.value() + self.rtotaldembt.value() + self.rpsbt.value() + self.rgrainscassébt.value() + self.rgrainpunaiséebt.value()
            self.btotalbiniEtrefactionbt.setValue(btotalBonietRefac)
            self.rtotalbiniEtrefactionbt.setValue(rtotalRefac)
            self.vtotalprembt.setValue(totalprembt)
            if self.vgrainscassébt.value() <= 4.00 and self.vgrainpunaiséebt.value() <= 2:
                self.vtotaldembt.setValue(totaldem)
            elif self.vgrainscassébt.value() >= 4.01 and self.vgrainpunaiséebt.value() >= 2.01:
                self.vtotaldembt.setValue(f + g + h + i)

        def clear_bt(self):
            self.vpsbt.clear()
            self.vpsbt.setValue(0)
            self.vhumiditebt.clear()
            self.vergotbt.clear()
            self.vmatiére20mmbt.clear()
            self.vmatiére20mmbt.setValue(0)
            self.vdébrisvébt.clear()
            self.vdébrisvébt.setValue(0)
            self.vgrainnuisiblebt.clear()
            self.vgrainnuisiblebt.setValue(0)
            self.vtotalprembt.clear()
            self.vtotalprembt.setValue(0)
            self.vgrainscassébt.clear()
            self.vgrainscassébt.setValue(0)
            self.vgrainpunaiséebt.clear()
            self.vgrainpunaiséebt.setValue(0)
            self.vgrainsfortementboutésbt.clear()
            self.vgrainsfortementboutésbt.setValue(0)
            self.vgrainsfaiblementboutésbt.clear()
            self.vgrainsfaiblementboutésbt.setValue(0)
            self.vgrainefortementmouchetesbt.clear()
            self.vgrainefortementmouchetesbt.setValue(0)
            self.vgrainetrangebt.clear()
            self.vgrainetrangebt.setValue(0)
            self.vtotaldembt.clear()
            self.vtotaldembt.setValue(0)
            self.rpsbt.clear()
            self.rpsbt.setValue(0)
            self.rergotbt.clear()
            self.rmatiére20mmbt.clear()
            self.rmatiére20mmbt.setValue(0)
            self.rdébrisvébt.clear()
            self.rdébrisvébt.setValue(0)
            self.rgrainnuisiblebt.clear()
            self.rgrainnuisiblebt.setValue(0)
            self.rtotalprembt.clear()
            self.rtotalprembt.setValue(0)
            self.rgrainscassébt.clear()
            self.rgrainscassébt.setValue(0)
            self.rgrainpunaiséebt.clear()
            self.rgrainpunaiséebt.setValue(0)
            self.rgrainetrangebt.clear()
            self.rgrainetrangebt.setValue(0)
            self.bpsbt.clear()
            self.bpsbt.setValue(0)
            self.bergotbt.clear()
            self.bmatiére20mmbt.clear()
            self.bmatiére20mmbt.setValue(0)
            self.btébrisvébt.clear()
            self.btébrisvébt.setValue(0)
            self.bgrainnuisiblebt.clear()
            self.bgrainnuisiblebt.setValue(0)
            self.btotalprembt.clear()
            self.btotalprembt.setValue(0)
            self.bgrainscassébt.clear()
            self.bgrainscassébt.setValue(0)
            self.bgrainpunaiséebt.clear()
            self.bgrainpunaiséebt.setValue(0)
            self.bgrainsfaiblementboutésbt.clear()
            self.bgrainsfaiblementboutésbt.setValue(0)
            self.bgrainefortementmouchetesbt.clear()
            self.bgrainefortementmouchetesbt.setValue(0)
            self.bgrainetrangebt.clear()
            self.bgrainetrangebt.setValue(0)
            self.btotaldembt.clear()
            self.btotaldembt.setValue(0)

        def calcul_bt(self):
            self.réfaction_ps_bt()
            self.bonification_ps_bt()
            self.refaction_impurté_er_bt()
            self.bonification_pimpurte_er_bt()
            self.refaction_impurte_eme_bt()
            self.rGrains_cassebt()
            self.grain_pounaiséebt()
            self.plus_bt()

        def bulletin_bt(self):
            try:
                nbulltinbt = self.n_bultinbt.text()
                nomproducteurbt = self.nome_du_producteurbt.text()
                perebt = self.pérebt.text()
                matriculbt = self.adressebt.text()
                poindecollectebt = self.pointdecollectbt.currentText()
                datebt = self.dattereceptiontbt.text()
                agreeurbt = self.agréeeurcombobt.currentText()
                quantitebt = self.quantitebt.value()
                vpsbt = self.vpsbt.value()
                bpsbt = self.bpsbt.text()
                rpsbt = self.rpsbt.text()
                vhumidite = self.vhumiditebt.value()
                vergot = self.vergotbt.value()
                vmatier20mm = self.vmatiére20mmbt.value()
                vdebrit = self.vdébrisvébt.value()
                vgrainnuisible = self.vgrainnuisiblebt.value()
                graincasse = self.vgrainscassébt.value()
                rgraicasse = self.rgrainscassébt.text()
                grainpunaise = self.vgrainpunaiséebt.value()
                grainfboute = self.vgrainsfortementboutésbt.value()
                grainfaboute = self.vgrainsfaiblementboutésbt.value()
                grainfortementmouchte = self.vgrainefortementmouchetesbt.value()
                grainetrange = self.vgrainetrangebt.value()
                rgrainpunaise = self.rgrainpunaiséebt.text()
                vtotalbtp = self.vtotalprembt.value()
                vtotalbtd = self.vtotaldembt.value()
                btotalbtp = self.btotalprembt.text()
                btotalbtd = self.btotaldembt.text()
                rtotalbtp = self.rtotalprembt.text()
                rtotalbtd = self.rtotaldembt.text()
                btotalbt = self.btotalbiniEtrefaction.text()
                rtotalbt = self.rtotalbiniEtrefaction.text()
                ndcartebt = self.n_cartebt.text()
                observationbt = self.obesrvationbt.toPlainText()
                date = QDate.currentDate()
                current_date = QDate.currentDate()
                current_day = current_date.day()
                current_month = current_date.month()
                current_year = current_date.year()
                current_time = QTime.currentTime()
                current_hour = current_time.hour()
                current_minute = current_time.minute()
                # Format the date as "dd-mm-yyyy"
                current_days = f"{current_day:02d}-{current_month:02d}-{current_year}-{current_hour}-{current_minute}"
                self.docbt = DocxTemplate("bulletin agreage/bulletin_blétendre/bulletin_blétendre.docx")
                self.docbt.render({
                    "nmp": nomproducteurbt,
                    "per": perebt,
                    "mat": matriculbt,
                    "ptt": poindecollectebt,
                    "dt": datebt,
                    "agr": agreeurbt,
                    "qtt": quantitebt,
                    "vps": vpsbt,
                    "bps": bpsbt,
                    "rps": rpsbt,
                    "vmh": vhumidite,
                    "vrg": vergot,
                    "vgs": vmatier20mm,
                    "vdb": vdebrit,
                    "vgn": vgrainnuisible,
                    "vgc": graincasse,
                    "vgfb": grainpunaise,
                    "vgr": grainfboute,
                    "vgfm": grainfaboute,
                    "vgpn": grainfortementmouchte,
                    "vgrp": grainetrange,
                    "rgp": rgrainpunaise,
                    "rgc": rgraicasse,
                    "vttp": vtotalbtp,
                    "bttp": btotalbtp,
                    "rtt": rtotalbtp,
                    "vtd": vtotalbtd,
                    "btd": btotalbtd,
                    "rtd": rtotalbtd,
                    "btb": btotalbt,
                    "rtr": rtotalbt,
                    "ncn": ndcartebt,
                    "oo": observationbt,
                    "nm": nbulltinbt, })
                databt = sqlite3.connect("data_bt.db")
                cursbt = databt.cursor()
                cursbt.execute(
                    "SELECT * FROM bttable_bulletin_sortie WHERE num_bulletinbts AND datebts=? AND nom_producteurbts = ? AND perebts = ?  AND quantites=?",
                    (datebt, nomproducteurbt, perebt, quantitebt))
                existing_data = cursbt.fetchone()
                if existing_data:
                    msgbox = QtWidgets.QMessageBox()
                    msgbox.setWindowTitle('Confirmation')
                    msgbox.setText('Le Bulletin sortie  déjà.')
                    msgbox.exec()
                    # tempfilebtd = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                    # self.docbt.save( tempfilebtd)
                    # os.startfile(tempfilebtd, "print")
                    # tempfilebt = nomproducteurbt + "-" + current_days + "B.T" + ".docx"
                    # self.docbt.save("bulletin agreage/bulletin_blétendre/" + tempfilebt)
                    # pathbt = os.path.abspath("bulletin agreage/bulletin_blétendre/" + tempfilebt)

                else:
                    doc_names = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                    doc_pdf = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False).name
                    self.docbt.save(doc_names)
                    try:
                        if doc_names:
                            a = self.progress_bar()
                            sys.stderr = open("consoleoutput.log", "w")
                            convert(doc_names, doc_pdf)
                            # Open the resulting .pdf file using the default associated application
                            # os.startfile(doc_pdf, 'open')
                            app_path = 'C:\\Program Files\\Okular\\bin\\okular.exe'
                            subprocess.Popen([app_path, doc_pdf])
                    except Exception as e:
                        print(e)
                    tempfilebt = nomproducteurbt + "-" + current_days + "B.T" + ".docx"
                    self.docbt.save("bulletin agreage/bulletin_blétendre/" + tempfilebt)
                    pathbt = os.path.abspath("bulletin agreage/bulletin_blétendre/" + tempfilebt)
                    databasebt = sqlite3.connect("data_bt.db")
                    curs = databasebt.cursor()
                    curs.execute(
                        "INSERT INTO bttable_bulletin_sortie (datebts, nom_producteurbts, perebts,quantites, pathbts,date_sortie) VALUES (?,?,?, ?, ?, ?)",
                        (datebt, nomproducteurbt, perebt, quantitebt, pathbt, current_day,))
                    databasebt.commit()
                    databasebt.close()
            except Exception as e:
                print(e)

        def add_datta_bt(self):
            try:
                nbulltinbt = self.n_bultinbt.text()
                nomproducteurbt = self.nome_du_producteurbt.text()
                perebt = self.pérebt.text()
                matriculbt = self.adressebt.text()
                poindecollectebt = self.pointdecollectbt.currentText()
                datebt = self.dattereceptiontbt.text()
                agreeurbt = self.agréeeurcombobt.currentText()
                quantitebt = self.quantitebt.value()
                vpsbt = self.vpsbt.value()
                bpsbt = self.bpsbt.text()
                rpsbt = self.rpsbt.text()
                vhumidite = self.vhumiditebt.value()
                vergot = self.vergotbt.value()
                vmatier20mm = self.vmatiére20mmbt.value()
                vdebrit = self.vdébrisvébt.value()
                vgrainnuisible = self.vgrainnuisiblebt.value()
                graincasse = self.vgrainscassébt.value()
                rgraicasse = self.rgrainscassébt.text()
                grainpunaise = self.vgrainpunaiséebt.value()
                grainfboute = self.vgrainsfortementboutésbt.value()
                grainfaboute = self.vgrainsfaiblementboutésbt.value()
                grainfortementmouchte = self.vgrainefortementmouchetesbt.value()
                grainetrange = self.vgrainetrangebt.value()
                rgrainpunaise = self.rgrainpunaiséebt.text()
                vtotalbtp = self.vtotalprembt.value()
                vtotalbtd = self.vtotaldembt.value()
                btotalbtp = self.btotalprembt.text()
                btotalbtd = self.btotaldembt.text()
                rtotalbtp = self.rtotalprembt.text()
                rtotalbtd = self.rtotaldembt.text()
                btotalbt = self.btotalbiniEtrefaction.text()
                rtotalbt = self.rtotalbiniEtrefaction.text()
                ndcartebt = self.n_cartebt.text()
                observationbt = self.obesrvationbt.toPlainText()
                date = QDate.currentDate()
                current_date = QDate.currentDate()
                current_day = current_date.day()
                current_month = current_date.month()
                current_year = current_date.year()
                current_time = QTime.currentTime()
                current_hour = current_time.hour()
                current_minute = current_time.minute()
                # Format the date as "dd-mm-yyyy"
                current_days = f"{current_day:02d}-{current_month:02d}-{current_year}-{current_hour}-{current_minute}"
                self.docbt = DocxTemplate("bulletin agreage/bulletin_blétendre/bulletin_blétendre.docx")
                self.docbt.render({
                    "nmp": nomproducteurbt,
                    "per": perebt,
                    "mat": matriculbt,
                    "ptt": poindecollectebt,
                    "dt": datebt,
                    "agr": agreeurbt,
                    "qtt": quantitebt,
                    "vps": vpsbt,
                    "bps": bpsbt,
                    "rps": rpsbt,
                    "vmh": vhumidite,
                    "vrg": vergot,
                    "vgs": vmatier20mm,
                    "vdb": vdebrit,
                    "vgn": vgrainnuisible,
                    "vgc": graincasse,
                    "vgfb": grainpunaise,
                    "vgr": grainfboute,
                    "vgfm": grainfaboute,
                    "vgpn": grainfortementmouchte,
                    "vgrp": grainetrange,
                    "rgp": rgrainpunaise,
                    "rgc": rgraicasse,
                    "vttp": vtotalbtp,
                    "bttp": btotalbtp,
                    "rtt": rtotalbtp,
                    "vtd": vtotalbtd,
                    "btd": btotalbtd,
                    "rtd": rtotalbtd,
                    "btb": btotalbt,
                    "rtr": rtotalbt,
                    "ncn": ndcartebt,
                    "oo": observationbt,
                    "nm": nbulltinbt
                })
                tempfilebt = nomproducteurbt + "-" + current_days + "B.T" + ".docx"
                self.docbt.save("bulletin agreage/bulletin_blétendre/" + tempfilebt)
                pathbt = os.path.abspath("bulletin agreage/bulletin_blétendre/" + tempfilebt)
                # Check if the data already exists
                datadb = sqlite3.connect("data_bt.db")
                cursbt = datadb.cursor()
                cursbt.execute(
                    "SELECT * FROM bttable WHERE datebt=? AND nom_producteurbt=? AND perebt=? AND matriculebt=? AND point_collectbt=? AND nom_agreeurbt=? AND quantitebt=? AND psbt=? AND totalbt=? AND pathbt=? ",
                    (datebt, nomproducteurbt, perebt, matriculbt, poindecollectebt, agreeurbt, quantitebt, vpsbt,
                     vtotalbtp, pathbt,))
                existing_data = cursbt.fetchone()

                if existing_data:
                    msgbox = QtWidgets.QMessageBox()
                    msgbox.setWindowTitle('confirmation')
                    msgbox.setText('Les données existent déjà.')
                    msgbox.exec()
                else:
                    cursbt.execute(
                        "INSERT INTO bttable (datebt ,nom_producteurbt ,perebt ,matriculebt ,point_collectbt ,nom_agreeurbt ,quantitebt ,psbt ,totalbt,pathbt  )  values(?,?,?,?,?,?,?,?,?,?)",
                        (datebt, nomproducteurbt, perebt, matriculbt, poindecollectebt, agreeurbt, quantitebt, vpsbt,
                         vtotalbtp, pathbt,))
                    datadb.commit()
                    datadb.close()

                    msgbox = QtWidgets.QMessageBox()
                    msgbox.setWindowTitle('confirmation')
                    msgbox.setText('Les données sont sauvegardées')
                    msgbox.exec()

            except sqlite3.Error as e:
                print("Error accessing the database:", e)

        def number_bulletin_bt(self):
            try:
                databasebt = sqlite3.connect("data_bt.db")
                cursor = databasebt.cursor()
                cursor.execute("SELECT num_bulletinbt FROM bttable ORDER BY num_bulletinbt DESC LIMIT 1")
                resultbt = cursor.fetchone()
                if resultbt is not None:
                    last_id = resultbt[0]
                    self.n_bultinbt.setText("DC-" + str(last_id))
                else:
                    # Handle the case when there are no records in the orgetable
                    self.n_bultinbt.setText("DC-")
            except sqlite3.Error as e:
                print("Error accessing the database:", e)
            pass

        def today_bt(self):
            try:
                dialog = QtWidgets.QMessageBox()
                dialog.setStyleSheet("""
                                          QWidget {
                                              color: #000000;
                                              background-color: #D8F9DB;
                                              border-width: 1px;
                                              border-color: #1e1e1e;
                                              border-style: solid;
                                              border-radius: 6;
                                              padding: 3px;
                                              font-size: 18px;
                                              padding-left: 5px;
                                              padding-right: 5px;
                                          }
                                          QWidget:item:hover {
                                              background-color: #3daee9;
                                              color: #eff0f1;
                                          }
                                          QWidget:item:selected {
                                              background-color: #3daee9;
                                          }
                                          QWidget:disabled {
                                              color: #454545;
                                              background-color: #31363b;
                                          }
                                          QPushButton {
                                              color: #000000;
                                              background-color:#84dbc8;
                                              border-width: 1px;
                                              border-color: #1e1e1e;
                                              border-style: solid;
                                              border-radius: 6;
                                              padding: 3px;
                                              font-size: 12px;
                                              padding-left: 5px;
                                              padding-right: 5px;
                                              min-width: 40px;
                                          }
                                          QPushButton:disabled {
                                              background-color: #31363b;
                                              border-width: 1px;
                                              border-color: #454545;
                                              border-style: solid;
                                              padding-top: 5px;
                                              padding-bottom: 5px;
                                              padding-left: 10px;
                                              padding-right: 10px;
                                              border-radius: 2px;
                                              color: #454545;
                                          }
                                          QPushButton:pressed {
                                              background-color: #3daee9;
                                              padding-top: -15px;
                                              padding-bottom: -17px;
                                          }
                                          QPushButton:hover {
                                              border: 1px solid #ff8c00;
                                              color: #eff0f1;
                                          }
                                          QLabel {
                                              font-size: 18px;
                                              border: 0px solid orange;
                                          }
                                      """)
                dialog.setWindowTitle("Sélectionnez la date")
                dialog.setText("Sélectionnez la date  \t\n")
                self.datebt = QtWidgets.QDateEdit(dialog)
                self.datebt.setDate(self.datedaytime)
                self.datebt.resize(180, 30)
                self.datebt.move(40, 50)
                ok_button = QtWidgets.QPushButton("OK", dialog)
                cancel_button = QtWidgets.QPushButton("Cancel", dialog)
                dialog.addButton(ok_button, QMessageBox.ButtonRole.ActionRole)
                dialog.addButton(cancel_button, QMessageBox.ButtonRole.ActionRole)
                dialog.exec()

                if dialog.clickedButton() == ok_button:
                    self.selected_datebt = self.datebt.text()
                    if self.datebt.text() == self.selected_datebt:

                        conn = sqlite3.connect('data_bt.db')
                        cursor = conn.cursor()
                        cursor.execute(
                            "SELECT num_bulletinbt,datebt ,nom_producteurbt ,perebt ,matriculebt ,point_collectbt  ,quantitebt  FROM bttable WHERE DATE(substr(datebt, 7, 4) || '-' || substr(datebt, 4, 2) || '-' || substr(datebt, 1, 2)) =?; ",
                            (self.selected_datebt,))
                        rows = cursor.fetchall()

                        ########################################################################
                        # Create a new document and add a table
                        self.doc = docx.Document()
                        section = self.doc.sections[0]
                        section.page_width = docx.shared.Cm(29.7)
                        section.page_height = docx.shared.Cm(21.0)
                        section.top_margin = docx.shared.Cm(1.5)
                        section.bottom_margin = docx.shared.Cm(1.5)
                        heading1 = self.doc.add_heading("\t\t\t\t\t\t les entre de Blé Dur ", level=1)
                        heading1.style.font.name = 'Times New Roman'
                        heading1.style.font.size = Pt(22)
                        heading1.style.font.bold = True
                        heading1.style.font.color.rgb = RGBColor(0, 0, 0)
                        # add second heading
                        paragraph = self.doc.add_paragraph()
                        paragraph.style.font.name = 'Times New Roman'
                        paragraph.style.font.size = Pt(14)
                        left_run = paragraph.add_run('CCLS RELIZANE\t\t\t\t\t\t\t\t\t\t\t\t')
                        left_run.bold = True
                        right_run = paragraph.add_run('Date:' + self.selected_datebt)
                        right_run.bold = True

                        table = self.doc.add_table(rows=1, cols=9)
                        table.style = "Table Grid"  # set the table style
                        hdr_cells = table.rows[0].cells
                        hdr_cells[0].text = 'N° Bulletin \nDC'
                        hdr_cells[1].text = 'Date'
                        hdr_cells[2].text = 'Nom et Prénom'
                        hdr_cells[3].text = 'Pére'
                        hdr_cells[4].text = 'Matricule'
                        hdr_cells[5].text = 'point de collect'
                        hdr_cells[6].text = 'Quantité'
                        hdr_cells[7].text = 'Date de sortie '
                        hdr_cells[8].text = 'N° de la carte '

                        # Set the width of the header cells
                        hdr_cells[0].width = Inches(1.5)  # set the width of the first column to 1.5 inches
                        hdr_cells[1].width = Inches(1.5)
                        hdr_cells[2].width = Inches(2)
                        hdr_cells[3].width = Inches(1)
                        hdr_cells[4].width = Inches(1.5)
                        hdr_cells[5].width = Inches(1.5)
                        hdr_cells[6].width = Inches(1)
                        hdr_cells[7].width = Inches(1.5)
                        hdr_cells[8].width = Inches(1.5)

                        # set hight of the column
                        hdr_cells[0].height = Inches(1)  # set the width of the first column to 1.5 inches
                        hdr_cells[1].height = Inches(1)
                        hdr_cells[2].height = Inches(1)
                        hdr_cells[3].height = Inches(1)
                        hdr_cells[4].height = Inches(1)
                        hdr_cells[5].height = Inches(1)
                        hdr_cells[6].height = Inches(1)
                        hdr_cells[7].width = Inches(1)
                        hdr_cells[8].width = Inches(1)

                        for cell in table.rows[0].cells:
                            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                            cell.paragraphs[0].runs[0].font.size = Pt(12)
                            cell.paragraphs[0].runs[0].font.bold = True
                        # Add the data to the table
                        table_rows = len(rows)
                        table_cols = len(rows[0])
                        for row in range(table_rows):
                            table.add_row()
                            for col in range(table_cols):
                                cell = table.cell(row + 1, col)
                                cell.text = str(rows[row][col])
                                # Set font properties
                                cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                                cell.paragraphs[0].runs[0].font.size = Pt(12)
                                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                        for row in table.rows:
                            row.height = Inches(0.3)
                        # Save and open the document for printing
                        doc_names = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                        doc_pdf = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False).name
                        self.doc.save(doc_names)
                        try:
                            if doc_names:
                                a = self.progress_bar()
                                sys.stderr = open("consoleoutput.log", "w")
                                convert(doc_names, doc_pdf)
                                # Open the resulting .pdf file using the default associated application
                                # os.startfile(doc_pdf, 'open')
                                app_path = 'C:\\Program Files\\Okular\\bin\\okular.exe'
                                subprocess.Popen([app_path, doc_pdf])
                        except Exception as e:
                            print(e)
                    else:
                        msgbox = QtWidgets.QMessageBox()
                        msgbox.setWindowTitle('Erreur')
                        msgbox.setText("Il n'yapas des données à cette date.")
                        msgbox.exec()
                if dialog.clickedButton() == cancel_button:
                    dialog.close()
            except Exception as e:
                print(e)

        def printBulltinProducteurbt(self):
            try:
                dialog = QtWidgets.QMessageBox()
                dialog.setFixedSize(300, 600)
                dialog.setStyleSheet("""
                                          QWidget {
                                              color: #000000;
                                              background-color: #D8F9DB;
                                              border-width: 1px;
                                              border-color: #1e1e1e;
                                              border-style: solid;
                                              border-radius: 6;
                                              padding: 3px;
                                              font-size: 18px;
                                              padding-left: 5px;
                                              padding-right: 5px;
                                          }
                                          QWidget:item:hover {
                                              background-color: #3daee9;
                                              color: #eff0f1;
                                          }
                                          QWidget:item:selected {
                                              background-color: #3daee9;
                                          }
                                          QWidget:disabled {
                                              color: #454545;
                                              background-color: #31363b;
                                          }
                                          QPushButton {
                                              color: #000000;
                                              background-color:#84dbc8;
                                              border-width: 1px;
                                              border-color: #1e1e1e;
                                              border-style: solid;
                                              border-radius: 6;
                                              padding: 3px;
                                              font-size: 12px;
                                              padding-left: 5px;
                                              padding-right: 5px;
                                              min-width: 40px;
                                          }
                                          QPushButton:disabled {
                                              background-color: #31363b;
                                              border-width: 1px;
                                              border-color: #454545;
                                              border-style: solid;
                                              padding-top: 5px;
                                              padding-bottom: 5px;
                                              padding-left: 10px;
                                              padding-right: 10px;
                                              border-radius: 2px;
                                              color: #454545;
                                          }
                                          QPushButton:pressed {
                                              background-color: #3daee9;
                                              padding-top: -15px;
                                              padding-bottom: -17px;
                                          }
                                          QPushButton:hover {
                                              border: 1px solid #ff8c00;
                                              color: #eff0f1;
                                          }
                                          QLabel {
                                              font-size: 18px;
                                              border: 0px solid orange;
                                          }
                                      """)
                dialog.setWindowTitle("Sélectionnez le Nom et Prénom et pére")
                dialog.setText("Sélectionnez le Nom et Prénom et pére  \t\t\t\t\t\t\n\t\t\t\t\t\t\t\t\t\t\n\t\t\t\t")
                self.nomproducture = QtWidgets.QLineEdit(dialog)
                self.nomproducture.setPlaceholderText("Nom et Prénom")
                self.nomproducture.resize(180, 30)
                self.nomproducture.move(40, 40)
                self.pére = QtWidgets.QLineEdit(dialog)
                self.pére.setPlaceholderText("Pére")
                self.pére.resize(180, 30)
                self.pére.move(240, 40)
                self.datetxt = QtWidgets.QLabel("Date entré", dialog)
                self.datetxt.setGeometry(QtCore.QRect(30, 80, 100, 20))
                self.datep = QDateTimeEdit(dialog)
                self.datep.setDisplayFormat("dd-MM-yyyy hh:mm")
                self.datep.resize(180, 30)
                self.datep.move(40, 105)
                self.quantiteptext = QtWidgets.QLabel("Quantité", dialog)
                self.quantiteptext.setGeometry(QtCore.QRect(230, 80, 100, 20))
                self.quantitep = QtWidgets.QDoubleSpinBox(dialog)
                self.quantitep.setRange(1, 10000)
                self.quantitep.setSuffix('  QX')
                self.quantitep.setSpecialValueText(" ")
                self.quantitep.setValue(0.0)
                self.quantitep.resize(180, 30)
                self.quantitep.move(240, 105)
                ok_button = QtWidgets.QPushButton("OK", dialog)
                cancel_button = QtWidgets.QPushButton("Cancel", dialog)
                dialog.addButton(ok_button, QMessageBox.ButtonRole.ActionRole)
                dialog.addButton(cancel_button, QMessageBox.ButtonRole.ActionRole)
                dialog.exec()
                date = QDate.currentDate()
                current_date = QDate.currentDate()
                current_day = current_date.day()
                current_month = current_date.month()
                current_year = current_date.year()

                # Format the date as "dd-mm-yyyy"
                current_day = f"{current_day:02d}-{current_month:02d}-{current_year}"

                if dialog.clickedButton() == ok_button:
                    producteur = self.nomproducture.text()
                    pere = self.pére.text()
                    dateentré = self.datep.text()
                    quantite = self.quantitep.value()
                    datadb = sqlite3.connect("data_bt.db")
                    cursbt = datadb.cursor()
                    cursbt.execute(
                        "SELECT * FROM bttable_bulletin_sortie WHERE num_bulletinbts AND datebts=? AND nom_producteurbts = ? AND perebts = ?  AND quantites=?",
                        (dateentré, producteur, pere, quantite))
                    existing_data = cursbt.fetchone()
                    if existing_data:
                        msgbox = QtWidgets.QMessageBox()
                        msgbox.setWindowTitle('Confirmation')
                        msgbox.setText('Le Bulletin sortie  déjà.')
                        msgbox.exec()
                    else:
                        curs = datadb.cursor()
                        curs.execute(
                            "SELECT pathbt FROM bttable WHERE datebt=? AND nom_producteurbt=? AND perebt=? AND quantitebt=?",
                            (dateentré, producteur, pere, quantite))
                        result = curs.fetchall()
                        if result:
                            # Iterate over the paths and open each file
                            for row in result:
                                file_path = row[0]
                                if file_path:
                                    doc_pdf = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False).name
                                    try:
                                        a = self.progress_bar()
                                        sys.stderr = open("consoleoutput.log", "w")
                                        convert(file_path, doc_pdf)
                                        # Open the resulting .pdf file using the default associated application
                                        # os.startfile(doc_pdf, 'open')
                                        app_path = 'C:\\Program Files\\Okular\\bin\\okular.exe'
                                        subprocess.Popen([app_path, doc_pdf])
                                    except Exception as e:
                                        print(e)
                                    curs.execute(
                                        "INSERT INTO bttable_bulletin_sortie (datebts, nom_producteurbts, perebts, quantites, pathbts, date_sortie) VALUES (?,?,?,?,?,?)",
                                        (dateentré, producteur, pere, quantite, file_path, current_day))
                                    datadb.commit()
                            datadb.close()

                        else:
                            msgbox = QtWidgets.QMessageBox()
                            msgbox.setWindowTitle('Confirmation')
                            msgbox.setText('Le Bulletin ne existe pas.')
                            msgbox.exec()
            except sqlite3.Error as e:
                print(e)

        def eticket_bletendre(self):
            try:
                dialogor = QtWidgets.QMessageBox()
                dialogor.setFixedSize(300, 600)
                dialogor.setStyleSheet("""
                                                  QWidget {
                                                      color: #000000;
                                                      background-color: #D8F9DB;
                                                      border-width: 1px;
                                                      border-color: #1e1e1e;
                                                      border-style: solid;
                                                      border-radius: 6;
                                                      padding: 3px;
                                                      font-size: 18px;
                                                      padding-left: 5px;
                                                      padding-right: 5px;
                                                  }
                                                  QWidget:item:hover {
                                                      background-color: #3daee9;
                                                      color: #eff0f1;
                                                  }
                                                  QWidget:item:selected {
                                                      background-color: #3daee9;
                                                  }
                                                  QWidget:disabled {
                                                      color: #454545;
                                                      background-color: #31363b;
                                                  }
                                                  QPushButton {
                                                      color: #000000;
                                                      background-color:#84dbc8;
                                                      border-width: 1px;
                                                      border-color: #1e1e1e;
                                                      border-style: solid;
                                                      border-radius: 6;
                                                      padding: 3px;
                                                      font-size: 12px;
                                                      padding-left: 5px;
                                                      padding-right: 5px;
                                                      min-width: 40px;
                                                  }
                                                  QPushButton:disabled {
                                                      background-color: #31363b;
                                                      border-width: 1px;
                                                      border-color: #454545;
                                                      border-style: solid;
                                                      padding-top: 5px;
                                                      padding-bottom: 5px;
                                                      padding-left: 10px;
                                                      padding-right: 10px;
                                                      border-radius: 2px;
                                                      color: #454545;
                                                  }
                                                  QPushButton:pressed {
                                                      background-color: #3daee9;
                                                      padding-top: -15px;
                                                      padding-bottom: -17px;
                                                  }
                                                  QPushButton:hover {
                                                      border: 1px solid #ff8c00;
                                                      color: #eff0f1;
                                                  }
                                                  QLabel {
                                                      font-size: 18px;
                                                      border: 0px solid orange;
                                                  }
                                                  QLineEdit
            {
                background-color: #ffffff;
                padding: 1px;
                border-style: solid;
                border: 1px solid #76797C;
                border-radius: 0px;
                color: #000000;
            }
            QDoubleSpinBox
            {
                background-color: #ffffff;
                padding: 1px;
                border-style: solid;
                border: 1px solid #76797C;
                border-radius: 0px;
                color:#000000;
                font-size: 11px;
                font-weight: bold;

            }
            QComboBox
            {
               background-color: #ffffff;
                border-style: solid;
                border: 1px solid #76797C;
                border-radius: 2px;
                min-width: 40px;
            }
                                              """)
                dialogor.setWindowTitle("Sélectionnez le Nom et Prénom et pére")
                dialogor.setText(
                    "BLE TENDRE \t\t\t\n\t\t\t\t\t\t\n\t\t\t\t\t\t\t\n\t\t\t\t\n\t\t\t\t\t\t\t\t\t\t\n\t\t\t\t")
                nomproductureor = QtWidgets.QLineEdit(dialogor)
                nomproductureor.setPlaceholderText("Nom et Prénom")
                nomproductureor.resize(180, 30)
                nomproductureor.move(40, 125)
                pereor = QtWidgets.QLineEdit(dialogor)
                pereor.setPlaceholderText("Pére")
                pereor.resize(180, 30)
                pereor.move(240, 125)
                matricultxt = QtWidgets.QLabel("Matricul", dialogor)
                matricultxt.setGeometry(QtCore.QRect(30, 165, 100, 20))
                mator = QtWidgets.QLineEdit(dialogor)
                mator.setInputMask("99999-999-99")
                mator.resize(180, 30)
                mator.move(40, 190)
                pointdecollector = QtWidgets.QComboBox(dialogor)
                pointdecollector.setPlaceholderText("Point de collect")
                pointdecollector.addItem('')
                pointdecollector.addItem('Dock central')
                pointdecollector.addItem('Magasin Zemmoura')
                pointdecollector.addItem('Magasin Kef-lazreg')
                pointdecollector.addItem('Magasin Messra')
                pointdecollector.addItem('Magasin Marche gros belacel')
                pointdecollector.resize(180, 30)
                pointdecollector.move(40, 65)
                # pstxt = QtWidgets.QLabel("PS", dialogor)
                # pstxt.setGeometry(QtCore.QRect(240, 40, 100, 20))
                psor = QLineEdit(dialogor)
                validator = QDoubleValidator(10, 1000, 2)
                psor.setValidator(validator)

                # psor.setInputMask("99.99")
                psor.setPlaceholderText("PS")
                psor.resize(180, 30)
                psor.move(240, 65)
                quantiteptext = QtWidgets.QLabel("Quantité", dialogor)
                quantiteptext.setGeometry(QtCore.QRect(230, 165, 100, 20))
                quantiteor = QtWidgets.QDoubleSpinBox(dialogor)
                quantiteor.setRange(1, 10000)
                quantiteor.setSuffix('  QX')
                quantiteor.setSpecialValueText(" ")
                quantiteor.setValue(0.0)
                quantiteor.resize(180, 30)
                quantiteor.move(240, 190)
                ok_button = QtWidgets.QPushButton("OK", dialogor)
                cancel_button = QtWidgets.QPushButton("Cancel", dialogor)
                dialogor.addButton(ok_button, QtWidgets.QMessageBox.ButtonRole.ActionRole)
                dialogor.addButton(cancel_button, QtWidgets.QMessageBox.ButtonRole.ActionRole)
                dialogor.exec()
                current_date = QtCore.QDate.currentDate().toString(QtCore.QDate.currentDate().toString("yyyy-MM-dd"))
                current_day = current_date.split("-")[2]
                current_month = current_date.split("-")[1]
                current_year = current_date.split("-")[0]
                current_daya = f"{current_day}-{current_month}-{current_year}"

                if dialogor.clickedButton() == ok_button:
                    producteur = nomproductureor.text()
                    pere = pereor.text()
                    quantite = quantiteor.text()
                    pointdecolect = pointdecollector.currentText()
                    mator = mator.text()
                    psor = psor.text()

                    eticketorge = DocxTemplate("eticket/Eticket_BleTendre/Eticket_BleTendre.docx")
                    eticketorge.render(
                        {"dt": current_date, "ptt": pointdecolect, "nmp": producteur, "mat": mator, "qtt": quantite,
                         "ps": psor, "per": pere})
                    name = f"B.D_{producteur}_{current_daya}.docx"
                    eticketorge.save("eticket/Eticket_BleDur/" + name)
                    doc_names = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
                    doc_pdf = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False).name
                    eticketorge.save(doc_names)
                    try:
                        if doc_names:
                            a = self.progress_bar()
                            sys.stderr = open("consoleoutput.log", "w")
                            convert(doc_names, doc_pdf)
                            # Open the resulting .pdf file using the default associated application
                            # os.startfile(doc_pdf, 'open')
                            app_path = 'C:\\Program Files\\Okular\\bin\\okular.exe'
                            subprocess.Popen([app_path, doc_pdf])
                    except Exception as e:
                        print(e)
            except Exception as e:
                print(e)

        def progress_bar(self):
            self.widgetprogress = QtWidgets.QDialog()
            self.widgetprogress.setStyleSheet(""" QWidget
                                   {
                                       color: #000000;
                                       background-color: #ffffff;
                                       border-width: 1px;
                                       border-color: #1e1e1e;
                                       border-style: solid;
                                       border-radius: 6;
                                       padding: 0px;
                                       font-size: 18px;
                                       padding-left: 1px;
                                       padding-right: 1px
                                   }
                                   QWidget:item:hover
                                   {
                                       background-color: #3daee9;
                                       color: #eff0f1;
                                   }
                                   QWidget:item:selected
                                   {
                                       background-color: #3daee9;
                                   }
                                   QWidget:disabled
                                   {
                                       color: #454545;
                                       background-color: #31363b;
                                   }
                                   QPushButton
                                   {
                                       color: #000000;
                                       background-color:#ade3e7;
                                       border-width: 1px;
                                       border-color: #1e1e1e;
                                       border-style: solid;
                                       border-radius: 6;
                                       padding: 3px;
                                       font-size: 12px;
                                       padding-left: 5px;
                                       padding-right: 5px;
                                       min-width: 40px
                                   }
                                   QPushButton:disabled
                                   {
                                       background-color: #31363b;
                                       border-width: 1px;
                                       border-color: #454545;
                                       border-style: solid;
                                       padding-top: 5px;
                                       padding-bottom: 5px;
                                       padding-left: 10px;
                                       padding-right: 10px;
                                       border-radius: 2px;
                                       color: #454545;
                                   }

                                   QPushButton:pressed
                                   {
                                       background-color: #3daee9;
                                       padding-top: -15px;
                                       padding-bottom: -17px;
                                   }
                                   QPushButton:hover
                                   {
                                       border: 1px solid #ff8c00;
                                       color: #000000;
                                   }
                                    QLabel
                                   {
                                       font-size: 18px;
                                       border: 0px solid orange;
                                   }

                               """)
            self.widgetprogress.setWindowTitle("جاري تحميل الملف يرجى الانتظار ")
            self.widgetprogress.setGeometry(550, 450, 250, 20)
            self.progressBar = QtWidgets.QProgressBar(self.widgetprogress)
            self.progressBar.setGeometry(10, 10, 200, 10)
            self.progressBar.setMinimum(0)
            self.progressBar.setMaximum(100)
            self.progressBar.setStyleSheet("""QProgressBar
           {
           border: solid grey;
           border-radius: 15px;
           color: black;
           }
           QProgressBar::chunk 
           {
           background-color: #05B8CC;
           border-radius :15px;
           }      """)
            self.progressBar.setAlignment(Qt.AlignmentFlag.AlignCenter)
            self.vbox = QVBoxLayout(self.widgetprogress)
            self.vbox.addWidget(self.progressBar)
            self.timer = QtCore.QTimer()
            self.timer.timeout.connect(self.update_progress)
            self.timer.start(5)  # Update progress every
            self.widgetprogress.show()

        def update_progress(self):
            # Simulate file download progress
            current_value = self.progressBar.value()
            if current_value < 100:
                new_value = current_value + 10
                self.progressBar.setValue(new_value)
                if current_value == 99:
                    self.timer.stop()
                    self.progressBar.close()
                    self.widgetprogress.close()

        def agréage(self, MainWindow):
            MainWindow.setObjectName("MainWindow")
            MainWindow.resize(1350, 700)

            self.centralwidget = QtWidgets.QWidget(MainWindow)
            self.centralwidget.setObjectName("centralwidget")
            self.verticalLayout = QtWidgets.QVBoxLayout(self.centralwidget)
            self.verticalLayout.setObjectName("verticalLayout")
            self.tabagreage = QtWidgets.QTabWidget(self.centralwidget)

            self.tabagreage.setTabShape(QtWidgets.QTabWidget.TabShape.Rounded)
            self.tabagreage.setIconSize(QtCore.QSize(20, 20))
            self.tabagreage.setObjectName("tabagreage")

            ####################################################################################
            ########################BLE DUR##############################
            #############################################
            self.tabbldur = QtWidgets.QWidget()
            self.tabbldur.setObjectName("tabbldur")
            self.tabbldur.setStyleSheet("""QToolTip
    {
        border: 1px solid #76797C;
        background-color:  #fff8b0;
        color: white;
        padding: 5px;
        opacity: 200;
    }

    QWidget
    {
        color: #000000;
        background-color: #ffc6c9;
        selection-background-color:#3daee9;
        selection-color: #eff0f1;
        background-clip: border;
        border-image: none;
        border: 0px transparent black;
        outline: 0;
    }

    QWidget:item:hover
    {
        background-color: #3daee9;
        color: #eff0f1;
    }

    QWidget:item:selected
    {
        background-color: #3daee9;
    }



    QWidget:disabled
    {
        color: #454545;
        background-color: #31363b;
    }

    QAbstractItemView
    {
        alternate-background-color: #31363b;
        color: #eff0f1;
        border: 1px solid 3A3939;
        border-radius: 2px;
    }

    QWidget:focus, QMenuBar:focus
    {
        border: 1px solid #3daee9;
    }

    QTabWidget:focus, QCheckBox:focus, QRadioButton:focus, QSlider:focus
    {
        border: none;
    }

    QLineEdit
    {
        background-color: #ffffff;
        padding: 1px;
        border-style: solid;
        border: 1px solid #76797C;
        border-radius: 2px;
        color: #000000;
        font-size:12px;
        font-weight:bold;
    }
    QDoubleSpinBox
    {
        background-color: #ffffff;
        padding: 1px;
        border-style: solid;
        border: 1px solid #76797C;
        border-radius: 2px;
        color:#000000;
        font-size:12px;
        font-weight:bold;

    }
    QDoubleSpinBox::drop-down
    {
        subcontrol-origin: padding;
        subcontrol-position: top right;
        width: 1px;

        border-left-width: 0px;
        border-left-color: #232629;
        border-left-style: solid;
        border-top-right-radius: 1px;
        border-bottom-right-radius: 1px;
    }



    QGroupBox {
        border:1px solid #76797C;
        border-radius: 2px;
        margin-top: 20px;
    }

    QGroupBox::title {
        subcontrol-origin: margin;
        subcontrol-position: top center;
        padding-left: 10px;
        padding-right: 10px;
        padding-top: 10px;
    }

    QAbstractScrollArea
    {
        border-radius: 2px;
        border: 1px solid #76797C;
        background-color: transparent;
    }

    QScrollBar:horizontal
    {
        height: 15px;
        margin: 3px 15px 3px 15px;
        border: 1px transparent #2A2929;
        border-radius: 4px;
        background-color: #2A2929;
    }

    QScrollBar::handle:horizontal
    {
        background-color: #605F5F;
        min-width: 5px;
        border-radius: 4px;
    }

    QScrollBar::add-line:horizontal
    {
        margin: 0px 3px 0px 3px;
        border-image: url(:/qss_icons/Dark_rc/right_arrow_disabled.png);
        width: 10px;
        height: 10px;
        subcontrol-position: right;
        subcontrol-origin: margin;
    }

    QScrollBar::sub-line:horizontal
    {
        margin: 0px 3px 0px 3px;
        border-image: url(:/qss_icons/Dark_rc/left_arrow_disabled.png);
        height: 10px;
        width: 10px;
        subcontrol-position: left;
        subcontrol-origin: margin;
    }

    QScrollBar::add-line:horizontal:hover,QScrollBar::add-line:horizontal:on
    {
        border-image: url(:/qss_icons/Dark_rc/right_arrow.png);
        height: 10px;
        width: 10px;
        subcontrol-position: right;
        subcontrol-origin: margin;
    }


    QScrollBar::sub-line:horizontal:hover, QScrollBar::sub-line:horizontal:on
    {
        border-image: url(:/qss_icons/Dark_rc/left_arrow.png);
        height: 10px;
        width: 10px;
        subcontrol-position: left;
        subcontrol-origin: margin;
    }

    QScrollBar::up-arrow:horizontal, QScrollBar::down-arrow:horizontal
    {
        background: none;
    }


    QScrollBar::add-page:horizontal, QScrollBar::sub-page:horizontal
    {
        background: none;
    }

    QScrollBar:vertical
    {
        background-color: #2A2929;
        width: 15px;
        margin: 15px 3px 15px 3px;
        border: 1px transparent #2A2929;
        border-radius: 4px;
    }

    QScrollBar::handle:vertical
    {
        background-color: #605F5F;
        min-height: 5px;
        border-radius: 4px;
    }

    QScrollBar::sub-line:vertical
    {
        margin: 3px 0px 3px 0px;
        border-image: url(:/qss_icons/Dark_rc/up_arrow_disabled.png);
        height: 10px;
        width: 10px;
        subcontrol-position: top;
        subcontrol-origin: margin;
    }

    QScrollBar::add-line:vertical
    {
        margin: 3px 0px 3px 0px;
        border-image: url(:/qss_icons/Dark_rc/down_arrow_disabled.png);
        height: 10px;
        width: 10px;
        subcontrol-position: bottom;
        subcontrol-origin: margin;
    }

    QScrollBar::sub-line:vertical:hover,QScrollBar::sub-line:vertical:on
    {

        border-image: url(:/qss_icons/Dark_rc/up_arrow.png);
        height: 10px;
        width: 10px;
        subcontrol-position: top;
        subcontrol-origin: margin;
    }


    QScrollBar::add-line:vertical:hover, QScrollBar::add-line:vertical:on
    {
        border-image: url(:/qss_icons/Dark_rc/down_arrow.png);
        height: 10px;
        width: 10px;
        subcontrol-position: bottom;
        subcontrol-origin: margin;
    }

    QScrollBar::up-arrow:vertical, QScrollBar::down-arrow:vertical
    {
        background: none;
    }


    QScrollBar::add-page:vertical, QScrollBar::sub-page:vertical
    {
        background: none;
    }

    QTextEdit
    {
        background-color: #fffff1;
        color: #000000;
        border: 1px solid #76797C;
        font-size:12px;
        font-weight:bold;
    }

    QPlainTextEdit
    {
        background-color: #232629;;
        color: #eff0f1;
        border-radius: 2px;
        border: 1px solid #76797C;
    }

    QHeaderView::section
    {
        background-color: #76797C;
        color: #eff0f1;
        padding: 5px;
        border: 1px solid #76797C;
    }

    QSizeGrip {
        image: url(:/qss_icons/Dark_rc/sizegrip.png);
        width: 12px;
        height: 12px;
    }


    QMainWindow::separator
    {
        background-color: #31363b;
        color: white;
        padding-left: 4px;
        spacing: 2px;
        border: 1px dashed #76797C;
    }

    QMainWindow::separator:hover
    {

        background-color: #787876;
        color: white;
        padding-left: 4px;
        border: 1px solid #76797C;
        spacing: 2px;
    }


    QMenu::separator
    {
        height: 1px;
        background-color: #76797C;
        color: white;
        padding-left: 4px;
        margin-left: 10px;
        margin-right: 5px;
    }


    QFrame
    {
        border-radius: 2px;
        border: 1px solid #76797C;
    }

    QFrame[frameShape="0"]
    {
        border-radius: 2px;
        border: 1px transparent #76797C;
    }

    QStackedWidget
    {
        border: 1px transparent black;
    }


    QPushButton
    {
        color: #000000;
        background-color:#84dbc8;
        border-width: 1px;
        border-color: #1e1e1e;
        border-style: solid;
        border-radius: 6;
        padding: 3px;
        font-size: 12px;
        padding-left: 5px;
        padding-right: 5px;
        min-width: 40px;

    }

    QPushButton:disabled
    {
        background-color: #31363b;
        border-width: 1px;
        border-color: #454545;
        border-style: solid;
        padding-top: 5px;
        padding-bottom: 5px;
        padding-left: 10px;
        padding-right: 10px;
        border-radius: 2px;
        color: #454545;
    }

    QPushButton:focus {
        background-color: #3daee9;
        color: white;
    }

    QPushButton:pressed
    {
        background-color: #3daee9;
        padding-top: -15px;
        padding-bottom: -17px;
    }

    

    QPushButton:checked{
        background-color: #76797C;
        border-color: #6A6969;
    }

    QComboBox:hover,QDoubleSpinBox:Hover,QPushButton:hover,QAbstractSpinBox:hover,QLineEdit:hover,QTextEdit:hover,QPlainTextEdit:hover,QAbstractView:hover,QTreeView:hover
    {
        border: 1px solid #ff8c00;
        color: #000000;
    }

    QComboBox {
    background-color: #ffffff;
    border: 1px solid #76797C;
    color:#000000;
    border-radius: 0.25em;
    padding: 0.10em 0.10em;
    font-size:12px;
    font-weight:bold;
    cursor: pointer;
}

QComboBox::drop-down {
    subcontrol-origin: padding;
    subcontrol-position: top right;
    width: 1.3em;
    border-left: 0px solid #777;
    border-radius: 0.25em;
}

QComboBox::drop-down::icon {
    image: url('E:/pythonProject_moullin-application.3.5/images/down-arroww.png');
}

        QComboBox:on
        {
            padding-top: 0px;
            padding-left: 0px;        
            selection-background-color: #e4f0f1;
        }
        QComboBox QAbstractItemView
        {
            background-color: #ffffff;
            border-radius: 2px;
            border: 1px solid #76797C;
            color:#000000;
            selection-background-color: #000000;
        }
                                                         
            QComboBox:hover,QDoubleSpinBox:Hover,QPushButton:hover,QAbstractSpinBox:hover,QLineEdit:hover,QTextEdit:hover,QPlainTextEdit:hover,QAbstractView:hover,QTreeView:hover
            {
                border: 1px solid #ff8c00;
                color: #eff0f1;
            }

    QLabel
    {
        border: 2px solid black;
        font-size:13px;
        font-weight:bold;
    }

    QTabWidget{
        border: 0px transparent black;
    }

    QTabWidget::pane {
        border: 1px solid #76797C;
        padding: 5px;
        margin: 0px;
    }

    QTabBar
    {
        qproperty-drawBase: 0;
        left: 5px; /* move to the right by 5px */
        border-radius: 3px;
    }

    QTabBar:focus
    {
        border: 0px transparent black;
    }

    QTabBar::close-button  {
        image: url(:/qss_icons/Dark_rc/close.png);
        background: transparent;
    }

    QTabBar::close-button:hover
    {
        image: url(:/qss_icons/Dark_rc/close-hover.png);
        background: transparent;
    }

    QTabBar::close-button:pressed {
        image: url(:/qss_icons/Dark_rc/close-pressed.png);
        background: transparent;
    }

    /* TOP TABS */
    QTabBar::tab:top {
        color: #eff0f1;
        border: 1px solid #76797C;
        border-bottom: 1px transparent black;
        background-color: #31363b;
        padding: 5px;
        min-width: 50px;
        border-top-left-radius: 2px;
        border-top-right-radius: 2px;
    }

    QTabBar::tab:top:!selected
    {
        color: #eff0f1;
        background-color: #54575B;
        border: 1px solid #76797C;
        border-bottom: 1px transparent black;
        border-top-left-radius: 2px;
        border-top-right-radius: 2px;    
    }

    QTabBar::tab:top:!selected:hover {
        background-color: #3daee9;
    }

    /* BOTTOM TABS */
    QTabBar::tab:bottom {
        color: #eff0f1;
        border: 1px solid #76797C;
        border-top: 1px transparent black;
        background-color: #31363b;
        padding: 5px;
        border-bottom-left-radius: 2px;
        border-bottom-right-radius: 2px;
        min-width: 50px;
    }

    QTabBar::tab:bottom:!selected
    {
        color: #eff0f1;
        background-color: #54575B;
        border: 1px solid #76797C;
        border-top: 1px transparent black;
        border-bottom-left-radius: 2px;
        border-bottom-right-radius: 2px;
    }

    QTabBar::tab:bottom:!selected:hover {
        background-color: #3daee9;
    }

    /* LEFT TABS */
    QTabBar::tab:left {
        color: #eff0f1;
        border: 1px solid #76797C;
        border-left: 1px transparent black;
        background-color: #31363b;
        padding: 5px;
        border-top-right-radius: 2px;
        border-bottom-right-radius: 2px;
        min-height: 50px;
    }

    QTabBar::tab:left:!selected
    {
        color: #eff0f1;
        background-color: #54575B;
        border: 1px solid #76797C;
        border-left: 1px transparent black;
        border-top-right-radius: 2px;
        border-bottom-right-radius: 2px;
    }

    QTabBar::tab:left:!selected:hover {
        background-color: #3daee9;
    }


    /* RIGHT TABS */
    QTabBar::tab:right {
        color: #eff0f1;
        border: 1px solid #76797C;
        border-right: 1px transparent black;
        background-color: #31363b;
        padding: 5px;
        border-top-left-radius: 2px;
        border-bottom-left-radius: 2px;
        min-height: 50px;
    }

    QTabBar::tab:right:!selected
    {
        color: #eff0f1;
        background-color: #54575B;
        border: 1px solid #76797C;
        border-right: 1px transparent black;
        border-top-left-radius: 2px;
        border-bottom-left-radius: 2px;
    }

    QTabBar::tab:right:!selected:hover {
        background-color: #3daee9;
    }

    QTabBar QToolButton::right-arrow:enabled {
         image: url(:/qss_icons/Dark_rc/right_arrow.png);
     }

     QTabBar QToolButton::left-arrow:enabled {
         image: url(:/qss_icons/Dark_rc/left_arrow.png);
     }

    QTabBar QToolButton::right-arrow:disabled {
         image: url(:/qss_icons/Dark_rc/right_arrow_disabled.png);
     }

     QTabBar QToolButton::left-arrow:disabled {
         image: url(:/qss_icons/Dark_rc/left_arrow_disabled.png);
     }


    QDockWidget {
        background: #31363b;
        border: 1px solid #403F3F;
        titlebar-close-icon: url(:/qss_icons/Dark_rc/close.png);
        titlebar-normal-icon: url(:/qss_icons/Dark_rc/undock.png);
    }

    QDockWidget::close-button, QDockWidget::float-button {
        border: 1px solid transparent;
        border-radius: 2px;
        background: transparent;
    }

    QDockWidget::close-button:hover, QDockWidget::float-button:hover {
        background: rgba(255, 255, 255, 10);
    }

    QDockWidget::close-button:pressed, QDockWidget::float-button:pressed {
        padding: 1px -1px -1px 1px;
        background: rgba(255, 255, 255, 10);
    }


    QSlider::groove:horizontal {
        border: 1px solid #565a5e;
        height: 4px;
        background: #565a5e;
        margin: 0px;
        border-radius: 2px;
    }

    QSlider::handle:horizontal {
        background: #232629;
        border: 1px solid #565a5e;
        width: 16px;
        height: 16px;
        margin: -8px 0;
        border-radius: 9px;
    }

    QSlider::groove:vertical {
        border: 1px solid #565a5e;
        width: 4px;
        background: #565a5e;
        margin: 0px;
        border-radius: 3px;
    }

    QSlider::handle:vertical {
        background: #232629;
        border: 1px solid #565a5e;
        width: 16px;
        height: 16px;
        margin: 0 -8px;
        border-radius: 9px;
    }

    QToolButton {
        background-color: transparent;
        border: 1px transparent #76797C;
        border-radius: 2px;
        margin: 3px;
        padding: 5px;
    }

    QToolButton[popupMode="1"] { /* only for MenuButtonPopup */
     padding-right: 20px; /* make way for the popup button */
     border: 1px #76797C;
     border-radius: 5px;
    }

    QToolButton[popupMode="2"] { /* only for InstantPopup */
     padding-right: 10px; /* make way for the popup button */
     border: 1px #76797C;
    }


    QToolButton:hover, QToolButton::menu-button:hover {
        background-color: transparent;
        border: 1px solid #3daee9;
        padding: 5px;
    }

    QToolButton:checked, QToolButton:pressed,
            QToolButton::menu-button:pressed {
        background-color: #3daee9;
        border: 1px solid #3daee9;
        padding: 5px;
    }

    /* the subcontrol below is used only in the InstantPopup or DelayedPopup mode */
    QToolButton::menu-indicator {
        background-color:ff8c00;
        top: -7px; left: -2px; /* shift it a bit */
    }

    /* the subcontrols below are used only in the MenuButtonPopup mode */
    QToolButton::menu-button {
        border: 1px transparent #76797C;
        border-top-right-radius: 6px;
        border-bottom-right-radius: 6px;
        /* 16px width + 4px for border = 20px allocated above */
        width: 16px;
        outline: none;
    }

    QToolButton::menu-arrow {
       background-color:ff8c00;
    }

    QToolButton::menu-arrow:open {
        border: 1px solid #76797C;
    }

    QPushButton::menu-indicator  {
        subcontrol-origin: padding;
        subcontrol-position: bottom right;
        left: 8px;
    }

    QTableView
    {
        border: 1px solid #76797C;
        gridline-color: #31363b;
        background-color: #232629;
    }


    QTableView, QHeaderView
    {
        border-radius: 0px;
    }

    QTableView::item:pressed, QListView::item:pressed, QTreeView::item:pressed  {
        background: #3daee9;
        color: #eff0f1;
    }

    QTableView::item:selected:active, QTreeView::item:selected:active, QListView::item:selected:active  {
        background: #3daee9;
        color: #eff0f1;
    }


    QHeaderView
    {
        background-color: #31363b;
        border: 1px transparent;
        border-radius: 0px;
        margin: 0px;
        padding: 0px;

    }

    QHeaderView::section  {
        background-color: #31363b;
        color: #eff0f1;
        padding: 5px;
        border: 1px solid #76797C;
        border-radius: 0px;
        text-align: center;
    }

    QHeaderView::section::vertical::first, QHeaderView::section::vertical::only-one
    {
        border-top: 1px solid #76797C;
    }

    QHeaderView::section::vertical
    {
        border-top: transparent;
    }

    QHeaderView::section::horizontal::first, QHeaderView::section::horizontal::only-one
    {
        border-left: 1px solid #76797C;
    }

    QHeaderView::section::horizontal
    {
        border-left: transparent;
    }


    QHeaderView::section:checked
     {
        color: white;
        background-color: #334e5e;
     }

     /* style the sort indicator */
    QHeaderView::down-arrow {
        image: url(:/qss_icons/Dark_rc/down_arrow.png);
    }

    QHeaderView::up-arrow {
        image: url(:/qss_icons/Dark_rc/up_arrow.png);
    }


    QTableCornerButton::section {
        background-color: #31363b;
        border: 1px transparent #76797C;
        border-radius: 0px;
    }

    QToolBox  {
        padding: 5px;
        border: 1px transparent black;
    }

    QToolBox::tab {
        color: #eff0f1;
        background-color: #31363b;
        border: 1px solid #76797C;
        border-bottom: 1px transparent #31363b;
        border-top-left-radius: 5px;
        border-top-right-radius: 5px;
    }

    QToolBox::tab:selected { /* italicize selected tabs */
        font: italic;
        background-color: #31363b;
        border-color: #3daee9;
     }

    QStatusBar::item {
        border: 0px transparent dark;
     }


    QFrame[height="3"], QFrame[width="3"] {
        background-color: #76797C;
    }




    QDateTimeEdit
    {
        background-color:#ffffff;
        border-style: solid;
        border: 1px solid #76797C;
        border-radius: 2px;
        padding: 1px;
        min-width: 75px;
    }

    QDateTimeEdit 
    {
        padding-top: 2px;
        padding-left: 2px;
        selection-background-color: #ffffff;
    }

    QDateTimeEdit QAbstractItemView
    {
        background-color: #ffffff;
        border-radius: 2px;
        border: 1px solid #3375A3;
        selection-background-color:ff8c00;
    }

    QDateTimeEdit::drop-down
    {
        subcontrol-origin: padding;
        subcontrol-position: top right;
        width: 15px;
        border-left-width: 0px;
        border-left-color: darkgray;
        border-left-style: solid;
        border-top-right-radius: 3px;
        border-bottom-right-radius: 3px;
    }""")
            self.tabagreage.addTab(self.tabbldur, "")
            self.cclstxt = QtWidgets.QLabel("<h2>COOPERATIVE DES CEREALES ET LEGUMES SECS DE RELIZANE<h2/>",
                                            self.tabbldur)
            self.cclstxt.resize(600, 40)
            self.cclstxt.move(400, 2)

            self.font = QtGui.QFont()
            self.font.setBold(True)
            self.font.setPointSize(10)

            self.txtpsfont = QtGui.QFont()
            self.txtpsfont.setBold(True)
            self.txtpsfont.setPointSize(9)

            self.rtotaldemfont = QtGui.QFont("color:black")
            self.rtotaldemfont.setBold(True)
            self.rtotaldemfont.setPointSize(12)

            self.bulletin_agréagetxt = QtWidgets.QLabel("<h2>Bulletin D'Agréage<h2/>", self.tabbldur)
            self.bulletin_agréagetxt.move(600, 30)
            self.bulletin_agréagetxt.resize(200, 40)

            self.éspécetxt = QtWidgets.QLabel("<h2>Espéce:Blé Dur<h2/>", self.tabbldur)
            self.éspécetxt.move(630, 60)

            self.n_bultinbdtxt = QtWidgets.QLabel("N° Bulletin:", self.tabbldur)
            self.n_bultinbdtxt.move(20, 60)
            self.n_bultinbd = QtWidgets.QLineEdit(self.tabbldur, readOnly=True)
            self.n_bultinbd.move(150, 60)
            self.n_bultinbd.resize(150, 30)

            self.nome_du_producteurbdtxt = QtWidgets.QLabel("Nom du Producteur:", self.tabbldur)
            self.nome_du_producteurbdtxt.move(20, 100)
            self.nome_du_producteurbd = QtWidgets.QLineEdit(self.tabbldur)
            self.nome_du_producteurbd.move(150, 98)
            self.nome_du_producteurbd.resize(150, 30)

            self.péretxtbd = QtWidgets.QLabel("Pére:", self.tabbldur)
            self.péretxtbd.setGeometry(QtCore.QRect(315, 98, 100, 20))
            self.pérebd = QtWidgets.QLineEdit(self.tabbldur)
            self.pérebd.setGeometry(QtCore.QRect(410, 98, 150, 30))

            self.n_cartebdtxt = QtWidgets.QLabel("N.C d'identité:", self.tabbldur)
            self.n_cartebdtxt.setGeometry(QtCore.QRect(315, 138, 100, 20))
            self.n_cartebd = QtWidgets.QLineEdit(self.tabbldur)
            self.n_cartebd.setGeometry(QtCore.QRect(410, 138, 150, 30))
            self.n_cartebd.setInputMask('99999999')

            self.imatriculebdtxt = QtWidgets.QLabel("Adresse:", self.tabbldur)
            self.imatriculebdtxt.move(20, 140)
            self.adressebd = QtWidgets.QLineEdit(self.tabbldur)
            self.adressebd.setInputMask("99999-999-99")
            self.adressebd.move(150, 138)
            self.adressebd.resize(150, 30)

            self.pointdecollectbdtxt = QtWidgets.QLabel("Point de collecte", self.tabbldur)
            self.pointdecollectbdtxt.move(20, 180)
            self.pointdecollectbd = QtWidgets.QComboBox(self.tabbldur)
            self.pointdecollectbd.setStyleSheet("background-color:#ffffff;color:#000000")
            self.pointdecollectbd.move(150, 178)
            self.pointdecollectbd.resize(150, 30)
            self.pointdecollectbd.addItem('')
            self.pointdecollectbd.addItem('Dock central')
            self.pointdecollectbd.addItem('Magasin Zemmoura')
            self.pointdecollectbd.addItem('Magasin Kef-lazreg')
            self.pointdecollectbd.addItem('Magasin Messra')
            self.pointdecollectbd.addItem('Marche gros belacel')
            self.pointdecollectbd.addItem('Station Mendes')
            self.pointdecollectbd.addItem('Nouvelles S.Mendes')

            self.dattereceptionbdtxt = QtWidgets.QLabel("Relizane le :", self.tabbldur)
            self.dattereceptionbdtxt.setGeometry(QtCore.QRect(880, 100, 150, 23))
            self.dattereceptiontbd = QtWidgets.QDateTimeEdit(self.tabbldur)
            self.dattereceptiontbd.setGeometry(QtCore.QRect(1000, 100, 150, 30))
            self.datedaytime = QDate.currentDate()
            self.dattereceptiontbd.setDisplayFormat("dd-MM-yyyy hh:mm")
            self.dattereceptiontbd.setDate(self.datedaytime)

            self.agréeeurbdtxt = QtWidgets.QLabel("Nom de l’Agréeur:", self.tabbldur)
            self.agréeeurbdtxt.setGeometry(QtCore.QRect(880, 135, 150, 23))
            self.agréeeurcombobd = QtWidgets.QComboBox(self.tabbldur, editable=True)
            self.agréeeurcombobd.setStyleSheet("background-color:#ffffff;color:#000000")
            self.agréeeurcombobd.addItem("")
            self.agréeeurcombobd.addItem("FELOUAH OMAR")
            self.agréeeurcombobd.addItem("BEKHEDDA AEK")
            self.agréeeurcombobd.addItem("BENAISSA YOUCEF")
            self.agréeeurcombobd.addItem("REZZAG SOFIANE ")
            self.agréeeurcombobd.addItem("BELBACHA M.NADIR")
            self.agréeeurcombobd.move(1000, 135)
            self.agréeeurcombobd.resize(150, 30)

            self.quantitebdtxt = QtWidgets.QLabel("Quantité", self.tabbldur)
            self.quantitebdtxt.setGeometry(QtCore.QRect(880, 180, 150, 23))

            self.quantitebd = QtWidgets.QDoubleSpinBox(self.tabbldur)
            self.quantitebd.setRange(1, 10000)
            self.quantitebd.setSuffix('  QX')
            self.quantitebd.setSpecialValueText(" ")
            self.quantitebd.setGeometry(QtCore.QRect(1000, 178, 150, 30))

            self.paramétrebd = QtWidgets.QLabel("Paramètre", self.tabbldur)
            self.paramétrebd.move(30, 205)
            self.paramétrebd.resize(90, 20)
            self.paramétrebd.setFont(self.font)
            self.txtpsfontbd = QtGui.QFont()
            self.txtpsfontbd.setBold(True)
            self.txtpsfontbd.setPointSize(9)
            ################Limites(sans bon ni réf)###############
            self.valeurbd = QtWidgets.QLabel("", self.tabbldur)
            self.valeurbd.move(170, 205)
            self.valeurbd.resize(145, 0)
            self.valeurbd.setFont(self.font)
            ######################Limites(sans bon ni réf)################
            self.psbd = QtWidgets.QLineEdit("Poids spécifique (kg/hl)", self.tabbldur, readOnly=True)
            self.psbd.resize(369, 20)
            self.psbd.move(30, 230)
            self.psbd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            self.psbd.setFont(self.txtpsfont)
            self.psbd.setStyleSheet(" border: 2px solid bleu ;border-radius: 4px;padding: 0px")
            ###############################humidite#############
            self.humiditebd = QtWidgets.QLineEdit("Teneur en eau(%)", self.tabbldur, readOnly=True)
            self.humiditebd.resize(369, 20)
            self.humiditebd.move(30, 251)
            self.humiditebd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            self.humiditebd.setFont(self.txtpsfont)

            #######################ergot#########################
            self.ergotbd = QtWidgets.QLineEdit("Ergo(%  ", self.tabbldur, readOnly=True)
            self.ergotbd.resize(369, 20)
            self.ergotbd.move(30, 272)
            self.ergotbd.setStyleSheet("background-color: #232629")
            self.ergotbd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            self.ergotbd.setFont(self.txtpsfont)

            #########################Graines nuisibles (%)##########
            self.matiére20mm = QtWidgets.QLineEdit("Matiéres qui passent à travers \nle tamis 20 mm x 2.1 mm ",
                                                   self.tabbldur, readOnly=True)
            self.matiére20mm.resize(369, 20)
            self.matiére20mm.move(30, 293)
            self.matiére20mm.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            self.matiére20mm.setFont(self.txtpsfont)
            #############################Débris végétaux (%)########
            self.débrisvébd = QtWidgets.QLineEdit("Les débris végétaux et \nles éléments minéreaux(%)", self.tabbldur,
                                                  readOnly=True)
            self.débrisvébd.resize(369, 20)
            self.débrisvébd.move(30, 314)
            self.débrisvébd.setFont(self.txtpsfont)
            self.débrisvébd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #########################Matière inerte (%)################
            self.grainnuisiblebd = QtWidgets.QLineEdit("Graines nuisibles(%)", self.tabbldur, readOnly=True)
            self.grainnuisiblebd.resize(369, 20)
            self.grainnuisiblebd.move(30, 335)
            self.grainnuisiblebd.setFont(self.txtpsfont)
            self.grainnuisiblebd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            ################################Grains chauffés (%)############################
            self.totalprembd = QtWidgets.QLineEdit("Total(%)", self.tabbldur, readOnly=True)
            self.totalprembd.resize(369, 20)
            self.totalprembd.move(30, 356)
            self.totalprembd.setFont(self.txtpsfont)
            self.totalprembd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            ########################################Grains sans valeur (%)#######################################
            self.grainscassébd = QtWidgets.QLineEdit("Grains cassés(%)", self.tabbldur, readOnly=True)
            self.grainscassébd.resize(369, 20)
            self.grainscassébd.move(30, 377)
            self.grainscassébd.setFont(self.txtpsfont)
            self.grainscassébd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #######################################Grains cariés##########################################
            self.grainsboutésbd = QtWidgets.QLineEdit("Grains fortement boutés", self.tabbldur, readOnly=True)
            self.grainsboutésbd.resize(369, 20)
            self.grainsboutésbd.move(30, 398)
            self.grainsboutésbd.setFont(self.txtpsfont)
            self.grainsboutésbd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")

            #######################################Total(%) 1er#####################################################
            self.grainsrouxbd = QtWidgets.QLineEdit("Grains Roux", self.tabbldur, readOnly=True)
            self.grainsrouxbd.resize(369, 20)
            self.grainsrouxbd.move(30, 419)
            self.grainsrouxbd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            self.grainsrouxbd.setFont(self.txtpsfont)
            ##############################################Grains cassés (%) #########################################################
            self.grainfortementmouchtébd = QtWidgets.QLineEdit("Grains fortement mouchtés(%)", self.tabbldur,
                                                               readOnly=True)
            self.grainfortementmouchtébd.move(30, 440)
            self.grainfortementmouchtébd.resize(369, 20)
            self.grainfortementmouchtébd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            self.grainfortementmouchtébd.setFont(self.txtpsfont)
            #########################################################Gains échaudés (%)#####################################################
            self.grainepunaisésbd = QtWidgets.QLineEdit("Gains punaisés(%)", self.tabbldur, readOnly=True)
            self.grainepunaisésbd.move(30, 461)
            self.grainepunaisésbd.resize(369, 20)
            self.grainepunaisésbd.setFont(self.txtpsfont)
            self.grainepunaisésbd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #####################################################Grains maigres (%)########################################################
            self.grainpiquebd = QtWidgets.QLineEdit("Grains piqués(%) ", self.tabbldur, readOnly=True)
            self.grainpiquebd.move(30, 482)
            self.grainpiquebd.resize(369, 20)
            self.grainpiquebd.setFont(self.txtpsfont)
            self.grainpiquebd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            ##########################################################Grains germés (%)###################################################
            self.totaldembd = QtWidgets.QLineEdit("Total (%)", self.tabbldur, readOnly=True)
            self.totaldembd.move(30, 503)
            self.totaldembd.resize(369, 20)
            self.totaldembd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            self.totaldembd.setFont(self.txtpsfont)
            ##########################################################Grain punaisés (%)#########################################################
            self.grainmitadinésbd = QtWidgets.QLineEdit("Grain mitadinés(%) ", self.tabbldur, readOnly=True)
            self.grainmitadinésbd.move(30, 524)
            self.grainmitadinésbd.resize(369, 20)
            self.grainmitadinésbd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            self.grainmitadinésbd.setFont(self.txtpsfont)

            #######################################################################Grains piqués (%)##########################################
            self.bletendredansbledurbd = QtWidgets.QLineEdit("Blé tendre dans blé dure (%)  ", self.tabbldur,
                                                             readOnly=True)
            self.bletendredansbledurbd.move(30, 545)
            self.bletendredansbledurbd.resize(369, 20)
            self.bletendredansbledurbd.setFont(self.txtpsfont)
            self.bletendredansbledurbd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #######################################################################Grains boutés « faible » (%)#######################################

            ####################################################################Grains boutés  « forte » (%)######################################
            self.totalmitadinésbd = QtWidgets.QLineEdit("Total (%)", self.tabbldur, readOnly=True)
            self.totalmitadinésbd.move(30, 566)
            self.totalmitadinésbd.resize(369, 20)
            self.totalmitadinésbd.setFont(self.txtpsfont)
            self.totalmitadinésbd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")

            self.totalbiniEtrefaction = QtWidgets.QLineEdit("\t\tTotal de bonification et réfaction", self.tabbldur,
                                                            readOnly=True)
            self.totalbiniEtrefaction.move(30, 587)
            self.totalbiniEtrefaction.resize(470, 30)
            self.totalbiniEtrefaction.setFont(self.txtpsfont)
            self.totalbiniEtrefaction.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")

            #################label valeure##############
            self.valeurbd = QtWidgets.QLabel("valeur", self.tabbldur)
            self.valeurbd.move(400, 205)
            self.valeurbd.resize(80, 20)
            self.valeurbd.setFont(self.font)
            ######################Limites(sans bon ni réf)################
            self.vpsbd = QtWidgets.QDoubleSpinBox(self.tabbldur)
            self.vpsbd.setDecimals(3)

            self.vpsbd.setSpecialValueText(' ')
            self.vpsbd.resize(100, 20)
            self.vpsbd.move(400, 230)
            self.vpsbd.setFont(self.txtpsfont)
            self.vpsbd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")
            ###############################humidite#############
            self.vhumiditebd = QtWidgets.QDoubleSpinBox(self.tabbldur)
            self.vhumiditebd.setRange(3, 17)
            self.vhumiditebd.resize(100, 20)
            self.vhumiditebd.setSpecialValueText(' ')
            self.vhumiditebd.move(400, 251)
            self.vhumiditebd.setFont(self.txtpsfont)
            self.vhumiditebd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #######################ergot#########################
            self.vergotbd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=False)
            self.vergotbd.setRange(0, 10)
            self.vergotbd.setSpecialValueText(' ')
            self.vergotbd.resize(100, 20)
            self.vergotbd.move(400, 272)
            self.vergotbd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #########################Graines nuisibles (%)##########
            self.vmatiére20mmbd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=False)
            self.vmatiére20mmbd.setRange(0, 10)
            self.vmatiére20mmbd.setSpecialValueText(' ')
            self.vmatiére20mmbd.resize(100, 20)
            self.vmatiére20mmbd.move(400, 293)
            self.vmatiére20mmbd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #############################Débris végétaux (%)########
            self.vdébrisvébd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=False)
            self.vdébrisvébd.setRange(0, 10)
            self.vdébrisvébd.setSpecialValueText(' ')
            self.vdébrisvébd.resize(100, 20)
            self.vdébrisvébd.move(400, 314)
            self.vdébrisvébd.setFont(self.txtpsfont)
            self.vdébrisvébd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #########################Matière inerte (%)################
            self.vgrainnuisiblebd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=False)
            self.vgrainnuisiblebd.setRange(0, 10)
            self.vgrainnuisiblebd.setSpecialValueText(' ')
            self.vgrainnuisiblebd.resize(100, 20)
            self.vgrainnuisiblebd.move(400, 335)
            self.vgrainnuisiblebd.setFont(self.txtpsfont)
            self.vgrainnuisiblebd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            ################################Grains chauffés (%)############################
            self.vtotalprembd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=True)
            self.vtotalprembd.setRange(0, 10)
            self.vtotalprembd.setDecimals(3)
            self.vtotalprembd.setSpecialValueText(' ')
            self.vtotalprembd.resize(100, 20)
            self.vtotalprembd.move(400, 356)
            self.vtotalprembd.setFont(self.txtpsfont)
            self.vtotalprembd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            ########################################Grains sans valeur (%)#######################################
            self.vgrainscassébd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=False)
            self.vgrainscassébd.setSpecialValueText(' ')
            self.vgrainscassébd.setRange(0, 30)
            self.vgrainscassébd.resize(100, 20)
            self.vgrainscassébd.move(400, 377)
            self.vgrainscassébd.setFont(self.txtpsfont)
            self.vgrainscassébd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #######################################Grains cariés##########################################
            self.vgrainsboutésbd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=False)
            self.vgrainsboutésbd.setSpecialValueText(' ')
            self.vgrainsboutésbd.setRange(0, 10)
            self.vgrainsboutésbd.resize(100, 20)
            self.vgrainsboutésbd.move(400, 398)
            self.vgrainsboutésbd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #######################################Total(%) 1er#####################################################
            self.vgrainsrouxbd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=False)
            self.vgrainsrouxbd.setSpecialValueText(' ')
            self.vgrainsrouxbd.setRange(0, 10)
            self.vgrainsrouxbd.resize(100, 20)
            self.vgrainsrouxbd.move(400, 419)
            self.vgrainsrouxbd.setFont(self.txtpsfont)
            self.vgrainsrouxbd.setStyleSheet(
                "background-color:#ffffff;color:#000000;border: 2px solid bleu;border-radius: 4px;padding: 0px")
            ##############################################Grains cassés (%) #########################################################
            self.vgrainfortementmouchtébd = QtWidgets.QDoubleSpinBox(self.tabbldur)
            self.vgrainfortementmouchtébd.move(400, 440)
            self.vgrainfortementmouchtébd.resize(100, 20)
            self.vgrainfortementmouchtébd.setRange(0, 10)
            self.vgrainfortementmouchtébd.setSpecialValueText(" ")
            self.vgrainfortementmouchtébd.setFont(self.txtpsfont)
            self.vgrainfortementmouchtébd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #########################################################Gains échaudés (%)#####################################################
            self.vgrainepunaisésbd = QtWidgets.QDoubleSpinBox(self.tabbldur)
            self.vgrainepunaisésbd.setSpecialValueText(" ")
            self.vgrainepunaisésbd.setRange(0, 10)
            self.vgrainepunaisésbd.move(400, 461)
            self.vgrainepunaisésbd.resize(100, 20)
            self.vgrainepunaisésbd.setFont(self.txtpsfont)
            self.vgrainepunaisésbd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")
            #####################################################Grains maigres (%)########################################################
            self.vgrainpiquebd = QtWidgets.QDoubleSpinBox(self.tabbldur)
            self.vgrainpiquebd.setRange(0, 10)
            self.vgrainpiquebd.setSpecialValueText(" ")
            self.vgrainpiquebd.move(400, 482)
            self.vgrainpiquebd.setFont(self.txtpsfont)
            self.vgrainpiquebd.resize(100, 20)
            self.vgrainpiquebd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")
            ##########################################################Grains germés (%)###################################################
            self.vtotaldembd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=True)
            self.vtotaldembd.move(400, 503)
            self.vtotaldembd.resize(100, 20)
            self.vtotaldembd.setSpecialValueText('  ')
            self.vtotaldembd.setFont(self.txtpsfont)
            self.vtotaldembd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")
            ##########################################################Grain punaisés (%)#########################################################
            self.vgrainmitadinésbd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=False)
            self.vgrainmitadinésbd.move(400, 524)
            self.vgrainmitadinésbd.setRange(0, 100)
            self.vgrainmitadinésbd.resize(100, 20)
            self.vgrainmitadinésbd.setSpecialValueText('   ')
            self.vgrainmitadinésbd.setFont(self.txtpsfont)
            self.vgrainmitadinésbd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")
            #######################################################################Grains piqués (%)##########################################
            self.vbletendredansbledurbd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=False)
            self.vbletendredansbledurbd.move(400, 545)
            self.vbletendredansbledurbd.resize(100, 20)
            self.vbletendredansbledurbd.setSpecialValueText('  ')
            self.vbletendredansbledurbd.setFont(self.txtpsfont)
            self.vbletendredansbledurbd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")
            #######################################################################Grains boutés « faible » (%)#######################################

            ####################################################################Grains boutés  « forte » (%)######################################
            self.vtotalmitadinésbd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=True)
            self.vtotalmitadinésbd.move(400, 566)
            self.vtotalmitadinésbd.resize(100, 20)
            self.vtotalmitadinésbd.setSpecialValueText('  ')
            self.vtotalmitadinésbd.setFont(self.txtpsfont)
            self.vtotalmitadinésbd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")

            #######################################################réfaction##############################################
            #################label valeure##############
            self.rvaleurbd = QtWidgets.QLabel("Réfaction (DA)", self.tabbldur)
            self.rvaleurbd.move(501, 205)
            self.rvaleurbd.resize(100, 20)
            self.rvaleurbd.setFont(self.font)
            ######################Limites(sans bon ni réf)################
            self.rpsbd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=True)

            self.rpsbd.setSpecialValueText(' ')
            self.rpsbd.resize(100, 20)
            self.rpsbd.move(501, 230)
            self.rpsbd.setFont(self.txtpsfont)
            self.rpsbd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")
            ###############################humidite#############
            self.rhumiditebd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=True)
            self.rhumiditebd.setRange(8, 14)
            self.rhumiditebd.resize(100, 20)
            self.rhumiditebd.setSpecialValueText(' ')
            self.rhumiditebd.move(501, 251)
            self.rhumiditebd.setFont(self.txtpsfont)
            self.rhumiditebd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #######################ergot#########################
            self.rergotbd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=True)
            self.rergotbd.setRange(0, 10)
            self.rergotbd.setSpecialValueText(' ')
            self.rergotbd.resize(100, 20)
            self.rergotbd.move(501, 272)
            self.rergotbd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #########################Graines nuisibles (%)##########
            self.rmatiére20mmbd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=True)
            self.rmatiére20mmbd.setRange(0, 10)
            self.rmatiére20mmbd.setSpecialValueText(' ')
            self.rmatiére20mmbd.resize(100, 20)
            self.rmatiére20mmbd.move(501, 293)
            self.rmatiére20mmbd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #############################Débris végétaux (%)########
            self.rdébrisvébd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=True)
            self.rdébrisvébd.setRange(0, 10)
            self.rdébrisvébd.setSpecialValueText(' ')
            self.rdébrisvébd.resize(100, 20)
            self.rdébrisvébd.move(501, 314)
            self.rdébrisvébd.setFont(self.txtpsfont)
            self.rdébrisvébd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #########################Matière inerte (%)################
            self.rgrainnuisiblebd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=True)
            self.rgrainnuisiblebd.setRange(0, 10)
            self.rgrainnuisiblebd.setSpecialValueText(' ')
            self.rgrainnuisiblebd.resize(100, 20)
            self.rgrainnuisiblebd.move(501, 335)
            self.rgrainnuisiblebd.setFont(self.txtpsfont)
            self.rgrainnuisiblebd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            ################################Grains chauffés (%)############################
            self.rtotalprembd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=True)
            self.rtotalprembd.setRange(0, 10)
            self.rtotalprembd.setDecimals(3)
            self.rtotalprembd.setSpecialValueText(' ')
            self.rtotalprembd.resize(100, 20)
            self.rtotalprembd.move(501, 356)
            self.rtotalprembd.setFont(self.txtpsfont)
            self.rtotalprembd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            ########################################Grains sans valeur (%)#######################################
            self.rgrainscassébd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=True)
            self.rgrainscassébd.setSpecialValueText(' ')
            self.rgrainscassébd.setDecimals(3)
            self.rgrainscassébd.resize(100, 20)
            self.rgrainscassébd.move(501, 377)
            self.rgrainscassébd.setFont(self.txtpsfont)
            self.rgrainscassébd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #######################################Grains cariés##########################################
            self.rgrainsboutésbd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=True)
            self.rgrainsboutésbd.setSpecialValueText(' ')
            self.rgrainsboutésbd.setRange(0, 10)
            self.rgrainsboutésbd.resize(100, 20)
            self.rgrainsboutésbd.move(501, 398)
            self.rgrainsboutésbd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #######################################Total(%) 1er#####################################################
            self.rgrainsrouxbd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=True)
            self.rgrainsrouxbd.setSpecialValueText(' ')
            self.rgrainsrouxbd.setRange(0, 10)
            self.rgrainsrouxbd.resize(100, 20)
            self.rgrainsrouxbd.move(501, 419)
            self.rgrainsrouxbd.setFont(self.txtpsfont)
            self.rgrainsrouxbd.setStyleSheet(
                "background-color:#ffffff;color:#000000;border: 2px solid bleu;border-radius: 4px;padding: 0px")
            ##############################################Grains cassés (%) #########################################################
            self.rgrainfortementmouchtébd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=True)
            self.rgrainfortementmouchtébd.move(501, 440)
            self.rgrainfortementmouchtébd.resize(100, 20)
            self.rgrainfortementmouchtébd.setRange(0, 10)
            self.rgrainfortementmouchtébd.setSpecialValueText(" ")
            self.rgrainfortementmouchtébd.setFont(self.txtpsfont)
            self.rgrainfortementmouchtébd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #########################################################Gains échaudés (%)#####################################################
            self.rgrainepunaisésbd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=True)
            self.rgrainepunaisésbd.setSpecialValueText(" ")
            self.rgrainepunaisésbd.setRange(0, 10)
            self.rgrainepunaisésbd.move(501, 461)
            self.rgrainepunaisésbd.resize(100, 20)
            self.rgrainepunaisésbd.setFont(self.txtpsfont)
            self.rgrainepunaisésbd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")
            #####################################################Grains maigres (%)########################################################
            self.rgrainpiquebd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=True)
            self.rgrainpiquebd.setRange(0, 10)
            self.rgrainpiquebd.setSpecialValueText(" ")
            self.rgrainpiquebd.move(501, 482)
            self.rgrainpiquebd.setFont(self.txtpsfont)
            self.rgrainpiquebd.resize(100, 20)
            self.rgrainpiquebd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")
            ##########################################################Grains germés (%)###################################################
            self.rtotaldembd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=True)
            self.rtotaldembd.move(501, 503)
            self.rtotaldembd.resize(100, 20)
            self.rtotaldembd.setSpecialValueText('  ')
            self.rtotaldembd.setFont(self.txtpsfont)
            self.rtotaldembd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")
            ##########################################################Grain punaisés (%)#########################################################
            self.rgrainmitadinésbd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=True)
            self.rgrainmitadinésbd.move(501, 524)
            self.rgrainmitadinésbd.resize(100, 20)
            self.rgrainmitadinésbd.setSpecialValueText('   ')
            self.rgrainmitadinésbd.setFont(self.txtpsfont)
            self.rgrainmitadinésbd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")
            #######################################################################Grains piqués (%)##########################################
            self.rbletendredansbledurbd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=True)
            self.rbletendredansbledurbd.move(501, 545)
            self.rbletendredansbledurbd.resize(100, 20)
            self.rbletendredansbledurbd.setSpecialValueText('  ')
            self.rbletendredansbledurbd.setFont(self.txtpsfont)
            self.rbletendredansbledurbd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")
            #######################################################################Grains boutés « faible » (%)#######################################

            ####################################################################Grains boutés  « forte » (%)######################################
            self.rtotalmitadinésbd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=True)
            self.rtotalmitadinésbd.move(501, 566)
            self.rtotalmitadinésbd.resize(100, 20)
            self.rtotalmitadinésbd.setSpecialValueText('  ')
            self.rtotalmitadinésbd.setFont(self.txtpsfont)
            self.rtotalmitadinésbd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")

            self.rtotalbiniEtrefaction = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=True)
            self.rtotalbiniEtrefaction.move(501, 587)
            self.rtotalbiniEtrefaction.resize(100, 30)
            self.rtotalbiniEtrefaction.setSpecialValueText('  ')
            self.rtotalbiniEtrefaction.setFont(self.txtpsfont)
            self.rtotalbiniEtrefaction.setStyleSheet(
                "background-color:#e22630;border: 2px solid bleu;border-radius: 4px;padding: 2px")

            ###################################################bonifécation#####################################################
            #################label valeure##############
            self.bvaleurbd = QtWidgets.QLabel("Bonification(DA)", self.tabbldur)
            self.bvaleurbd.move(602, 205)
            self.bvaleurbd.resize(103, 20)
            self.bvaleurbd.setFont(self.font)
            ######################Limites(sans bon ni réf)################
            self.bpsbd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=True)
            self.bpsbd.setSpecialValueText(' ')
            self.bpsbd.resize(100, 20)
            self.bpsbd.move(602, 230)
            self.bpsbd.setFont(self.txtpsfont)
            self.bpsbd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")
            ###############################humidite#############
            self.bhumiditebd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=True)
            self.bhumiditebd.setRange(8, 14)
            self.bhumiditebd.resize(100, 20)
            self.bhumiditebd.setSpecialValueText(' ')
            self.bhumiditebd.move(602, 251)
            self.bhumiditebd.setFont(self.txtpsfont)
            self.bhumiditebd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #######################ergot#########################
            self.bergotbd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=True)
            self.bergotbd.setRange(0, 10)
            self.bergotbd.setSpecialValueText(' ')
            self.bergotbd.resize(100, 20)
            self.bergotbd.move(602, 272)
            self.bergotbd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #########################Graines nuisibles (%)##########
            self.bmatiére20mmbd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=True)
            self.bmatiére20mmbd.setRange(0, 10)
            self.bmatiére20mmbd.setSpecialValueText(' ')
            self.bmatiére20mmbd.resize(100, 20)
            self.bmatiére20mmbd.move(602, 293)
            self.bmatiére20mmbd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #############################Débris végétaux (%)########
            self.bdébrisvébd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=True)
            self.bdébrisvébd.setSpecialValueText(' ')
            self.bdébrisvébd.resize(100, 20)
            self.bdébrisvébd.move(602, 314)
            self.bdébrisvébd.setFont(self.txtpsfont)
            self.bdébrisvébd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #########################Matière inerte (%)################
            self.bgrainnuisiblebd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=True)
            self.bgrainnuisiblebd.setSpecialValueText(' ')
            self.bgrainnuisiblebd.resize(100, 20)
            self.bgrainnuisiblebd.move(602, 335)
            self.bgrainnuisiblebd.setFont(self.txtpsfont)
            self.bgrainnuisiblebd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            ################################Grains chauffés (%)############################
            self.btotalprembd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=True)
            self.btotalprembd.setSpecialValueText(' ')
            self.btotalprembd.resize(100, 20)
            self.btotalprembd.move(602, 356)
            self.btotalprembd.setFont(self.txtpsfont)
            self.btotalprembd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            ########################################Grains sans valeur (%)#######################################
            self.bgrainscassébd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=True)
            self.bgrainscassébd.setSpecialValueText(' ')
            self.bgrainscassébd.resize(100, 20)
            self.bgrainscassébd.move(602, 377)
            self.bgrainscassébd.setFont(self.txtpsfont)
            self.bgrainscassébd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #######################################Grains cariés##########################################
            self.bgrainsboutésbd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=True)
            self.bgrainsboutésbd.setSpecialValueText(' ')
            self.bgrainsboutésbd.resize(100, 20)
            self.bgrainsboutésbd.move(602, 398)
            self.bgrainsboutésbd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #######################################Total(%) 1er#####################################################
            self.bgrainsrouxbd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=True)
            self.bgrainsrouxbd.setSpecialValueText(' ')
            self.bgrainsrouxbd.resize(100, 20)
            self.bgrainsrouxbd.move(602, 419)
            self.bgrainsrouxbd.setFont(self.txtpsfont)
            self.bgrainsrouxbd.setStyleSheet(
                "background-color:#ffffff;color:#000000;border: 2px solid bleu;border-radius: 4px;padding: 0px")
            ##############################################Grains cassés (%) #########################################################
            self.bgrainfortementmouchtébd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=True)
            self.bgrainfortementmouchtébd.move(602, 440)
            self.bgrainfortementmouchtébd.resize(100, 20)
            self.bgrainfortementmouchtébd.setSpecialValueText(" ")
            self.bgrainfortementmouchtébd.setFont(self.txtpsfont)
            self.bgrainfortementmouchtébd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #########################################################Gains échaudés (%)#####################################################
            self.bgrainepunaisésbd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=True)
            self.bgrainepunaisésbd.setSpecialValueText(" ")
            self.bgrainepunaisésbd.move(602, 461)
            self.bgrainepunaisésbd.resize(100, 20)
            self.bgrainepunaisésbd.setFont(self.txtpsfont)
            self.bgrainepunaisésbd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")
            #####################################################Grains maigres (%)########################################################
            self.bgrainpiquebd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=True)
            self.bgrainpiquebd.setSpecialValueText(" ")
            self.bgrainpiquebd.move(602, 482)
            self.bgrainpiquebd.setFont(self.txtpsfont)
            self.bgrainpiquebd.resize(100, 20)
            self.bgrainpiquebd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")
            ##########################################################Grains germés (%)###################################################
            self.btotaldembd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=True)
            self.btotaldembd.move(602, 503)
            self.btotaldembd.resize(100, 20)
            self.btotaldembd.setSpecialValueText('  ')
            self.btotaldembd.setFont(self.txtpsfont)
            self.btotaldembd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")
            ##########################################################Grain punaisés (%)#########################################################
            self.bgrainmitadinésbd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=True)
            self.bgrainmitadinésbd.move(602, 524)
            self.bgrainmitadinésbd.resize(100, 20)
            self.bgrainmitadinésbd.setSpecialValueText('   ')
            self.bgrainmitadinésbd.setFont(self.txtpsfont)
            self.bgrainmitadinésbd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")
            #######################################################################Grains piqués (%)##########################################
            self.bbletendredansbledurbd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=True)
            self.bbletendredansbledurbd.move(602, 545)
            self.bbletendredansbledurbd.resize(100, 20)
            self.bbletendredansbledurbd.setSpecialValueText('  ')
            self.bbletendredansbledurbd.setFont(self.txtpsfont)
            self.bbletendredansbledurbd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")
            #######################################################################Grains boutés « faible » (%)#######################################

            ####################################################################Grains boutés  « forte » (%)######################################
            self.btotalmitadinésbd = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=True)
            self.btotalmitadinésbd.move(602, 566)
            self.btotalmitadinésbd.resize(100, 20)
            self.btotalmitadinésbd.setSpecialValueText('  ')
            self.btotalmitadinésbd.setFont(self.txtpsfont)
            self.btotalmitadinésbd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")

            self.btotalbiniEtrefaction = QtWidgets.QDoubleSpinBox(self.tabbldur, readOnly=True)
            self.btotalbiniEtrefaction.move(602, 587)
            self.btotalbiniEtrefaction.resize(100, 30)
            self.btotalbiniEtrefaction.setSpecialValueText('  ')
            self.btotalbiniEtrefaction.setFont(self.txtpsfont)
            self.btotalbiniEtrefaction.setStyleSheet(
                "background-color:#88ffaa;border: 2px solid bleu;border-radius: 4px;padding: 2px")

            self.obesrvationbdtxt = QtWidgets.QLabel("Observation", self.tabbldur)
            self.obesrvationbdtxt.setGeometry(QtCore.QRect(710, 205, 100, 20))

            self.obesrvationbd = QtWidgets.QTextEdit(self.tabbldur)
            self.obesrvationbd.setGeometry(QtCore.QRect(704, 230, 100, 386))
            self.obesrvationbd.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")

            self.btnsavebd = QtWidgets.QPushButton("ENREGISTRER", self.tabbldur, clicked=lambda: self.add_datta_bd())
            self.btnsavebd.move(820, 230)
            self.btnsavebd.resize(180, 85)

            self.btnimprimejour = QtWidgets.QPushButton("IMPRIMER \nLE JOURNALIER", self.tabbldur,
                                                        clicked=lambda: self.print_day_bd())
            self.btnimprimejour.move(820, 430)
            self.btnimprimejour.resize(180, 85)

            self.btnprintbd = QtWidgets.QPushButton("IMPRIMER", self.tabbldur, clicked=lambda: self.bulletin_bd())
            self.btnprintbd.move(820, 330)
            self.btnprintbd.resize(180, 85)

            self.btnefaceor = QtWidgets.QPushButton("EFACER", self.tabbldur, clicked=lambda: self.clear_bd())
            self.btnefaceor.move(820, 530)
            self.btnefaceor.resize(180, 85)

            self.btnprintbulletindocx = QtWidgets.QPushButton("SELECTIONNE \nBULLETIN", self.tabbldur,
                                                              clicked=lambda: self.printBulltinProducteur())
            self.btnprintbulletindocx.move(1020, 230)
            self.btnprintbulletindocx.resize(180, 85)

            self.btnimprimeticket = QtWidgets.QPushButton("IMPRIMER \nLES TICKET", self.tabbldur,
                                                          clicked=lambda: self.eticket_bledur())
            self.btnimprimeticket.move(1020, 330)
            self.btnimprimeticket.resize(180, 85)

            self.timercalculbd = QTimer()
            self.timercalculbd.timeout.connect(self.calcul_bd)
            self.timercalculbd.setInterval(1000)
            self.timercalculbd.start()

            self.timernumbulltinBd = QTimer()
            self.timernumbulltinBd.timeout.connect(self.number_bulletin_bd)
            self.timernumbulltinBd.setInterval(1000)
            self.timernumbulltinBd.start()

            #####################################################################################
            ######################BLE TENDRE##################################
            ###############################################
            self.tabbltendre = QtWidgets.QWidget()
            self.tabbltendre.setObjectName("tabbltendre")
            self.tabbltendre.setStyleSheet("""QToolTip
    {
        border: 1px solid #76797C;
        background-color:  #fff8b0;
        color: white;
        padding: 5px;
        opacity: 200;
    }

    QWidget
    {
        color: #000000;
        background-color: #D8F9DB;
        selection-background-color:#3daee9;
        selection-color: #eff0f1;
        background-clip: border;
        border-image: none;
        border: 0px transparent black;
        outline: 0;
    }

    QWidget:item:hover
    {
        background-color: #3daee9;
        color: #eff0f1;
    }

    QWidget:item:selected
    {
        background-color: #3daee9;
    }



    QWidget:disabled
    {
        color: #454545;
        background-color: #31363b;
    }

    QAbstractItemView
    {
        alternate-background-color: #31363b;
        color: #eff0f1;
        border: 1px solid 3A3939;
        border-radius: 2px;
    }

    QWidget:focus, QMenuBar:focus
    {
        border: 1px solid #3daee9;
    }

    QTabWidget:focus, QCheckBox:focus, QRadioButton:focus, QSlider:focus
    {
        border: none;
    }

    QLineEdit
    {
        background-color: #ffffff;
        padding: 1px;
        border-style: solid;
        border: 1px solid #76797C;
        border-radius: 2px;
        color: #000000;
        font-size:12px;
        font-weight:bold;
    }
    QDoubleSpinBox
    {
        background-color: #ffffff;
        padding: 1px;
        border-style: solid;
        border: 1px solid #76797C;
        border-radius: 2px;
        color:#000000;
        font-size:12px;
        font-weight:bold;

    }
    QDoubleSpinBox::drop-down
    {
        subcontrol-origin: padding;
        subcontrol-position: top right;
        width: 1px;

        border-left-width: 0px;
        border-left-color: #232629;
        border-left-style: solid;
        border-top-right-radius: 1px;
        border-bottom-right-radius: 1px;
    }



    QGroupBox {
        border:1px solid #76797C;
        border-radius: 2px;
        margin-top: 20px;
    }

    QGroupBox::title {
        subcontrol-origin: margin;
        subcontrol-position: top center;
        padding-left: 10px;
        padding-right: 10px;
        padding-top: 10px;
    }

    QAbstractScrollArea
    {
        border-radius: 2px;
        border: 1px solid #76797C;
        background-color: transparent;
    }

    QScrollBar:horizontal
    {
        height: 15px;
        margin: 3px 15px 3px 15px;
        border: 1px transparent #2A2929;
        border-radius: 4px;
        background-color: #2A2929;
    }

    QScrollBar::handle:horizontal
    {
        background-color: #605F5F;
        min-width: 5px;
        border-radius: 4px;
    }

    QScrollBar::add-line:horizontal
    {
        margin: 0px 3px 0px 3px;
        border-image: url(:/qss_icons/Dark_rc/right_arrow_disabled.png);
        width: 10px;
        height: 10px;
        subcontrol-position: right;
        subcontrol-origin: margin;
    }

    QScrollBar::sub-line:horizontal
    {
        margin: 0px 3px 0px 3px;
        border-image: url(:/qss_icons/Dark_rc/left_arrow_disabled.png);
        height: 10px;
        width: 10px;
        subcontrol-position: left;
        subcontrol-origin: margin;
    }

    QScrollBar::add-line:horizontal:hover,QScrollBar::add-line:horizontal:on
    {
        border-image: url(:/qss_icons/Dark_rc/right_arrow.png);
        height: 10px;
        width: 10px;
        subcontrol-position: right;
        subcontrol-origin: margin;
    }


    QScrollBar::sub-line:horizontal:hover, QScrollBar::sub-line:horizontal:on
    {
        border-image: url(:/qss_icons/Dark_rc/left_arrow.png);
        height: 10px;
        width: 10px;
        subcontrol-position: left;
        subcontrol-origin: margin;
    }

    QScrollBar::up-arrow:horizontal, QScrollBar::down-arrow:horizontal
    {
        background: none;
    }


    QScrollBar::add-page:horizontal, QScrollBar::sub-page:horizontal
    {
        background: none;
    }

    QScrollBar:vertical
    {
        background-color: #2A2929;
        width: 15px;
        margin: 15px 3px 15px 3px;
        border: 1px transparent #2A2929;
        border-radius: 4px;
    }

    QScrollBar::handle:vertical
    {
        background-color: #605F5F;
        min-height: 5px;
        border-radius: 4px;
    }

    QScrollBar::sub-line:vertical
    {
        margin: 3px 0px 3px 0px;
        border-image: url(:/qss_icons/Dark_rc/up_arrow_disabled.png);
        height: 10px;
        width: 10px;
        subcontrol-position: top;
        subcontrol-origin: margin;
    }

    QScrollBar::add-line:vertical
    {
        margin: 3px 0px 3px 0px;
        border-image: url(:/qss_icons/Dark_rc/down_arrow_disabled.png);
        height: 10px;
        width: 10px;
        subcontrol-position: bottom;
        subcontrol-origin: margin;
    }

    QScrollBar::sub-line:vertical:hover,QScrollBar::sub-line:vertical:on
    {

        border-image: url(:/qss_icons/Dark_rc/up_arrow.png);
        height: 10px;
        width: 10px;
        subcontrol-position: top;
        subcontrol-origin: margin;
    }


    QScrollBar::add-line:vertical:hover, QScrollBar::add-line:vertical:on
    {
        border-image: url(:/qss_icons/Dark_rc/down_arrow.png);
        height: 10px;
        width: 10px;
        subcontrol-position: bottom;
        subcontrol-origin: margin;
    }

    QScrollBar::up-arrow:vertical, QScrollBar::down-arrow:vertical
    {
        background: none;
    }


    QScrollBar::add-page:vertical, QScrollBar::sub-page:vertical
    {
        background: none;
    }

    QTextEdit
    {
        background-color: #fffff1;
        color: #000000;
        border: 1px solid #76797C;
        font-size:12px;
        font-weight:bold;
    }

    QPlainTextEdit
    {
        background-color: #232629;;
        color: #eff0f1;
        border-radius: 2px;
        border: 1px solid #76797C;
    }

    QHeaderView::section
    {
        background-color: #76797C;
        color: #eff0f1;
        padding: 5px;
        border: 1px solid #76797C;
    }

    QSizeGrip {
        image: url(:/qss_icons/Dark_rc/sizegrip.png);
        width: 12px;
        height: 12px;
    }


    QMainWindow::separator
    {
        background-color: #31363b;
        color: white;
        padding-left: 4px;
        spacing: 2px;
        border: 1px dashed #76797C;
    }

    QMainWindow::separator:hover
    {

        background-color: #787876;
        color: white;
        padding-left: 4px;
        border: 1px solid #76797C;
        spacing: 2px;
    }


    QMenu::separator
    {
        height: 1px;
        background-color: #76797C;
        color: white;
        padding-left: 4px;
        margin-left: 10px;
        margin-right: 5px;
    }


    QFrame
    {
        border-radius: 2px;
        border: 1px solid #76797C;
    }

    QFrame[frameShape="0"]
    {
        border-radius: 2px;
        border: 1px transparent #76797C;
    }

    QStackedWidget
    {
        border: 1px transparent black;
    }


    QPushButton
    {
        color: #000000;
        background-color:#84dbc8;
        border-width: 1px;
        border-color: #1e1e1e;
        border-style: solid;
        border-radius: 6;
        padding: 3px;
        font-size: 12px;
        padding-left: 5px;
        padding-right: 5px;
        min-width: 40px;

    }

    QPushButton:disabled
    {
        background-color: #31363b;
        border-width: 1px;
        border-color: #454545;
        border-style: solid;
        padding-top: 5px;
        padding-bottom: 5px;
        padding-left: 10px;
        padding-right: 10px;
        border-radius: 2px;
        color: #454545;
    }

    QPushButton:focus {
        background-color: #3daee9;
        color: white;
    }

    QPushButton:pressed
    {
        background-color: #3daee9;
        padding-top: -15px;
        padding-bottom: -17px;
    }

   

    QPushButton:checked{
        background-color: #76797C;
        border-color: #6A6969;
    }

    QComboBox:hover,QDoubleSpinBox:Hover,QPushButton:hover,QAbstractSpinBox:hover,QLineEdit:hover,QTextEdit:hover,QPlainTextEdit:hover,QAbstractView:hover,QTreeView:hover
    {
        border: 1px solid #ff8c00;
        color: #000000;
    }

   QComboBox {
    background-color: #ffffff;
    border: 1px solid #76797C;
    color:#000000;
    border-radius: 0.25em;
    padding: 0.10em 0.10em;
    font-size:12px;
    font-weight:bold;
    cursor: pointer;
}

QComboBox::drop-down {
    subcontrol-origin: padding;
    subcontrol-position: top right;
    width: 1.3em;
    border-left: 0px solid #777;
    border-radius: 0.25em;
}

QComboBox::drop-down::icon {
    image: url('E:/pythonProject_moullin-application.3.5/images/down-arroww.png');
}

        QComboBox:on
        {
            padding-top: 0px;
            padding-left: 0px;        
            selection-background-color: #e4f0f1;
        }
        QComboBox QAbstractItemView
        {
            background-color: #ffffff;
            border-radius: 2px;
            border: 1px solid #76797C;
            color:#000000;
            selection-background-color: #000000;
        }
                                                         
            QComboBox:hover,QDoubleSpinBox:Hover,QPushButton:hover,QAbstractSpinBox:hover,QLineEdit:hover,QTextEdit:hover,QPlainTextEdit:hover,QAbstractView:hover,QTreeView:hover
            {
                border: 1px solid #ff8c00;
                color: #eff0f1;
            }


    QLabel
    {
        border: 2px solid black;
        font-size:13px;
        font-weight:bold;
    }

    QTabWidget{
        border: 0px transparent black;
    }

    QTabWidget::pane {
        border: 1px solid #76797C;
        padding: 5px;
        margin: 0px;
    }

    QTabBar
    {
        qproperty-drawBase: 0;
        left: 5px; /* move to the right by 5px */
        border-radius: 3px;
    }

    QTabBar:focus
    {
        border: 0px transparent black;
    }

    QTabBar::close-button  {
        image: url(:/qss_icons/Dark_rc/close.png);
        background: transparent;
    }

    QTabBar::close-button:hover
    {
        image: url(:/qss_icons/Dark_rc/close-hover.png);
        background: transparent;
    }

    QTabBar::close-button:pressed {
        image: url(:/qss_icons/Dark_rc/close-pressed.png);
        background: transparent;
    }

    /* TOP TABS */
    QTabBar::tab:top {
        color: #eff0f1;
        border: 1px solid #76797C;
        border-bottom: 1px transparent black;
        background-color: #31363b;
        padding: 5px;
        min-width: 50px;
        border-top-left-radius: 2px;
        border-top-right-radius: 2px;
    }

    QTabBar::tab:top:!selected
    {
        color: #eff0f1;
        background-color: #54575B;
        border: 1px solid #76797C;
        border-bottom: 1px transparent black;
        border-top-left-radius: 2px;
        border-top-right-radius: 2px;    
    }

    QTabBar::tab:top:!selected:hover {
        background-color: #3daee9;
    }

    /* BOTTOM TABS */
    QTabBar::tab:bottom {
        color: #eff0f1;
        border: 1px solid #76797C;
        border-top: 1px transparent black;
        background-color: #31363b;
        padding: 5px;
        border-bottom-left-radius: 2px;
        border-bottom-right-radius: 2px;
        min-width: 50px;
    }

    QTabBar::tab:bottom:!selected
    {
        color: #eff0f1;
        background-color: #54575B;
        border: 1px solid #76797C;
        border-top: 1px transparent black;
        border-bottom-left-radius: 2px;
        border-bottom-right-radius: 2px;
    }

    QTabBar::tab:bottom:!selected:hover {
        background-color: #3daee9;
    }

    /* LEFT TABS */
    QTabBar::tab:left {
        color: #eff0f1;
        border: 1px solid #76797C;
        border-left: 1px transparent black;
        background-color: #31363b;
        padding: 5px;
        border-top-right-radius: 2px;
        border-bottom-right-radius: 2px;
        min-height: 50px;
    }

    QTabBar::tab:left:!selected
    {
        color: #eff0f1;
        background-color: #54575B;
        border: 1px solid #76797C;
        border-left: 1px transparent black;
        border-top-right-radius: 2px;
        border-bottom-right-radius: 2px;
    }

    QTabBar::tab:left:!selected:hover {
        background-color: #3daee9;
    }


    /* RIGHT TABS */
    QTabBar::tab:right {
        color: #eff0f1;
        border: 1px solid #76797C;
        border-right: 1px transparent black;
        background-color: #31363b;
        padding: 5px;
        border-top-left-radius: 2px;
        border-bottom-left-radius: 2px;
        min-height: 50px;
    }

    QTabBar::tab:right:!selected
    {
        color: #eff0f1;
        background-color: #54575B;
        border: 1px solid #76797C;
        border-right: 1px transparent black;
        border-top-left-radius: 2px;
        border-bottom-left-radius: 2px;
    }

    QTabBar::tab:right:!selected:hover {
        background-color: #3daee9;
    }

    QTabBar QToolButton::right-arrow:enabled {
         image: url(:/qss_icons/Dark_rc/right_arrow.png);
     }

     QTabBar QToolButton::left-arrow:enabled {
         image: url(:/qss_icons/Dark_rc/left_arrow.png);
     }

    QTabBar QToolButton::right-arrow:disabled {
         image: url(:/qss_icons/Dark_rc/right_arrow_disabled.png);
     }

     QTabBar QToolButton::left-arrow:disabled {
         image: url(:/qss_icons/Dark_rc/left_arrow_disabled.png);
     }


    QDockWidget {
        background: #31363b;
        border: 1px solid #403F3F;
        titlebar-close-icon: url(:/qss_icons/Dark_rc/close.png);
        titlebar-normal-icon: url(:/qss_icons/Dark_rc/undock.png);
    }

    QDockWidget::close-button, QDockWidget::float-button {
        border: 1px solid transparent;
        border-radius: 2px;
        background: transparent;
    }

    QDockWidget::close-button:hover, QDockWidget::float-button:hover {
        background: rgba(255, 255, 255, 10);
    }

    QDockWidget::close-button:pressed, QDockWidget::float-button:pressed {
        padding: 1px -1px -1px 1px;
        background: rgba(255, 255, 255, 10);
    }


    QSlider::groove:horizontal {
        border: 1px solid #565a5e;
        height: 4px;
        background: #565a5e;
        margin: 0px;
        border-radius: 2px;
    }

    QSlider::handle:horizontal {
        background: #232629;
        border: 1px solid #565a5e;
        width: 16px;
        height: 16px;
        margin: -8px 0;
        border-radius: 9px;
    }

    QSlider::groove:vertical {
        border: 1px solid #565a5e;
        width: 4px;
        background: #565a5e;
        margin: 0px;
        border-radius: 3px;
    }

    QSlider::handle:vertical {
        background: #232629;
        border: 1px solid #565a5e;
        width: 16px;
        height: 16px;
        margin: 0 -8px;
        border-radius: 9px;
    }

    QToolButton {
        background-color: transparent;
        border: 1px transparent #76797C;
        border-radius: 2px;
        margin: 3px;
        padding: 5px;
    }

    QToolButton[popupMode="1"] { /* only for MenuButtonPopup */
     padding-right: 20px; /* make way for the popup button */
     border: 1px #76797C;
     border-radius: 5px;
    }

    QToolButton[popupMode="2"] { /* only for InstantPopup */
     padding-right: 10px; /* make way for the popup button */
     border: 1px #76797C;
    }


    QToolButton:hover, QToolButton::menu-button:hover {
        background-color: transparent;
        border: 1px solid #3daee9;
        padding: 5px;
    }

    QToolButton:checked, QToolButton:pressed,
            QToolButton::menu-button:pressed {
        background-color: #3daee9;
        border: 1px solid #3daee9;
        padding: 5px;
    }

    /* the subcontrol below is used only in the InstantPopup or DelayedPopup mode */
    QToolButton::menu-indicator {
        background-color:ff8c00;
        top: -7px; left: -2px; /* shift it a bit */
    }

    /* the subcontrols below are used only in the MenuButtonPopup mode */
    QToolButton::menu-button {
        border: 1px transparent #76797C;
        border-top-right-radius: 6px;
        border-bottom-right-radius: 6px;
        /* 16px width + 4px for border = 20px allocated above */
        width: 16px;
        outline: none;
    }

    QToolButton::menu-arrow {
       background-color:ff8c00;
    }

    QToolButton::menu-arrow:open {
        border: 1px solid #76797C;
    }

    QPushButton::menu-indicator  {
        subcontrol-origin: padding;
        subcontrol-position: bottom right;
        left: 8px;
    }

    QTableView
    {
        border: 1px solid #76797C;
        gridline-color: #31363b;
        background-color: #232629;
    }


    QTableView, QHeaderView
    {
        border-radius: 0px;
    }

    QTableView::item:pressed, QListView::item:pressed, QTreeView::item:pressed  {
        background: #3daee9;
        color: #eff0f1;
    }

    QTableView::item:selected:active, QTreeView::item:selected:active, QListView::item:selected:active  {
        background: #3daee9;
        color: #eff0f1;
    }


    QHeaderView
    {
        background-color: #31363b;
        border: 1px transparent;
        border-radius: 0px;
        margin: 0px;
        padding: 0px;

    }

    QHeaderView::section  {
        background-color: #31363b;
        color: #eff0f1;
        padding: 5px;
        border: 1px solid #76797C;
        border-radius: 0px;
        text-align: center;
    }

    QHeaderView::section::vertical::first, QHeaderView::section::vertical::only-one
    {
        border-top: 1px solid #76797C;
    }

    QHeaderView::section::vertical
    {
        border-top: transparent;
    }

    QHeaderView::section::horizontal::first, QHeaderView::section::horizontal::only-one
    {
        border-left: 1px solid #76797C;
    }

    QHeaderView::section::horizontal
    {
        border-left: transparent;
    }


    QHeaderView::section:checked
     {
        color: white;
        background-color: #334e5e;
     }

     /* style the sort indicator */
    QHeaderView::down-arrow {
        image: url(:/qss_icons/Dark_rc/down_arrow.png);
    }

    QHeaderView::up-arrow {
        image: url(:/qss_icons/Dark_rc/up_arrow.png);
    }


    QTableCornerButton::section {
        background-color: #31363b;
        border: 1px transparent #76797C;
        border-radius: 0px;
    }

    QToolBox  {
        padding: 5px;
        border: 1px transparent black;
    }

    QToolBox::tab {
        color: #eff0f1;
        background-color: #31363b;
        border: 1px solid #76797C;
        border-bottom: 1px transparent #31363b;
        border-top-left-radius: 5px;
        border-top-right-radius: 5px;
    }

    QToolBox::tab:selected { /* italicize selected tabs */
        font: italic;
        background-color: #31363b;
        border-color: #3daee9;
     }

    QStatusBar::item {
        border: 0px transparent dark;
     }


    QFrame[height="3"], QFrame[width="3"] {
        background-color: #76797C;
    }




    QDateTimeEdit
    {
        background-color:#ffffff;
        border-style: solid;
        border: 1px solid #76797C;
        border-radius: 2px;
        padding: 1px;
        min-width: 75px;
    }

    QDateEdit:on
    {
        padding-top: 2px;
        padding-left: 2px;
        selection-background-color: #4a4a4a;
    }

    QDateEdit QAbstractItemView
    {
        background-color: #ff8c00;
        border-radius: 2px;
        border: 1px solid #3375A3;
        selection-background-color:ff8c00;
    }

    QDateEdit::drop-down
    {
        subcontrol-origin: padding;
        subcontrol-position: top right;
        width: 15px;
        border-left-width: 0px;
        border-left-color: darkgray;
        border-left-style: solid;
        border-top-right-radius: 3px;
        border-bottom-right-radius: 3px;
    }""")
            self.tabagreage.addTab(self.tabbltendre, "")
            self.cclstxt = QtWidgets.QLabel("<h2>COOPERATIVE DES CEREALES ET LEGUMES SECS DE RELIZANE<h2/>",
                                            self.tabbltendre)
            self.cclstxt.resize(600, 40)
            self.cclstxt.move(400, 2)

            self.font = QtGui.QFont()
            self.font.setBold(True)
            self.font.setPointSize(10)

            self.txtpsfont = QtGui.QFont()
            self.txtpsfont.setBold(True)
            self.txtpsfont.setPointSize(9)

            self.rtotaldemfont = QtGui.QFont("color:black")
            self.rtotaldemfont.setBold(True)
            self.rtotaldemfont.setPointSize(12)

            self.bulletin_agréagetxt = QtWidgets.QLabel("<h2>Bulletin D'Agréage<h2/>", self.tabbltendre)
            self.bulletin_agréagetxt.move(600, 30)
            self.bulletin_agréagetxt.resize(200, 40)

            self.éspécetxt = QtWidgets.QLabel("<h2>Espéce:Blé Tendre<h2/>", self.tabbltendre)
            self.éspécetxt.move(630, 60)

            self.n_bultinbttxt = QtWidgets.QLabel("N° Bulletin:", self.tabbltendre)
            self.n_bultinbttxt.move(20, 60)
            self.n_bultinbt = QtWidgets.QLineEdit(self.tabbltendre, readOnly=True)
            self.n_bultinbt.move(150, 60)
            self.n_bultinbt.resize(150, 25)

            self.nome_du_producteurbttxt = QtWidgets.QLabel("Nom du Producteur:", self.tabbltendre)
            self.nome_du_producteurbttxt.move(20, 100)
            self.nome_du_producteurbt = QtWidgets.QLineEdit(self.tabbltendre)
            self.nome_du_producteurbt.move(150, 98)
            self.nome_du_producteurbt.resize(150, 25)

            self.péretxtbt = QtWidgets.QLabel("Pére:", self.tabbltendre)
            self.péretxtbt.setGeometry(QtCore.QRect(315, 98, 100, 20))
            self.pérebt = QtWidgets.QLineEdit(self.tabbltendre)
            self.pérebt.setGeometry(QtCore.QRect(410, 98, 150, 23))

            self.n_cartebttxt = QtWidgets.QLabel("N.C d'identité:", self.tabbltendre)
            self.n_cartebttxt.setGeometry(QtCore.QRect(315, 138, 100, 20))
            self.n_cartebt = QtWidgets.QLineEdit(self.tabbltendre)
            self.n_cartebt.setGeometry(QtCore.QRect(410, 138, 150, 23))
            self.n_cartebt.setInputMask('99999999')

            self.imatriculebttxt = QtWidgets.QLabel("Adresse:", self.tabbltendre)
            self.imatriculebttxt.move(20, 140)
            self.adressebt = QtWidgets.QLineEdit(self.tabbltendre)
            self.adressebt.setInputMask("99999-999-99")
            self.adressebt.move(150, 138)
            self.adressebt.resize(150, 25)

            self.pointdecollectbttxt = QtWidgets.QLabel("Point de collecte", self.tabbltendre)
            self.pointdecollectbttxt.move(20, 180)
            self.pointdecollectbt = QtWidgets.QComboBox(self.tabbltendre)
            self.pointdecollectbt.setStyleSheet("background-color:#ffffff;color:#000000")
            self.pointdecollectbt.move(150, 178)
            self.pointdecollectbt.resize(150, 25)
            self.pointdecollectbt.addItem('')
            self.pointdecollectbt.addItem('Dock central')
            self.pointdecollectbt.addItem('Magasin Zemmoura')
            self.pointdecollectbt.addItem('Magasin Kef-lazreg')
            self.pointdecollectbt.addItem('Magasin Messra')
            self.pointdecollectbt.addItem('Marche gros belacel')
            self.pointdecollectbt.addItem('Station Mendes')
            self.pointdecollectbt.addItem('Nouvelles S.Mendes')

            self.dattereceptionbttxt = QtWidgets.QLabel("Relizane le :", self.tabbltendre)
            self.dattereceptionbttxt.setGeometry(QtCore.QRect(880, 100, 150, 23))
            self.dattereceptiontbt = QtWidgets.QDateTimeEdit(self.tabbltendre)
            self.dattereceptiontbt.setGeometry(QtCore.QRect(1000, 100, 150, 23))
            self.dattereceptiontbt.setDisplayFormat("dd-MM-yyyy hh:mm")
            self.dattereceptiontbt.setDate(self.datedaytime)

            self.agréeeurbttxt = QtWidgets.QLabel("Nom de l’Agréeur:", self.tabbltendre)
            self.agréeeurbttxt.setGeometry(QtCore.QRect(880, 135, 150, 23))
            self.agréeeurcombobt = QtWidgets.QComboBox(self.tabbltendre, editable=True)
            self.agréeeurcombobt.setStyleSheet("background-color:#ffffff;color:#000000")
            self.agréeeurcombobt.addItem("")
            self.agréeeurcombobt.addItem("FELOUAH OMAR")
            self.agréeeurcombobt.addItem("BEKHEDDA AEK")
            self.agréeeurcombobt.addItem("BENAISSA YOUCEF")
            self.agréeeurcombobt.addItem("REZZAG SOFIANE ")
            self.agréeeurcombobt.addItem("BELBACHA M.NADIR")
            self.agréeeurcombobt.move(1000, 135)
            self.agréeeurcombobt.resize(150, 23)

            self.quantitebttxt = QtWidgets.QLabel("Quantité", self.tabbltendre)
            self.quantitebttxt.setGeometry(QtCore.QRect(880, 180, 150, 23))

            self.quantitebt = QtWidgets.QDoubleSpinBox(self.tabbltendre)
            self.quantitebt.setRange(1, 10000)
            self.quantitebt.setSuffix('  QX')
            self.quantitebt.setSpecialValueText(" ")
            self.quantitebt.setGeometry(QtCore.QRect(1000, 178, 150, 23))

            self.paramétrebt = QtWidgets.QLabel("Paramètre", self.tabbltendre)
            self.paramétrebt.move(30, 205)
            self.paramétrebt.resize(80, 20)
            self.paramétrebt.setFont(self.font)
            self.txtpsfontbt = QtGui.QFont()
            self.txtpsfontbt.setBold(True)
            self.txtpsfontbt.setPointSize(9)
            ################Limites(sans bon ni réf)###############
            self.valeurbt = QtWidgets.QLabel("", self.tabbltendre)
            self.valeurbt.move(170, 205)
            self.valeurbt.resize(145, 30)
            self.valeurbt.setFont(self.font)
            ######################Limites(sans bon ni réf)################
            self.psbt = QtWidgets.QLineEdit("Poids spécifique (kg/hl)", self.tabbltendre, readOnly=True)
            self.psbt.resize(369, 20)
            self.psbt.move(30, 230)
            self.psbt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            self.psbt.setFont(self.txtpsfont)
            self.psbt.setStyleSheet(" border: 2px solid bleu ;border-radius: 4px;padding: 0px")
            ###############################humidite#############
            self.humiditebt = QtWidgets.QLineEdit("Teneur en eau(%)", self.tabbltendre, readOnly=True)
            self.humiditebt.resize(369, 20)
            self.humiditebt.move(30, 251)
            self.humiditebt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            self.humiditebt.setFont(self.txtpsfont)

            #######################ergot#########################
            self.ergotbt = QtWidgets.QLineEdit("Ergo(%  ", self.tabbltendre, readOnly=True)
            self.ergotbt.resize(369, 20)
            self.ergotbt.move(30, 272)
            self.ergotbt.setStyleSheet("background-color: #232629")
            self.ergotbt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            self.ergotbt.setFont(self.txtpsfont)

            #########################Graines nuisibles (%)##########
            self.matiére20mm = QtWidgets.QLineEdit("Matiéres qui passent à travers \nle tamis 20 mm x 2.1 mm ",
                                                   self.tabbltendre, readOnly=True)
            self.matiére20mm.resize(369, 20)
            self.matiére20mm.move(30, 293)
            self.matiére20mm.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            self.matiére20mm.setFont(self.txtpsfont)
            #############################Débris végétaux (%)########
            self.débrisvébt = QtWidgets.QLineEdit("Les débris végétaux et \nles éléments minéreaux(%)",
                                                  self.tabbltendre,
                                                  readOnly=True)
            self.débrisvébt.resize(369, 20)
            self.débrisvébt.move(30, 314)
            self.débrisvébt.setFont(self.txtpsfont)
            self.débrisvébt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #########################Matière inerte (%)################
            self.grainnuisiblebt = QtWidgets.QLineEdit("Graines nuisibles(%)", self.tabbltendre, readOnly=True)
            self.grainnuisiblebt.resize(369, 20)
            self.grainnuisiblebt.move(30, 335)
            self.grainnuisiblebt.setFont(self.txtpsfont)
            self.grainnuisiblebt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            ################################Grains chauffés (%)############################
            self.totalprembt = QtWidgets.QLineEdit("Total(%)", self.tabbltendre, readOnly=True)
            self.totalprembt.resize(369, 20)
            self.totalprembt.move(30, 356)
            self.totalprembt.setFont(self.txtpsfont)
            self.totalprembt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            ########################################Grains sans valeur (%)#######################################
            self.grainscassébt = QtWidgets.QLineEdit("Grains cassés(%)", self.tabbltendre, readOnly=True)
            self.grainscassébt.resize(369, 20)
            self.grainscassébt.move(30, 377)
            self.grainscassébt.setFont(self.txtpsfont)
            self.grainscassébt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #######################################Grains cariés##########################################
            self.grainpunaiséebt = QtWidgets.QLineEdit("Grains punaisés", self.tabbltendre, readOnly=True)
            self.grainpunaiséebt.resize(369, 20)
            self.grainpunaiséebt.move(30, 398)
            self.grainpunaiséebt.setFont(self.txtpsfont)
            self.grainpunaiséebt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")

            #######################################Total(%) 1er#####################################################
            self.grainsfortementboutésbt = QtWidgets.QLineEdit("Grains fortement boutés(%)", self.tabbltendre,
                                                               readOnly=True)
            self.grainsfortementboutésbt.resize(369, 20)
            self.grainsfortementboutésbt.move(30, 419)
            self.grainsfortementboutésbt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            self.grainsfortementboutésbt.setFont(self.txtpsfont)
            ##############################################Grains cassés (%) #########################################################
            self.grainsfaiblementboutésbt = QtWidgets.QLineEdit("Grains faiblement boutés(%)", self.tabbltendre,
                                                                readOnly=True)
            self.grainsfaiblementboutésbt.move(30, 440)
            self.grainsfaiblementboutésbt.resize(369, 20)
            self.grainsfaiblementboutésbt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            self.grainsfaiblementboutésbt.setFont(self.txtpsfont)
            #########################################################Gains échaudés (%)#####################################################
            self.grainefortementmouchetesbt = QtWidgets.QLineEdit("Gains fortement mouchetés(%)", self.tabbltendre,
                                                                  readOnly=True)
            self.grainefortementmouchetesbt.move(30, 461)
            self.grainefortementmouchetesbt.resize(369, 20)
            self.grainefortementmouchetesbt.setFont(self.txtpsfont)
            self.grainefortementmouchetesbt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #####################################################Grains maigres (%)########################################################
            self.grainetrangebt = QtWidgets.QLineEdit("Grains étrangers utilisable pour le bétail(%) ",
                                                      self.tabbltendre, readOnly=True)
            self.grainetrangebt.move(30, 482)
            self.grainetrangebt.resize(369, 20)
            self.grainetrangebt.setFont(self.txtpsfont)
            self.grainetrangebt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            ##########################################################Grains germés (%)###################################################
            self.totaldembt = QtWidgets.QLineEdit("Total (%)", self.tabbltendre, readOnly=True)
            self.totaldembt.move(30, 503)
            self.totaldembt.resize(369, 20)
            self.totaldembt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            self.totaldembt.setFont(self.txtpsfont)
            ##########################################################Grain punaisés (%)#########################################################
            self.totalbiniEtrefaction = QtWidgets.QLineEdit("\t\tTotal de bonification et réfaction", self.tabbltendre,
                                                            readOnly=True)
            self.totalbiniEtrefaction.move(30, 524)
            self.totalbiniEtrefaction.resize(470, 30)
            self.totalbiniEtrefaction.setFont(self.txtpsfont)
            self.totalbiniEtrefaction.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")

            #################label valeure##############
            self.valeurbt = QtWidgets.QLabel("valeur", self.tabbltendre)
            self.valeurbt.move(400, 205)
            self.valeurbt.resize(100, 20)
            self.valeurbt.setFont(self.font)
            ######################Limites(sans bon ni réf)################
            self.vpsbt = QtWidgets.QDoubleSpinBox(self.tabbltendre)
            self.vpsbt.setDecimals(3)

            self.vpsbt.setSpecialValueText(' ')
            self.vpsbt.resize(100, 20)
            self.vpsbt.move(400, 230)
            self.vpsbt.setFont(self.txtpsfont)
            self.vpsbt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")
            ###############################humidite#############
            self.vhumiditebt = QtWidgets.QDoubleSpinBox(self.tabbltendre)
            self.vhumiditebt.setRange(3, 17)
            self.vhumiditebt.resize(100, 20)
            self.vhumiditebt.setSpecialValueText(' ')
            self.vhumiditebt.move(400, 251)
            self.vhumiditebt.setFont(self.txtpsfont)
            self.vhumiditebt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #######################ergot#########################
            self.vergotbt = QtWidgets.QDoubleSpinBox(self.tabbltendre, readOnly=False)
            self.vergotbt.setRange(0, 10)
            self.vergotbt.setSpecialValueText(' ')
            self.vergotbt.resize(100, 20)
            self.vergotbt.move(400, 272)
            self.vergotbt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #########################Graines nuisibles (%)##########
            self.vmatiére20mmbt = QtWidgets.QDoubleSpinBox(self.tabbltendre, readOnly=False)
            self.vmatiére20mmbt.setRange(0, 10)
            self.vmatiére20mmbt.setSpecialValueText(' ')
            self.vmatiére20mmbt.resize(100, 20)
            self.vmatiére20mmbt.move(400, 293)
            self.vmatiére20mmbt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #############################Débris végétaux (%)########
            self.vdébrisvébt = QtWidgets.QDoubleSpinBox(self.tabbltendre, readOnly=False)
            self.vdébrisvébt.setRange(0, 10)
            self.vdébrisvébt.setSpecialValueText(' ')
            self.vdébrisvébt.resize(100, 20)
            self.vdébrisvébt.move(400, 314)
            self.vdébrisvébt.setFont(self.txtpsfont)
            self.vdébrisvébt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #########################Matière inerte (%)################
            self.vgrainnuisiblebt = QtWidgets.QDoubleSpinBox(self.tabbltendre, readOnly=False)
            self.vgrainnuisiblebt.setRange(0, 10)
            self.vgrainnuisiblebt.setSpecialValueText(' ')
            self.vgrainnuisiblebt.resize(100, 20)
            self.vgrainnuisiblebt.move(400, 335)
            self.vgrainnuisiblebt.setFont(self.txtpsfont)
            self.vgrainnuisiblebt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            ################################Grains chauffés (%)############################
            self.vtotalprembt = QtWidgets.QDoubleSpinBox(self.tabbltendre, readOnly=True)
            self.vtotalprembt.setRange(0, 10)
            self.vtotalprembt.setDecimals(3)
            self.vtotalprembt.setSpecialValueText(' ')
            self.vtotalprembt.resize(100, 20)
            self.vtotalprembt.move(400, 356)
            self.vtotalprembt.setFont(self.txtpsfont)
            self.vtotalprembt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            ########################################Grains sans valeur (%)#######################################
            self.vgrainscassébt = QtWidgets.QDoubleSpinBox(self.tabbltendre, readOnly=False)
            self.vgrainscassébt.setSpecialValueText(' ')
            self.vgrainscassébt.setRange(0, 30)
            self.vgrainscassébt.resize(100, 20)
            self.vgrainscassébt.move(400, 377)
            self.vgrainscassébt.setFont(self.txtpsfont)
            self.vgrainscassébt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #######################################Grains cariés##########################################
            self.vgrainpunaiséebt = QtWidgets.QDoubleSpinBox(self.tabbltendre, readOnly=False)
            self.vgrainpunaiséebt.setSpecialValueText(' ')
            self.vgrainpunaiséebt.setRange(0, 10)
            self.vgrainpunaiséebt.resize(100, 20)
            self.vgrainpunaiséebt.move(400, 398)
            self.vgrainpunaiséebt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #######################################Total(%) 1er#####################################################
            self.vgrainsfortementboutésbt = QtWidgets.QDoubleSpinBox(self.tabbltendre, readOnly=False)
            self.vgrainsfortementboutésbt.setSpecialValueText(' ')
            self.vgrainsfortementboutésbt.setRange(0, 10)
            self.vgrainsfortementboutésbt.resize(100, 20)
            self.vgrainsfortementboutésbt.move(400, 419)
            self.vgrainsfortementboutésbt.setFont(self.txtpsfont)
            self.vgrainsfortementboutésbt.setStyleSheet(
                "background-color:#ffffff;color:#000000;border: 2px solid bleu;border-radius: 4px;padding: 0px")
            ##############################################Grains cassés (%) #########################################################
            self.vgrainsfaiblementboutésbt = QtWidgets.QDoubleSpinBox(self.tabbltendre)
            self.vgrainsfaiblementboutésbt.move(400, 440)
            self.vgrainsfaiblementboutésbt.resize(100, 20)
            self.vgrainsfaiblementboutésbt.setRange(0, 10)
            self.vgrainsfaiblementboutésbt.setSpecialValueText(" ")
            self.vgrainsfaiblementboutésbt.setFont(self.txtpsfont)
            self.vgrainsfaiblementboutésbt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #########################################################Gains échaudés (%)#####################################################
            self.vgrainefortementmouchetesbt = QtWidgets.QDoubleSpinBox(self.tabbltendre)
            self.vgrainefortementmouchetesbt.setSpecialValueText(" ")
            self.vgrainefortementmouchetesbt.setRange(0, 10)
            self.vgrainefortementmouchetesbt.move(400, 461)
            self.vgrainefortementmouchetesbt.resize(100, 20)
            self.vgrainefortementmouchetesbt.setFont(self.txtpsfont)
            self.vgrainefortementmouchetesbt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")
            #####################################################Grains maigres (%)########################################################
            self.vgrainetrangebt = QtWidgets.QDoubleSpinBox(self.tabbltendre)
            self.vgrainetrangebt.setRange(0, 10)
            self.vgrainetrangebt.setSpecialValueText(" ")
            self.vgrainetrangebt.move(400, 482)
            self.vgrainetrangebt.setFont(self.txtpsfont)
            self.vgrainetrangebt.resize(100, 20)
            self.vgrainetrangebt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")
            ##########################################################Grains germés (%)###################################################
            self.vtotaldembt = QtWidgets.QDoubleSpinBox(self.tabbltendre, readOnly=True)
            self.vtotaldembt.move(400, 503)
            self.vtotaldembt.resize(100, 20)
            self.vtotaldembt.setSpecialValueText('  ')
            self.vtotaldembt.setFont(self.txtpsfont)
            self.vtotaldembt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")
            ###############################################réfaction##############################################
            #################label valeure##############
            self.rvaleurbt = QtWidgets.QLabel("Réfaction (DA)", self.tabbltendre)
            self.rvaleurbt.move(501, 205)
            self.rvaleurbt.resize(100, 20)
            self.rvaleurbt.setFont(self.font)
            ######################Limites(sans bon ni réf)################
            self.rpsbt = QtWidgets.QDoubleSpinBox(self.tabbltendre, readOnly=True)

            self.rpsbt.setSpecialValueText(' ')
            self.rpsbt.resize(100, 20)
            self.rpsbt.move(501, 230)
            self.rpsbt.setFont(self.txtpsfont)
            self.rpsbt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")
            ###############################humidite#############
            self.rhumiditebt = QtWidgets.QDoubleSpinBox(self.tabbltendre, readOnly=True)
            self.rhumiditebt.setRange(8, 14)
            self.rhumiditebt.resize(100, 20)
            self.rhumiditebt.setSpecialValueText(' ')
            self.rhumiditebt.move(501, 251)
            self.rhumiditebt.setFont(self.txtpsfont)
            self.rhumiditebt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #######################ergot#########################
            self.rergotbt = QtWidgets.QDoubleSpinBox(self.tabbltendre, readOnly=True)
            self.rergotbt.setRange(0, 10)
            self.rergotbt.setSpecialValueText(' ')
            self.rergotbt.resize(100, 20)
            self.rergotbt.move(501, 272)
            self.rergotbt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #########################Graines nuisibles (%)##########
            self.rmatiére20mmbt = QtWidgets.QDoubleSpinBox(self.tabbltendre, readOnly=True)
            self.rmatiére20mmbt.setRange(0, 10)
            self.rmatiére20mmbt.setSpecialValueText(' ')
            self.rmatiére20mmbt.resize(100, 20)
            self.rmatiére20mmbt.move(501, 293)
            self.rmatiére20mmbt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #############################Débris végétaux (%)########
            self.rdébrisvébt = QtWidgets.QDoubleSpinBox(self.tabbltendre, readOnly=True)
            self.rdébrisvébt.setRange(0, 10)
            self.rdébrisvébt.setSpecialValueText(' ')
            self.rdébrisvébt.resize(100, 20)
            self.rdébrisvébt.move(501, 314)
            self.rdébrisvébt.setFont(self.txtpsfont)
            self.rdébrisvébt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #########################Matière inerte (%)################
            self.rgrainnuisiblebt = QtWidgets.QDoubleSpinBox(self.tabbltendre, readOnly=True)
            self.rgrainnuisiblebt.setRange(0, 10)
            self.rgrainnuisiblebt.setSpecialValueText(' ')
            self.rgrainnuisiblebt.resize(100, 20)
            self.rgrainnuisiblebt.move(501, 335)
            self.rgrainnuisiblebt.setFont(self.txtpsfont)
            self.rgrainnuisiblebt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            ################################Grains chauffés (%)############################
            self.rtotalprembt = QtWidgets.QDoubleSpinBox(self.tabbltendre, readOnly=True)
            self.rtotalprembt.setRange(0, 10)
            self.rtotalprembt.setDecimals(2)
            self.rtotalprembt.setSpecialValueText(' ')
            self.rtotalprembt.resize(100, 20)
            self.rtotalprembt.move(501, 356)
            self.rtotalprembt.setFont(self.txtpsfont)
            self.rtotalprembt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            ########################################Grains sans valeur (%)#######################################
            self.rgrainscassébt = QtWidgets.QDoubleSpinBox(self.tabbltendre, readOnly=True)
            self.rgrainscassébt.setSpecialValueText(' ')
            self.rgrainscassébt.setDecimals(2)
            self.rgrainscassébt.resize(100, 20)
            self.rgrainscassébt.move(501, 377)
            self.rgrainscassébt.setFont(self.txtpsfont)
            self.rgrainscassébt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #######################################Grains cariés##########################################
            self.rgrainpunaiséebt = QtWidgets.QDoubleSpinBox(self.tabbltendre, readOnly=True)
            self.rgrainpunaiséebt.setSpecialValueText(' ')
            self.rgrainpunaiséebt.setRange(0, 10)
            self.rgrainpunaiséebt.resize(100, 20)
            self.rgrainpunaiséebt.move(501, 398)
            self.rgrainpunaiséebt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #######################################Total(%) 1er#####################################################
            self.rgrainsfortementboutésbt = QtWidgets.QDoubleSpinBox(self.tabbltendre, readOnly=True)
            self.rgrainsfortementboutésbt.setSpecialValueText(' ')
            self.rgrainsfortementboutésbt.setRange(0, 10)
            self.rgrainsfortementboutésbt.resize(100, 20)
            self.rgrainsfortementboutésbt.move(501, 419)
            self.rgrainsfortementboutésbt.setFont(self.txtpsfont)
            self.rgrainsfortementboutésbt.setStyleSheet(
                "background-color:#ffffff;color:#000000;border: 2px solid bleu;border-radius: 4px;padding: 0px")
            ##############################################Grains cassés (%) #########################################################
            self.rgrainsfaiblementboutésbt = QtWidgets.QDoubleSpinBox(self.tabbltendre, readOnly=True)
            self.rgrainsfaiblementboutésbt.move(501, 440)
            self.rgrainsfaiblementboutésbt.resize(100, 20)
            self.rgrainsfaiblementboutésbt.setRange(0, 10)
            self.rgrainsfaiblementboutésbt.setSpecialValueText(" ")
            self.rgrainsfaiblementboutésbt.setFont(self.txtpsfont)
            self.rgrainsfaiblementboutésbt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #########################################################Gains échaudés (%)#####################################################
            self.rgrainefortementmouchetesbt = QtWidgets.QDoubleSpinBox(self.tabbltendre, readOnly=True)
            self.rgrainefortementmouchetesbt.setSpecialValueText(" ")
            self.rgrainefortementmouchetesbt.setRange(0, 10)
            self.rgrainefortementmouchetesbt.move(501, 461)
            self.rgrainefortementmouchetesbt.resize(100, 20)
            self.rgrainefortementmouchetesbt.setFont(self.txtpsfont)
            self.rgrainefortementmouchetesbt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")
            #####################################################Grains maigres (%)########################################################
            self.rgrainetrangebt = QtWidgets.QDoubleSpinBox(self.tabbltendre, readOnly=True)
            self.rgrainetrangebt.setRange(0, 10)
            self.rgrainetrangebt.setSpecialValueText(" ")
            self.rgrainetrangebt.move(501, 482)
            self.rgrainetrangebt.setFont(self.txtpsfont)
            self.rgrainetrangebt.resize(100, 20)
            self.rgrainetrangebt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")
            ##########################################################Grains germés (%)###################################################
            self.rtotaldembt = QtWidgets.QDoubleSpinBox(self.tabbltendre, readOnly=True)
            self.rtotaldembt.move(501, 503)
            self.rtotaldembt.resize(100, 20)
            self.rtotaldembt.setSpecialValueText('  ')
            self.rtotaldembt.setFont(self.txtpsfont)
            self.rtotaldembt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")
            #################Grains piqués (%)##########################################
            self.rtotalbiniEtrefactionbt = QtWidgets.QDoubleSpinBox(self.tabbltendre, readOnly=True)
            self.rtotalbiniEtrefactionbt.move(501, 524)
            self.rtotalbiniEtrefactionbt.resize(100, 30)
            self.rtotalbiniEtrefactionbt.setSpecialValueText('  ')
            self.rtotalbiniEtrefactionbt.setFont(self.txtpsfont)
            self.rtotalbiniEtrefactionbt.setStyleSheet(
                "background-color:#e22630;border: 2px solid bleu;border-radius: 4px;padding: 2px")

            ###################################################bonifécation#####################################################
            #################label valeure##############
            self.bvaleurbt = QtWidgets.QLabel("Bonification(DA)", self.tabbltendre)
            self.bvaleurbt.move(602, 205)
            self.bvaleurbt.resize(103, 20)
            self.bvaleurbt.setFont(self.font)
            ######################Limites(sans bon ni réf)################
            self.bpsbt = QtWidgets.QDoubleSpinBox(self.tabbltendre, readOnly=True)
            self.bpsbt.setSpecialValueText(' ')
            self.bpsbt.resize(100, 20)
            self.bpsbt.move(602, 230)
            self.bpsbt.setFont(self.txtpsfont)
            self.bpsbt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")
            ###############################humidite#############
            self.bhumiditebt = QtWidgets.QDoubleSpinBox(self.tabbltendre, readOnly=True)
            self.bhumiditebt.setRange(8, 14)
            self.bhumiditebt.resize(100, 20)
            self.bhumiditebt.setSpecialValueText(' ')
            self.bhumiditebt.move(602, 251)
            self.bhumiditebt.setFont(self.txtpsfont)
            self.bhumiditebt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #######################ergot#########################
            self.bergotbt = QtWidgets.QDoubleSpinBox(self.tabbltendre, readOnly=True)
            self.bergotbt.setRange(0, 10)
            self.bergotbt.setSpecialValueText(' ')
            self.bergotbt.resize(100, 20)
            self.bergotbt.move(602, 272)
            self.bergotbt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #########################Graines nuisibles (%)##########
            self.bmatiére20mmbt = QtWidgets.QDoubleSpinBox(self.tabbltendre, readOnly=True)
            self.bmatiére20mmbt.setRange(0, 10)
            self.bmatiére20mmbt.setSpecialValueText(' ')
            self.bmatiére20mmbt.resize(100, 20)
            self.bmatiére20mmbt.move(602, 293)
            self.bmatiére20mmbt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #############################Débris végétaux (%)########
            self.btébrisvébt = QtWidgets.QDoubleSpinBox(self.tabbltendre, readOnly=True)
            self.btébrisvébt.setSpecialValueText(' ')
            self.btébrisvébt.resize(100, 20)
            self.btébrisvébt.move(602, 314)
            self.btébrisvébt.setFont(self.txtpsfont)
            self.btébrisvébt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #########################Matière inerte (%)################
            self.bgrainnuisiblebt = QtWidgets.QDoubleSpinBox(self.tabbltendre, readOnly=True)
            self.bgrainnuisiblebt.setSpecialValueText(' ')
            self.bgrainnuisiblebt.resize(100, 20)
            self.bgrainnuisiblebt.move(602, 335)
            self.bgrainnuisiblebt.setFont(self.txtpsfont)
            self.bgrainnuisiblebt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            ################################Grains chauffés (%)############################
            self.btotalprembt = QtWidgets.QDoubleSpinBox(self.tabbltendre, readOnly=True)
            self.btotalprembt.setSpecialValueText(' ')
            self.btotalprembt.resize(100, 20)
            self.btotalprembt.move(602, 356)
            self.btotalprembt.setFont(self.txtpsfont)
            self.btotalprembt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            ########################################Grains sans valeur (%)#######################################
            self.bgrainscassébt = QtWidgets.QDoubleSpinBox(self.tabbltendre, readOnly=True)
            self.bgrainscassébt.setSpecialValueText(' ')
            self.bgrainscassébt.resize(100, 20)
            self.bgrainscassébt.move(602, 377)
            self.bgrainscassébt.setFont(self.txtpsfont)
            self.bgrainscassébt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #######################################Grains cariés##########################################
            self.bgrainpunaiséebt = QtWidgets.QDoubleSpinBox(self.tabbltendre, readOnly=True)
            self.bgrainpunaiséebt.setSpecialValueText(' ')
            self.bgrainpunaiséebt.resize(100, 20)
            self.bgrainpunaiséebt.move(602, 398)
            self.bgrainpunaiséebt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #######################################Total(%) 1er#####################################################
            self.brgrainsfortementboutésbt = QtWidgets.QDoubleSpinBox(self.tabbltendre, readOnly=True)
            self.brgrainsfortementboutésbt.setSpecialValueText(' ')
            self.brgrainsfortementboutésbt.resize(100, 20)
            self.brgrainsfortementboutésbt.move(602, 419)
            self.brgrainsfortementboutésbt.setFont(self.txtpsfont)
            self.brgrainsfortementboutésbt.setStyleSheet(
                "background-color:#ffffff;color:#000000;border: 2px solid bleu;border-radius: 4px;padding: 0px")
            ##############################################Grains cassés (%) #########################################################
            self.bgrainsfaiblementboutésbt = QtWidgets.QDoubleSpinBox(self.tabbltendre, readOnly=True)
            self.bgrainsfaiblementboutésbt.move(602, 440)
            self.bgrainsfaiblementboutésbt.resize(100, 20)
            self.bgrainsfaiblementboutésbt.setSpecialValueText(" ")
            self.bgrainsfaiblementboutésbt.setFont(self.txtpsfont)
            self.bgrainsfaiblementboutésbt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #########################################################Gains échaudés (%)#####################################################
            self.bgrainefortementmouchetesbt = QtWidgets.QDoubleSpinBox(self.tabbltendre, readOnly=True)
            self.bgrainefortementmouchetesbt.setSpecialValueText(" ")
            self.bgrainefortementmouchetesbt.move(602, 461)
            self.bgrainefortementmouchetesbt.resize(100, 20)
            self.bgrainefortementmouchetesbt.setFont(self.txtpsfont)
            self.bgrainefortementmouchetesbt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")
            #####################################################Grains maigres (%)########################################################
            self.bgrainetrangebt = QtWidgets.QDoubleSpinBox(self.tabbltendre, readOnly=True)
            self.bgrainetrangebt.setSpecialValueText(" ")
            self.bgrainetrangebt.move(602, 482)
            self.bgrainetrangebt.setFont(self.txtpsfont)
            self.bgrainetrangebt.resize(100, 20)
            self.bgrainetrangebt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")
            ##########################################################Grains germés (%)###################################################
            self.btotaldembt = QtWidgets.QDoubleSpinBox(self.tabbltendre, readOnly=True)
            self.btotaldembt.move(602, 503)
            self.btotaldembt.resize(100, 20)
            self.btotaldembt.setSpecialValueText('  ')
            self.btotaldembt.setFont(self.txtpsfont)
            self.btotaldembt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")
            ##########################################################Grain punaisés (%)#########################################################

            ####################################################################Grains boutés  « forte » (%)######################################

            self.btotalbiniEtrefactionbt = QtWidgets.QDoubleSpinBox(self.tabbltendre, readOnly=True)
            self.btotalbiniEtrefactionbt.move(602, 524)
            self.btotalbiniEtrefactionbt.resize(100, 30)
            self.btotalbiniEtrefactionbt.setSpecialValueText('  ')
            self.btotalbiniEtrefactionbt.setFont(self.txtpsfont)
            self.btotalbiniEtrefactionbt.setStyleSheet(
                "background-color:#88ffaa;border: 2px solid bleu;border-radius: 4px;padding: 2px")

            self.obesrvationbttxt = QtWidgets.QLabel("Observation", self.tabbltendre)
            self.obesrvationbttxt.setGeometry(QtCore.QRect(710, 205, 100, 20))

            self.obesrvationbt = QtWidgets.QTextEdit(self.tabbltendre)
            self.obesrvationbt.setGeometry(QtCore.QRect(704, 230, 100, 325))
            self.obesrvationbt.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")

            self.btnsavebt = QtWidgets.QPushButton("ENREGISTRER", self.tabbltendre, clicked=lambda: self.add_datta_bt())
            self.btnsavebt.move(820, 230)
            self.btnsavebt.resize(180, 70)

            self.btnimprimejour = QtWidgets.QPushButton("IMPRIMER \nLE JOURNALIER", self.tabbltendre,
                                                        clicked=lambda: self.today_bt())
            self.btnimprimejour.move(820, 400)
            self.btnimprimejour.resize(180, 70)

            self.btnprintbt = QtWidgets.QPushButton("IMPRIMER", self.tabbltendre, clicked=lambda: self.bulletin_bt())
            self.btnprintbt.move(820, 315)
            self.btnprintbt.resize(180, 70)

            self.btnefacebt = QtWidgets.QPushButton("EFACER", self.tabbltendre, clicked=lambda: self.clear_bt())
            self.btnefacebt.move(820, 485)
            self.btnefacebt.resize(180, 70)

            self.btnprintbulletindocx = QtWidgets.QPushButton("SELECTIONNE \nBULLETIN", self.tabbltendre,
                                                              clicked=lambda: self.printBulltinProducteurbt())
            self.btnprintbulletindocx.move(1020, 230)
            self.btnprintbulletindocx.resize(180, 70)

            self.btnimprimeticket = QtWidgets.QPushButton("IMPRIMER \nLES TICKET", self.tabbltendre,
                                                          clicked=lambda: self.eticket_bletendre())
            self.btnimprimeticket.move(1020, 315)
            self.btnimprimeticket.resize(180, 70)

            self.timercalculbt = QTimer()
            self.timercalculbt.timeout.connect(self.calcul_bt)
            self.timercalculbt.setInterval(1000)
            self.timercalculbt.start()

            self.timernumbulltinbt = QTimer()
            self.timernumbulltinbt.timeout.connect(self.number_bulletin_bt)
            self.timernumbulltinbt.setInterval(1000)
            self.timernumbulltinbt.start()

            ###################################################################################
            #################################ORGE###############################
            ######################################
            self.taborge = QtWidgets.QWidget()
            self.taborge.setStyleSheet("""QToolTip
    {
        border: 1px solid #76797C;
        background-color:  #fff8b0;
        color: white;
        padding: 5px;
        opacity: 200;
    }

    QWidget
    {
        color: #000000;
        background-color:  #feffbd;
        selection-background-color:#3daee9;
        selection-color: #eff0f1;
        background-clip: border;
        border-image: none;
        border: 0px transparent black;
        outline: 0;
    }

    QWidget:item:hover
    {
        background-color: #3daee9;
        color: #eff0f1;
    }

    QWidget:item:selected
    {
        background-color: #3daee9;
    }



    QWidget:disabled
    {
        color: #454545;
        background-color: #31363b;
    }

    QAbstractItemView
    {
        alternate-background-color: #31363b;
        color: #eff0f1;
        border: 1px solid 3A3939;
        border-radius: 2px;
    }

    QWidget:focus, QMenuBar:focus
    {
        border: 1px solid #3daee9;
    }

    QTabWidget:focus, QCheckBox:focus, QRadioButton:focus, QSlider:focus
    {
        border: none;
    }

    QLineEdit
    {
        background-color: #ffffff;
        padding: 1px;
        border-style: solid;
        border: 1px solid #76797C;
        border-radius: 2px;
        color: #000000;
        font-size:12px;
        font-weight:bold;
    }
    QDoubleSpinBox
    {
        background-color: #ffffff;
        padding: 0px;
        border-style: solid;
        border: 1px solid #76797C;
        border-radius: 2px;
        color:#000000;
        font-size:12px;
        font-weight:bold;

    }
    QDoubleSpinBox::drop-down
    {
        subcontrol-origin: padding;
        subcontrol-position: top right;
        width: 1px;

        border-left-width: 0px;
        border-left-color: #232629;
        border-left-style: solid;
        border-top-right-radius: 1px;
        border-bottom-right-radius: 1px;
    }



    QGroupBox {
        border:1px solid #76797C;
        border-radius: 2px;
        margin-top: 20px;
    }

    QGroupBox::title {
        subcontrol-origin: margin;
        subcontrol-position: top center;
        padding-left: 10px;
        padding-right: 10px;
        padding-top: 10px;
    }

    QAbstractScrollArea
    {
        border-radius: 2px;
        border: 1px solid #76797C;
        background-color: transparent;
    }

    QScrollBar:horizontal
    {
        height: 15px;
        margin: 3px 15px 3px 15px;
        border: 1px transparent #2A2929;
        border-radius: 4px;
        background-color: #2A2929;
    }

    QScrollBar::handle:horizontal
    {
        background-color: #605F5F;
        min-width: 5px;
        border-radius: 4px;
    }

    QScrollBar::add-line:horizontal
    {
        margin: 0px 3px 0px 3px;
        border-image: url(:/qss_icons/Dark_rc/right_arrow_disabled.png);
        width: 10px;
        height: 10px;
        subcontrol-position: right;
        subcontrol-origin: margin;
    }

    QScrollBar::sub-line:horizontal
    {
        margin: 0px 3px 0px 3px;
        border-image: url(:/qss_icons/Dark_rc/left_arrow_disabled.png);
        height: 10px;
        width: 10px;
        subcontrol-position: left;
        subcontrol-origin: margin;
    }

    QScrollBar::add-line:horizontal:hover,QScrollBar::add-line:horizontal:on
    {
        border-image: url(:/qss_icons/Dark_rc/right_arrow.png);
        height: 10px;
        width: 10px;
        subcontrol-position: right;
        subcontrol-origin: margin;
    }


    QScrollBar::sub-line:horizontal:hover, QScrollBar::sub-line:horizontal:on
    {
        border-image: url(:/qss_icons/Dark_rc/left_arrow.png);
        height: 10px;
        width: 10px;
        subcontrol-position: left;
        subcontrol-origin: margin;
    }

    QScrollBar::up-arrow:horizontal, QScrollBar::down-arrow:horizontal
    {
        background: none;
    }


    QScrollBar::add-page:horizontal, QScrollBar::sub-page:horizontal
    {
        background: none;
    }

    QScrollBar:vertical
    {
        background-color: #2A2929;
        width: 15px;
        margin: 15px 3px 15px 3px;
        border: 1px transparent #2A2929;
        border-radius: 4px;
    }

    QScrollBar::handle:vertical
    {
        background-color: #605F5F;
        min-height: 5px;
        border-radius: 4px;
    }

    QScrollBar::sub-line:vertical
    {
        margin: 3px 0px 3px 0px;
        border-image: url(:/qss_icons/Dark_rc/up_arrow_disabled.png);
        height: 10px;
        width: 10px;
        subcontrol-position: top;
        subcontrol-origin: margin;
    }

    QScrollBar::add-line:vertical
    {
        margin: 3px 0px 3px 0px;
        border-image: url(:/qss_icons/Dark_rc/down_arrow_disabled.png);
        height: 10px;
        width: 10px;
        subcontrol-position: bottom;
        subcontrol-origin: margin;
    }

    QScrollBar::sub-line:vertical:hover,QScrollBar::sub-line:vertical:on
    {

        border-image: url(:/qss_icons/Dark_rc/up_arrow.png);
        height: 10px;
        width: 10px;
        subcontrol-position: top;
        subcontrol-origin: margin;
    }


    QScrollBar::add-line:vertical:hover, QScrollBar::add-line:vertical:on
    {
        border-image: url(:/qss_icons/Dark_rc/down_arrow.png);
        height: 10px;
        width: 10px;
        subcontrol-position: bottom;
        subcontrol-origin: margin;
    }

    QScrollBar::up-arrow:vertical, QScrollBar::down-arrow:vertical
    {
        background: none;
    }


    QScrollBar::add-page:vertical, QScrollBar::sub-page:vertical
    {
        background: none;
    }

    QTextEdit
    {
        background-color: #ffffff;
        color: #000000;
        border: 1px solid #76797C;
        font-size:12px;
        font-weight:bold;
    }

    QPlainTextEdit
    {
        background-color: #232629;;
        color: #eff0f1;
        border-radius: 2px;
        border: 1px solid #76797C;
    }

    QHeaderView::section
    {
        background-color: #76797C;
        color: #eff0f1;
        padding: 5px;
        border: 1px solid #76797C;
    }

    QSizeGrip {
        image: url(:/qss_icons/Dark_rc/sizegrip.png);
        width: 12px;
        height: 12px;
    }


    QMainWindow::separator
    {
        background-color: #31363b;
        color: white;
        padding-left: 4px;
        spacing: 2px;
        border: 1px dashed #76797C;
    }

    QMainWindow::separator:hover
    {

        background-color: #787876;
        color: white;
        padding-left: 4px;
        border: 1px solid #76797C;
        spacing: 2px;
    }


    QMenu::separator
    {
        height: 1px;
        background-color: #76797C;
        color: white;
        padding-left: 4px;
        margin-left: 10px;
        margin-right: 5px;
    }


    QFrame
    {
        border-radius: 2px;
        border: 1px solid #76797C;
    }

    QFrame[frameShape="0"]
    {
        border-radius: 2px;
        border: 1px transparent #76797C;
    }

    QStackedWidget
    {
        border: 1px transparent black;
    }


    QPushButton
    {
        color: #000000;
        background-color:#84dbc8;
        border-width: 1px;
        border-color: #1e1e1e;
        border-style: solid;
        border-radius: 6;
        padding: 3px;
        font-size: 12px;
        padding-left: 5px;
        padding-right: 5px;
        min-width: 40px;

    }

    QPushButton:disabled
    {
        background-color: #31363b;
        border-width: 1px;
        border-color: #454545;
        border-style: solid;
        padding-top: 5px;
        padding-bottom: 5px;
        padding-left: 10px;
        padding-right: 10px;
        border-radius: 2px;
        color: #454545;
    }

    QPushButton:focus {
        background-color: #3daee9;
        color: white;
    }

    QPushButton:pressed
    {
        background-color: #3daee9;
        padding-top: -15px;
        padding-bottom: -17px;
    }

    

    QPushButton:checked{
        background-color: #76797C;
        border-color: #6A6969;
    }

    QComboBox:hover,QDoubleSpinBox:Hover,QPushButton:hover,QAbstractSpinBox:hover,QLineEdit:hover,QTextEdit:hover,QPlainTextEdit:hover,QAbstractView:hover,QTreeView:hover
    {
        border: 1px solid #ff8c00;
        color: #000000;
    }

    QComboBox {
    background-color: #ffffff;
    border: 1px solid #76797C;
    color:#000000;
    border-radius: 0.25em;
    padding: 0.10em 0.10em;
    font-size:12px;
    font-weight:bold;
    cursor: pointer;
}

QComboBox::drop-down {
    subcontrol-origin: padding;
    subcontrol-position: top right;
    width: 1.3em;
    border-left: 0px solid #777;
    border-radius: 0.25em;
}

QComboBox::drop-down::icon {
    image: url('E:/pythonProject_moullin-application.3.5/images/down-arroww.png');
}

        QComboBox:on
        {
            padding-top: 0px;
            padding-left: 0px;        
            selection-background-color: #e4f0f1;
        }
        QComboBox QAbstractItemView
        {
            background-color: #ffffff;
            border-radius: 2px;
            border: 1px solid #76797C;
            color:#000000;
            selection-background-color: #000000;
        }


    QLabel
    {
        border: 2px solid black;
        font-size:13px;
        font-weight:bold;
    }

    QTabWidget{
        border: 0px transparent black;
    }

    QTabWidget::pane {
        border: 1px solid #76797C;
        padding: 5px;
        margin: 0px;
    }

    QTabBar
    {
        qproperty-drawBase: 0;
        left: 5px; /* move to the right by 5px */
        border-radius: 3px;
    }

    QTabBar:focus
    {
        border: 0px transparent black;
    }

    QTabBar::close-button  {
        image: url(:/qss_icons/Dark_rc/close.png);
        background: transparent;
    }

    QTabBar::close-button:hover
    {
        image: url(:/qss_icons/Dark_rc/close-hover.png);
        background: transparent;
    }

    QTabBar::close-button:pressed {
        image: url(:/qss_icons/Dark_rc/close-pressed.png);
        background: transparent;
    }

    /* TOP TABS */
    QTabBar::tab:top {
        color: #eff0f1;
        border: 1px solid #76797C;
        border-bottom: 1px transparent black;
        background-color: #31363b;
        padding: 5px;
        min-width: 50px;
        border-top-left-radius: 2px;
        border-top-right-radius: 2px;
    }

    QTabBar::tab:top:!selected
    {
        color: #eff0f1;
        background-color: #54575B;
        border: 1px solid #76797C;
        border-bottom: 1px transparent black;
        border-top-left-radius: 2px;
        border-top-right-radius: 2px;    
    }

    QTabBar::tab:top:!selected:hover {
        background-color: #3daee9;
    }

    /* BOTTOM TABS */
    QTabBar::tab:bottom {
        color: #eff0f1;
        border: 1px solid #76797C;
        border-top: 1px transparent black;
        background-color: #31363b;
        padding: 5px;
        border-bottom-left-radius: 2px;
        border-bottom-right-radius: 2px;
        min-width: 50px;
    }

    QTabBar::tab:bottom:!selected
    {
        color: #eff0f1;
        background-color: #54575B;
        border: 1px solid #76797C;
        border-top: 1px transparent black;
        border-bottom-left-radius: 2px;
        border-bottom-right-radius: 2px;
    }

    QTabBar::tab:bottom:!selected:hover {
        background-color: #3daee9;
    }

    /* LEFT TABS */
    QTabBar::tab:left {
        color: #eff0f1;
        border: 1px solid #76797C;
        border-left: 1px transparent black;
        background-color: #31363b;
        padding: 5px;
        border-top-right-radius: 2px;
        border-bottom-right-radius: 2px;
        min-height: 50px;
    }

    QTabBar::tab:left:!selected
    {
        color: #eff0f1;
        background-color: #54575B;
        border: 1px solid #76797C;
        border-left: 1px transparent black;
        border-top-right-radius: 2px;
        border-bottom-right-radius: 2px;
    }

    QTabBar::tab:left:!selected:hover {
        background-color: #3daee9;
    }


    /* RIGHT TABS */
    QTabBar::tab:right {
        color: #eff0f1;
        border: 1px solid #76797C;
        border-right: 1px transparent black;
        background-color: #31363b;
        padding: 5px;
        border-top-left-radius: 2px;
        border-bottom-left-radius: 2px;
        min-height: 50px;
    }

    QTabBar::tab:right:!selected
    {
        color: #eff0f1;
        background-color: #54575B;
        border: 1px solid #76797C;
        border-right: 1px transparent black;
        border-top-left-radius: 2px;
        border-bottom-left-radius: 2px;
    }

    QTabBar::tab:right:!selected:hover {
        background-color: #3daee9;
    }

    QTabBar QToolButton::right-arrow:enabled {
         image: url(:/qss_icons/Dark_rc/right_arrow.png);
     }

     QTabBar QToolButton::left-arrow:enabled {
         image: url(:/qss_icons/Dark_rc/left_arrow.png);
     }

    QTabBar QToolButton::right-arrow:disabled {
         image: url(:/qss_icons/Dark_rc/right_arrow_disabled.png);
     }

     QTabBar QToolButton::left-arrow:disabled {
         image: url(:/qss_icons/Dark_rc/left_arrow_disabled.png);
     }


    QDockWidget {
        background: #31363b;
        border: 1px solid #403F3F;
        titlebar-close-icon: url(:/qss_icons/Dark_rc/close.png);
        titlebar-normal-icon: url(:/qss_icons/Dark_rc/undock.png);
    }

    QDockWidget::close-button, QDockWidget::float-button {
        border: 1px solid transparent;
        border-radius: 2px;
        background: transparent;
    }

    QDockWidget::close-button:hover, QDockWidget::float-button:hover {
        background: rgba(255, 255, 255, 10);
    }

    QDockWidget::close-button:pressed, QDockWidget::float-button:pressed {
        padding: 1px -1px -1px 1px;
        background: rgba(255, 255, 255, 10);
    }


    QSlider::groove:horizontal {
        border: 1px solid #565a5e;
        height: 4px;
        background: #565a5e;
        margin: 0px;
        border-radius: 2px;
    }

    QSlider::handle:horizontal {
        background: #232629;
        border: 1px solid #565a5e;
        width: 16px;
        height: 16px;
        margin: -8px 0;
        border-radius: 9px;
    }

    QSlider::groove:vertical {
        border: 1px solid #565a5e;
        width: 4px;
        background: #565a5e;
        margin: 0px;
        border-radius: 3px;
    }

    QSlider::handle:vertical {
        background: #232629;
        border: 1px solid #565a5e;
        width: 16px;
        height: 16px;
        margin: 0 -8px;
        border-radius: 9px;
    }

    QToolButton {
        background-color: transparent;
        border: 1px transparent #76797C;
        border-radius: 2px;
        margin: 3px;
        padding: 5px;
    }

    QToolButton[popupMode="1"] { /* only for MenuButtonPopup */
     padding-right: 20px; /* make way for the popup button */
     border: 1px #76797C;
     border-radius: 5px;
    }

    QToolButton[popupMode="2"] { /* only for InstantPopup */
     padding-right: 10px; /* make way for the popup button */
     border: 1px #76797C;
    }


    QToolButton:hover, QToolButton::menu-button:hover {
        background-color: transparent;
        border: 1px solid #3daee9;
        padding: 5px;
    }

    QToolButton:checked, QToolButton:pressed,
            QToolButton::menu-button:pressed {
        background-color: #3daee9;
        border: 1px solid #3daee9;
        padding: 5px;
    }

    /* the subcontrol below is used only in the InstantPopup or DelayedPopup mode */
    QToolButton::menu-indicator {
        background-color:ff8c00;
        top: -7px; left: -2px; /* shift it a bit */
    }

    /* the subcontrols below are used only in the MenuButtonPopup mode */
    QToolButton::menu-button {
        border: 1px transparent #76797C;
        border-top-right-radius: 6px;
        border-bottom-right-radius: 6px;
        /* 16px width + 4px for border = 20px allocated above */
        width: 16px;
        outline: none;
    }

    QToolButton::menu-arrow {
       background-color:ff8c00;
    }

    QToolButton::menu-arrow:open {
        border: 1px solid #76797C;
    }

    QPushButton::menu-indicator  {
        subcontrol-origin: padding;
        subcontrol-position: bottom right;
        left: 8px;
    }

    QTableView
    {
        border: 1px solid #76797C;
        gridline-color: #31363b;
        background-color: #232629;
    }


    QTableView, QHeaderView
    {
        border-radius: 0px;
    }

    QTableView::item:pressed, QListView::item:pressed, QTreeView::item:pressed  {
        background: #3daee9;
        color: #eff0f1;
    }

    QTableView::item:selected:active, QTreeView::item:selected:active, QListView::item:selected:active  {
        background: #3daee9;
        color: #eff0f1;
    }


    QHeaderView
    {
        background-color: #31363b;
        border: 1px transparent;
        border-radius: 0px;
        margin: 0px;
        padding: 0px;

    }

    QHeaderView::section  {
        background-color: #31363b;
        color: #eff0f1;
        padding: 5px;
        border: 1px solid #76797C;
        border-radius: 0px;
        text-align: center;
    }

    QHeaderView::section::vertical::first, QHeaderView::section::vertical::only-one
    {
        border-top: 1px solid #76797C;
    }

    QHeaderView::section::vertical
    {
        border-top: transparent;
    }

    QHeaderView::section::horizontal::first, QHeaderView::section::horizontal::only-one
    {
        border-left: 1px solid #76797C;
    }

    QHeaderView::section::horizontal
    {
        border-left: transparent;
    }


    QHeaderView::section:checked
     {
        color: white;
        background-color: #334e5e;
     }

     /* style the sort indicator */
    QHeaderView::down-arrow {
        image: url(:/qss_icons/Dark_rc/down_arrow.png);
    }

    QHeaderView::up-arrow {
        image: url(:/qss_icons/Dark_rc/up_arrow.png);
    }


    QTableCornerButton::section {
        background-color: #31363b;
        border: 1px transparent #76797C;
        border-radius: 0px;
    }

    QToolBox  {
        padding: 5px;
        border: 1px transparent black;
    }

    QToolBox::tab {
        color: #eff0f1;
        background-color: #31363b;
        border: 1px solid #76797C;
        border-bottom: 1px transparent #31363b;
        border-top-left-radius: 5px;
        border-top-right-radius: 5px;
    }

    QToolBox::tab:selected { /* italicize selected tabs */
        font: italic;
        background-color: #31363b;
        border-color: #3daee9;
     }

    QStatusBar::item {
        border: 0px transparent dark;
     }


    QFrame[height="3"], QFrame[width="3"] {
        background-color: #76797C;
    }




    QDateTimeEdit
    {
        background-color:#ffffff;
        border-style: solid;
        border: 1px solid #76797C;
        border-radius: 2px;
        padding: 1px;
        min-width: 75px;
    }

    QDateEdit:on
    {
        padding-top: 2px;
        padding-left: 2px;
        selection-background-color: #4a4a4a;
    }

    QDateEdit QAbstractItemView
    {
        background-color: #ff8c00;
        border-radius: 2px;
        border: 1px solid #3375A3;
        selection-background-color:ff8c00;
    }

    QDateEdit::drop-down
    {
        subcontrol-origin: padding;
        subcontrol-position: top right;
        width: 15px;
        border-left-width: 0px;
        border-left-color: darkgray;
        border-left-style: solid;
        border-top-right-radius: 3px;
        border-bottom-right-radius: 3px;
    }""")
            self.taborge.setObjectName("taborge")
            self.tabagreage.addTab(self.taborge, "")
            self.cclstxt = QtWidgets.QLabel("<h2>COOPERATIVE DES CEREALES ET LEGUMES SECS DE RELIZANE<h2/>",
                                            self.taborge)
            self.cclstxt.resize(600, 40)
            self.cclstxt.move(400, 2)

            self.bulletin_agréagetxt = QtWidgets.QLabel("<h2>Bulletin D'Agréage<h2/>", self.taborge)
            self.bulletin_agréagetxt.move(600, 30)
            self.bulletin_agréagetxt.resize(200, 40)

            self.éspécetxt = QtWidgets.QLabel("<h2>Espéce:Orge<h2/>", self.taborge)
            self.éspécetxt.move(630, 60)

            self.n_bultintxt = QtWidgets.QLabel("N° Bulletin:", self.taborge)
            self.n_bultintxt.move(20, 60)
            self.n_bultin = QtWidgets.QLineEdit(self.taborge, readOnly=True)
            self.n_bultin.move(150, 60)
            self.n_bultin.resize(150, 25)

            self.nome_du_producteurtxt = QtWidgets.QLabel("Nom du Producteur:", self.taborge)
            self.nome_du_producteurtxt.move(20, 100)
            self.nome_du_producteur = QtWidgets.QLineEdit(self.taborge)
            self.nome_du_producteur.move(150, 98)
            self.nome_du_producteur.resize(150, 25)

            self.péretxt = QtWidgets.QLabel("Pére:", self.taborge)
            self.péretxt.setGeometry(QtCore.QRect(315, 98, 100, 20))
            self.pére = QtWidgets.QLineEdit(self.taborge)
            self.pére.setGeometry(QtCore.QRect(410, 98, 150, 23))

            self.n_cartetxt = QtWidgets.QLabel("N.C d'identité:", self.taborge)
            self.n_cartetxt.setGeometry(QtCore.QRect(315, 138, 100, 20))
            self.n_carte = QtWidgets.QLineEdit(self.taborge)
            self.n_carte.setGeometry(QtCore.QRect(410, 138, 150, 23))
            self.n_carte.setInputMask('99999999')

            self.imatriculetxt = QtWidgets.QLabel("Adresse:", self.taborge)
            self.imatriculetxt.move(20, 140)
            self.adresse = QtWidgets.QLineEdit(self.taborge)
            self.adresse.setInputMask("99999-999-99")
            self.adresse.move(150, 138)
            self.adresse.resize(150, 25)

            self.pointdecollecttxt = QtWidgets.QLabel("Point de collecte", self.taborge)
            self.pointdecollecttxt.move(20, 180)
            self.pointdecollect = QtWidgets.QComboBox(self.taborge)
            self.pointdecollect.setStyleSheet("background-color:#ffffff;color:#000000")
            self.pointdecollect.move(150, 178)
            self.pointdecollect.resize(150, 25)
            self.pointdecollect.addItem('')
            self.pointdecollect.addItem('Dock central')
            self.pointdecollect.addItem('Magasin Zemmoura')
            self.pointdecollect.addItem('Magasin Kef-lazreg')
            self.pointdecollect.addItem('Magasin Messra')
            self.pointdecollect.addItem('Marche gros belacel')
            self.pointdecollect.addItem('Station Mendes')
            self.pointdecollect.addItem('Nouvelles S.Mendes')

            self.dattereceptiontxt = QtWidgets.QLabel("Relizane le :", self.taborge)
            self.dattereceptiontxt.setGeometry(QtCore.QRect(880, 100, 150, 23))
            self.dattereceptiont = QtWidgets.QDateTimeEdit(self.taborge)
            self.dattereceptiont.setDisplayFormat("dd-MM-yyyy hh:mm")
            self.dattereceptiont.setGeometry(QtCore.QRect(1000, 100, 150, 23))
            self.dattereceptiont.setDate(self.datedaytime)


            self.agréeeurtxt = QtWidgets.QLabel("Nom de l’Agréeur:", self.taborge)
            self.agréeeurtxt.setGeometry(QtCore.QRect(880, 135, 150, 23))
            self.agréeeurcomboorge = QtWidgets.QComboBox(self.taborge, editable=True)
            self.agréeeurcomboorge.setStyleSheet("background-color:#ffffff;color:#000000")
            self.agréeeurcomboorge.addItem("")
            self.agréeeurcomboorge.addItem("FELOUAH OMAR")
            self.agréeeurcomboorge.addItem("BEKHEDDA AEK")
            self.agréeeurcomboorge.addItem("BENAISSA YOUCEF")
            self.agréeeurcomboorge.addItem("REZZAG SOFIANE ")
            self.agréeeurcomboorge.addItem("BELBACHA M.NADIR")
            self.agréeeurcomboorge.move(1000, 135)
            self.agréeeurcomboorge.resize(150, 23)

            self.quantitetxt = QtWidgets.QLabel("Quantité", self.taborge)
            self.quantitetxt.setGeometry(QtCore.QRect(880, 180, 150, 23))

            self.quantiteorge = QtWidgets.QDoubleSpinBox(self.taborge)
            self.quantiteorge.setRange(1, 10000)
            self.quantiteorge.setSuffix('  QX')
            self.quantiteorge.setSpecialValueText(" ")
            self.quantiteorge.setGeometry(QtCore.QRect(1000, 178, 150, 23))

            self.paramétreor = QtWidgets.QLabel("Paramètre", self.taborge)
            self.paramétreor.move(20, 250)
            self.paramétreor.resize(80, 20)

            self.txtpsfont = QtGui.QFont()
            self.txtpsfont.setBold(True)
            self.txtpsfont.setPointSize(9)
            ################Limites(sans bon ni réf)###############
            self.valeuror = QtWidgets.QLabel("""Limite-ssans-bon-ni-réf)""", self.taborge)
            self.valeuror.move(170, 250)
            self.valeuror.resize(145, 20)

            ######################Limites(sans bon ni réf)################
            self.psor = QtWidgets.QLineEdit("Poids spécifique (kg/hl):\t(58-62)", self.taborge, readOnly=True)
            self.psor.resize(319, 40)
            self.psor.move(20, 280)
            self.psor.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            self.psor.setFont(self.txtpsfont)
            self.psor.setStyleSheet(" border: 2px solid bleu ;border-radius: 4px;padding: 0px")
            ###############################humidite#############
            self.ergotor = QtWidgets.QLineEdit("Ergot(%):        \t\t<= 1 ", self.taborge, readOnly=True)
            self.ergotor.resize(319, 40)
            self.ergotor.move(20, 322)
            self.ergotor.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            self.ergotor.setFont(self.txtpsfont)

            #######################ergot#########################
            self.Grainsanvaleuror = QtWidgets.QLineEdit("Grains sans valeurs (%) ", self.taborge, readOnly=True)
            self.Grainsanvaleuror.resize(319, 40)
            self.Grainsanvaleuror.move(20, 364)
            self.Grainsanvaleuror.setStyleSheet("background-color: #232629")
            self.Grainsanvaleuror.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            self.Grainsanvaleuror.setFont(self.txtpsfont)

            #########################Graines nuisibles (%)##########
            self.matierinertor = QtWidgets.QLineEdit("Matiéres inertes (%):", self.taborge, readOnly=True)
            self.matierinertor.resize(319, 40)
            self.matierinertor.move(20, 406)
            self.matierinertor.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            self.matierinertor.setFont(self.txtpsfont)
            #############################Débris végétaux (%)########
            self.totalor = QtWidgets.QLineEdit("Total (%):        \t\t<= 2 ", self.taborge, readOnly=True)
            self.totalor.resize(319, 40)
            self.totalor.move(20, 448)
            self.totalor.setFont(self.txtpsfont)
            self.totalor.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")

            self.totalor = QtWidgets.QLineEdit("\tTotal des Bonifications et Réfactions:", self.taborge, readOnly=True)
            self.totalor.resize(450, 40)
            self.totalor.move(20, 490)
            self.totalor.setFont(self.txtpsfont)
            self.totalor.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")

            #########################Matière inerte (%)################

            self.valeuror = QtWidgets.QLabel("valeur", self.taborge)
            self.valeuror.move(350, 250)
            self.valeuror.resize(100, 20)

            ######################Limites(sans bon ni réf)################
            self.vpsor = QtWidgets.QDoubleSpinBox(self.taborge)
            self.vpsor.setRange(0, 72.00)
            self.vpsor.setSpecialValueText(' ')
            self.vpsor.resize(120, 40)
            self.vpsor.move(350, 280)
            self.vpsor.setFont(self.txtpsfont)
            self.vpsor.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")
            ###############################humidite#############
            self.vergotor = QtWidgets.QDoubleSpinBox(self.taborge)
            self.vergotor.setRange(0, 1)
            self.vergotor.resize(120, 40)
            self.vergotor.setSpecialValueText(' ')
            self.vergotor.move(350, 322)
            self.vergotor.setFont(self.txtpsfont)
            self.vergotor.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #######################ergot#########################
            self.vGrainsanvaleuror = QtWidgets.QDoubleSpinBox(self.taborge, readOnly=False)
            self.vGrainsanvaleuror.setRange(0, 10)
            self.vGrainsanvaleuror.setSpecialValueText(' ')
            self.vGrainsanvaleuror.resize(120, 40)
            self.vGrainsanvaleuror.move(350, 364)
            self.vGrainsanvaleuror.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #########################Graines nuisibles (%)##########
            self.vmatierinertor = QtWidgets.QDoubleSpinBox(self.taborge, readOnly=False)
            self.vmatierinertor.setRange(0, 10)
            self.vmatierinertor.setSpecialValueText(' ')
            self.vmatierinertor.resize(120, 40)
            self.vmatierinertor.move(350, 406)
            self.vmatierinertor.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")
            #############################Débris végétaux (%)########
            self.vtotalor = QtWidgets.QDoubleSpinBox(self.taborge, readOnly=True)
            self.vtotalor.setRange(0, 20)
            self.vtotalor.setSpecialValueText(' ')
            self.vtotalor.resize(120, 40)
            self.vtotalor.move(350, 448)
            self.vtotalor.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 0px")

            self.bonior = QtWidgets.QLabel("Bonification", self.taborge)
            self.bonior.move(480, 250)
            self.bonior.resize(80, 20)
            ######################Limites(sans bon ni réf)################
            self.bpsor = QtWidgets.QDoubleSpinBox(self.taborge, readOnly=True)
            self.bpsor.resize(120, 40)
            self.bpsor.move(480, 280)
            self.bpsor.setSpecialValueText(" ")
            self.bpsor.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")
            ###############################humidite#############
            self.bergotor = QtWidgets.QDoubleSpinBox(self.taborge, readOnly=True)
            self.bergotor.resize(120, 40)
            self.bergotor.move(480, 322)
            self.bergotor.setSpecialValueText(" ")
            self.bergotor.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")
            #######################ergot#########################
            self.bGrainsanvaleuror = QtWidgets.QDoubleSpinBox(self.taborge, readOnly=True)
            self.bGrainsanvaleuror.resize(120, 40)
            self.bGrainsanvaleuror.move(480, 364)
            self.bGrainsanvaleuror.setSpecialValueText(" ")
            self.bGrainsanvaleuror.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")
            #########################Graines nuisibles (%)##########
            self.bmatierinertor = QtWidgets.QDoubleSpinBox(self.taborge, readOnly=True)
            self.bmatierinertor.resize(120, 40)
            self.bmatierinertor.move(480, 406)
            self.bmatierinertor.setSpecialValueText(" ")
            self.bmatierinertor.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")
            #############################Débris végétaux (%)########
            self.btotalor = QtWidgets.QDoubleSpinBox(self.taborge, readOnly=True)
            self.btotalor.resize(120, 40)
            self.btotalor.move(480, 448)
            self.btotalor.setSpecialValueText(" ")
            self.btotalor.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")
            #########################Matière inerte (%)################
            self.bglobaltotalor = QtWidgets.QDoubleSpinBox(self.taborge, readOnly=True)
            self.bglobaltotalor.resize(120, 40)
            self.bglobaltotalor.move(480, 490)
            self.bglobaltotalor.setSpecialValueText(" ")
            self.bglobaltotalor.setStyleSheet(
                "background-color:#88ffaa;color:#000000;border: 2px solid bleu;border-radius: 4px;padding: 0px")

            self.refactionor = QtWidgets.QLabel("Réfaction", self.taborge)
            self.refactionor.move(610, 250)
            self.refactionor.resize(80, 20)
            ######################Limites(sans bon ni réf)################
            self.rpsor = QtWidgets.QDoubleSpinBox(self.taborge, readOnly=True)
            self.rpsor.resize(120, 40)
            self.rpsor.move(610, 280)
            self.rpsor.setSpecialValueText(" ")
            self.rpsor.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")
            ###############################humidite#############
            self.rergotor = QtWidgets.QDoubleSpinBox(self.taborge, readOnly=True)
            self.rergotor.resize(120, 40)
            self.rergotor.move(610, 322)
            self.rergotor.setSpecialValueText(" ")
            self.rergotor.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")
            #######################ergot#########################
            self.rGrainsanvaleuror = QtWidgets.QDoubleSpinBox(self.taborge, readOnly=True)
            self.rGrainsanvaleuror.resize(120, 40)
            self.rGrainsanvaleuror.move(610, 364)
            self.rGrainsanvaleuror.setSpecialValueText(" ")
            self.rGrainsanvaleuror.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")
            #########################Graines nuisibles (%)##########
            self.rmatierinertor = QtWidgets.QDoubleSpinBox(self.taborge, readOnly=True)
            self.rmatierinertor.resize(120, 40)
            self.rmatierinertor.move(610, 406)
            self.rmatierinertor.setSpecialValueText(" ")
            self.rmatierinertor.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")
            #############################Débris végétaux (%)########
            self.rtotalor = QtWidgets.QDoubleSpinBox(self.taborge, readOnly=True)
            self.rtotalor.resize(120, 40)
            self.rtotalor.move(610, 448)
            self.rtotalor.setSpecialValueText(" ")
            self.rtotalor.setStyleSheet("border: 2px solid bleu;border-radius: 4px;padding: 2px")
            #########################Matière inerte (%)################
            self.rglobaltotalor = QtWidgets.QDoubleSpinBox(self.taborge, readOnly=True)
            self.rglobaltotalor.setStyleSheet(
                "background-color:#e22630;color:000000;border: 2px solid bleu;border-radius: 4px;padding: 0px")
            self.rglobaltotalor.resize(120, 40)
            self.rglobaltotalor.setSpecialValueText(" ")
            self.rglobaltotalor.move(610, 490)

            self.observationortxt = QtWidgets.QLabel("Observation", self.taborge)
            self.observationortxt.setGeometry(QtCore.QRect(743, 250, 100, 20))

            self.observationor = QtWidgets.QTextEdit(self.taborge)
            self.observationor.setGeometry(QtCore.QRect(743, 280, 120, 250))

            self.timer = QTimer()
            # Connect the timeout signal of the timer to the refresh function
            self.timer.timeout.connect(self.number_bulletin_orge)
            # Set the interval in milliseconds (e.g., 5000 ms = 5 seconds)
            self.timer.setInterval(1000)  # Adjust the interval as per your requirement
            # Start the timer
            self.timer.start()

            self.timercalculorge = QTimer()
            self.timercalculorge.timeout.connect(self.calcul_orge)
            self.timercalculorge.setInterval(1000)
            self.timercalculorge.start()

            self.btnsavebd = QtWidgets.QPushButton("ENREGISTRER", self.taborge, clicked=lambda: self.add_datta_orge())
            self.btnsavebd.move(880, 280)
            self.btnsavebd.resize(125, 50)

            self.btnimprimejour = QtWidgets.QPushButton("IMPRIMER \nLE JOURNALIER", self.taborge,
                                                        clicked=lambda: self.print_day())
            self.btnimprimejour.move(880, 410)
            self.btnimprimejour.resize(125, 50)

            self.btnimprimeticket = QtWidgets.QPushButton("IMPRIMER \nLES TICKET", self.taborge,
                                                          clicked=lambda: self.eticket_orge())
            self.btnimprimeticket.move(1030, 340)
            self.btnimprimeticket.resize(125, 50)

            self.btnselectionblt = QtWidgets.QPushButton("SELECTIONNE \nBULLETIN", self.taborge,
                                                         clicked=lambda: self.printBulltinProducteurOrge())
            self.btnselectionblt.move(1030, 280)
            self.btnselectionblt.resize(125, 50)

            self.btnprintbd = QtWidgets.QPushButton("IMPRIMER", self.taborge,
                                                    clicked=lambda: self.bulletin_orge_print())
            self.btnprintbd.move(880, 340)
            self.btnprintbd.resize(125, 50)

            self.btnefaceor = QtWidgets.QPushButton("EFACER", self.taborge, clicked=lambda: self.clear_orge())
            self.btnefaceor.move(880, 480)
            self.btnefaceor.resize(125, 50)
            #########################################################################################################
            ##################################AVOINE###################################################
            #########################################################################################################
            ###self.tabavoine = QtWidgets.QWidget()
            # self.tabavoine.setObjectName("tabavoine")
            # self.tabagreage.addTab(self.tabavoine, "")
            # self.cclstxt = QtWidgets.QLabel("<h2>COOPERATIVE DES CEREALES ET LEGUMES SECS DE RELIZANE<h2/>", self.tabavoine)
            # self.cclstxt.resize(600, 40)
            # self.cclstxt.move(400, 2)

            # self.cclstxt = QtWidgets.QLabel("<h2>SERA DISPONIBLE PLUS TARD ................<h2/>", self.tabavoine)
            # self.cclstxt.resize(600, 100)
            # self.cclstxt.move(500, 100)

            # self.font = QtGui.QFont()
            # self.font.setBold(True)
            # self.font.setPointSize(10)

            # self.txtpsfont = QtGui.QFont()
            # self.txtpsfont.setBold(True)
            # self.txtpsfont.setPointSize(9)

            # self.rtotaldemfont = QtGui.QFont("color:black")
            # self.rtotaldemfont.setBold(True)
            # self.rtotaldemfont.setPointSize(12)

            # self.bulletin_agréagetxt = QtWidgets.QLabel("<h2>Bulletin D'Agréage<h2/>", self.tabavoine)
            # self.bulletin_agréagetxt.move(600, 30)
            # self.bulletin_agréagetxt.resize(200, 40)

            # self.éspécetxt = QtWidgets.QLabel("<h2>Espéce:Avoine <h2/>", self.tabavoine)
            # self.éspécetxt.move(630, 60)

            self.verticalLayout.addWidget(self.tabagreage)
            MainWindow.setCentralWidget(self.centralwidget)
            self.statusbar = QtWidgets.QStatusBar(MainWindow)
            self.statusbar.setObjectName("statusbar")
            MainWindow.setStatusBar(self.statusbar)

            self.retranslateUi(MainWindow)
            self.tabagreage.setCurrentIndex(3)
            QtCore.QMetaObject.connectSlotsByName(MainWindow)

        def retranslateUi(self, MainWindow):
            _translate = QtCore.QCoreApplication.translate
            MainWindow.setWindowTitle(_translate("MainWindow", "MainWindow"))
            self.tabagreage.setTabText(self.tabagreage.indexOf(self.tabbldur), _translate("MainWindow", "Blé Dur"))
            self.tabagreage.setTabText(self.tabagreage.indexOf(self.tabbltendre),
                                       _translate("MainWindow", "Blé Tendre"))
            self.tabagreage.setTabText(self.tabagreage.indexOf(self.taborge), _translate("MainWindow", "Orge"))
            # self.tabagreage.setTabText(self.tabagreage.indexOf(self.tabavoine), _translate("MainWindow", "Avoine"))


    if __name__ == "__main__":
        import sys
        app = QtWidgets.QApplication(sys.argv)
        MainWindow = QtWidgets.QMainWindow()
        ui = Agréage_Window()
        ui.agréage(MainWindow)
        MainWindow.show()
        sys.exit(app.exec())

except Exception as e:
    print(e)



